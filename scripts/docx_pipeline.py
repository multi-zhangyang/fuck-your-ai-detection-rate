"""Utilities for preserving DOCX structure while rewriting text.

The core AIGC reduction still operates on plain text, but this module now
extracts a stable editable-text snapshot from the original DOCX and can write
the rewritten text back onto a copy of the original Word file. This preserves
images, cover pages, tables, page settings, and most existing layout details.
"""

from __future__ import annotations

import argparse
from difflib import SequenceMatcher
import hashlib
import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Iterable, Iterator, Sequence

from path_utils import build_document_artifact_stem

try:
    from docx import Document  # type: ignore[import]
    from docx.document import Document as DocxDocument  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # type: ignore[import]
    from docx.oxml import OxmlElement  # type: ignore[import]
    from docx.oxml.ns import qn  # type: ignore[import]
    from docx.oxml.table import CT_Tbl  # type: ignore[import]
    from docx.oxml.text.paragraph import CT_P  # type: ignore[import]
    from docx.shared import Cm, Pt  # type: ignore[import]
    from docx.table import Table, _Cell  # type: ignore[import]
    from docx.text.paragraph import Paragraph  # type: ignore[import]
    from docx.text.run import Run  # type: ignore[import]
except ImportError as exc:  # pragma: no cover - import guard
    raise SystemExit(
        "Missing dependency python-docx. Install it with: pip install python-docx"
    ) from exc


ROOT_DIR = Path(__file__).resolve().parents[1]
INTERMEDIATE_DIR = ROOT_DIR / "finish" / "intermediate"
# Version 22 freezes standalone university-template writing instructions and
# acknowledgement guidance. Cached v21 scope is deliberately not migratable:
# the source OOXML must be parsed again so directive prose cannot retain an old
# ``body_prose`` allow decision merely because it contains sentence punctuation.
DOCX_SNAPSHOT_VERSION = 22
DOCX_SCOPE_DIAGNOSTICS_VERSION = 5
DOCX_STRUCTURAL_ROLE_POLICY_VERSION = 5
DOCX_STRUCTURAL_INVENTORY_VERSION = 3
DOCX_EDITABLE_STRUCTURAL_ROLES = frozenset({"abstract_body", "body_prose", "body_list"})
SCOPE_DIAGNOSTIC_TEXT_PREVIEW_CHARS = 180
DEFAULT_CN_FONT = "宋体"
DEFAULT_EN_FONT = "Times New Roman"
DEFAULT_BODY_FONT_SIZE_PT = 12
DEFAULT_BODY_LINE_SPACING_PT = 20
DEFAULT_BODY_FIRST_LINE_INDENT_CM = 0.74
DEFAULT_PAGE_MARGINS_CM = {
    "top_margin": 2.54,
    "bottom_margin": 2.54,
    "left_margin": 3.17,
    "right_margin": 3.17,
}
ABSTRACT_MARKERS = {"摘要", "中文摘要", "论文摘要", "内容摘要", "abstract"}
ACKNOWLEDGEMENT_MARKERS = {
    "致谢",
    "致谢辞",
    "谢辞",
    "鸣谢",
    "acknowledgement",
    "acknowledgements",
    "acknowledgment",
    "acknowledgments",
}
TOC_MARKERS = {"目录", "tableofcontents"}
REFERENCE_MARKERS = {
    "参考文献",
    "主要参考文献",
    "参考文献目录",
    "参考文献列表",
    "参考资料",
    "主要参考资料",
    "参考资料目录",
    "参考资料列表",
    "参考书目",
    "引用文献",
    "主要引用文献",
    "文献目录",
    "文献列表",
    "文献来源",
    "资料来源",
    "参考来源",
    "reference",
    "references",
    "referencelist",
    "listofreferences",
    "referencematerials",
    "selectedreferences",
    "citedreferences",
    "bibliography",
    "selectedbibliography",
    "workscited",
    "worksconsulted",
    "worksreferenced",
    "sourcescited",
    "sourcesconsulted",
    "sourcelist",
    "listofsources",
    "literaturecited",
    "citedliterature",
    "literaturereferences",
    "furtherreading",
}
BACK_MATTER_MARKERS = {
    "附录",
    "appendix",
    "appendices",
    "声明",
    "诚信声明",
    "承诺书",
    "独创性声明",
    "作者简介",
    "个人简历",
    "任务书",
    "开题报告",
    "指导教师评阅表",
    "评阅教师评阅表",
    "答辩记录",
    "答辩委员会评语",
    "成绩评定表",
    "版权使用授权书",
    "学位论文版权使用授权书",
    "毕业论文版权使用授权书",
    "毕业设计论文任务书",
    "毕业设计（论文）任务书",
    "原创性声明",
    "学位论文原创性声明",
    "学位论文使用授权声明",
    "攻读学位期间取得的研究成果",
    "攻读学位期间发表的学术论文",
    "在学期间研究成果",
    "外文资料",
    "外文原文",
    "外文译文",
    "译文",
    "原文",
    "学术诚信承诺书",
    "诚信承诺书",
    "论文原创性声明",
    "论文使用授权声明",
    "毕业设计论文诚信承诺书",
    "毕业设计论文原创性声明",
    "封底",
}
KEYWORD_PREFIXES = ("关键词", "关键字", "keywords", "keyword")
BODY_START_RE = re.compile(r"^(第[一二三四五六七八九十0-9]+章|[1-9]\d*(?:\.\d+)*[^\d].*|绪论|引言|前言)$")
HEADING_NUMBER_RE = re.compile(r"^(第[一二三四五六七八九十百0-9]+章.*|[1-9]\d*(?:\.\d+){0,3}[^\d].*)$")
CAPTION_RE = re.compile(r"^(图|表|fig(?:ure)?[.]?|table)(?:[a-z]?\d+|[一二三四五六七八九十]+)(?:[-.．－—]\d+)*")
NUMBERED_BODY_MARKER_RE = re.compile(
    r"^\s*(?:[（(]?(?:\d{1,3}|[一二三四五六七八九十百]+)[）)]|(?:\d{1,3}|[一二三四五六七八九十百]+|[A-Za-z])[.．、)])\s*(?P<body>.+)$"
)
NUMBERED_SPACE_BODY_MARKER_RE = re.compile(r"^\s*(?:\d{1,3}|[一二三四五六七八九十百]+)\s+(?P<body>.+)$")
DATE_LINE_RE = re.compile(r"^\d{4}年\d{1,2}月\d{1,2}日$")
LATIN_CHAR_RE = re.compile(r"[A-Za-z]")
DIGIT_CHAR_RE = re.compile(r"\d")
FORMULA_SYMBOL_RE = re.compile(r"[=≈≠≤≥±×÷∑Σ√∫∞α-ωΑ-ΩμσλθπφψωΔδβγ_^\[\]{}<>]")
FORMULA_STYLE_HINTS = ("equation", "公式")
SHORT_HEADING_RE = re.compile(r"^[A-Za-z0-9\u4e00-\u9fff\s（）()、：:·/+\-]{2,36}$")
REFERENCE_BRACKET_ENTRY_RE = re.compile(r"^\s*\[\s*\d{1,4}\s*\]")
REFERENCE_NUMBERED_ENTRY_RE = re.compile(r"^\s*\d{1,4}\s*[.．]\s+")
REFERENCE_YEAR_RE = re.compile(r"(?<!\d)(?:19|20)\d{2}[a-z]?(?!\d)", re.IGNORECASE)
REFERENCE_AUTHOR_YEAR_PAREN_RE = re.compile(
    r"^\s*.{2,96}?[（(]\s*(?:19|20)\d{2}[a-z]?\s*[）)]\s*[.．]",
    re.IGNORECASE,
)
REFERENCE_EN_AUTHOR_LEAD_RE = re.compile(
    r"^\s*[A-Z][A-Za-z'’\-]+"
    r"(?:,\s*(?:[A-Z](?:\.)?|[A-Z][A-Za-z'’\-]+))?"
    r"(?:\s*(?:&|and)\s*[A-Z][A-Za-z'’\-]+)?",
)
REFERENCE_CN_AUTHOR_LEAD_RE = re.compile(
    r"^\s*(?:"
    r"[\u3400-\u9fff]{2,4}(?:[，,、]\s*[\u3400-\u9fff]{2,4}){1,5}[.．，,]"
    r"|[\u3400-\u9fff]{2,4}[.．]"
    r"|[\u3400-\u9fff]{2,4}\s*[（(](?:19|20)\d{2}"
    r")",
)
REFERENCE_BIBLIOGRAPHIC_CUE_RE = re.compile(
    r"(?:"
    r"doi\s*[:：]|https?://|isbn\s*[:：]?|"
    r"\[(?:j|m|c|d|r|s|eb|ol)\]|"
    r"\b(?:journal|press|proceedings|conference|vol\.?|no\.?|pp\.?)\b|"
    r"(?:学报|期刊|出版社|会议论文|硕士学位论文|博士学位论文)"
    r")",
    re.IGNORECASE,
)
POST_ACKNOWLEDGEMENT_TAIL_MARKER_RE = re.compile(
    r"(?:参考文献|主要参考文献|参考资料|主要参考资料|参考书目|引用文献|文献目录|文献列表|资料来源|"
    r"附录|appendix|references?|referencelist|listofreferences|bibliography|selectedbibliography|"
    r"workscited|worksconsulted|worksreferenced|sourcescited|sourcesconsulted|literaturecited|"
    r"声明|承诺书|作者简介|个人简历|任务书|开题报告)"
)
HEADING_KEYWORDS = (
    "研究背景",
    "研究意义",
    "国内外研究",
    "研究现状",
    "相关理论",
    "关键技术",
    "需求分析",
    "总体设计",
    "系统设计",
    "系统实现",
    "模型构建",
    "模型训练",
    "实验设计",
    "实验结果",
    "结果分析",
    "特征工程",
    "数据预处理",
    "系统测试",
    "总结",
    "展望",
)
SHORT_NUMBERED_HEADING_BODIES = {
    "摘要",
    "abstract",
    "绪论",
    "引言",
    "前言",
    "结论",
    "结语",
    "总结",
    "展望",
    "数据与方法",
    "材料与方法",
    "研究方法",
    "实验方法",
    "参考文献",
    "致谢",
}
TEMPLATE_INSTRUCTION_PREFIX_RE = re.compile(
    r"^\s*(?:"
    r"注意|提示|模板(?:说明|要求)|撰写(?:说明|要求)|编写(?:说明|要求)|填写(?:说明|要求)|"
    r"内容要求|写作要求"
    r")\s*[:：]",
    re.IGNORECASE,
)
TEMPLATE_DOCUMENT_AUTHORING_CUE_RE = re.compile(
    r"(?:"
    r"毕业设计(?:[（(]论文[）)])?|毕业论文|学位论文|"
    r"论文(?:中|正文|内容|写作|撰写|编写|格式)|"
    r"本(?:章|节)(?:中|内容|应|需|须|要求)|课题性质|"
    r"每位同学|学生(?:应|须|需要)|作者(?:应|须|需要)"
    r")",
    re.IGNORECASE,
)
TEMPLATE_DIRECTIVE_CUE_RE = re.compile(
    r"(?:应当|应|须|需要|必须|不得|请|要求|建议|给出|列明|体现|包含|包括|填写|撰写|描述|替换|删除|修改)",
    re.IGNORECASE,
)
ACKNOWLEDGEMENT_GUIDANCE_DEFINITION_RE = re.compile(
    r"^\s*(?:致谢|谢辞|鸣谢)(?:部分|内容|章节)?\s*(?:是|用于|主要用于|应当?|须|需要)",
    re.IGNORECASE,
)
ACKNOWLEDGEMENT_GUIDANCE_AUDIENCE_RE = re.compile(
    r"(?:作者|学生|同学|撰写者|本人|个人|实际情况|自行)",
    re.IGNORECASE,
)
ACKNOWLEDGEMENT_GUIDANCE_ACTION_RE = re.compile(
    r"(?:填写|撰写|描述|替换|删除|修改|表达感谢|说明)",
    re.IGNORECASE,
)
SEMANTIC_RANGE_MARKER_ACTIONS = {
    qn("w:bookmarkStart"): ("bookmark", "start"),
    qn("w:bookmarkEnd"): ("bookmark", "end"),
    qn("w:commentRangeStart"): ("comment", "start"),
    qn("w:commentRangeEnd"): ("comment", "end"),
}


@dataclass
class DocxTextUnit:
    unit_index: int
    target: dict[str, Any]
    text: str
    style_name: str
    leading_whitespace: str = ""
    trailing_whitespace: str = ""
    # ``editable`` is a compatibility projection only.  Fresh parsing and cache
    # loading derive it from the role/evidence fields below; a missing field is
    # therefore protected rather than implicitly trusted.
    editable: bool = False
    protect_reason: str | None = None
    structural_role: str = "unknown"
    edit_eligibility: str = "protected"
    edit_eligibility_evidence: dict[str, Any] = field(default_factory=dict)
    presentation_signals: dict[str, Any] = field(default_factory=dict)
    has_field_code: bool = False
    has_drawing: bool = False
    has_math: bool = False
    has_complex_inline: bool = False
    has_semantic_range_anchor: bool = False
    inside_semantic_range: bool = False
    has_bookmark_range_anchor: bool = False
    has_comment_range_anchor: bool = False
    inside_bookmark_range: bool = False
    inside_comment_range: bool = False
    has_semantic_point_reference: bool = False
    has_numbering: bool = False
    numbering_level: int | None = None
    outline_level: int | None = None
    format_anchors: list[dict[str, Any]] = field(default_factory=list)
    format_anchor_ambiguous: bool = False

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    def source_text(self) -> str:
        """Return the literal source text represented by this model-text unit."""

        return f"{self.leading_whitespace}{self.text}{self.trailing_whitespace}"


@dataclass
class DocxSnapshot:
    version: int
    source_path: str
    source_size: int
    source_mtime_ns: int
    source_sha256: str
    editable_unit_count: int
    total_text_unit_count: int
    units: list[DocxTextUnit]
    semantic_range_count: int = 0
    bookmark_range_count: int = 0
    comment_range_count: int = 0
    semantic_range_topology_valid: bool = True
    semantic_range_issue_count: int = 0
    semantic_range_issue_codes: list[str] = field(default_factory=list)
    structural_role_policy_version: int = DOCX_STRUCTURAL_ROLE_POLICY_VERSION
    structural_inventory_version: int = DOCX_STRUCTURAL_INVENTORY_VERSION
    protected_structural_units: list[DocxTextUnit] = field(default_factory=list)

    def editable_units(self) -> list[DocxTextUnit]:
        return [
            unit
            for unit in self.units
            if unit.editable
            and unit.structural_role in DOCX_EDITABLE_STRUCTURAL_ROLES
            and unit.edit_eligibility == "eligible"
            and validate_docx_unit_edit_eligibility_evidence(unit)
        ]

    def editable_texts(self) -> list[str]:
        return [unit.text for unit in self.editable_units()]

    def to_dict(self) -> dict[str, Any]:
        return {
            "version": self.version,
            "source_path": self.source_path,
            "source_size": self.source_size,
            "source_mtime_ns": self.source_mtime_ns,
            "source_sha256": self.source_sha256,
            "editable_unit_count": self.editable_unit_count,
            "total_text_unit_count": self.total_text_unit_count,
            "units": [unit.to_dict() for unit in self.units],
            "semantic_range_count": self.semantic_range_count,
            "bookmark_range_count": self.bookmark_range_count,
            "comment_range_count": self.comment_range_count,
            "semantic_range_topology_valid": self.semantic_range_topology_valid,
            "semantic_range_issue_count": self.semantic_range_issue_count,
            "semantic_range_issue_codes": list(self.semantic_range_issue_codes),
            "structural_role_policy_version": self.structural_role_policy_version,
            "structural_inventory_version": self.structural_inventory_version,
            "protected_structural_units": [
                unit.to_dict()
                for unit in self.protected_structural_units
            ],
        }


def get_docx_extracted_text_path(source_path: Path) -> Path:
    artifact_stem = build_document_artifact_stem(root_dir=ROOT_DIR, source_path=source_path)
    return INTERMEDIATE_DIR / f"{artifact_stem}_extracted.txt"


def get_docx_snapshot_path(source_path: Path) -> Path:
    artifact_stem = build_document_artifact_stem(root_dir=ROOT_DIR, source_path=source_path)
    return INTERMEDIATE_DIR / f"{artifact_stem}_docx_snapshot.json"


def get_docx_scope_diagnostics_path(source_path: Path) -> Path:
    artifact_stem = build_document_artifact_stem(root_dir=ROOT_DIR, source_path=source_path)
    return INTERMEDIATE_DIR / f"{artifact_stem}_scope_diagnostics.json"


def _get_legacy_docx_snapshot_path(source_path: Path) -> Path:
    """Return the pre-identity snapshot path for read-only migration."""

    return INTERMEDIATE_DIR / f"{source_path.stem}_docx_snapshot.json"


def read_docx_text(path: Path) -> str:
    """Read a DOCX file as editable text blocks joined by blank lines."""
    snapshot = build_docx_snapshot(path)
    return "\n\n".join(snapshot.editable_texts())


def read_docx_paragraphs(path: Path) -> list[str]:
    """Read a DOCX file as editable text blocks in order."""
    snapshot = build_docx_snapshot(path)
    return snapshot.editable_texts()


def write_docx_text(lines: Iterable[str], path: Path) -> None:
    """Write an iterable of text blocks into a brand-new .docx file."""
    document = Document()
    _apply_default_document_layout(document)
    for block in lines:
        paragraph = document.add_paragraph(block)
        _apply_default_paragraph_layout(paragraph)
    document.save(str(path))


def write_docx_paragraphs(paragraphs: Iterable[str], path: Path) -> None:
    write_docx_text(paragraphs, path)


def build_docx_snapshot(path: Path) -> DocxSnapshot:
    source_path = path.resolve()
    document = Document(str(source_path))
    semantic_range_scan = _scan_document_semantic_ranges(document)
    units = _collect_text_units(
        document,
        semantic_range_paragraph_indexes=set(semantic_range_scan["paragraphIndexes"]),
        bookmark_range_interior_paragraph_indexes=set(
            semantic_range_scan["bookmarkInteriorParagraphIndexes"]
        ),
        comment_range_interior_paragraph_indexes=set(
            semantic_range_scan["commentInteriorParagraphIndexes"]
        ),
    )
    _apply_protection_rules(
        units,
        semantic_range_topology_valid=bool(semantic_range_scan["topologyValid"]),
    )
    protected_structural_units = _collect_protected_table_units(
        document,
        first_unit_index=len(units),
    )
    source_stat = source_path.stat()
    source_sha256 = _sha256_file(source_path)
    return DocxSnapshot(
        version=DOCX_SNAPSHOT_VERSION,
        source_path=str(source_path),
        source_size=source_stat.st_size,
        source_mtime_ns=source_stat.st_mtime_ns,
        source_sha256=source_sha256,
        editable_unit_count=sum(1 for unit in units if unit.editable),
        total_text_unit_count=len(units),
        units=units,
        semantic_range_count=int(semantic_range_scan["rangeCount"]),
        bookmark_range_count=int(semantic_range_scan["bookmarkRangeCount"]),
        comment_range_count=int(semantic_range_scan["commentRangeCount"]),
        semantic_range_topology_valid=bool(semantic_range_scan["topologyValid"]),
        semantic_range_issue_count=int(semantic_range_scan["issueCount"]),
        semantic_range_issue_codes=[str(code) for code in semantic_range_scan["issueCodes"]],
        structural_role_policy_version=DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
        structural_inventory_version=DOCX_STRUCTURAL_INVENTORY_VERSION,
        protected_structural_units=protected_structural_units,
    )


def _docx_snapshot_derivation_digest(snapshot: DocxSnapshot) -> str:
    """Hash the complete source-derived snapshot, including its frozen scope.

    Source size/mtime/sha only prove which file a cache *claims* to describe.
    They do not prove that the cached units, protection flags, or targets were
    actually derived from that file.  This digest deliberately covers the full
    deterministic snapshot payload so a forged/stale editable allowlist cannot
    become the authority for later model input or export audits.
    """

    payload = json.dumps(
        snapshot.to_dict(),
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def get_docx_unit_edit_eligibility_evidence_digest(unit: DocxTextUnit) -> str:
    evidence = unit.edit_eligibility_evidence
    if not isinstance(evidence, dict):
        return ""
    return str(evidence.get("evidenceDigest", "") or "").strip().lower()


def validate_docx_unit_edit_eligibility_evidence(unit: DocxTextUnit) -> bool:
    evidence = unit.edit_eligibility_evidence
    if not isinstance(evidence, dict):
        return False
    actual_digest = str(evidence.get("evidenceDigest", "") or "").strip().lower()
    if not actual_digest:
        return False
    unsigned = {key: value for key, value in evidence.items() if key != "evidenceDigest"}
    return bool(
        actual_digest == _sha256_json_payload(unsigned)
        and int(evidence.get("policyVersion", 0) or 0) == DOCX_STRUCTURAL_ROLE_POLICY_VERSION
        and str(evidence.get("structuralRole", "")) == unit.structural_role
        and str(evidence.get("editEligibility", "")) == unit.edit_eligibility
        and str(evidence.get("sourceTextSha256", ""))
        == hashlib.sha256(unit.source_text().encode("utf-8")).hexdigest()
        and str(evidence.get("targetDigest", "")) == _sha256_json_payload(unit.target)
        and str(evidence.get("presentationSignalsDigest", ""))
        == _sha256_json_payload(unit.presentation_signals)
    )


def get_docx_structural_role_map_digest(snapshot: DocxSnapshot) -> str:
    """Bind every model-facing and explicitly protected structural decision."""

    unit_payloads = []
    for unit in (*snapshot.units, *snapshot.protected_structural_units):
        unit_payloads.append(
            {
                "unitIndex": int(unit.unit_index),
                "target": dict(unit.target),
                "structuralRole": str(unit.structural_role),
                "editEligibility": str(unit.edit_eligibility),
                "editable": bool(unit.editable),
                "evidenceDigest": get_docx_unit_edit_eligibility_evidence_digest(unit),
                "sourceTextSha256": hashlib.sha256(unit.source_text().encode("utf-8")).hexdigest(),
            }
        )
    payload = {
        "policyVersion": int(snapshot.structural_role_policy_version),
        "structuralInventoryVersion": int(snapshot.structural_inventory_version),
        "units": unit_payloads,
    }
    encoded = json.dumps(
        payload,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def verify_docx_snapshot_derivation(
    snapshot: DocxSnapshot,
    source_path: Path,
) -> tuple[dict[str, Any], DocxSnapshot]:
    """Re-derive the authoritative scope from the DOCX and compare the cache.

    The persisted JSON snapshot is an acceleration artifact, never an authority
    boundary.  Hard gates call this function before model use/export and reject
    any difference between the cached allowlist and a fresh parse of the source
    OOXML.  Reports expose only hashes and unit indexes, never document text.
    """

    normalized_source = source_path.resolve()
    authoritative_snapshot = build_docx_snapshot(normalized_source)
    cached_digest = _docx_snapshot_derivation_digest(snapshot)
    authoritative_digest = _docx_snapshot_derivation_digest(authoritative_snapshot)
    cached_units = [
        unit.to_dict()
        for unit in (*snapshot.units, *snapshot.protected_structural_units)
    ]
    authoritative_units = [
        unit.to_dict()
        for unit in (
            *authoritative_snapshot.units,
            *authoritative_snapshot.protected_structural_units,
        )
    ]
    mismatch_unit_indexes = [
        index
        for index in range(max(len(cached_units), len(authoritative_units)))
        if (
            index >= len(cached_units)
            or index >= len(authoritative_units)
            or cached_units[index] != authoritative_units[index]
        )
    ]
    authoritative_current = _is_snapshot_current(authoritative_snapshot, normalized_source)
    return (
        {
            "ok": bool(authoritative_current and cached_digest == authoritative_digest),
            "sourcePath": str(normalized_source),
            "cachedDigest": cached_digest,
            "authoritativeDigest": authoritative_digest,
            "cachedUnitCount": len(cached_units),
            "authoritativeUnitCount": len(authoritative_units),
            "mismatchUnitIndexes": mismatch_unit_indexes[:40],
            "truncatedMismatchUnitIndexes": max(0, len(mismatch_unit_indexes) - 40),
            "authoritativeSourceCurrent": authoritative_current,
        },
        authoritative_snapshot,
    )


def build_docx_scope_diagnostics(
    snapshot: DocxSnapshot,
    *,
    snapshot_path: Path | None = None,
) -> dict[str, Any]:
    units = list(snapshot.units)
    start_index = _find_rewrite_scope_start_index(units)
    end_index = _find_rewrite_scope_end_index(units, start_index) if start_index is not None else None
    acknowledgement_index = (
        next(
            (
                index
                for index, unit in enumerate(units[start_index:], start=start_index)
                if _looks_like_acknowledgement_heading(unit.text)
            ),
            None,
        )
        if start_index is not None
        else None
    )
    boundary_index = (
        _find_rewrite_scope_boundary_index(
            units,
            start_index=start_index,
            acknowledgement_index=acknowledgement_index,
        )
        if start_index is not None
        else None
    )
    reason_counts: dict[str, int] = {}
    for unit in units:
        reason = "editable" if unit.editable else str(unit.protect_reason or "protected")
        reason_counts[reason] = reason_counts.get(reason, 0) + 1

    issues = _find_scope_diagnostic_issues(
        units,
        start_index=start_index,
        end_index=end_index,
        acknowledgement_index=acknowledgement_index,
        boundary_index=boundary_index,
    )
    if not snapshot.semantic_range_topology_valid:
        issues.insert(
            0,
            {
                "code": "semantic_range_topology_invalid",
                "severity": "error",
                "message": "Bookmark/comment range markers are unmatched, duplicated, reversed, or outside a paragraph; model scope is fail-closed.",
                "issueCount": snapshot.semantic_range_issue_count,
                "issueCodes": list(snapshot.semantic_range_issue_codes),
            },
        )
    error_count = sum(1 for issue in issues if issue.get("severity") == "error")
    warning_count = sum(1 for issue in issues if issue.get("severity") == "warning")
    return {
        "version": DOCX_SCOPE_DIAGNOSTICS_VERSION,
        "ok": error_count == 0,
        "sourcePath": snapshot.source_path,
        "snapshotPath": str(snapshot_path.resolve()) if snapshot_path is not None else "",
        "totalTextUnitCount": snapshot.total_text_unit_count,
        "editableUnitCount": snapshot.editable_unit_count,
        "protectedUnitCount": max(0, snapshot.total_text_unit_count - snapshot.editable_unit_count),
        "semanticRangeCount": snapshot.semantic_range_count,
        "bookmarkRangeCount": snapshot.bookmark_range_count,
        "commentRangeCount": snapshot.comment_range_count,
        "semanticRangeTopologyValid": snapshot.semantic_range_topology_valid,
        "semanticRangeIssueCount": snapshot.semantic_range_issue_count,
        "semanticRangeIssueCodes": list(snapshot.semantic_range_issue_codes),
        "semanticRangeCoveredUnitCount": sum(
            1 for unit in units if bool(unit.inside_comment_range)
        ),
        "editableSemanticRangeCoveredUnitCount": sum(
            1 for unit in units if bool(unit.inside_comment_range and unit.editable)
        ),
        "bookmarkRangeInteriorUnitCount": sum(
            1 for unit in units if bool(unit.inside_bookmark_range)
        ),
        "editableBookmarkRangeInteriorUnitCount": sum(
            1 for unit in units if bool(unit.inside_bookmark_range and unit.editable)
        ),
        "commentRangeInteriorUnitCount": sum(
            1 for unit in units if bool(unit.inside_comment_range)
        ),
        "editableCommentRangeInteriorUnitCount": sum(
            1 for unit in units if bool(unit.inside_comment_range and unit.editable)
        ),
        "semanticRangeAnchorUnitCount": sum(
            1 for unit in units if bool(unit.has_semantic_range_anchor)
        ),
        "editableSemanticRangeAnchorUnitCount": sum(
            1 for unit in units if bool(unit.has_semantic_range_anchor and unit.editable)
        ),
        "bookmarkRangeAnchorUnitCount": sum(
            1 for unit in units if bool(unit.has_bookmark_range_anchor)
        ),
        "commentRangeAnchorUnitCount": sum(
            1 for unit in units if bool(unit.has_comment_range_anchor)
        ),
        "structuralRolePolicyVersion": snapshot.structural_role_policy_version,
        "structuralInventoryVersion": snapshot.structural_inventory_version,
        "protectedStructuralUnitCount": len(snapshot.protected_structural_units),
        "protectedTableParagraphCount": sum(
            1
            for unit in snapshot.protected_structural_units
            if unit.structural_role == "table_content"
        ),
        "templateInstructionUnitCount": sum(
            1 for unit in units if unit.structural_role == "template_instruction"
        ),
        "editableTemplateInstructionUnitCount": sum(
            1
            for unit in units
            if unit.structural_role == "template_instruction" and unit.editable
        ),
        "reasonCounts": dict(sorted(reason_counts.items())),
        "scope": {
            "startIndex": start_index,
            "startReason": _scope_start_reason(units[start_index]) if start_index is not None else "",
            "startUnit": _scope_unit_summary(units[start_index]) if start_index is not None else None,
            "endIndex": end_index,
            "endReason": "before_back_matter_boundary" if boundary_index is not None else "document_end",
            "endUnit": _scope_unit_summary(units[end_index]) if end_index is not None else None,
            "acknowledgementIndex": acknowledgement_index,
            "acknowledgementUnit": _scope_unit_summary(units[acknowledgement_index]) if acknowledgement_index is not None else None,
            "postAcknowledgementBoundaryIndex": boundary_index,
            "postAcknowledgementBoundaryUnit": _scope_unit_summary(units[boundary_index]) if boundary_index is not None else None,
        },
        "issueCount": len(issues),
        "errorCount": error_count,
        "warningCount": warning_count,
        "issues": issues[:80],
        "truncatedIssues": max(0, len(issues) - 80),
        "units": [_scope_unit_summary(unit) for unit in units],
    }


def ensure_docx_processing_assets(
    source_path: Path,
    *,
    extracted_path: Path | None = None,
    snapshot_path: Path | None = None,
    scope_diagnostics_path: Path | None = None,
) -> tuple[Path, Path, DocxSnapshot]:
    normalized_source = source_path.resolve()
    extracted_path = extracted_path or get_docx_extracted_text_path(normalized_source)
    snapshot_path = snapshot_path or get_docx_snapshot_path(normalized_source)
    scope_diagnostics_path = scope_diagnostics_path or get_docx_scope_diagnostics_path(normalized_source)

    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None or not _is_snapshot_current(snapshot, normalized_source):
        # Pre-identity assets used only the basename and could therefore belong
        # to a different same-named document. The embedded snapshot provenance
        # is the sole safe migration signal; legacy extracted/scope files do not
        # carry enough identity and are deliberately never reused.
        legacy_snapshot_path = _get_legacy_docx_snapshot_path(normalized_source)
        legacy_snapshot = _load_docx_snapshot(legacy_snapshot_path)
        if legacy_snapshot is not None and _is_snapshot_current(legacy_snapshot, normalized_source):
            snapshot = legacy_snapshot
        else:
            snapshot = build_docx_snapshot(normalized_source)
        snapshot_path.parent.mkdir(parents=True, exist_ok=True)
        snapshot_path.write_text(
            json.dumps(snapshot.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    scope_diagnostics_path.parent.mkdir(parents=True, exist_ok=True)
    scope_diagnostics_path.write_text(
        json.dumps(
            build_docx_scope_diagnostics(snapshot, snapshot_path=snapshot_path),
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    snapshot_mtime_ns = snapshot_path.stat().st_mtime_ns if snapshot_path.exists() else 0
    should_refresh_extracted = (
        not extracted_path.exists()
        or extracted_path.stat().st_mtime_ns < snapshot_mtime_ns
    )
    if should_refresh_extracted:
        extracted_text = "\n\n".join(snapshot.editable_texts())
        extracted_path.parent.mkdir(parents=True, exist_ok=True)
        extracted_path.write_text(extracted_text, encoding="utf-8")
    return extracted_path, snapshot_path, snapshot


def rebuild_docx_from_snapshot(
    rewritten_paragraphs: Sequence[str],
    *,
    source_path: Path,
    snapshot_path: Path,
    export_path: Path,
    preserve_format: bool = False,
) -> None:
    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        raise ValueError(f"DOCX snapshot not found: {snapshot_path}")

    derivation_report, authoritative_snapshot = verify_docx_snapshot_derivation(
        snapshot,
        source_path,
    )
    if not bool(derivation_report.get("ok")):
        mismatch_indexes = derivation_report.get("mismatchUnitIndexes", [])
        raise ValueError(
            "DOCX snapshot does not match a fresh source-derived structural-role map. "
            "Hard failure: refusing snapshot-based export. "
            f"Mismatched unit indexes: {mismatch_indexes!r}."
        )

    # The persisted snapshot is only a cache.  Even after equality has been
    # established, resolve export targets from the freshly parsed authority so
    # every direct caller (including the CLI) shares the production trust
    # boundary and cannot promote a self-consistent forged editable unit.
    editable_units = authoritative_snapshot.editable_units()
    if len(rewritten_paragraphs) != len(editable_units):
        raise ValueError(
            "Rewritten paragraph count does not match the DOCX editable unit count. "
            f"Expected {len(editable_units)}, got {len(rewritten_paragraphs)}. "
            "Hard failure: do not silently fall back to the source text."
        )

    document = Document(str(source_path.resolve()))
    expected_targets: list[tuple[dict[str, Any], str]] = []
    for unit, rewritten_text in zip(editable_units, rewritten_paragraphs):
        paragraph = _resolve_target_paragraph(document, unit.target)
        restored_text = _restore_rewritten_text_with_source_whitespace(
            rewritten_text,
            original_text=unit.text,
            leading_whitespace=unit.leading_whitespace,
            trailing_whitespace=unit.trailing_whitespace,
        )
        _replace_paragraph_text(paragraph, restored_text)
        _polish_rewritten_paragraph(paragraph, preserve_format=preserve_format)
        _verify_rewritten_paragraph_text(paragraph, restored_text, target=unit.target)
        expected_targets.append((dict(unit.target), restored_text))

    export_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(export_path))
    _verify_saved_rewritten_targets(export_path, expected_targets)


def rebuild_docx_from_body_map_units(
    body_map_units: Sequence[Any],
    *,
    source_path: Path,
    export_path: Path,
    preserve_format: bool = False,
) -> None:
    if not body_map_units:
        raise ValueError("DOCX body map has no editable units to export.")

    # Treat a body map as an untrusted transport object. Direct callers do not
    # necessarily pass through the Web contract gate, so re-derive the complete
    # editable allowlist from source OOXML and require an exact, ordered binding
    # before opening a mutable output document.
    authoritative_snapshot = build_docx_snapshot(source_path.resolve())
    authoritative_units = authoritative_snapshot.editable_units()
    if len(body_map_units) != len(authoritative_units):
        raise ValueError(
            "DOCX body map does not match the fresh source-derived editable unit count. "
            "Hard failure: refusing body-map export."
        )

    for position, (unit, authoritative_unit) in enumerate(
        zip(body_map_units, authoritative_units)
    ):
        target = getattr(unit, "target", None)
        target_matches = bool(
            isinstance(target, dict)
            and json.dumps(target, ensure_ascii=False, sort_keys=True)
            == json.dumps(authoritative_unit.target, ensure_ascii=False, sort_keys=True)
        )
        original_text = str(getattr(unit, "original_text", ""))
        leading_whitespace = str(getattr(unit, "leading_whitespace", ""))
        trailing_whitespace = str(getattr(unit, "trailing_whitespace", ""))
        structural_role = str(getattr(unit, "structural_role", "") or "").strip()
        edit_eligibility = str(getattr(unit, "edit_eligibility", "") or "").strip()
        eligibility_evidence_digest = str(
            getattr(unit, "edit_eligibility_evidence_digest", "") or ""
        ).strip()
        if not (
            int(getattr(unit, "unit_index", -1)) == authoritative_unit.unit_index
            and target_matches
            and original_text == authoritative_unit.text
            and leading_whitespace == authoritative_unit.leading_whitespace
            and trailing_whitespace == authoritative_unit.trailing_whitespace
            and structural_role == authoritative_unit.structural_role
            and edit_eligibility == authoritative_unit.edit_eligibility
            and eligibility_evidence_digest
            == get_docx_unit_edit_eligibility_evidence_digest(authoritative_unit)
        ):
            raise ValueError(
                "DOCX body map unit is not bound to the fresh source-derived prose authority. "
                f"Hard failure at editable position {position}."
            )

    document = Document(str(source_path.resolve()))
    seen_targets: set[str] = set()
    expected_targets: list[tuple[dict[str, Any], str]] = []
    for unit in body_map_units:
        target = getattr(unit, "target", None)
        if not isinstance(target, dict):
            raise ValueError("DOCX body map contains an invalid target.")
        structural_role = str(getattr(unit, "structural_role", "") or "").strip()
        edit_eligibility = str(getattr(unit, "edit_eligibility", "") or "").strip()
        eligibility_evidence_digest = str(
            getattr(unit, "edit_eligibility_evidence_digest", "") or ""
        ).strip()
        if (
            structural_role not in DOCX_EDITABLE_STRUCTURAL_ROLES
            or edit_eligibility != "eligible"
            or not eligibility_evidence_digest
        ):
            raise ValueError(
                "DOCX body map export may only rewrite source-certified prose units."
            )
        if str(target.get("kind", "")) != "paragraph":
            raise ValueError("DOCX body map export may only rewrite top-level body paragraphs.")
        target_key = json.dumps(target, ensure_ascii=False, sort_keys=True)
        if target_key in seen_targets:
            raise ValueError(f"DOCX body map contains a duplicate target: {target}")
        seen_targets.add(target_key)
        current_text = str(getattr(unit, "current_text", ""))
        if "\n" in current_text or "\r" in current_text:
            raise ValueError("DOCX body map export rejected a rewritten paragraph containing inline line breaks.")
        restored_text = _restore_rewritten_text_with_source_whitespace(
            current_text,
            original_text=str(getattr(unit, "original_text", "")),
            leading_whitespace=str(getattr(unit, "leading_whitespace", "")),
            trailing_whitespace=str(getattr(unit, "trailing_whitespace", "")),
        )
        paragraph = _resolve_target_paragraph(document, target)
        _replace_paragraph_text(paragraph, restored_text)
        _polish_rewritten_paragraph(paragraph, preserve_format=preserve_format)
        _verify_rewritten_paragraph_text(paragraph, restored_text, target=target)
        expected_targets.append((dict(target), restored_text))

    export_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(export_path))
    _verify_saved_rewritten_targets(export_path, expected_targets)


def _split_text_into_blocks(text: str) -> list[str]:
    blocks: list[str] = []
    current: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")
        if not line.strip():
            if current:
                blocks.append(" ".join(current).strip())
                current = []
            continue
        current.append(line.strip())
    if current:
        blocks.append(" ".join(current).strip())
    return blocks


def _read_paragraphs_file(path: Path) -> list[str]:
    if path.suffix.lower() == ".json":
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, list) or not all(isinstance(item, str) for item in data):
            raise SystemExit("Paragraph json must be a string array.")
        return data
    text = path.read_text(encoding="utf-8")
    return _split_text_into_blocks(text)


def _load_docx_snapshot(path: Path) -> DocxSnapshot | None:
    if not path.exists():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(data, dict):
        return None
    raw_units = data.get("units")
    if not isinstance(raw_units, list):
        return None
    units = [
        unit
        for index, raw_unit in enumerate(raw_units)
        if isinstance(raw_unit, dict)
        and (unit := _docx_text_unit_from_payload(raw_unit, fallback_index=index)) is not None
    ]
    raw_protected_structural_units = data.get("protected_structural_units", [])
    protected_structural_units = [
        unit
        for index, raw_unit in enumerate(raw_protected_structural_units)
        if isinstance(raw_unit, dict)
        and (
            unit := _docx_text_unit_from_payload(
                raw_unit,
                fallback_index=len(units) + index,
            )
        ) is not None
    ] if isinstance(raw_protected_structural_units, list) else []
    return DocxSnapshot(
        version=int(data.get("version", 0)),
        source_path=str(data.get("source_path", "")),
        source_size=int(data.get("source_size", 0)),
        source_mtime_ns=int(data.get("source_mtime_ns", 0)),
        source_sha256=str(data.get("source_sha256", "") or ""),
        editable_unit_count=sum(1 for unit in units if unit.editable),
        total_text_unit_count=int(data.get("total_text_unit_count", len(units))),
        units=units,
        semantic_range_count=int(data.get("semantic_range_count", 0)),
        bookmark_range_count=int(data.get("bookmark_range_count", 0)),
        comment_range_count=int(data.get("comment_range_count", 0)),
        semantic_range_topology_valid=bool(data.get("semantic_range_topology_valid", True)),
        semantic_range_issue_count=int(data.get("semantic_range_issue_count", 0)),
        semantic_range_issue_codes=[
            str(code)
            for code in data.get("semantic_range_issue_codes", [])
            if str(code)
        ] if isinstance(data.get("semantic_range_issue_codes", []), list) else [],
        structural_role_policy_version=int(data.get("structural_role_policy_version", 0)),
        structural_inventory_version=int(data.get("structural_inventory_version", 0)),
        protected_structural_units=protected_structural_units,
    )


def _docx_text_unit_from_payload(raw_unit: dict[str, Any], *, fallback_index: int) -> DocxTextUnit | None:
    unit = DocxTextUnit(
        unit_index=int(raw_unit.get("unit_index", fallback_index)),
        target=dict(raw_unit.get("target", {})) if isinstance(raw_unit.get("target"), dict) else {},
        text=str(raw_unit.get("text", "")),
        style_name=str(raw_unit.get("style_name", "")),
        leading_whitespace=str(raw_unit.get("leading_whitespace", "")),
        trailing_whitespace=str(raw_unit.get("trailing_whitespace", "")),
        editable=False,
        protect_reason=str(raw_unit.get("protect_reason")) if raw_unit.get("protect_reason") is not None else None,
        structural_role=str(raw_unit.get("structural_role", "unknown") or "unknown"),
        edit_eligibility=str(raw_unit.get("edit_eligibility", "protected") or "protected"),
        edit_eligibility_evidence=(
            dict(raw_unit.get("edit_eligibility_evidence", {}))
            if isinstance(raw_unit.get("edit_eligibility_evidence"), dict)
            else {}
        ),
        presentation_signals=(
            dict(raw_unit.get("presentation_signals", {}))
            if isinstance(raw_unit.get("presentation_signals"), dict)
            else {}
        ),
        has_field_code=bool(raw_unit.get("has_field_code", False)),
        has_drawing=bool(raw_unit.get("has_drawing", False)),
        has_math=bool(raw_unit.get("has_math", False)),
        has_complex_inline=bool(raw_unit.get("has_complex_inline", False)),
        has_semantic_range_anchor=bool(raw_unit.get("has_semantic_range_anchor", False)),
        inside_semantic_range=bool(raw_unit.get("inside_semantic_range", False)),
        has_bookmark_range_anchor=bool(raw_unit.get("has_bookmark_range_anchor", False)),
        has_comment_range_anchor=bool(raw_unit.get("has_comment_range_anchor", False)),
        inside_bookmark_range=bool(raw_unit.get("inside_bookmark_range", False)),
        inside_comment_range=bool(raw_unit.get("inside_comment_range", False)),
        has_semantic_point_reference=bool(raw_unit.get("has_semantic_point_reference", False)),
        has_numbering=bool(raw_unit.get("has_numbering", False)),
        numbering_level=_as_optional_int(raw_unit.get("numbering_level")),
        outline_level=_as_optional_int(raw_unit.get("outline_level")),
        format_anchors=[
            dict(item)
            for item in raw_unit.get("format_anchors", [])
            if isinstance(item, dict)
        ] if isinstance(raw_unit.get("format_anchors", []), list) else [],
        format_anchor_ambiguous=bool(raw_unit.get("format_anchor_ambiguous", False)),
    )
    unit.editable = bool(
        unit.structural_role in DOCX_EDITABLE_STRUCTURAL_ROLES
        and unit.edit_eligibility == "eligible"
        and validate_docx_unit_edit_eligibility_evidence(unit)
    )
    if unit.editable:
        unit.protect_reason = None
    return unit


def _is_snapshot_current(snapshot: DocxSnapshot, source_path: Path) -> bool:
    try:
        stat = source_path.stat()
    except OSError:
        return False
    return (
        snapshot.version == DOCX_SNAPSHOT_VERSION
        and
        snapshot.structural_role_policy_version == DOCX_STRUCTURAL_ROLE_POLICY_VERSION
        and snapshot.structural_inventory_version == DOCX_STRUCTURAL_INVENTORY_VERSION
        and all(
            validate_docx_unit_edit_eligibility_evidence(unit)
            for unit in (*snapshot.units, *snapshot.protected_structural_units)
        )
        and
        Path(snapshot.source_path).resolve() == source_path.resolve()
        and snapshot.source_size == stat.st_size
        and snapshot.source_mtime_ns == stat.st_mtime_ns
        and bool(snapshot.source_sha256)
        and snapshot.source_sha256 == _sha256_file(source_path)
    )


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _scan_document_semantic_ranges(document: DocxDocument) -> dict[str, Any]:
    """Resolve kind-aware document ranges without reading prose.

    Bookmark and comment ranges have different rewrite semantics. A comment
    range denotes reviewed text, so its marker-free cross-paragraph interior
    remains protected. A bookmark is primarily a named navigation boundary:
    marker paragraphs stay frozen, while a marker-free interior paragraph can
    be rewritten without moving either boundary node. The scanner inventories
    both kinds separately so policy never infers this distinction from names or
    document text.

    The returned evidence contains only counts, issue codes and paragraph
    indexes. No bookmark name, comment body or document text is serialized.
    """

    body = document._element.body
    range_kinds = ("bookmark", "comment")
    active: dict[tuple[str, str], None] = {}
    seen_starts: set[tuple[str, str]] = set()
    seen_ends: set[tuple[str, str]] = set()
    covered_by_kind = {kind: set() for kind in range_kinds}
    interior_by_kind = {kind: set() for kind in range_kinds}
    anchor_by_kind = {kind: set() for kind in range_kinds}
    completed_by_kind = {kind: 0 for kind in range_kinds}
    issues: list[str] = []
    top_level_paragraph_index = 0

    def add_issue(code: str) -> None:
        issues.append(code)

    def marker_has_paragraph_ancestor(node: Any) -> bool:
        parent = node.getparent()
        while parent is not None and parent is not body:
            if parent.tag == qn("w:p"):
                return True
            parent = parent.getparent()
        return False

    for block in list(body):
        is_top_level_paragraph = block.tag == qn("w:p")
        active_before = set(active)
        marker_kinds: set[str] = set()
        for node in block.iter():
            marker = SEMANTIC_RANGE_MARKER_ACTIONS.get(node.tag)
            if marker is None:
                continue
            kind, action = marker
            marker_kinds.add(kind)
            if not marker_has_paragraph_ancestor(node):
                add_issue("semantic_range_marker_outside_paragraph")

            marker_id = str(node.get(qn("w:id")) or "")
            if not marker_id:
                add_issue("semantic_range_missing_id")
                continue
            key = (kind, marker_id)

            if action == "start":
                if key in seen_starts or key in active:
                    add_issue("semantic_range_duplicate_start")
                    continue
                if key in seen_ends:
                    add_issue("semantic_range_reversed")
                    continue
                seen_starts.add(key)
                active[key] = None
                continue

            if key in seen_ends:
                add_issue("semantic_range_duplicate_end")
                continue
            seen_ends.add(key)
            if key not in active:
                add_issue("semantic_range_reversed")
                continue
            active.pop(key, None)
            completed_by_kind[kind] += 1

        if is_top_level_paragraph:
            active_after = set(active)
            for kind in range_kinds:
                active_before_kind = any(key[0] == kind for key in active_before)
                active_after_kind = any(key[0] == kind for key in active_after)
                if kind in marker_kinds:
                    anchor_by_kind[kind].add(top_level_paragraph_index)
                if kind in marker_kinds or active_before_kind or active_after_kind:
                    covered_by_kind[kind].add(top_level_paragraph_index)
                if (
                    kind not in marker_kinds
                    and active_before_kind
                    and active_after_kind
                ):
                    interior_by_kind[kind].add(top_level_paragraph_index)
            top_level_paragraph_index += 1

    if active:
        issues.extend("semantic_range_unmatched_start" for _key in active)

    covered_paragraph_indexes = set().union(*covered_by_kind.values())
    issue_codes = sorted(set(issues))
    return {
        "topologyValid": not issues,
        "rangeCount": sum(completed_by_kind.values()),
        "bookmarkRangeCount": completed_by_kind["bookmark"],
        "commentRangeCount": completed_by_kind["comment"],
        "paragraphIndexes": sorted(covered_paragraph_indexes),
        "bookmarkParagraphIndexes": sorted(covered_by_kind["bookmark"]),
        "commentParagraphIndexes": sorted(covered_by_kind["comment"]),
        "bookmarkInteriorParagraphIndexes": sorted(interior_by_kind["bookmark"]),
        "commentInteriorParagraphIndexes": sorted(interior_by_kind["comment"]),
        "bookmarkAnchorParagraphIndexes": sorted(anchor_by_kind["bookmark"]),
        "commentAnchorParagraphIndexes": sorted(anchor_by_kind["comment"]),
        "issueCount": len(issues),
        "issueCodes": issue_codes,
    }


def _sha256_json_payload(value: Any) -> str:
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _assign_unit_edit_decision(
    unit: DocxTextUnit,
    *,
    structural_role: str,
    edit_eligibility: str,
    reason_codes: Sequence[str],
    protect_reason: str | None,
) -> None:
    normalized_role = str(structural_role or "unknown").strip().lower() or "unknown"
    requested_eligibility = str(edit_eligibility or "protected").strip().lower()
    normalized_eligibility = (
        "eligible"
        if requested_eligibility == "eligible" and normalized_role in DOCX_EDITABLE_STRUCTURAL_ROLES
        else "protected"
    )
    normalized_reasons = sorted({str(code).strip() for code in reason_codes if str(code).strip()})
    presentation_digest = _sha256_json_payload(unit.presentation_signals)
    target_digest = _sha256_json_payload(unit.target)
    evidence = {
        "schemaVersion": 1,
        "policyVersion": DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
        "structuralRole": normalized_role,
        "editEligibility": normalized_eligibility,
        "reasonCodes": normalized_reasons,
        "sourceTextSha256": hashlib.sha256(unit.source_text().encode("utf-8")).hexdigest(),
        "targetDigest": target_digest,
        "presentationSignalsDigest": presentation_digest,
    }
    evidence["evidenceDigest"] = _sha256_json_payload(evidence)
    unit.structural_role = normalized_role
    unit.edit_eligibility = normalized_eligibility
    unit.edit_eligibility_evidence = evidence
    unit.editable = normalized_eligibility == "eligible"
    unit.protect_reason = None if unit.editable else (protect_reason or normalized_role)


def _length_pt(value: Any) -> float | None:
    if value is None:
        return None
    try:
        points = float(value.pt)
    except (AttributeError, TypeError, ValueError):
        return None
    return round(points, 3)


def _style_chain_identity(style: Any) -> str:
    """Return a stable key for a python-docx style proxy.

    ``base_style`` may create a short-lived proxy on every access. Using the
    proxy's Python ``id`` as a visited key is therefore nondeterministic because
    CPython can reuse the released proxy address while walking the same chain.
    Style IDs are document-stable and unique; name/element identity are only
    defensive fallbacks for malformed custom styles.
    """

    style_id = str(getattr(style, "style_id", "") or "").strip()
    if style_id:
        return f"style-id:{style_id}"
    style_name = str(getattr(style, "name", "") or "").strip()
    if style_name:
        return f"style-name:{style_name}"
    element = getattr(style, "_element", None)
    return f"style-element:{id(element) if element is not None else id(style)}"


def _effective_paragraph_format_property(paragraph: Paragraph, property_name: str) -> Any:
    try:
        value = getattr(paragraph.paragraph_format, property_name, None)
    except Exception:
        value = None
    if value is not None:
        return value
    style = paragraph.style
    visited: set[str] = set()
    while style is not None and _style_chain_identity(style) not in visited:
        visited.add(_style_chain_identity(style))
        try:
            value = getattr(style.paragraph_format, property_name, None)
        except Exception:
            value = None
        if value is not None:
            return value
        style = style.base_style
    return None


def _effective_paragraph_style_font_property(paragraph: Paragraph, property_name: str) -> Any:
    style = paragraph.style
    visited: set[str] = set()
    while style is not None and _style_chain_identity(style) not in visited:
        visited.add(_style_chain_identity(style))
        value = getattr(getattr(style, "font", None), property_name, None)
        if value is not None:
            return value
        style = style.base_style
    return None


def _paragraph_style_chain(paragraph: Paragraph) -> list[str]:
    values: list[str] = []
    style = paragraph.style
    visited: set[str] = set()
    while style is not None and _style_chain_identity(style) not in visited:
        visited.add(_style_chain_identity(style))
        style_id = str(getattr(style, "style_id", "") or "").strip()
        style_name = str(getattr(style, "name", "") or "").strip()
        values.append(f"{style_id}:{style_name}")
        style = style.base_style
    return values


def _paragraph_presentation_signals(paragraph: Paragraph) -> dict[str, Any]:
    """Return text-independent paragraph/run presentation evidence."""

    alignment_value = _effective_paragraph_format_property(paragraph, "alignment")
    try:
        alignment = int(alignment_value) if alignment_value is not None else None
    except (TypeError, ValueError):
        alignment = None

    text_runs = _iter_paragraph_text_runs(paragraph)
    visible_runs = [(run, value) for run, value in text_runs if value]
    size_coverage: dict[float, int] = {}
    bold_values: list[bool] = []
    paragraph_style_size = _effective_paragraph_style_font_property(paragraph, "size")
    paragraph_style_bold = _effective_paragraph_style_font_property(paragraph, "bold")
    for run, value in visible_runs:
        size = _effective_run_font_property(run, "size")
        if size is None:
            size = paragraph_style_size
        size_pt = _length_pt(size)
        if size_pt is not None:
            size_coverage[size_pt] = size_coverage.get(size_pt, 0) + max(1, len(value))
        bold = _effective_run_font_property(run, "bold")
        if bold is None:
            bold = paragraph_style_bold
        bold_values.append(bool(bold is True))

    dominant_size = None
    if size_coverage:
        dominant_size = max(size_coverage, key=lambda size: (size_coverage[size], size))
    return {
        "alignment": alignment,
        "keepWithNext": bool(_effective_paragraph_format_property(paragraph, "keep_with_next") is True),
        "keepTogether": bool(_effective_paragraph_format_property(paragraph, "keep_together") is True),
        "pageBreakBefore": bool(_effective_paragraph_format_property(paragraph, "page_break_before") is True),
        "firstLineIndentPt": _length_pt(_effective_paragraph_format_property(paragraph, "first_line_indent")),
        "leftIndentPt": _length_pt(_effective_paragraph_format_property(paragraph, "left_indent")),
        "rightIndentPt": _length_pt(_effective_paragraph_format_property(paragraph, "right_indent")),
        "spaceBeforePt": _length_pt(_effective_paragraph_format_property(paragraph, "space_before")),
        "spaceAfterPt": _length_pt(_effective_paragraph_format_property(paragraph, "space_after")),
        "lineSpacingPt": _length_pt(_effective_paragraph_format_property(paragraph, "line_spacing")),
        "dominantFontSizePt": dominant_size,
        "maximumFontSizePt": max(size_coverage) if size_coverage else None,
        "allVisibleRunsBold": bool(bold_values and all(bold_values)),
        "anyVisibleRunBold": any(bold_values),
        "visibleRunCount": len(visible_runs),
        "styleChain": _paragraph_style_chain(paragraph),
    }


def _collect_text_units(
    document: DocxDocument,
    *,
    semantic_range_paragraph_indexes: set[int] | None = None,
    bookmark_range_interior_paragraph_indexes: set[int] | None = None,
    comment_range_interior_paragraph_indexes: set[int] | None = None,
) -> list[DocxTextUnit]:
    units: list[DocxTextUnit] = []
    top_level_paragraph_index = 0
    covered_paragraph_indexes = semantic_range_paragraph_indexes or set()
    bookmark_interior_indexes = bookmark_range_interior_paragraph_indexes or set()
    comment_interior_indexes = comment_range_interior_paragraph_indexes or set()

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            unit = _build_text_unit(
                block,
                {
                    "kind": "paragraph",
                    "paragraph_index": top_level_paragraph_index,
                },
                unit_index=len(units),
            )
            if unit is not None:
                unit.inside_semantic_range = top_level_paragraph_index in covered_paragraph_indexes
                unit.inside_bookmark_range = top_level_paragraph_index in bookmark_interior_indexes
                unit.inside_comment_range = top_level_paragraph_index in comment_interior_indexes
                units.append(unit)
            top_level_paragraph_index += 1

    return units


def _collect_protected_table_units(
    document: DocxDocument,
    *,
    first_unit_index: int,
) -> list[DocxTextUnit]:
    """Recursively inventory table-cell paragraphs without polluting body scope.

    ``python-docx`` exposes only top-level tables through ``document.tables``.
    Nested tables are reached through their owning cell and receive a stable
    source-derived ``table_path``. Merged cells are deduplicated independently
    at each table depth by their underlying ``w:tc`` identity.
    """

    units: list[DocxTextUnit] = []
    table_body_block_indexes = [
        index
        for index, block in enumerate(list(document.element.body))
        if block.tag == qn("w:tbl")
    ]

    def collect_table(
        table: Table,
        *,
        body_block_index: int,
        top_level_table_index: int,
        table_path: list[dict[str, int]],
    ) -> None:
        # Keep the actual lxml ``w:tc`` objects alive while walking the table.
        # Retaining only ``id(_tc)`` permits CPython to reuse an address for a
        # later proxy and makes merged-cell deduplication nondeterministic.
        seen_cells: set[Any] = set()
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_identity = cell._tc
                if cell_identity in seen_cells:
                    continue
                seen_cells.add(cell_identity)
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    unit = _build_text_unit(
                        paragraph,
                        {
                            "kind": "table_cell_paragraph",
                            "body_block_index": body_block_index,
                            "table_index": top_level_table_index,
                            "table_path": [dict(step) for step in table_path],
                            "table_depth": max(0, len(table_path) - 1),
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "paragraph_index": paragraph_index,
                        },
                        unit_index=first_unit_index + len(units),
                    )
                    if unit is None:
                        continue
                    _assign_unit_edit_decision(
                        unit,
                        structural_role="table_content",
                        edit_eligibility="protected",
                        reason_codes=(
                            "inside_table",
                            "nested_table_content" if len(table_path) > 1 else "top_level_table_content",
                        ),
                        protect_reason="table_content",
                    )
                    units.append(unit)

                for nested_table_index, nested_table in enumerate(cell.tables):
                    collect_table(
                        nested_table,
                        body_block_index=body_block_index,
                        top_level_table_index=top_level_table_index,
                        table_path=[
                            *table_path,
                            {
                                "row_index": row_index,
                                "cell_index": cell_index,
                                "table_index": nested_table_index,
                            },
                        ],
                    )

    for table_index, table in enumerate(document.tables):
        body_block_index = (
            table_body_block_indexes[table_index]
            if table_index < len(table_body_block_indexes)
            else -1
        )
        collect_table(
            table,
            body_block_index=body_block_index,
            top_level_table_index=table_index,
            table_path=[{"table_index": table_index}],
        )
    return units


def _split_boundary_whitespace(text: str) -> tuple[str, str, str]:
    """Split literal text into exact leading whitespace, model core, and suffix.

    Python's Unicode whitespace classification covers the boundary characters
    Word users commonly employ as manual indentation, including ASCII spaces,
    NBSP (U+00A0), full-width spaces (U+3000), and tabs.  Keeping these values
    outside the model-facing core prevents prompt/output normalization from
    silently changing source layout semantics.
    """

    value = str(text)
    start = 0
    while start < len(value) and value[start].isspace():
        start += 1
    end = len(value)
    while end > start and value[end - 1].isspace():
        end -= 1
    return value[:start], value[start:end], value[end:]


def _build_text_unit(paragraph: Paragraph, target: dict[str, Any], *, unit_index: int) -> DocxTextUnit | None:
    leading_whitespace, text, trailing_whitespace = _split_boundary_whitespace(paragraph.text)
    if not text:
        return None
    style_name = paragraph.style.name if paragraph.style is not None else ""
    numbering_level = _paragraph_numbering_level(paragraph)
    outline_level = _paragraph_outline_level(paragraph)
    format_anchors = _extract_format_sensitive_anchors(paragraph)
    presentation_signals = _paragraph_presentation_signals(paragraph)
    return DocxTextUnit(
        unit_index=unit_index,
        target=target,
        text=text,
        style_name=style_name,
        leading_whitespace=leading_whitespace,
        trailing_whitespace=trailing_whitespace,
        editable=False,
        protect_reason="unclassified",
        structural_role="unknown",
        edit_eligibility="protected",
        edit_eligibility_evidence={},
        presentation_signals=presentation_signals,
        has_field_code=_paragraph_has_field_code(paragraph),
        has_drawing=_paragraph_has_drawing(paragraph),
        has_math=_paragraph_has_math(paragraph),
        has_complex_inline=_paragraph_has_complex_inline(paragraph),
        has_semantic_range_anchor=_paragraph_has_semantic_range_anchor(paragraph),
        has_bookmark_range_anchor=_paragraph_has_bookmark_range_anchor(paragraph),
        has_comment_range_anchor=_paragraph_has_comment_range_anchor(paragraph),
        has_semantic_point_reference=_paragraph_has_semantic_point_reference(paragraph),
        has_numbering=numbering_level is not None,
        numbering_level=numbering_level,
        outline_level=outline_level,
        format_anchors=format_anchors,
        format_anchor_ambiguous=any(bool(anchor.get("ambiguous")) for anchor in format_anchors),
    )


def _unit_has_positive_prose_evidence(unit: DocxTextUnit) -> bool:
    text = (unit.text or "").strip()
    if not text:
        return False
    # A long unpunctuated label is not affirmative prose evidence.  This is
    # intentional: formatting-poor university templates frequently encode long
    # headings as Normal paragraphs, so absence of a structural regex may never
    # become an implicit allow decision.
    if _has_sentence_punctuation(text):
        return True
    # A colon commonly separates a thesis heading from its subtitle and is not
    # sentence evidence on its own. For punctuation-poor prose, require at
    # least two comma-delimited clauses; a single weak separator still fails
    # closed as an ambiguous non-prose paragraph.
    comma_count = text.count("，") + text.count(",")
    return len(text) >= 24 and comma_count >= 2


def _median(values: Sequence[float]) -> float | None:
    ordered = sorted(float(value) for value in values)
    if not ordered:
        return None
    middle = len(ordered) // 2
    if len(ordered) % 2:
        return ordered[middle]
    return (ordered[middle - 1] + ordered[middle]) / 2


def _derive_body_font_size_baseline(
    units: Sequence[DocxTextUnit],
    *,
    start_index: int | None,
    end_index: int | None,
) -> float | None:
    if start_index is None:
        return None
    candidates: list[float] = []
    for index, unit in enumerate(units):
        if index < start_index or (end_index is not None and index > end_index):
            continue
        if not _unit_has_positive_prose_evidence(unit):
            continue
        if (
            unit.has_field_code
            or unit.has_drawing
            or unit.has_math
            or unit.has_complex_inline
            or _looks_like_heading(
                unit.text,
                style_name=unit.style_name,
                has_numbering=unit.has_numbering,
                outline_level=unit.outline_level,
            )
            or _looks_like_caption(unit.text, style_name=unit.style_name)
            or _looks_like_note(unit.text)
            or _looks_like_keyword_line(unit.text)
        ):
            continue
        signals = unit.presentation_signals
        if signals.get("alignment") in {1, 2}:
            continue
        size = signals.get("dominantFontSizePt")
        if isinstance(size, (int, float)) and float(size) > 0:
            candidates.append(float(size))
    baseline = _median(candidates)
    return round(baseline, 3) if baseline is not None else None


def _presentation_heading_reason_codes(
    unit: DocxTextUnit,
    *,
    body_font_size_baseline: float | None,
) -> list[str]:
    signals = unit.presentation_signals
    alignment = signals.get("alignment")
    centered_or_right = alignment in {1, 2}
    keep_with_next = bool(signals.get("keepWithNext"))
    page_break_before = bool(signals.get("pageBreakBefore"))
    all_bold = bool(signals.get("allVisibleRunsBold"))
    maximum_size = signals.get("maximumFontSizePt")
    size_value = float(maximum_size) if isinstance(maximum_size, (int, float)) else None
    large_relative_to_body = bool(
        size_value is not None
        and (
            (body_font_size_baseline is not None and size_value >= body_font_size_baseline + 2.0)
            or (body_font_size_baseline is None and size_value >= 14.0)
        )
    )
    has_prose = _unit_has_positive_prose_evidence(unit)

    reasons: list[str] = []
    if keep_with_next:
        reasons.append("paragraph_keep_with_next")
    if page_break_before:
        reasons.append("paragraph_page_break_before")
    if centered_or_right and (not has_prose or all_bold or large_relative_to_body):
        reasons.append("paragraph_centered_or_right_structural")
    if large_relative_to_body and (all_bold or not has_prose):
        reasons.append("font_size_above_body_baseline")
    if all_bold and (centered_or_right or large_relative_to_body or not has_prose):
        reasons.append("all_visible_runs_bold_structural")
    return reasons


def _unit_has_high_confidence_heading_context(unit: DocxTextUnit) -> bool:
    """Recognize a neighboring section label without trusting short text alone."""

    if (
        _looks_like_acknowledgement_heading(unit.text)
        or _unit_looks_like_references_heading(unit)
        or _looks_like_back_matter_heading(unit.text)
        or _looks_like_body_start(unit.text, style_name=unit.style_name)
    ):
        return True
    normalized_style = _normalize_style_name(unit.style_name)
    if unit.outline_level is not None and 0 <= unit.outline_level <= 8:
        return True
    if (
        normalized_style.startswith("heading")
        or normalized_style.startswith("标题")
        or any(
            hint in normalized_style
            for hint in ("chapterheading", "sectionheading", "title", "章标题", "节标题")
        )
    ):
        return True
    return _looks_like_numbered_structure_label(unit.text)


def _has_adjacent_structural_heading_context(
    units: Sequence[DocxTextUnit],
    index: int,
) -> bool:
    return any(
        0 <= neighbor_index < len(units)
        and _unit_has_high_confidence_heading_context(units[neighbor_index])
        for neighbor_index in (index - 1, index + 1)
    )


def _template_instruction_reason_codes(
    units: Sequence[DocxTextUnit],
    index: int,
) -> list[str]:
    """Return evidence for a standalone template-writing directive.

    A bare ``注意`` is deliberately insufficient. University-template prose is
    frozen only when four independent signals agree: an explicit paragraph
    prefix, document-authoring semantics, directive modality, and immediate
    section-heading context. This keeps ordinary academic observations such as
    ``注意：实验过程中需要保持温度稳定`` inside the editable body.
    """

    if index < 0 or index >= len(units):
        return []
    text = (units[index].text or "").strip()
    if not TEMPLATE_INSTRUCTION_PREFIX_RE.search(text):
        return []
    if not TEMPLATE_DOCUMENT_AUTHORING_CUE_RE.search(text):
        return []
    if not TEMPLATE_DIRECTIVE_CUE_RE.search(text):
        return []
    if not _has_adjacent_structural_heading_context(units, index):
        return []
    return [
        "template_instruction_prefix",
        "template_document_authoring_cue",
        "template_directive_cue",
        "adjacent_structural_heading",
    ]


def _acknowledgement_guidance_reason_codes(
    units: Sequence[DocxTextUnit],
    index: int,
) -> list[str]:
    """Recognize meta-guidance immediately following a real acknowledgements heading."""

    if index <= 0 or index >= len(units):
        return []
    text = (units[index].text or "").strip()
    if not _looks_like_acknowledgement_heading(units[index - 1].text):
        return []
    if not ACKNOWLEDGEMENT_GUIDANCE_DEFINITION_RE.search(text):
        return []
    if not ACKNOWLEDGEMENT_GUIDANCE_AUDIENCE_RE.search(text):
        return []
    if not ACKNOWLEDGEMENT_GUIDANCE_ACTION_RE.search(text):
        return []
    return [
        "acknowledgement_guidance",
        "adjacent_acknowledgement_heading",
        "template_document_authoring_cue",
        "template_directive_cue",
    ]


def _apply_protection_rules(
    units: list[DocxTextUnit],
    *,
    semantic_range_topology_valid: bool = True,
) -> None:
    if not units:
        return

    if not semantic_range_topology_valid:
        for unit in units:
            _assign_unit_edit_decision(
                unit,
                structural_role="complex_container",
                edit_eligibility="protected",
                reason_codes=("semantic_range_topology_invalid",),
                protect_reason="semantic_range_topology_invalid",
            )
        return

    start_index = _find_rewrite_scope_start_index(units)
    end_index = _find_rewrite_scope_end_index(units, start_index) if start_index is not None else None
    reference_entry_run_starts = _find_reference_entry_run_starts(units, start_index=start_index or 0)
    body_font_size_baseline = _derive_body_font_size_baseline(
        units,
        start_index=start_index,
        end_index=end_index,
    )
    phase = (
        "abstract"
        if start_index is not None and _looks_like_abstract_start(units[start_index].text)
        else "body"
    )

    for index, unit in enumerate(units):
        role = "unknown"
        eligibility = "protected"
        protect_reason: str | None = "ambiguous_non_prose"
        reason_codes: list[str] = []

        inside_confirmed_scope = bool(
            start_index is not None
            and index >= start_index
            and (end_index is None or index <= end_index)
        )
        is_acknowledgement_heading = _looks_like_acknowledgement_heading(unit.text)
        is_references_heading = _unit_looks_like_references_heading(unit)
        is_back_matter_heading = _looks_like_back_matter_heading(unit.text)

        # Structural phase is document context, not a side effect of whichever
        # per-paragraph protection branch wins. In particular, a real ``致谢``
        # heading may carry a bookmark anchor and be classified as a complex
        # container; its following paragraphs must still enter the frozen
        # acknowledgements phase.
        if inside_confirmed_scope:
            if is_acknowledgement_heading:
                phase = "acknowledgements"
            elif is_references_heading:
                phase = "references"
            elif is_back_matter_heading:
                phase = "back_matter"
            elif phase in {"abstract", "body"} and _looks_like_abstract_start(unit.text):
                phase = "abstract"
            elif (
                phase in {"abstract", "body"}
                and _looks_like_body_start(unit.text, style_name=unit.style_name)
                and not _looks_like_abstract_start(unit.text)
            ):
                phase = "body"

        if start_index is None or index < start_index:
            role = "front_matter"
            protect_reason = "front_matter"
            reason_codes.append("outside_confirmed_body_start")
        elif end_index is not None and index > end_index:
            role = "back_matter"
            protect_reason = "outside_body_scope"
            reason_codes.append("outside_confirmed_body_end")
        elif unit.has_semantic_range_anchor:
            role = "complex_container"
            protect_reason = "semantic_range_anchor"
            reason_codes.append("semantic_range_anchor")
            if unit.has_bookmark_range_anchor:
                reason_codes.append("bookmark_range_anchor")
            if unit.has_comment_range_anchor:
                reason_codes.append("comment_range_anchor")
        elif unit.inside_comment_range:
            role = "complex_container"
            protect_reason = "semantic_range_span"
            reason_codes.extend(("inside_semantic_range", "inside_comment_range"))
        elif unit.has_semantic_point_reference:
            role = "complex_container"
            protect_reason = "semantic_point_reference"
            reason_codes.append("semantic_point_reference")
        elif _looks_like_toc_heading(unit.text):
            role = "toc_heading"
            protect_reason = "generated_field"
            reason_codes.append("toc_heading")
        elif unit.has_field_code or _looks_like_toc_entry(unit.text):
            role = "toc_entry"
            protect_reason = "generated_field"
            reason_codes.append("generated_or_toc_field")
        elif _is_table_unit(unit):
            role = "table_content"
            protect_reason = "table_content"
            reason_codes.append("inside_table")
        elif unit.has_math or _looks_like_formula_paragraph(unit.text, style_name=unit.style_name):
            role = "equation"
            protect_reason = "formula"
            reason_codes.append("equation_or_formula")
        elif unit.has_drawing:
            role = "complex_container"
            protect_reason = "graphic_anchor"
            reason_codes.append("graphic_anchor")
        elif unit.has_complex_inline:
            role = "complex_container"
            protect_reason = "complex_inline"
            reason_codes.append("unsupported_complex_inline")
        elif unit.format_anchor_ambiguous:
            role = "complex_container"
            protect_reason = "ambiguous_format_anchor"
            reason_codes.append("ambiguous_format_anchor")
        elif _format_anchors_cover_source_text(unit.source_text(), unit.format_anchors):
            role = "complex_container"
            protect_reason = "format_sensitive_text"
            reason_codes.append("format_sensitive_text_covers_paragraph")
        elif is_back_matter_heading:
            phase = "back_matter"
            role = "back_matter"
            protect_reason = "back_matter"
            reason_codes.append("back_matter_heading")
        elif is_references_heading:
            phase = "references"
            role = "references_heading"
            protect_reason = "references"
            reason_codes.append("references_heading")
        elif is_acknowledgement_heading:
            phase = "acknowledgements"
            role = "acknowledgement_heading"
            protect_reason = "heading"
            reason_codes.append("acknowledgement_heading")
        elif index in reference_entry_run_starts:
            phase = "references"
            role = "reference_entry"
            protect_reason = "references"
            reason_codes.append("reference_entry_run")
        elif phase == "references":
            role = "reference_entry"
            protect_reason = "references"
            reason_codes.append("inside_references_phase")
        elif phase == "back_matter":
            role = "back_matter"
            protect_reason = "back_matter"
            reason_codes.append("inside_back_matter_phase")
        elif phase == "acknowledgements":
            acknowledgement_guidance_reasons = _acknowledgement_guidance_reason_codes(
                units,
                index,
            )
            if acknowledgement_guidance_reasons:
                role = "template_instruction"
                protect_reason = "template_instruction"
                reason_codes.extend(
                    ("inside_acknowledgement_phase", *acknowledgement_guidance_reasons)
                )
            else:
                role = "acknowledgement_body"
                protect_reason = "acknowledgement_body"
                reason_codes.append("inside_acknowledgement_phase")
        elif _looks_like_heading(
            unit.text,
            style_name=unit.style_name,
            has_numbering=unit.has_numbering,
            outline_level=unit.outline_level,
        ):
            role = "heading"
            protect_reason = "heading"
            reason_codes.append("semantic_heading")
        elif _looks_like_caption(unit.text, style_name=unit.style_name):
            role = "caption"
            protect_reason = "caption"
            reason_codes.append("caption")
        elif _looks_like_note(unit.text):
            role = "note"
            protect_reason = "caption"
            reason_codes.append("figure_or_table_note")
        elif _looks_like_keyword_line(unit.text):
            role = "keywords"
            protect_reason = "structured_field"
            reason_codes.append("keyword_field")
        else:
            template_instruction_reasons = _template_instruction_reason_codes(
                units,
                index,
            )
            presentation_reasons = (
                []
                if template_instruction_reasons
                else _presentation_heading_reason_codes(
                    unit,
                    body_font_size_baseline=body_font_size_baseline,
                )
            )
            if template_instruction_reasons:
                role = "template_instruction"
                protect_reason = "template_instruction"
                reason_codes.extend(template_instruction_reasons)
            elif presentation_reasons:
                role = "heading"
                protect_reason = "heading"
                reason_codes.extend(("presentation_structural_heading", *presentation_reasons))
            elif not _unit_has_positive_prose_evidence(unit):
                role = "ambiguous_non_prose"
                protect_reason = "ambiguous_non_prose"
                reason_codes.append("insufficient_positive_body_evidence")
            elif unit.has_numbering or _looks_like_numbered_body_item(unit.text):
                role = "body_list"
                eligibility = "eligible"
                protect_reason = None
                reason_codes.extend(("inside_confirmed_body_scope", "numbered_sentence_prose"))
            elif phase == "abstract":
                role = "abstract_body"
                eligibility = "eligible"
                protect_reason = None
                reason_codes.extend(("inside_abstract_scope", "sentence_prose_evidence"))
            else:
                role = "body_prose"
                eligibility = "eligible"
                protect_reason = None
                reason_codes.extend(("inside_confirmed_body_scope", "sentence_prose_evidence"))

        if (
            eligibility == "eligible"
            and unit.inside_bookmark_range
            and not unit.has_semantic_range_anchor
        ):
            reason_codes.append("marker_free_bookmark_interior")

        _assign_unit_edit_decision(
            unit,
            structural_role=role,
            edit_eligibility=eligibility,
            reason_codes=reason_codes,
            protect_reason=protect_reason,
        )


def _scope_text_preview(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    if len(normalized) <= SCOPE_DIAGNOSTIC_TEXT_PREVIEW_CHARS:
        return normalized
    return normalized[:SCOPE_DIAGNOSTIC_TEXT_PREVIEW_CHARS].rstrip()


def _scope_unit_summary(unit: DocxTextUnit) -> dict[str, Any]:
    raw_reasons = unit.edit_eligibility_evidence.get("reasonCodes", []) if isinstance(unit.edit_eligibility_evidence, dict) else []
    return {
        "unitIndex": unit.unit_index,
        "target": unit.target,
        "targetKind": str(unit.target.get("kind", "")),
        "styleName": unit.style_name,
        "editable": bool(unit.editable),
        "protectReason": unit.protect_reason or "",
        "structuralRole": unit.structural_role,
        "editEligibility": unit.edit_eligibility,
        "editEligibilityEvidenceDigest": get_docx_unit_edit_eligibility_evidence_digest(unit),
        "editEligibilityReasonCodes": [str(code) for code in raw_reasons if str(code)],
        "presentationSignals": dict(unit.presentation_signals),
        "textLength": len(unit.text),
        "textPreview": _scope_text_preview(unit.text),
        "hasFieldCode": bool(unit.has_field_code),
        "hasDrawing": bool(unit.has_drawing),
        "hasMath": bool(unit.has_math),
        "hasComplexInline": bool(unit.has_complex_inline),
        "hasSemanticRangeAnchor": bool(unit.has_semantic_range_anchor),
        "insideSemanticRange": bool(unit.inside_semantic_range),
        "hasBookmarkRangeAnchor": bool(unit.has_bookmark_range_anchor),
        "hasCommentRangeAnchor": bool(unit.has_comment_range_anchor),
        "insideBookmarkRange": bool(unit.inside_bookmark_range),
        "insideCommentRange": bool(unit.inside_comment_range),
        "hasSemanticPointReference": bool(unit.has_semantic_point_reference),
        "hasNumbering": bool(unit.has_numbering),
        "numberingLevel": unit.numbering_level,
        "outlineLevel": unit.outline_level,
        "formatAnchorCount": len(unit.format_anchors),
        "formatAnchorAmbiguous": bool(unit.format_anchor_ambiguous),
        "flags": _scope_unit_flags(unit),
    }


def _scope_unit_flags(unit: DocxTextUnit) -> dict[str, bool]:
    return {
        "abstractStart": _looks_like_abstract_start(unit.text),
        "bodyStart": _looks_like_body_start(unit.text, style_name=unit.style_name),
        "acknowledgementHeading": _looks_like_acknowledgement_heading(unit.text),
        "referencesHeading": _unit_looks_like_references_heading(unit),
        "referenceEntry": _looks_like_reference_entry(unit.text),
        "backMatterHeading": _looks_like_back_matter_heading(unit.text),
        "tocHeading": _looks_like_toc_heading(unit.text),
        "tocEntry": _looks_like_toc_entry(unit.text),
        "heading": _looks_like_heading(
            unit.text,
            style_name=unit.style_name,
            has_numbering=unit.has_numbering,
            outline_level=unit.outline_level,
        ),
        "numberedBodyItem": _looks_like_numbered_body_item(unit.text),
        "keywordLine": _looks_like_keyword_line(unit.text),
        "caption": _looks_like_caption(unit.text, style_name=unit.style_name),
        "note": _looks_like_note(unit.text),
        "formula": bool(unit.has_math) or _looks_like_formula_paragraph(unit.text, style_name=unit.style_name),
        "templateInstruction": unit.structural_role == "template_instruction",
        "semanticRangeCovered": bool(unit.inside_semantic_range),
        "bookmarkRangeInterior": bool(unit.inside_bookmark_range),
        "commentRangeInterior": bool(unit.inside_comment_range),
    }


def _scope_start_reason(unit: DocxTextUnit) -> str:
    if _looks_like_abstract_start(unit.text):
        return "abstract_marker"
    if _looks_like_body_start(unit.text, style_name=unit.style_name):
        return "body_start_marker"
    return "fallback"


def _add_scope_issue(
    issues: list[dict[str, Any]],
    *,
    code: str,
    severity: str,
    message: str,
    unit: DocxTextUnit | None = None,
) -> None:
    issue: dict[str, Any] = {
        "code": code,
        "severity": severity,
        "message": message,
    }
    if unit is not None:
        issue["unit"] = _scope_unit_summary(unit)
    issues.append(issue)


def _has_sentence_punctuation(text: str) -> bool:
    return any(mark in (text or "") for mark in "。！？；.!?;")


def _looks_like_long_body_text(unit: DocxTextUnit) -> bool:
    stripped = (unit.text or "").strip()
    if len(stripped) < 80:
        return False
    if _looks_like_toc_entry(stripped):
        return False
    if _looks_like_caption(stripped, style_name=unit.style_name) or _looks_like_note(stripped) or _looks_like_formula_paragraph(stripped, style_name=unit.style_name):
        return False
    return _has_sentence_punctuation(stripped)


def _find_scope_diagnostic_issues(
    units: list[DocxTextUnit],
    *,
    start_index: int | None,
    end_index: int | None,
    acknowledgement_index: int | None,
    boundary_index: int | None,
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    if not units:
        _add_scope_issue(
            issues,
            code="empty_docx_text_units",
            severity="error",
            message="No non-empty text units were extracted from the DOCX.",
        )
        return issues
    if start_index is None:
        _add_scope_issue(
            issues,
            code="missing_rewrite_scope_start",
            severity="error",
            message="No abstract or body-start marker was found; the document would have no editable body scope.",
        )
        return issues
    if end_index is None or end_index < start_index:
        _add_scope_issue(
            issues,
            code="invalid_rewrite_scope_end",
            severity="error",
            message="The rewrite scope end is missing or appears before the start marker.",
        )
    if not any(unit.editable for unit in units):
        _add_scope_issue(
            issues,
            code="empty_editable_scope",
            severity="error",
            message="The protection rules produced no editable body paragraphs.",
        )
    if acknowledgement_index is None:
        _add_scope_issue(
            issues,
            code="missing_acknowledgement_marker",
            severity="warning",
            message="No acknowledgement heading was found; the body scope ends at the first back-matter marker or the document end.",
        )
    elif boundary_index is None:
        risky_tail_units = _find_unbounded_post_acknowledgement_tail_units(units, acknowledgement_index)
        if risky_tail_units:
            _add_scope_issue(
                issues,
                code="unbounded_post_acknowledgement_tail",
                severity="error",
                message="Acknowledgements were found, but later text looks like references, appendix, declaration, or other protected back matter.",
                unit=risky_tail_units[0],
            )

    for unit in units:
        index = unit.unit_index
        evidence_digest = get_docx_unit_edit_eligibility_evidence_digest(unit)
        if unit.editable and (
            unit.structural_role not in DOCX_EDITABLE_STRUCTURAL_ROLES
            or unit.edit_eligibility != "eligible"
            or not evidence_digest
        ):
            _add_scope_issue(
                issues,
                code="illegal_editable_structural_role",
                severity="error",
                message="An editable unit lacks an allowed prose role or complete eligibility evidence.",
                unit=unit,
            )
        elif unit.structural_role == "ambiguous_non_prose":
            _add_scope_issue(
                issues,
                code="ambiguous_non_prose_protected",
                severity="warning",
                message="A paragraph without affirmative prose evidence was protected from model input.",
                unit=unit,
            )
        if unit.editable and (index < start_index or (end_index is not None and index > end_index)):
            _add_scope_issue(
                issues,
                code="editable_outside_rewrite_scope",
                severity="error",
                message="A unit outside the computed body scope is editable.",
                unit=unit,
            )
        if not unit.editable and start_index <= index and (end_index is None or index <= end_index):
            reason = str(unit.protect_reason or "")
            if reason == "heading" and _looks_like_long_body_text(unit):
                _add_scope_issue(
                    issues,
                    code="body_like_text_protected_as_heading",
                    severity="warning",
                    message="A long sentence-like unit inside the body scope was protected as a heading.",
                    unit=unit,
                )
        if unit.editable and (
            _unit_looks_like_references_heading(unit)
            or _looks_like_back_matter_heading(unit.text)
            or _looks_like_keyword_line(unit.text)
            or _looks_like_caption(unit.text, style_name=unit.style_name)
            or _looks_like_formula_paragraph(unit.text, style_name=unit.style_name)
        ):
            _add_scope_issue(
                issues,
                code="structural_unit_marked_editable",
                severity="warning",
                message="A structural-looking unit is editable and should be reviewed.",
                unit=unit,
            )
    return issues


def _find_rewrite_scope_start_index(units: list[DocxTextUnit]) -> int | None:
    abstract_index = next(
        (index for index, unit in enumerate(units) if _looks_like_abstract_start(unit.text)),
        None,
    )
    if abstract_index is not None:
        return abstract_index
    return next(
        (index for index, unit in enumerate(units) if _looks_like_body_start(unit.text, style_name=unit.style_name)),
        None,
    )


def _find_rewrite_scope_boundary_index(
    units: list[DocxTextUnit],
    *,
    start_index: int,
    acknowledgement_index: int | None,
) -> int | None:
    """Find the first trusted back-matter boundary after the active body tail."""

    search_from = acknowledgement_index if acknowledgement_index is not None else start_index
    reference_entry_run_starts = _find_reference_entry_run_starts(
        units,
        start_index=search_from + 1,
    )
    for index, unit in enumerate(units[search_from + 1:], start=search_from + 1):
        if index in reference_entry_run_starts:
            return index
        if acknowledgement_index is not None:
            if _looks_like_post_acknowledgement_boundary(unit):
                return index
            continue
        if _unit_looks_like_references_heading(unit) or _looks_like_back_matter_heading(unit.text):
            return index
    return None


def _find_rewrite_scope_end_index(units: list[DocxTextUnit], start_index: int) -> int:
    acknowledgement_index = next(
        (index for index, unit in enumerate(units[start_index:], start=start_index) if _looks_like_acknowledgement_heading(unit.text)),
        None,
    )
    back_matter_index = _find_rewrite_scope_boundary_index(
        units,
        start_index=start_index,
        acknowledgement_index=acknowledgement_index,
    )
    if back_matter_index is not None:
        return max(start_index, back_matter_index - 1)
    return len(units) - 1


def _looks_like_post_acknowledgement_boundary(unit: DocxTextUnit) -> bool:
    if _unit_looks_like_references_heading(unit):
        return True
    if _looks_like_back_matter_heading(unit.text):
        return True
    normalized_style = _normalize_style_name(unit.style_name)
    if normalized_style.startswith("heading") or normalized_style.startswith("标题"):
        return _looks_like_heading(
            unit.text,
            style_name=unit.style_name,
            outline_level=unit.outline_level,
        )
    return False


def _find_unbounded_post_acknowledgement_tail_units(
    units: list[DocxTextUnit],
    acknowledgement_index: int,
) -> list[DocxTextUnit]:
    risky_units: list[DocxTextUnit] = []
    for unit in units[acknowledgement_index + 1:]:
        text = (unit.text or "").strip()
        normalized = _normalize_marker_text(text)
        if not normalized:
            continue
        if _unit_looks_like_references_heading(unit) or _looks_like_back_matter_heading(text):
            risky_units.append(unit)
        elif _looks_like_reference_entry(text):
            risky_units.append(unit)
        elif (
            len(normalized) <= 48
            and not any(mark in text for mark in "。！？；!?;")
            and POST_ACKNOWLEDGEMENT_TAIL_MARKER_RE.search(normalized)
        ):
            risky_units.append(unit)
        if len(risky_units) >= 5:
            break
    return risky_units


def _looks_like_abstract_start(text: str) -> bool:
    stripped = (text or "").strip()
    if _looks_like_toc_entry(stripped):
        return False
    normalized = _normalize_marker_text(text)
    return (
        normalized in ABSTRACT_MARKERS
        or (normalized.startswith("摘要") and len(normalized) <= 12)
        or (normalized.endswith("摘要") and len(normalized) <= 12)
    )


def _looks_like_toc_entry(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    # Dot leaders alone are not sufficient: ordinary thesis prose and template
    # placeholders can legitimately contain repeated dots/ellipses.  A trusted
    # visible TOC row needs a trailing Arabic/Roman page token after the leader.
    if re.search(
        r"(?:\.{3,}|…{2,}|·{3,})\s*(?:[ivxlcdmIVXLCDM]+|\d+)\s*$",
        stripped,
    ):
        return True
    return bool(re.search(r"\s+(?:[ivxlcdmIVXLCDM]+|\d+)\s*$", stripped)) and len(stripped) <= 80


def _looks_like_body_start(text: str, *, style_name: str = "") -> bool:
    if _looks_like_toc_entry(text):
        return False
    normalized = _normalize_marker_text(text)
    if normalized in ABSTRACT_MARKERS:
        return True
    if DATE_LINE_RE.match(normalized):
        return False
    if normalized in {"绪论", "引言", "前言"}:
        return True
    if normalized.startswith("第") and "章" in normalized:
        return True
    style_prefix = style_name.strip().lower()
    if style_prefix.startswith("heading") and BODY_START_RE.match(normalized):
        return True
    if BODY_START_RE.match(normalized):
        return True
    return False


def _looks_like_toc_heading(text: str) -> bool:
    return _normalize_marker_text(text) in TOC_MARKERS


def _looks_like_acknowledgement_heading(text: str) -> bool:
    # A generated table-of-contents entry such as ``致 谢 22`` carries the
    # same visible marker as the real section heading.  Treating it as the
    # acknowledgement boundary truncates the editable scope before the thesis
    # body starts.  The page-number/leader shape is already recognized by the
    # shared TOC predicate, so reject it before normalizing whitespace.
    if _looks_like_toc_entry(text):
        return False
    normalized = _normalize_marker_text(text)
    return normalized in ACKNOWLEDGEMENT_MARKERS or (normalized.startswith(("致谢", "谢辞", "鸣谢")) and len(normalized) <= 12)


def _looks_like_references_heading(text: str) -> bool:
    # Do not let ``参考文献 23`` (or its dotted-leader equivalent) in a TOC act
    # as a trusted back-matter boundary.  Only the real section heading may
    # close the editable body range.
    if _looks_like_toc_entry(text):
        return False
    normalized = _normalize_marker_text(text)
    return normalized in REFERENCE_MARKERS or ("参考文献" in normalized and len(normalized) <= 24)


def _has_reference_heading_semantics(text: str) -> bool:
    """Recognize reference-section semantics without broadening all headings."""

    stripped = (text or "").strip()
    normalized = _normalize_marker_text(stripped)
    if not normalized or len(normalized) > 48:
        return False

    if "参考" in normalized and any(term in normalized for term in ("文献", "资料", "书目", "来源")):
        return True
    if any(term in normalized for term in ("引用文献", "文献目录", "文献列表", "资料来源")):
        return True

    english_words = set(re.findall(r"[a-z]+", stripped.casefold()))
    if english_words.intersection({"bibliography", "bibliographies", "bibliographic"}):
        return True
    if english_words.intersection({"reference", "references"}) and english_words.intersection(
        {"list", "cited", "consulted", "selected", "materials", "sources"}
    ):
        return True
    if "sources" in english_words and english_words.intersection({"cited", "consulted", "reference", "references"}):
        return True
    if "works" in english_words and english_words.intersection({"cited", "consulted", "referenced"}):
        return True
    return False


def _unit_looks_like_references_heading(unit: DocxTextUnit) -> bool:
    # The broader semantic fallback below intentionally recognizes customized
    # reference headings.  It must still respect the stronger TOC-entry signal
    # or a generated ``参考文献 23`` row will switch the whole following body
    # into the references phase.
    if _looks_like_toc_entry(unit.text):
        return False
    if _looks_like_references_heading(unit.text):
        return True
    if not _has_reference_heading_semantics(unit.text):
        return False
    return _looks_like_heading(
        unit.text,
        style_name=unit.style_name,
        has_numbering=unit.has_numbering,
        outline_level=unit.outline_level,
    )


def _looks_like_reference_entry(text: str) -> bool:
    stripped = (text or "").strip()
    if len(stripped) < 8:
        return False
    if REFERENCE_BRACKET_ENTRY_RE.match(stripped):
        return True
    if stripped.casefold().startswith(("doi:", "doi：", "http://", "https://")):
        return True
    if REFERENCE_AUTHOR_YEAR_PAREN_RE.match(stripped):
        return True

    has_year = bool(REFERENCE_YEAR_RE.search(stripped[:220]))
    has_bibliographic_cue = bool(REFERENCE_BIBLIOGRAPHIC_CUE_RE.search(stripped))
    if REFERENCE_NUMBERED_ENTRY_RE.match(stripped):
        return has_year or has_bibliographic_cue
    if has_year and (
        REFERENCE_EN_AUTHOR_LEAD_RE.match(stripped)
        or REFERENCE_CN_AUTHOR_LEAD_RE.match(stripped)
    ):
        return True
    return False


def _find_reference_entry_run_starts(
    units: Sequence[DocxTextUnit],
    *,
    start_index: int = 0,
    minimum_run_length: int = 2,
) -> set[int]:
    """Return starts of consecutive reference-like runs used as fail-safe gates."""

    starts: set[int] = set()
    index = max(0, int(start_index))
    while index < len(units):
        if not _looks_like_reference_entry(units[index].text):
            index += 1
            continue
        run_start = index
        while index < len(units) and _looks_like_reference_entry(units[index].text):
            index += 1
        if index - run_start >= minimum_run_length:
            starts.add(run_start)
    return starts


def _looks_like_back_matter_heading(text: str) -> bool:
    if _looks_like_toc_entry(text):
        return False
    normalized = _normalize_marker_text(text)
    return (
        normalized in BACK_MATTER_MARKERS
        or normalized.startswith("附录")
        or normalized.startswith("appendix")
        or normalized.endswith("封底格式")
    )


def _looks_like_heading(
    text: str,
    *,
    style_name: str = "",
    has_numbering: bool = False,
    outline_level: int | None = None,
) -> bool:
    normalized = _normalize_marker_text(text)
    if not normalized or DATE_LINE_RE.match(normalized):
        return False
    if (
        normalized in ABSTRACT_MARKERS
        or normalized in ACKNOWLEDGEMENT_MARKERS
        or _looks_like_references_heading(text)
        or _looks_like_back_matter_heading(text)
        or normalized in {"绪论", "引言", "前言", "结论", "结语", "总结", "总结与展望", "结论与展望"}
    ):
        return True

    if _looks_like_numbered_structure_label(text):
        return True

    style_prefix = _normalize_style_name(style_name)
    if outline_level is not None and 0 <= outline_level <= 8:
        return True
    if (
        style_prefix.startswith("heading")
        or style_prefix.startswith("标题")
        or any(hint in style_prefix for hint in ("chapterheading", "sectionheading", "title", "章标题", "节标题"))
    ):
        return True

    if has_numbering:
        # Word's automatic numbering is not part of ``paragraph.text``.  A
        # Normal/custom-style ``w:numPr`` paragraph therefore reaches this
        # branch as bare text such as “系统设计”.  Short standalone noun
        # phrases are ambiguous but format-sensitive structural labels, so
        # fail closed; punctuated/long list prose remains editable.
        return _looks_like_short_structure_label_body(text)

    if _looks_like_numbered_body_item(text):
        return False

    if len(normalized) <= 48 and bool(HEADING_NUMBER_RE.match(normalized)):
        return True

    return _looks_like_short_unnumbered_heading(text)


def _looks_like_numbered_body_item(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False

    marker_match = NUMBERED_BODY_MARKER_RE.match(stripped)
    marker_kind = "punctuated"
    if marker_match is None:
        marker_match = NUMBERED_SPACE_BODY_MARKER_RE.match(stripped)
        marker_kind = "space"
    if marker_match is None:
        return False

    body = marker_match.group("body").strip()
    if not body:
        return False
    if re.match(r"^\d", body):
        return False

    normalized_body = _normalize_marker_text(body)
    # A visible numbered marker followed by a noun phrase but no sentence/list
    # punctuation is a structural label, even when the template incorrectly
    # uses Normal style. Body list items remain editable when they carry real
    # sentence punctuation. Ambiguous short labels fail closed as structure.
    if marker_kind == "punctuated" and _looks_like_numbered_structure_label(stripped):
        return False
    if len(normalized_body) <= 8 and not any(mark in body for mark in ":：，,；;。"):
        if marker_kind == "punctuated" and not _looks_like_short_numbered_heading_body(normalized_body):
            return True
        return False
    if any(mark in body for mark in ":：，,；;。"):
        return True
    if len(normalized_body) >= 16:
        return True
    return False


def _looks_like_numbered_structure_label(text: str) -> bool:
    """Recognize a visible numbered subheading independently of Word style."""

    marker_match = NUMBERED_BODY_MARKER_RE.match((text or "").strip())
    if marker_match is None:
        return False
    return _looks_like_short_structure_label_body(marker_match.group("body"))


def _looks_like_short_structure_label_body(text: str) -> bool:
    """Fail closed for a short standalone label, not real numbered prose."""

    body = (text or "").strip()
    normalized_body = _normalize_marker_text(body)
    if not normalized_body or len(normalized_body) > 64:
        return False
    return not any(mark in body for mark in ":：，,；;。！？!?")


def _looks_like_short_numbered_heading_body(normalized_body: str) -> bool:
    if normalized_body in SHORT_NUMBERED_HEADING_BODIES:
        return True
    return any(keyword == normalized_body or keyword in normalized_body for keyword in HEADING_KEYWORDS)


def _looks_like_short_unnumbered_heading(text: str) -> bool:
    stripped = (text or "").strip()
    normalized = _normalize_marker_text(stripped)
    if not normalized or len(normalized) > 36:
        return False
    if any(mark in stripped for mark in "。！？；!?;"):
        return False
    if any(keyword in normalized for keyword in HEADING_KEYWORDS):
        return True
    if len(normalized) <= 14 and SHORT_HEADING_RE.match(stripped):
        return True
    return False


def _looks_like_caption(text: str, *, style_name: str = "") -> bool:
    normalized_style = _normalize_style_name(style_name)
    if any(hint in normalized_style for hint in ("caption", "图注", "表注", "题注")):
        return True
    return bool(CAPTION_RE.match(_normalize_marker_text(text)))


def _looks_like_note(text: str) -> bool:
    stripped = (text or "").strip()
    normalized = _normalize_marker_text(stripped)
    if normalized.startswith(("图注", "表注")):
        return True
    return bool(re.match(r"^(?:注|说明|备注)\s*[:：]", stripped, flags=re.IGNORECASE))


def _looks_like_keyword_line(text: str) -> bool:
    normalized = _normalize_marker_text(text)
    return any(normalized.startswith(prefix) for prefix in KEYWORD_PREFIXES)


def _looks_like_formula_paragraph(text: str, *, style_name: str = "") -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False

    normalized_style = _normalize_style_name(style_name)
    if any(hint in normalized_style for hint in FORMULA_STYLE_HINTS):
        return True

    if len(stripped) > 220:
        return False

    formula_symbol_count = len(FORMULA_SYMBOL_RE.findall(stripped))
    digit_count = len(DIGIT_CHAR_RE.findall(stripped))
    latin_count = len(LATIN_CHAR_RE.findall(stripped))
    has_formula_operator = any(symbol in stripped for symbol in ("=", "≈", "≤", "≥", "∑", "Σ", "√", "×", "÷"))
    has_formula_hint = any(keyword in stripped for keyword in ("计算方式", "计算公式", "公式如下", "定义如下"))

    if has_formula_operator and formula_symbol_count >= 2:
        return True
    if (
        "=" in stripped
        and latin_count >= 2
        and len(stripped) <= 100
        and not any(mark in stripped for mark in "。！？；!?;")
    ):
        return True
    if has_formula_hint and formula_symbol_count >= 1:
        return True
    if formula_symbol_count >= 4 and digit_count >= 1 and latin_count >= 1:
        return True
    return False


def _normalize_marker_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "").casefold()


def _normalize_style_name(style_name: str) -> str:
    return re.sub(r"\s+", "", style_name or "").casefold()


def _is_table_unit(unit: DocxTextUnit) -> bool:
    return str(unit.target.get("kind", "")) == "table_cell_paragraph"


def _paragraph_has_field_code(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:fldChar | .//w:instrText | .//w:fldSimple"))


def _paragraph_has_drawing(paragraph: Paragraph) -> bool:
    # Cover modern DrawingML as well as legacy VML/OLE containers.  Rewriting
    # prose in the same paragraph could otherwise move an anchored object or
    # detach its caption semantics.
    return bool(paragraph._element.xpath(".//w:drawing | .//w:pict | .//w:object"))


def _paragraph_has_math(paragraph: Paragraph) -> bool:
    # OMML uses the officeDocument math namespace.  local-name keeps this
    # robust across namespace-prefix choices made by different Word versions.
    return bool(
        paragraph._element.xpath(
            ".//*[local-name()='oMath' or local-name()='oMathPara']"
        )
    )


def _paragraph_has_semantic_range_anchor(paragraph: Paragraph) -> bool:
    """Return True for zero-width range markers bound to exact prose.

    Bookmark and comment-range nodes are preserved as XML siblings while the
    rewrite mapper redistributes text across direct runs.  Keeping the nodes is
    therefore insufficient: their enclosed text can change even when paragraph
    text and non-text OOXML structure appear stable.
    """

    return bool(
        _paragraph_has_bookmark_range_anchor(paragraph)
        or _paragraph_has_comment_range_anchor(paragraph)
    )


def _paragraph_has_bookmark_range_anchor(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:bookmarkStart | .//w:bookmarkEnd"))


def _paragraph_has_comment_range_anchor(paragraph: Paragraph) -> bool:
    return bool(
        paragraph._element.xpath(".//w:commentRangeStart | .//w:commentRangeEnd")
    )


def _paragraph_has_semantic_point_reference(paragraph: Paragraph) -> bool:
    """Return True for comment/note references attached at one text offset."""

    return bool(
        paragraph._element.xpath(
            ".//w:commentReference | .//w:footnoteReference | .//w:endnoteReference"
        )
    )


def _paragraph_has_complex_inline(paragraph: Paragraph) -> bool:
    """Return True when visible text sits outside supported run containers.

    Rewriting is deliberately limited to direct ``w:t`` text in direct
    ``w:r`` and direct hyperlink runs. Inline content controls, tracked
    revisions, smart tags, ruby text, bidi wrappers, AlternateContent, tabs,
    and manual line breaks must be protected wholesale so the model never
    rewrites only the visible fragments that python-docx exposes.  In
    particular, ``Paragraph.text`` serializes ``w:tab``/``w:br`` as literal
    characters while the text-run replacer preserves their OOXML nodes; if
    such a paragraph were editable, a rewrite could duplicate or relocate the
    structural character during export.
    """

    if (
        _paragraph_has_semantic_range_anchor(paragraph)
        or _paragraph_has_semantic_point_reference(paragraph)
    ):
        return True

    complex_container = paragraph._element.xpath(
        ".//*["
        "local-name()='sdt' or local-name()='customXml' or "
        "local-name()='smartTag' or local-name()='ins' or local-name()='del' or "
        "local-name()='moveFrom' or local-name()='moveTo' or local-name()='ruby' or "
        "local-name()='dir' or local-name()='bdo' or local-name()='AlternateContent'"
        "]"
    )
    if complex_container:
        return True

    if paragraph._element.xpath(".//w:tab | .//w:br | .//w:cr"):
        return True

    paragraph_element = paragraph._p
    for text_node in paragraph_element.xpath(".//w:t"):
        run = text_node.getparent()
        if run is None or run.tag != qn("w:r"):
            return True
        container = run.getparent()
        if container is paragraph_element:
            continue
        if (
            container is not None
            and container.tag == qn("w:hyperlink")
            and container.getparent() is paragraph_element
        ):
            continue
        return True
    return False


def _paragraph_numbering_level(paragraph: Paragraph) -> int | None:
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        ilvl = p_pr.numPr.ilvl
        if ilvl is not None and ilvl.val is not None:
            return _as_optional_int(ilvl.val) or 0
        return 0

    style = paragraph.style
    style_name = style.name if style is not None else ""
    if _looks_like_numbered_paragraph_style(style_name):
        return 0
    try:
        if style is not None and style._element.xpath(".//w:numPr"):
            return 0
    except Exception:
        return None
    return None


def _paragraph_outline_level(paragraph: Paragraph) -> int | None:
    candidates: list[Any] = []
    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is not None:
        candidates.extend(paragraph_properties.xpath("./w:outlineLvl"))

    style = paragraph.style
    visited: set[str] = set()
    while style is not None and _style_chain_identity(style) not in visited:
        visited.add(_style_chain_identity(style))
        try:
            candidates.extend(style._element.xpath(".//w:pPr/w:outlineLvl"))
        except Exception:
            pass
        style = style.base_style

    for candidate in candidates:
        raw_value = candidate.get(qn("w:val"))
        value = _as_optional_int(raw_value)
        if value is not None and 0 <= value <= 8:
            return value
    return None


def _looks_like_numbered_paragraph_style(style_name: str) -> bool:
    normalized_style = _normalize_style_name(style_name)
    if normalized_style.startswith(("heading", "标题")):
        return False
    return any(
        hint in normalized_style
        for hint in ("listnumber", "numbered", "列表编号", "编号列表", "项目编号", "自动编号")
    )


def _as_optional_int(value: Any) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:  # pragma: no cover - internal contract
        raise TypeError(f"Unsupported parent type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _resolve_target_paragraph(document: DocxDocument, target: dict[str, Any]) -> Paragraph:
    target_kind = str(target.get("kind", ""))
    if target_kind == "paragraph":
        paragraph_index = int(target["paragraph_index"])
        return document.paragraphs[paragraph_index]

    if target_kind == "table_cell_paragraph":
        raw_table_path = target.get("table_path")
        if isinstance(raw_table_path, list) and raw_table_path:
            first_step = raw_table_path[0]
            if not isinstance(first_step, dict):
                raise ValueError("Invalid DOCX table target path root.")
            table = document.tables[int(first_step["table_index"])]
            for raw_step in raw_table_path[1:]:
                if not isinstance(raw_step, dict):
                    raise ValueError("Invalid DOCX nested table target path step.")
                parent_cell = table.rows[int(raw_step["row_index"])].cells[int(raw_step["cell_index"])]
                table = parent_cell.tables[int(raw_step["table_index"])]
        else:
            # Read-only compatibility for structural inventory v1. Such
            # snapshots fail ``_is_snapshot_current`` and are re-derived before
            # any production model/export authority use.
            table = document.tables[int(target["table_index"])]
        cell = table.rows[int(target["row_index"])].cells[int(target["cell_index"])]
        return cell.paragraphs[int(target["paragraph_index"])]

    raise ValueError(f"Unsupported DOCX target kind: {target_kind}")


def _snap_split_index(text: str, raw_index: int, lower_bound: int, upper_bound: int) -> int:
    if raw_index <= lower_bound:
        return lower_bound
    if raw_index >= upper_bound:
        return upper_bound

    preferred_boundaries = set(" \t\n,.;:!?)]}，。；：！？）】、")
    search_start = max(lower_bound, raw_index - 12)
    search_end = min(upper_bound, raw_index + 12)
    best_index = raw_index
    best_distance = None

    for index in range(search_start, search_end + 1):
        left_char = text[index - 1] if index > 0 else ""
        right_char = text[index] if index < len(text) else ""
        if left_char not in preferred_boundaries and right_char not in preferred_boundaries:
            continue
        distance = abs(index - raw_index)
        if best_distance is None or distance < best_distance:
            best_distance = distance
            best_index = index
            if distance == 0:
                break

    return max(lower_bound, min(best_index, upper_bound))


def _split_text_for_runs_by_weight(text: str, run_texts: list[str]) -> list[str]:
    if not run_texts:
        return []
    if len(run_texts) == 1:
        return [text]

    total_weight = sum(max(len(run_text), 1) for run_text in run_texts)
    splits: list[int] = []
    previous_index = 0
    consumed_weight = 0

    for run_text in run_texts[:-1]:
        consumed_weight += max(len(run_text), 1)
        raw_index = round(len(text) * consumed_weight / total_weight)
        split_index = _snap_split_index(text, raw_index, previous_index, len(text))
        splits.append(split_index)
        previous_index = split_index

    parts: list[str] = []
    start = 0
    for split_index in splits:
        parts.append(text[start:split_index])
        start = split_index
    parts.append(text[start:])
    return parts


def _map_original_boundary_to_rewrite(
    boundary: int,
    *,
    original_length: int,
    rewritten_length: int,
    opcodes: Sequence[tuple[str, int, int, int, int]],
) -> int:
    if boundary <= 0:
        return 0
    if boundary >= original_length:
        return rewritten_length

    # Prefer an unchanged block when a run boundary sits exactly at an edit
    # edge. This keeps an unchanged styled token (for example a bold model
    # name or a hyperlink label) inside the same run after nearby prose moves.
    for tag, source_start, source_end, rewrite_start, _rewrite_end in opcodes:
        if tag == "equal" and source_start <= boundary <= source_end:
            return rewrite_start + (boundary - source_start)

    for _tag, source_start, source_end, rewrite_start, rewrite_end in opcodes:
        if source_start <= boundary <= source_end:
            source_span = source_end - source_start
            if source_span <= 0:
                return rewrite_end
            ratio = (boundary - source_start) / source_span
            return round(rewrite_start + ratio * (rewrite_end - rewrite_start))
    return round(rewritten_length * boundary / max(original_length, 1))


def _split_text_for_runs(text: str, run_texts: list[str]) -> list[str]:
    if not run_texts:
        return []
    if len(run_texts) == 1:
        return [text]

    original_text = "".join(run_texts)
    if not original_text:
        return _split_text_for_runs_by_weight(text, run_texts)

    matcher = SequenceMatcher(None, original_text, text, autojunk=False)
    matching_chars = sum(block.size for block in matcher.get_matching_blocks())
    if matching_chars < max(4, round(len(original_text) * 0.15)):
        return _split_text_for_runs_by_weight(text, run_texts)

    opcodes = matcher.get_opcodes()
    source_boundaries: list[int] = []
    consumed = 0
    for run_text in run_texts[:-1]:
        consumed += len(run_text)
        source_boundaries.append(consumed)

    rewrite_boundaries: list[int] = []
    previous = 0
    for boundary in source_boundaries:
        mapped = _map_original_boundary_to_rewrite(
            boundary,
            original_length=len(original_text),
            rewritten_length=len(text),
            opcodes=opcodes,
        )
        mapped = max(previous, min(mapped, len(text)))
        rewrite_boundaries.append(mapped)
        previous = mapped

    parts: list[str] = []
    start = 0
    for boundary in rewrite_boundaries:
        parts.append(text[start:boundary])
        start = boundary
    parts.append(text[start:])
    return parts


def _run_text_nodes(run: Run) -> list[Any]:
    return list(run._r.findall(qn("w:t")))


def _iter_paragraph_text_runs(paragraph: Paragraph) -> list[tuple[Run, str]]:
    """Return text-bearing run slots in exact visible OOXML order.

    Empty ``w:t`` slots are retained because they can safely receive rewritten
    prose without creating/removing runs.  Structural runs without a ``w:t``
    (fields, drawings, tabs, breaks) are deliberately excluded and are already
    protected by the paragraph-level complex-inline gates.
    """

    text_runs: list[tuple[Run, str]] = []
    for item in paragraph.iter_inner_content():
        candidates = [item] if isinstance(item, Run) else list(getattr(item, "runs", ()))
        for run in candidates:
            text_nodes = _run_text_nodes(run)
            run_text = "".join(str(node.text or "") for node in text_nodes)
            if text_nodes:
                text_runs.append((run, run_text))
    return text_runs


def _iter_run_style_chain(run: Run) -> Iterator[Any]:
    try:
        style = run.style
    except Exception:
        style = None
    visited: set[str] = set()
    while style is not None and _style_chain_identity(style) not in visited:
        visited.add(_style_chain_identity(style))
        yield style
        style = getattr(style, "base_style", None)


def _effective_run_font_property(run: Run, property_name: str) -> Any:
    value = getattr(run.font, property_name, None)
    if value is not None:
        return value
    for style in _iter_run_style_chain(run):
        value = getattr(getattr(style, "font", None), property_name, None)
        if value is not None:
            return value
    return None


def _run_colour_signature(run: Run) -> str:
    colour_candidates = [run.font.color]
    colour_candidates.extend(
        getattr(getattr(style, "font", None), "color", None)
        for style in _iter_run_style_chain(run)
    )
    for colour in colour_candidates:
        if colour is None or getattr(colour, "type", None) is None:
            continue
        values: list[str] = [str(colour.type)]
        for attribute in ("rgb", "theme_color", "brightness"):
            try:
                value = getattr(colour, attribute, None)
            except (AttributeError, ValueError):
                value = None
            if value is not None:
                values.append(f"{attribute}={value}")
        return "|".join(values)
    return ""


def _nondefault_character_style_name(run: Run) -> str:
    try:
        style_name = str(run.style.name or "").strip()
    except Exception:
        return ""
    normalized = _normalize_style_name(style_name)
    if normalized in {
        "",
        "defaultparagraphfont",
        "默认段落字体",
        "默认字符字体",
        "无样式",
    }:
        return ""
    return style_name


def _iter_run_rpr_sources(run: Run) -> Iterator[Any]:
    """Yield direct/character-style rPr nodes in effective precedence order."""

    seen: set[int] = set()
    direct = getattr(run._r, "rPr", None)
    if direct is not None:
        seen.add(id(direct))
        yield direct
    for style in _iter_run_style_chain(run):
        style_element = getattr(style, "_element", None)
        r_pr = getattr(style_element, "rPr", None)
        if r_pr is None and style_element is not None:
            r_pr = style_element.find(qn("w:rPr"))
        if r_pr is not None and id(r_pr) not in seen:
            seen.add(id(r_pr))
            yield r_pr


def _effective_rpr_attribute_values(
    run: Run,
    element_name: str,
    attribute_names: Sequence[str],
) -> tuple[tuple[str, str], ...]:
    """Resolve inheritable rPr attributes without serializing unrelated XML."""

    resolved: dict[str, str] = {}
    for r_pr in _iter_run_rpr_sources(run):
        element = r_pr.find(qn(f"w:{element_name}"))
        if element is None:
            continue
        for attribute_name in attribute_names:
            if attribute_name in resolved:
                continue
            raw_value = element.get(qn(f"w:{attribute_name}"))
            if raw_value is not None:
                resolved[attribute_name] = str(raw_value).strip()
    return tuple((name, resolved[name]) for name in attribute_names if name in resolved)


def _effective_rpr_element(run: Run, element_name: str) -> Any | None:
    for r_pr in _iter_run_rpr_sources(run):
        element = r_pr.find(qn(f"w:{element_name}"))
        if element is not None:
            return element
    return None


def _normalize_ooxml_on_off(value: Any) -> str:
    normalized = str(value or "").strip().casefold()
    return "false" if normalized in {"0", "false", "off", "no"} else "true"


def _effective_rpr_on_off(run: Run, element_name: str) -> str | None:
    element = _effective_rpr_element(run, element_name)
    if element is None:
        return None
    return _normalize_ooxml_on_off(element.get(qn("w:val")))


def _canonical_rpr_element_signature(run: Run, element_name: str) -> str:
    element = _effective_rpr_element(run, element_name)
    if element is None:
        return ""
    attributes = sorted(
        (
            str(attribute_name).rsplit("}", 1)[-1],
            str(attribute_value).strip(),
        )
        for attribute_name, attribute_value in element.attrib.items()
    )
    if not attributes:
        return "present"
    return "|".join(f"{name}={value}" for name, value in attributes)


def _format_sensitive_run_descriptor(run: Run) -> tuple[tuple[tuple[str, str], ...], tuple[str, ...]] | None:
    """Describe inline formatting whose *text attachment* must stay frozen."""

    values: list[tuple[str, str]] = []
    kinds: list[str] = []

    def add(kind: str, value: Any = True) -> None:
        kinds.append(kind)
        values.append((kind, str(value)))

    font_family = _effective_rpr_attribute_values(
        run,
        "rFonts",
        (
            "ascii",
            "hAnsi",
            "eastAsia",
            "cs",
            "asciiTheme",
            "hAnsiTheme",
            "eastAsiaTheme",
            "cstheme",
            "hint",
        ),
    )
    if font_family:
        add("font_family", "|".join(f"{name}={value}" for name, value in font_family))

    font_size = (
        *_effective_rpr_attribute_values(run, "sz", ("val",)),
        *((f"cs_{name}", value) for name, value in _effective_rpr_attribute_values(run, "szCs", ("val",))),
    )
    if font_size:
        add("font_size", "|".join(f"{name}={value}" for name, value in font_size))

    scalar_properties = (
        ("spacing", "character_spacing"),
        ("position", "character_position"),
        ("kern", "kerning"),
    )
    for element_name, kind in scalar_properties:
        attributes = _effective_rpr_attribute_values(run, element_name, ("val",))
        if attributes:
            add(kind, attributes[0][1])

    fit_text = _effective_rpr_attribute_values(run, "fitText", ("val", "id"))
    if fit_text:
        add("fit_text", "|".join(f"{name}={value}" for name, value in fit_text))

    for kind, element_names in (
        ("bold", ("b", "bCs")),
        ("italic", ("i", "iCs")),
    ):
        effective_values = [
            (element_name, value)
            for element_name in element_names
            if (value := _effective_rpr_on_off(run, element_name)) is not None
        ]
        if effective_values:
            add(kind, "|".join(f"{name}={value}" for name, value in effective_values))

    underline = _canonical_rpr_element_signature(run, "u")
    if underline:
        # Keep explicit w:val="none": it can locally cancel inherited
        # underlining and therefore still defines a meaningful range boundary.
        add("underline", underline)

    vertical_alignment = _effective_rpr_attribute_values(run, "vertAlign", ("val",))
    if vertical_alignment:
        alignment_value = vertical_alignment[0][1]
        if alignment_value == "superscript":
            add("superscript", alignment_value)
        elif alignment_value == "subscript":
            add("subscript", alignment_value)
        else:
            # w:vertAlign="baseline" is an explicit cancellation of an
            # inherited superscript/subscript and must not disappear.
            add("vertical_align", alignment_value)

    on_off_properties = (
        ("strike", "strike"),
        ("dstrike", "double_strike"),
        ("caps", "caps"),
        ("smallCaps", "small_caps"),
        ("outline", "outline"),
        ("shadow", "shadow"),
        ("emboss", "emboss"),
        ("imprint", "imprint"),
        ("vanish", "hidden"),
        ("webHidden", "web_hidden"),
    )
    for element_name, kind in on_off_properties:
        value = _effective_rpr_on_off(run, element_name)
        if value is not None:
            add(kind, value)

    text_border = _canonical_rpr_element_signature(run, "bdr")
    if text_border:
        add("text_border", text_border)
    text_shading = _canonical_rpr_element_signature(run, "shd")
    if text_shading:
        add("text_shading", text_shading)

    colour = _run_colour_signature(run)
    if colour:
        add("color", colour)
    highlight = _effective_run_font_property(run, "highlight_color")
    if highlight is not None:
        add("highlight", highlight)

    character_style = _nondefault_character_style_name(run)
    if character_style:
        add("character_style", character_style)

    parent = run._r.getparent()
    if parent is not None and parent.tag == qn("w:hyperlink"):
        relationship_id = str(parent.get(qn("r:id")) or "")
        anchor = str(parent.get(qn("w:anchor")) or "")
        add("hyperlink", f"{relationship_id}|{anchor}")

    if not values:
        return None
    return tuple(values), tuple(sorted(set(kinds)))


def _text_occurrence_positions(text: str, needle: str) -> list[int]:
    if not needle:
        return []
    positions: list[int] = []
    cursor = 0
    while cursor <= len(text) - len(needle):
        position = text.find(needle, cursor)
        if position < 0:
            break
        positions.append(position)
        cursor = position + 1
    return positions


def _extract_format_sensitive_anchors(paragraph: Paragraph) -> list[dict[str, Any]]:
    """Capture styled inline text that must remain attached to its source runs.

    Uniform paragraph-wide emphasis does not create an anchor: rewriting a
    single all-bold run remains all-bold and cannot move an emphasis boundary.
    Local emphasis, hyperlinks, and vertical-script text do create anchors.
    Consecutive runs are merged only when their sensitive signatures match.
    """

    text_runs = _iter_paragraph_text_runs(paragraph)
    visible = [(index, run, value) for index, (run, value) in enumerate(text_runs) if value]
    if not visible:
        return []

    descriptors: dict[int, tuple[tuple[tuple[str, str], ...], tuple[str, ...]] | None] = {
        index: _format_sensitive_run_descriptor(run)
        for index, run, _value in visible
    }
    structural_kind_names = {"hyperlink", "superscript", "subscript"}

    def signature_for(index: int) -> tuple[tuple[str, str], ...]:
        descriptor = descriptors.get(index)
        return descriptor[0] if descriptor is not None else ()

    structural_indexes = {
        index
        for index, descriptor in descriptors.items()
        if descriptor is not None and structural_kind_names.intersection(descriptor[1])
    }
    unique_signatures = {signature_for(index) for index, _run, _value in visible}
    if len(unique_signatures) == 1 and not structural_indexes:
        # Multiple identically formatted runs do not contain a meaningful
        # character-range boundary.  Let the ordinary text mapper redistribute
        # prose while preserving the unchanged OOXML run graph.
        return []

    nonstructural_indexes = [
        index
        for index, _run, _value in visible
        if index not in structural_indexes
    ]
    signature_stats: dict[tuple[tuple[str, str], ...], tuple[int, int]] = {}
    visible_text_by_index = {index: value for index, _run, value in visible}
    for index in nonstructural_indexes:
        signature = signature_for(index)
        character_count, run_count = signature_stats.get(signature, (0, 0))
        signature_stats[signature] = (
            character_count + len(visible_text_by_index[index]),
            run_count + 1,
        )

    baseline_ambiguous = False
    baseline_signature: tuple[tuple[str, str], ...] = ()
    if signature_stats:
        best_score = max(signature_stats.values())
        leaders = [signature for signature, score in signature_stats.items() if score == best_score]
        if len(leaders) == 1:
            baseline_signature = leaders[0]
        elif () in leaders:
            # An unformatted run is the safest ordinary slot when run/length
            # evidence ties.  This preserves the long-standing two-run case of
            # one plain span plus one locally styled span.
            baseline_signature = ()
        else:
            baseline_ambiguous = True

    if baseline_ambiguous:
        selected_indexes = {index for index, _run, _value in visible}
    else:
        selected_indexes = set(structural_indexes)
        selected_indexes.update(
            index
            for index in nonstructural_indexes
            if signature_for(index) != baseline_signature
        )

    if not selected_indexes:
        return []

    offsets: dict[int, tuple[int, int]] = {}
    cursor = 0
    for index, (_run, run_text) in enumerate(text_runs):
        offsets[index] = (cursor, cursor + len(run_text))
        cursor += len(run_text)

    raw_groups: list[dict[str, Any]] = []
    for index, _run, run_text in visible:
        if index not in selected_indexes:
            continue
        descriptor = descriptors.get(index)
        signature, kinds = descriptor if descriptor is not None else ((), ("plain_text_boundary",))
        if baseline_ambiguous:
            kinds = tuple(sorted(set(kinds) | {"format_signature_ambiguous"}))
        if (
            raw_groups
            and int(raw_groups[-1]["run_end_index"]) + 1 == index
        ):
            # Adjacent sensitive runs with different properties still have no
            # ordinary run slot between them. Treat the complete visible span
            # as one composite anchor so a model cannot insert prose between a
            # bold token and an immediately adjacent superscript/italic token.
            raw_groups[-1]["run_end_index"] = index
            raw_groups[-1]["text"] += run_text
            raw_groups[-1]["source_end"] = offsets[index][1]
            raw_groups[-1]["kinds"] = sorted(set(raw_groups[-1]["kinds"]) | set(kinds))
            continue
        raw_groups.append(
            {
                "run_start_index": index,
                "run_end_index": index,
                "text": run_text,
                "source_start": offsets[index][0],
                "source_end": offsets[index][1],
                "kinds": list(kinds),
                "signature": signature,
                "baseline_ambiguous": baseline_ambiguous,
            }
        )

    source_text = paragraph.text
    anchors: list[dict[str, Any]] = []
    for anchor_index, group in enumerate(raw_groups):
        anchor_text = str(group["text"])
        occurrences = _text_occurrence_positions(source_text, anchor_text)
        anchors.append(
            {
                "anchor_index": anchor_index,
                "run_start_index": int(group["run_start_index"]),
                "run_end_index": int(group["run_end_index"]),
                "source_start": int(group["source_start"]),
                "source_end": int(group["source_end"]),
                "text": anchor_text,
                "text_sha256": hashlib.sha256(anchor_text.encode("utf-8")).hexdigest(),
                "kinds": list(group["kinds"]),
                "source_occurrence_count": len(occurrences),
                "source_paragraph_length": len(source_text),
                "ambiguous": bool(group.get("baseline_ambiguous")) or len(occurrences) != 1,
            }
        )
    return anchors


def _format_anchors_cover_source_text(source_text: str, anchors: Sequence[dict[str, Any]]) -> bool:
    if not anchors or not source_text:
        return False
    covered = [False] * len(source_text)
    for anchor in anchors:
        start = max(0, min(int(anchor.get("source_start", 0)), len(source_text)))
        end = max(start, min(int(anchor.get("source_end", start)), len(source_text)))
        for index in range(start, end):
            covered[index] = True
    return all(is_covered or character.isspace() for is_covered, character in zip(covered, source_text))


def validate_format_sensitive_anchors(
    rewritten_text: str,
    anchors: Sequence[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Return non-content-leaking hard issues for styled-anchor drift."""

    value = str(rewritten_text)
    issues: list[dict[str, Any]] = []
    positions: list[tuple[int, int, int, dict[str, Any]]] = []

    def anchor_evidence(anchor: dict[str, Any], anchor_index: int) -> dict[str, Any]:
        anchor_text = str(anchor.get("text", ""))
        anchor_hash = str(anchor.get("text_sha256", "")) or hashlib.sha256(anchor_text.encode("utf-8")).hexdigest()
        return {
            "anchorIndex": anchor_index,
            "anchorSha256": anchor_hash,
            "kinds": [str(item) for item in anchor.get("kinds", [])] if isinstance(anchor.get("kinds", []), list) else [],
        }

    for fallback_index, anchor in enumerate(anchors):
        anchor_index = int(anchor.get("anchor_index", fallback_index))
        anchor_text = str(anchor.get("text", ""))
        evidence = anchor_evidence(anchor, anchor_index)
        if not anchor_text or bool(anchor.get("ambiguous")) or int(anchor.get("source_occurrence_count", 1)) != 1:
            issues.append(
                {
                    "code": "format_anchor_ambiguous",
                    "message": "A format-sensitive source anchor is ambiguous and cannot be mapped safely.",
                    **evidence,
                }
            )
            continue

        occurrences = _text_occurrence_positions(value, anchor_text)
        if not occurrences:
            issues.append(
                {
                    "code": "format_anchor_missing",
                    "message": "Rewritten text removed or changed a format-sensitive anchor.",
                    **evidence,
                }
            )
            continue
        if len(occurrences) != 1:
            issues.append(
                {
                    "code": "format_anchor_duplicated",
                    "message": "Rewritten text duplicated a format-sensitive anchor, so run mapping is unsafe.",
                    "actualOccurrenceCount": len(occurrences),
                    **evidence,
                }
            )
            continue
        position = occurrences[0]
        positions.append((anchor_index, position, position + len(anchor_text), anchor))

    if len(positions) > 1:
        positions_by_anchor = [position for _index, position, _end, _anchor in sorted(positions)]
        if positions_by_anchor != sorted(positions_by_anchor):
            ordered_by_anchor = sorted(positions)
            issues.append(
                {
                    "code": "format_anchor_order_changed",
                    "message": "Rewritten text changed the order of format-sensitive anchors.",
                    "anchorIndexes": [index for index, _position, _end, _anchor in ordered_by_anchor],
                    "anchors": [
                        anchor_evidence(anchor, index)
                        for index, _position, _end, anchor in ordered_by_anchor
                    ],
                }
            )
    if positions:
        ordered_positions = sorted(positions)
        first_index, first_position, _first_end, first_anchor = ordered_positions[0]
        if int(first_anchor.get("source_start", -1)) == 0 and first_position != 0:
            issues.append(
                {
                    "code": "format_anchor_boundary_shifted",
                    "message": "Rewritten text inserted content before a paragraph-leading format anchor without an ordinary run slot.",
                    **anchor_evidence(first_anchor, first_index),
                }
            )
        last_index, _last_position, last_end, last_anchor = ordered_positions[-1]
        source_length = int(last_anchor.get("source_paragraph_length", -1))
        if source_length >= 0 and int(last_anchor.get("source_end", -1)) == source_length and last_end != len(value):
            issues.append(
                {
                    "code": "format_anchor_boundary_shifted",
                    "message": "Rewritten text appended content after a paragraph-trailing format anchor without an ordinary run slot.",
                    **anchor_evidence(last_anchor, last_index),
                }
            )
        for previous, current in zip(ordered_positions, ordered_positions[1:]):
            _previous_index, _previous_position, previous_end, previous_anchor = previous
            current_index, current_position, _current_end, current_anchor = current
            if (
                int(previous_anchor.get("source_end", -1)) == int(current_anchor.get("source_start", -2))
                and previous_end != current_position
            ):
                issues.append(
                    {
                        "code": "format_anchor_gap_unmappable",
                        "message": "Rewritten text inserted content between adjacent format anchors without an ordinary run slot.",
                        **anchor_evidence(current_anchor, current_index),
                    }
                )
    return issues


def _assert_format_sensitive_anchors(rewritten_text: str, anchors: Sequence[dict[str, Any]]) -> None:
    issues = validate_format_sensitive_anchors(rewritten_text, anchors)
    if not issues:
        return
    codes = ", ".join(sorted({str(issue.get("code", "format_anchor_invalid")) for issue in issues}))
    raise ValueError(f"DOCX format-sensitive text anchor validation failed: {codes}.")


def _set_text_node_value(node: Any, text: str) -> None:
    node.text = text
    xml_space = qn("xml:space")
    # ``xml:space=preserve`` is required for XML whitespace (most notably
    # manually typed ASCII indentation), but NBSP/full-width spaces are literal
    # Unicode characters and do not need it.  Do not remove an existing source
    # attribute when it is unnecessary: fidelity-lock mode treats the original
    # OOXML wrapper as immutable formatting evidence.
    if text[:1] in {" ", "\t", "\r", "\n"} or text[-1:] in {" ", "\t", "\r", "\n"}:
        node.set(xml_space, "preserve")


def _replace_run_text_preserving_nontext(run: Run, text: str) -> None:
    text_nodes = _run_text_nodes(run)
    if not text_nodes:
        raise ValueError("DOCX text run has no w:t node to update safely.")
    original_parts = [str(node.text or "") for node in text_nodes]
    replacement_parts = _split_text_for_runs(text, original_parts)
    for node, replacement in zip(text_nodes, replacement_parts):
        _set_text_node_value(node, replacement)


def _replace_text_around_format_anchors(
    paragraph: Paragraph,
    replacement_text: str,
    *,
    text_runs: Sequence[tuple[Run, str]],
    anchors: Sequence[dict[str, Any]],
) -> None:
    """Rewrite only ordinary run regions while styled anchors stay untouched."""

    _assert_format_sensitive_anchors(replacement_text, anchors)
    ordered_anchors = sorted(anchors, key=lambda item: int(item.get("anchor_index", 0)))
    replacement_positions: list[tuple[int, int]] = []
    for anchor in ordered_anchors:
        anchor_text = str(anchor.get("text", ""))
        occurrences = _text_occurrence_positions(replacement_text, anchor_text)
        # Validation above guarantees one exact occurrence; keep this explicit
        # so future validator changes cannot make the mapper permissive.
        if len(occurrences) != 1:
            raise ValueError("DOCX format-sensitive anchor mapping is not unique.")
        start = occurrences[0]
        replacement_positions.append((start, start + len(anchor_text)))

    anchor_run_indexes: set[int] = set()
    previous_run_end = -1
    for anchor in ordered_anchors:
        run_start = int(anchor.get("run_start_index", -1))
        run_end = int(anchor.get("run_end_index", -1))
        if run_start < 0 or run_end < run_start or run_end >= len(text_runs) or run_start <= previous_run_end:
            raise ValueError("DOCX format-sensitive anchor run mapping is invalid.")
        anchor_run_indexes.update(range(run_start, run_end + 1))
        previous_run_end = run_end

    segments: list[tuple[str, range]] = []
    replacement_cursor = 0
    run_cursor = 0
    for anchor, (replacement_start, replacement_end) in zip(ordered_anchors, replacement_positions):
        run_start = int(anchor["run_start_index"])
        run_end = int(anchor["run_end_index"])
        segments.append((replacement_text[replacement_cursor:replacement_start], range(run_cursor, run_start)))
        replacement_cursor = replacement_end
        run_cursor = run_end + 1
    segments.append((replacement_text[replacement_cursor:], range(run_cursor, len(text_runs))))

    assignments: dict[int, str] = {}
    for segment_text, candidate_indexes in segments:
        mutable_indexes = [
            index
            for index in candidate_indexes
            if index not in anchor_run_indexes
        ]
        if not mutable_indexes:
            if segment_text:
                raise ValueError(
                    "DOCX rewrite has text outside a format-sensitive anchor but no ordinary run slot can hold it."
                )
            continue
        original_parts = [text_runs[index][1] for index in mutable_indexes]
        replacement_parts = _split_text_for_runs(segment_text, original_parts)
        assignments.update(zip(mutable_indexes, replacement_parts))

    # All mapping checks occur before this mutation point, so anchor failures
    # cannot leave a partially rewritten in-memory document.
    for index, replacement in assignments.items():
        _replace_run_text_preserving_nontext(text_runs[index][0], replacement)

    for anchor in ordered_anchors:
        run_start = int(anchor["run_start_index"])
        run_end = int(anchor["run_end_index"])
        actual_anchor_text = "".join(text_runs[index][0].text for index in range(run_start, run_end + 1))
        if actual_anchor_text != str(anchor.get("text", "")):
            raise ValueError("DOCX format-sensitive anchor text moved away from its frozen run group.")


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    replacement_text = str(text)
    text_runs = _iter_paragraph_text_runs(paragraph)
    if not text_runs:
        if replacement_text:
            paragraph.add_run(replacement_text)
        return

    format_anchors = _extract_format_sensitive_anchors(paragraph)
    if format_anchors:
        _replace_text_around_format_anchors(
            paragraph,
            replacement_text,
            text_runs=text_runs,
            anchors=format_anchors,
        )
    else:
        split_parts = _split_text_for_runs(replacement_text, [run_text for _run, run_text in text_runs])
        for (run, _run_text), part in zip(text_runs, split_parts):
            _replace_run_text_preserving_nontext(run, part)
    if paragraph.text != replacement_text:
        raise ValueError(
            "DOCX paragraph contains unsupported inline text structure; "
            "refusing to flatten hyperlinks, fields, or non-text run content."
        )


def _verify_rewritten_paragraph_text(paragraph: Paragraph, expected_text: str, *, target: dict[str, Any]) -> None:
    expected = str(expected_text)
    actual = paragraph.text
    if actual != expected:
        raise ValueError(
            "DOCX export text verification failed for "
            f"{target}: expected {_preview_text(expected)!r}, got {_preview_text(actual)!r}."
        )


def _verify_saved_rewritten_targets(export_path: Path, expected_targets: Sequence[tuple[dict[str, Any], str]]) -> None:
    saved_document = Document(str(export_path.resolve()))
    try:
        for target, expected_text in expected_targets:
            paragraph = _resolve_target_paragraph(saved_document, target)
            _verify_rewritten_paragraph_text(paragraph, expected_text, target=target)
    except Exception:
        try:
            export_path.unlink(missing_ok=True)
        except OSError:
            pass
        raise


def _preview_text(text: str, limit: int = 120) -> str:
    compact = re.sub(r"\s+", " ", str(text or "")).strip()
    return compact[:limit] + ("..." if len(compact) > limit else "")


def _apply_default_document_layout(document: DocxDocument) -> None:
    for section in document.sections:
        for attribute, value_cm in DEFAULT_PAGE_MARGINS_CM.items():
            setattr(section, attribute, Cm(value_cm))

    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = DEFAULT_EN_FONT
    normal_style.font.size = Pt(DEFAULT_BODY_FONT_SIZE_PT)
    _set_run_fonts(normal_style._element.get_or_add_rPr(), cn_font=DEFAULT_CN_FONT, en_font=DEFAULT_EN_FONT)


def _apply_default_paragraph_layout(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Cm(DEFAULT_BODY_FIRST_LINE_INDENT_CM)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(DEFAULT_BODY_LINE_SPACING_PT)
    for run in paragraph.runs:
        _apply_default_run_layout(run)


def _polish_rewritten_paragraph(paragraph: Paragraph, *, preserve_format: bool = False) -> None:
    if preserve_format:
        # Preserve-original mode: the source run formatting is the single source
        # of truth, so we must NOT backfill default font/size — only the text was
        # changed, formatting (incl. inherited styles) stays exactly as-is.
        return
    for run in paragraph.runs:
        if run.text:
            _ensure_run_font_layout(run)


def _normalize_rewritten_text(text: str) -> str:
    normalized = text.replace("\u00a0", " ").replace("\u3000", " ")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\s+([，。；：！？、,.!?;:])", r"\1", normalized)
    normalized = re.sub(r"([（【《])\s+", r"\1", normalized)
    normalized = re.sub(r"\s+([）】》])", r"\1", normalized)
    return normalized.strip()


def _restore_rewritten_text_with_source_whitespace(
    rewritten_text: str,
    *,
    original_text: str,
    leading_whitespace: str = "",
    trailing_whitespace: str = "",
) -> str:
    """Normalize model prose while restoring the source's literal boundaries.

    ``rewritten_text`` is model-facing content and is deliberately normalized.
    Boundary whitespace never enters that normalization path.  If normalization
    shows the rewrite is an identity operation, retain the exact original core
    as well; this makes identity round-trips byte-for-byte literal at the text
    layer even when the model-output cleaner folded NBSP or repeated spaces.
    """

    prefix = str(leading_whitespace)
    suffix = str(trailing_whitespace)
    if any(not character.isspace() for character in prefix + suffix):
        raise ValueError("DOCX boundary whitespace evidence contains non-whitespace characters.")
    if "\n" in prefix + suffix or "\r" in prefix + suffix:
        raise ValueError("DOCX boundary whitespace evidence may not contain paragraph line breaks.")

    source_core = str(original_text)
    normalized_source_core = _normalize_rewritten_text(source_core)
    normalized_rewrite_core = _normalize_rewritten_text(str(rewritten_text))
    restored_core = source_core if normalized_rewrite_core == normalized_source_core else normalized_rewrite_core
    return f"{prefix}{restored_core}{suffix}"


def _ensure_run_font_layout(run: Any) -> None:
    if run.font.size is None:
        run.font.size = Pt(DEFAULT_BODY_FONT_SIZE_PT)
    if not str(run.font.name or "").strip():
        run.font.name = DEFAULT_EN_FONT
    _set_run_fonts(run._element.get_or_add_rPr(), cn_font=DEFAULT_CN_FONT, en_font=DEFAULT_EN_FONT, preserve_existing=True)


def _apply_default_run_layout(run: Any) -> None:
    run.font.name = DEFAULT_EN_FONT
    run.font.size = Pt(DEFAULT_BODY_FONT_SIZE_PT)
    _set_run_fonts(run._element.get_or_add_rPr(), cn_font=DEFAULT_CN_FONT, en_font=DEFAULT_EN_FONT)


def _set_run_fonts(r_pr: Any, *, cn_font: str, en_font: str, preserve_existing: bool = False) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    font_values = {
        "eastAsia": cn_font,
        "ascii": en_font,
        "hAnsi": en_font,
        "cs": en_font,
    }
    for key, value in font_values.items():
        attribute = qn(f"w:{key}")
        if preserve_existing and str(r_fonts.get(attribute) or "").strip():
            continue
        r_fonts.set(attribute, value)


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="DOCX <-> text helper with layout-preserving roundtrip support")
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract_parser = subparsers.add_parser(
        "extract", help="Extract editable plain text from a .docx file",
    )
    extract_parser.add_argument("input", type=Path, help="Path to input .docx file")

    extract_to_file_parser = subparsers.add_parser(
        "extract-to-file", help="Extract editable plain text from a .docx file into a UTF-8 text file",
    )
    extract_to_file_parser.add_argument("input", type=Path, help="Path to input .docx file")
    extract_to_file_parser.add_argument("output", type=Path, help="Path to output .txt file")
    extract_to_file_parser.add_argument("--snapshot-output", type=Path, default=None, help="Optional output path for the snapshot json.")
    extract_to_file_parser.add_argument("--scope-diagnostics-output", type=Path, default=None, help="Optional output path for body-scope diagnostics json.")

    extract_paragraphs_parser = subparsers.add_parser(
        "extract-paragraphs",
        help="Extract editable text blocks from a .docx file into a JSON array",
    )
    extract_paragraphs_parser.add_argument("input", type=Path, help="Path to input .docx file")
    extract_paragraphs_parser.add_argument("output", type=Path, help="Path to output .json file")

    build_parser = subparsers.add_parser(
        "build", help="Build a brand-new .docx file from a plain text file",
    )
    build_parser.add_argument("input", type=Path, help="Path to input .txt file")
    build_parser.add_argument("output", type=Path, help="Path to output .docx file")

    build_paragraphs_parser = subparsers.add_parser(
        "build-paragraphs",
        help="Build a brand-new .docx file from a paragraph JSON array or block text file",
    )
    build_paragraphs_parser.add_argument("input", type=Path, help="Path to paragraph json/txt file")
    build_paragraphs_parser.add_argument("output", type=Path, help="Path to output .docx file")

    roundtrip_build_parser = subparsers.add_parser(
        "build-from-source",
        help="Write rewritten text back onto the original DOCX layout",
    )
    roundtrip_build_parser.add_argument("source_docx", type=Path, help="Path to original .docx file")
    roundtrip_build_parser.add_argument("input", type=Path, help="Path to rewritten .txt file")
    roundtrip_build_parser.add_argument("output", type=Path, help="Path to output .docx file")
    roundtrip_build_parser.add_argument("--snapshot", type=Path, default=None, help="Optional snapshot json path.")

    diagnostics_parser = subparsers.add_parser(
        "scope-diagnostics",
        help="Write body-scope protection diagnostics for a .docx file",
    )
    diagnostics_parser.add_argument("input", type=Path, help="Path to input .docx file")
    diagnostics_parser.add_argument("output", type=Path, help="Path to output diagnostics json")
    diagnostics_parser.add_argument("--snapshot-output", type=Path, default=None, help="Optional output path for the snapshot json.")

    args = parser.parse_args(argv)

    if args.command == "extract":
        text = read_docx_text(args.input)
        print(text)
    elif args.command == "extract-to-file":
        extracted_path, snapshot_path, _ = ensure_docx_processing_assets(
            args.input,
            extracted_path=args.output,
            snapshot_path=args.snapshot_output,
            scope_diagnostics_path=args.scope_diagnostics_output,
        )
        if extracted_path != args.output:
            args.output.write_text(extracted_path.read_text(encoding="utf-8"), encoding="utf-8")
        if args.snapshot_output is not None and snapshot_path != args.snapshot_output:
            args.snapshot_output.write_text(snapshot_path.read_text(encoding="utf-8"), encoding="utf-8")
    elif args.command == "extract-paragraphs":
        paragraphs = read_docx_paragraphs(args.input)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(
            json.dumps(paragraphs, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    elif args.command == "build":
        text = args.input.read_text(encoding="utf-8")
        blocks = _split_text_into_blocks(text)
        write_docx_text(blocks, args.output)
    elif args.command == "build-paragraphs":
        paragraphs = _read_paragraphs_file(args.input)
        write_docx_paragraphs(paragraphs, args.output)
    elif args.command == "build-from-source":
        snapshot_path = args.snapshot or get_docx_snapshot_path(args.source_docx.resolve())
        if not snapshot_path.exists():
            ensure_docx_processing_assets(args.source_docx, snapshot_path=snapshot_path)
        text = args.input.read_text(encoding="utf-8")
        rebuild_docx_from_snapshot(
            _split_text_into_blocks(text),
            source_path=args.source_docx,
            snapshot_path=snapshot_path,
            export_path=args.output,
        )
    elif args.command == "scope-diagnostics":
        snapshot_output = args.snapshot_output or get_docx_snapshot_path(args.input.resolve())
        _, _, snapshot = ensure_docx_processing_assets(
            args.input,
            snapshot_path=snapshot_output,
            scope_diagnostics_path=args.output,
        )
        if not args.output.exists():
            args.output.parent.mkdir(parents=True, exist_ok=True)
            args.output.write_text(
                json.dumps(
                    build_docx_scope_diagnostics(snapshot, snapshot_path=snapshot_output),
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
    else:  # pragma: no cover - argparse guarantees command
        parser.error("Unknown command")


if __name__ == "__main__":  # pragma: no cover - CLI entry point
    main()
