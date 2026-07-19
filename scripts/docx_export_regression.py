from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from dataclasses import replace
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

try:
    from docx import Document  # type: ignore[import]
    from docx.enum.section import WD_SECTION  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
    from docx.shared import Cm, Pt  # type: ignore[import]
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency python-docx. Install it with: pip install python-docx") from exc

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR / "scripts"))

import app_service  # noqa: E402
from app_service import export_round_output  # noqa: E402
from docx_audit import audit_docx_export, audit_docx_ooxml_integrity, get_docx_audit_report_path, get_docx_ooxml_audit_report_path  # noqa: E402
from docx_bodymap import load_docx_body_map, validate_docx_body_map  # noqa: E402
from docx_export_guard import run_docx_pre_export_guard  # noqa: E402
from docx_pipeline import (  # noqa: E402
    build_docx_scope_diagnostics,
    build_docx_snapshot,
    _load_docx_snapshot,
    _looks_like_acknowledgement_heading,
    _looks_like_back_matter_heading,
    _looks_like_caption,
    _looks_like_formula_paragraph,
    _looks_like_heading,
    _looks_like_keyword_line,
    _looks_like_note,
    _looks_like_references_heading,
    _normalize_rewritten_text,
    _polish_rewritten_paragraph,
    _replace_paragraph_text,
    _restore_rewritten_text_with_source_whitespace,
    get_docx_scope_diagnostics_path,
    get_docx_snapshot_path,
)
from fyadr_round_service import _extract_required_terms, find_english_spacing_corruptions, find_sentence_surface_issues, validate_chunk_output  # noqa: E402
from round_helper import run_document_round  # noqa: E402

REGRESSION_DIR = ROOT_DIR / "finish" / "regression"
DEFAULT_SAMPLE_PATH = REGRESSION_DIR / "fyadr_regression_sample.docx"
DEFAULT_EXPORT_PATH = REGRESSION_DIR / "fyadr_regression_export.docx"
DEFAULT_REPORT_PATH = REGRESSION_DIR / "fyadr_regression_report.json"
AUTO_NUMBERED_BODY_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)*|[一二三四五六七八九十]+)[\.．、)]\s*\S")


def zh(*codes: int) -> str:
    return "".join(chr(code) for code in codes)


def _set_paragraph_font(paragraph: Any, *, size_pt: int = 12, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), zh(0x5b8b, 0x4f53))


try:
    from docx.oxml.ns import qn  # type: ignore[import]
except ImportError:  # pragma: no cover
    qn = None  # type: ignore[assignment]


def add_paragraph(document: Any, text: str, *, style: str | None = None, align: int | None = None, size_pt: int = 12, bold: bool | None = None) -> Any:
    paragraph = document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    if align is not None:
        paragraph.alignment = align
    _set_paragraph_font(paragraph, size_pt=size_pt, bold=bold)
    return paragraph


def _run_english_run_spacing_regression() -> dict[str, Any]:
    document = Document()
    paragraph = document.add_paragraph()
    for run_text in ("old ", "text ", "with ", "many ", "runs ", "here ", "please"):
        paragraph.add_run(run_text)
    expected = (
        "Using Qwen2.5-1.5B-Instruct as the base model, a LoRA adapter is then constructed "
        "with approach, 500 samples employing 4-bit QLoRA. In addition, Key words: LoRA; QLoRA"
    )
    failures: list[str] = []
    _replace_paragraph_text(paragraph, expected)
    after_replace = paragraph.text
    _polish_rewritten_paragraph(paragraph)
    actual = paragraph.text
    normalized_expected = _normalize_rewritten_text(expected)
    if after_replace != normalized_expected:
        failures.append("replace step changed English paragraph text")
    if actual != normalized_expected:
        failures.append("polish step changed English paragraph text")
    required_fragments = (
        "Using Qwen2.5-1.5B-Instruct",
        "as the",
        "is then constructed",
        "approach, 500",
        "employing 4-bit QLoRA. In addition",
        "Key words: LoRA; QLoRA",
    )
    forbidden_fragments = (
        "UsingQwen",
        "asthe",
        "isthen",
        "approach,500",
        "employing4-bit",
        "QLoRA.In",
    )
    missing_fragments = [fragment for fragment in required_fragments if fragment not in actual]
    leaked_fragments = [fragment for fragment in forbidden_fragments if fragment in actual]
    if missing_fragments:
        failures.append(f"missing fragments: {missing_fragments}")
    if leaked_fragments:
        failures.append(f"bad fragments leaked: {leaked_fragments}")
    validate_chunk_output(normalized_expected, actual, "english-run-spacing")
    bad_output = (
        "UsingQwen2.5-1.5B-Instruct asthe base model, a LoRA adapter isthen constructed "
        "with approach,500 samples employing4-bit QLoRA.In addition, Key words: LoRA; QL"
    )
    corruptions = find_english_spacing_corruptions(normalized_expected, bad_output)
    if len(corruptions) < 5:
        failures.append(f"spacing guard missed corruptions: {corruptions}")
    terms = _extract_required_terms(normalized_expected)
    if "LoRA" not in terms or "QLoRA" not in terms or "Qwen2.5-1.5B-Instruct" not in terms:
        failures.append(f"term guard missed technical terms: {sorted(terms)}")
    try:
        validate_chunk_output(normalized_expected, bad_output, "english-run-spacing-bad")
        failures.append("validation accepted corrupted English output")
    except ValueError:
        pass
    # Use a unique protected term for the truncation probe.  QLoRA appears more
    # than once in this fixture and occurrence counts are intentionally allowed
    # to change during rewriting; truncating the unique versioned model token
    # must still be rejected as a missing required term.
    bad_term_output = normalized_expected.replace("Qwen2.5-1.5B-Instruct", "Qwen2.5-1.5B", 1)
    try:
        validate_chunk_output(normalized_expected, bad_term_output, "english-term-count-bad")
        failures.append("validation accepted truncated protected technical term")
    except ValueError:
        pass
    return {
        "ok": not failures,
        "failures": failures,
        "afterReplace": after_replace,
        "actual": actual,
        "detectedCorruptions": corruptions,
        "requiredTerms": sorted(terms),
    }


def _run_sentence_surface_integrity_regression() -> dict[str, Any]:
    source = "本文采用YOLOv8模型进行识别，并在移动端完成部署。"
    valid_output = "本文选用YOLOv8模型完成识别任务，并在移动端实现部署。"
    cases = {
        "truncated": "本文采用YOLOv8模型进行识别，并在移动端完成",
        "wrapper": "改写后：本文采用YOLOv8模型进行识别，并在移动端完成部署。",
        "unbalanced": "本文采用（YOLOv8模型进行识别，并在移动端完成部署。",
        "repeated_punctuation": "本文采用YOLOv8模型进行识别，，并在移动端完成部署。",
    }
    failures: list[str] = []
    clean_issues = find_sentence_surface_issues(source, valid_output)
    if clean_issues:
        failures.append(f"surface guard rejected valid output: {clean_issues}")
    validate_chunk_output(source, valid_output, "sentence-surface-valid")
    detected: dict[str, list[str]] = {}
    for name, output in cases.items():
        issues = find_sentence_surface_issues(source, output)
        detected[name] = [str(issue.get("code", "")) for issue in issues]
        if not issues:
            failures.append(f"surface guard missed {name}")
        try:
            validate_chunk_output(source, output, f"sentence-surface-{name}")
            failures.append(f"chunk validation accepted {name}")
        except ValueError:
            pass
    return {
        "ok": not failures,
        "failures": failures,
        "validIssues": clean_issues,
        "detected": detected,
    }


def _build_ack_boundary_sample(path: Path, tail: list[str]) -> dict[str, Any]:
    document = Document()
    add_paragraph(document, "测试论文", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, "摘 要", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, "本文测试正文边界识别能力，确保摘要到致谢结束范围内的正文可以改写。")
    add_paragraph(document, "致 谢", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, "感谢导师与同学在论文写作过程中的帮助。")
    for item in tail:
        add_paragraph(document, item)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    snapshot = build_docx_snapshot(path)
    diagnostics = build_docx_scope_diagnostics(snapshot)
    return {
        "path": str(path),
        "diagnostics": diagnostics,
        "editableTexts": [unit.text for unit in snapshot.units if unit.editable],
        "protectedTexts": [unit.text for unit in snapshot.units if not unit.editable],
    }


def _run_acknowledgement_boundary_regression() -> dict[str, Any]:
    failures: list[str] = []
    doc_end_case = _build_ack_boundary_sample(REGRESSION_DIR / "ack_boundary_doc_end.docx", [])
    doc_end_issues = [
        str(issue.get("code", ""))
        for issue in doc_end_case["diagnostics"].get("issues", [])
        if isinstance(issue, dict)
    ]
    if doc_end_case["diagnostics"].get("errorCount") != 0:
        failures.append(f"acknowledgement-at-document-end should not be an error: {doc_end_issues}")
    if "missing_post_acknowledgement_boundary" in doc_end_issues:
        failures.append("acknowledgement-at-document-end should not warn about a missing post-ack boundary")

    references_case = _build_ack_boundary_sample(
        REGRESSION_DIR / "ack_boundary_references.docx",
        ["参考文献", "[1] Zhang A. A sample reference entry."],
    )
    references_scope = references_case["diagnostics"].get("scope", {})
    if references_scope.get("postAcknowledgementBoundaryIndex") is None:
        failures.append("normal-style references heading after acknowledgement should be detected as a boundary")
    if any("参考文献" in text for text in references_case["editableTexts"]):
        failures.append("references heading after acknowledgement should not be editable")
    if any("[1] Zhang" in text for text in references_case["editableTexts"]):
        failures.append("reference entries after acknowledgement should not be editable")

    variant_boundaries: dict[str, int | None] = {}
    for case_name, heading, entries in (
        (
            "reference_list",
            "Reference List",
            [
                "[1] Smith J. Reliable control[J]. Systems Journal, 2021, 12(3): 44-59.",
                "[2] Brown A. Boundary evaluation[M]. Academic Press, 2022.",
            ],
        ),
        (
            "sources_consulted",
            "Sources Consulted",
            [
                "Smith, J. (2021). Reliable Control. Systems Journal, 12(3), 44-59.",
                "Brown, A. (2022). Boundary Evaluation. Academic Press.",
            ],
        ),
        (
            "major_reference_materials_cn",
            "主要参考资料",
            [
                "张三，李四. 时序约束研究[J]. 控制学报，2021，12(3)：44-52.",
                "王五. 异构控制边界分析[M]. 北京：科学出版社，2022.",
            ],
        ),
    ):
        variant_case = _build_ack_boundary_sample(
            REGRESSION_DIR / f"ack_boundary_{case_name}.docx",
            [heading, *entries],
        )
        variant_scope = variant_case["diagnostics"].get("scope", {})
        variant_boundaries[case_name] = variant_scope.get("postAcknowledgementBoundaryIndex")
        if variant_boundaries[case_name] is None:
            failures.append(f"{heading!r} should be recognized as a post-ack reference boundary")
        for protected_text in (heading, *entries):
            if protected_text in variant_case["editableTexts"]:
                failures.append(f"{case_name}: reference heading/entry became editable: {protected_text}")

    entry_run_case = _build_ack_boundary_sample(
        REGRESSION_DIR / "ack_boundary_reference_entry_run.docx",
        [
            "[1] Taylor K. Robust sequence estimation[J]. Automation Journal, 2020, 6(1): 1-12.",
            "[2] Clark M. Constraint propagation in practice[M]. Technical Press, 2021.",
        ],
    )
    entry_run_scope = entry_run_case["diagnostics"].get("scope", {})
    if entry_run_scope.get("postAcknowledgementBoundaryIndex") is None:
        failures.append("two consecutive reference entries should establish a fail-safe post-ack boundary")
    if any("Taylor K." in text or "Clark M." in text for text in entry_run_case["editableTexts"]):
        failures.append("consecutive reference entries without a heading should not be editable")
    entry_run_issue_codes = [
        str(issue.get("code", ""))
        for issue in entry_run_case["diagnostics"].get("issues", [])
        if isinstance(issue, dict)
    ]
    if "unbounded_post_acknowledgement_tail" in entry_run_issue_codes:
        failures.append("a confirmed consecutive reference run should be treated as a boundary, not an unbounded tail")

    author_year_run_case = _build_ack_boundary_sample(
        REGRESSION_DIR / "ack_boundary_author_year_reference_run.docx",
        [
            "Taylor, K. (2020). Robust Sequence Estimation. Automation Journal, 6(1), 1-12.",
            "Clark, M. (2021). Constraint Propagation in Practice. Technical Press.",
        ],
    )
    author_year_run_scope = author_year_run_case["diagnostics"].get("scope", {})
    if author_year_run_scope.get("postAcknowledgementBoundaryIndex") is None:
        failures.append("two consecutive author-year references should establish a fail-safe post-ack boundary")
    if any("Taylor, K." in text or "Clark, M." in text for text in author_year_run_case["editableTexts"]):
        failures.append("consecutive author-year references without a heading should not be editable")

    risky_tail_case = _build_ack_boundary_sample(
        REGRESSION_DIR / "ack_boundary_unbounded_tail.docx",
        ["[1] Zhang A. A sample reference entry without an explicit heading."],
    )
    risky_codes = [
        str(issue.get("code", ""))
        for issue in risky_tail_case["diagnostics"].get("issues", [])
        if isinstance(issue, dict)
    ]
    if "unbounded_post_acknowledgement_tail" not in risky_codes:
        failures.append(f"unbounded post-ack reference-like tail should be an error: {risky_codes}")

    return {
        "ok": not failures,
        "failures": failures,
        "docEndIssueCodes": doc_end_issues,
        "referencesBoundaryIndex": references_scope.get("postAcknowledgementBoundaryIndex"),
        "variantBoundaryIndexes": variant_boundaries,
        "entryRunBoundaryIndex": entry_run_scope.get("postAcknowledgementBoundaryIndex"),
        "authorYearRunBoundaryIndex": author_year_run_scope.get("postAcknowledgementBoundaryIndex"),
        "riskyTailIssueCodes": risky_codes,
    }


def _run_compare_fallback_spacing_regression() -> dict[str, Any]:
    payload = {
        "paragraphCount": 1,
        "chunkCount": 3,
        "chunks": [
            {
                "chunkId": "p0_c0",
                "paragraphIndex": 0,
                "chunkIndex": 0,
                "inputText": "Using Qwen2.5-1.5B-Instruct as",
                "outputText": "Using Qwen2.5-1.5B-Instruct as",
            },
            {
                "chunkId": "p0_c1",
                "paragraphIndex": 0,
                "chunkIndex": 1,
                "inputText": "the base model,",
                "outputText": "the base model,",
            },
            {
                "chunkId": "p0_c2",
                "paragraphIndex": 0,
                "chunkIndex": 2,
                "inputText": "500 samples employing 4-bit QLoRA.",
                "outputText": "500 samples employing 4-bit QLoRA.",
            },
        ],
    }
    paragraphs = app_service._build_paragraphs_from_compare_payload(
        REGRESSION_DIR / "missing_manifest_round.txt",
        payload,
    )
    actual = paragraphs[0] if paragraphs else ""
    failures: list[str] = []
    for fragment in ("as the", "model, 500", "4-bit QLoRA"):
        if fragment not in actual:
            failures.append(f"missing fallback spacing fragment: {fragment}")
    for fragment in ("asthe", "model,500", "4-bitQLoRA"):
        if fragment in actual:
            failures.append(f"fallback compare restore glued English text: {fragment}")
    return {
        "ok": not failures,
        "failures": failures,
        "actual": actual,
    }


def _run_compare_payload_integrity_regression(
    *,
    source_path: Path,
    snapshot_path: Path,
    body_map: Any | None,
) -> dict[str, Any]:
    failures: list[str] = []
    fallback_output = REGRESSION_DIR / "compare_integrity_output.txt"
    fallback_payload = {
        "paragraphCount": 3,
        "chunkCount": 2,
        "chunks": [
            {"chunkId": "p0_c0", "paragraphIndex": 0, "chunkIndex": 0, "inputText": "A", "outputText": "A"},
            {"chunkId": "p2_c0", "paragraphIndex": 2, "chunkIndex": 0, "inputText": "C", "outputText": "C"},
        ],
    }
    if app_service._build_paragraphs_from_compare_payload(fallback_output, fallback_payload) is not None:
        failures.append("compare fallback accepted a payload with a missing paragraph")

    duplicate_payload = {
        "paragraphCount": 2,
        "chunkCount": 3,
        "chunks": [
            {"chunkId": "p0_c0a", "paragraphIndex": 0, "chunkIndex": 0, "inputText": "A", "outputText": "A"},
            {"chunkId": "p0_c0b", "paragraphIndex": 0, "chunkIndex": 0, "inputText": "B", "outputText": "B"},
            {"chunkId": "p1_c0", "paragraphIndex": 1, "chunkIndex": 0, "inputText": "C", "outputText": "C"},
        ],
    }
    if app_service._build_paragraphs_from_compare_payload(fallback_output, duplicate_payload) is not None:
        failures.append("compare fallback accepted duplicate paragraph chunk positions")

    if body_map is not None and len(body_map.units) >= 2:
        paragraphs = body_map.current_texts()
        fallback_output.write_text("\n\n".join(paragraphs), encoding="utf-8")
        guard_payload = {
            "paragraphCount": len(paragraphs),
            "chunkCount": len(paragraphs) - 1,
            "chunks": [
                {
                    "chunkId": f"p{index}_c0",
                    "paragraphIndex": index,
                    "chunkIndex": 0,
                    "inputText": text,
                    "outputText": text,
                }
                for index, text in enumerate(paragraphs)
                if index != 1
            ],
        }
        report = run_docx_pre_export_guard(
            paragraphs,
            source_path=source_path,
            snapshot_path=snapshot_path,
            export_path=REGRESSION_DIR / "compare_integrity_export.docx",
            output_path=fallback_output,
            body_map=body_map,
            compare_payload=guard_payload,
            mode="regression",
            paragraph_source="compare-integrity",
        )
        issue_codes = {
            str(issue.get("code", ""))
            for issue in report.get("blockingIssues", [])
            if isinstance(issue, dict)
        }
        if report.get("ok") or "compare_paragraph_missing" not in issue_codes:
            failures.append(f"pre-export guard missed missing compare paragraph: {sorted(issue_codes)}")

    return {
        "ok": not failures,
        "failures": failures,
    }


def create_regression_sample(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    add_paragraph(document, zh(0x57fa, 0x4e8e, 0x56fe, 0x50cf, 0x5206, 0x5272, 0x7684, 0x5178, 0x578b, 0x70df, 0x53f6, 0x75c5, 0x866b, 0x5bb3, 0x76ee, 0x6807, 0x68c0, 0x6d4b, 0x7b97, 0x6cd5), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=18, bold=True)
    add_paragraph(document, "Typical Tobacco Disease and Pest Detection Based on Image Segmentation", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=18, bold=True)
    add_paragraph(document, zh(0x5b66, 0x751f, 0x59d3, 0x540d, 0xff1a, 0x5f20, 0x4e09), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=15)
    add_paragraph(document, zh(0x4e8c, 0x25cb, 0x4e8c, 0x516d, 0x5e74, 0x56db, 0x6708), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=15)

    document.add_page_break()
    add_paragraph(document, zh(0x76ee, 0x20, 0x20, 0x5f55), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, zh(0x6458, 0x8981) + " ................................................................ I")
    add_paragraph(document, zh(0x82f1, 0x6587, 0x6458, 0x8981) + " .................................................... II")
    add_paragraph(document, "1 " + zh(0x5f15, 0x8a00) + " .............................................................. 1")
    add_paragraph(document, zh(0x53c2, 0x8003, 0x6587, 0x732e) + " ...................................................... 8")
    add_paragraph(document, zh(0x81f4, 0x8c22) + " .............................................................. 9")

    document.add_page_break()
    add_paragraph(document, zh(0x6458, 0x20, 0x20, 0x8981), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, zh(0x672c, 0x6587, 0x56f4, 0x7ed5, 0x70df, 0x53f6, 0x75c5, 0x866b, 0x5bb3, 0x56fe, 0x50cf, 0x68c0, 0x6d4b, 0x4efb, 0x52a1, 0x5c55, 0x5f00, 0xff0c, 0x901a, 0x8fc7, 0x5206, 0x5272, 0x7ed3, 0x679c, 0x4e0e, 0x76ee, 0x6807, 0x68c0, 0x6d4b, 0x6a21, 0x578b, 0x8fdb, 0x884c, 0x7ed3, 0x5408, 0xff0c, 0x63d0, 0x9ad8, 0x7530, 0x95f4, 0x8bc6, 0x522b, 0x6548, 0x7387, 0x3002))
    add_paragraph(document, zh(0x5173, 0x952e, 0x8bcd, 0xff1a, 0x56fe, 0x50cf, 0x5206, 0x5272, 0xff1b, 0x75c5, 0x866b, 0x5bb3, 0x68c0, 0x6d4b, 0xff1b, 0x6df1, 0x5ea6, 0x5b66, 0x4e60))
    add_paragraph(document, zh(0x82f1, 0x6587, 0x6458, 0x8981), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, "The English abstract boundary paragraph stays in the editable body scope and keeps the same export formatting.")
    add_paragraph(document, "Abstract", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, "This paper studies a detection workflow for tobacco diseases and pests, combining segmentation outputs with a lightweight detector to improve field recognition robustness.")
    add_paragraph(document, "Key words: image segmentation; pest detection; deep learning")

    add_paragraph(document, zh(0x5f15, 0x8a00), style="Heading 1")
    add_paragraph(document, zh(0x70df, 0x53f6, 0x75c5, 0x866b, 0x5bb3, 0x7684, 0x65e9, 0x671f, 0x8bc6, 0x522b, 0x5bf9, 0x751f, 0x4ea7, 0x7ba1, 0x7406, 0x5177, 0x6709, 0x76f4, 0x63a5, 0x4ef7, 0x503c, 0x3002) + zh(0x76f8, 0x5173, 0x7814, 0x7a76) + zh(0xff08, 0x90ed, 0x6c34, 0x826f, 0xff0c) + "1999" + zh(0xff1b, 0x94b1, 0x5b8f, 0xff0c) + "1990" + zh(0xff09, 0x8868, 0x660e, 0xff0c, 0x5408, 0x7406, 0x7684, 0x56fe, 0x50cf, 0x9884, 0x5904, 0x7406, 0x80fd, 0x591f, 0x964d, 0x4f4e, 0x8bef, 0x68c0, 0x7387, 0x3002))
    add_paragraph(document, "1 " + zh(0x6570, 0x636e, 0x4e0e, 0x65b9, 0x6cd5), style="Heading 1")
    preserved_spacing_paragraph = add_paragraph(document, zh(0x672c, 0x7814, 0x7a76, 0x91c7, 0x7528) + "YOLOv8" + zh(0x4f5c, 0x4e3a, 0x57fa, 0x7840, 0x6a21, 0x578b, 0xff0c, 0x8bad, 0x7ec3, 0x56fe, 0x50cf, 0x5c3a, 0x5bf8, 0x63a7, 0x5236, 0x5728) + "640px" + zh(0xff0c, 0x5b66, 0x4e60, 0x7387, 0x8303, 0x56f4, 0x4e3a) + "0.001?0.01" + zh(0x3002))
    preserved_spacing_paragraph.paragraph_format.space_before = Pt(7)
    preserved_spacing_paragraph.paragraph_format.space_after = Pt(9)
    add_paragraph(document, "1.1 " + zh(0x7cfb, 0x7edf, 0x5b9e, 0x73b0), style="Heading 2")
    add_paragraph(document, zh(0x7cfb, 0x7edf, 0x5b9e, 0x73b0, 0x8fb9, 0x754c, 0x6bb5, 0x843d, 0x63cf, 0x8ff0, 0x4e86, 0x79fb, 0x52a8, 0x7aef, 0x63a8, 0x7406, 0x4e0e, 0x7ed3, 0x679c, 0x5c55, 0x793a, 0x6d41, 0x7a0b, 0xff0c, 0x5e94, 0x5f53, 0x4f5c, 0x4e3a, 0x6b63, 0x6587, 0x53c2, 0x4e0e, 0x6539, 0x5199, 0x3002))
    add_paragraph(document, "3." + zh(0x7ed3, 0x679c, 0x5206, 0x6790))
    add_paragraph(document, zh(0x7ed3, 0x679c, 0x5206, 0x6790, 0x540e, 0x7684, 0x8fb9, 0x754c, 0x6bb5, 0x843d, 0x5bf9, 0x6a21, 0x578b, 0x8bef, 0x68c0, 0x60c5, 0x51b5, 0x8fdb, 0x884c, 0x8865, 0x5145, 0x8bf4, 0x660e, 0x3002))
    # Keep one auto-numbered body paragraph with a measurable, safely removable
    # style risk.  The targeted-rerun smoke later removes the content-light
    # transition while preserving the numbering, facts and DOCX anchors.  This
    # makes the production selector publish the fixture for a real combined
    # style gain instead of relying on a zero-penalty synonym tie.
    add_paragraph(
        document,
        "1. 登录注册：具体而言，用户通过用户名或手机号创建账户并登录系统，以访问平台的各项管理功能。",
    )
    add_paragraph(document, "2." + zh(0x6743, 0x9650, 0x7ba1, 0x7406, 0x529f, 0x80fd, 0x8bf4, 0x660e))
    add_paragraph(document, zh(0x767b, 0x5f55, 0x6ce8, 0x518c, 0x6a21, 0x5757, 0x56f4, 0x7ed5, 0x8d26, 0x53f7, 0x521b, 0x5efa, 0x3001, 0x8eab, 0x4efd, 0x6821, 0x9a8c, 0x548c, 0x4f1a, 0x8bdd, 0x4fdd, 0x6301, 0x5c55, 0x5f00, 0x3002), style="List Number")
    add_paragraph(document, zh(0x56fe, 0x20) + "1 " + zh(0x70df, 0x53f6, 0x75c5, 0x6591, 0x533a, 0x57df, 0x793a, 0x610f), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(document, zh(0x8868, 0x20) + "1 " + zh(0x5b9e, 0x9a8c, 0x7ed3, 0x679c, 0x5bf9, 0x6bd4), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = zh(0x6a21, 0x578b)
    table.cell(0, 1).text = "mAP"
    table.cell(0, 2).text = zh(0x901f, 0x5ea6)
    table.cell(1, 0).text = "YOLOv8"
    table.cell(1, 1).text = "87.2%"
    table.cell(1, 2).text = "35 FPS"
    table.cell(2, 0).text = "DeepLabV3+"
    table.cell(2, 1).text = "84.5%"
    table.cell(2, 2).text = "18 FPS"
    add_paragraph(document, zh(0x6ce8, 0xff1a, 0x8868, 0x4e2d, 0x6570, 0x636e, 0x4ec5, 0x7528, 0x4e8e, 0x56de, 0x5f52, 0x9a8c, 0x6536, 0x6837, 0x672c, 0x3002))
    add_paragraph(document, zh(0x53c2, 0x8003, 0x6587, 0x732e), style="Heading 1")
    add_paragraph(document, "[1] " + zh(0x90ed, 0x6c34, 0x826f) + ". " + zh(0x751f, 0x6001, 0x7cfb, 0x7edf, 0x7814, 0x7a76) + ". 1999.")
    add_paragraph(document, zh(0x81f4, 0x20, 0x20, 0x8c22), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True)
    add_paragraph(document, zh(0x611f, 0x8c22, 0x5bfc, 0x5e08, 0x548c, 0x540c, 0x5b66, 0x5728, 0x8bba, 0x6587, 0x5b8c, 0x6210, 0x8fc7, 0x7a0b, 0x4e2d, 0x63d0, 0x4f9b, 0x7684, 0x5e2e, 0x52a9, 0x3002))
    add_paragraph(document, zh(0x81f4, 0x8c22, 0x7b2c, 0x4e8c, 0x6bb5, 0x4ecd, 0x5904, 0x4e8e, 0x6b63, 0x6587, 0x8fb9, 0x754c, 0x5185, 0xff0c, 0x5e94, 0x5f53, 0x5141, 0x8bb8, 0x5c40, 0x90e8, 0x6539, 0x5199, 0x5e76, 0x4fdd, 0x7559, 0x5bfc, 0x5e08, 0x4e0e, 0x540c, 0x5b66, 0x7684, 0x5e2e, 0x52a9, 0x4fe1, 0x606f, 0x3002))
    add_paragraph(document, zh(0x53c2, 0x8003, 0x6587, 0x732e), style="Heading 1")
    add_paragraph(document, "[2] " + zh(0x81f4, 0x8c22, 0x4e4b, 0x540e, 0x7684, 0x53c2, 0x8003, 0x6587, 0x732e, 0x5e94, 0x5f53, 0x88ab, 0x4fdd, 0x62a4, 0x3002))
    add_paragraph(document, zh(0x540e, 0x7f6e, 0x68c0, 0x67e5, 0x9875), style="Heading 1")
    add_paragraph(document, "post-ack protected page must not be rewritten.")
    add_paragraph(document, zh(0x9644, 0x5f55) + "A " + zh(0x7cfb, 0x7edf, 0x622a, 0x56fe), style="Heading 1")
    add_paragraph(document, zh(0x9644, 0x5f55, 0x5185, 0x5bb9, 0x4ec5, 0x7528, 0x4e8e, 0x4fdd, 0x62a4, 0x533a, 0x56de, 0x5f52, 0xff0c, 0x4e0d, 0x5e94, 0x8fdb, 0x5165, 0x6539, 0x5199, 0x6d41, 0x7a0b, 0x3002))
    add_paragraph(document, zh(0x8bda, 0x4fe1, 0x58f0, 0x660e), style="Heading 1")
    add_paragraph(document, zh(0x8bda, 0x4fe1, 0x58f0, 0x660e, 0x9875, 0x5c5e, 0x4e8e, 0x81f4, 0x8c22, 0x4e4b, 0x540e, 0x7684, 0x4fdd, 0x62a4, 0x533a, 0xff0c, 0x4e0d, 0x53c2, 0x4e0e, 0x6539, 0x5199, 0x3002))

    document.save(str(path))
    return path


def identity_transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
    return chunk_text


def _pt_value(value: Any) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        return float(value.pt)
    if isinstance(value, (int, float)):
        return float(value)
    return None


def _first_text_run(paragraph: Any) -> Any | None:
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def _audit_exported_editable_format(export_path: Path, snapshot_path: Path) -> dict[str, Any]:
    document = Document(str(export_path.resolve()))
    snapshot = _load_docx_snapshot(snapshot_path)
    checks: list[dict[str, Any]] = []
    issues: list[dict[str, Any]] = []
    if snapshot is None:
        return {"ok": False, "checked": 0, "issues": [{"type": "missing_snapshot"}]}

    acknowledgement_editable_unit_indexes: set[int] = set()
    in_acknowledgements = False
    for unit in snapshot.units:
        if _looks_like_acknowledgement_heading(unit.text):
            in_acknowledgements = True
            continue
        if in_acknowledgements and (unit.protect_reason == "outside_body_scope" or _looks_like_back_matter_heading(unit.text)):
            in_acknowledgements = False
        if in_acknowledgements and unit.editable:
            acknowledgement_editable_unit_indexes.add(unit.unit_index)

    for unit in snapshot.units:
        if not unit.editable or str(unit.target.get("kind", "")) != "paragraph":
            continue
        try:
            paragraph_index = int(unit.target["paragraph_index"])
            paragraph = document.paragraphs[paragraph_index]
        except (KeyError, TypeError, ValueError, IndexError):
            issues.append({"type": "missing_editable_paragraph", "unitIndex": unit.unit_index})
            continue
        run = _first_text_run(paragraph)
        font_size = _pt_value(run.font.size) if run is not None else None
        line_spacing = _pt_value(paragraph.paragraph_format.line_spacing)
        space_before = _pt_value(paragraph.paragraph_format.space_before)
        space_after = _pt_value(paragraph.paragraph_format.space_after)
        alignment = int(paragraph.alignment) if paragraph.alignment is not None else None
        paragraph_text = paragraph.text.strip()
        expected_font_size = 12.0 if unit.unit_index in acknowledgement_editable_unit_indexes else 10.5
        check = {
            "unitIndex": unit.unit_index,
            "paragraphIndex": paragraph_index,
            "text": paragraph_text[:80],
            "fontSizePt": font_size,
            "expectedFontSizePt": expected_font_size,
            "lineSpacingPt": line_spacing,
            "spaceBeforePt": space_before,
            "spaceAfterPt": space_after,
            "alignment": alignment,
        }
        checks.append(check)
        if font_size is None or abs(font_size - expected_font_size) > 0.2:
            issues.append({"type": "editable_font_size", "expected": expected_font_size, **check})
        if line_spacing is None or abs(line_spacing - 20.0) > 0.5:
            issues.append({"type": "editable_line_spacing", "expected": 20.0, **check})
        if "YOLOv8" in paragraph_text and "640px" in paragraph_text:
            if space_before is None or abs(space_before - 7.0) > 0.2:
                issues.append({"type": "unspecified_space_before_preserved", "expected": 7.0, **check})
            if space_after is None or abs(space_after - 9.0) > 0.2:
                issues.append({"type": "unspecified_space_after_preserved", "expected": 9.0, **check})

    first_section = document.sections[0]
    margins = {
        "topMarginCm": round(float(first_section.top_margin.cm), 2),
        "bottomMarginCm": round(float(first_section.bottom_margin.cm), 2),
        "leftMarginCm": round(float(first_section.left_margin.cm), 2),
        "rightMarginCm": round(float(first_section.right_margin.cm), 2),
    }
    expected_margins = {"topMarginCm": 2.5, "bottomMarginCm": 2.5, "leftMarginCm": 3.0, "rightMarginCm": 3.0}
    for key, expected in expected_margins.items():
        if abs(margins[key] - expected) > 0.08:
            issues.append({"type": "page_margin", "key": key, "expected": expected, "actual": margins[key]})
    return {"ok": not issues, "checked": len(checks), "issues": issues, "sampleChecks": checks[:8], "margins": margins}


def _audit_exported_text_integrity(export_path: Path, body_map: Any) -> dict[str, Any]:
    if body_map is None:
        return {"ok": False, "checked": 0, "issues": [{"type": "missing_body_map"}]}
    document = Document(str(export_path.resolve()))
    issues: list[dict[str, Any]] = []
    checks: list[dict[str, Any]] = []
    for unit_index, unit in enumerate(getattr(body_map, "units", []) or []):
        target = getattr(unit, "target", {})
        if not isinstance(target, dict) or str(target.get("kind", "")) != "paragraph":
            issues.append({"type": "unsupported_target", "unitIndex": unit_index, "target": target})
            continue
        paragraph_index = int(target.get("paragraph_index", -1))
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            issues.append({"type": "target_out_of_range", "unitIndex": unit_index, "target": target})
            continue
        expected = _restore_rewritten_text_with_source_whitespace(
            str(getattr(unit, "current_text", "")),
            original_text=str(getattr(unit, "original_text", "")),
            leading_whitespace=str(getattr(unit, "leading_whitespace", "")),
            trailing_whitespace=str(getattr(unit, "trailing_whitespace", "")),
        )
        actual = document.paragraphs[paragraph_index].text
        check = {
            "unitIndex": unit_index,
            "paragraphIndex": paragraph_index,
            "expectedPreview": expected[:90],
            "actualPreview": actual[:90],
        }
        checks.append(check)
        if actual != expected:
            issues.append({"type": "exported_text_changed", **check})
    return {
        "ok": not issues,
        "checked": len(checks),
        "issues": issues[:20],
        "sampleChecks": checks[:8],
    }


def _run_export_text_integrity_block_regression(output_path: Path, export_path: Path) -> dict[str, Any]:
    failures: list[str] = []
    body_map_path = output_path.with_name(f"{output_path.stem}_body_map.json")
    body_map = load_docx_body_map(body_map_path)
    if body_map is None or not getattr(body_map, "units", None):
        return {"ok": False, "failures": [f"missing body map for integrity block check: {body_map_path}"]}

    first_unit = body_map.units[0]
    target = getattr(first_unit, "target", {})
    blocked_export_path = export_path.with_name(f"{export_path.stem}_text_integrity_blocked{export_path.suffix}")
    original_rebuild = app_service.rebuild_docx_from_body_map_units

    def corrupting_rebuild(*args: Any, **kwargs: Any) -> Any:
        result = original_rebuild(*args, **kwargs)
        target_path = Path(str(kwargs.get("export_path", blocked_export_path)))
        document = Document(str(target_path.resolve()))
        paragraph = app_service._resolve_target_paragraph(document, target)
        _replace_paragraph_text(paragraph, f"{paragraph.text} TEXT_INTEGRITY_CORRUPTION")
        document.save(target_path)
        return result

    try:
        app_service.rebuild_docx_from_body_map_units = corrupting_rebuild
        try:
            export_round_output(str(output_path), str(blocked_export_path), "docx", "school_rules")
            failures.append("export accepted a corrupted rewritten DOCX target")
            export_failure: dict[str, Any] = {}
        except app_service.ExportRoundError as exc:
            export_failure = exc.export_failure
            if export_failure.get("stage") != "text-integrity":
                failures.append(f"unexpected export failure stage: {export_failure.get('stage')}")
            report_path = Path(str(export_failure.get("reportPath", ""))) if export_failure.get("reportPath") else None
            if report_path is None or not report_path.exists():
                failures.append("text integrity block did not write a report")
            else:
                report = _read_json(report_path)
                issue_codes = {
                    str(issue.get("code", ""))
                    for issue in report.get("issues", [])
                    if isinstance(issue, dict)
                }
                if "docx_exported_text_changed" not in issue_codes:
                    failures.append(f"text integrity report missed changed-text issue: {sorted(issue_codes)}")
        except Exception as exc:
            failures.append(f"unexpected text integrity block exception: {type(exc).__name__}: {exc}")
            export_failure = {}
    finally:
        app_service.rebuild_docx_from_body_map_units = original_rebuild
        try:
            blocked_export_path.unlink(missing_ok=True)
        except OSError:
            pass

    return {
        "ok": not failures,
        "failures": failures,
        "stage": export_failure.get("stage", ""),
        "reportPath": str(export_failure.get("reportPath", "")),
    }


def _audit_snapshot_protection_scope(snapshot: Any | None, *, strict_sample_expectations: bool) -> dict[str, Any]:
    if snapshot is None:
        return {"ok": False, "issues": [{"type": "missing_snapshot"}]}

    issues: list[dict[str, Any]] = []
    units = list(snapshot.units)
    for unit in units:
        if not unit.editable:
            continue
        target_kind = str(unit.target.get("kind", ""))
        if target_kind != "paragraph":
            issues.append({"type": "editable_non_paragraph_target", "unitIndex": unit.unit_index, "target": unit.target})
        if _looks_like_heading(unit.text, style_name=unit.style_name, has_numbering=bool(getattr(unit, "has_numbering", False))):
            issues.append({"type": "editable_heading", "unitIndex": unit.unit_index, "sample": unit.text[:80]})
        if _looks_like_keyword_line(unit.text):
            issues.append({"type": "editable_structured_keyword", "unitIndex": unit.unit_index, "sample": unit.text[:80]})
        if _looks_like_caption(unit.text) or _looks_like_note(unit.text):
            issues.append({"type": "editable_caption_or_note", "unitIndex": unit.unit_index, "sample": unit.text[:80]})
        if _looks_like_references_heading(unit.text):
            issues.append({"type": "editable_references_heading", "unitIndex": unit.unit_index, "sample": unit.text[:80]})
        if _looks_like_formula_paragraph(unit.text, style_name=unit.style_name):
            issues.append({"type": "editable_formula", "unitIndex": unit.unit_index, "sample": unit.text[:80]})

    if not strict_sample_expectations:
        return {
            "ok": not issues,
            "issues": issues,
            "editableUnitIndexes": [unit.unit_index for unit in units if unit.editable],
            "protectedUnitIndexes": [unit.unit_index for unit in units if not unit.editable],
        }

    abstract_heading_index = next(
        (unit.unit_index for unit in units if unit.text.replace(" ", "") == zh(0x6458, 0x8981)),
        None,
    )
    if abstract_heading_index is None:
        issues.append({"type": "missing_abstract_heading"})
    else:
        leaked_before = [unit.unit_index for unit in units if unit.unit_index < abstract_heading_index and unit.editable]
        if leaked_before:
            issues.append({"type": "front_matter_editable", "unitIndexes": leaked_before})

    expected_editable_samples = [
        zh(0x672c, 0x6587, 0x56f4, 0x7ed5),
        "This paper studies",
        "The English abstract boundary paragraph",
        zh(0x70df, 0x53f6, 0x75c5, 0x866b),
        "YOLOv8",
        zh(0x7cfb, 0x7edf, 0x5b9e, 0x73b0, 0x8fb9, 0x754c, 0x6bb5, 0x843d),
        zh(0x7ed3, 0x679c, 0x5206, 0x6790, 0x540e, 0x7684, 0x8fb9, 0x754c, 0x6bb5, 0x843d),
        zh(0x767b, 0x5f55, 0x6ce8, 0x518c),
        zh(0x767b, 0x5f55, 0x6ce8, 0x518c, 0x6a21, 0x5757),
    ]
    for sample in expected_editable_samples:
        if not any(sample in unit.text and unit.editable for unit in units):
            issues.append({"type": "expected_body_scope_not_editable", "sample": sample})

    expected_protected_samples = [
        zh(0x82f1, 0x6587, 0x6458, 0x8981),
        zh(0x5173, 0x952e, 0x8bcd),
        zh(0x53c2, 0x8003, 0x6587, 0x732e),
        zh(0x6570, 0x636e, 0x4e0e, 0x65b9, 0x6cd5),
        "1.1 " + zh(0x7cfb, 0x7edf, 0x5b9e, 0x73b0),
        "3." + zh(0x7ed3, 0x679c, 0x5206, 0x6790),
        "2." + zh(0x6743, 0x9650, 0x7ba1, 0x7406, 0x529f, 0x80fd, 0x8bf4, 0x660e),
        zh(0x56fe, 0x20) + "1",
        zh(0x6ce8, 0xff1a),
        zh(0x611f, 0x8c22, 0x5bfc, 0x5e08),
        zh(0x81f4, 0x8c22, 0x7b2c, 0x4e8c, 0x6bb5),
        zh(0x81f4, 0x8c22, 0x4e4b, 0x540e, 0x7684, 0x53c2, 0x8003, 0x6587, 0x732e),
        "post-ack protected page",
        zh(0x9644, 0x5f55) + "A " + zh(0x7cfb, 0x7edf, 0x622a, 0x56fe),
        zh(0x8bda, 0x4fe1, 0x58f0, 0x660e, 0x9875),
    ]
    for sample in expected_protected_samples:
        if not any(sample in unit.text and not unit.editable for unit in units):
            issues.append({"type": "expected_protected_scope_editable", "sample": sample})

    leaked_after_ack = [
        {"unitIndex": unit.unit_index, "text": unit.text[:80], "reason": unit.protect_reason}
        for unit in units
        if (zh(0x9644, 0x5f55) in unit.text or zh(0x8bda, 0x4fe1, 0x58f0, 0x660e) in unit.text) and unit.editable
    ]
    if leaked_after_ack:
        issues.append({"type": "post_acknowledgement_editable", "units": leaked_after_ack})

    ack_heading_position = next(
        (index for index, unit in enumerate(units) if _looks_like_acknowledgement_heading(unit.text)),
        None,
    )
    if ack_heading_position is not None:
        post_ack_boundary_position = next(
            (
                index
                for index, unit in enumerate(units[ack_heading_position + 1 :], start=ack_heading_position + 1)
                if _looks_like_references_heading(unit.text) or _looks_like_back_matter_heading(unit.text)
            ),
            None,
        )
        if post_ack_boundary_position is not None:
            leaked_tail = [
                {"unitIndex": unit.unit_index, "text": unit.text[:80], "reason": unit.protect_reason}
                for unit in units[post_ack_boundary_position:]
                if unit.editable
            ]
            if leaked_tail:
                issues.append({"type": "post_acknowledgement_tail_editable", "units": leaked_tail})

    auto_numbered_editable = [
        unit.unit_index
        for unit in units
        if getattr(unit, "has_numbering", False) and zh(0x767b, 0x5f55, 0x6ce8, 0x518c, 0x6a21, 0x5757) in unit.text and unit.editable
    ]
    if not auto_numbered_editable:
        issues.append({"type": "auto_numbered_body_item_not_editable"})

    return {
        "ok": not issues,
        "issues": issues,
        "editableUnitIndexes": [unit.unit_index for unit in units if unit.editable],
        "protectedUnitIndexes": [unit.unit_index for unit in units if not unit.editable],
    }


def _audit_body_map_scope_contract(body_map: Any | None, *, source_path: Path, snapshot_path: Path) -> dict[str, Any]:
    if body_map is None:
        return {"ok": False, "issues": [{"type": "missing_body_map"}]}

    issues: list[dict[str, Any]] = []
    scope_signature = body_map.scope_signature if isinstance(getattr(body_map, "scope_signature", None), dict) else {}
    if not str(scope_signature.get("fingerprint", "")).strip():
        issues.append({"type": "missing_scope_fingerprint"})

    validation_report = validate_docx_body_map(body_map, source_path=source_path, snapshot_path=snapshot_path)
    if not bool(validation_report.get("ok")):
        issues.append({"type": "body_map_validation_failed", "blockingIssues": validation_report.get("blockingIssues", [])})

    tamper_detected = True
    if len(body_map.units) >= 2:
        tampered_units = [body_map.units[1], body_map.units[0], *body_map.units[2:]]
        tampered = replace(body_map, units=tampered_units)
        tampered_report = validate_docx_body_map(tampered, source_path=source_path, snapshot_path=snapshot_path)
        tampered_codes = {
            str(issue.get("code", ""))
            for issue in tampered_report.get("blockingIssues", [])
            if isinstance(issue, dict)
        }
        tamper_detected = "body_map_scope_signature_mismatch" in tampered_codes
    if not tamper_detected:
        issues.append({"type": "tampered_scope_signature_not_blocked"})

    return {
        "ok": not issues,
        "issues": issues,
        "fingerprint": str(scope_signature.get("fingerprint", "")),
        "editableUnitCount": len(body_map.units),
        "validationWarnings": validation_report.get("warnings", []) if isinstance(validation_report, dict) else [],
    }


def _read_json(path: str | Path | None) -> dict[str, Any]:
    if not path:
        return {}
    json_path = Path(path)
    if not json_path.exists():
        return {}
    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return data if isinstance(data, dict) else {}


def _audit_certified_export_bundle(export_result: dict[str, Any], latest_alias_path: Path) -> dict[str, Any]:
    failures: list[str] = []
    certified_path = Path(str(export_result.get("path", "") or ""))
    attempt_id = str(export_result.get("exportAttemptId", "") or "")
    artifact_sha256 = str(export_result.get("artifactSha256", "") or "")
    manifest_path = Path(str(export_result.get("evidenceManifestPath", "") or ""))
    if not certified_path.exists() or certified_path == latest_alias_path:
        failures.append("certified export must use an existing immutable attempt path distinct from the latest alias")
    if not attempt_id or attempt_id not in certified_path.name:
        failures.append("certified export path is not bound to its attempt id")
    actual_sha256 = hashlib.sha256(certified_path.read_bytes()).hexdigest() if certified_path.exists() else ""
    if not artifact_sha256 or artifact_sha256 != actual_sha256:
        failures.append("certified export SHA-256 does not match the immutable DOCX bytes")
    if not latest_alias_path.exists() or (certified_path.exists() and latest_alias_path.read_bytes() != certified_path.read_bytes()):
        failures.append("latest compatibility alias does not match the certified attempt bytes")

    manifest = _read_json(manifest_path)
    if not manifest:
        failures.append("certified export evidence manifest is missing")
    else:
        if manifest.get("status") != "passed":
            failures.append("certified export manifest status is not passed")
        if manifest.get("exportAttemptId") != attempt_id:
            failures.append("certified export manifest attempt id mismatch")
        if manifest.get("artifactPath") != str(certified_path.resolve()):
            failures.append("certified export manifest path mismatch")
        if manifest.get("artifactSha256") != artifact_sha256:
            failures.append("certified export manifest hash mismatch")
        reports = manifest.get("reports") if isinstance(manifest.get("reports"), dict) else {}
        if not reports:
            failures.append("certified export manifest omitted reports")
        for key, raw_report_path in reports.items():
            report_path = Path(str(raw_report_path))
            report = _read_json(report_path)
            if not report:
                failures.append(f"certified export report missing or invalid: {key}")
                continue
            if report.get("publishedArtifactPath") != str(certified_path.resolve()):
                failures.append(f"certified export report path mismatch: {key}")
            if report.get("artifactSha256") != artifact_sha256:
                failures.append(f"certified export report hash mismatch: {key}")
    return {
        "ok": not failures,
        "failures": failures,
        "certifiedPath": str(certified_path),
        "latestAliasPath": str(latest_alias_path),
        "attemptId": attempt_id,
        "artifactSha256": artifact_sha256,
        "manifestPath": str(manifest_path),
    }


def _chunk_text(chunk: dict[str, Any]) -> str:
    return str(chunk.get("outputText") or chunk.get("inputText") or "").strip()


def _find_auto_numbered_body_chunk(compare_payload: dict[str, Any]) -> dict[str, Any] | None:
    chunks = compare_payload.get("chunks")
    if not isinstance(chunks, list):
        return None
    auto_numbered_chunks = [
        chunk
        for chunk in chunks
        if isinstance(chunk, dict) and AUTO_NUMBERED_BODY_RE.search(_chunk_text(chunk))
    ]
    if not auto_numbered_chunks:
        return None
    return max(auto_numbered_chunks, key=lambda chunk: len(_chunk_text(chunk)))


def _build_auto_numbered_rerun_feedback(chunk: dict[str, Any]) -> str:
    excerpt = " ".join(_chunk_text(chunk).split())[:360]
    return (
        "AUTO_NUMBERED_RERUN_USER_FEEDBACK: rewrite this auto-numbered body paragraph.\n"
        f"Excerpt: {excerpt}\n"
        "Preserve the numbering prefix, facts, citations, and document formatting."
    )


def _run_tampered_targeted_rerun_contract_smoke(output_path: Path) -> dict[str, Any]:
    failures: list[str] = []
    compare_path = app_service._find_compare_path_for_output(output_path)
    original_compare = compare_path.read_bytes()
    compare_payload = json.loads(original_compare.decode("utf-8"))
    chunks = compare_payload.get("chunks")
    if not isinstance(chunks, list) or not chunks or not isinstance(chunks[0], dict):
        return {"ok": False, "failures": ["no compare chunk was available for contract tamper smoke"]}
    target_chunk = chunks[0]
    target_chunk_id = str(target_chunk.get("chunkId", "") or "")
    target_chunk["inputText"] = "论文标题不允许进入定点重跑模型输入"
    compare_path.write_text(json.dumps(compare_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    builder_calls = 0
    model_calls = 0
    original_builder = app_service._build_transform_from_model_config

    def fake_builder(_model_config: dict[str, Any]):
        nonlocal builder_calls
        builder_calls += 1

        def should_never_run(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
            nonlocal model_calls
            model_calls += 1
            return chunk_text

        return should_never_run, "online"

    blocked_message = ""
    try:
        app_service._build_transform_from_model_config = fake_builder
        try:
            app_service.rerun_compare_chunk(
                str(output_path),
                target_chunk_id,
                {"baseUrl": "http://localhost", "apiKey": "smoke", "model": "smoke-model"},
            )
        except ValueError as exc:
            blocked_message = str(exc)
        else:
            failures.append("tampered DOCX targeted rerun was not blocked")
    finally:
        app_service._build_transform_from_model_config = original_builder
        compare_path.write_bytes(original_compare)

    if "冻结 chunk manifest" not in blocked_message:
        failures.append(f"unexpected targeted-rerun contract block: {blocked_message}")
    if builder_calls != 0 or model_calls != 0:
        failures.append(f"targeted-rerun contract must block before model setup/call: builder={builder_calls}, model={model_calls}")

    missing_body_map_builder_calls = 0
    missing_body_map_model_calls = 0
    missing_body_map_blocked_message = ""
    body_map_path = app_service._find_body_map_path_for_output(output_path)
    if body_map_path is None or not body_map_path.exists():
        failures.append("DOCX round did not expose a body map for the missing-map contract smoke")
    else:
        original_body_map = body_map_path.read_bytes()

        def missing_map_builder(_model_config: dict[str, Any]):
            nonlocal missing_body_map_builder_calls
            missing_body_map_builder_calls += 1

            def should_never_run(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
                nonlocal missing_body_map_model_calls
                missing_body_map_model_calls += 1
                return chunk_text

            return should_never_run, "online"

        try:
            body_map_path.unlink()
            app_service._build_transform_from_model_config = missing_map_builder
            try:
                app_service.rerun_compare_chunk(
                    str(output_path),
                    target_chunk_id,
                    {"baseUrl": "http://localhost", "apiKey": "smoke", "model": "smoke-model"},
                )
            except ValueError as exc:
                missing_body_map_blocked_message = str(exc)
            else:
                failures.append("DOCX targeted rerun without its frozen body map was not blocked")
        finally:
            app_service._build_transform_from_model_config = original_builder
            body_map_path.write_bytes(original_body_map)

        if "冻结 body map 缺失" not in missing_body_map_blocked_message:
            failures.append(f"unexpected missing-body-map contract block: {missing_body_map_blocked_message}")
        if missing_body_map_builder_calls != 0 or missing_body_map_model_calls != 0:
            failures.append(
                "missing body-map contract must block before model setup/call: "
                f"builder={missing_body_map_builder_calls}, model={missing_body_map_model_calls}"
            )
    return {
        "ok": not failures,
        "failures": failures,
        "targetChunkId": target_chunk_id,
        "builderCalls": builder_calls,
        "modelCalls": model_calls,
        "blockedMessage": blocked_message,
        "missingBodyMap": {
            "builderCalls": missing_body_map_builder_calls,
            "modelCalls": missing_body_map_model_calls,
            "blockedMessage": missing_body_map_blocked_message,
        },
    }


def _run_atomic_docx_export_commit_smoke(output_path: Path, export_path: Path) -> dict[str, Any]:
    failures: list[str] = []
    if not export_path.exists():
        return {"ok": False, "failures": ["successful baseline export is missing"]}

    certified_bytes = export_path.read_bytes()
    staging_pattern = f".{export_path.stem}.*.tmp{export_path.suffix}"
    staging_before = {path.resolve() for path in export_path.parent.glob(staging_pattern)}
    original_audit = app_service.audit_docx_export
    blocked_stage = ""

    def force_post_build_audit_failure(
        _export_path: Path,
        *,
        source_path: Path,
        snapshot_path: Path,
        expected_source_sha256: str | None = None,
        provenance_source_path: Path | None = None,
        report_path: Path,
    ) -> dict[str, Any]:
        return {
            "ok": False,
            "issueCount": 1,
            "issues": [
                {
                    "code": "forced_atomic_commit_regression",
                    "message": "forced post-build audit failure",
                }
            ],
            "sourcePath": str(source_path),
            "sourceSha256": str(expected_source_sha256 or ""),
            "expectedSourceSha256": str(expected_source_sha256 or ""),
            "sourceGenerationStable": True,
            "provenanceSourcePath": str(provenance_source_path or ""),
            "snapshotPath": str(snapshot_path),
            "reportPath": str(report_path),
        }

    try:
        app_service.audit_docx_export = force_post_build_audit_failure
        try:
            app_service.export_round_output(str(output_path), str(export_path), "docx", "school_rules")
        except app_service.ExportRoundError as exc:
            failure = exc.export_failure if isinstance(exc.export_failure, dict) else {}
            blocked_stage = str(failure.get("stage", "") or "")
        else:
            failures.append("forced post-build DOCX audit failure did not block export")
    finally:
        app_service.audit_docx_export = original_audit

    if blocked_stage != "audit":
        failures.append(f"forced atomic export failure stopped at unexpected stage: {blocked_stage}")
    if not export_path.exists() or export_path.read_bytes() != certified_bytes:
        failures.append("failed DOCX export replaced or removed the previous certified file")
    staging_after = {path.resolve() for path in export_path.parent.glob(staging_pattern)}
    leaked_staging = sorted(str(path) for path in staging_after - staging_before)
    if leaked_staging:
        failures.append(f"failed DOCX export leaked staging files: {leaked_staging}")

    return {
        "ok": not failures,
        "failures": failures,
        "blockedStage": blocked_stage,
        "previousCertifiedFilePreserved": export_path.exists() and export_path.read_bytes() == certified_bytes,
        "leakedStagingPaths": leaked_staging,
    }


def _run_auto_numbered_targeted_rerun_smoke(output_path: Path, export_path: Path, snapshot_path: Path) -> dict[str, Any]:
    failures: list[str] = []
    removable_transition = "具体而言，"
    required_body_phrase = "用户通过用户名或手机号创建账户"
    rerun_marker = "登录注册：用户通过"
    compare_payload = app_service.read_round_compare(str(output_path))
    target_chunk = _find_auto_numbered_body_chunk(compare_payload)
    if target_chunk is None:
        return {"ok": False, "failures": ["no auto-numbered body compare chunk was available"]}
    target_chunk_id = str(target_chunk.get("chunkId", ""))
    user_feedback = _build_auto_numbered_rerun_feedback(target_chunk)

    prompts: list[str] = []
    original_builder = app_service._build_transform_from_model_config

    def fake_builder(_model_config: dict[str, Any]):
        def smoke_transform(chunk_text: str, prompt_input: str, _round: int, _chunk_id: str) -> str:
            prompts.append(prompt_input)
            # Exercise a real, minimal style repair by removing only a
            # content-light mechanical transition.  The source therefore has a
            # positive production style penalty and the fact-preserving
            # candidate has a measurable combined gain.
            if removable_transition not in chunk_text:
                raise AssertionError(
                    f"auto-numbered rerun fixture is missing {removable_transition!r}"
                )
            if required_body_phrase not in chunk_text:
                raise AssertionError(
                    f"auto-numbered rerun fixture is missing {required_body_phrase!r}"
                )
            return chunk_text.replace(removable_transition, "", 1)

        return smoke_transform, "online"

    try:
        app_service._build_transform_from_model_config = fake_builder
        rerun_result = app_service.rerun_compare_chunk(
            str(output_path),
            target_chunk_id,
            {"baseUrl": "http://localhost", "apiKey": "smoke", "model": "smoke-model"},
            user_feedback,
        )
    finally:
        app_service._build_transform_from_model_config = original_builder

    rerun_chunk = rerun_result.get("chunk") if isinstance(rerun_result, dict) else {}
    output_text = str(rerun_chunk.get("outputText", "")) if isinstance(rerun_chunk, dict) else ""
    selection = rerun_chunk.get("candidateSelection") if isinstance(rerun_chunk, dict) else None
    selection = selection if isinstance(selection, dict) else {}
    candidates = selection.get("candidates") if isinstance(selection.get("candidates"), list) else []
    baseline_candidate = next(
        (
            candidate
            for candidate in candidates
            if isinstance(candidate, dict) and candidate.get("origin") == "baseline"
        ),
        {},
    )
    selected_candidate = next(
        (
            candidate
            for candidate in candidates
            if isinstance(candidate, dict)
            and str(candidate.get("candidateId", "")) == str(selection.get("selectedCandidateId", ""))
        ),
        {},
    )
    output_sha256 = hashlib.sha256(output_text.encode("utf-8")).hexdigest()
    if not AUTO_NUMBERED_BODY_RE.search(output_text):
        failures.append("targeted rerun did not preserve the auto-numbered paragraph prefix")
    if removable_transition in output_text:
        failures.append("auto-numbered targeted rerun did not remove the measured mechanical transition")
    if rerun_marker not in output_text:
        failures.append("auto-numbered targeted rerun did not publish the minimal transition removal")
    if not prompts or "AUTO_NUMBERED_RERUN_USER_FEEDBACK" not in prompts[0]:
        failures.append("auto-numbered targeted rerun prompt did not include the user feedback")
    if prompts and "[DETECTOR MICRO-REPAIR MODE]" in prompts[0]:
        failures.append("auto-numbered targeted rerun should not enter removed detection-report mode")
    if isinstance(rerun_chunk, dict) and ("rerunCandidateCount" in rerun_chunk or "rerunSelectedCandidate" in rerun_chunk):
        failures.append("auto-numbered targeted rerun should not emit legacy candidate metadata")
    if isinstance(rerun_chunk, dict) and "rerunDetectorProfile" in rerun_chunk:
        failures.append("auto-numbered targeted rerun should not persist removed detection-report metadata")
    if selection.get("schema") != "fyadr.chunk-candidate-selection" or selection.get("schemaVersion") != 2:
        failures.append("auto-numbered targeted rerun did not persist the production candidate-selection schema")
    if selection.get("decision") != "generated_selected":
        failures.append(f"auto-numbered production selector did not publish the repaired candidate: {selection.get('decision')}")
    if selection.get("publishedRewrite") is not True or selection.get("selectedOrigin") != "model":
        failures.append("auto-numbered production selector evidence does not identify a published model rewrite")
    if selection.get("runFailed") is not False:
        failures.append("auto-numbered production selector incorrectly marked the published rerun as failed")
    if "combined_style_penalty_improved" not in list(selection.get("reasonCodes") or []):
        failures.append("auto-numbered production selector did not record a measurable combined style gain")
    if len(prompts) != 1 or selection.get("modelAttemptCount") != 1 or selection.get("conditionalRetryCount") != 0:
        failures.append(
            "auto-numbered measurable-gain candidate should publish on the first bounded attempt: "
            f"prompts={len(prompts)}, modelAttempts={selection.get('modelAttemptCount')}, "
            f"conditionalRetries={selection.get('conditionalRetryCount')}"
        )
    if not baseline_candidate or not selected_candidate:
        failures.append("auto-numbered production selector omitted baseline/selected candidate evidence")
    else:
        try:
            baseline_penalty = float(baseline_candidate.get("stylePenalty"))
            selected_penalty = float(selected_candidate.get("stylePenalty"))
        except (TypeError, ValueError):
            failures.append("auto-numbered production selector emitted invalid style penalties")
        else:
            if selected_penalty > baseline_penalty - 0.05:
                failures.append(
                    "auto-numbered selected candidate did not achieve the required style-penalty gain: "
                    f"baseline={baseline_penalty}, selected={selected_penalty}"
                )
        if selected_candidate.get("hardValid") is not True or selected_candidate.get("safetyEligible") is not True:
            failures.append("auto-numbered selected candidate did not pass the production hard/safety gates")
        if selected_candidate.get("factualGuardPassed") is not True:
            failures.append("auto-numbered selected candidate did not pass the factual-relation guard")
        if selected_candidate.get("readabilityGuardPassed") is not True:
            failures.append("auto-numbered selected candidate did not pass the academic-readability delta guard")
    for hash_field in ("selectedTextSha256", "resultTextSha256", "publishedTextSha256"):
        if selection.get(hash_field) != output_sha256:
            failures.append(f"auto-numbered production selector {hash_field} is not bound to outputText")
    for count_field in ("selectedCharCount", "resultCharCount", "publishedCharCount"):
        if selection.get(count_field) != len(output_text):
            failures.append(f"auto-numbered production selector {count_field} is not bound to outputText")

    post_export_path = export_path.with_name(f"{export_path.stem}_auto_numbered_post_rerun{export_path.suffix}")
    post_export = app_service.export_round_output(str(output_path), str(post_export_path), "docx", "preserve_original")
    post_format_audit = _audit_exported_editable_format(post_export_path, snapshot_path)
    post_body_map = load_docx_body_map(output_path.with_name(f"{output_path.stem}_body_map.json"))
    post_text_integrity = _audit_exported_text_integrity(post_export_path, post_body_map)
    if int(post_export.get("auditIssueCount", 0) or 0) != 0:
        failures.append(f"auto-numbered post-rerun export audit issues: {post_export.get('auditIssueCount')}")
    if int(post_export.get("ooxmlAuditIssueCount", 0) or 0) != 0:
        failures.append(f"auto-numbered post-rerun export OOXML audit issues: {post_export.get('ooxmlAuditIssueCount')}")
    if str(post_export.get("formatMode", "")) != "preserve_original":
        failures.append(f"auto-numbered post-rerun format mode escaped fidelity lock: {post_export.get('formatMode')}")
    if int(post_export.get("formatLockIssueCount", 0) or 0) != 0:
        failures.append(f"auto-numbered post-rerun format-lock issues: {post_export.get('formatLockIssueCount')}")
    if not bool(post_export.get("contentContractReady")):
        failures.append("auto-numbered post-rerun content contract did not pass")
    if int(post_export.get("editableHeadingCount", 0) or 0) != 0:
        failures.append(f"auto-numbered post-rerun exposed editable headings: {post_export.get('editableHeadingCount')}")
    if not bool(post_text_integrity.get("ok")):
        failures.append(f"auto-numbered post-rerun text integrity issues: {len(post_text_integrity.get('issues', []) or [])}")
    exported_text = "\n".join(paragraph.text for paragraph in Document(str(post_export_path.resolve())).paragraphs)
    if rerun_marker not in exported_text:
        failures.append("auto-numbered post-rerun export did not contain the rerun output text")

    return {
        "ok": not failures,
        "failures": failures,
        "targetChunkId": target_chunk_id,
        "targetText": _chunk_text(target_chunk),
        "promptContainsUserFeedback": bool(prompts and "AUTO_NUMBERED_RERUN_USER_FEEDBACK" in prompts[0]),
        "prefixPreserved": bool(AUTO_NUMBERED_BODY_RE.search(output_text)),
        "candidateSelection": selection,
        "measuredTransitionRemoved": removable_transition not in output_text,
        "postRerunExport": post_export,
        "postRerunFormatAudit": post_format_audit,
        "postRerunTextIntegrity": post_text_integrity,
        "postRerunOoxmlAuditIssueCount": int(post_export.get("ooxmlAuditIssueCount", 0) or 0),
        "postRerunMarkerExported": rerun_marker in exported_text,
    }


def run_regression(
    sample_path: Path,
    export_path: Path,
    report_path: Path,
    *,
    rebuild_sample: bool,
    strict_sample_scope: bool = True,
) -> dict[str, Any]:
    if rebuild_sample or not sample_path.exists():
        create_regression_sample(sample_path)

    auto_numbered_rerun = {"ok": True, "skipped": True, "reason": "strict sample scope is disabled"}
    tampered_targeted_rerun = {"ok": True, "skipped": True, "reason": "strict sample scope is disabled"}
    atomic_export_commit = {"ok": True, "skipped": True, "reason": "strict sample scope is disabled"}
    certified_export_bundle = {"ok": True, "skipped": True, "reason": "export not completed"}
    export_text_integrity_block = {"ok": True, "skipped": True, "reason": "strict sample scope is disabled"}
    body_map = None
    round_result = run_document_round(sample_path, identity_transform, round_number=1, prompt_profile="cn")
    output_path = Path(str(round_result["output_path"]))
    export_result = export_round_output(str(output_path), str(export_path), "docx", "preserve_original")
    certified_export_bundle = _audit_certified_export_bundle(export_result, export_path)
    body_map = load_docx_body_map(Path(str(round_result.get("body_map_path", ""))))
    if strict_sample_scope:
        atomic_export_commit = _run_atomic_docx_export_commit_smoke(output_path, export_path)
        export_text_integrity_block = _run_export_text_integrity_block_regression(output_path, export_path)
        tampered_targeted_rerun = _run_tampered_targeted_rerun_contract_smoke(output_path)
        auto_numbered_rerun = _run_auto_numbered_targeted_rerun_smoke(
            output_path,
            export_path,
            get_docx_snapshot_path(sample_path),
        )

    snapshot_path = get_docx_snapshot_path(sample_path)
    snapshot = _load_docx_snapshot(snapshot_path)
    scope_diagnostics_path = get_docx_scope_diagnostics_path(sample_path)
    scope_diagnostics = _read_json(scope_diagnostics_path)
    if body_map is None:
        body_map = load_docx_body_map(Path(str(round_result.get("body_map_path", ""))))
    audit_report = _read_json(export_result.get("auditPath"))
    if not audit_report:
        audit_report = audit_docx_export(
            export_path,
            source_path=sample_path,
            snapshot_path=snapshot_path,
            report_path=get_docx_audit_report_path(export_path),
        )
    ooxml_audit_report = _read_json(export_result.get("ooxmlAuditPath"))
    if not ooxml_audit_report:
        ooxml_audit_report = audit_docx_ooxml_integrity(
            export_path,
            source_path=sample_path,
            snapshot_path=snapshot_path,
            report_path=get_docx_ooxml_audit_report_path(export_path),
        )
    format_audit = _audit_exported_editable_format(export_path, snapshot_path)
    exported_text_integrity = _audit_exported_text_integrity(export_path, body_map)
    protection_scope_audit = _audit_snapshot_protection_scope(snapshot, strict_sample_expectations=strict_sample_scope)
    body_map_scope_contract = _audit_body_map_scope_contract(body_map, source_path=sample_path, snapshot_path=snapshot_path)
    english_run_spacing = _run_english_run_spacing_regression()
    sentence_surface_integrity = _run_sentence_surface_integrity_regression()
    acknowledgement_boundary = _run_acknowledgement_boundary_regression()
    compare_fallback_spacing = _run_compare_fallback_spacing_regression()
    compare_payload_integrity = _run_compare_payload_integrity_regression(
        source_path=sample_path,
        snapshot_path=snapshot_path,
        body_map=body_map,
    )

    failures: list[str] = []
    failures.extend(f"english run spacing: {failure}" for failure in english_run_spacing.get("failures", []) or [])
    failures.extend(f"sentence surface integrity: {failure}" for failure in sentence_surface_integrity.get("failures", []) or [])
    failures.extend(f"acknowledgement boundary: {failure}" for failure in acknowledgement_boundary.get("failures", []) or [])
    failures.extend(f"compare fallback spacing: {failure}" for failure in compare_fallback_spacing.get("failures", []) or [])
    failures.extend(f"compare payload integrity: {failure}" for failure in compare_payload_integrity.get("failures", []) or [])
    failures.extend(f"export text integrity block: {failure}" for failure in export_text_integrity_block.get("failures", []) or [])
    if export_result.get("layoutMode") != "body-map-roundtrip":
        failures.append(f"unexpected layout mode: {export_result.get('layoutMode')}")
    if int(export_result.get("auditIssueCount", 0) or 0) != 0:
        failures.append(f"audit issues: {export_result.get('auditIssueCount')}")
    if int(export_result.get("ooxmlAuditIssueCount", 0) or 0) != 0:
        failures.append(f"OOXML audit issues: {export_result.get('ooxmlAuditIssueCount')}")
    if not bool(ooxml_audit_report.get("ok")):
        failures.append(f"OOXML audit report issues: {ooxml_audit_report.get('issueCount')}")
    if not str(export_result.get("ooxmlAuditPath", "")).strip():
        failures.append("missing OOXML audit path")
    elif not Path(str(export_result.get("ooxmlAuditPath"))).exists():
        failures.append("OOXML audit report file does not exist")
    if snapshot is None:
        failures.append("missing docx snapshot")
    elif snapshot.editable_unit_count <= 0:
        failures.append("snapshot has no editable units")
    if body_map is None:
        failures.append("missing body map")
    elif snapshot is not None and len(body_map.units) != snapshot.editable_unit_count:
        failures.append(f"body map count mismatch: {len(body_map.units)} != {snapshot.editable_unit_count}")
    # ``format_audit`` is a source-relative diagnostic for editable paragraph
    # styling. Fidelity is asserted by the dedicated format-lock contract.
    if str(export_result.get("formatMode", "")) != "preserve_original":
        failures.append(f"unexpected product format mode: {export_result.get('formatMode')}")
    if int(export_result.get("evidenceVersion", 0) or 0) != 1:
        failures.append(f"unexpected export evidence version: {export_result.get('evidenceVersion')}")
    if str(export_result.get("overallStatus", "")) != "passed":
        failures.append(f"unexpected export evidence status: {export_result.get('overallStatus')}")
    if str(export_result.get("sourceKind", "")) != "original_docx":
        failures.append(f"unexpected export source kind: {export_result.get('sourceKind')}")
    if str(export_result.get("contentContractStatus", "")) != "passed":
        failures.append(f"content-contract evidence was not explicit: {export_result.get('contentContractStatus')}")
    if str(export_result.get("formatLockStatus", "")) != "passed":
        failures.append(f"format-lock evidence was not explicit: {export_result.get('formatLockStatus')}")
    performed_checks = {
        str(item)
        for item in (export_result.get("checksPerformed") or [])
        if str(item).strip()
    }
    required_checks = {
        "document_generation",
        "pre_export_guard",
        "content_contract",
        "text_integrity",
        "protected_text_audit",
        "ooxml_integrity",
        "format_lock",
        "post_export_contract",
    }
    if not required_checks.issubset(performed_checks):
        failures.append(f"export evidence omitted performed checks: {sorted(required_checks - performed_checks)}")
    if int(export_result.get("formatLockIssueCount", 0) or 0) != 0:
        failures.append(f"format-lock issues: {export_result.get('formatLockIssueCount')}")
    if not str(export_result.get("formatLockPath", "")).strip():
        failures.append("missing format-lock audit path")
    if not bool(export_result.get("contentContractReady")):
        failures.append("content-only contract did not pass")
    if int(export_result.get("contentContractIssueCount", 0) or 0) != 0:
        failures.append(f"content-contract issues: {export_result.get('contentContractIssueCount')}")
    if int(export_result.get("editableHeadingCount", 0) or 0) != 0:
        failures.append(f"editable headings leaked into contract: {export_result.get('editableHeadingCount')}")
    if not bool(export_result.get("modelInputMatchesEditableUnits")):
        failures.append("model input does not match frozen editable units")
    if not bool(exported_text_integrity.get("ok")):
        failures.append(f"exported text integrity issues: {len(exported_text_integrity.get('issues', []) or [])}")
    if not bool(protection_scope_audit.get("ok")):
        failures.append(f"protection scope audit issues: {len(protection_scope_audit.get('issues', []) or [])}")
    if not bool(body_map_scope_contract.get("ok")):
        failures.append(f"body map scope contract issues: {len(body_map_scope_contract.get('issues', []) or [])}")
    if not scope_diagnostics:
        failures.append("missing DOCX scope diagnostics")
    else:
        if not bool(scope_diagnostics.get("ok", False)):
            failures.append(f"scope diagnostics errors: {scope_diagnostics.get('errorCount')}")
        if int(scope_diagnostics.get("editableUnitCount", -1) or -1) != (snapshot.editable_unit_count if snapshot is not None else -2):
            failures.append("scope diagnostics editable count mismatch")
        if strict_sample_scope and int(scope_diagnostics.get("issueCount", 0) or 0) != 0:
            failures.append(f"scope diagnostics unexpected issues: {scope_diagnostics.get('issueCount')}")
        scope_payload = scope_diagnostics.get("scope") if isinstance(scope_diagnostics.get("scope"), dict) else {}
        if strict_sample_scope and scope_payload.get("postAcknowledgementBoundaryIndex") is None:
            failures.append("scope diagnostics missing post-acknowledgement boundary")
    if not bool(auto_numbered_rerun.get("ok", True)):
        failures.extend(f"auto-numbered targeted rerun: {failure}" for failure in auto_numbered_rerun.get("failures", []) or [])
    if not bool(tampered_targeted_rerun.get("ok", True)):
        failures.extend(f"tampered targeted rerun contract: {failure}" for failure in tampered_targeted_rerun.get("failures", []) or [])
    if not bool(atomic_export_commit.get("ok", True)):
        failures.extend(f"atomic DOCX export commit: {failure}" for failure in atomic_export_commit.get("failures", []) or [])
    if not bool(certified_export_bundle.get("ok", True)):
        failures.extend(f"certified DOCX export bundle: {failure}" for failure in certified_export_bundle.get("failures", []) or [])
    for sample_key in ("guardIssueSamples", "auditIssueSamples", "ooxmlAuditIssueSamples"):
        if sample_key not in export_result:
            failures.append(f"missing export issue sample field: {sample_key}")
        elif not isinstance(export_result.get(sample_key), list):
            failures.append(f"export issue sample field must be a list: {sample_key}")

    report = {
        "ok": not failures,
        "createdAt": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "samplePath": str(sample_path.resolve()),
        "outputPath": str(output_path.resolve()),
        "exportPath": str(export_path.resolve()),
        "reportPath": str(report_path.resolve()),
        "failures": failures,
        "round": {
            "round": round_result.get("round"),
            "chunkCount": round_result.get("chunk_count", round_result.get("chunkCount", round_result.get("input_segment_count"))),
            "bodyMapPath": round_result.get("body_map_path", ""),
            "validationPath": round_result.get("validation_path", ""),
            "comparePath": round_result.get("compare_path", ""),
            "qualityPath": round_result.get("quality_path", ""),
        },
        "snapshot": {
            "path": str(snapshot_path.resolve()),
            "editableUnitCount": snapshot.editable_unit_count if snapshot is not None else 0,
            "totalTextUnitCount": snapshot.total_text_unit_count if snapshot is not None else 0,
            "protectedUnitCount": (snapshot.total_text_unit_count - snapshot.editable_unit_count) if snapshot is not None else 0,
        },
        "scopeDiagnostics": {
            "path": str(scope_diagnostics_path.resolve()),
            "ok": bool(scope_diagnostics.get("ok")) if scope_diagnostics else False,
            "issueCount": int(scope_diagnostics.get("issueCount", 0) or 0) if scope_diagnostics else 0,
            "errorCount": int(scope_diagnostics.get("errorCount", 0) or 0) if scope_diagnostics else 0,
            "warningCount": int(scope_diagnostics.get("warningCount", 0) or 0) if scope_diagnostics else 0,
            "scope": scope_diagnostics.get("scope", {}) if isinstance(scope_diagnostics.get("scope"), dict) else {},
            "reasonCounts": scope_diagnostics.get("reasonCounts", {}) if isinstance(scope_diagnostics.get("reasonCounts"), dict) else {},
            "sampleIssues": (scope_diagnostics.get("issues", []) or [])[:5] if isinstance(scope_diagnostics.get("issues"), list) else [],
        },
        "export": export_result,
        "exportIssueSamples": {
            "guard": len(export_result.get("guardIssueSamples", []) or []),
            "audit": len(export_result.get("auditIssueSamples", []) or []),
            "ooxmlAudit": len(export_result.get("ooxmlAuditIssueSamples", []) or []),
        },
        "formatAudit": format_audit,
        "exportedTextIntegrity": exported_text_integrity,
        "protectionScopeAudit": protection_scope_audit,
        "bodyMapScopeContract": body_map_scope_contract,
        "englishRunSpacing": english_run_spacing,
        "sentenceSurfaceIntegrity": sentence_surface_integrity,
        "acknowledgementBoundary": acknowledgement_boundary,
        "compareFallbackSpacing": compare_fallback_spacing,
        "comparePayloadIntegrity": compare_payload_integrity,
        "exportTextIntegrityBlock": export_text_integrity_block,
        "atomicExportCommit": atomic_export_commit,
        "certifiedExportBundle": certified_export_bundle,
        "tamperedTargetedRerun": tampered_targeted_rerun,
        "autoNumberedRerun": auto_numbered_rerun,
        "audit": {
            "ok": bool(audit_report.get("ok")),
            "issueCount": int(audit_report.get("issueCount", 0) or 0),
            "protectedChecked": int(audit_report.get("protectedChecked", 0) or 0),
            "tableIssueCount": len(audit_report.get("tableStructureIssues", []) or []),
        },
        "ooxmlAudit": {
            "ok": bool(ooxml_audit_report.get("ok")),
            "issueCount": int(ooxml_audit_report.get("issueCount", 0) or 0),
            "path": str(export_result.get("ooxmlAuditPath", "")),
        },
    }
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run FYADR DOCX export regression without calling an LLM.")
    parser.add_argument("--sample", type=Path, default=DEFAULT_SAMPLE_PATH)
    parser.add_argument("--export", type=Path, default=DEFAULT_EXPORT_PATH)
    parser.add_argument("--report", type=Path, default=DEFAULT_REPORT_PATH)
    parser.add_argument("--rebuild-sample", action="store_true", help="Recreate the regression sample DOCX before running.")
    args = parser.parse_args(argv)

    report = run_regression(
        args.sample.resolve(),
        args.export.resolve(),
        args.report.resolve(),
        rebuild_sample=args.rebuild_sample,
    )
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
