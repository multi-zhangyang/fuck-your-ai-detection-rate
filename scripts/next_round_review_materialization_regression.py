#!/usr/bin/env python3
"""P0 regression: downstream rounds consume immutable reviewed content."""

from __future__ import annotations

import json
import sys
import tempfile
from contextlib import contextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterator

from docx import Document  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import app_service  # noqa: E402
import web_app  # noqa: E402
from chunking import build_manifest  # noqa: E402
from docx_bodymap import load_docx_body_map, save_docx_body_map, update_docx_body_map_texts  # noqa: E402
from fyadr_round_service import (  # noqa: E402
    _build_candidate_selection_event,
    _build_checkpoint_signature,
    _evaluate_rewrite_candidate,
    _is_checkpoint_compatible,
    _sha256_text,
)
from round_helper import build_round_context, run_document_round  # noqa: E402


PROFILE = "cn_custom"
SEQUENCE = ["prewrite", "round1"]
SOURCES = [
    "第一段说明系统在本地保存原始记录，用户可以随时核对处理结果。",
    "第二段记录队列按顺序处理任务，失败项目会保留具体原因。",
    "第三段说明缓存容量保持不变，实验只比较两种调度方式。",
    "第四段汇总人工复核要求，未经确认的高风险文本应保留原文。",
]
CANDIDATES = [
    "第一段说明系统会在本地保留原始记录，用户能够随时核对处理结果。",
    "第二段记录队列依次处理任务，失败项目仍会保留具体原因。",
    "第三段说明缓存容量维持不变，实验仅比较两种调度方式。",
    "第四段汇总人工复核要求，未经确认的高风险内容应当保留原文。",
]
CUSTOM = "第二段显示任务由队列依次处理；失败项目的具体原因会被完整保留。"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _candidate_transform(_chunk_text: str, _prompt_input: str, _round: int, chunk_id: str) -> str:
    paragraph_index = int(chunk_id.split("_", 1)[0].removeprefix("p"))
    return CANDIDATES[paragraph_index]


def _identity_transform(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
    return chunk_text


def _snapshot_provenance(snapshot: dict[str, Any], parent_output_path: Path) -> dict[str, str]:
    return {
        "parentOutputPath": str(parent_output_path.resolve()),
        "parentCompareRevision": str(snapshot["compareRevision"]),
        "parentContentRevision": str(snapshot["contentRevision"]),
        "parentReviewRevision": str(snapshot["reviewRevision"]),
        "effectiveInputSha256": str(snapshot["effectiveTextSha256"]),
        "materializationSource": "review_materialized_compare",
        "parentArtifactSnapshotDigest": str(snapshot["artifactSnapshotDigest"]),
    }


def _parent_input_binding(snapshot: dict[str, Any]) -> dict[str, str]:
    return {
        "compareRevision": str(snapshot["compareRevision"]),
        "reviewRevision": str(snapshot["reviewRevision"]),
        "contentRevision": str(snapshot["contentRevision"]),
        "artifactSnapshotDigest": str(snapshot["artifactSnapshotDigest"]),
        "effectiveTextSha256": str(snapshot["effectiveTextSha256"]),
    }


def _approval_request_fields(snapshot: dict[str, Any]) -> dict[str, str]:
    binding = _parent_input_binding(snapshot)
    return {
        "expectedPreviousCompareRevision": binding["compareRevision"],
        "expectedPreviousReviewRevision": binding["reviewRevision"],
        "expectedPreviousContentRevision": binding["contentRevision"],
        "expectedPreviousArtifactSnapshotDigest": binding["artifactSnapshotDigest"],
        "expectedPreviousEffectiveTextSha256": binding["effectiveTextSha256"],
    }


def _published_selection(
    input_text: str,
    output_text: str,
    chunk_id: str,
    *,
    global_style_profile: dict[str, object],
) -> dict[str, object]:
    neutral = {"id": "neutral", "primaryMetric": ""}
    baseline = _evaluate_rewrite_candidate(
        input_text=input_text,
        output_text=input_text,
        candidate_id="baseline",
        origin="baseline",
        attempt=0,
        hard_valid=True,
        round_dimension=neutral,
        global_style_profile=global_style_profile,
    )
    candidate = _evaluate_rewrite_candidate(
        input_text=input_text,
        output_text=output_text,
        candidate_id="model-attempt-1",
        origin="model",
        attempt=1,
        hard_valid=True,
        round_dimension=neutral,
        global_style_profile=global_style_profile,
    )
    _assert(candidate.get("safetyEligible") is True, f"fixture {chunk_id} rewrite is not release eligible")
    return _build_candidate_selection_event(
        chunk_id=chunk_id,
        round_number=1,
        candidates=[baseline, candidate],
        selected=candidate,
        reason_codes=["fixture_production_selected"],
        conditional_retry_count=0,
    )


def _rewrite_parent_artifacts(
    round_result: dict[str, Any],
    *,
    docx: bool,
) -> tuple[Path, Path, list[str], dict[str, Any]]:
    output_path = Path(str(round_result["output_path"])).resolve()
    compare_path = Path(str(round_result["compare_path"])).resolve()
    compare = json.loads(compare_path.read_text(encoding="utf-8"))
    chunk_ids = [str(chunk["chunkId"]) for chunk in compare["chunks"]]
    global_style_profile = app_service.build_global_style_profile_from_texts(SOURCES)
    source_pattern_profile = global_style_profile["documentPatternBaseline"]
    compare["sourcePatternProfiles"] = {
        str(source_pattern_profile["profileSha256"]): source_pattern_profile,
    }
    compare["chunks"][3]["quality"] = {
        "needsReview": True,
        "flags": ["machine_like_expression"],
        "advisoryFlags": [],
    }
    certified_chunk = compare["chunks"][2]
    certified_chunk["outputText"] = CANDIDATES[2]
    certified_chunk["outputCharCount"] = len(CANDIDATES[2])
    certified_chunk["candidateSelection"] = _published_selection(
        str(certified_chunk["inputText"]),
        CANDIDATES[2],
        str(certified_chunk["chunkId"]),
        global_style_profile=global_style_profile,
    )
    compare["updatedAt"] = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    compare_path.write_text(json.dumps(compare, ensure_ascii=False, indent=2), encoding="utf-8")

    # Simulate a previously safe-materialized artifact: the confirmed rewrite
    # exists only in compare until the next-round snapshot applies decisions.
    raw_paragraphs = list(CANDIDATES)
    raw_paragraphs[2] = SOURCES[2]
    output_path.write_text("\n\n".join(raw_paragraphs), encoding="utf-8")
    if docx:
        body_map_path = Path(str(round_result["body_map_path"])).resolve()
        body_map = load_docx_body_map(body_map_path)
        _assert(body_map is not None, "DOCX parent body map is unavailable")
        save_docx_body_map(
            update_docx_body_map_texts(body_map, raw_paragraphs, round_number=1),
            body_map_path,
        )

    decisions = {
        chunk_ids[0]: "source_confirmed",
        chunk_ids[1]: {"mode": "custom", "text": CUSTOM, "source": "manual", "confirmed": True},
        chunk_ids[2]: "rewrite_confirmed",
    }
    app_service.save_review_decisions(str(output_path), decisions)
    expected = [SOURCES[0], CUSTOM, CANDIDATES[2], SOURCES[3]]
    snapshot = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
    _assert(snapshot["_internal"]["effectiveParagraphs"] == expected, "fixture did not materialize expected review state")
    return output_path, compare_path, expected, snapshot


@contextmanager
def _txt_parent() -> Iterator[tuple[Path, dict[str, Any], Path, Path, list[str], dict[str, Any]]]:
    origin_root = ROOT_DIR / "origin"
    origin_root.mkdir(parents=True, exist_ok=True)
    context = None
    with tempfile.TemporaryDirectory(prefix="next-round-txt-", dir=origin_root) as temp_dir:
        source_path = Path(temp_dir) / "reviewed.txt"
        source_path.write_text("\n\n".join(SOURCES), encoding="utf-8")
        round1 = run_document_round(
            source_path,
            _candidate_transform,
            round_number=1,
            prompt_profile=PROFILE,
            prompt_sequence=SEQUENCE,
        )
        context = build_round_context(source_path, round_number=2, prompt_profile=PROFILE, prompt_sequence=SEQUENCE)
        output_path, compare_path, expected, snapshot = _rewrite_parent_artifacts(round1, docx=False)
        try:
            yield source_path, round1, output_path, compare_path, expected, snapshot
        finally:
            try:
                app_service.delete_document_history(context.doc_id)
            except Exception:
                pass


@contextmanager
def _docx_parent() -> Iterator[tuple[Path, dict[str, Any], Path, Path, list[str], dict[str, Any]]]:
    origin_root = ROOT_DIR / "origin"
    origin_root.mkdir(parents=True, exist_ok=True)
    context = None
    with tempfile.TemporaryDirectory(prefix="next-round-docx-", dir=origin_root) as temp_dir:
        source_path = Path(temp_dir) / "reviewed.docx"
        document = Document()
        document.add_heading("下一轮输入契约测试", 0)
        document.add_paragraph("摘要")
        for text in SOURCES:
            document.add_paragraph(text)
        document.save(source_path)
        round1 = run_document_round(
            source_path,
            _candidate_transform,
            round_number=1,
            prompt_profile=PROFILE,
            prompt_sequence=SEQUENCE,
        )
        context = build_round_context(source_path, round_number=2, prompt_profile=PROFILE, prompt_sequence=SEQUENCE)
        output_path, compare_path, expected, snapshot = _rewrite_parent_artifacts(round1, docx=True)
        try:
            yield source_path, round1, output_path, compare_path, expected, snapshot
        finally:
            try:
                app_service.delete_document_history(context.doc_id)
            except Exception:
                pass


def _checkpoint_signature(context: Any, snapshot: dict[str, Any]) -> dict[str, object]:
    effective_text = str(snapshot["_internal"]["effectiveText"])
    manifest = build_manifest(effective_text)
    return _build_checkpoint_signature(
        doc_id=context.doc_id,
        round_number=2,
        prompt_profile=PROFILE,
        prompt_sequence=SEQUENCE,
        input_path=context.input_text_path,
        output_path=context.output_text_path,
        manifest_path=context.manifest_path,
        manifest_chunk_ids=[chunk.chunk_id for chunk in manifest.chunks],
        input_sha256=_sha256_text(effective_text),
        prompt_sha256="prompt-sha",
        chunk_limit=1800,
        chunk_metric="char",
        checkpoint_metadata={"parent_input_provenance": _snapshot_provenance(snapshot, context.parent_output_path)},
    )


def run_regression() -> dict[str, Any]:
    checks: list[str] = []

    with _txt_parent() as (source_path, _round1, output_path, _compare_path, expected, snapshot_a):
        context = build_round_context(source_path, round_number=2, prompt_profile=PROFILE, prompt_sequence=SEQUENCE)
        _assert(context.input_text_path != output_path, "TXT round 2 still points at mutable parent output")
        _assert(context.input_text_path.name.endswith("_round2_input.txt"), "TXT round 2 has no immutable input path")

        decisions = snapshot_a["review"]["decisions"]
        app_service.save_review_decisions(str(output_path), decisions)
        snapshot_b = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
        _assert(snapshot_a["effectiveTextSha256"] == snapshot_b["effectiveTextSha256"], "same decisions changed effective text")
        _assert(snapshot_a["compareRevision"] != snapshot_b["compareRevision"], "review save did not advance compare revision")
        _assert(
            not _is_checkpoint_compatible(_checkpoint_signature(context, snapshot_a), _checkpoint_signature(context, snapshot_b)),
            "parent provenance change did not invalidate checkpoint signature",
        )
        checks.append("parent compare/content/review provenance invalidates stale checkpoints even when effective text is unchanged")

        client = web_app.app.test_client()
        model_config = {
            "baseUrl": "http://127.0.0.1:9",
            "apiKey": "regression",
            "model": "regression",
            "promptProfile": PROFILE,
            "promptSequence": SEQUENCE,
        }
        missing = client.post("/api/run-round", json={"sourcePath": str(source_path), "modelConfig": model_config})
        _assert(missing.status_code == 428 and missing.get_json().get("code") == "round_input_revision_required", f"missing revision was not blocked: {missing.status_code} {missing.get_json()}")
        partial = client.post(
            "/api/run-round",
            json={
                "sourcePath": str(source_path),
                "modelConfig": model_config,
                "expectedPreviousCompareRevision": snapshot_b["compareRevision"],
            },
        )
        _assert(
            partial.status_code == 428 and partial.get_json().get("code") == "round_input_revision_required",
            f"partial parent binding was not blocked: {partial.status_code} {partial.get_json()}",
        )
        stale = client.post(
            "/api/run-round",
            json={
                "sourcePath": str(source_path),
                "modelConfig": model_config,
                **_approval_request_fields(snapshot_a),
            },
        )
        _assert(stale.status_code == 409 and stale.get_json().get("code") == "stale_round_input", f"stale HTTP preflight drifted: {stale.status_code} {stale.get_json()}")

        model_calls = 0
        original_completion = app_service.llm_completion

        def should_never_call(*_args: Any, **_kwargs: Any) -> str:
            nonlocal model_calls
            model_calls += 1
            return ""

        app_service.llm_completion = should_never_call
        try:
            try:
                app_service.run_round_for_app(
                    str(source_path),
                    model_config,
                    expected_previous_compare_revision=str(snapshot_a["compareRevision"]),
                    expected_parent_input_binding=_parent_input_binding(snapshot_a),
                )
                raise AssertionError("stale worker revision unexpectedly ran")
            except app_service.StaleRoundInputError:
                pass
            review_path = Path(str(snapshot_b["_internal"]["reviewPath"]))
            review_bytes = review_path.read_bytes()
            try:
                review_path.write_text("{corrupt", encoding="utf-8")
                try:
                    app_service.run_round_for_app(
                        str(source_path),
                        model_config,
                        expected_previous_compare_revision=str(snapshot_b["compareRevision"]),
                        expected_parent_input_binding=_parent_input_binding(snapshot_b),
                    )
                    raise AssertionError("corrupt parent review evidence unexpectedly ran")
                except app_service.StaleRoundInputError as exc:
                    _assert("round_snapshot_review_corrupt" in exc.mismatch_codes, f"corrupt evidence code drifted: {exc.mismatch_codes}")
            finally:
                review_path.write_bytes(review_bytes)
        finally:
            app_service.llm_completion = original_completion
        _assert(model_calls == 0, "stale/corrupt worker path called the model")
        original_thread = web_app.threading.Thread
        captured_worker_args: tuple[Any, ...] | None = None
        worker_construct_count = 0

        class CapturedThread:
            def __init__(self, *, target: Any, args: tuple[Any, ...], daemon: bool) -> None:
                nonlocal captured_worker_args, worker_construct_count
                _assert(target is web_app.run_round_async and daemon is True, "valid run registered an unexpected worker")
                worker_construct_count += 1
                captured_worker_args = args

            def start(self) -> None:
                return None

        web_app.threading.Thread = CapturedThread
        valid_run_id = ""
        snapshot_for_round2 = snapshot_b
        try:
            drift_fields = {
                "expectedPreviousCompareRevision": "parent_compare_revision_mismatch",
                "expectedPreviousReviewRevision": "parent_review_revision_mismatch",
                "expectedPreviousContentRevision": "parent_content_revision_mismatch",
                "expectedPreviousArtifactSnapshotDigest": "parent_artifact_snapshot_digest_mismatch",
                "expectedPreviousEffectiveTextSha256": "parent_effective_text_sha256_mismatch",
            }
            for request_field, mismatch_code in drift_fields.items():
                drifted_binding = _approval_request_fields(snapshot_b)
                drifted_binding[request_field] = f"stale-{drifted_binding[request_field]}"
                drifted = client.post(
                    "/api/run-round",
                    json={
                        "sourcePath": str(source_path),
                        "modelConfig": model_config,
                        **drifted_binding,
                    },
                )
                drifted_payload = drifted.get_json()
                _assert(
                    drifted.status_code == 409
                    and drifted_payload.get("code") == "stale_round_input"
                    and mismatch_code in (drifted_payload.get("mismatchCodes") or []),
                    f"{request_field} drift was not blocked precisely: {drifted.status_code} {drifted_payload}",
                )
            _assert(worker_construct_count == 0, "stale approval constructed a downstream worker")

            valid = client.post(
                "/api/run-round",
                json={
                    "sourcePath": str(source_path),
                    "modelConfig": model_config,
                    **_approval_request_fields(snapshot_b),
                },
            )
            valid_payload = valid.get_json()
            _assert(valid.status_code == 202 and valid_payload.get("alreadyActive") is False, f"valid bound run was rejected: {valid.status_code} {valid_payload}")
            valid_run_id = str(valid_payload.get("runId", ""))
            _assert(worker_construct_count == 1, "valid approval did not construct exactly one worker")
            _assert(
                captured_worker_args is not None
                and captured_worker_args[3] == snapshot_b["compareRevision"]
                and captured_worker_args[4] == _parent_input_binding(snapshot_b),
                "worker lost the complete bound parent generation",
            )

            # Simulate the exact HTTP-202 -> worker-start race. A same-decision
            # save advances review/content/artifact revisions while leaving the
            # effective text unchanged. The worker must fresh-read and reject
            # before the first provider call.
            app_service.save_review_decisions(str(output_path), snapshot_b["review"]["decisions"])
            snapshot_c = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
            _assert(
                snapshot_c["effectiveTextSha256"] == snapshot_b["effectiveTextSha256"]
                and snapshot_c["artifactSnapshotDigest"] != snapshot_b["artifactSnapshotDigest"],
                "worker-race fixture did not advance the parent generation",
            )
            snapshot_for_round2 = snapshot_c
            app_service.llm_completion = should_never_call
            try:
                web_app.run_round_async(*captured_worker_args)
            finally:
                app_service.llm_completion = original_completion
            worker_state = web_app.RUN_STATES.get(valid_run_id)
            _assert(
                worker_state is not None
                and worker_state.completed
                and worker_state.status == "failed"
                and "approval no longer matches" in str(worker_state.error or ""),
                "stale bound worker did not fail closed after HTTP 202",
            )
            _assert(model_calls == 0, "HTTP-to-worker parent drift called the model")
        finally:
            web_app.threading.Thread = original_thread
            if valid_run_id:
                web_app.RUN_STATES.pop(valid_run_id, None)
                web_app.ACTIVE_RUNS_BY_SOURCE.pop(str(source_path.resolve()), None)
                web_app.run_round_state_path(valid_run_id).unlink(missing_ok=True)
        checks.append(
            "round>1 approval binds compare/review/content/artifact/effective hashes; every drift and the HTTP-to-worker race make zero worker/model calls"
        )

        round2 = run_document_round(
            source_path,
            _identity_transform,
            round_number=2,
            prompt_profile=PROFILE,
            prompt_sequence=SEQUENCE,
            parent_artifact_snapshot=snapshot_for_round2,
        )
        input_path = Path(str(round2["round_context"]["input_text_path"]))
        _assert(input_path.read_text(encoding="utf-8") == "\n\n".join(expected), "TXT immutable input is not review-materialized")
        compare2 = json.loads(Path(str(round2["compare_path"])).read_text(encoding="utf-8"))
        _assert([chunk["inputText"] for chunk in compare2["chunks"]] == expected, "TXT round2 compare inputs ignore decisions/defaults")
        for key in (
            "parentOutputPath",
            "parentCompareRevision",
            "parentContentRevision",
            "parentReviewRevision",
            "effectiveInputSha256",
            "materializationSource",
        ):
            _assert(compare2.get(key) == round2["input_provenance"].get(key), f"TXT compare lost {key}")
        checks.append("TXT source/custom/rewrite/default-source choices become an immutable round2 input with provenance")

    with _docx_parent() as (source_path, round1, output_path, _compare_path, expected, snapshot):
        parent_body_map = load_docx_body_map(Path(str(round1["body_map_path"])))
        _assert(parent_body_map is not None, "DOCX parent body map disappeared")
        source_bytes = source_path.read_bytes()
        round2 = run_document_round(
            source_path,
            _identity_transform,
            round_number=2,
            prompt_profile=PROFILE,
            prompt_sequence=SEQUENCE,
            parent_artifact_snapshot=snapshot,
        )
        input_path = Path(str(round2["round_context"]["input_text_path"]))
        _assert(input_path != output_path and input_path.read_text(encoding="utf-8") == "\n\n".join(expected), "DOCX immutable input drifted")
        child_body_map = load_docx_body_map(Path(str(round2["body_map_path"])))
        _assert(child_body_map is not None and child_body_map.current_texts() == expected, "DOCX round2 body map ignored effective paragraphs")
        _assert(child_body_map.source_sha256 == parent_body_map.source_sha256, "DOCX source identity changed")
        _assert(child_body_map.scope_signature == parent_body_map.scope_signature, "DOCX scope signature changed")
        _assert(
            [unit.target for unit in child_body_map.units] == [unit.target for unit in parent_body_map.units],
            "DOCX frozen targets changed",
        )
        _assert(source_path.read_bytes() == source_bytes, "DOCX source bytes changed")
        _assert(bool(round2.get("edit_contract", {}).get("ready")), "DOCX post-round edit contract is not ready")
        _assert("下一轮输入契约测试" not in input_path.read_text(encoding="utf-8"), "protected title entered model input")
        checks.append("DOCX snapshot body-map clone/retag applies only effective body paragraphs and preserves source/scope/targets")

    return {"ok": True, "checks": checks}


def main() -> int:
    report = run_regression()
    report_path = ROOT_DIR / "finish" / "regression" / "next_round_review_materialization_regression_report.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
