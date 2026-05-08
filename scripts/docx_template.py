from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document  # type: ignore[import]
from docx.document import Document as DocxDocument  # type: ignore[import]
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # type: ignore[import]
from docx.oxml import OxmlElement  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]
from docx.shared import Cm, Pt  # type: ignore[import]
from docx.styles.style import BaseStyle  # type: ignore[import]
from docx.text.paragraph import Paragraph  # type: ignore[import]
from docx.text.run import Run  # type: ignore[import]

from docx_pipeline import (
    ACKNOWLEDGEMENT_MARKERS,
    ABSTRACT_MARKERS,
    INTERMEDIATE_DIR,
    KEYWORD_PREFIXES,
    ROOT_DIR,
    _iter_block_items,
    _load_docx_snapshot,
    _looks_like_back_matter_heading,
    _looks_like_caption,
    _looks_like_formula_paragraph,
    _looks_like_numbered_body_item,
    _looks_like_references_heading,
    _looks_like_toc_heading,
    _normalize_marker_text,
    _paragraph_has_field_code,
)
from format_rules import PAGE_FORMAT_FIELDS, STYLE_FORMAT_FIELDS, load_active_format_rules


DOCX_TEMPLATE_PROFILE_VERSION = 5
FONT_CN_SONG = "宋体"
FONT_CN_HEI = "黑体"
FONT_CN_KAI = "楷体"
FONT_EN_TIMES = "Times New Roman"
SIZE_XIAO_2 = 18
SIZE_3 = 16
SIZE_XIAO_3 = 15
SIZE_4 = 14
SIZE_XIAO_4 = 12
SIZE_5 = 10.5
SIZE_XIAO_5 = 9
SCHOOL_PAGE_MARGINS_CM = {
    "top_margin": 2.5,
    "bottom_margin": 2.5,
    "left_margin": 3.0,
    "right_margin": 3.0,
}
ALIGNMENT_BY_NAME = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
CJK_CHAR_RE = re.compile(r"[\u3400-\u4DBF\u4E00-\u9FFF]")
LATIN_CHAR_RE = re.compile(r"[A-Za-z]")
LEVEL_1_HEADING_RE = re.compile(r"^(?:第[一二三四五六七八九十百零0-9]+章\s*.+|[1-9]\d*\s+\S.+)$")
LEVEL_2_HEADING_RE = re.compile(r"^[1-9]\d*\.[1-9]\d*\s+\S.+$")
LEVEL_3_HEADING_RE = re.compile(r"^[1-9]\d*\.[1-9]\d*\.[1-9]\d*\s+\S.+$")
LEVEL_4_HEADING_RE = re.compile(r"^[1-9]\d*\.[1-9]\d*\.[1-9]\d*\.[1-9]\d*\s+\S.+$")
EN_TITLE_RE = re.compile(r"^[A-Za-z][A-Za-z0-9 ,;:()'\"./&-]{18,}$")
EN_KEYWORDS_PREFIXES = ("keywords", "keyword", "keywords:", "keyword:", "keywords：", "keyword：", "keywords ", "key words")
HEADING_TERMINAL_PUNCTUATION_RE = re.compile(r"[\u3002\uff1b;\uff0c,\u3001\uff1a:]$")
HEADING_BODY_PUNCTUATION_RE = re.compile(r"[\u3002\uff1b;\uff0c,\u3001\uff1a:]")
MIXED_LABEL_ROLES = {"cn_abstract_lead", "en_abstract_lead", "cn_keywords", "en_keywords"}
ROLE_FONT_FALLBACKS = {
    "cn_abstract_lead": ("cn_abstract_body", "body_text"),
    "en_abstract_lead": ("en_abstract_body", "body_text"),
    "cn_keywords": ("cn_abstract_body", "body_text"),
    "en_keywords": ("en_abstract_body", "body_text"),
    "ack_body": ("body_text",),
    "numbered_body": ("body_text",),
}
STRUCTURAL_FORMAT_ROLES = {"toc_heading", "references_heading", "references_body", "ack_heading"}
CONTENT_LOCKED_FORMAT_ROLES = {
    "toc_heading",
    "cn_abstract_lead",
    "en_abstract_lead",
    "cn_keywords",
    "en_keywords",
    "heading_1",
    "heading_2",
    "heading_3",
    "heading_4",
    "caption",
    "note",
    "references_heading",
    "references_body",
    "ack_heading",
    "ack_body",
}
DOCX_TEMPLATE_APPLY_MODES = {"safe_template_fill", "full_template_profile"}


@dataclass
class TemplateParagraphProfile:
    role: str
    style_name: str
    source_text: str
    alignment: int | None = None
    first_line_indent_pt: float | None = None
    left_indent_pt: float | None = None
    right_indent_pt: float | None = None
    space_before_pt: float | None = None
    space_after_pt: float | None = None
    line_spacing_rule: int | None = None
    line_spacing_pt: float | None = None
    line_spacing_multiple: float | None = None
    cn_font: str | None = None
    en_font: str | None = None
    font_size_pt: float | None = None
    bold: bool | None = None
    italic: bool | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class DocxTemplateProfile:
    version: int
    template_path: str
    template_size: int
    template_mtime_ns: int
    start_marker_index: int
    first_body_index: int
    references_index: int | None
    acknowledgement_index: int | None
    roles: dict[str, TemplateParagraphProfile]

    def to_dict(self) -> dict[str, Any]:
        return {
            "version": self.version,
            "template_path": self.template_path,
            "template_size": self.template_size,
            "template_mtime_ns": self.template_mtime_ns,
            "start_marker_index": self.start_marker_index,
            "first_body_index": self.first_body_index,
            "references_index": self.references_index,
            "acknowledgement_index": self.acknowledgement_index,
            "roles": {key: value.to_dict() for key, value in self.roles.items()},
        }


def get_docx_template_profile_path(template_path: Path) -> Path:
    return INTERMEDIATE_DIR / f"{template_path.stem}_template_profile.json"


def ensure_docx_template_profile(
    template_path: Path,
    *,
    profile_path: Path | None = None,
) -> tuple[Path, DocxTemplateProfile]:
    normalized_template_path = template_path.resolve()
    profile_path = profile_path or get_docx_template_profile_path(normalized_template_path)
    profile = _load_docx_template_profile(profile_path)
    if profile is None or not _is_template_profile_current(profile, normalized_template_path):
        profile = build_docx_template_profile(normalized_template_path)
        profile_path.parent.mkdir(parents=True, exist_ok=True)
        profile_path.write_text(
            json.dumps(profile.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    return profile_path, profile


def build_docx_template_profile(template_path: Path) -> DocxTemplateProfile:
    normalized_template_path = template_path.resolve()
    document = Document(str(normalized_template_path))
    paragraphs = _collect_top_level_nonempty_paragraphs(document)
    if not paragraphs:
        raise ValueError(f"Template document has no top-level paragraphs: {normalized_template_path}")

    start_index = _find_body_start_index(paragraphs)
    first_body_index = _find_first_body_heading_index(paragraphs, start_index)
    references_index = _find_references_index(paragraphs, start_index)
    acknowledgement_index = _find_acknowledgement_index(paragraphs, start_index)
    role_candidates = _collect_role_candidates(
        paragraphs,
        start_index=start_index,
        first_body_index=first_body_index,
        references_index=references_index,
        acknowledgement_index=acknowledgement_index,
    )
    roles = {
        role: _extract_paragraph_profile(role, _select_role_candidate(role, candidates))
        for role, candidates in role_candidates.items()
        if candidates
    }

    template_stat = normalized_template_path.stat()
    return DocxTemplateProfile(
        version=DOCX_TEMPLATE_PROFILE_VERSION,
        template_path=str(normalized_template_path),
        template_size=template_stat.st_size,
        template_mtime_ns=template_stat.st_mtime_ns,
        start_marker_index=start_index,
        first_body_index=first_body_index,
        references_index=references_index,
        acknowledgement_index=acknowledgement_index,
        roles=roles,
    )


def apply_docx_template(
    export_path: Path,
    *,
    template_path: Path,
    profile_path: Path | None = None,
    snapshot_path: Path | None = None,
    mode: str = "full_template_profile",
    copy_section_layout: bool = False,
) -> dict[str, Any]:
    normalized_export_path = export_path.resolve()
    normalized_mode = str(mode or "full_template_profile").strip().lower() or "full_template_profile"
    if normalized_mode not in DOCX_TEMPLATE_APPLY_MODES:
        raise ValueError(f"Unsupported DOCX template apply mode: {mode}")
    _, profile = ensure_docx_template_profile(template_path, profile_path=profile_path)
    template_document = Document(str(template_path.resolve()))
    document = Document(str(normalized_export_path))
    if copy_section_layout:
        _apply_template_section_layout(
            document,
            template_document,
            preserve_existing=False,
        )
    _apply_school_section_layout(document)
    paragraphs = _collect_top_level_nonempty_paragraphs(document)
    if not paragraphs:
        document.save(str(normalized_export_path))
        return {
            "templatePath": str(Path(profile.template_path)),
            "profilePath": str((profile_path or get_docx_template_profile_path(template_path)).resolve()),
            "appliedCount": 0,
            "bodyParagraphCount": 0,
            "mode": normalized_mode,
            "sectionLayoutApplied": bool(copy_section_layout),
            "roles": sorted(profile.roles),
        }

    start_index = _find_body_start_index(paragraphs)
    first_body_index = _find_first_body_heading_index(paragraphs, start_index)
    references_index = _find_references_index(paragraphs, start_index)
    acknowledgement_index = _find_acknowledgement_index(paragraphs, start_index)
    editable_top_level_indexes = _load_editable_top_level_paragraph_indexes(snapshot_path)
    paragraph_indexes_by_element = {id(paragraph._element): index for index, paragraph in enumerate(document.paragraphs)}
    applied_count = 0
    body_paragraph_count = 0
    applied_profiles: list[dict[str, Any]] = []

    for index, paragraph in enumerate(paragraphs):
        role = _classify_template_role(
            paragraph,
            index=index,
            first_body_index=first_body_index,
            references_index=references_index,
            acknowledgement_index=acknowledgement_index,
        )
        if role is None:
            continue
        if index < start_index and role not in STRUCTURAL_FORMAT_ROLES:
            continue
        actual_paragraph_index = paragraph_indexes_by_element.get(id(paragraph._element))
        if (
            editable_top_level_indexes is not None
            and actual_paragraph_index is not None
            and actual_paragraph_index not in editable_top_level_indexes
            and role not in STRUCTURAL_FORMAT_ROLES
        ):
            continue
        body_paragraph_count += 1
        school_profile = _build_school_profile_for_role(role, paragraph)
        if school_profile is not None:
            _apply_paragraph_profile(
                paragraph,
                school_profile,
                font_profile=school_profile,
                preserve_existing=False,
            )
        else:
            paragraph_profile = _resolve_profile_for_role(role, profile.roles)
            if paragraph_profile is None:
                continue
            font_profile = _resolve_font_profile_for_role(role, profile.roles, paragraph_profile)
            _apply_paragraph_profile(
                paragraph,
                paragraph_profile,
                font_profile=font_profile,
                preserve_existing=normalized_mode == "safe_template_fill",
            )
        applied_count += 1

    _apply_school_table_layout(document)

    document.save(str(normalized_export_path))
    return {
        "templatePath": profile.template_path,
        "profilePath": str((profile_path or get_docx_template_profile_path(template_path)).resolve()),
        "appliedCount": applied_count,
        "bodyParagraphCount": body_paragraph_count,
        "mode": normalized_mode,
        "sectionLayoutApplied": bool(copy_section_layout),
        "roles": sorted(profile.roles),
    }


def apply_school_format_rules(
    export_path: Path,
    *,
    snapshot_path: Path | None = None,
    rules: dict[str, Any] | None = None,
) -> dict[str, Any]:
    normalized_export_path = export_path.resolve()
    active_rules = rules or load_active_format_rules()
    document = Document(str(normalized_export_path))
    content_fingerprint_before = _collect_document_text_fingerprint(document)
    editable_top_level_indexes = _load_editable_top_level_paragraph_indexes(snapshot_path)
    body_scope_top_level_indexes = _load_body_scope_top_level_paragraph_indexes(snapshot_path)
    editable_only_mode = editable_top_level_indexes is not None
    _apply_school_section_layout(document, active_rules)
    paragraphs = _collect_top_level_nonempty_paragraphs(document)
    if not paragraphs:
        document.save(str(normalized_export_path))
        return {"appliedCount": 0, "bodyParagraphCount": 0, "mode": "school_rules"}

    start_index = _find_body_start_index(paragraphs)
    first_body_index = _find_first_body_heading_index(paragraphs, start_index)
    references_index = _find_references_index(paragraphs, start_index)
    acknowledgement_index = _find_acknowledgement_index(paragraphs, start_index)
    paragraph_indexes_by_element = {id(paragraph._element): index for index, paragraph in enumerate(document.paragraphs)}
    applied_count = 0
    body_paragraph_count = 0
    content_locked_style_count = 0
    applied_profiles: list[dict[str, Any]] = []

    for index, paragraph in enumerate(paragraphs):
        actual_paragraph_index = paragraph_indexes_by_element.get(id(paragraph._element))
        is_editable_top_level = (
            not editable_only_mode
            or (
                actual_paragraph_index is not None
                and actual_paragraph_index in editable_top_level_indexes
            )
        )
        role = _classify_template_role(
            paragraph,
            index=index,
            first_body_index=first_body_index,
            references_index=references_index,
            acknowledgement_index=acknowledgement_index,
        )
        if editable_only_mode and is_editable_top_level:
            if role is None:
                role = "body_text"
            elif role in {"heading_1", "heading_2", "heading_3", "heading_4", "caption", "note", "references_heading", "references_body", "toc_heading"}:
                role = "numbered_body" if _looks_like_numbered_body_item(paragraph.text.strip()) else "body_text"
        if role is None:
            continue
        if editable_only_mode:
            if (
                body_scope_top_level_indexes is None
                or actual_paragraph_index is None
                or actual_paragraph_index not in body_scope_top_level_indexes
            ):
                continue
        elif index < start_index and role not in CONTENT_LOCKED_FORMAT_ROLES:
            continue
        school_profile = _build_rule_profile_for_role(role, paragraph, active_rules)
        if school_profile is None:
            school_profile = _build_school_profile_for_role(role, paragraph)
        if school_profile is None:
            continue
        body_paragraph_count += 1
        if editable_only_mode and not is_editable_top_level:
            content_locked_style_count += 1
        _apply_paragraph_profile(
            paragraph,
            school_profile,
            font_profile=school_profile,
            preserve_existing=False,
        )
        applied_profiles.append({
            "index": actual_paragraph_index if actual_paragraph_index is not None else index,
            "role": school_profile.role,
            "text": paragraph.text.strip(),
            "fontSizePt": school_profile.font_size_pt,
            "cnFont": school_profile.cn_font,
            "enFont": school_profile.en_font,
        })
        applied_count += 1

    table_layout_stats = {"tableCount": 0, "tableParagraphCount": 0, "borderedTableCount": 0}
    if not editable_only_mode or _rules_allow_content_locked_table_format(active_rules):
        table_layout_stats = _apply_school_table_layout(document, active_rules)

    preflight_report = _write_school_format_preflight_report(normalized_export_path, applied_profiles, active_rules)
    if preflight_report["blockingIssues"]:
        raise ValueError(
            "DOCX formatting aborted: school format preflight found high-risk layout issues. "
            f"Report: {preflight_report['path']}"
        )

    content_fingerprint_after = _collect_document_text_fingerprint(document)
    if content_fingerprint_after != content_fingerprint_before:
        raise ValueError("DOCX formatting aborted: format rules changed document text content.")
    document.save(str(normalized_export_path))
    return {
        "appliedCount": applied_count,
        "bodyParagraphCount": body_paragraph_count,
        "contentLockedStyleCount": content_locked_style_count,
        "tableStyleCount": int(table_layout_stats.get("tableParagraphCount", 0) or 0),
        "tableBorderCount": int(table_layout_stats.get("borderedTableCount", 0) or 0),
        "mode": "school_rules",
        "formatScope": "body_scope_style_only" if editable_only_mode else "full_generated_document",
        "schoolName": str(active_rules.get("schoolName", "")),
        "preflightPath": str(preflight_report.get("path", "")),
        "preflightIssueCount": int(preflight_report.get("issueCount", 0) or 0),
    }


def _collect_document_text_fingerprint(document: DocxDocument) -> list[tuple[str, str]]:
    fingerprint: list[tuple[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs):
        fingerprint.append((f"p:{index}", paragraph.text))
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    fingerprint.append((f"t:{table_index}:{row_index}:{cell_index}:{paragraph_index}", paragraph.text))
    return fingerprint


def _collect_top_level_nonempty_paragraphs(document: DocxDocument) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    for block in _iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue
        if not block.text.strip():
            continue
        paragraphs.append(block)
    return paragraphs


def _load_editable_top_level_paragraph_indexes(snapshot_path: Path | None) -> set[int] | None:
    if snapshot_path is None:
        return None

    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        return None

    editable_indexes: set[int] = set()
    for unit in snapshot.units:
        if not unit.editable:
            continue
        if str(unit.target.get("kind", "")) != "paragraph":
            continue
        try:
            editable_indexes.add(int(unit.target["paragraph_index"]))
        except (KeyError, TypeError, ValueError):
            continue
    return editable_indexes


def _load_body_scope_top_level_paragraph_indexes(snapshot_path: Path | None) -> set[int] | None:
    if snapshot_path is None:
        return None

    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        return None

    outside_reasons = {"front_matter", "outside_body_scope", "back_matter", "generated_field"}
    body_scope_indexes: set[int] = set()
    for unit in snapshot.units:
        if str(unit.target.get("kind", "")) != "paragraph":
            continue
        if str(unit.protect_reason or "") in outside_reasons:
            continue
        try:
            body_scope_indexes.add(int(unit.target["paragraph_index"]))
        except (KeyError, TypeError, ValueError):
            continue
    return body_scope_indexes


def _find_body_start_index(paragraphs: list[Paragraph]) -> int:
    for index, paragraph in enumerate(paragraphs):
        normalized = _normalize_marker_text(paragraph.text)
        if normalized.startswith("摘要"):
            return index
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if _is_level_1_heading(text):
            return index
    return 0


def _find_first_body_heading_index(paragraphs: list[Paragraph], start_index: int) -> int:
    for index in range(start_index, len(paragraphs)):
        if _is_level_1_heading(paragraphs[index].text.strip()):
            return index
    return len(paragraphs)


def _find_acknowledgement_index(paragraphs: list[Paragraph], start_index: int) -> int | None:
    for index in range(start_index, len(paragraphs)):
        if _normalize_marker_text(paragraphs[index].text) in ACKNOWLEDGEMENT_MARKERS:
            return index
    return None


def _find_references_index(paragraphs: list[Paragraph], start_index: int) -> int | None:
    for index in range(start_index, len(paragraphs)):
        if _looks_like_references_heading(paragraphs[index].text):
            return index
    return None


def _collect_role_candidates(
    paragraphs: list[Paragraph],
    *,
    start_index: int,
    first_body_index: int,
    references_index: int | None,
    acknowledgement_index: int | None,
) -> dict[str, list[Paragraph]]:
    candidates: dict[str, list[Paragraph]] = {}
    for index in range(start_index, len(paragraphs)):
        paragraph = paragraphs[index]
        role = _classify_template_role(
            paragraph,
            index=index,
            first_body_index=first_body_index,
            references_index=references_index,
            acknowledgement_index=acknowledgement_index,
        )
        if role is None:
            continue
        candidates.setdefault(role, []).append(paragraph)
    return candidates


def _classify_template_role(
    paragraph: Paragraph,
    *,
    index: int,
    first_body_index: int,
    references_index: int | None,
    acknowledgement_index: int | None,
) -> str | None:
    text = paragraph.text.strip()
    normalized = _normalize_marker_text(text)
    normalized_style = _normalize_marker_text(paragraph.style.name if paragraph.style is not None else "")
    if not normalized:
        return None
    if _looks_like_references_heading(text):
        return "references_heading"
    if _looks_like_back_matter_heading(text):
        return None
    if _looks_like_caption(text):
        return "caption"
    if normalized.startswith("图注") or normalized.startswith("表注"):
        return "note"
    if _looks_like_toc_heading(text):
        return "toc_heading"
    if normalized_style.startswith("toc") or normalized_style.startswith("目录"):
        return None
    if _paragraph_has_field_code(paragraph):
        return None
    has_numbering = _paragraph_has_numbering(paragraph)
    if _looks_like_formula_paragraph(text, style_name=paragraph.style.name if paragraph.style is not None else ""):
        return None

    if references_index is not None and index >= references_index and (acknowledgement_index is None or index < acknowledgement_index):
        return "references_heading" if index == references_index else "references_body"

    if acknowledgement_index is not None and index >= acknowledgement_index:
        if index == acknowledgement_index or normalized in ACKNOWLEDGEMENT_MARKERS:
            return "ack_heading"
        return "ack_body"

    if index < first_body_index:
        if normalized.startswith("摘要"):
            return "cn_abstract_lead"
        if _is_cn_keywords_line(normalized):
            return "cn_keywords"
        if normalized.startswith("abstract"):
            return "en_abstract_lead"
        if _is_en_keywords_line(normalized):
            return "en_keywords"
        if _looks_mostly_english(text):
            return "en_abstract_body"
        if _looks_like_english_title(paragraph):
            return "en_title"
        if _looks_like_english_meta(paragraph):
            return "en_meta"
        return "cn_abstract_body"

    if _is_level_3_heading(text):
        return "heading_3"
    if _is_level_2_heading(text):
        return "heading_2"
    if _is_level_1_heading(text):
        return "heading_1"
    if has_numbering and (normalized_style.startswith("heading") or normalized_style.startswith("标题")):
        return None
    if has_numbering:
        return "numbered_body"
    return "body_text"


def _select_role_candidate(role: str, candidates: list[Paragraph]) -> Paragraph:
    if len(candidates) == 1:
        return candidates[0]
    if role in {"body_text", "numbered_body", "cn_abstract_body", "en_abstract_body", "ack_body"}:
        return max(candidates, key=_score_body_candidate)
    return candidates[0]


def _score_body_candidate(paragraph: Paragraph) -> tuple[int, int, int]:
    text = paragraph.text.strip()
    fmt = paragraph.paragraph_format
    indent_score = 1 if fmt.first_line_indent is not None else 0
    latin_penalty = -1 if _looks_mostly_english(text) else 0
    return (
        indent_score,
        len(text),
        latin_penalty,
    )


def _extract_paragraph_profile(role: str, paragraph: Paragraph) -> TemplateParagraphProfile:
    paragraph_format = paragraph.paragraph_format
    line_spacing_pt: float | None = None
    line_spacing_multiple: float | None = None
    line_spacing_value = paragraph_format.line_spacing
    if hasattr(line_spacing_value, "pt"):
        line_spacing_pt = float(line_spacing_value.pt)
    elif isinstance(line_spacing_value, (int, float)):
        line_spacing_multiple = float(line_spacing_value)

    cn_font, en_font = _resolve_paragraph_fonts(paragraph)
    font_size_pt = _resolve_paragraph_font_size(paragraph)
    bold = None if role in MIXED_LABEL_ROLES else _resolve_uniform_text_flag(paragraph, "bold")
    italic = None if role in MIXED_LABEL_ROLES else _resolve_uniform_text_flag(paragraph, "italic")

    return TemplateParagraphProfile(
        role=role,
        style_name=paragraph.style.name if paragraph.style is not None else "",
        source_text=paragraph.text.strip(),
        alignment=int(paragraph.alignment) if paragraph.alignment is not None else None,
        first_line_indent_pt=_length_to_pt(paragraph_format.first_line_indent),
        left_indent_pt=_length_to_pt(paragraph_format.left_indent),
        right_indent_pt=_length_to_pt(paragraph_format.right_indent),
        space_before_pt=_length_to_pt(paragraph_format.space_before),
        space_after_pt=_length_to_pt(paragraph_format.space_after),
        line_spacing_rule=int(paragraph_format.line_spacing_rule) if paragraph_format.line_spacing_rule is not None else None,
        line_spacing_pt=line_spacing_pt,
        line_spacing_multiple=line_spacing_multiple,
        cn_font=cn_font,
        en_font=en_font,
        font_size_pt=font_size_pt,
        bold=bold,
        italic=italic,
    )


def _resolve_profile_for_role(role: str, roles: dict[str, TemplateParagraphProfile]) -> TemplateParagraphProfile | None:
    profile = roles.get(role)
    if profile is not None:
        return profile
    for fallback_role in ROLE_FONT_FALLBACKS.get(role, ()):
        profile = roles.get(fallback_role)
        if profile is not None:
            return profile
    return None


def _resolve_font_profile_for_role(
    role: str,
    roles: dict[str, TemplateParagraphProfile],
    paragraph_profile: TemplateParagraphProfile,
) -> TemplateParagraphProfile:
    if paragraph_profile.cn_font or paragraph_profile.en_font or paragraph_profile.font_size_pt is not None:
        return paragraph_profile
    for fallback_role in ROLE_FONT_FALLBACKS.get(role, ()):
        candidate = roles.get(fallback_role)
        if candidate is None:
            continue
        if candidate.cn_font or candidate.en_font or candidate.font_size_pt is not None:
            return candidate
    return paragraph_profile


def _build_school_profile_for_role(role: str, paragraph: Paragraph) -> TemplateParagraphProfile | None:
    text = paragraph.text.strip()
    normalized = _normalize_marker_text(text)
    base = {
        "role": role,
        "style_name": paragraph.style.name if paragraph.style is not None else "",
        "source_text": text,
        "space_before_pt": 0.0,
        "space_after_pt": 0.0,
    }

    if role == "toc_heading":
        return TemplateParagraphProfile(
            **base,
            alignment=int(WD_ALIGN_PARAGRAPH.CENTER),
            space_after_pt=18.0,
            line_spacing_multiple=1.5,
            cn_font=FONT_CN_HEI,
            en_font=FONT_EN_TIMES,
            font_size_pt=SIZE_3,
            bold=False,
        )
    if role == "cn_abstract_lead":
        return _school_center_heading_profile(role, text, cn_font=FONT_CN_HEI, en_font=FONT_EN_TIMES, size_pt=SIZE_3, bold=False)
    if role == "en_abstract_lead":
        return _school_center_heading_profile(role, text, cn_font=FONT_EN_TIMES, en_font=FONT_EN_TIMES, size_pt=SIZE_3, bold=True)
    if role == "cn_keywords":
        return _school_body_profile(role, text, cn_font=FONT_CN_SONG, en_font=FONT_EN_TIMES, bold=None, first_line_indent_pt=0.0)
    if role == "en_keywords":
        return _school_body_profile(role, text, cn_font=FONT_EN_TIMES, en_font=FONT_EN_TIMES, bold=None, first_line_indent_pt=0.0)
    if role == "cn_abstract_body":
        return _school_body_profile(role, text, cn_font=FONT_CN_SONG, en_font=FONT_EN_TIMES)
    if role == "en_abstract_body":
        return _school_body_profile(role, text, cn_font=FONT_EN_TIMES, en_font=FONT_EN_TIMES)
    if role == "heading_1":
        return _school_left_heading_profile(role, text, size_pt=SIZE_4)
    if role == "heading_2":
        return _school_left_heading_profile(role, text, size_pt=SIZE_XIAO_4)
    if role == "heading_3":
        return _school_left_heading_profile(role, text, size_pt=SIZE_5)
    if _is_level_4_heading(text):
        return _school_left_heading_profile("heading_4", text, size_pt=SIZE_5)
    if role == "references_heading":
        return TemplateParagraphProfile(
            **base,
            alignment=int(WD_ALIGN_PARAGRAPH.CENTER),
            line_spacing_multiple=1.5,
            cn_font=FONT_CN_HEI,
            en_font=FONT_EN_TIMES,
            font_size_pt=SIZE_4,
            bold=False,
        )
    if role == "references_body":
        return _school_body_profile(role, text, cn_font=FONT_CN_SONG, en_font=FONT_EN_TIMES, first_line_indent_pt=0.0)
    if role == "ack_heading":
        return _school_center_heading_profile(role, text, cn_font=FONT_CN_HEI, en_font=FONT_EN_TIMES, size_pt=SIZE_3, bold=False)
    if role == "ack_body":
        return _school_body_profile(role, text, cn_font=FONT_CN_SONG, en_font=FONT_EN_TIMES, font_size_pt=SIZE_XIAO_4)
    if _looks_like_caption(text):
        return _school_caption_profile(role, text)
    if role == "numbered_body":
        return _school_body_profile(role, text, cn_font=FONT_CN_SONG, en_font=FONT_EN_TIMES, first_line_indent_pt=None)
    if role == "body_text":
        if normalized.startswith("图注") or normalized.startswith("表注"):
            return _school_note_profile(role, text)
        return _school_body_profile(role, text, cn_font=FONT_CN_SONG, en_font=FONT_EN_TIMES)
    return None


def _build_rule_profile_for_role(role: str, paragraph: Paragraph, rules: dict[str, Any]) -> TemplateParagraphProfile | None:
    style_key = _resolve_rule_style_key(role, paragraph)
    styles = rules.get("styles")
    if not isinstance(styles, dict):
        return None
    raw_style = styles.get(style_key)
    if not isinstance(raw_style, dict):
        return None
    explicit_fields = _get_explicit_rule_fields(rules, style_key)
    alignment = ALIGNMENT_BY_NAME.get(str(_explicit_rule_value(raw_style, explicit_fields, "alignment") or "").strip().lower())
    first_line_indent_pt = _as_optional_float(_explicit_rule_value(raw_style, explicit_fields, "firstLineIndentPt"))
    if role == "numbered_body":
        first_line_indent_pt = None
    line_spacing_pt = _as_optional_float(_explicit_rule_value(raw_style, explicit_fields, "lineSpacingPt"))
    return TemplateParagraphProfile(
        role=style_key,
        style_name=paragraph.style.name if paragraph.style is not None else "",
        source_text=paragraph.text.strip(),
        alignment=int(alignment) if alignment is not None else None,
        first_line_indent_pt=first_line_indent_pt,
        space_before_pt=_as_optional_float(_explicit_rule_value(raw_style, explicit_fields, "spaceBeforePt")),
        space_after_pt=_as_optional_float(_explicit_rule_value(raw_style, explicit_fields, "spaceAfterPt")),
        line_spacing_rule=int(WD_LINE_SPACING.EXACTLY) if line_spacing_pt is not None else None,
        line_spacing_pt=line_spacing_pt,
        line_spacing_multiple=_as_optional_float(_explicit_rule_value(raw_style, explicit_fields, "lineSpacingMultiple")),
        cn_font=_as_optional_str(_explicit_rule_value(raw_style, explicit_fields, "cnFont")),
        en_font=_as_optional_str(_explicit_rule_value(raw_style, explicit_fields, "enFont")),
        font_size_pt=_as_optional_float(_explicit_rule_value(raw_style, explicit_fields, "fontSizePt")),
        bold=_as_optional_bool(_explicit_rule_value(raw_style, explicit_fields, "bold")),
        italic=_as_optional_bool(_explicit_rule_value(raw_style, explicit_fields, "italic")),
    )


def _resolve_rule_style_key(role: str, paragraph: Paragraph) -> str:
    text = paragraph.text.strip()
    normalized = _normalize_marker_text(text)
    if role == "numbered_body":
        return "body_text"
    if _is_level_4_heading(text):
        return "heading_4"
    if _looks_like_caption(text):
        return "caption"
    if normalized.startswith("图注") or normalized.startswith("表注"):
        return "note"
    return role


def _get_explicit_rule_fields(rules: dict[str, Any], style_key: str) -> set[str] | None:
    meta = rules.get("styleMeta")
    if not isinstance(meta, dict):
        return None
    style_meta = meta.get(style_key)
    if not isinstance(style_meta, dict) and meta:
        return set()
    if not isinstance(style_meta, dict) or "explicitFields" not in style_meta:
        return None
    raw_fields = style_meta.get("explicitFields")
    if not isinstance(raw_fields, (list, tuple, set)):
        return set()
    return {str(item) for item in raw_fields if str(item) in STYLE_FORMAT_FIELDS}


def _explicit_rule_value(raw_style: dict[str, Any], explicit_fields: set[str] | None, key: str) -> Any:
    if explicit_fields is not None and key not in explicit_fields:
        return None
    return raw_style.get(key)


def _school_center_heading_profile(
    role: str,
    text: str,
    *,
    cn_font: str,
    en_font: str,
    size_pt: float,
    bold: bool,
) -> TemplateParagraphProfile:
    return TemplateParagraphProfile(
        role=role,
        style_name="",
        source_text=text,
        alignment=int(WD_ALIGN_PARAGRAPH.CENTER),
        space_before_pt=12.0,
        space_after_pt=0.0,
        line_spacing_multiple=1.5,
        cn_font=cn_font,
        en_font=en_font,
        font_size_pt=size_pt,
        bold=bold,
    )


def _school_left_heading_profile(role: str, text: str, *, size_pt: float) -> TemplateParagraphProfile:
    return TemplateParagraphProfile(
        role=role,
        style_name="",
        source_text=text,
        alignment=int(WD_ALIGN_PARAGRAPH.LEFT),
        first_line_indent_pt=0.0,
        space_before_pt=0.0,
        space_after_pt=0.0,
        line_spacing_pt=20.0,
        line_spacing_rule=int(WD_LINE_SPACING.EXACTLY),
        cn_font=FONT_CN_HEI,
        en_font=FONT_EN_TIMES,
        font_size_pt=size_pt,
        bold=False,
    )


def _school_body_profile(
    role: str,
    text: str,
    *,
    cn_font: str,
    en_font: str,
    font_size_pt: float = SIZE_5,
    bold: bool | None = False,
    first_line_indent_pt: float | None = 21.0,
) -> TemplateParagraphProfile:
    return TemplateParagraphProfile(
        role=role,
        style_name="",
        source_text=text,
        alignment=int(WD_ALIGN_PARAGRAPH.JUSTIFY),
        first_line_indent_pt=first_line_indent_pt,
        space_before_pt=0.0,
        space_after_pt=0.0,
        line_spacing_rule=int(WD_LINE_SPACING.EXACTLY),
        line_spacing_pt=20.0,
        cn_font=cn_font,
        en_font=en_font,
        font_size_pt=font_size_pt,
        bold=bold,
    )


def _school_caption_profile(role: str, text: str) -> TemplateParagraphProfile:
    return TemplateParagraphProfile(
        role=role,
        style_name="",
        source_text=text,
        alignment=int(WD_ALIGN_PARAGRAPH.CENTER),
        first_line_indent_pt=0.0,
        space_before_pt=0.0,
        space_after_pt=0.0,
        line_spacing_rule=int(WD_LINE_SPACING.EXACTLY),
        line_spacing_pt=20.0,
        cn_font=FONT_CN_HEI,
        en_font=FONT_EN_TIMES,
        font_size_pt=SIZE_XIAO_5,
        bold=False,
    )


def _school_note_profile(role: str, text: str) -> TemplateParagraphProfile:
    return TemplateParagraphProfile(
        role=role,
        style_name="",
        source_text=text,
        alignment=int(WD_ALIGN_PARAGRAPH.LEFT),
        first_line_indent_pt=21.0,
        space_before_pt=0.0,
        space_after_pt=0.0,
        line_spacing_rule=int(WD_LINE_SPACING.EXACTLY),
        line_spacing_pt=20.0,
        cn_font=FONT_CN_KAI,
        en_font=FONT_EN_TIMES,
        font_size_pt=SIZE_XIAO_5,
        bold=False,
    )


def _apply_paragraph_profile(
    paragraph: Paragraph,
    profile: TemplateParagraphProfile,
    *,
    font_profile: TemplateParagraphProfile,
    preserve_existing: bool = False,
) -> None:
    paragraph_format = paragraph.paragraph_format
    if profile.alignment is not None and (not preserve_existing or paragraph.alignment is None):
        paragraph.alignment = WD_ALIGN_PARAGRAPH(profile.alignment)
    if profile.first_line_indent_pt is not None and (not preserve_existing or paragraph_format.first_line_indent is None):
        paragraph_format.first_line_indent = Pt(profile.first_line_indent_pt)
    if profile.left_indent_pt is not None and (not preserve_existing or paragraph_format.left_indent is None):
        paragraph_format.left_indent = Pt(profile.left_indent_pt)
    if profile.right_indent_pt is not None and (not preserve_existing or paragraph_format.right_indent is None):
        paragraph_format.right_indent = Pt(profile.right_indent_pt)
    if profile.space_before_pt is not None and (not preserve_existing or paragraph_format.space_before is None):
        paragraph_format.space_before = Pt(profile.space_before_pt)
    if profile.space_after_pt is not None and (not preserve_existing or paragraph_format.space_after is None):
        paragraph_format.space_after = Pt(profile.space_after_pt)

    if profile.line_spacing_rule is not None and (not preserve_existing or paragraph_format.line_spacing_rule is None):
        paragraph_format.line_spacing_rule = WD_LINE_SPACING(profile.line_spacing_rule)
    if profile.line_spacing_pt is not None and (not preserve_existing or paragraph_format.line_spacing is None):
        paragraph_format.line_spacing = Pt(profile.line_spacing_pt)
    elif profile.line_spacing_multiple is not None and (not preserve_existing or paragraph_format.line_spacing is None):
        paragraph_format.line_spacing = profile.line_spacing_multiple

    runs = [run for run in paragraph.runs if run.text]
    if not runs:
        if not paragraph.text.strip():
            return
        runs = [paragraph.add_run(paragraph.text)]

    for run in runs:
        _apply_run_font_settings(
            run,
            cn_font=font_profile.cn_font,
            en_font=font_profile.en_font,
            font_size_pt=font_profile.font_size_pt,
            bold=profile.bold,
            italic=profile.italic,
            preserve_existing=preserve_existing,
        )


def _apply_run_font_settings(
    run: Run,
    *,
    cn_font: str | None,
    en_font: str | None,
    font_size_pt: float | None,
    bold: bool | None,
    italic: bool | None,
    preserve_existing: bool = False,
) -> None:
    if font_size_pt is not None and (not preserve_existing or run.font.size is None):
        run.font.size = Pt(font_size_pt)
    if bold is not None and (not preserve_existing or run.font.bold is None):
        run.font.bold = bold
    if italic is not None and (not preserve_existing or run.font.italic is None):
        run.font.italic = italic
    existing_font_name = str(run.font.name or "").strip()
    if en_font and (not preserve_existing or not existing_font_name):
        run.font.name = en_font
    elif cn_font and (not preserve_existing or not existing_font_name):
        run.font.name = cn_font

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    existing_east_asia = str(r_fonts.get(qn("w:eastAsia")) or "").strip()
    existing_ascii = str(r_fonts.get(qn("w:ascii")) or "").strip()
    existing_hansi = str(r_fonts.get(qn("w:hAnsi")) or "").strip()
    existing_cs = str(r_fonts.get(qn("w:cs")) or "").strip()
    if cn_font and (not preserve_existing or not existing_east_asia):
        r_fonts.set(qn("w:eastAsia"), cn_font)
    if en_font and (not preserve_existing or not existing_ascii):
        r_fonts.set(qn("w:ascii"), en_font)
    if en_font and (not preserve_existing or not existing_hansi):
        r_fonts.set(qn("w:hAnsi"), en_font)
    if en_font and (not preserve_existing or not existing_cs):
        r_fonts.set(qn("w:cs"), en_font)



def _school_format_preflight_report_path(export_path: Path) -> Path:
    INTERMEDIATE_DIR.mkdir(parents=True, exist_ok=True)
    return INTERMEDIATE_DIR / f"{export_path.stem}_format_preflight.json"


def _write_school_format_preflight_report(
    export_path: Path,
    applied_profiles: list[dict[str, Any]],
    rules: dict[str, Any],
) -> dict[str, Any]:
    issues: list[dict[str, Any]] = []
    body_style = rules.get("styles", {}).get("body_text", {}) if isinstance(rules.get("styles"), dict) else {}
    expected_body_size = _as_optional_float(body_style.get("fontSizePt")) if isinstance(body_style, dict) else None
    for item in applied_profiles:
        role = str(item.get("role", ""))
        text = str(item.get("text", "")).strip()
        font_size = _as_optional_float(item.get("fontSizePt"))
        if role.startswith("heading_") and not _is_safe_heading_text(text, max_chars=44):
            issues.append({
                "severity": "blocker",
                "code": "long_heading_candidate",
                "paragraphIndex": item.get("index"),
                "role": role,
                "message": "A long paragraph was classified as a heading.",
                "sample": text[:160],
            })
        if role == "body_text" and expected_body_size is not None and font_size is not None and font_size > expected_body_size + 1.0:
            issues.append({
                "severity": "warning",
                "code": "body_font_size_unusual",
                "paragraphIndex": item.get("index"),
                "role": role,
                "message": "Body paragraph font size is larger than expected.",
                "sample": text[:160],
            })
    blocking = [issue for issue in issues if issue.get("severity") == "blocker"]
    report_path = _school_format_preflight_report_path(export_path)
    report = {
        "ok": not blocking,
        "path": str(report_path),
        "issueCount": len(issues),
        "blockingIssues": blocking,
        "issues": issues,
        "appliedCount": len(applied_profiles),
    }
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def _apply_school_section_layout(document: DocxDocument, rules: dict[str, Any] | None = None) -> None:
    page_rules = rules.get("page", {}) if isinstance(rules, dict) else {}
    raw_explicit_fields = rules.get("pageExplicitFields") if isinstance(rules, dict) else None
    explicit_fields = {str(item) for item in raw_explicit_fields if str(item) in PAGE_FORMAT_FIELDS} if isinstance(raw_explicit_fields, (list, tuple, set)) else None
    for section in document.sections:
        margins = {
            "top_margin": page_rules.get("topMarginCm", SCHOOL_PAGE_MARGINS_CM["top_margin"]),
            "bottom_margin": page_rules.get("bottomMarginCm", SCHOOL_PAGE_MARGINS_CM["bottom_margin"]),
            "left_margin": page_rules.get("leftMarginCm", SCHOOL_PAGE_MARGINS_CM["left_margin"]),
            "right_margin": page_rules.get("rightMarginCm", SCHOOL_PAGE_MARGINS_CM["right_margin"]),
        }
        field_by_attribute = {
            "top_margin": "topMarginCm",
            "bottom_margin": "bottomMarginCm",
            "left_margin": "leftMarginCm",
            "right_margin": "rightMarginCm",
        }
        for attribute, value_cm in margins.items():
            if explicit_fields is not None and field_by_attribute[attribute] not in explicit_fields:
                continue
            setattr(section, attribute, Cm(value_cm))


def _rules_allow_content_locked_table_format(rules: dict[str, Any] | None) -> bool:
    if not isinstance(rules, dict):
        return False
    if str(rules.get("schoolName", "")).strip().lower() == "default":
        return True
    meta = rules.get("styleMeta")
    if isinstance(meta, dict):
        table_meta = meta.get("table_text")
        if isinstance(table_meta, dict) and not bool(table_meta.get("isInferred")):
            return True
    quality = rules.get("quality")
    if isinstance(quality, dict):
        explicit_roles = quality.get("explicitRoles")
        if isinstance(explicit_roles, list) and "table_text" in explicit_roles:
            return True
    return False


def _apply_school_table_layout(document: DocxDocument, rules: dict[str, Any] | None = None) -> dict[str, int]:
    table_style = None
    if isinstance(rules, dict) and isinstance(rules.get("styles"), dict):
        raw_table_style = rules["styles"].get("table_text")
        if isinstance(raw_table_style, dict):
            table_style = raw_table_style
    table_count = 0
    table_paragraph_count = 0
    bordered_table_count = 0
    for table in document.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    table_paragraph_count += 1
                    paragraph.alignment = paragraph.alignment or WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    paragraph_format.line_spacing = Pt(20)
                    for run in paragraph.runs:
                        _apply_run_font_settings(
                            run,
                            cn_font=_as_optional_str(table_style.get("cnFont")) if table_style else FONT_CN_SONG,
                            en_font=_as_optional_str(table_style.get("enFont")) if table_style else FONT_EN_TIMES,
                            font_size_pt=_as_optional_float(table_style.get("fontSizePt")) if table_style else SIZE_XIAO_5,
                            bold=_as_optional_bool(table_style.get("bold")) if table_style else None,
                            italic=_as_optional_bool(table_style.get("italic")) if table_style else None,
                            preserve_existing=False,
                        )
        _apply_three_line_table_borders(table)
        bordered_table_count += 1
    return {
        "tableCount": table_count,
        "tableParagraphCount": table_paragraph_count,
        "borderedTableCount": bordered_table_count,
    }


def _apply_three_line_table_borders(table: Any) -> None:
    table_pr = table._tbl.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if edge in {"left", "right", "insideV"}:
            element.set(qn("w:val"), "nil")
            continue
        element.set(qn("w:val"), "single")
        element.set(qn("w:color"), "000000")
        element.set(qn("w:space"), "0")
        element.set(qn("w:sz"), "12" if edge in {"top", "bottom"} else "4")


def _apply_template_section_layout(
    document: DocxDocument,
    template_document: DocxDocument,
    *,
    preserve_existing: bool,
) -> None:
    if not document.sections or not template_document.sections:
        return

    template_section = template_document.sections[0]
    for section in document.sections:
        _copy_section_length(section, "top_margin", template_section.top_margin, preserve_existing=preserve_existing)
        _copy_section_length(section, "bottom_margin", template_section.bottom_margin, preserve_existing=preserve_existing)
        _copy_section_length(section, "left_margin", template_section.left_margin, preserve_existing=preserve_existing)
        _copy_section_length(section, "right_margin", template_section.right_margin, preserve_existing=preserve_existing)
        _copy_section_length(section, "header_distance", template_section.header_distance, preserve_existing=preserve_existing)
        _copy_section_length(section, "footer_distance", template_section.footer_distance, preserve_existing=preserve_existing)
        _copy_section_length(section, "gutter", template_section.gutter, preserve_existing=preserve_existing)
        _copy_section_length(section, "page_width", template_section.page_width, preserve_existing=preserve_existing)
        _copy_section_length(section, "page_height", template_section.page_height, preserve_existing=preserve_existing)
        if not preserve_existing or getattr(section, "orientation", None) is None:
            section.orientation = template_section.orientation
        if not preserve_existing or not section.different_first_page_header_footer:
            section.different_first_page_header_footer = template_section.different_first_page_header_footer


def _copy_section_length(section: Any, attribute: str, value: Any, *, preserve_existing: bool) -> None:
    if value is None:
        return
    if preserve_existing and getattr(section, attribute) is not None:
        return
    setattr(section, attribute, value)


def _paragraph_has_numbering(paragraph: Paragraph) -> bool:
    p_pr = getattr(paragraph._element, "pPr", None)
    if p_pr is None:
        return False
    return p_pr.numPr is not None


def _resolve_paragraph_fonts(paragraph: Paragraph) -> tuple[str | None, str | None]:
    cn_font = None
    en_font = None
    base_fonts = _resolve_style_fonts(paragraph.style)
    first_run_fonts: dict[str, str] = {}

    for run in paragraph.runs:
        text = run.text.strip()
        if not text:
            continue
        run_fonts = _resolve_run_fonts(run, paragraph=paragraph)
        if not first_run_fonts:
            first_run_fonts = dict(run_fonts)
        if cn_font is None and _contains_cjk(text):
            cn_font = run_fonts.get("eastAsia") or run_fonts.get("ascii") or base_fonts.get("eastAsia")
        if en_font is None and _contains_latin(text):
            en_font = run_fonts.get("ascii") or run_fonts.get("hAnsi") or base_fonts.get("ascii")

    if cn_font is None:
        cn_font = first_run_fonts.get("eastAsia") or base_fonts.get("eastAsia") or first_run_fonts.get("ascii") or base_fonts.get("ascii")
    if en_font is None:
        en_font = (
            first_run_fonts.get("ascii")
            or first_run_fonts.get("hAnsi")
            or base_fonts.get("ascii")
            or base_fonts.get("hAnsi")
            or first_run_fonts.get("eastAsia")
            or base_fonts.get("eastAsia")
        )
    return cn_font, en_font


def _resolve_paragraph_font_size(paragraph: Paragraph) -> float | None:
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        size_pt = _resolve_run_font_size_pt(run, paragraph=paragraph)
        if size_pt is not None:
            return size_pt
    return _resolve_style_font_size_pt(paragraph.style)


def _resolve_uniform_text_flag(paragraph: Paragraph, attribute: str) -> bool | None:
    values: list[bool] = []
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        value = _resolve_run_text_flag(run, paragraph=paragraph, attribute=attribute)
        if value is None:
            continue
        values.append(bool(value))
    if not values:
        return None
    if all(item == values[0] for item in values):
        return values[0]
    return None


def _resolve_run_text_flag(run: Run, *, paragraph: Paragraph, attribute: str) -> bool | None:
    value = getattr(run.font, attribute)
    if value is not None:
        return bool(value)
    if run.style is not None:
        for style in _iter_style_chain(run.style):
            font_value = getattr(style.font, attribute)
            if font_value is not None:
                return bool(font_value)
    if paragraph.style is not None:
        for style in _iter_style_chain(paragraph.style):
            font_value = getattr(style.font, attribute)
            if font_value is not None:
                return bool(font_value)
    return None


def _resolve_run_font_size_pt(run: Run, *, paragraph: Paragraph) -> float | None:
    if run.font.size is not None:
        return float(run.font.size.pt)
    if run.style is not None:
        size_pt = _resolve_style_font_size_pt(run.style)
        if size_pt is not None:
            return size_pt
    return _resolve_style_font_size_pt(paragraph.style)


def _resolve_style_font_size_pt(style: BaseStyle | None) -> float | None:
    if style is None:
        return None
    for style_item in _iter_style_chain(style):
        if style_item.font.size is not None:
            return float(style_item.font.size.pt)
    return None


def _resolve_run_fonts(run: Run, *, paragraph: Paragraph) -> dict[str, str]:
    fonts: dict[str, str] = {}
    _merge_missing_font_values(fonts, _get_fonts_from_rpr(run._element.rPr))
    if run.style is not None:
        _merge_missing_font_values(fonts, _resolve_style_fonts(run.style))
    _merge_missing_font_values(fonts, _resolve_style_fonts(paragraph.style))
    if run.font.name:
        fonts.setdefault("ascii", run.font.name)
        fonts.setdefault("hAnsi", run.font.name)
    return fonts


def _resolve_style_fonts(style: BaseStyle | None) -> dict[str, str]:
    fonts: dict[str, str] = {}
    if style is None:
        return fonts
    for style_item in _iter_style_chain(style):
        _merge_missing_font_values(fonts, _get_fonts_from_rpr(style_item.element.rPr))
        if style_item.font.name:
            fonts.setdefault("ascii", style_item.font.name)
            fonts.setdefault("hAnsi", style_item.font.name)
    return fonts


def _iter_style_chain(style: BaseStyle | None) -> Iterable[BaseStyle]:
    current = style
    while current is not None:
        yield current
        current = current.base_style


def _get_fonts_from_rpr(r_pr: Any) -> dict[str, str]:
    if r_pr is None or r_pr.rFonts is None:
        return {}
    r_fonts = r_pr.rFonts
    font_values: dict[str, str] = {}
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        value = r_fonts.get(qn(f"w:{key}"))
        if value:
            font_values[key] = str(value)
    return font_values


def _merge_missing_font_values(target: dict[str, str], source: dict[str, str]) -> None:
    for key, value in source.items():
        if value and key not in target:
            target[key] = value


def _length_to_pt(value: Any) -> float | None:
    if value is None:
        return None
    if hasattr(value, "pt"):
        return float(value.pt)
    return None


def _is_safe_heading_text(text: str, *, max_chars: int) -> bool:
    normalized = re.sub(r"\s+", " ", text.strip())
    if not normalized:
        return False
    cjk_count = len(CJK_CHAR_RE.findall(normalized))
    if len(normalized) > max_chars or cjk_count > max_chars:
        return False
    if HEADING_TERMINAL_PUNCTUATION_RE.search(normalized):
        return False
    if len(HEADING_BODY_PUNCTUATION_RE.findall(normalized)) >= 2:
        return False
    if normalized.count("?") + normalized.count(",") >= 1 and len(normalized) > 18:
        return False
    return True


def _is_level_1_heading(text: str) -> bool:
    normalized = text.strip()
    if _looks_like_numbered_body_item(normalized):
        return False
    return bool(LEVEL_1_HEADING_RE.match(normalized)) and _is_safe_heading_text(normalized, max_chars=32)


def _is_level_2_heading(text: str) -> bool:
    normalized = text.strip()
    if _looks_like_numbered_body_item(normalized):
        return False
    return bool(LEVEL_2_HEADING_RE.match(normalized)) and _is_safe_heading_text(normalized, max_chars=36)


def _is_level_3_heading(text: str) -> bool:
    normalized = text.strip()
    if _looks_like_numbered_body_item(normalized):
        return False
    return bool(LEVEL_3_HEADING_RE.match(normalized)) and _is_safe_heading_text(normalized, max_chars=40)


def _is_level_4_heading(text: str) -> bool:
    normalized = text.strip()
    if _looks_like_numbered_body_item(normalized):
        return False
    return bool(LEVEL_4_HEADING_RE.match(normalized)) and _is_safe_heading_text(normalized, max_chars=44)


def _looks_like_english_title(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False
    if paragraph.alignment not in {None, WD_ALIGN_PARAGRAPH.CENTER}:
        return False
    return bool(EN_TITLE_RE.match(text)) and _looks_mostly_english(text)


def _looks_like_english_meta(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False
    if not _looks_mostly_english(text):
        return False
    if len(text) > 40:
        return False
    return paragraph.alignment in {None, WD_ALIGN_PARAGRAPH.CENTER}


def _is_cn_keywords_line(normalized_text: str) -> bool:
    return any(normalized_text.startswith(prefix) for prefix in KEYWORD_PREFIXES)


def _is_en_keywords_line(normalized_text: str) -> bool:
    return any(normalized_text.startswith(prefix) for prefix in EN_KEYWORDS_PREFIXES)


def _looks_mostly_english(text: str) -> bool:
    latin_count = len(LATIN_CHAR_RE.findall(text))
    cjk_count = len(CJK_CHAR_RE.findall(text))
    return latin_count >= 12 and latin_count >= max(12, cjk_count * 3)


def _contains_cjk(text: str) -> bool:
    return bool(CJK_CHAR_RE.search(text))


def _contains_latin(text: str) -> bool:
    return bool(LATIN_CHAR_RE.search(text))


def _load_docx_template_profile(path: Path) -> DocxTemplateProfile | None:
    if not path.exists():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(data, dict):
        return None
    raw_roles = data.get("roles")
    if not isinstance(raw_roles, dict):
        return None
    roles: dict[str, TemplateParagraphProfile] = {}
    for role, raw_profile in raw_roles.items():
        if not isinstance(role, str) or not isinstance(raw_profile, dict):
            continue
        roles[role] = TemplateParagraphProfile(
            role=role,
            style_name=str(raw_profile.get("style_name", "")),
            source_text=str(raw_profile.get("source_text", "")),
            alignment=_as_optional_int(raw_profile.get("alignment")),
            first_line_indent_pt=_as_optional_float(raw_profile.get("first_line_indent_pt")),
            left_indent_pt=_as_optional_float(raw_profile.get("left_indent_pt")),
            right_indent_pt=_as_optional_float(raw_profile.get("right_indent_pt")),
            space_before_pt=_as_optional_float(raw_profile.get("space_before_pt")),
            space_after_pt=_as_optional_float(raw_profile.get("space_after_pt")),
            line_spacing_rule=_as_optional_int(raw_profile.get("line_spacing_rule")),
            line_spacing_pt=_as_optional_float(raw_profile.get("line_spacing_pt")),
            line_spacing_multiple=_as_optional_float(raw_profile.get("line_spacing_multiple")),
            cn_font=_as_optional_str(raw_profile.get("cn_font")),
            en_font=_as_optional_str(raw_profile.get("en_font")),
            font_size_pt=_as_optional_float(raw_profile.get("font_size_pt")),
            bold=_as_optional_bool(raw_profile.get("bold")),
            italic=_as_optional_bool(raw_profile.get("italic")),
        )
    return DocxTemplateProfile(
        version=int(data.get("version", DOCX_TEMPLATE_PROFILE_VERSION)),
        template_path=str(data.get("template_path", "")),
        template_size=int(data.get("template_size", 0)),
        template_mtime_ns=int(data.get("template_mtime_ns", 0)),
        start_marker_index=int(data.get("start_marker_index", 0)),
        first_body_index=int(data.get("first_body_index", 0)),
        references_index=_as_optional_int(data.get("references_index")),
        acknowledgement_index=_as_optional_int(data.get("acknowledgement_index")),
        roles=roles,
    )


def _is_template_profile_current(profile: DocxTemplateProfile, template_path: Path) -> bool:
    try:
        template_stat = template_path.stat()
    except OSError:
        return False
    return (
        profile.version == DOCX_TEMPLATE_PROFILE_VERSION
        and Path(profile.template_path).resolve() == template_path.resolve()
        and profile.template_size == template_stat.st_size
        and profile.template_mtime_ns == template_stat.st_mtime_ns
    )


def _as_optional_int(value: Any) -> int | None:
    if value is None:
        return None
    return int(value)


def _as_optional_float(value: Any) -> float | None:
    if value is None:
        return None
    return float(value)


def _as_optional_str(value: Any) -> str | None:
    if value is None:
        return None
    normalized = str(value).strip()
    return normalized or None


def _as_optional_bool(value: Any) -> bool | None:
    if value is None:
        return None
    return bool(value)


def main(argv: list[str] | None = None) -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Analyze or apply a DOCX school template")
    subparsers = parser.add_subparsers(dest="command", required=True)

    analyze_parser = subparsers.add_parser("analyze", help="Analyze a DOCX template into a cached JSON profile")
    analyze_parser.add_argument("template", type=Path, help="Path to the DOCX school template")

    apply_parser = subparsers.add_parser("apply", help="Apply a DOCX template profile to an exported DOCX")
    apply_parser.add_argument("template", type=Path, help="Path to the DOCX school template")
    apply_parser.add_argument("target", type=Path, help="Path to the DOCX file that should receive the template")
    apply_parser.add_argument(
        "--mode",
        choices=sorted(DOCX_TEMPLATE_APPLY_MODES),
        default="full_template_profile",
        help="How aggressively template formatting should be applied",
    )
    apply_parser.add_argument(
        "--copy-section-layout",
        action="store_true",
        help="Copy page layout settings from the template before paragraph styling",
    )

    args = parser.parse_args(argv)

    if args.command == "analyze":
        profile_path, profile = ensure_docx_template_profile(args.template)
        print(json.dumps({"profilePath": str(profile_path), **profile.to_dict()}, ensure_ascii=False, indent=2))
        return

    if args.command == "apply":
        result = apply_docx_template(
            args.target,
            template_path=args.template,
            mode=args.mode,
            copy_section_layout=bool(args.copy_section_layout),
        )
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    parser.error("Unknown command")


if __name__ == "__main__":
    main()
