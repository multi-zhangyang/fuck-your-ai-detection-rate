from __future__ import annotations

import copy
import json
import sys
import tempfile
from pathlib import Path
from typing import Any, Callable


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402

import app_service  # noqa: E402
from docx_bodymap import get_body_map_unit_model_format_anchors, load_docx_body_map  # noqa: E402
from docx_export_regression import create_regression_sample, identity_transform  # noqa: E402
from round_helper import run_document_round  # noqa: E402


REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_targeted_format_anchor_regression_report.json"
ANCHOR_TEXT = "基础模型"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _add_styled_anchor(source_path: Path) -> None:
    document = Document(str(source_path))
    paragraph = next(
        item
        for item in document.paragraphs
        if ANCHOR_TEXT in item.text and "YOLOv8" in item.text
    )
    original_text = paragraph.text
    prefix, suffix = original_text.split(ANCHOR_TEXT, 1)
    paragraph.clear()
    paragraph.add_run(prefix)
    styled_run = paragraph.add_run(ANCHOR_TEXT)
    styled_run.bold = True
    paragraph.add_run(suffix)
    document.save(str(source_path))


def _target_fixture(output_path: Path, compare_payload: dict[str, Any]) -> tuple[str, int, str]:
    body_map_path = app_service._find_body_map_path_for_output(output_path)
    _assert(body_map_path is not None, "DOCX targeted fixture has no body-map path")
    body_map = load_docx_body_map(body_map_path)
    _assert(body_map is not None, "DOCX targeted fixture body map is unreadable")
    target_paragraph_index = -1
    anchors: list[str] = []
    for paragraph_index, unit in enumerate(body_map.units):
        unit_anchors = get_body_map_unit_model_format_anchors(unit)
        if ANCHOR_TEXT in unit_anchors:
            target_paragraph_index = paragraph_index
            anchors = unit_anchors
            break
    _assert(target_paragraph_index >= 0, "styled anchor did not reach an editable body-map unit")
    _assert(anchors == [ANCHOR_TEXT], f"fixture anchor plan drifted: {anchors}")

    raw_chunks = compare_payload.get("chunks")
    _assert(isinstance(raw_chunks, list), "DOCX targeted fixture compare has no chunks")
    target_chunk = next(
        (
            item
            for item in raw_chunks
            if isinstance(item, dict)
            and item.get("paragraphIndex") == target_paragraph_index
            and str(item.get("inputText", "") or "").count(ANCHOR_TEXT) == 1
        ),
        None,
    )
    _assert(isinstance(target_chunk, dict), "styled anchor was not uniquely assigned to one frozen chunk")
    return str(target_chunk.get("chunkId", "") or ""), target_paragraph_index, ANCHOR_TEXT


def _materialized_target_text(
    output_path: Path,
    compare_payload: dict[str, Any],
    chunk_id: str,
) -> str:
    materialized = app_service._materialize_rate_audit_output(
        output_path,
        compare_payload=compare_payload,
    )
    target = next(
        (
            item
            for item in (materialized.get("chunks") or [])
            if isinstance(item, dict) and str(item.get("chunkId", "") or "") == chunk_id
        ),
        None,
    )
    _assert(isinstance(target, dict), "review materialization lost the styled target chunk")
    text = str(target.get("text", "") or "")
    _assert(text.count(ANCHOR_TEXT) == 1, "styled target is not unique before negative probes")
    return text


def _builder_probe(
    constructor_events: list[dict[str, Any]],
    model_events: list[dict[str, Any]],
) -> Callable[[dict[str, Any]], tuple[Callable[..., str], str]]:
    def builder(model_config: dict[str, Any]) -> tuple[Callable[..., str], str]:
        constructor_events.append({"modelConfig": copy.deepcopy(model_config)})

        def transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
            model_events.append(
                {
                    "chunkText": chunk_text,
                    "promptInput": prompt_input,
                    "round": round_number,
                    "chunkId": chunk_id,
                }
            )
            return chunk_text

        return transform, "targeted-format-anchor-regression"

    return builder


def main() -> int:
    checks: list[str] = []
    work_root = ROOT_DIR / "finish" / "web_exports"
    work_root.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="docx-targeted-anchor-", dir=work_root) as temp_dir:
        source_path = Path(temp_dir) / "source.docx"
        create_regression_sample(source_path)
        _add_styled_anchor(source_path)
        round_result = run_document_round(
            source_path,
            identity_transform,
            round_number=1,
            prompt_profile="cn",
        )
        output_path = Path(str(round_result["output_path"]))
        compare_payload = app_service.read_round_compare(str(output_path))
        chunk_id, paragraph_index, anchor_text = _target_fixture(output_path, compare_payload)
        effective_text = _materialized_target_text(output_path, compare_payload, chunk_id)

        compare_path = app_service._find_compare_path_for_output(output_path)
        review_path = app_service._find_review_decisions_path_for_output(output_path)
        compare_before_review_probe = compare_path.read_bytes()
        review_before_review_probe = review_path.read_bytes() if review_path.exists() else None
        for label, invalid_custom in (
            ("missing", effective_text.replace(anchor_text, "", 1)),
            ("duplicated", effective_text.replace(anchor_text, anchor_text + anchor_text, 1)),
        ):
            try:
                app_service.save_review_decisions(
                    str(output_path),
                    {
                        chunk_id: {
                            "mode": "custom",
                            "text": invalid_custom,
                            "source": "review_editor",
                            "confirmed": True,
                        }
                    },
                )
            except ValueError as exc:
                _assert("格式锚点" in str(exc) or "format-sensitive" in str(exc), f"custom {label} error lost anchor reason: {exc}")
            else:
                raise AssertionError(f"custom review save accepted a {label} styled anchor")
            _assert(compare_path.read_bytes() == compare_before_review_probe, f"custom {label} save mutated compare before failure")
            if review_before_review_probe is None:
                _assert(not review_path.exists(), f"custom {label} save published a review sidecar")
            else:
                _assert(review_path.read_bytes() == review_before_review_probe, f"custom {label} save mutated review sidecar")
        checks.append("custom review save blocks styled-anchor drift atomically")

        original_materializer = app_service._materialize_rate_audit_output
        original_builder = app_service._build_transform_from_model_config
        for label, invalid_text in (
            ("missing", effective_text.replace(anchor_text, "", 1)),
            ("ambiguous", effective_text.replace(anchor_text, anchor_text + anchor_text, 1)),
        ):
            constructors: list[dict[str, Any]] = []
            model_calls: list[dict[str, Any]] = []

            def materialize_invalid(*args: Any, **kwargs: Any) -> dict[str, Any]:
                payload = copy.deepcopy(original_materializer(*args, **kwargs))
                for item in payload.get("chunks") or []:
                    if isinstance(item, dict) and str(item.get("chunkId", "") or "") == chunk_id:
                        item["text"] = invalid_text
                return payload

            try:
                app_service._materialize_rate_audit_output = materialize_invalid
                app_service._build_transform_from_model_config = _builder_probe(constructors, model_calls)
                try:
                    app_service.rerun_compare_chunk(str(output_path), chunk_id, {})
                except ValueError as exc:
                    _assert("格式锚点" in str(exc), f"ordinary {label} anchor failure lost its hard-block reason: {exc}")
                else:
                    raise AssertionError(f"ordinary targeted rerun accepted a {label} format anchor")
            finally:
                app_service._materialize_rate_audit_output = original_materializer
                app_service._build_transform_from_model_config = original_builder
            _assert(not constructors, f"ordinary {label} anchor reached transform construction")
            _assert(not model_calls, f"ordinary {label} anchor reached a model call")
        checks.append("ordinary targeted rerun blocks missing/ambiguous styled anchors before transform construction")

        ordinary_constructors: list[dict[str, Any]] = []
        ordinary_calls: list[dict[str, Any]] = []
        try:
            app_service._build_transform_from_model_config = _builder_probe(
                ordinary_constructors,
                ordinary_calls,
            )
            app_service.rerun_compare_chunk(str(output_path), chunk_id, {})
        finally:
            app_service._build_transform_from_model_config = original_builder
        _assert(len(ordinary_constructors) == 1, "ordinary targeted rerun did not construct exactly one transform")
        _assert(1 <= len(ordinary_calls) <= 2, "ordinary targeted rerun escaped its single bounded retry")
        for ordinary_call in ordinary_calls:
            ordinary_model_text = str(ordinary_call.get("chunkText", ""))
            ordinary_prompt_text = str(ordinary_call.get("promptInput", ""))
            _assert(anchor_text not in ordinary_model_text, "ordinary targeted model saw raw styled-anchor text")
            _assert(anchor_text not in ordinary_prompt_text, "ordinary targeted prompt leaked raw styled-anchor text")
            _assert(ordinary_model_text.count("@@FYADR_FMT_") == 1, "ordinary targeted model lost its FMT placeholder")
            _assert(ordinary_prompt_text.count("@@FYADR_FMT_") == 1, "ordinary targeted prompt duplicated its FMT placeholder")
        refreshed_body_map_path = app_service._find_body_map_path_for_output(output_path)
        refreshed_body_map = load_docx_body_map(refreshed_body_map_path) if refreshed_body_map_path else None
        _assert(refreshed_body_map is not None, "ordinary targeted commit lost its body map")
        _assert(
            refreshed_body_map.units[paragraph_index].current_text.count(anchor_text) == 1,
            "ordinary targeted commit did not restore the styled anchor exactly once",
        )
        checks.append("ordinary targeted rerun keeps one FMT placeholder through its bounded retry and restores exact text")

        strategy_compare = app_service.read_round_compare(str(output_path))
        strategy_effective_text = _materialized_target_text(output_path, strategy_compare, chunk_id)
        strategy_effective_texts = {
            str(item.get("chunkId", "") or ""): (
                strategy_effective_text
                if str(item.get("chunkId", "") or "") == chunk_id
                else str(item.get("outputText", item.get("inputText", "")) or "")
            )
            for item in strategy_compare.get("chunks", [])
            if isinstance(item, dict) and str(item.get("chunkId", "") or "")
        }
        strategy_global_style_profile = app_service.build_global_style_profile_from_texts(
            [
                text
                for text in strategy_effective_texts.values()
                if str(text or "").strip()
            ]
        )
        original_strategy_contract = app_service._assert_rate_audit_strategy_model_contract
        original_dimension_converged = app_service._dimension_converged
        try:
            app_service._assert_rate_audit_strategy_model_contract = lambda **_kwargs: "Bound RateAudit prompt"
            app_service._dimension_converged = lambda *_args, **_kwargs: (True, {"regression": True})

            for label, invalid_text in (
                ("missing", strategy_effective_text.replace(anchor_text, "", 1)),
                (
                    "ambiguous",
                    strategy_effective_text.replace(anchor_text, anchor_text + anchor_text, 1),
                ),
            ):
                constructors = []
                model_calls = []
                try:
                    app_service._build_transform_from_model_config = _builder_probe(constructors, model_calls)
                    try:
                        app_service._rerun_rate_audit_strategy_chunk_unlocked(
                            source_path=source_path,
                            output_path=output_path,
                            compare_payload=copy.deepcopy(strategy_compare),
                            chunk_id=chunk_id,
                            effective_input_text=invalid_text,
                            effective_document_texts={
                                **strategy_effective_texts,
                                chunk_id: invalid_text,
                            },
                            dimension_id="transitions",
                            recommended_prompt_id="round2",
                            expected_binding={"planDigest": "regression"},
                            model_config={},
                            global_style_profile=strategy_global_style_profile,
                        )
                    except app_service.StaleRateAuditStrategyPlanError as exc:
                        _assert(
                            "docx_format_anchor_binding_invalid" in exc.mismatch_codes,
                            f"RateAudit {label} anchor failure lost its mismatch code: {exc.mismatch_codes}",
                        )
                    else:
                        raise AssertionError(f"RateAudit strategy accepted a {label} format anchor")
                finally:
                    app_service._build_transform_from_model_config = original_builder
                _assert(not constructors, f"RateAudit {label} anchor reached transform construction")
                _assert(not model_calls, f"RateAudit {label} anchor reached a model call")

            strategy_constructors: list[dict[str, Any]] = []
            strategy_calls: list[dict[str, Any]] = []
            try:
                app_service._build_transform_from_model_config = _builder_probe(
                    strategy_constructors,
                    strategy_calls,
                )
                app_service._rerun_rate_audit_strategy_chunk_unlocked(
                    source_path=source_path,
                    output_path=output_path,
                    compare_payload=copy.deepcopy(strategy_compare),
                    chunk_id=chunk_id,
                    effective_input_text=strategy_effective_text,
                    effective_document_texts=strategy_effective_texts,
                    dimension_id="transitions",
                    recommended_prompt_id="round2",
                    expected_binding={"planDigest": "regression"},
                    model_config={},
                    global_style_profile=strategy_global_style_profile,
                )
            finally:
                app_service._build_transform_from_model_config = original_builder
        finally:
            app_service._assert_rate_audit_strategy_model_contract = original_strategy_contract
            app_service._dimension_converged = original_dimension_converged
            app_service._build_transform_from_model_config = original_builder

        _assert(len(strategy_constructors) == 1, "RateAudit strategy did not construct exactly one transform")
        _assert(1 <= len(strategy_calls) <= 2, "RateAudit strategy escaped its single bounded retry")
        for strategy_call in strategy_calls:
            strategy_model_text = str(strategy_call.get("chunkText", ""))
            strategy_prompt_text = str(strategy_call.get("promptInput", ""))
            _assert(anchor_text not in strategy_model_text, "RateAudit strategy model saw raw styled-anchor text")
            _assert(anchor_text not in strategy_prompt_text, "RateAudit strategy prompt leaked raw styled-anchor text")
            _assert(strategy_model_text.count("@@FYADR_FMT_") == 1, "RateAudit strategy lost its FMT placeholder")
            _assert(strategy_prompt_text.count("@@FYADR_FMT_") == 1, "RateAudit strategy prompt duplicated its FMT placeholder")
        checks.append("RateAudit strategy shares the same pre-construction anchor gate and FMT protection")

    report = {"ok": True, "checks": checks}
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
