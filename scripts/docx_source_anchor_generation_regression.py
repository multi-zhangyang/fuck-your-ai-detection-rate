from __future__ import annotations

import hashlib
import json
import os
import stat
import sys
import tempfile
from pathlib import Path
from typing import Any


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402
from docx.shared import Cm  # noqa: E402

import app_service  # noqa: E402
from docx_audit import _paragraph_format_signature  # noqa: E402
from docx_export_regression import create_regression_sample, identity_transform  # noqa: E402
from round_helper import run_document_round  # noqa: E402


REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_source_anchor_generation_regression_report.json"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _build_replacement_generation(source_path: Path, replacement_path: Path) -> None:
    document = Document(str(source_path))
    document.paragraphs[0].text = "并发替换代际 B：该标题绝不能进入 A 的导出结果"
    document.paragraphs[0].runs[0].bold = False
    document.sections[0].top_margin = Cm(5.0)
    document.save(str(replacement_path))


def _read_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    _assert(isinstance(payload, dict), f"JSON evidence is not an object: {path}")
    return payload


def main() -> int:
    checks: list[str] = []
    regression_root = ROOT_DIR / "finish" / "web_exports"
    regression_root.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="docx-source-anchor-", dir=regression_root) as temp_dir:
        work_dir = Path(temp_dir)
        source_path = work_dir / "source-a.docx"
        replacement_path = work_dir / "source-b.docx"
        latest_alias = work_dir / "paper.docx"
        create_regression_sample(source_path)
        _build_replacement_generation(source_path, replacement_path)

        source_a_bytes = source_path.read_bytes()
        source_b_bytes = replacement_path.read_bytes()
        source_a_sha256 = _sha256(source_path)
        source_b_sha256 = _sha256(replacement_path)
        _assert(source_a_sha256 != source_b_sha256, "A and B DOCX fixtures unexpectedly have the same hash")

        round_result = run_document_round(
            source_path,
            identity_transform,
            round_number=1,
            prompt_profile="cn",
        )
        output_path = Path(str(round_result["output_path"]))
        original_source_stat = source_path.stat()
        attempts_before = {path.resolve() for path in work_dir.glob("paper.*.docx")}
        evidence_before = {path.resolve() for path in work_dir.glob("paper.*.evidence.json")}

        original_rebuild = app_service.rebuild_docx_from_body_map_units
        observed_anchor: dict[str, Any] = {}

        def rebuild_during_a_b_a_swap(*args: Any, **kwargs: Any) -> Any:
            anchor_path = Path(str(kwargs.get("source_path", ""))).resolve()
            observed_anchor["path"] = str(anchor_path)
            observed_anchor["sha256"] = _sha256(anchor_path)
            _assert(
                anchor_path.parent.name == "docx_source_anchors",
                "DOCX rebuild read the mutable provenance path instead of a content-addressed anchor",
            )
            _assert(anchor_path.stem == source_a_sha256, "content-addressed anchor name is not generation A")
            _assert(observed_anchor["sha256"] == source_a_sha256, "content-addressed anchor bytes are not generation A")

            source_path.write_bytes(source_b_bytes)
            try:
                return original_rebuild(*args, **kwargs)
            finally:
                # Restore byte-identical A and its original mtime. The ctime is
                # intentionally impossible to restore from user space, so the
                # generation seal must still detect the A -> B -> A exchange.
                source_path.write_bytes(source_a_bytes)
                os.utime(
                    source_path,
                    ns=(original_source_stat.st_atime_ns, original_source_stat.st_mtime_ns),
                )

        try:
            app_service.rebuild_docx_from_body_map_units = rebuild_during_a_b_a_swap
            try:
                app_service.export_round_output(
                    str(output_path),
                    str(latest_alias),
                    "docx",
                )
            except app_service.ExportRoundError as exc:
                failure = exc.export_failure
                _assert(
                    str(failure.get("stage", "")) == "source-anchor",
                    f"A -> B -> A exchange failed at the wrong stage: {failure}",
                )
            else:
                raise AssertionError("A -> B -> A source exchange published a DOCX")
        finally:
            app_service.rebuild_docx_from_body_map_units = original_rebuild
            if source_path.read_bytes() != source_a_bytes:
                source_path.write_bytes(source_a_bytes)
                os.utime(
                    source_path,
                    ns=(original_source_stat.st_atime_ns, original_source_stat.st_mtime_ns),
                )

        _assert(observed_anchor, "negative export never reached the anchored rebuild")
        _assert(not latest_alias.exists(), "blocked source exchange published the mutable latest alias")
        attempts_after = {path.resolve() for path in work_dir.glob("paper.*.docx")}
        evidence_after = {path.resolve() for path in work_dir.glob("paper.*.evidence.json")}
        _assert(attempts_after == attempts_before, "blocked source exchange published an immutable attempt DOCX")
        _assert(evidence_after == evidence_before, "blocked source exchange published a passed evidence manifest")
        _assert(
            not list(work_dir.glob(".*.tmp.docx")),
            "blocked source exchange leaked a staging DOCX",
        )
        checks.append("A -> B -> A source replacement is blocked before immutable publication")

        control = app_service.export_round_output(
            str(output_path),
            str(latest_alias),
            "docx",
        )
        control_path = Path(str(control.get("path", "")))
        manifest_path = Path(str(control.get("evidenceManifestPath", "")))
        anchor_path = Path(str(control.get("sourceAnchorPath", "")))
        _assert(control_path.exists(), "stable control export did not publish an immutable DOCX")
        _assert(manifest_path.exists(), "stable control export did not publish evidence")
        _assert(anchor_path.exists(), "stable control export did not retain its source anchor")
        _assert(_sha256(anchor_path) == source_a_sha256, "stable control anchor does not contain generation A")
        if os.name != "nt":
            _assert(
                stat.S_IMODE(anchor_path.stat().st_mode) == 0o400,
                "stable control anchor must be readable only by its owner",
            )
        _assert(control.get("sourceSha256") == source_a_sha256, "control result is not bound to generation A")
        _assert("source_generation_anchor" in (control.get("checksPerformed") or []), "control omitted its generation-anchor check")

        manifest = _read_json(manifest_path)
        _assert(manifest.get("sourceSha256") == source_a_sha256, "manifest is not bound to generation A")
        _assert(
            manifest.get("provenanceSourcePath") == str(source_path.resolve()),
            "manifest lost the original provenance path",
        )
        report_paths = manifest.get("reports") if isinstance(manifest.get("reports"), dict) else {}
        for report_key in ("auditPath", "ooxmlAuditPath", "formatLockPath", "contentContractPath"):
            report_path = Path(str(report_paths.get(report_key, "")))
            _assert(report_path.exists(), f"control evidence report is missing: {report_key}")
            report = _read_json(report_path)
            _assert(report.get("sourceSha256") == source_a_sha256, f"{report_key} is not bound to generation A")
            _assert(
                report.get("expectedSourceSha256") == source_a_sha256,
                f"{report_key} lost its expected generation binding",
            )
            _assert(report.get("sourceGenerationStable") is True, f"{report_key} observed an unstable source anchor")
            _assert(
                report.get("provenanceSourcePath") == str(source_path.resolve()),
                f"{report_key} lost provenance for generation A",
            )

        source_document = Document(str(source_path))
        export_document = Document(str(control_path))
        _assert(
            export_document.paragraphs[0].text == source_document.paragraphs[0].text,
            "stable control exported protected title text from generation B",
        )
        _assert(
            _paragraph_format_signature(export_document.paragraphs[0]._p)
            == _paragraph_format_signature(source_document.paragraphs[0]._p),
            "stable control exported protected title format from generation B",
        )
        _assert(
            "并发替换代际 B" not in "\n".join(paragraph.text for paragraph in export_document.paragraphs),
            "generation B protected text leaked into the stable generation-A export",
        )
        checks.append("stable control binds rebuild, all audits, contract and evidence to one generation-A anchor")

    report = {"ok": True, "checks": checks}
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
