#!/usr/bin/env python3
"""Security and closure regression for bound RateAudit strategy execution."""

from __future__ import annotations

import copy
import json
import sys
import tempfile
import time
from contextlib import contextmanager
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterator

from docx import Document  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
ASYNC_STATUS_TIMEOUT_SECONDS = 60.0
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import app_service  # noqa: E402
import web_app  # noqa: E402
from chunking import build_manifest, save_manifest  # noqa: E402
from docx_bodymap import (  # noqa: E402
    build_docx_body_map,
    save_docx_body_map,
    update_docx_body_map_texts,
)
from docx_pipeline import ensure_docx_processing_assets  # noqa: E402
from prompt_library import resolve_prompt_path  # noqa: E402


NATURAL_TEXT = (
    "系统把重复校验结果写入缓存，下一次请求可以直接读取已有状态。"
    "队列只处理尚未完成的任务，失败项会保留原因并等待人工确认。"
    "管理员据此定位具体步骤，不必重新执行整条流程。"
)
TRANSITION_RISK_TEXT = (
    "首先，该系统能够提高处理效率。其次，该系统能够优化业务流程。"
    "此外，该方案具有十分重要的现实意义。综上所述，该方案具有重要意义。"
    "第一，系统提升效率；第二，系统优化流程；第三，系统促进发展。"
)
TRANSITION_REPAIRED_TEXT = (
    "该系统能够提高处理效率，也能够优化业务流程。"
    "该方案具有十分重要的现实意义，其重要性由处理结果体现。"
    "系统提升效率、优化流程并促进发展。"
)
RHYTHM_RISK_TEXT = "。".join(["该方法用于完成公开数据集上的模型效果验证"] * 6) + "。"
RHYTHM_STYLE_DELTA_REJECTED_TEXT = (
    "该方法用于完成公开数据集上的模型效果验证。"
    "公开数据集为模型效果验证提供了样本基础。"
    "测试过程围绕同一目标展开，并记录各项结果。"
    "不同批次的结果随后用于核对模型表现。"
    "验证结束后，研究人员汇总了公开数据集上的表现。"
    "上述步骤共同完成了模型效果验证。"
)
RHYTHM_REPAIRED_TEXT = (
    "该方法用于完成公开数据集上的模型效果验证。"
    "基于公开数据集，可以验证该方法的模型效果。"
    "模型效果的验证依托公开数据集，并由该方法完成。"
    "在公开数据集上，该方法完成模型效果验证。"
    "通过该方法，可以在公开数据集上完成模型效果验证。"
    "对于公开数据集上的模型效果，验证工作由该方法完成。"
)
TEMPLATE_RISK_TEXT = (
    "系统把重复校验结果写入缓存，这一做法具有重要意义。"
    "队列只处理尚未完成的任务，为后续工作提供了有力支持。"
    "管理员依据失败原因定位具体步骤，由此为流程复核奠定了坚实基础。"
)
TEMPLATE_REPAIRED_TEXT = (
    "系统把重复校验结果写入缓存，后续请求可以读取已有状态。"
    "队列只处理尚未完成的任务。管理员依据失败原因定位具体步骤并复核流程。"
)
MIXED_MANUAL_EXECUTABLE_RISK_TEXT = (
    "首先，此外，系统依次执行（1）读取数据，（2）校验数据，（3）写入结果，（4）复核记录。"
    "其次，最后，指标包括：准确率；召回率；响应时间；失败数量。"
)
@dataclass
class Fixture:
    source_path: Path
    output_path: Path
    compare_path: Path
    report: dict[str, Any]

    @property
    def binding(self) -> dict[str, Any]:
        value = self.report.get("strategyBinding")
        assert isinstance(value, dict)
        return value

    def request(self) -> dict[str, Any]:
        binding = self.binding
        return {
            "sourcePath": str(self.source_path),
            "outputPath": str(self.output_path),
            "dimensionId": binding["dimensionId"],
            "recommendedPromptId": binding["recommendedPromptId"],
            "compareRevision": binding["compareRevision"],
            "scopeDigest": binding["scopeDigest"],
            "formatDigest": binding["formatDigest"],
            "sourceSha256": binding["sourceSha256"],
            "targetChunkIds": list(binding["targetChunkIds"]),
            "planDigest": binding["planDigest"],
            "modelConfig": {},
        }


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _wait_for_terminal_async_status(client: Any, run_id: str) -> dict[str, Any]:
    deadline = time.monotonic() + ASYNC_STATUS_TIMEOUT_SECONDS
    last_status: dict[str, Any] = {}
    while True:
        response = client.get(f"/api/batch-rerun-status/{run_id}")
        _assert(response.status_code == 200, f"strategy status API returned {response.status_code}: {response.get_data(as_text=True)}")
        payload = response.get_json()
        _assert(isinstance(payload, dict), f"strategy status API returned a non-object payload: {payload!r}")
        _assert(payload.get("runId") == run_id, f"strategy status API returned the wrong runId: {payload}")
        last_status = payload
        if payload.get("completed") is True:
            return payload
        if time.monotonic() >= deadline:
            break
        time.sleep(0.02)
    raise AssertionError(
        f"strategy task did not reach a terminal status within {ASYNC_STATUS_TIMEOUT_SECONDS:.0f}s: {last_status}"
    )


@contextmanager
def _txt_fixture(kind: str = "transitions", *, chunk_count: int = 1) -> Iterator[Fixture]:
    origin_root = ROOT_DIR / "origin"
    finish_root = ROOT_DIR / "finish"
    origin_root.mkdir(parents=True, exist_ok=True)
    finish_root.mkdir(parents=True, exist_ok=True)
    risk_text = {
        "transitions": TRANSITION_RISK_TEXT,
        "rhythm": RHYTHM_RISK_TEXT,
        "templates": TEMPLATE_RISK_TEXT,
        "mixed": MIXED_MANUAL_EXECUTABLE_RISK_TEXT,
    }[kind]
    round_number = 2 if kind in {"transitions", "templates", "mixed"} else 3
    with tempfile.TemporaryDirectory(prefix="rate-strategy-origin-", dir=origin_root) as origin_temp, tempfile.TemporaryDirectory(
        prefix="rate-strategy-output-", dir=finish_root
    ) as output_temp:
        source_path = Path(origin_temp) / f"{kind}.txt"
        output_path = Path(output_temp) / f"{kind}_round.txt"
        source_path.write_text("\n\n".join([NATURAL_TEXT] * chunk_count), encoding="utf-8")
        output_path.write_text("\n\n".join([risk_text] * chunk_count), encoding="utf-8")
        source_status = app_service.get_document_status(str(source_path))
        compare_path = app_service._find_compare_path_for_output(output_path)
        chunks = [
            {
                "chunkId": f"p{index:04d}-c00",
                "paragraphIndex": index,
                "chunkIndex": 0,
                # This fixture represents a later-round identity baseline: the
                # risky current text is already the frozen input for this
                # round, while source_path remains the original document used
                # by document-level RateAudit deltas.
                "inputText": risk_text,
                "outputText": risk_text,
            }
            for index in range(chunk_count)
        ]
        compare_path.write_text(
            json.dumps(
                {
                    "version": 1,
                    "docId": source_status["docId"],
                    "round": round_number,
                    "promptProfile": "cn_custom",
                    "promptSequence": ["prewrite", "round1", "round2"],
                    "updatedAt": "2026-07-18T00:00:00Z",
                    "chunkCount": len(chunks),
                    "paragraphCount": chunk_count,
                    "chunks": chunks,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        report = app_service.get_document_rate_audit(str(source_path), str(output_path))
        expected_dimension = "transitions" if kind == "mixed" else kind
        expected_prompt = {
            "transitions": "round2",
            "rhythm": "round1",
            "templates": "template-repair",
            "mixed": "round2",
        }[kind]
        _assert(report["strategyPlan"]["decision"] == "targeted_rerun", f"{kind} fixture is not targeted")
        _assert(report["strategyPlan"]["dimensionId"] == expected_dimension, f"{kind} dimension drifted")
        _assert(report["strategyPlan"]["recommendedPromptId"] == expected_prompt, f"{kind} prompt drifted")
        _assert(report["strategyBinding"]["ready"] is True, f"{kind} strategy binding is not ready: {report['strategyBinding']}")
        yield Fixture(source_path, output_path, compare_path, report)


@contextmanager
def _docx_fixture() -> Iterator[Fixture]:
    origin_root = ROOT_DIR / "origin"
    finish_root = ROOT_DIR / "finish"
    origin_root.mkdir(parents=True, exist_ok=True)
    finish_root.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="rate-strategy-docx-origin-", dir=origin_root) as origin_temp, tempfile.TemporaryDirectory(
        prefix="rate-strategy-docx-output-", dir=finish_root
    ) as output_temp:
        source_path = Path(origin_temp) / "contract.docx"
        document = Document()
        document.add_heading("测试论文标题", 0)
        document.add_heading("一、引言", level=1)
        document.add_paragraph("摘要：这是摘要内容。")
        document.add_paragraph(NATURAL_TEXT)
        document.save(source_path)

        extracted_path, snapshot_path, _ = ensure_docx_processing_assets(source_path)
        body_map = build_docx_body_map(source_path, snapshot_path=snapshot_path, round_number=1)
        _assert(len(body_map.units) == 2, "DOCX fixture must expose exactly two editable body units")
        manifest = build_manifest(extracted_path.read_text(encoding="utf-8"))
        _assert(bool(manifest.chunks), "DOCX fixture manifest is empty")

        output_path = Path(output_temp) / "contract_round1.txt"
        output_path.write_text("\n\n".join([TRANSITION_RISK_TEXT] * len(body_map.units)), encoding="utf-8")
        manifest_path = output_path.with_name(f"{output_path.stem}_manifest.json")
        body_map_path = output_path.with_name(f"{output_path.stem}_body_map.json")
        save_manifest(manifest, manifest_path)
        save_docx_body_map(
            update_docx_body_map_texts(
                body_map,
                [TRANSITION_RISK_TEXT] * len(body_map.units),
                round_number=1,
            ),
            body_map_path,
        )

        source_status = app_service.get_document_status(str(source_path))
        compare_chunks: list[dict[str, Any]] = []
        for chunk in manifest.chunks:
            paragraph_indexes = list(chunk.paragraph_indices or [chunk.paragraph_index])
            compare_chunks.append(
                {
                    "chunkId": chunk.chunk_id,
                    "paragraphIndex": chunk.paragraph_index,
                    "chunkIndex": chunk.chunk_index,
                    "inputText": "\n\n".join([TRANSITION_RISK_TEXT] * len(paragraph_indexes)),
                    "outputText": "\n\n".join([TRANSITION_RISK_TEXT] * len(paragraph_indexes)),
                }
            )
        compare_path = app_service._find_compare_path_for_output(output_path)
        compare_path.write_text(
            json.dumps(
                {
                    "version": 1,
                    "docId": source_status["docId"],
                    "round": 1,
                    "promptProfile": "cn_custom",
                    "promptSequence": ["prewrite", "round1", "round2"],
                    "updatedAt": "2026-07-18T00:00:00Z",
                    "inputPath": str(extracted_path),
                    "manifestPath": str(manifest_path),
                    "chunkCount": len(compare_chunks),
                    "paragraphCount": manifest.paragraph_count,
                    "chunks": compare_chunks,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        report = app_service.get_document_rate_audit(str(source_path), str(output_path))
        _assert(report["contentContract"]["ready"] is True, "DOCX fixture contract is not ready")
        _assert(report["strategyBinding"]["ready"] is True, f"DOCX fixture binding is not ready: {report['strategyBinding']}")
        yield Fixture(source_path, output_path, compare_path, report)


def _fake_builder(output_factory: Any, capture: list[dict[str, Any]]):
    def builder(_model_config: dict[str, Any]):
        capture.append({"event": "builder"})

        def transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
            capture.append(
                {
                    "event": "model",
                    "chunkText": chunk_text,
                    "promptInput": prompt_input,
                    "round": round_number,
                    "chunkId": chunk_id,
                }
            )
            return output_factory(chunk_text, chunk_id)

        return transform, "regression"

    return builder


def _post_stale(client: Any, payload: dict[str, Any], *, expected_code: str = "stale_strategy_plan") -> dict[str, Any]:
    response = client.post("/api/rate-audit/execute", json=payload)
    body = response.get_json()
    _assert(response.status_code == 409, f"stale strategy request returned HTTP {response.status_code}: {body}")
    _assert(isinstance(body, dict) and body.get("code") == expected_code, f"stale response code drifted: {body}")
    _assert(isinstance(body.get("mismatchCodes"), list) and body["mismatchCodes"], f"stale response lost mismatch codes: {body}")
    return body


def _exercise_stale_gates() -> list[str]:
    checks: list[str] = []
    client = web_app.app.test_client()
    with _txt_fixture("transitions") as fixture:
        request_payload = fixture.request()
        captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        app_service._build_transform_from_model_config = _fake_builder(
            lambda _text, _chunk_id: TRANSITION_REPAIRED_TEXT,
            captures,
        )
        try:
            mutations = {
                "compare": {"compareRevision": "stale-compare-revision"},
                "scope": {"scopeDigest": "0" * 64},
                "format": {"formatDigest": "0" * 64},
                "source": {"sourceSha256": "0" * 64},
                "dimension": {"dimensionId": "rhythm"},
                "prompt": {"recommendedPromptId": "round1"},
                "targets": {"targetChunkIds": ["unknown-target"]},
                "plan": {"planDigest": "0" * 64},
            }
            for label, mutation in mutations.items():
                stale_request = copy.deepcopy(request_payload)
                stale_request.update(mutation)
                body = _post_stale(client, stale_request)
                _assert(body["mismatchCodes"], f"{label} stale request did not identify a mismatch")

            invalid_requests: list[dict[str, Any]] = []
            missing_plan = copy.deepcopy(request_payload)
            missing_plan.pop("planDigest")
            invalid_requests.append(missing_plan)
            for key, value in (
                ("formatDigest", None),
                ("targetChunkIds", [123]),
                ("dimensionId", 123),
                ("compareRevision", 123),
                ("sourcePath", 123),
                ("outputPath", 456),
            ):
                invalid_request = copy.deepcopy(request_payload)
                invalid_request[key] = value
                invalid_requests.append(invalid_request)
            for invalid_request in invalid_requests:
                invalid_response = client.post("/api/rate-audit/execute", json=invalid_request)
                invalid_body = invalid_response.get_json()
                _assert(invalid_response.status_code == 400, f"invalid typed field did not return 400: {invalid_body}")
                _assert(invalid_body.get("code") == "invalid_strategy_request", f"invalid request code drifted: {invalid_body}")

            original_compare = fixture.compare_path.read_bytes()
            compare_payload = json.loads(original_compare.decode("utf-8"))
            compare_payload["staleProbe"] = "content changed without updatedAt"
            fixture.compare_path.write_text(json.dumps(compare_payload, ensure_ascii=False, indent=2), encoding="utf-8")
            try:
                body = _post_stale(client, request_payload)
                _assert("plan_digest_mismatch" in body["mismatchCodes"], "compare byte drift did not invalidate planDigest")
            finally:
                fixture.compare_path.write_bytes(original_compare)
        finally:
            app_service._build_transform_from_model_config = original_builder
        _assert(not captures, f"stale requests reached model construction/call: {captures}")
    checks.append("compare/scope/format/source/dimension/prompt/targets/plan drift returns 409 before model construction")

    with _txt_fixture("transitions") as fixture:
        original_resolver = app_service.resolve_prompt_path
        captures = []
        prompt_copy = fixture.output_path.with_name("bound-round2-prompt.md")
        prompt_copy.write_bytes(
            resolve_prompt_path("round2").read_text(encoding="utf-8").replace("\n", "\r\n").encode("utf-8")
        )

        def resolve_bound_prompt(prompt_id: object) -> Path:
            return prompt_copy if str(prompt_id) == "round2" else original_resolver(prompt_id)

        original_builder = app_service._build_transform_from_model_config
        try:
            app_service.resolve_prompt_path = resolve_bound_prompt
            report = app_service.get_document_rate_audit(str(fixture.source_path), str(fixture.output_path))
            rebound_fixture = Fixture(fixture.source_path, fixture.output_path, fixture.compare_path, report)
            compare_payload = app_service.read_round_compare(str(fixture.output_path))
            app_service._assert_rate_audit_strategy_model_contract(
                source_path=fixture.source_path,
                output_path=fixture.output_path,
                compare_payload=compare_payload,
                target_chunk=compare_payload["chunks"][0],
                expected_binding=rebound_fixture.binding,
                recommended_prompt_id="round2",
            )
            prompt_copy.write_bytes(prompt_copy.read_bytes() + b"\r\n<!-- changed -->\r\n")
            app_service._build_transform_from_model_config = _fake_builder(
                lambda _text, _chunk_id: TRANSITION_REPAIRED_TEXT,
                captures,
            )
            body = _post_stale(client, rebound_fixture.request())
            _assert("plan_digest_mismatch" in body["mismatchCodes"], "prompt content drift did not invalidate planDigest")
        finally:
            app_service.resolve_prompt_path = original_resolver
            app_service._build_transform_from_model_config = original_builder
        _assert(not captures, "prompt file drift reached model construction")
    checks.append("prompt content SHA is part of the bound plan and stale prompt bytes fail closed")

    with _txt_fixture("transitions") as fixture:
        original_request = fixture.request()
        app_service.save_review_decisions(
            str(fixture.output_path),
            {fixture.binding["targetChunkIds"][0]: "rewrite_confirmed"},
        )
        fresh = app_service.get_document_rate_audit(str(fixture.source_path), str(fixture.output_path))
        _assert(fresh["strategyBinding"]["ready"] is False, "review-confirmed target remained executable")
        _assert(fresh["strategyBinding"]["blockedReason"] == "review_locked_target", "review lock reason is not explicit")
        captures = []
        original_builder = app_service._build_transform_from_model_config
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda _text, _chunk_id: TRANSITION_REPAIRED_TEXT,
                captures,
            )
            body = _post_stale(client, original_request)
            _assert("review_locked_target" in body["mismatchCodes"], "stale response omitted review lock")
        finally:
            app_service._build_transform_from_model_config = original_builder
        _assert(not captures, "review-locked target reached model construction")
    checks.append("saved reviewer choices are immutable strategy targets until explicitly cleared")
    return checks


def _exercise_real_prompt_and_evaluator() -> list[str]:
    checks: list[str] = []
    cases = (
        ("transitions", "round2", "connector_detail", TRANSITION_REPAIRED_TEXT),
        ("rhythm", "round1", "sentence_structure", RHYTHM_REPAIRED_TEXT),
        ("templates", "template-repair", "template_expression", TEMPLATE_REPAIRED_TEXT),
    )
    for kind, expected_prompt_id, expected_evaluator, repaired_text in cases:
        with _txt_fixture(kind) as fixture:
            original_output = fixture.output_path.read_text(encoding="utf-8")
            captures: list[dict[str, Any]] = []
            candidate_outputs = (
                [RHYTHM_STYLE_DELTA_REJECTED_TEXT, repaired_text]
                if kind == "rhythm"
                else [repaired_text]
            )
            generated_count = 0

            def next_candidate(_text: str, _chunk_id: str) -> str:
                nonlocal generated_count
                if generated_count >= len(candidate_outputs):
                    raise AssertionError(
                        f"{kind} strategy exceeded its expected bounded candidate attempts"
                    )
                candidate = candidate_outputs[generated_count]
                generated_count += 1
                return candidate

            original_builder = app_service._build_transform_from_model_config
            try:
                app_service._build_transform_from_model_config = _fake_builder(
                    next_candidate,
                    captures,
                )
                result = app_service.execute_rate_audit_strategy(fixture.request(), {})
            finally:
                app_service._build_transform_from_model_config = original_builder

            model_calls = [item for item in captures if item.get("event") == "model"]
            _assert(result["successCount"] == 1 and result["failureCount"] == 0, f"{kind} strategy did not converge: {result}")
            expected_model_calls = 2 if kind == "rhythm" else 1
            _assert(
                len(model_calls) == expected_model_calls == generated_count,
                f"{kind} strategy did not honor its bounded candidate sequence: {captures}",
            )
            real_prompt = resolve_prompt_path(expected_prompt_id).read_text(encoding="utf-8").strip()
            _assert(
                all(real_prompt in item["promptInput"] for item in model_calls),
                f"{kind} did not load the real {expected_prompt_id} prompt on every bounded attempt",
            )
            expected_input = {
                "transitions": TRANSITION_RISK_TEXT,
                "rhythm": RHYTHM_RISK_TEXT,
                "templates": TEMPLATE_RISK_TEXT,
            }[kind]
            _assert(model_calls[0]["chunkText"] == expected_input, f"{kind} did not use effective materialized text")
            result_chunk = result["compare"]["chunks"][0]
            selection = result_chunk.get("candidateSelection") or {}
            _assert(
                selection.get("publishedRewrite") is True,
                f"{kind} did not publish a selector-certified candidate",
            )
            _assert(
                selection.get("modelAttemptCount") == expected_model_calls,
                f"{kind} candidate-selection evidence lost its bounded attempt count: {selection}",
            )
            _assert(
                selection.get("conditionalRetryCount") == expected_model_calls - 1,
                f"{kind} candidate-selection retry evidence drifted: {selection}",
            )
            if kind == "rhythm":
                first_prompt = str(model_calls[0].get("promptInput", ""))
                retry_prompt = str(model_calls[1].get("promptInput", ""))
                _assert(
                    "[CANDIDATE SELECTION RETRY]" not in first_prompt,
                    "rhythm initial candidate was mislabeled as a retry",
                )
                _assert(
                    "[CANDIDATE SELECTION RETRY]" in retry_prompt
                    and "source_relative_style_delta_failed" in retry_prompt
                    and "no_safe_changed_generated_candidate" in retry_prompt,
                    "rhythm retry did not expose the first candidate's source-relative rejection",
                )
                _assert(
                    selection.get("selectedCandidateId") == "model-attempt-2",
                    f"rhythm selector did not publish the bounded repair candidate: {selection}",
                )
                rejected_candidates = [
                    candidate
                    for candidate in selection.get("candidates", [])
                    if isinstance(candidate, dict)
                    and candidate.get("candidateId") == "model-attempt-1"
                ]
                _assert(
                    len(rejected_candidates) == 1
                    and rejected_candidates[0].get("sourceRelativeStyleGuardPassed") is False
                    and "repeated_opening_family_introduced"
                    in (
                        rejected_candidates[0]
                        .get("sourceRelativeStyleDelta", {})
                        .get("blockingIssueCodes", [])
                    ),
                    f"rhythm fixture did not retain the rejected first-attempt evidence: {selection}",
                )
            _assert(result_chunk["rateAuditStrategyPromptId"] == expected_prompt_id, f"{kind} prompt evidence drifted")
            _assert(result_chunk["rateAuditStrategyEvaluatorDimensionId"] == expected_evaluator, f"{kind} evaluator evidence drifted")
            _assert(result_chunk["rerunDimensionConverged"] is True, f"{kind} did not record same-dimension convergence")
            _assert(result_chunk["rateAuditStrategyReviewRequired"] is True, f"{kind} candidate skipped Diff confirmation")
            _assert(fixture.output_path.read_text(encoding="utf-8") != repaired_text, f"{kind} candidate was auto-exported before review")
            _assert(
                fixture.output_path.read_text(encoding="utf-8") == expected_input,
                f"{kind} frozen round baseline was not retained pending review",
            )
            _assert(
                original_output == fixture.output_path.read_text(encoding="utf-8"),
                f"{kind} pending review unexpectedly rewrote the materialized baseline",
            )

            fresh = app_service.get_document_rate_audit(str(fixture.source_path), str(fixture.output_path))
            _assert(fresh["strategyBinding"]["ready"] is False, f"{kind} pending candidate remained executable")
            _assert(fresh["strategyBinding"]["blockedReason"] == "review_pending_target", f"{kind} pending review reason drifted")

            client = web_app.app.test_client()
            post_success_captures: list[dict[str, Any]] = []
            original_builder = app_service._build_transform_from_model_config
            try:
                app_service._build_transform_from_model_config = _fake_builder(
                    lambda _text, _chunk_id: repaired_text,
                    post_success_captures,
                )
                body = _post_stale(client, fixture.request())
                _assert("review_pending_target" in body["mismatchCodes"], f"{kind} stale plan omitted pending review")
            finally:
                app_service._build_transform_from_model_config = original_builder
            _assert(not post_success_captures, f"{kind} pending review allowed a second model call")

            confirmed_decision = "rewrite_confirmed" if kind == "transitions" else "source_confirmed"
            app_service.save_review_decisions(
                str(fixture.output_path),
                {result_chunk["chunkId"]: confirmed_decision},
            )
            confirmed_compare = app_service.read_round_compare(str(fixture.output_path))
            _assert(
                confirmed_compare["chunks"][0].get("rateAuditStrategyReviewRequired") is not True,
                f"{kind} explicit review did not clear the pending flag",
            )
            confirmed = app_service._materialize_rate_audit_output(fixture.output_path)
            confirmed_chunks = confirmed.get("chunks") or []
            expected_confirmed_text = repaired_text if kind == "transitions" else expected_input
            _assert(
                confirmed_chunks and confirmed_chunks[0]["text"] == expected_confirmed_text,
                f"{kind} explicit {confirmed_decision} did not materialize the selected text",
            )
            confirmed_export_path = fixture.output_path.with_name(f"{kind}_confirmed_export.txt")
            app_service.export_round_output(
                str(fixture.output_path),
                str(confirmed_export_path),
                "txt",
            )
            _assert(
                confirmed_export_path.read_text(encoding="utf-8") == expected_confirmed_text,
                f"{kind} explicit {confirmed_decision} did not control the exported text",
            )
    checks.append(
        "transitions loads round2/connector, rhythm loads round1/sentence-structure, "
        "and templates loads template-repair/template-expression"
    )
    checks.append("successful candidates remain review-pending and cannot run/export again before explicit confirmation")
    return checks


def _exercise_manual_and_executable_dual_track() -> list[str]:
    checks: list[str] = []
    with _txt_fixture("mixed", chunk_count=4) as fixture:
        plan = fixture.report.get("strategyPlan") or {}
        _assert(plan.get("decision") == "targeted_rerun", f"mixed plan did not select an executable dimension: {plan}")
        _assert(plan.get("dimensionId") == "transitions", f"manual structure risk starved transitions: {plan}")
        executable_ids = [item.get("dimensionId") for item in plan.get("executableQueue", [])]
        _assert(
            executable_ids
            and executable_ids[0] == "transitions"
            and "structure" not in executable_ids
            and len(executable_ids) == len(set(executable_ids)),
            f"mixed plan execution queue lost priority or admitted a manual-only dimension: {plan}",
        )
        _assert(
            [item.get("dimensionId") for item in plan.get("blockingManualDimensions", [])] == ["structure"],
            f"mixed plan hid its manual-only structure risk: {plan}",
        )
        _assert(plan.get("manualReviewRequired") is True, "mixed plan did not require manual review")
        _assert(plan.get("manualReviewStillRequired") is True, "mixed plan implied the manual risk was resolved")
        _assert(fixture.binding.get("dimensionId") == "transitions", "binding selected more than the queued transition dimension")
        _assert(
            list(fixture.binding.get("targetChunkIds") or []) == list(plan.get("targetChunkIds") or []),
            "binding target set diverged from the single selected executable dimension",
        )

        captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda text, _chunk_id: text.replace("首先，此外，", "").replace("其次，最后，", ""),
                captures,
            )
            result = app_service.execute_rate_audit_strategy(fixture.request(), {})
        finally:
            app_service._build_transform_from_model_config = original_builder

        model_calls = [item for item in captures if item.get("event") == "model"]
        _assert(result.get("successCount") == 4 and result.get("failureCount") == 0, f"mixed transition repair failed: {result}")
        _assert(len(model_calls) == 4, f"mixed plan did not execute exactly one dimension over four targets: {captures}")
        _assert(
            all("[BOUND DIMENSION] transitions / connector_detail / connectorDensity" in item.get("promptInput", "") for item in model_calls),
            "mixed plan sent a manual-only or second executable dimension to the model",
        )
        _assert(result.get("manualReviewRequired") is True, "execution result cleared the pre-existing manual review requirement")
        _assert(result.get("manualReviewStillRequired") is True, "execution result claimed the manual-only risk was resolved")
        _assert(
            [item.get("dimensionId") for item in result.get("blockingManualDimensions", [])] == ["structure"],
            "execution result lost the unresolved manual-only dimension",
        )
    checks.append(
        "a higher-scored manual-only structure risk cannot starve transitions; one bound dimension executes while manual review remains explicit"
    )
    return checks


def _exercise_non_convergence_and_transaction() -> list[str]:
    checks: list[str] = []
    with _txt_fixture("transitions") as fixture:
        original_output = fixture.output_path.read_bytes()
        original_compare = fixture.compare_path.read_bytes()
        captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda text, _chunk_id: text,
                captures,
            )
            result = app_service.execute_rate_audit_strategy(fixture.request(), {})
        finally:
            app_service._build_transform_from_model_config = original_builder
        model_calls = [item for item in captures if item.get("event") == "model"]
        _assert(len(model_calls) == 2, f"non-converged strategy ignored maxAttempts=2: {captures}")
        _assert(
            result["successCount"] == 1 and result["failureCount"] == 0,
            f"preserved-baseline result shape drifted: {result}",
        )
        _assert(fixture.output_path.read_bytes() == original_output, "non-convergence overwrote output text")
        persisted_compare = app_service.read_round_compare(str(fixture.output_path))
        original_compare_payload = json.loads(original_compare.decode("utf-8"))
        _assert(
            persisted_compare["chunks"][0]["outputText"] == original_compare_payload["chunks"][0]["outputText"],
            "non-convergence overwrote compare outputText",
        )
        preserved_chunk = persisted_compare["chunks"][0]
        _assert(
            preserved_chunk == original_compare_payload["chunks"][0],
            "soft no-op mutated the authoritative compare chunk instead of preserving it byte-for-byte semantically",
        )
        _assert(
            "candidateSelection" not in preserved_chunk and "rerunStatus" not in preserved_chunk,
            "soft no-op persisted non-authoritative candidate/status evidence onto the accepted chunk",
        )
        _assert(
            preserved_chunk.get("rateAuditStrategyReviewRequired") is not True,
            "preserved baseline incorrectly created a pending review candidate",
        )
        attempt_entry = persisted_compare.get("rateAuditStrategyAttempts", {}).get(
            f"transitions:{fixture.binding['targetChunkIds'][0]}"
        )
        _assert(isinstance(attempt_entry, dict) and attempt_entry.get("attemptCount") == 2, "attempt ledger did not persist maxAttempts")
        fresh = app_service.get_document_rate_audit(str(fixture.source_path), str(fixture.output_path))
        _assert(fresh["strategyBinding"]["ready"] is False, "attempt-exhausted strategy remained executable")
        _assert(fresh["strategyBinding"]["blockedReason"] == "strategy_attempt_limit", "attempt-limit block reason drifted")
        plateau_plan = fresh.get("strategyPlan") or {}
        _assert(plateau_plan.get("decision") == "manual_review", "attempt exhaustion left the public plan as targeted_rerun")
        _assert(plateau_plan.get("canExecute") is False, "attempt-exhausted public plan remained executable")
        _assert(plateau_plan.get("recommendedPromptId") == "", "attempt-exhausted public plan still recommended a model prompt")
        _assert(plateau_plan.get("hardStop") is True, "attempt exhaustion did not materialize a hard stop")
        _assert(plateau_plan.get("plateauReached") is True, "attempt exhaustion did not materialize plateauReached")
        _assert(plateau_plan.get("plateauReason") == "strategy_attempt_limit", "attempt plateau reason drifted")
        _assert(plateau_plan.get("manualReviewRequired") is True, "attempt plateau did not require manual review")
        _assert(plateau_plan.get("manualReviewStillRequired") is True, "attempt plateau hid its unresolved review state")
        _assert(plateau_plan.get("executableQueue") == [], "attempt plateau still advertised an executable queue")
        _assert((fresh.get("plateau") or {}).get("preservedPreviousText") is True, "plateau report lost preserve-previous evidence")
        _assert(result.get("resultingStrategyDecision") == "manual_review", "execution result did not expose the terminal plan")
        _assert(result.get("plateauReached") is True, "execution result did not expose the terminal plateau")
        _assert(result.get("plateauReason") == "strategy_attempt_limit", "execution result plateau reason drifted")
        _assert(result.get("manualReviewRequired") is True, "execution result hid required manual review")
        post_limit_captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda text, _chunk_id: text,
                post_limit_captures,
            )
            body = _post_stale(web_app.app.test_client(), fixture.request())
            _assert("strategy_attempt_limit" in body["mismatchCodes"], "attempt-limit stale response lost its code")
        finally:
            app_service._build_transform_from_model_config = original_builder
        _assert(not post_limit_captures, "attempt-exhausted retry reached model construction")
    checks.append(
        "attempt-limit non-convergence preserves text, materializes a hard-stop/manual-review plateau, "
        "and blocks later model calls"
    )

    with _txt_fixture("transitions", chunk_count=2) as fixture:
        target_ids = list(fixture.binding["targetChunkIds"])
        _assert(len(target_ids) == 2, f"multi-target fixture did not bind two targets: {target_ids}")
        original_output = fixture.output_path.read_bytes()
        original_compare = fixture.compare_path.read_bytes()
        captures = []
        original_builder = app_service._build_transform_from_model_config

        def mixed_output(text: str, chunk_id: str) -> str:
            return TRANSITION_REPAIRED_TEXT if chunk_id == target_ids[0] else text

        try:
            app_service._build_transform_from_model_config = _fake_builder(mixed_output, captures)
            result = app_service.execute_rate_audit_strategy(fixture.request(), {})
        finally:
            app_service._build_transform_from_model_config = original_builder
        _assert(result["successCount"] == 2, f"mixed candidate decisions did not complete atomically: {result}")
        _assert(result["failureCount"] == 0, f"preserved baseline was mislabeled as a hard failure: {result}")
        # The converged candidate remains pending explicit review, so the
        # materialized file keeps its safe source default for that target. The
        # non-converged target must keep the current review-materialized
        # baseline, not the attempted model output.
        expected_output = "\n\n".join([TRANSITION_RISK_TEXT, TRANSITION_RISK_TEXT])
        _assert(
            fixture.output_path.read_text(encoding="utf-8") == expected_output,
            "multi-target publication did not materialize pending-review default plus preserved baseline",
        )
        persisted_compare = app_service.read_round_compare(str(fixture.output_path))
        original_compare_payload = json.loads(original_compare.decode("utf-8"))
        _assert(
            [item["outputText"] for item in persisted_compare["chunks"]]
            == [TRANSITION_REPAIRED_TEXT, original_compare_payload["chunks"][1]["outputText"]],
            "multi-target compare did not match the atomic selected outputs",
        )
        first_selection = persisted_compare["chunks"][0].get("candidateSelection") or {}
        second_selection = persisted_compare["chunks"][1].get("candidateSelection") or {}
        _assert(first_selection.get("publishedRewrite") is True, "converged target was not published")
        _assert(
            not second_selection
            and persisted_compare["chunks"][1] == original_compare_payload["chunks"][1],
            "non-converged target did not remain an authoritative soft no-op",
        )
        _assert(
            persisted_compare["chunks"][1].get("rateAuditStrategyReviewRequired") is not True,
            "preserved target incorrectly became review-pending",
        )
        second_attempt = persisted_compare.get("rateAuditStrategyAttempts", {}).get(
            f"transitions:{target_ids[1]}"
        )
        _assert(
            isinstance(second_attempt, dict) and second_attempt.get("attemptCount") == 2,
            "preserved target did not advance its bound attempt ledger",
        )
    checks.append(
        "multi-target execution atomically commits each selector decision: converged candidates publish and non-converged targets preserve baseline"
    )
    return checks


def _exercise_docx_contract_failures() -> list[str]:
    checks: list[str] = []
    client = web_app.app.test_client()
    with _docx_fixture() as fixture:
        request_payload = fixture.request()
        body_map_path = fixture.output_path.with_name(f"{fixture.output_path.stem}_body_map.json")
        original_body_map = body_map_path.read_bytes()
        manifest_path = Path(json.loads(fixture.compare_path.read_text(encoding="utf-8"))["manifestPath"])
        original_manifest = manifest_path.read_bytes()
        captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda _text, _chunk_id: TRANSITION_REPAIRED_TEXT,
                captures,
            )
            body_map_path.unlink()
            body = _post_stale(client, request_payload)
            _assert("docx_body_map_missing" in body["mismatchCodes"], f"missing body-map reason drifted: {body}")
            body_map_path.write_bytes(original_body_map)

            tampered = json.loads(original_body_map.decode("utf-8"))
            tampered["units"][0]["target"] = {"kind": "paragraph", "paragraph_index": 0}
            body_map_path.write_text(json.dumps(tampered, ensure_ascii=False, indent=2), encoding="utf-8")
            body = _post_stale(client, request_payload)
            _assert(
                any(code in body["mismatchCodes"] for code in ("strategy_not_ready", "strategy_recompute_failed", "content_contract_not_ready")),
                f"tampered body-map did not fail through the contract gate: {body}",
            )
            body_map_path.write_bytes(original_body_map)

            tampered_manifest = json.loads(original_manifest.decode("utf-8"))
            tampered_manifest["strategyTamperProbe"] = "manifest bytes changed"
            manifest_path.write_text(json.dumps(tampered_manifest, ensure_ascii=False, indent=2), encoding="utf-8")
            body = _post_stale(client, request_payload)
            _assert("plan_digest_mismatch" in body["mismatchCodes"], f"manifest SHA drift did not invalidate plan: {body}")
        finally:
            app_service._build_transform_from_model_config = original_builder
            body_map_path.write_bytes(original_body_map)
            manifest_path.write_bytes(original_manifest)
        _assert(not captures, f"DOCX missing/tampered body-map reached model construction/call: {captures}")
    checks.append("DOCX missing/retargeted body-map or changed manifest returns stale-plan 409 with zero model setup")
    return checks


def _exercise_async_status_api() -> list[str]:
    checks: list[str] = []
    client = web_app.app.test_client()
    with _txt_fixture("transitions") as fixture:
        captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        run_id = ""
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda _text, _chunk_id: TRANSITION_REPAIRED_TEXT,
                captures,
            )
            response = client.post("/api/rate-audit/execute", json=fixture.request())
            payload = response.get_json()
            _assert(response.status_code == 202 and isinstance(payload, dict), f"valid strategy did not return 202: {response.status_code} {payload}")
            run_id = str(payload.get("runId", "") or "")
            _assert(bool(run_id), "valid strategy 202 response omitted runId")
            status = _wait_for_terminal_async_status(client, run_id)
            _assert(status.get("status") == "completed", f"strategy task terminal status drifted: {status}")
            last_event = status.get("lastEvent")
            _assert(
                isinstance(last_event, dict) and last_event.get("phase") == "batch-complete",
                f"strategy task terminal event drifted: {status}",
            )
            _assert(
                status.get("completedCount") == status.get("totalCount") == status.get("successCount") == 1
                and status.get("failureCount") == 0,
                f"strategy task terminal counters drifted: {status}",
            )
            result = status.get("result")
            _assert(isinstance(result, dict), f"BatchRerun-compatible result missing: {status}")
            _assert(
                result.get("completedCount") == result.get("totalCount") == result.get("successCount") == 1
                and result.get("failureCount") == 0,
                f"strategy task result counters drifted: {status}",
            )
            _assert(result.get("runId") == run_id, "strategy task result runId drifted")
            _assert(isinstance(result.get("postAudit"), dict), "strategy task did not return a fresh postAudit")
            public_post_audit = json.dumps(result.get("postAudit"), ensure_ascii=False, sort_keys=True)
            _assert('"excerpt"' not in public_post_audit, "strategy task persisted a RateAudit body excerpt")
            _assert(
                TRANSITION_RISK_TEXT[:24] not in public_post_audit,
                "strategy task postAudit leaked the accepted/source body",
            )
            _assert(result.get("compareRevisionBefore") != result.get("compareRevisionAfter"), "strategy compare revision did not advance")
        finally:
            app_service._build_transform_from_model_config = original_builder
            if run_id:
                with web_app.RUN_REGISTRY_LOCK:
                    web_app.BATCH_RERUN_STATES.pop(run_id, None)
                    web_app.ACTIVE_BATCH_RERUNS_BY_OUTPUT.pop(str(fixture.output_path.resolve()), None)
                web_app.batch_rerun_state_path(run_id).unlink(missing_ok=True)
        _assert(len([item for item in captures if item.get("event") == "model"]) == 1, f"async strategy call count drifted: {captures}")
    checks.append("valid POST returns 202 and existing batch status API exposes the completed compatible result")

    with _txt_fixture("transitions") as fixture:
        race_run_id = ""
        captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        output_lock = app_service.get_output_rerun_lock(fixture.output_path)
        output_lock.acquire()
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda _text, _chunk_id: TRANSITION_REPAIRED_TEXT,
                captures,
            )
            race_response = client.post("/api/rate-audit/execute", json=fixture.request())
            race_payload = race_response.get_json()
            _assert(race_response.status_code == 202, f"race fixture did not register: {race_payload}")
            race_run_id = str(race_payload.get("runId", "") or "")
            compare_payload = json.loads(fixture.compare_path.read_text(encoding="utf-8"))
            compare_payload["workerRaceProbe"] = "changed after 202"
            fixture.compare_path.write_text(json.dumps(compare_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        finally:
            output_lock.release()

        try:
            race_status = _wait_for_terminal_async_status(client, race_run_id)
            _assert(race_status.get("status") == "failed", f"worker did not reject post-202 stale plan: {race_status}")
            last_event = race_status.get("lastEvent") or {}
            _assert(last_event.get("phase") == "strategy-stale", f"worker stale event lost its phase: {race_status}")
            _assert(last_event.get("code") == "stale_strategy_plan", f"worker stale event lost its code: {race_status}")
            _assert(
                "plan_digest_mismatch" in (last_event.get("mismatchCodes") or []),
                f"worker stale event lost its safe mismatch code: {race_status}",
            )
        finally:
            app_service._build_transform_from_model_config = original_builder
            if race_run_id:
                with web_app.RUN_REGISTRY_LOCK:
                    web_app.BATCH_RERUN_STATES.pop(race_run_id, None)
                    web_app.ACTIVE_BATCH_RERUNS_BY_OUTPUT.pop(str(fixture.output_path.resolve()), None)
                web_app.batch_rerun_state_path(race_run_id).unlink(missing_ok=True)
        _assert(not captures, "post-202 stale plan reached model construction")
    checks.append("worker revalidates the immutable plan after taking the output lock and before model setup")

    with _txt_fixture("transitions") as fixture:
        conflict_run_id = ""
        captures: list[dict[str, Any]] = []
        original_builder = app_service._build_transform_from_model_config
        try:
            app_service._build_transform_from_model_config = _fake_builder(
                lambda _text, _chunk_id: TRANSITION_REPAIRED_TEXT,
                captures,
            )
            conflict_run_id, _ = web_app.register_batch_rerun(str(fixture.output_path), 1)
            conflict_response = client.post("/api/rate-audit/execute", json=fixture.request())
            conflict_body = conflict_response.get_json()
            _assert(conflict_response.status_code == 409, f"active batch conflict did not return 409: {conflict_body}")
            _assert(conflict_body.get("code") == "strategy_execution_conflict", f"conflict code drifted: {conflict_body}")
        finally:
            app_service._build_transform_from_model_config = original_builder
            if conflict_run_id:
                with web_app.RUN_REGISTRY_LOCK:
                    web_app.BATCH_RERUN_STATES.pop(conflict_run_id, None)
                    web_app.ACTIVE_BATCH_RERUNS_BY_OUTPUT.pop(str(fixture.output_path.resolve()), None)
                web_app.batch_rerun_state_path(conflict_run_id).unlink(missing_ok=True)
        _assert(not captures, "strategy conflict reached model construction")
    checks.append("an active ordinary batch rerun returns 409 strategy_execution_conflict with zero model setup")
    return checks


def main() -> int:
    checks = [
        *_exercise_stale_gates(),
        *_exercise_real_prompt_and_evaluator(),
        *_exercise_manual_and_executable_dual_track(),
        *_exercise_non_convergence_and_transaction(),
        *_exercise_docx_contract_failures(),
        *_exercise_async_status_api(),
    ]
    report = {"ok": True, "checks": checks}
    report_path = ROOT_DIR / "finish" / "regression" / "rate_audit_strategy_execution_regression_report.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
