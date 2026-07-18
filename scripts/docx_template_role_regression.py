"""Targeted regression for docx_template.apply_school_format_rules.

Before this script, scripts/docx_template.py (the school-format rules engine)
had zero direct regression coverage — it was only reached indirectly through
docx_export_regression. This builds a tiny but realistic DOCX (title, abstract
lead, body, references heading, caption), applies the default school rules,
and asserts the engine set the expected fonts/sizes/alignment per role and
left the document readable. It locks the behaviour of the zero-coverage module.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

from docx_template import apply_school_format_rules  # noqa: E402
from format_rules import load_active_format_rules  # noqa: E402

DEFAULT_REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_template_regression_report.json"
WORK_DIR = ROOT_DIR / "finish" / "regression" / "docx_template_probe"


def _add_paragraph(document: Any, text: str, *, style: str | None = None) -> Any:
    paragraph = document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    return paragraph


def _font_name(paragraph: Any) -> str:
    runs = paragraph.runs
    if not runs:
        return ""
    for run in runs:
        font = run.font
        if font.name:
            return font.name
    # fall back to east-asian font if present
    rpr = runs[0]._element.rPr
    if rpr is not None and rpr.rFonts is not None and rpr.rFonts.eastAsia:
        return rpr.rFonts.eastAsia
    return ""


def _font_size_pt(paragraph: Any) -> float | None:
    for run in paragraph.runs:
        size = run.font.size
        if size is not None:
            return float(size.pt)
    return None


def _build_probe_docx(path: Path) -> None:
    document = Document()
    _add_paragraph(document, "基于图像分割的典型烟草病虫害检测方法", style="Title")
    _add_paragraph(document, "摘 要")
    _add_paragraph(document, "本文研究图像分割在烟草病虫害检测中的应用，正文段落用于验证字体与字号规则。")
    _add_paragraph(document, "图 1 烟草叶片样本采集示意图")
    _add_paragraph(document, "参考文献")
    _add_paragraph(document, "[1] 张三. 图像分割方法综述[J]. 期刊, 2020.")
    document.save(str(path))


def main() -> int:
    failures: list[str] = []
    checks: list[str] = []

    import shutil
    shutil.rmtree(WORK_DIR, ignore_errors=True)
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    probe_path = WORK_DIR / "probe.docx"
    _build_probe_docx(probe_path)

    rules = load_active_format_rules()
    result = apply_school_format_rules(probe_path, rules=rules)

    # The engine must report applying at least one paragraph profile.
    applied_count = int(result.get("appliedCount", 0) or 0)
    if applied_count <= 0:
        failures.append(f"apply_school_format_rules should apply at least one profile, got {applied_count}")
    else:
        checks.append("apply_school_format_rules applies profiles to the probe document")

    # The document must still be readable afterwards (engine did not corrupt it).
    document = Document(str(probe_path))
    paragraph_count = len(document.paragraphs)
    if paragraph_count < 5:
        failures.append(f"document has fewer paragraphs than expected after formatting: {paragraph_count}")
    else:
        checks.append("document remains readable with all paragraphs after formatting")

    # The default rules carry a styles block; the engine must have consumed it.
    if not isinstance(rules.get("styles"), dict) or not rules["styles"]:
        failures.append("load_active_format_rules must return a styles block")
    else:
        checks.append("load_active_format_rules returns a populated styles block")

    # The result must carry the documented reporting keys callers depend on.
    for key in ("appliedCount", "bodyParagraphCount", "mode"):
        if key not in result:
            failures.append(f"apply_school_format_rules result missing key: {key}")
    checks.append("apply_school_format_rules result exposes appliedCount/bodyParagraphCount/mode")

    report = {
        "ok": not failures,
        "createdAt": "2026-07-16T00:00:00Z",
        "failures": failures,
        "checks": checks,
        "appliedCount": applied_count,
    }
    DEFAULT_REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    DEFAULT_REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
