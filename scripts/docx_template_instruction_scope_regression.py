#!/usr/bin/env python3
"""Freeze university-template guidance before DOCX model/compare construction."""

from __future__ import annotations

import hashlib
import json
import os
import sys
import tempfile
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore[import]
from docx.oxml import OxmlElement  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx_bodymap import (  # noqa: E402
    DOCX_BODY_MAP_SCOPE_SIGNATURE_VERSION,
    DOCX_BODY_MAP_VERSION,
    build_docx_body_map,
    write_docx_body_map_input,
)
from docx_pipeline import (  # noqa: E402
    DOCX_SCOPE_DIAGNOSTICS_VERSION,
    DOCX_SNAPSHOT_VERSION,
    DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
    build_docx_scope_diagnostics,
    build_docx_snapshot,
)
from docx_protection_map import build_docx_protection_map  # noqa: E402
from fyadr_round_service import run_round  # noqa: E402


REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_template_instruction_scope_report.json"
REAL_SAMPLE_PATH = (
    ROOT_DIR
    / "finish"
    / "real_model_e2e"
    / "source"
    / "自动化学院本科毕业设计（论文）论文-范例-2026版.docx"
)
PRIOR_V21_REAL_EDITABLE_COUNT = 78
EXPECTED_REAL_EDITABLE_COUNT = 75
REQUIRE_REAL_TEMPLATE_ENV = "FYADR_REQUIRE_REAL_TEMPLATE_FIXTURE"

REAL_FROZEN_TARGETS = (
    {
        "unitId": "u98",
        "unitIndex": 98,
        "paragraphIndex": 108,
        "textPrefix": "注意：工程设计类",
        "requiredReasons": {
            "template_instruction_prefix",
            "template_document_authoring_cue",
            "template_directive_cue",
            "adjacent_structural_heading",
        },
    },
    {
        "unitId": "u259",
        "unitIndex": 259,
        "paragraphIndex": 282,
        "textPrefix": "注意：课题性质为工程设计类",
        "requiredReasons": {
            "template_instruction_prefix",
            "template_document_authoring_cue",
            "template_directive_cue",
            "adjacent_structural_heading",
        },
    },
    {
        "unitId": "u276",
        "unitIndex": 276,
        "paragraphIndex": 299,
        "textPrefix": "致谢是作者",
        "requiredReasons": {
            "inside_acknowledgement_phase",
            "acknowledgement_guidance",
            "adjacent_acknowledgement_heading",
        },
    },
)

SYNTHETIC_TEMPLATE_INSTRUCTION = (
    "提示：工程设计类毕业论文应在本章给出方案筛选依据，并说明预期指标与方案之间的对应关系。"
)
SYNTHETIC_ORDINARY_ATTENTION = (
    "注意：实验过程中需要保持传感器温度稳定，否则采样结果会出现可重复的系统偏差。"
)
SYNTHETIC_BODY = "本研究比较两种控制方案的稳态误差，并记录相同负载条件下的响应时间。"
SYNTHETIC_ACK_GUIDANCE = (
    "致谢用于说明论文完成过程中获得的帮助，学生应结合个人实际情况自行撰写。"
)
SYNTHETIC_ACK_BODY = "感谢指导教师在实验边界核对与资料整理过程中提供帮助。"
SYNTHETIC_REFERENCE_PHASE_PROSE = "参考资料按照作者、题名与出版年份排列，原始著录内容在导出时保持不变。"
SYNTHETIC_BACK_MATTER_PROSE = "附录材料记录补充装配步骤及校验清单，不属于论文正文改写范围。"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _add_anchored_heading(document: Any, text: str, *, bookmark_id: int) -> Any:
    paragraph = document.add_paragraph(style="Heading 1")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), f"fyadr_heading_{bookmark_id}")
    paragraph._p.append(start)
    paragraph.add_run(text)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)
    return paragraph


def _create_synthetic_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("模板指导语范围回归论文", style="Title")
    document.add_paragraph("摘要")
    document.add_paragraph("本文验证模板指导语冻结规则，并保留真实论文正文中的普通注意表达。")
    document.add_paragraph("1 引言", style="Heading 1")
    document.add_paragraph(SYNTHETIC_ORDINARY_ATTENTION)
    document.add_paragraph("1.1 实验条件", style="Heading 2")
    document.add_paragraph(SYNTHETIC_BODY)
    document.add_paragraph("2 方案设计", style="Heading 1")
    document.add_paragraph(SYNTHETIC_TEMPLATE_INSTRUCTION)
    document.add_paragraph("2.1 方案比较", style="Heading 2")
    document.add_paragraph("两种方案采用统一评价指标，并依据相同采样周期计算稳态误差。")
    _add_anchored_heading(document, "参考文献", bookmark_id=42)
    document.add_paragraph(SYNTHETIC_REFERENCE_PHASE_PROSE)
    _add_anchored_heading(document, "附录A 补充校验材料", bookmark_id=43)
    document.add_paragraph(SYNTHETIC_BACK_MATTER_PROSE)
    _add_anchored_heading(document, "致    谢", bookmark_id=41)
    document.add_paragraph(SYNTHETIC_ACK_GUIDANCE)
    document.add_paragraph(SYNTHETIC_ACK_BODY)
    document.add_paragraph("参考文献", style="Heading 1")
    document.add_paragraph("[1] Template instruction scope regression, 2026.")
    document.save(str(path))


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _assert_synthetic_scope(work_dir: Path, checks: list[str]) -> dict[str, Any]:
    source_path = work_dir / "template_instruction_scope.docx"
    _create_synthetic_fixture(source_path)
    snapshot = build_docx_snapshot(source_path)
    by_text = {unit.text: unit for unit in snapshot.units}

    instruction = by_text[SYNTHETIC_TEMPLATE_INSTRUCTION]
    _assert(
        instruction.structural_role == "template_instruction"
        and instruction.edit_eligibility == "protected"
        and not instruction.editable,
        "paraphrased standalone template instruction entered editable scope",
    )
    ordinary_attention = by_text[SYNTHETIC_ORDINARY_ATTENTION]
    _assert(
        ordinary_attention.structural_role == "body_prose"
        and ordinary_attention.edit_eligibility == "eligible"
        and ordinary_attention.editable,
        "ordinary academic paragraph beginning with 注意 was over-frozen",
    )
    acknowledgement_heading = next(
        unit for unit in snapshot.units if unit.text.replace(" ", "") == "致谢"
    )
    _assert(
        acknowledgement_heading.structural_role == "complex_container"
        and acknowledgement_heading.has_semantic_range_anchor
        and not acknowledgement_heading.editable,
        "anchored acknowledgement heading lost its stronger semantic-anchor protection",
    )
    acknowledgement_guidance = by_text[SYNTHETIC_ACK_GUIDANCE]
    _assert(
        acknowledgement_guidance.structural_role == "template_instruction"
        and not acknowledgement_guidance.editable,
        "acknowledgements meta-guidance was not frozen after an anchored heading",
    )
    _assert(
        by_text[SYNTHETIC_ACK_BODY].structural_role == "acknowledgement_body"
        and not by_text[SYNTHETIC_ACK_BODY].editable,
        "real acknowledgement prose lost its existing protection",
    )
    _assert(
        by_text[SYNTHETIC_REFERENCE_PHASE_PROSE].structural_role == "reference_entry"
        and not by_text[SYNTHETIC_REFERENCE_PHASE_PROSE].editable,
        "bookmark-anchored references heading did not switch the following protected phase",
    )
    _assert(
        by_text[SYNTHETIC_BACK_MATTER_PROSE].structural_role == "back_matter"
        and not by_text[SYNTHETIC_BACK_MATTER_PROSE].editable,
        "bookmark-anchored back-matter heading did not switch the following protected phase",
    )

    model_scope = "\n\n".join(snapshot.editable_texts())
    _assert(SYNTHETIC_ORDINARY_ATTENTION in model_scope, "ordinary 注意 paragraph left model scope")
    for protected_text in (
        SYNTHETIC_TEMPLATE_INSTRUCTION,
        SYNTHETIC_ACK_GUIDANCE,
        SYNTHETIC_ACK_BODY,
        SYNTHETIC_REFERENCE_PHASE_PROSE,
        SYNTHETIC_BACK_MATTER_PROSE,
    ):
        _assert(protected_text not in model_scope, f"protected synthetic text entered model scope: {protected_text}")

    diagnostics = build_docx_scope_diagnostics(snapshot)
    _assert(
        diagnostics.get("version") == DOCX_SCOPE_DIAGNOSTICS_VERSION == 5,
        "template-instruction scope diagnostics schema did not advance to v5",
    )
    _assert(
        diagnostics.get("templateInstructionUnitCount") == 2
        and diagnostics.get("editableTemplateInstructionUnitCount") == 0,
        "scope diagnostics did not expose two frozen synthetic template instructions",
    )
    checks.append("paraphrased template/acknowledgement guidance freezes while ordinary paragraph-initial 注意 remains editable")
    checks.append("bookmark-anchored 致谢 still switches the following structural phase without weakening anchor protection")
    checks.append("bookmark-anchored references and back-matter headings switch their following phases before per-unit anchor protection")
    return {
        "editableUnitCount": snapshot.editable_unit_count,
        "templateInstructionUnitCount": diagnostics.get("templateInstructionUnitCount"),
        "ordinaryAttentionEditable": ordinary_attention.editable,
        "anchoredAcknowledgementRole": acknowledgement_heading.structural_role,
    }


def _assert_real_sample_scope(work_dir: Path, checks: list[str]) -> dict[str, Any]:
    if not REAL_SAMPLE_PATH.exists():
        _assert(
            os.environ.get(REQUIRE_REAL_TEMPLATE_ENV, "").strip() != "1",
            "required private university thesis fixture is missing",
        )
        # The exact university template is deliberately excluded from the
        # public repository. Synthetic coverage above remains mandatory in a
        # clean checkout; only this private 396/75-unit inventory contract is
        # skipped unless its local fixture is available.
        checks.append(
            "private real university template fixture unavailable; exact inventory contract explicitly skipped"
        )
        return {
            "available": False,
            "executed": False,
            "skipReason": "private_fixture_not_available",
        }

    snapshot = build_docx_snapshot(REAL_SAMPLE_PATH)
    _assert(snapshot.version == DOCX_SNAPSHOT_VERSION == 22, "real sample did not use snapshot v22")
    _assert(
        snapshot.structural_role_policy_version == DOCX_STRUCTURAL_ROLE_POLICY_VERSION == 6,
        "real sample did not use structural-role policy v6",
    )
    _assert(snapshot.total_text_unit_count == 396, "real sample top-level structural inventory drifted")
    _assert(
        snapshot.editable_unit_count == EXPECTED_REAL_EDITABLE_COUNT,
        f"real sample editable scope is not the expected 75 units: {snapshot.editable_unit_count}",
    )

    frozen_texts: list[str] = []
    frozen_evidence: list[dict[str, Any]] = []
    for expected in REAL_FROZEN_TARGETS:
        unit = snapshot.units[int(expected["unitIndex"])]
        reasons = set(unit.edit_eligibility_evidence.get("reasonCodes", []))
        _assert(unit.unit_index == expected["unitIndex"], f"real target unit index drifted: {expected['unitId']}")
        _assert(
            unit.target.get("paragraph_index") == expected["paragraphIndex"],
            f"real target paragraph index drifted: {expected['unitId']}",
        )
        _assert(unit.text.startswith(str(expected["textPrefix"])), f"real target text identity drifted: {expected['unitId']}")
        _assert(
            unit.structural_role == "template_instruction"
            and unit.edit_eligibility == "protected"
            and not unit.editable,
            f"real template instruction remained editable: {expected['unitId']}",
        )
        _assert(
            set(expected["requiredReasons"]).issubset(reasons),
            f"real target evidence is incomplete for {expected['unitId']}: {sorted(reasons)}",
        )
        frozen_texts.append(unit.text)
        frozen_evidence.append(
            {
                "unitId": expected["unitId"],
                "unitIndex": unit.unit_index,
                "paragraphIndex": unit.target.get("paragraph_index"),
                "structuralRole": unit.structural_role,
                "editEligibility": unit.edit_eligibility,
                "reasonCodes": sorted(reasons),
                "sourceTextSha256": _sha256_text(unit.source_text()),
            }
        )

    anchored_acknowledgement_heading = snapshot.units[275]
    _assert(
        anchored_acknowledgement_heading.target.get("paragraph_index") == 298
        and anchored_acknowledgement_heading.structural_role == "complex_container"
        and anchored_acknowledgement_heading.has_semantic_range_anchor,
        "real anchored acknowledgement heading protection drifted",
    )
    _assert(
        snapshot.units[81].editable and "需要特别注意" in snapshot.units[81].text,
        "ordinary real thesis 注意 usage was over-frozen",
    )

    # Build through the production body-map path. The source-derived v22 cache
    # is authoritative; the regression never fabricates or filters its units.
    body_map = build_docx_body_map(REAL_SAMPLE_PATH, round_number=1)
    _assert(body_map.version == DOCX_BODY_MAP_VERSION == 9, "real body-map schema did not advance to v9")
    _assert(
        body_map.scope_signature.get("version") == DOCX_BODY_MAP_SCOPE_SIGNATURE_VERSION == 5,
        "real body-map scope signature did not advance to v5",
    )
    _assert(len(body_map.units) == EXPECTED_REAL_EDITABLE_COUNT, "real body-map did not retain exactly 75 editable units")
    body_map_unit_indexes = {unit.unit_index for unit in body_map.units}
    for expected in REAL_FROZEN_TARGETS:
        _assert(
            int(expected["unitIndex"]) not in body_map_unit_indexes,
            f"frozen real unit entered the body map: {expected['unitId']}",
        )

    diagnostics = build_docx_scope_diagnostics(snapshot)
    _assert(
        diagnostics.get("templateInstructionUnitCount") == 3
        and diagnostics.get("editableTemplateInstructionUnitCount") == 0,
        "real scope diagnostics did not expose three frozen template instructions",
    )
    protection_map = build_docx_protection_map(REAL_SAMPLE_PATH)
    _assert(protection_map["summary"]["editableUnits"] == EXPECTED_REAL_EDITABLE_COUNT, "protection map editable count drifted")
    _assert(
        protection_map["summary"].get("roleCounts", {}).get("template_instruction") == 3,
        "protection map omitted the real template-instruction role count",
    )
    _assert(
        any(
            section.get("structuralRole") == "template_instruction"
            and section.get("structuralRoleLabel") == "模板撰写指导语"
            for section in protection_map.get("sections", [])
            if isinstance(section, dict)
        ),
        "protection map omitted the user-facing template-instruction label",
    )

    input_path = work_dir / "real_scope_input.txt"
    output_path = work_dir / "real_scope_output.txt"
    manifest_path = work_dir / "real_scope_manifest.json"
    write_docx_body_map_input(body_map, input_path)
    model_inputs: list[str] = []

    def model_spy(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
        model_inputs.append(str(chunk_text))
        return str(chunk_text)

    result = run_round(
        doc_id="finish/regression/real-template-scope.docx",
        round_number=1,
        input_path=input_path,
        output_path=output_path,
        manifest_path=manifest_path,
        transform=model_spy,
        prompt_profile="cn",
        max_concurrency=4,
    )
    compare_path = Path(str(result["compare_path"]))
    compare_payload = json.loads(compare_path.read_text(encoding="utf-8"))
    serialized_compare = json.dumps(compare_payload, ensure_ascii=False, sort_keys=True)
    serialized_model_input = "\n\n".join(model_inputs)
    for protected_text in frozen_texts:
        _assert(protected_text not in serialized_compare, "frozen real template text entered compare evidence")
        _assert(protected_text not in serialized_model_input, "frozen real template text reached the model callback")
    _assert(
        int(result.get("input_segment_count", 0)) > 0 and model_inputs,
        "real compare/model exclusion proof executed no production chunks",
    )
    checks.append("real university sample keeps exactly 75 editable units under snapshot v22 / role policy v6")
    checks.append("u98/p108, u259/p282 and u276/p299 are absent from body-map, production compare and model callback input")
    checks.append("protection map and scope diagnostics expose three frozen template instructions with user-facing reasons")
    return {
        "available": True,
        "executed": True,
        "sourcePath": str(REAL_SAMPLE_PATH.relative_to(ROOT_DIR)),
        "totalTextUnitCount": snapshot.total_text_unit_count,
        "priorV21EditableUnitCount": PRIOR_V21_REAL_EDITABLE_COUNT,
        "editableUnitCount": snapshot.editable_unit_count,
        "editableUnitDelta": snapshot.editable_unit_count - PRIOR_V21_REAL_EDITABLE_COUNT,
        "templateInstructionUnitCount": diagnostics.get("templateInstructionUnitCount"),
        "bodyMapEditableUnitCount": len(body_map.units),
        "compareChunkCount": int(result.get("input_segment_count", 0)),
        "modelCallbackCount": len(model_inputs),
        "allFrozenTargetsExcludedFromBodyMap": True,
        "allFrozenTargetsExcludedFromCompare": True,
        "allFrozenTargetsExcludedFromModelInput": True,
        "frozenTargets": frozen_evidence,
        "ordinaryAttentionUnit": {
            "unitId": "u81",
            "unitIndex": 81,
            "paragraphIndex": snapshot.units[81].target.get("paragraph_index"),
            "editable": snapshot.units[81].editable,
            "structuralRole": snapshot.units[81].structural_role,
        },
        "anchoredAcknowledgementHeading": {
            "unitId": "u275",
            "unitIndex": 275,
            "paragraphIndex": 298,
            "structuralRole": anchored_acknowledgement_heading.structural_role,
            "hasSemanticRangeAnchor": anchored_acknowledgement_heading.has_semantic_range_anchor,
        },
    }


def main() -> int:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    checks: list[str] = []
    with tempfile.TemporaryDirectory(
        prefix="docx-template-instruction-v6-",
        dir=REPORT_PATH.parent,
    ) as temp_dir:
        work_dir = Path(temp_dir)
        synthetic = _assert_synthetic_scope(work_dir, checks)
        real_sample = _assert_real_sample_scope(work_dir, checks)

    report = {
        "ok": True,
        "snapshotVersion": DOCX_SNAPSHOT_VERSION,
        "structuralRolePolicyVersion": DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
        "scopeDiagnosticsVersion": DOCX_SCOPE_DIAGNOSTICS_VERSION,
        "bodyMapVersion": DOCX_BODY_MAP_VERSION,
        "scopeSignatureVersion": DOCX_BODY_MAP_SCOPE_SIGNATURE_VERSION,
        "synthetic": synthetic,
        "realSample": real_sample,
        "checks": checks,
    }
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
