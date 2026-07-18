#!/usr/bin/env python3
"""Regression for legacy rerun review materialization and atomic publication."""

from __future__ import annotations

import json
import sys
import tempfile
from contextlib import contextmanager
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterator

from docx import Document  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import app_service  # noqa: E402
from chunking import build_manifest, save_manifest  # noqa: E402
from docx_bodymap import (  # noqa: E402
    build_docx_body_map,
    load_docx_body_map,
    save_docx_body_map,
    update_docx_body_map_texts,
)
from docx_pipeline import ensure_docx_processing_assets  # noqa: E402
from fyadr_round_service import (  # noqa: E402
    _build_candidate_selection_event,
    _evaluate_rewrite_candidate,
)


SOURCE_TEXTS = [
    "第一组实验记录了十二次独立运行，平均耗时为三分钟，所有原始日志均已归档。",
    "第二组实验沿用同一数据集和评价口径，研究人员逐项核对了召回率与误差分布。",
    "第三组实验只调整缓存容量，其余参数保持不变，结果用于说明资源约束的影响。",
    "第四组实验保留为人工复核样本，尚未确认的候选不能被普通重跑覆盖。",
]
OLD_REWRITES = [
    "第一组实验包含十二次独立运行，平均耗时三分钟，所有相关原始日志均已归档。",
    "第二组实验继续采用相同数据集和评价口径，并逐项检查召回率及误差分布。",
    "第三组实验仅改变缓存容量，其他参数维持原值，以观察资源约束造成的影响。",
    "第四组实验作为人工复核样本保留，未确认候选不得由普通重跑直接覆盖。",
]
CUSTOM_TEXT = "第二组实验仍使用原数据集与评价口径；召回率和误差分布由研究人员分别核对。"
NEW_REWRITES = [
    "第一组实验共独立运行十二次，平均每次耗时三分钟；所有原始日志随后均已统一归档。",
    "数据集与评价口径在第二组实验中没有变化，研究人员分别核对召回率和误差分布。",
    "第三组实验把缓存容量作为唯一变量，其他参数保持不变，用于观察资源约束的影响。",
    "第四组实验完成普通重跑后生成新的候选文本。",
]


@dataclass
class Fixture:
    source_path: Path
    output_path: Path
    compare_path: Path
    manifest_path: Path
    body_map_path: Path
    decisions_path: Path
    chunk_ids: list[str]


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _published_selection(
    input_text: str,
    output_text: str,
    chunk_id: str,
    *,
    global_style_profile: dict[str, object],
) -> dict[str, object]:
    neutral = {"id": "neutral", "primaryMetric": ""}
    baseline = _evaluate_rewrite_candidate(
        input_text=input_text,
        output_text=input_text,
        candidate_id="baseline",
        origin="baseline",
        attempt=0,
        hard_valid=True,
        round_dimension=neutral,
        global_style_profile=global_style_profile,
    )
    candidate = _evaluate_rewrite_candidate(
        input_text=input_text,
        output_text=output_text,
        candidate_id="model-attempt-1",
        origin="model",
        attempt=1,
        hard_valid=True,
        round_dimension=neutral,
        global_style_profile=global_style_profile,
    )
    _assert(candidate.get("safetyEligible") is True, f"fixture {chunk_id} is not release eligible")
    return _build_candidate_selection_event(
        chunk_id=chunk_id,
        round_number=1,
        candidates=[baseline, candidate],
        selected=candidate,
        reason_codes=["fixture_production_selected"],
        conditional_retry_count=0,
    )


@contextmanager
def _fixture() -> Iterator[Fixture]:
    origin_root = ROOT_DIR / "origin"
    finish_root = ROOT_DIR / "finish"
    origin_root.mkdir(parents=True, exist_ok=True)
    finish_root.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="legacy-review-origin-", dir=origin_root) as origin_temp, tempfile.TemporaryDirectory(
        prefix="legacy-review-output-",
        dir=finish_root,
    ) as output_temp:
        source_path = Path(origin_temp) / "legacy_review_contract.docx"
        document = Document()
        document.add_heading("实验结果分析", 0)
        document.add_heading("一、引言", level=1)
        document.add_paragraph("摘要")
        for text in SOURCE_TEXTS:
            document.add_paragraph(text)
        document.save(source_path)

        extracted_path, snapshot_path, _ = ensure_docx_processing_assets(source_path)
        body_map = build_docx_body_map(source_path, snapshot_path=snapshot_path, round_number=1)
        _assert(len(body_map.units) == len(SOURCE_TEXTS), f"fixture editable-unit count drifted: {len(body_map.units)}")
        manifest = build_manifest(extracted_path.read_text(encoding="utf-8"))
        _assert(len(manifest.chunks) == len(SOURCE_TEXTS), f"fixture chunk count drifted: {len(manifest.chunks)}")
        global_style_profile = app_service.build_global_style_profile_from_texts(
            [chunk.text for chunk in manifest.chunks]
        )
        source_pattern_profile = global_style_profile["documentPatternBaseline"]

        output_path = Path(output_temp) / "round1.txt"
        compare_path = app_service._find_compare_path_for_output(output_path)
        manifest_path = output_path.with_name(f"{output_path.stem}_manifest.json")
        body_map_path = output_path.with_name(f"{output_path.stem}_body_map.json")
        decisions_path = compare_path.with_name(f"{compare_path.stem}_review_decisions.json")
        initial_materialized = [OLD_REWRITES[0], CUSTOM_TEXT, SOURCE_TEXTS[2], SOURCE_TEXTS[3]]
        output_path.write_text("\n\n".join(initial_materialized), encoding="utf-8")
        save_manifest(manifest, manifest_path)
        save_docx_body_map(
            update_docx_body_map_texts(body_map, initial_materialized, round_number=1),
            body_map_path,
        )

        compare_chunks: list[dict[str, Any]] = []
        for index, chunk in enumerate(manifest.chunks):
            payload: dict[str, Any] = {
                "chunkId": chunk.chunk_id,
                "paragraphIndex": chunk.paragraph_index,
                "chunkIndex": chunk.chunk_index,
                "inputText": chunk.text,
                "outputText": OLD_REWRITES[index],
                "quality": {
                    "needsReview": False,
                    "flags": [],
                    "advisoryFlags": ["dimension_direction_not_effective"] if index == 1 else [],
                },
            }
            if index == 0:
                payload["candidateBaselineText"] = chunk.text
                payload["candidateSelection"] = _published_selection(
                    chunk.text,
                    OLD_REWRITES[index],
                    chunk.chunk_id,
                    global_style_profile=global_style_profile,
                )
                payload.update(
                    {
                        "rateAuditStrategyReviewRequired": True,
                        "rateAuditStrategyPlanDigest": "old-plan-digest",
                        "rateAuditStrategyPromptId": "round1",
                        "rateAuditStrategyEvaluatorDimensionId": "rhythm",
                        "rateAuditStrategyInputSource": "review_materialized_compare",
                        "rateAuditStrategyEffectiveInputSha256": "old-effective-sha",
                        "rerunDimensionConverged": True,
                        "rerunDimensionConvergeDirections": [{"ok": True}],
                    }
                )
            if index == 3:
                payload["rateAuditStrategyReviewRequired"] = True
            compare_chunks.append(payload)

        compare_payload = {
            "version": 2,
            "docId": "legacy-rerun-review-contract",
            "round": 1,
            "promptProfile": "cn_custom",
            "promptSequence": ["prewrite", "round1", "round2"],
            "updatedAt": "2026-07-18T00:00:00Z",
            "inputPath": str(extracted_path),
            "manifestPath": str(manifest_path),
            "outputPath": str(output_path),
            "chunkCount": len(compare_chunks),
            "paragraphCount": manifest.paragraph_count,
            "chunks": compare_chunks,
            "qualitySummary": {},
            "validationEvents": [],
            "sourcePatternProfiles": {
                str(source_pattern_profile["profileSha256"]): source_pattern_profile,
            },
            "sourceRelativeDocumentDelta": app_service.assess_source_relative_document_delta(
                [chunk.text for chunk in manifest.chunks],
                OLD_REWRITES,
            ),
        }
        compare_path.write_text(
            json.dumps(compare_payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        chunk_ids = [chunk.chunk_id for chunk in manifest.chunks]
        app_service.save_review_decisions(
            str(output_path),
            {
                chunk_ids[0]: "rewrite_confirmed",
                chunk_ids[1]: {
                    "mode": "custom",
                    "text": CUSTOM_TEXT,
                    "source": "manual",
                    "confirmed": True,
                },
                chunk_ids[2]: "source_confirmed",
            },
        )
        saved_compare = json.loads(compare_path.read_text(encoding="utf-8"))
        _assert(saved_compare["chunks"][0].get("rateAuditStrategyReviewRequired") is None, "confirmed strategy flag was not cleared")
        _assert(saved_compare["chunks"][3].get("rateAuditStrategyReviewRequired") is True, "pending strategy fixture lost its flag")
        yield Fixture(
            source_path=source_path,
            output_path=output_path,
            compare_path=compare_path,
            manifest_path=manifest_path,
            body_map_path=body_map_path,
            decisions_path=decisions_path,
            chunk_ids=chunk_ids,
        )


def _artifact_bytes(fixture: Fixture) -> dict[Path, bytes]:
    return {
        path: path.read_bytes()
        for path in (
            fixture.compare_path,
            fixture.decisions_path,
            fixture.output_path,
            fixture.body_map_path,
        )
    }


def run_regression() -> dict[str, Any]:
    checks: list[str] = []
    with _fixture() as fixture:
        manifest_before = fixture.manifest_path.read_bytes()
        source_before = fixture.source_path.read_bytes()
        model_inputs: dict[str, list[str]] = {}
        validation_inputs: dict[str, list[str]] = {}
        score_inputs: dict[str, list[str]] = {}
        quality_inputs: dict[str, list[str]] = {}
        dimension_inputs: dict[str, list[str]] = {}
        fail_validation_ids: set[str] = set()
        builder_count = 0
        outputs_by_chunk = {
            chunk_id: NEW_REWRITES[index]
            for index, chunk_id in enumerate(fixture.chunk_ids)
        }

        originals = {
            "_build_transform_from_model_config": app_service._build_transform_from_model_config,
            "validate_chunk_output": app_service.validate_chunk_output,
            "_score_rewrite_output": app_service._score_rewrite_output,
            "_build_chunk_quality": app_service._build_chunk_quality,
            "_dimension_converged": app_service._dimension_converged,
        }

        def fake_builder(_model_config: dict[str, Any]):
            nonlocal builder_count
            builder_count += 1

            def transform(chunk_text: str, _prompt_input: str, _round: int, chunk_id: str) -> str:
                model_inputs.setdefault(chunk_id, []).append(chunk_text)
                return outputs_by_chunk[chunk_id]

            return transform, "regression"

        def capture_validation(input_text: str, _output_text: str, chunk_id: str) -> None:
            validation_inputs.setdefault(chunk_id, []).append(input_text)
            if chunk_id in fail_validation_ids:
                raise ValueError("forced validation failure")

        def capture_score(input_text: str, output_text: str) -> float:
            effective_input_owner = {
                OLD_REWRITES[0]: fixture.chunk_ids[0],
                CUSTOM_TEXT: fixture.chunk_ids[1],
                SOURCE_TEXTS[2]: fixture.chunk_ids[2],
            }
            active_chunk_id = effective_input_owner.get(input_text) or next(reversed(model_inputs))
            score_inputs.setdefault(active_chunk_id, []).append(input_text)
            return 1.0 if input_text == output_text else 0.0

        def capture_quality(input_text: str, _output_text: str, *, round_dimension: Any = None) -> dict[str, Any]:
            active_chunk_id = next(reversed(model_inputs))
            quality_inputs.setdefault(active_chunk_id, []).append(input_text)
            hard_risk = active_chunk_id == fixture.chunk_ids[0]
            return {
                "needsReview": hard_risk,
                "flags": ["machine_like_expression"] if hard_risk else [],
                "advisoryFlags": [],
                "roundDimension": round_dimension or {},
            }

        def capture_dimension(input_text: str, _output_text: str, _dimension: dict[str, Any]):
            active_chunk_id = next(reversed(model_inputs))
            dimension_inputs.setdefault(active_chunk_id, []).append(input_text)
            return True, {"ok": True, "satisfied": True}

        app_service._build_transform_from_model_config = fake_builder
        app_service.validate_chunk_output = capture_validation
        app_service._score_rewrite_output = capture_score
        app_service._build_chunk_quality = capture_quality
        app_service._dimension_converged = capture_dimension
        model_config = {
            "promptProfile": "cn_custom",
            "promptSequence": ["prewrite", "round1", "round2"],
        }
        try:
            first_id, second_id, failed_id, pending_id = fixture.chunk_ids

            # This is the same sequential core used by an ordinary batch task.
            first_result = app_service.rerun_compare_chunk(str(fixture.output_path), first_id, model_config)
            _assert(model_inputs[first_id] == [OLD_REWRITES[0]], "rewrite_confirmed text was not sent to the model")
            _assert(
                OLD_REWRITES[0] in validation_inputs[first_id],
                "hard validation did not use rewrite_confirmed text",
            )
            _assert(
                score_inputs[first_id] and set(score_inputs[first_id]) == {OLD_REWRITES[0]},
                "baseline/candidate scoring did not use rewrite_confirmed text",
            )
            _assert(quality_inputs[first_id] == [OLD_REWRITES[0]], "quality scoring did not use rewrite_confirmed text")
            first_review = app_service.load_review_decisions(str(fixture.output_path))["decisions"]
            _assert(first_id not in first_review, "successful target decision was not invalidated")
            _assert(first_review.get(second_id, {}).get("text") == CUSTOM_TEXT, "another chunk's custom decision was lost")
            _assert(first_review.get(failed_id) == "source_confirmed", "another chunk's source decision was lost")
            first_compare = json.loads(fixture.compare_path.read_text(encoding="utf-8"))
            first_target = next(chunk for chunk in first_compare["chunks"] if chunk["chunkId"] == first_id)
            _assert(not any(key.startswith("rateAuditStrategy") for key in first_target), "new legacy candidate retained current RateAudit evidence")
            archived_strategy = first_target.get("supersededRateAuditStrategyEvidence") or {}
            _assert(archived_strategy.get("rateAuditStrategyPlanDigest") == "old-plan-digest", "old strategy evidence was not archived")
            _assert(
                first_target.get("rerunDimensionConverged") is True
                and first_target.get("candidateSelection", {}).get("publishedRewrite") is True,
                "unified selector did not retain current convergence/selection evidence",
            )
            _assert(first_target.get("rerunDefaultDecision") == "source", "hard-risk target did not persist safe source default")
            first_review_raw = json.loads(fixture.decisions_path.read_text(encoding="utf-8"))
            first_body_raw = json.loads(fixture.body_map_path.read_text(encoding="utf-8"))
            _assert(first_compare["updatedAt"] == first_result["updatedAt"], "compare revision differs from rerun result")
            _assert(first_review_raw["updatedAt"] == first_result["updatedAt"], "review sidecar revision differs from rerun result")
            _assert(first_body_raw["updatedAt"] == first_result["updatedAt"], "body-map revision differs from rerun result")
            expected_after_first = [SOURCE_TEXTS[0], CUSTOM_TEXT, SOURCE_TEXTS[2], SOURCE_TEXTS[3]]
            _assert(fixture.output_path.read_text(encoding="utf-8") == "\n\n".join(expected_after_first), "output ignored remaining decisions")
            first_body_map = load_docx_body_map(fixture.body_map_path)
            _assert(first_body_map is not None and first_body_map.current_texts() == expected_after_first, "body map ignored remaining decisions")
            checks.append("rewrite_confirmed effective text drives model/validation/scoring and success invalidates only its decision")

            second_result = app_service.rerun_compare_chunk(str(fixture.output_path), second_id, model_config)
            _assert(model_inputs[second_id] == [CUSTOM_TEXT], "confirmed custom text was not sent to the model")
            _assert(
                CUSTOM_TEXT in validation_inputs[second_id],
                "hard validation did not use confirmed custom text",
            )
            _assert(
                score_inputs[second_id] and set(score_inputs[second_id]) == {CUSTOM_TEXT},
                "baseline/candidate scoring did not use confirmed custom text",
            )
            _assert(quality_inputs[second_id][-1] == CUSTOM_TEXT, "published quality did not use confirmed custom text")
            second_review = app_service.load_review_decisions(str(fixture.output_path))["decisions"]
            _assert(set(second_review) == {failed_id}, f"sequential reruns left incorrect decisions: {second_review}")
            second_compare = json.loads(fixture.compare_path.read_text(encoding="utf-8"))
            second_target = next(chunk for chunk in second_compare["chunks"] if chunk["chunkId"] == second_id)
            _assert(
                second_target.get("candidateSelection", {}).get("publishedRewrite") is True,
                "confirmed custom baseline did not produce an authoritative selected-candidate event",
            )
            _assert(second_target.get("rerunDefaultDecision") == "rewrite", "safe candidate did not persist rewrite default")
            expected_after_second = [SOURCE_TEXTS[0], NEW_REWRITES[1], SOURCE_TEXTS[2], SOURCE_TEXTS[3]]
            _assert(fixture.output_path.read_text(encoding="utf-8") == "\n\n".join(expected_after_second), "second rerun materialization drifted")
            second_body_map = load_docx_body_map(fixture.body_map_path)
            _assert(second_body_map is not None and second_body_map.current_texts() == expected_after_second, "second body-map materialization drifted")
            _assert(second_result["updatedAt"] != first_result["updatedAt"], "sequential reruns did not advance revision")
            checks.append("custom effective text drives same-dimension closure and sequential multi-chunk publication remains compatible")

            before_failure_output = fixture.output_path.read_bytes()
            before_failure_body = fixture.body_map_path.read_bytes()
            before_failure_review = fixture.decisions_path.read_bytes()
            before_failure_compare = json.loads(fixture.compare_path.read_text(encoding="utf-8"))
            candidate_before_failure = next(
                chunk["outputText"] for chunk in before_failure_compare["chunks"] if chunk["chunkId"] == failed_id
            )
            fail_validation_ids.add(failed_id)
            try:
                app_service.rerun_compare_chunk(str(fixture.output_path), failed_id, model_config)
                raise AssertionError("forced hard-validation failure unexpectedly succeeded")
            except ValueError as exc:
                _assert("failed hard validation" in str(exc), f"unexpected validation failure: {exc}")
            finally:
                fail_validation_ids.discard(failed_id)
            _assert(model_inputs[failed_id][0] == SOURCE_TEXTS[2], "source_confirmed text was not the failed rerun baseline")
            _assert(app_service.load_review_decisions(str(fixture.output_path))["decisions"].get(failed_id) == "source_confirmed", "failure cleared target decision")
            _assert(fixture.output_path.read_bytes() == before_failure_output, "failure changed output text")
            _assert(fixture.body_map_path.read_bytes() == before_failure_body, "failure changed body map")
            _assert(fixture.decisions_path.read_bytes() == before_failure_review, "failure changed review sidecar")
            after_failure_compare = json.loads(fixture.compare_path.read_text(encoding="utf-8"))
            candidate_after_failure = next(
                chunk["outputText"] for chunk in after_failure_compare["chunks"] if chunk["chunkId"] == failed_id
            )
            _assert(candidate_after_failure == candidate_before_failure, "failure replaced the existing compare candidate")
            checks.append("failed rerun preserves target decision, output and body map while retaining compare diagnostics")

            pending_before = _artifact_bytes(fixture)
            builders_before_pending = builder_count
            try:
                app_service.rerun_compare_chunk(str(fixture.output_path), pending_id, model_config)
                raise AssertionError("pending RateAudit strategy candidate was not blocked")
            except ValueError as exc:
                _assert("尚未确认的 RateAudit 定点候选" in str(exc), f"pending strategy error was unclear: {exc}")
            _assert(builder_count == builders_before_pending, "pending strategy target constructed a model")
            _assert(_artifact_bytes(fixture) == pending_before, "pending strategy gate mutated artifacts")
            checks.append("pending RateAudit candidate is blocked before model construction with no provenance mutation")

            # Force each publication stage to fail in turn.  Earlier writes in
            # the transaction, if any, must return to their exact prior bytes.
            original_replace = app_service._replace_file_bytes_atomically
            for target_publish_path in (
                fixture.body_map_path,
                fixture.output_path,
                fixture.decisions_path,
                fixture.compare_path,
            ):
                before_atomic_failure = _artifact_bytes(fixture)
                injected = False

                def flaky_replace(path: Path, payload: bytes, *, expected_path: Path = target_publish_path) -> None:
                    nonlocal injected
                    if Path(path) == expected_path and not injected:
                        injected = True
                        raise OSError(f"forced {expected_path.name} publish failure")
                    original_replace(path, payload)

                app_service._replace_file_bytes_atomically = flaky_replace
                try:
                    try:
                        app_service.rerun_compare_chunk(str(fixture.output_path), failed_id, model_config)
                        raise AssertionError("forced atomic publication failure unexpectedly succeeded")
                    except OSError as exc:
                        _assert("publish failure" in str(exc), f"unexpected atomic failure: {exc}")
                finally:
                    app_service._replace_file_bytes_atomically = original_replace
                _assert(injected, f"atomic failure was not injected for {target_publish_path.name}")
                _assert(
                    _artifact_bytes(fixture) == before_atomic_failure,
                    f"atomic rollback did not restore all four artifacts after {target_publish_path.name}",
                )
                _assert(
                    app_service.load_review_decisions(str(fixture.output_path))["decisions"].get(failed_id) == "source_confirmed",
                    f"rollback lost target decision after {target_publish_path.name}",
                )
            checks.append("failure at every publication stage rolls back compare/review/output/body-map byte-for-byte")

            third_result = app_service.rerun_compare_chunk(str(fixture.output_path), failed_id, model_config)
            _assert(
                failed_id not in app_service.load_review_decisions(str(fixture.output_path))["decisions"],
                "successful source_confirmed rerun did not invalidate its target decision",
            )
            _assert(third_result["chunk"].get("rerunDefaultDecision") == "rewrite", "third candidate default decision drifted")
            expected_after_third = [SOURCE_TEXTS[0], NEW_REWRITES[1], NEW_REWRITES[2], SOURCE_TEXTS[3]]
            _assert(fixture.output_path.read_text(encoding="utf-8") == "\n\n".join(expected_after_third), "source-confirmed success output drifted")
            third_body_map = load_docx_body_map(fixture.body_map_path)
            _assert(third_body_map is not None and third_body_map.current_texts() == expected_after_third, "source-confirmed success body map drifted")
            checks.append("source_confirmed success invalidates its decision and publishes the new safe-default candidate consistently")

            _assert(fixture.manifest_path.read_bytes() == manifest_before, "legacy reruns modified the frozen manifest")
            _assert(fixture.source_path.read_bytes() == source_before, "legacy reruns modified the source DOCX")
            final_compare = json.loads(fixture.compare_path.read_text(encoding="utf-8"))
            for index, chunk in enumerate(final_compare["chunks"]):
                _assert(chunk["inputText"] == SOURCE_TEXTS[index], "legacy rerun modified frozen inputText")
            checks.append("frozen inputText, manifest, source DOCX, scope and format provenance remain unchanged")
        finally:
            for name, value in originals.items():
                setattr(app_service, name, value)

    return {"ok": True, "checks": checks}


def main() -> int:
    report = run_regression()
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
