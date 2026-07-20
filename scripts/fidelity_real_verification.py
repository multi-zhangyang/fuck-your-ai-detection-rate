#!/usr/bin/env python3
"""Real fidelity verification: complex docx, real rewrite, independent byte-level proof.

This is NOT the toy docx_fidelity_lock_regression. It builds a genuinely complex
source docx (multi-run mixed fonts, CJK+Latin, style-inherited paragraphs with no
direct rPr, table + caption, formula placeholder, references, auto-numbering,
image placeholder), runs a real rewrite round, exports once normally and once
with the deprecated school_rules override, then independently (without trusting audit_docx_format_lock)
serializes every editable paragraph's pPr + each run's rPr via lxml, strips w:t,
and sha256-compares source vs both exports. The proof: every editable paragraph's
format signature equals the source byte-for-byte and the deprecated override is
migrated instead of reopening a formatting-mutating product path.
"""

from __future__ import annotations

import hashlib
import json
import os
import sys
import tempfile
from pathlib import Path
from typing import Any

ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402
from lxml import etree  # noqa: E402

import app_service  # noqa: E402
from docx_pipeline import get_docx_snapshot_path  # noqa: E402
from round_helper import run_document_round  # noqa: E402

os.environ.setdefault("FYADR_APP_CONFIG_DIR", tempfile.mkdtemp(prefix="fyadr_real_"))

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


def _set_run_font(run: Any, *, size_pt: float, ascii_font: str, ea_font: str, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), ea_font)
    if bold:
        run.bold = True


def _add_multi_run_body(document: Any, runs: list[tuple[str, float, str, str]]) -> Any:
    """Body paragraph with explicit non-standard 18pt-fixed spacing, first-line indent,
    and multiple runs of mixed fonts/sizes (mimics real CJK+Latin academic text)."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.95)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for text, size, ascii_font, ea_font in runs:
        _set_run_font(p.add_run(text), size_pt=size, ascii_font=ascii_font, ea_font=ea_font)
    return p


def _add_inherited_body(document: Any, text: str) -> Any:
    """A body paragraph that relies on STYLE INHERITANCE — no direct rPr/pPr at all.
    This is the dangerous case: preserve_existing=True used to inject explicit
    bold=False/cs-font, changing bytes. In true preserve mode it must stay empty."""
    return document.add_paragraph(text)


def _create_complex_sample(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    # Non-standard margins (school spec would NOT produce these).
    section.top_margin = Cm(3.6)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.9)

    # --- Front matter (protected) ---
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("复杂格式真实样本论文")
    trun.font.size = Pt(20)
    trun.bold = True

    document.add_paragraph("摘 要", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Editable body: multi-run mixed CJK+Latin+code, explicit non-standard format ---
    _add_multi_run_body(
        document,
        [
            # Deliberately retain one registered mechanical connector in the
            # source.  The fidelity test must exercise a rewrite that the
            # bounded selector can prefer for a real, measurable style gain;
            # a marker-only synonym swap is correctly kept as baseline.
            ("首先，本节研究 ", 14, "FangSong", "仿宋"),
            ("YOLOv8", 14, "Times New Roman", "仿宋"),
            (" 在小目标检测上的迁移效果，训练图像尺寸控制在 ", 14, "FangSong", "仿宋"),
            ("640px", 14, "Consolas", "仿宋"),
            ("，学习率范围为 0.001–0.01。", 14, "FangSong", "仿宋"),
        ],
    )
    # Editable body: style-INHERITED paragraph (no direct rPr) — the trap case.
    _add_inherited_body(document, "这一段没有任何直接格式，完全依赖 Normal 样式继承，是保真模式最危险的边界。")
    # Editable body: long-ish real academic sentence
    _add_multi_run_body(
        document,
        [
            ("实验数据显示，在引入结构偏移与句长突发性扰动后，模型对中文文本的分类准确率从 ", 14, "FangSong", "仿宋"),
            ("87.2%", 14, "Times New Roman", "仿宋"),
            (" 提升到 ", 14, "FangSong", "仿宋"),
            ("91.5%", 14, "Times New Roman", "仿宋"),
            ("，但推理延迟略有上升。", 14, "FangSong", "仿宋"),
        ],
    )

    # --- Auto-numbered body ---
    p = document.add_paragraph("1. 登录注册：通过用户名或手机号创建账户并登录系统，以访问平台的各项管理功能。")
    p.paragraph_format.first_line_indent = Cm(0.95)

    # --- Caption (protected) ---
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run("图 1 检测结果对比")
    crun.bold = True

    # --- Table (protected) ---
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "模型"
    table.cell(0, 1).text = "mAP"
    table.cell(1, 0).text = "YOLOv8"
    table.cell(1, 1).text = "87.2%"

    # --- Formula placeholder (protected) ---
    document.add_paragraph("E = mc²（公式占位）")

    # --- References (protected) ---
    document.add_paragraph("参考文献", style="Heading 1")
    document.add_paragraph("[1] 作者. 标题. 期刊, 2020.")

    # --- Acknowledgement (ends editable body scope) ---
    document.add_paragraph("致 谢", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("感谢导师与同学在论文写作过程中的帮助。")

    document.save(str(path))
    return path


def _paragraph_format_sig(paragraph_element: Any) -> tuple[str, list[str]]:
    """Independent serialization. Returns (pPr_sig, [rPr_sig per run]).

    Uses lxml directly and strips ALL w:t descendants so text replacement alone
    never changes the signature. This is independent of docx_audit to avoid
    circular trust.
    """
    p_el = paragraph_element

    def strip_text(el: Any) -> Any:
        clone = etree.Element(el.tag, attrib=dict(el.attrib))
        for child in el:
            tag_local = etree.QName(child.tag).localname
            if tag_local == "t":
                continue
            clone.append(strip_text(child))
        return clone

    ppr = p_el.find(f"{W}pPr")
    ppr_sig = hashlib.sha256(etree.tostring(strip_text(ppr), encoding="utf-8")).hexdigest() if ppr is not None else "none"
    rpr_sigs: list[str] = []
    for run in p_el.findall(f"{W}r"):
        rpr = run.find(f"{W}rPr")
        if rpr is None:
            rpr_sigs.append("none")
        else:
            rpr_sigs.append(hashlib.sha256(etree.tostring(strip_text(rpr), encoding="utf-8")).hexdigest())
    return ppr_sig, rpr_sigs


def _collect_editable_signatures(doc_path: Path, snapshot_path: Path) -> list[dict[str, Any]]:
    from docx_pipeline import _load_docx_snapshot
    snapshot = _load_docx_snapshot(snapshot_path)
    document = Document(str(doc_path))
    out = []
    for unit in snapshot.units:
        if not unit.editable:
            continue
        paragraph = app_service._resolve_target_paragraph(document, unit.target)
        ppr_sig, rpr_sigs = _paragraph_format_sig(paragraph._p)
        out.append({
            "unit_index": unit.unit_index,
            "paragraph_index": unit.target.get("paragraph_index"),
            "text": paragraph.text,
            "ppr_sig": ppr_sig,
            "rpr_sigs": rpr_sigs,
        })
    return out


def _rewrite_transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
    text = str(chunk_text)
    for src, dst in [
        # Remove the fixture's redundant mechanical transition while keeping
        # the technical claim and every protected fact intact.
        # Keep the candidate out of the document's already-saturated cn.plain
        # opening family.  "该研究" is a bounded demonstrative-subject family,
        # so removing the mechanical sequence marker remains a real style gain
        # without manufacturing a fourth repeated plain opening.
        ("首先，本节研究", "该研究考察"),
        ("实验数据显示", "结果显示"),
        ("在引入结构偏移", "在引入结构化扰动"),
        ("略有上升", "略有增加"),
    ]:
        text = text.replace(src, dst)
    return text or chunk_text


def run() -> dict[str, Any]:
    failures: list[str] = []
    work = ROOT_DIR / "finish" / "regression" / "fidelity_real"
    work.mkdir(parents=True, exist_ok=True)
    sample = work / "complex_source.docx"
    _create_complex_sample(sample)

    # Force config to preserve_original for this run.
    import app_config as ac
    ac.save_app_config({**ac.load_app_config(), "formatMode": "preserve_original"})

    round_result = run_document_round(sample, _rewrite_transform, round_number=1, prompt_profile="cn")
    output_path = Path(str(round_result["output_path"]))
    snapshot_path = get_docx_snapshot_path(sample)

    # Prove that this format-fidelity fixture actually traversed the published
    # rewrite path.  The compare chunk references the authoritative
    # candidate-selection event emitted by the production selector.
    compare_path = Path(str(round_result["compare_path"]))
    compare_payload = json.loads(compare_path.read_text(encoding="utf-8"))
    target_chunks = [
        chunk
        for chunk in compare_payload.get("chunks", [])
        if isinstance(chunk, dict) and "首先，本节研究" in str(chunk.get("inputText", ""))
    ]
    if len(target_chunks) != 1:
        failures.append(
            f"candidate selection: expected one mechanical-risk fixture chunk, found {len(target_chunks)}"
        )
    else:
        target_chunk = target_chunks[0]
        selection = target_chunk.get("candidateSelection")
        if not isinstance(selection, dict):
            failures.append("candidate selection: target chunk has no authoritative decision evidence")
        elif selection.get("publishedRewrite") is not True:
            failures.append(
                "candidate selection: target chunk did not publish the style-improving rewrite "
                f"(decision={selection.get('decision')!r})"
            )
        else:
            candidates = [item for item in selection.get("candidates", []) if isinstance(item, dict)]
            baseline = next((item for item in candidates if item.get("origin") == "baseline"), None)
            selected = next(
                (item for item in candidates if item.get("candidateId") == selection.get("selectedCandidateId")),
                None,
            )
            if not isinstance(baseline, dict) or not isinstance(selected, dict):
                failures.append("candidate selection: baseline/selected candidate evidence is incomplete")
            else:
                baseline_penalty = baseline.get("stylePenalty")
                selected_penalty = selected.get("stylePenalty")
                if not isinstance(baseline_penalty, (int, float)) or not isinstance(selected_penalty, (int, float)):
                    failures.append("candidate selection: style-penalty evidence is not numeric")
                elif selected_penalty >= baseline_penalty:
                    failures.append(
                        "candidate selection: published candidate did not lower measured style risk "
                        f"({baseline_penalty} -> {selected_penalty})"
                    )

    source_sigs = _collect_editable_signatures(sample, snapshot_path)
    if not source_sigs:
        failures.append("no editable units found in complex sample")
        return {"ok": False, "failures": failures}

    # --- MODE 1: preserve_original ---
    export_preserve = work / "complex_export_preserve.docx"
    res_p = app_service.export_round_output(str(output_path), str(export_preserve), "docx")
    if str(res_p.get("formatMode", "")) != "preserve_original":
        failures.append(f"preserve export formatMode={res_p.get('formatMode')!r}, expected preserve_original")
    preserve_sigs = _collect_editable_signatures(export_preserve, snapshot_path)

    sig_mismatches = 0
    text_changed = 0
    for src, exp in zip(source_sigs, preserve_sigs):
        if src["ppr_sig"] != exp["ppr_sig"]:
            sig_mismatches += 1
            failures.append(
                f"preserve: pPr signature changed for unit {src['unit_index']} "
                f"(idx {src['paragraph_index']}): {src['ppr_sig'][:12]} != {exp['ppr_sig'][:12]}"
            )
        if src["rpr_sigs"] != exp["rpr_sigs"]:
            sig_mismatches += 1
            failures.append(
                f"preserve: rPr signatures changed for unit {src['unit_index']} "
                f"(idx {src['paragraph_index']}): {src['rpr_sigs']} != {exp['rpr_sigs']}"
            )
        if src["text"] != exp["text"] and "该研究考察" in exp["text"]:
            text_changed += 1
    if text_changed == 0:
        failures.append("preserve: body text was NOT actually rewritten")
    if sig_mismatches == 0:
        print(f"  [preserve] {len(source_sigs)} editable paragraphs, all pPr/rPr signatures byte-identical to source; {text_changed} text(s) rewritten.")

    # --- Legacy override: school_rules must be migrated to fidelity lock ---
    export_school = work / "complex_export_school.docx"
    res_s = app_service.export_round_output(str(output_path), str(export_school), "docx", "school_rules")
    if str(res_s.get("formatMode", "")) != "preserve_original":
        failures.append(f"legacy override formatMode={res_s.get('formatMode')!r}, expected preserve_original")

    school_sigs = _collect_editable_signatures(export_school, snapshot_path)
    diverged = 0
    for src, sch in zip(source_sigs, school_sigs):
        if src["ppr_sig"] != sch["ppr_sig"] or src["rpr_sigs"] != sch["rpr_sigs"]:
            diverged += 1
    if diverged:
        failures.append(
            f"legacy school_rules override changed {diverged}/{len(source_sigs)} editable paragraph formats; "
            "the product must always preserve source formatting."
        )
    else:
        print("  [legacy override] explicit school_rules request was migrated; all source format signatures stayed identical.")

    # --- Concrete font-size check: both paths must retain explicit 14pt ---
    src_doc = Document(str(sample))
    sch_doc = Document(str(export_school))
    pres_doc = Document(str(export_preserve))
    # Find the multi-run paragraph by the exact published rewrite marker.
    def _first_body_size(doc: Any) -> float | None:
        for p in doc.paragraphs:
            if "该研究考察" in p.text:
                for r in p.runs:
                    if r.text.strip() and r.font.size is not None:
                        return r.font.size.pt
        return None
    pres_size = _first_body_size(pres_doc)
    sch_size = _first_body_size(sch_doc)
    if pres_size is None:
        failures.append(f"preserve: could not read body run size (pres_size={pres_size})")
    elif abs(pres_size - 14.0) > 0.2:
        failures.append(f"preserve: body run size drifted to {pres_size}pt, expected 14pt (source was 14pt)")
    else:
        print(f"  [preserve] body run size = {pres_size}pt (kept source 14pt).")
    if sch_size is None:
        failures.append("legacy override: could not read body run size")
    elif abs(sch_size - 14.0) > 0.2:
        failures.append(f"legacy override changed body run size to {sch_size}pt; expected source 14pt")
    else:
        print(f"  [legacy override] body run size = {sch_size}pt (kept source 14pt).")

    # --- Section margins: every product export keeps the source ---
    def _margins(doc: Any) -> tuple[float, float, float, float]:
        s = doc.sections[0]
        return (float(s.top_margin.cm), float(s.bottom_margin.cm), float(s.left_margin.cm), float(s.right_margin.cm))
    src_m = _margins(src_doc)
    pres_m = _margins(pres_doc)
    if any(abs(a - b) > 0.05 for a, b in zip(src_m, pres_m)):
        failures.append(f"preserve: section margins changed {src_m} -> {pres_m}")
    else:
        print(f"  [preserve] section margins kept = {src_m}")
    return {"ok": not failures, "failures": failures, "editableCount": len(source_sigs), "diverged": diverged}


if __name__ == "__main__":
    r = run()
    if r["ok"]:
        print(f"\nREAL FIDELITY VERIFICATION: PASS (editable={r['editableCount']}, legacy_override_drift={r['diverged']})")
        sys.exit(0)
    print("\nREAL FIDELITY VERIFICATION: FAIL")
    for f in r["failures"]:
        print(f"  - {f}")
    sys.exit(1)
