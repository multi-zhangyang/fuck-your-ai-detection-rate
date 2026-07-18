#!/usr/bin/env python3
"""DOCX fidelity-lock regression.

Asserts that the "preserve original" export mode locks the source document's
format-bearing OOXML (pPr + rPr) byte-for-byte per editable paragraph and only
replaces the text — the core user requirement that paper formatting stays
identical before and after rewriting.

This regression deliberately uses a NON-STANDARD source docx (仿宋 / 四号 14pt /
fixed 18pt line spacing / non-standard 3.5/2.2/2.8/2.8 cm margins) so it would
fail under the school-rules path (which would normalize to 10.5pt / 20pt). In
preserve mode, the exported paragraphs must keep the source 14pt / 18pt and the
format-lock signature must match the source exactly.
"""

from __future__ import annotations

import json
import os
import shutil
import sys
import tempfile
from pathlib import Path
from typing import Any

ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402

from docx_audit import (  # noqa: E402
    _paragraph_format_signature,
    audit_docx_export,
    audit_docx_format_lock,
    audit_docx_ooxml_integrity,
)
from docx_bodymap import load_docx_body_map, validate_docx_body_map  # noqa: E402
from docx_pipeline import (  # noqa: E402
    _load_docx_snapshot,
    _replace_paragraph_text,
    get_docx_snapshot_path,
    rebuild_docx_from_snapshot,
)
import app_service  # noqa: E402
from round_helper import run_document_round  # noqa: E402

# Force a clean app config dir so formatMode defaults to preserve_original.
os.environ.setdefault("FYADR_APP_CONFIG_DIR", tempfile.mkdtemp(prefix="fyadr_fidelity_"))


WHITESPACE_CASES = (
    (
        "full_width_leading",
        "\u3000\u3000全角空格缩进段用于验证源 Word 的边界字符不会进入模型，也不会在恒等改写和导出过程中被折叠。",
    ),
    (
        "nbsp_boundaries",
        "\u00a0\u00a0不间断空格边界段用于验证 NBSP 作为原始排版语义被完整记录，并在正文回填后逐字恢复。\u00a0",
    ),
    (
        "multiple_trailing_spaces",
        "多个尾随空格段用于验证连续半角空格不会被共同信任同一规范化逻辑的文本审计漏掉。   ",
    ),
    (
        "manual_first_line_indent",
        "    人工首行缩进段使用四个手工输入的半角空格，验证它们与段落格式属性分离后仍能原样往返。",
    ),
)


def _add_body_paragraph(document: Any, text: str) -> Any:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Cm(0.9)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(18)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    run.font.name = "FangSong"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "仿宋")
    return paragraph


def _create_nonstandard_sample(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # Front matter (protected): title + 摘要 heading.
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("非标准格式样本论文")
    trun.font.size = Pt(18)
    trun.bold = True

    abs_heading = document.add_paragraph()
    abs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abs_heading.add_run("摘 要").bold = True

    # Editable body paragraphs with explicit non-standard format.
    _add_body_paragraph(
        document,
        "本文研究深度学习模型在中文文本分类任务上的实际表现，并提出一种结合结构偏移与风格迁移的改写方法。"
        "综上所述，相关分析用于验证该方法。",
    )
    _add_body_paragraph(
        document,
        "实验表明，仅依赖同义词替换的改写在新型检测器下基本失效，需要在句长突发性与论证结构上同时扰动。",
    )

    # Back matter (protected): 致谢.
    ack_heading = document.add_paragraph()
    ack_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ack_heading.add_run("致 谢").bold = True
    document.add_paragraph("感谢导师与同学在论文写作过程中的帮助。")

    document.save(str(path))
    return path


def _add_literal_whitespace_paragraph(document: Any, text: str) -> Any:
    """Add a body paragraph whose indentation semantics live in literal text."""

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = None
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(18)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    run.font.name = "FangSong"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "仿宋")
    return paragraph


def _create_whitespace_sample(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("正文空白保真样本", style="Title")
    document.add_paragraph("摘 要")
    for _name, literal_text in WHITESPACE_CASES:
        _add_literal_whitespace_paragraph(document, literal_text)
    document.add_paragraph("致 谢")
    document.save(str(path))
    return path


def _run_whitespace_identity_regression(work_dir: Path) -> dict[str, Any]:
    failures: list[str] = []
    sample_path = work_dir / "whitespace_source.docx"
    snapshot_identity_path = work_dir / "whitespace_snapshot_identity.docx"
    snapshot_rewrite_path = work_dir / "whitespace_snapshot_rewrite.docx"
    export_path = work_dir / "whitespace_identity_export.docx"
    tampered_path = work_dir / "whitespace_stripped_negative.docx"
    tampered_map_path = work_dir / "whitespace_evidence_tampered_body_map.json"
    _create_whitespace_sample(sample_path)

    source_document = Document(str(sample_path))
    source_literals = [literal_text for _name, literal_text in WHITESPACE_CASES]
    source_body_texts = [paragraph.text for paragraph in source_document.paragraphs[2:6]]
    if source_body_texts != source_literals:
        failures.append(f"whitespace fixture did not survive initial DOCX save: {source_body_texts!r}")

    model_inputs: list[str] = []

    def identity_transform(chunk_text: str, _prompt_input: str, _round_number: int, _chunk_id: str) -> str:
        model_inputs.append(str(chunk_text))
        return str(chunk_text)

    round_result = run_document_round(sample_path, identity_transform, round_number=1, prompt_profile="cn")
    snapshot_path = get_docx_snapshot_path(sample_path)
    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        return {"ok": False, "failures": ["whitespace snapshot was not created"]}
    editable_units = snapshot.editable_units()
    if len(editable_units) != len(WHITESPACE_CASES):
        failures.append(f"expected {len(WHITESPACE_CASES)} whitespace units, got {len(editable_units)}")

    for index, (case_name, literal_text) in enumerate(WHITESPACE_CASES):
        if index >= len(editable_units):
            break
        unit = editable_units[index]
        if unit.source_text() != literal_text:
            failures.append(f"snapshot lost literal whitespace for {case_name}: {unit.source_text()!r}")
        if unit.text != literal_text.strip():
            failures.append(f"snapshot model core was not boundary-normalized for {case_name}: {unit.text!r}")

    rebuild_docx_from_snapshot(
        [unit.text for unit in editable_units],
        source_path=sample_path,
        snapshot_path=snapshot_path,
        export_path=snapshot_identity_path,
        preserve_format=True,
    )
    snapshot_identity_document = Document(str(snapshot_identity_path))
    for index, (case_name, literal_text) in enumerate(WHITESPACE_CASES):
        if index >= len(editable_units):
            break
        paragraph_index = int(editable_units[index].target["paragraph_index"])
        if snapshot_identity_document.paragraphs[paragraph_index].text != literal_text:
            failures.append(f"snapshot-only identity rebuild changed {case_name}")

    rewritten_cores = [unit.text.replace("验证", "确认", 1) for unit in editable_units]
    rebuild_docx_from_snapshot(
        rewritten_cores,
        source_path=sample_path,
        snapshot_path=snapshot_path,
        export_path=snapshot_rewrite_path,
        preserve_format=True,
    )
    snapshot_rewrite_document = Document(str(snapshot_rewrite_path))
    for index, (case_name, _literal_text) in enumerate(WHITESPACE_CASES):
        if index >= len(editable_units):
            break
        unit = editable_units[index]
        paragraph_index = int(unit.target["paragraph_index"])
        expected_rewrite = f"{unit.leading_whitespace}{rewritten_cores[index]}{unit.trailing_whitespace}"
        if snapshot_rewrite_document.paragraphs[paragraph_index].text != expected_rewrite:
            failures.append(f"snapshot rewrite did not restore source boundaries for {case_name}")

    input_path = Path(str(round_result["round_context"]["input_text_path"]))
    expected_model_input = "\n\n".join(unit.text for unit in editable_units)
    actual_model_input = input_path.read_text(encoding="utf-8")
    if actual_model_input != expected_model_input:
        failures.append("round input did not contain exactly the boundary-free snapshot cores")
    if any(text[:1].isspace() or text[-1:].isspace() for text in actual_model_input.split("\n\n")):
        failures.append("boundary whitespace leaked into the serialized model input")
    joined_model_calls = "\n\n".join(model_inputs)
    if "\u3000" in joined_model_calls or "\u00a0" in joined_model_calls:
        failures.append("full-width or NBSP boundary evidence leaked into an actual model call")
    model_call_paragraphs = [
        paragraph
        for chunk_text in model_inputs
        for paragraph in chunk_text.split("\n\n")
        if paragraph
    ]
    if any(paragraph[:1].isspace() or paragraph[-1:].isspace() for paragraph in model_call_paragraphs):
        failures.append("manual indentation or trailing spaces leaked into an actual model call")

    body_map_path = Path(str(round_result.get("body_map_path", "")))
    body_map = load_docx_body_map(body_map_path)
    if body_map is None:
        return {"ok": False, "failures": [*failures, "whitespace body map was not created"]}
    for index, (case_name, literal_text) in enumerate(WHITESPACE_CASES):
        if index >= len(body_map.units):
            break
        unit = body_map.units[index]
        if unit.source_text() != literal_text:
            failures.append(f"body map lost whitespace evidence for {case_name}: {unit.source_text()!r}")

    tampered_map_payload = body_map.to_dict()
    for raw_unit in tampered_map_payload.get("units", []):
        if isinstance(raw_unit, dict):
            raw_unit["leading_whitespace"] = ""
            raw_unit["trailing_whitespace"] = ""
    tampered_map_path.write_text(
        json.dumps(tampered_map_payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    tampered_map = load_docx_body_map(tampered_map_path)
    if tampered_map is None:
        failures.append("tampered whitespace body map did not remain parseable for the negative check")
    else:
        tampered_map_report = validate_docx_body_map(
            tampered_map,
            source_path=sample_path,
            snapshot_path=snapshot_path,
        )
        tampered_map_issue_codes = {
            str(issue.get("code", ""))
            for issue in tampered_map_report.get("blockingIssues", [])
            if isinstance(issue, dict)
        }
        if bool(tampered_map_report.get("ok")):
            failures.append("body-map validation accepted stripped boundary-whitespace evidence")
        if not {
            "body_map_scope_signature_mismatch",
            "snapshot_scope_signature_drift",
        }.intersection(tampered_map_issue_codes):
            failures.append(f"tampered whitespace evidence lacked scope-signature proof: {tampered_map_issue_codes}")

    output_path = Path(str(round_result["output_path"]))
    export_result = app_service.export_round_output(str(output_path), str(export_path), "docx")
    if str(export_result.get("overallStatus", "")) != "passed":
        failures.append(f"whitespace identity export did not pass: {export_result}")
    for key in ("auditIssueCount", "ooxmlAuditIssueCount", "formatLockIssueCount", "contentContractIssueCount"):
        if int(export_result.get(key, 0) or 0) != 0:
            failures.append(f"whitespace identity export reported {key}={export_result.get(key)}")

    exported_document = Document(str(export_path))
    for index, (case_name, literal_text) in enumerate(WHITESPACE_CASES):
        if index >= len(editable_units):
            break
        paragraph_index = int(editable_units[index].target["paragraph_index"])
        actual_text = exported_document.paragraphs[paragraph_index].text
        if actual_text != literal_text:
            failures.append(f"identity roundtrip changed {case_name}: {literal_text!r} -> {actual_text!r}")

    for case_index in (2, 3):
        if case_index >= len(editable_units):
            continue
        paragraph_index = int(editable_units[case_index].target["paragraph_index"])
        text_nodes = exported_document.paragraphs[paragraph_index]._p.xpath(".//w:t")
        if not any(node.get(qn("xml:space")) == "preserve" for node in text_nodes):
            failures.append(f"ASCII boundary whitespace case {WHITESPACE_CASES[case_index][0]} lacks xml:space=preserve")

    protected_report = audit_docx_export(
        export_path,
        source_path=sample_path,
        snapshot_path=snapshot_path,
        report_path=work_dir / "whitespace_protected_audit.json",
    )
    ooxml_report = audit_docx_ooxml_integrity(
        export_path,
        source_path=sample_path,
        snapshot_path=snapshot_path,
        report_path=work_dir / "whitespace_ooxml_audit.json",
    )
    format_report = audit_docx_format_lock(
        export_path,
        source_path=sample_path,
        snapshot_path=snapshot_path,
        report_path=work_dir / "whitespace_format_lock.json",
    )
    for label, report in (
        ("protected", protected_report),
        ("ooxml", ooxml_report),
        ("format-lock", format_report),
    ):
        if not bool(report.get("ok")):
            failures.append(f"whitespace identity {label} audit failed: {report.get('issues', [])[:1]}")

    shutil.copy2(export_path, tampered_path)
    tampered_document = Document(str(tampered_path))
    for unit in editable_units:
        paragraph_index = int(unit.target["paragraph_index"])
        paragraph = tampered_document.paragraphs[paragraph_index]
        _replace_paragraph_text(paragraph, paragraph.text.strip())
    tampered_document.save(str(tampered_path))
    negative_targets = app_service._build_docx_text_targets_from_body_map(body_map)
    negative_audit = app_service._audit_exported_docx_text_targets(tampered_path, negative_targets)
    if bool(negative_audit.get("ok")):
        failures.append("literal text-integrity audit accepted a whitespace-stripped DOCX")
    if int(negative_audit.get("issueCount", 0) or 0) < len(WHITESPACE_CASES):
        failures.append(f"whitespace negative audit missed cases: {negative_audit}")

    return {
        "ok": not failures,
        "failures": failures,
        "editableCount": len(editable_units),
        "negativeIssueCount": int(negative_audit.get("issueCount", 0) or 0),
        "modelInputBoundaryFree": actual_model_input == expected_model_input,
    }


def _rewrite_transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
    """Reword deterministically without changing terms/numbers so text differs but meaning holds."""

    text = str(chunk_text)
    replacements = [
        ("本文研究", "本工作探讨"),
        ("提出一种", "构建一种"),
        ("实验表明", "结果显示"),
        ("基本失效", "明显失效"),
        ("综上所述，", ""),
    ]
    for src, dst in replacements:
        text = text.replace(src, dst)
    return text or chunk_text


def _editable_paragraph_index_to_element(document: Any, index: int) -> Any:
    return document.paragraphs[index]._p


def run_regression() -> dict[str, Any]:
    failures: list[str] = []
    work_dir = ROOT_DIR / "finish" / "regression" / "fidelity_lock"
    work_dir.mkdir(parents=True, exist_ok=True)
    sample_path = work_dir / "nonstandard_source.docx"
    export_path = work_dir / "nonstandard_export.docx"

    _create_nonstandard_sample(sample_path)

    # Confirm app config resolves to preserve_original by default.
    resolved_mode = app_service._resolve_docx_format_mode()
    if resolved_mode != "preserve_original":
        failures.append(f"expected default formatMode preserve_original, got {resolved_mode}")

    round_result = run_document_round(sample_path, _rewrite_transform, round_number=1, prompt_profile="cn")
    output_path = Path(str(round_result["output_path"]))

    export_result = app_service.export_round_output(str(output_path), str(export_path), "docx")
    reported_mode = str(export_result.get("formatMode", ""))
    if reported_mode != "preserve_original":
        failures.append(f"export formatMode should be preserve_original, got {reported_mode!r}")

    # 1) Format-lock audit must pass: every editable paragraph keeps source pPr/rPr signature.
    snapshot_path = get_docx_snapshot_path(sample_path)
    lock_report = audit_docx_format_lock(
        export_path,
        source_path=sample_path,
        snapshot_path=snapshot_path,
        report_path=work_dir / "format_lock.json",
    )
    if not bool(lock_report.get("ok")):
        issues = lock_report.get("issues", [])
        failures.append(
            f"format-lock audit failed: {lock_report.get('issueCount')} issues; first={issues[0] if issues else None}"
        )
    if int(export_result.get("formatLockIssueCount", 0) or 0) != 0:
        failures.append(
            f"export reported formatLockIssueCount>0: {export_result.get('formatLockIssueCount')}"
        )

    # 2) Text was actually rewritten (not identity).
    exported = Document(str(export_path))
    exported_texts = [p.text for p in exported.paragraphs]
    source = Document(str(sample_path))
    source_texts = [p.text for p in source.paragraphs]
    body_changed = False
    for src_text, exp_text in zip(source_texts, exported_texts):
        if src_text and exp_text and src_text != exp_text and "本工作探讨" in exp_text:
            body_changed = True
            break
    if not body_changed:
        failures.append("fidelity export did not apply rewritten body text")

    # 3) Non-standard format must NOT be normalized to 10.5pt / 20pt (which school_rules would do).
    for paragraph in exported.paragraphs:
        if not paragraph.text.strip() or "本工作探讨" not in paragraph.text:
            continue
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            size = run.font.size
            if size is None or abs(size.pt - 14.0) > 0.2:
                failures.append(f"fidelity run font size drifted from 14pt: {size}")
            break
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing is None or abs(line_spacing.pt - 18.0) > 0.2:
            failures.append(f"fidelity line spacing drifted from 18pt: {line_spacing}")
        break

    # 4) Page margins (section layout) must be unchanged in preserve mode.
    exported_section = exported.sections[0]
    source_section = source.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        src_cm = float(getattr(source_section, attr).cm)
        exp_cm = float(getattr(exported_section, attr).cm)
        if abs(src_cm - exp_cm) > 0.05:
            failures.append(f"fidelity section {attr} drifted: {src_cm}cm -> {exp_cm}cm")

    # 5) Direct signature sanity: a protected (title) paragraph keeps signature too.
    title_index = 0
    src_sig = _paragraph_format_signature(source.paragraphs[title_index]._p)
    exp_sig = _paragraph_format_signature(exported.paragraphs[title_index]._p)
    if src_sig != exp_sig:
        failures.append("protected title paragraph format signature changed")

    whitespace_report = _run_whitespace_identity_regression(work_dir)
    failures.extend(str(item) for item in whitespace_report.get("failures", []))

    return {
        "ok": not failures,
        "failures": failures,
        "formatLockReport": lock_report,
        "whitespaceReport": whitespace_report,
    }


def main() -> int:
    report = run_regression()
    if report["ok"]:
        whitespace_report = report.get("whitespaceReport", {})
        print(
            "docx_fidelity_lock_regression: PASS "
            f"(whitespace_editable={whitespace_report.get('editableCount', 0)}, "
            f"stripped_negative_issues={whitespace_report.get('negativeIssueCount', 0)})"
        )
        return 0
    print("docx_fidelity_lock_regression: FAIL")
    for failure in report["failures"]:
        print(f"  - {failure}")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
