#!/usr/bin/env python3
"""Offline smoke for the bounded real-thesis E2E runner.

No network call is made.  The fake completion deliberately fails the first
validation attempt for each target and returns a conservative valid edit on
the retry, exercising the report's text-free ``failedAttempts`` diagnostics.
"""

from __future__ import annotations

import copy
import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import sys
import tempfile
import uuid

from docx import Document  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import app_config  # noqa: E402
import app_service  # noqa: E402
import real_thesis_model_e2e as runner  # noqa: E402
import fyadr_round_service as round_service  # noqa: E402


ENV_VALUES = {
    runner.REAL_RUN_ENV: "1",
    runner.BASE_URL_ENV: "https://offline-provider.invalid/v1",
    runner.API_KEY_ENV: "offline-nonsecret-credential-value",
    runner.MODEL_ENV: "offline-fake-model",
    runner.API_TYPE_ENV: "chat_completions",
}

REAL_TEMPLATE_PATH = (
    ROOT_DIR
    / "finish"
    / "real_model_e2e"
    / "source"
    / "自动化学院本科毕业设计（论文）论文-范例-2026版.docx"
)
REAL_TEMPLATE_EXPECTED_TOTAL_UNITS = 396
REAL_TEMPLATE_EXPECTED_EDITABLE_UNITS = 75
REAL_TEMPLATE_EXPECTED_EDITABLE_BOOKMARK_INTERIORS = 60
REQUIRE_REAL_TEMPLATE_ENV = "FYADR_REQUIRE_REAL_TEMPLATE_FIXTURE"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _assert_no_private_report_fields(value: object, *, path: str = "report") -> None:
    forbidden_keys = {
        "outputtext",
        "inputtext",
        "candidatetext",
        "matchedtext",
        "excerpt",
        "preview",
        "error",
        "hardvalidationerror",
        "providermessage",
        "endpoint",
        "reasoning",
        "thinking",
        "analysis",
        "thought",
        "chainofthought",
        "rawprompt",
        "prompttext",
        "promptbody",
        "modelbody",
        "providerbody",
        "rawresponse",
        "responsebody",
    }
    if isinstance(value, dict):
        for raw_key, child in value.items():
            key = str(raw_key)
            normalized_key = re.sub(r"[^a-z0-9]", "", key.casefold())
            _assert(
                normalized_key not in forbidden_keys,
                f"private report field {path}.{key} was retained",
            )
            _assert_no_private_report_fields(child, path=f"{path}.{key}")
    elif isinstance(value, list):
        for index, child in enumerate(value):
            _assert_no_private_report_fields(child, path=f"{path}[{index}]")


def _selection_event(
    baseline: str,
    output: str,
    *,
    published: bool = True,
) -> dict[str, object]:
    neutral_dimension: dict[str, object] = {"id": "neutral", "primaryMetric": ""}
    global_style_profile = app_service.build_global_style_profile_from_texts([baseline])
    baseline_candidate = app_service._evaluate_rewrite_candidate(
        input_text=baseline,
        output_text=baseline,
        candidate_id="baseline",
        origin="baseline",
        attempt=0,
        hard_valid=True,
        round_dimension=neutral_dimension,
        global_style_profile=global_style_profile,
    )
    generated_candidate = app_service._evaluate_rewrite_candidate(
        input_text=baseline,
        output_text=output,
        candidate_id="model-attempt-1",
        origin="model",
        attempt=1,
        hard_valid=True,
        round_dimension=neutral_dimension,
        global_style_profile=global_style_profile,
    )
    if published:
        # For the readability-negative fixture we intentionally publish the
        # generated shape even though its readability flags remain false.  The
        # consumer must reject that exact v2 event by recomputation and by the
        # canonical production release assessment.
        if generated_candidate.get("safetyEligible") is True:
            selected_candidate, _reason_codes = app_service._select_rewrite_candidate(
                [baseline_candidate, generated_candidate],
                round_dimension=neutral_dimension,
            )
            _assert(
                selected_candidate is generated_candidate,
                "valid candidate fixture did not win the production selector",
            )
        else:
            generated_candidate["safetyEligible"] = True
            selected_candidate = generated_candidate
        candidates = [baseline_candidate, generated_candidate]
    else:
        # Exact stale-evidence failure shape: the selector preserved its
        # baseline, while the surrounding chunk later claims a changed output.
        selected_candidate = baseline_candidate
        candidates = [baseline_candidate, generated_candidate]
    event = app_service._build_candidate_selection_event(
        chunk_id="p0_c0",
        round_number=1,
        candidates=candidates,
        selected=selected_candidate,
        reason_codes=["real-e2e-offline-fixture"],
        conditional_retry_count=0,
    )
    _assert(
        event.get("schema") == runner.CANDIDATE_SELECTION_SCHEMA
        and event.get("schemaVersion") == runner.CANDIDATE_SELECTION_VERSION,
        "candidate fixture did not use the production v2 schema",
    )
    _assert(
        event.get("publishedRewrite") is published,
        "candidate fixture publication decision drifted",
    )
    return event


def _candidate_chunk(
    baseline: str,
    output: str,
    *,
    published: bool = True,
    chunk_id: str = "p0_c0",
) -> dict[str, object]:
    return {
        "chunkId": chunk_id,
        "paragraphIndex": 0,
        "chunkIndex": 0,
        "inputText": baseline,
        "outputText": output,
        "candidateBaselineText": baseline.strip(),
        "quality": {"needsReview": False, "flags": []},
        "candidateSelection": _selection_event(baseline, output, published=published),
    }


def _candidate_snapshot(chunk: dict[str, object]) -> dict[str, object]:
    candidate_baseline = str(chunk.get("candidateBaselineText", chunk.get("inputText", "")) or "")
    profile = app_service.build_global_style_profile_from_texts(
        [candidate_baseline]
    )["documentPatternBaseline"]
    return {
        "compare": {
            "chunks": [chunk],
            "sourcePatternProfiles": {
                str(profile["profileSha256"]): profile,
            },
            "sourceRelativeDocumentDelta": app_service.assess_source_relative_document_delta(
                [str(chunk.get("inputText", "") or "")],
                [str(chunk.get("outputText", "") or "")],
            ),
        },
        "_internal": {
            "effectiveChunks": [
                {
                    "chunkId": str(chunk["chunkId"]),
                    "paragraphIndex": 0,
                    "chunkIndex": 0,
                    "text": str(chunk["outputText"]),
                }
            ]
        },
    }


def _assert_snapshot_decision_coverage_contract() -> None:
    baseline = (
        "首先，系统读取用户提交的论文段落并建立任务记录。"
        "其次，服务按照段落边界生成改写单元并保留原有编号。"
        "此外，校验模块核对术语、数值和引用标记。"
        "因此，只有通过事实与格式检查的文本才会进入结果页，审阅者仍需确认每处修改。"
    )
    output = (
        "系统读取用户提交的论文段落并建立任务记录，服务再按照段落边界生成改写单元，同时保留原有编号。"
        "校验模块负责核对术语、数值和引用标记；文本仅在通过事实与格式检查后进入结果页，"
        "审阅者仍需确认每处修改。"
    )
    frozen_text = "第二章介绍实验数据集、模型框架与评估指标。"
    snapshot = _candidate_snapshot(_candidate_chunk(baseline, output))
    snapshot["compare"]["chunks"].append(
        {
            "chunkId": "p1_c0",
            "paragraphIndex": 1,
            "chunkIndex": 0,
            "inputText": frozen_text,
            "outputText": frozen_text,
            "quality": {"needsReview": False, "flags": []},
        }
    )
    snapshot["compare"]["validationEvents"] = [
        {
            "event": "chunk-frozen",
            "chunkId": "p1_c0",
            "reasonCode": "structure_or_metadata_preserved",
        }
    ]
    snapshot["compare"]["sourceRelativeDocumentDelta"] = (
        app_service.assess_source_relative_document_delta(
            [baseline, frozen_text],
            [output, frozen_text],
        )
    )

    summary = runner._summarize_snapshot(snapshot)
    runner._require_v2_snapshot_candidate_evidence(summary)
    evidence = summary.get("candidateSelectionEvidence") or {}
    _assert(int(evidence.get("selectionCount", 0)) == 1, "candidate coverage drifted")
    _assert(int(evidence.get("frozenIdentityCount", 0)) == 1, "frozen identity was not counted")
    _assert(int(evidence.get("decisionCoverageCount", 0)) == 2, "mixed decision coverage is incomplete")
    _assert(evidence.get("frozenIdentityBound") is True, "valid frozen identity failed closed")

    invalid_cases = []
    changed_output = copy.deepcopy(snapshot)
    changed_output["compare"]["chunks"][1]["outputText"] = frozen_text + " "
    invalid_cases.append(("changed frozen output", changed_output))
    missing_event = copy.deepcopy(snapshot)
    missing_event["compare"]["validationEvents"] = []
    invalid_cases.append(("missing frozen event", missing_event))
    duplicate_event = copy.deepcopy(snapshot)
    duplicate_event["compare"]["validationEvents"].append(
        copy.deepcopy(duplicate_event["compare"]["validationEvents"][0])
    )
    invalid_cases.append(("duplicate frozen event", duplicate_event))
    wrong_reason = copy.deepcopy(snapshot)
    wrong_reason["compare"]["validationEvents"][0]["reasonCode"] = "forged-reason"
    invalid_cases.append(("wrong frozen reason", wrong_reason))

    for label, invalid_snapshot in invalid_cases:
        try:
            runner._require_v2_snapshot_candidate_evidence(
                runner._summarize_snapshot(invalid_snapshot)
            )
        except runner.E2EContractError as exc:
            _assert(
                exc.code == "candidate_selection_evidence_incomplete",
                f"{label} failed with the wrong code",
            )
        else:
            raise AssertionError(f"{label} was accepted as a frozen source identity")


def _assert_candidate_release_contract() -> None:
    baseline = (
        "首先，系统读取用户提交的论文段落并建立任务记录。"
        "其次，服务按照段落边界生成改写单元并保留原有编号。"
        "此外，校验模块核对术语、数值和引用标记。"
        "因此，只有通过事实与格式检查的文本才会进入结果页，审阅者仍需确认每处修改。"
    )
    readable_output = (
        "系统读取用户提交的论文段落并建立任务记录，服务再按照段落边界生成改写单元，同时保留原有编号。"
        "校验模块负责核对术语、数值和引用标记；文本仅在通过事实与格式检查后进入结果页，"
        "审阅者仍需确认每处修改。"
    )
    unreadable_baseline = "本研究针对控制系统的稳定性开展仿真分析，并依据实验结果评估参数变化的影响。"
    unreadable_output = "本研究针对控制系统的稳定性仿真分析，并依据实验结果评估参数变化的影响。"

    valid_chunk = _candidate_chunk(baseline, readable_output)
    valid_snapshot = _candidate_snapshot(valid_chunk)
    valid_evidence = runner._candidate_release_evidence(
        valid_chunk,
        compare=valid_snapshot["compare"],
    )
    _assert(valid_evidence.get("eligible") is True, f"valid production selection was rejected: {valid_evidence}")
    _assert(valid_evidence.get("candidateSelectionV2") is True, "valid candidate lost v2 evidence")
    _assert(valid_evidence.get("sourceRelativeStylePassed") is True, "valid candidate lost source-relative proof")
    _assert(valid_evidence.get("sourcePatternProfileResolved") is True, "valid candidate profile was not resolved")
    _assert(
        valid_evidence.get("productionReleaseAssessmentPassed") is True,
        "valid candidate failed the canonical release assessment",
    )
    valid_decisions, valid_summary = runner._build_review_decisions(
        valid_snapshot,
        baseline_mode="identity",
        real_candidate_ids={"p0_c0"},
        executions=[{"chunkId": "p0_c0", "sameDimension": {"available": False, "converged": None}}],
    )
    _assert(valid_decisions == {"p0_c0": "rewrite_confirmed"}, "valid production candidate was not confirmed")
    _assert(valid_summary.get("reviewActor") == "automated_e2e", "automated review actor is missing")
    _assert(valid_summary.get("isHumanReview") is False, "automated review was mislabeled as human")
    release_summary = runner._validate_automated_review_release(valid_snapshot, valid_decisions)
    _assert(
        release_summary.get("allRewriteConfirmationsProductionSelected") is True,
        "valid automated release lost production-selection proof",
    )

    stale_chunk = _candidate_chunk(baseline, readable_output, published=False)
    stale_snapshot = _candidate_snapshot(stale_chunk)
    stale_evidence = runner._candidate_release_evidence(
        stale_chunk,
        compare=stale_snapshot["compare"],
    )
    _assert(stale_evidence.get("eligible") is False, "stale preserved-baseline evidence approved changed output")
    _assert(stale_evidence.get("baselinePreservedSafely") is False, "changed output was mislabeled as a safe preserved baseline")
    stale_codes = set(stale_evidence.get("reasonCodes") or [])
    _assert("candidate_selection_not_published" in stale_codes, "stale decision did not fail publishedRewrite")
    _assert("candidate_selection_output_hash_mismatch" in stale_codes, "stale selected hash did not fail output binding")
    stale_decisions, stale_summary = runner._build_review_decisions(
        stale_snapshot,
        baseline_mode="identity",
        # Simulate the v7 bug's old hardValid-only real_candidate_ids set.  The
        # review builder must independently fail closed despite that stale set.
        real_candidate_ids={"p0_c0"},
        executions=[{"chunkId": "p0_c0", "sameDimension": {"available": False, "converged": None}}],
    )
    _assert(stale_decisions == {"p0_c0": "source_confirmed"}, "stale preserved baseline reached automated rewrite confirmation")
    _assert(int(stale_summary.get("changedRewriteConfirmedCount", -1)) == 0, "stale candidate was counted as changed success")
    try:
        runner._validate_automated_review_release(stale_snapshot, {"p0_c0": "rewrite_confirmed"})
    except runner.E2EContractError as exc:
        _assert(exc.code == "automated_review_candidate_evidence_invalid", "stale export gate failed with the wrong code")
    else:
        raise AssertionError("stale preserved-baseline/output-changed candidate reached the export release gate")

    unreadable_chunk = _candidate_chunk(unreadable_baseline, unreadable_output)
    unreadable_snapshot = _candidate_snapshot(unreadable_chunk)
    unreadable_evidence = runner._candidate_release_evidence(
        unreadable_chunk,
        compare=unreadable_snapshot["compare"],
    )
    _assert(unreadable_evidence.get("eligible") is False, "predicate-completeness regression was accepted")
    _assert(unreadable_evidence.get("readabilityPassed") is False, "readability failure was not explicit")
    _assert(
        "predicate_completeness_regression" in (unreadable_evidence.get("readabilityIssueCodes") or []),
        "readability failure lost its text-free issue code",
    )
    unreadable_decisions, unreadable_summary = runner._build_review_decisions(
        unreadable_snapshot,
        baseline_mode="identity",
        real_candidate_ids={"p0_c0"},
        executions=[{"chunkId": "p0_c0", "sameDimension": {"available": False, "converged": None}}],
    )
    _assert(unreadable_decisions == {"p0_c0": "source_confirmed"}, "readability failure reached automated review")
    _assert(
        int(unreadable_summary.get("academicReadabilityFailureCount", 0)) == 1,
        "readability rejection was not counted",
    )

    missing_profile_snapshot = copy.deepcopy(valid_snapshot)
    missing_profile_snapshot["compare"].pop("sourcePatternProfiles", None)
    missing_profile = runner._candidate_release_evidence(
        missing_profile_snapshot["compare"]["chunks"][0],
        compare=missing_profile_snapshot["compare"],
    )
    _assert(missing_profile.get("eligible") is False, "missing document profile failed open")
    _assert(
        "candidate_source_pattern_profile_unresolved" in (missing_profile.get("reasonCodes") or []),
        "missing document profile did not produce a stable reason",
    )

    stale_delta_snapshot = copy.deepcopy(valid_snapshot)
    stale_delta_selection = stale_delta_snapshot["compare"]["chunks"][0]["candidateSelection"]
    stale_delta_selection["resultSourceRelativeStyleDelta"]["binding"]["candidateTextSha256"] = "0" * 64
    stale_delta = runner._candidate_release_evidence(
        stale_delta_snapshot["compare"]["chunks"][0],
        compare=stale_delta_snapshot["compare"],
    )
    _assert(stale_delta.get("eligible") is False, "stale source-relative delta failed open")
    _assert(
        stale_delta.get("productionReleaseAssessmentPassed") is False,
        "stale source-relative delta bypassed canonical release assessment",
    )


def _build_fixture(path: Path) -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "某大学本科毕业设计（论文）"
    document.sections[0].footer.paragraphs[0].text = "内部格式保真测试页脚"
    document.add_heading("基于可信队列的任务处理系统研究", 0)
    document.add_heading("目录", level=1)
    document.add_paragraph("1 绪论........................1")
    document.add_paragraph("2 系统设计....................4")
    document.add_heading("摘要", level=1)
    document.add_paragraph(
        "本文围绕任务处理系统的状态一致性展开研究。该系统能够记录请求、校验结果和失败原因，"
        "并依据明确的状态边界决定后续操作。实验过程保持输入条件不变，所有结论均来自本文记录的测试结果。"
    )
    document.add_heading("1 绪论", level=1)
    technical_paragraph = document.add_paragraph()
    technical_paragraph.add_run(
        "首先，实验环境采用Python 3.11与PostgreSQL 16，测试集包含1200条请求。其次，系统在验证集上的成功率为91.2%，"
        "平均响应时间为128 ms"
    )
    citation_run = technical_paragraph.add_run("[3]")
    citation_run.font.superscript = True
    technical_paragraph.add_run(
        "。该系统能够保留事务编号、时间戳和错误类别，便于复核指标与原始记录之间的对应关系。"
    )
    document.add_heading("2 系统设计", level=1)
    document.add_paragraph(
        "首先，系统读取待处理任务并检查状态。其次，队列按照创建时间选择记录。再次，工作线程写入处理结果。"
        "最后，审计模块核对事务边界。该系统能够维持接口一致性，也能够在失败后保留可追踪证据。"
    )
    document.add_paragraph(
        "控制组与实验组使用相同数据，差异仅来自调度策略。研究人员逐项核对日志、数据库记录和导出文档，"
        "确认处理顺序没有改变，引用编号与数值单位也保持一致。该系统能够在不扩展事实范围的情况下完成局部修正。"
    )
    document.add_paragraph("E = mc^2")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "成功率"
    table.cell(1, 1).text = "91.2%"
    document.add_heading("致谢", level=1)
    document.add_paragraph("感谢指导教师对实验边界和论文格式提出的建议。")
    document.add_heading("参考文献", level=1)
    document.add_paragraph("[1] 张三. 任务队列一致性研究[J]. 软件学报, 2024, 35(2): 1-10.")
    document.add_paragraph("[2] Smith J. Transactional Queue Processing[J]. Systems, 2023, 8(1): 11-20.")
    document.save(path)


def _assert_real_template_kind_aware_scope(temp_dir: Path) -> dict[str, int | bool | str]:
    if not REAL_TEMPLATE_PATH.exists():
        _assert(
            os.environ.get(REQUIRE_REAL_TEMPLATE_ENV, "").strip() != "1",
            "required private university thesis fixture is missing",
        )
        # The authoritative university template is deliberately untracked and
        # must never be copied into a public CI checkout.  Dedicated synthetic
        # semantic-boundary regressions still cover bookmark/comment topology;
        # this exact 396/75/60 inventory remains an opt-in local contract.
        return {
            "available": False,
            "executed": False,
            "skipReason": "private_fixture_not_available",
            "kindAwareSemanticRangePolicy": True,
        }
    isolated_source = temp_dir / "authoritative-university-thesis-template.docx"
    shutil.copy2(REAL_TEMPLATE_PATH, isolated_source)
    contract, freeze, _extracted_path, _snapshot_path, authoritative_snapshot = (
        runner._fresh_docx_preflight(isolated_source)
    )
    _assert(
        int(authoritative_snapshot.total_text_unit_count)
        == REAL_TEMPLATE_EXPECTED_TOTAL_UNITS,
        "authoritative real template total text-unit count drifted",
    )
    _assert(
        int(authoritative_snapshot.editable_unit_count)
        == REAL_TEMPLATE_EXPECTED_EDITABLE_UNITS,
        "authoritative real template editable-unit count drifted",
    )
    for evidence in (contract, freeze):
        _assert(
            int(evidence.get("editableUnitCount", -1))
            == REAL_TEMPLATE_EXPECTED_EDITABLE_UNITS,
            "kind-aware contract/freezing summary lost the 75 editable body units",
        )
        _assert(
            int(evidence.get("editableBookmarkRangeInteriorUnitCount", -1))
            == REAL_TEMPLATE_EXPECTED_EDITABLE_BOOKMARK_INTERIORS,
            "marker-free bookmark interiors were not certified as 60 editable units",
        )
        _assert(
            int(evidence.get("editableSemanticRangeAnchorUnitCount", -1)) == 0,
            "a bookmark/comment endpoint entered the real template model scope",
        )
        _assert(
            int(evidence.get("editableSemanticRangeCoveredUnitCount", -1)) == 0,
            "a comment-range interior entered the real template model scope",
        )
        _assert(
            evidence.get("kindAwareSemanticRangePolicy") is True,
            "kind-aware semantic-range policy disclosure is missing",
        )
    _assert(
        int(contract.get("structuralInventoryVersion", 0) or 0) >= 3,
        "real template used a pre kind-aware structural inventory",
    )
    _assert(
        int(freeze.get("totalTextUnitCount", -1))
        == REAL_TEMPLATE_EXPECTED_TOTAL_UNITS,
        "public freeze summary lost the authoritative 396-unit inventory",
    )
    return {
        "available": True,
        "executed": True,
        "authoritativeTotalTextUnitCount": int(authoritative_snapshot.total_text_unit_count),
        "authoritativeEditableUnitCount": int(authoritative_snapshot.editable_unit_count),
        "editableBookmarkRangeInteriorUnitCount": int(
            freeze.get("editableBookmarkRangeInteriorUnitCount", 0) or 0
        ),
        "editableSemanticRangeAnchorUnitCount": int(
            freeze.get("editableSemanticRangeAnchorUnitCount", 0) or 0
        ),
        "editableCommentRangeInteriorUnitCount": int(
            freeze.get("editableSemanticRangeCoveredUnitCount", 0) or 0
        ),
        "kindAwareSemanticRangePolicy": True,
    }


class FakeCompletion:
    PRIVATE_STREAM_EVENT_BODY = "PRIVATE_STREAM_EVENT_BODY_MUST_NOT_PERSIST"
    PRIVATE_REASONING_BODY = "PRIVATE_REASONING_BODY_MUST_NOT_PERSIST"

    def __init__(self) -> None:
        self.attempts: dict[str, int] = {}
        self.transport_policies: list[tuple[int, int]] = []
        self.stream_requests: list[object] = []
        self.stream_callback_presence: list[bool] = []

    def __call__(self, *args, **kwargs) -> str:  # noqa: ANN002,ANN003
        self.transport_policies.append(
            (int(kwargs.get("max_retries", -1)), int(kwargs.get("timeout", -1)))
        )
        self.stream_requests.append(kwargs.get("stream"))
        stream_callback = kwargs.get("stream_callback")
        self.stream_callback_presence.append(callable(stream_callback))
        if callable(stream_callback):
            stream_callback(
                {
                    "eventCount": 1,
                    "done": False,
                    "delta": self.PRIVATE_STREAM_EVENT_BODY,
                    "reasoning": self.PRIVATE_REASONING_BODY,
                    "thinking": self.PRIVATE_REASONING_BODY,
                }
            )
            stream_callback(
                {
                    "eventCount": 2,
                    "done": True,
                    "providerEvent": {
                        "text": self.PRIVATE_STREAM_EVENT_BODY,
                        "analysis": self.PRIVATE_REASONING_BODY,
                    },
                }
            )
        prompt = str(args[0] if args else kwargs.get("prompt", "") or "")
        chunk_match = re.search(r"\[CHUNK\s+([^\]]+)\]", prompt)
        chunk_id = chunk_match.group(1) if chunk_match else "unknown"
        attempt = self.attempts.get(chunk_id, 0) + 1
        self.attempts[chunk_id] = attempt
        if attempt == 1:
            return "无效输出"
        _assert("[INPUT TEXT]" in prompt, "fake provider did not receive an FYADR input marker")
        text = prompt.rsplit("[INPUT TEXT]", 1)[1].lstrip("\r\n").strip()
        replacements = (
            ("首先，", ""),
            ("其次，", "随后，"),
            ("再次，", "之后，"),
            ("最后，", "处理完成后，"),
            ("该系统能够", "该系统可以"),
        )
        for before, after in replacements:
            text = text.replace(before, after)
        return text


class FakeProviderFailureCompletion:
    PRIVATE_BODY = "UPSTREAM_PRIVATE_BODY_MUST_NOT_PERSIST"
    PRIVATE_ENDPOINT = "https://example.com/v1/chat/completions"

    def __init__(self, status_code: int) -> None:
        self.status_code = int(status_code)
        self.call_count = 0
        self.stream_requests: list[object] = []

    def __call__(self, *_args, **kwargs) -> str:  # noqa: ANN002,ANN003
        self.call_count += 1
        self.stream_requests.append(kwargs.get("stream"))
        raise runner.LLMRequestError(
            self.PRIVATE_BODY,
            category="auth",
            status_code=self.status_code,
            retryable=False,
            endpoint=self.PRIVATE_ENDPOINT,
            provider_message=self.PRIVATE_BODY,
        )


def _assert_completion_auditor_stream_contract() -> None:
    prompt_body = "PRIVATE_AUDITOR_PROMPT_BODY_MUST_NOT_PERSIST"
    output_body = "PRIVATE_AUDITOR_OUTPUT_BODY_MUST_NOT_PERSIST"
    stream_body = "PRIVATE_AUDITOR_STREAM_BODY_MUST_NOT_PERSIST"
    chained_events: list[dict[str, object]] = []

    def upstream_callback(event: dict[str, object]) -> None:
        chained_events.append(event)

    def delegate(prompt: str, **kwargs: object) -> str:
        _assert(prompt == prompt_body, "auditor delegate lost the original prompt")
        _assert(kwargs.get("stream") is True, "auditor changed stream=True")
        callback = kwargs.get("stream_callback")
        _assert(callable(callback), "auditor did not install a stream callback")
        callback({"eventCount": 4, "done": False, "delta": stream_body})
        callback({"eventCount": 5, "done": True, "text": stream_body})
        return output_body

    auditor = runner.CompletionCallAuditor(delegate, max_calls=1)
    result = auditor(
        prompt_body,
        stream=True,
        stream_callback=upstream_callback,
    )
    _assert(result == output_body, "auditor changed the delegate result")
    _assert(len(chained_events) == 2, "auditor dropped the existing callback chain")
    _assert(
        chained_events[-1].get("text") == stream_body,
        "existing callback did not receive the original event",
    )
    summary = auditor.public_summary()
    _assert(summary.get("allRealCallsRequestedStreaming") is True, "stream request proof is false")
    _assert(int(summary.get("streamCompletedCallCount", 0)) == 1, "stream completion was not counted")
    _assert(int(summary.get("nonStreamingCallCount", -1)) == 0, "streaming call was counted as non-streaming")
    calls = summary.get("calls") or []
    _assert(len(calls) == 1, "auditor summary lost the call")
    _assert(int(calls[0].get("streamEventCount", 0)) == 5, "safe eventCount metadata drifted")
    _assert(calls[0].get("streamDone") is True, "safe done metadata drifted")
    encoded = json.dumps(summary, ensure_ascii=False, sort_keys=True)
    for forbidden in (prompt_body, output_body, stream_body, "delta", '"text"'):
        _assert(forbidden not in encoded, "auditor summary retained prompt/output/stream event content")

    # A terminal event observed on a failed call must not compensate for a
    # successful call that never emitted its own terminal stream event.
    def mixed_delegate(prompt: str, **kwargs: object) -> str:
        callback = kwargs.get("stream_callback")
        _assert(callable(callback), "mixed auditor did not install a stream callback")
        if prompt == "successful-without-done":
            callback({"eventCount": 1, "done": False})
            return "safe-final"
        callback({"eventCount": 1, "done": True})
        raise runner.LLMRequestError(
            "safe synthetic failure",
            category="server",
            status_code=503,
            retryable=False,
        )

    mixed_auditor = runner.CompletionCallAuditor(mixed_delegate, max_calls=2)
    _assert(
        mixed_auditor("successful-without-done", stream=True) == "safe-final",
        "mixed auditor changed successful output",
    )
    try:
        mixed_auditor("failed-with-done", stream=True)
    except runner.LLMRequestError:
        pass
    else:
        raise AssertionError("mixed auditor synthetic provider failure did not propagate")
    mixed_summary = mixed_auditor.public_summary()
    _assert(
        int(mixed_summary.get("streamDoneObservedCallCount", 0)) == 1,
        "mixed auditor lost the errored call's terminal event",
    )
    _assert(
        int(mixed_summary.get("successfulStreamCompletedCallCount", -1)) == 0
        and int(mixed_summary.get("successfulStreamIncompleteCallCount", -1)) == 1
        and mixed_summary.get("allSuccessfulCallsStreamCompleted") is False,
        "errored-call done event incorrectly satisfied successful-call stream evidence",
    )


def _assert_full_round_model_failure_audit_contract() -> None:
    private_prompt = "PRIVATE_FULL_ROUND_FAILURE_PROMPT_MUST_NOT_PERSIST"
    private_output = "PRIVATE_FULL_ROUND_FAILURE_OUTPUT_MUST_NOT_PERSIST"

    def delegate(_prompt: str, **kwargs: object) -> str:
        callback = kwargs.get("stream_callback")
        _assert(callable(callback), "full-round failure auditor received no stream callback")
        callback({"eventCount": 2, "done": True, "text": private_output})
        return private_output

    auditor = runner.CompletionCallAuditor(delegate, max_calls=4)
    _assert(auditor(private_prompt, stream=True) == private_output, "failure auditor changed provider output")
    failed_attempt = {
        "event": "validation-retry",
        "schema": round_service.FAILED_ATTEMPT_EVIDENCE_SCHEMA,
        "schemaVersion": round_service.FAILED_ATTEMPT_EVIDENCE_VERSION,
        "round": 1,
        "chunkId": "p10_c0",
        "paragraphIndex": 10,
        "chunkIndex": 0,
        "attempt": 2,
        "guardCategory": "factual",
        "issueCodes": ["factual_scope_qualifier_changed"],
        "outputCharCount": len(private_output),
        "outputTextSha256": runner._sha256_text(private_output),
        "textStored": False,
        "errorStored": False,
        "reasoningSuppressed": True,
        "providerContentStored": False,
        "truncated": False,
    }
    terminal = round_service._build_candidate_selection_event(
        chunk_id="p10_c0",
        round_number=1,
        candidates=[],
        selected={
            "candidateId": "baseline",
            "origin": "baseline",
            "textSha256": runner._sha256_text("baseline"),
            "charCount": len("baseline"),
            "sourceRelativeStyleDelta": {},
        },
        reason_codes=[
            "all_model_candidates_failed_hard_validation",
            "baseline_preserved_but_round_failed",
        ],
        conditional_retry_count=0,
        decision="hard_failure_preserved_baseline",
        run_failed=True,
    )
    validation_error = ValueError("PRIVATE_VALIDATION_ERROR_MUST_NOT_PERSIST")
    setattr(validation_error, "validation_events", [failed_attempt, terminal])
    pipeline_error = RuntimeError("PRIVATE_PIPELINE_ERROR_MUST_NOT_PERSIST")
    pipeline_error.__cause__ = validation_error

    failure = runner._model_output_validation_failure_for_error(pipeline_error)
    _assert(isinstance(failure, dict), "full-round hard failure was not recognized from structured events")
    _assert(
        failure
        == {
            "status": "model_output_failure",
            "category": "model_output_failure",
            "code": "model_output_hard_validation_failed",
            "chunkId": "p10_c0",
            "modelAttemptCount": 0,
            "guardCategories": ["factual"],
            "issueCodes": ["factual_scope_qualifier_changed"],
            "failedCandidateTextStored": False,
            "validationErrorStored": False,
            "providerContentStored": False,
        },
        f"full-round model failure projection drifted: {failure}",
    )
    summary = {**auditor.public_summary(), "callScope": "full_document_baseline"}
    _assert(int(summary.get("callCount", 0)) == 1, "full-round failure lost the completed real call")
    _assert(summary.get("allSuccessfulCallsStreamCompleted") is True, "full-round failure lost stream completion")
    encoded = json.dumps({"failure": failure, "realModel": summary}, ensure_ascii=False, sort_keys=True)
    for forbidden in (
        private_prompt,
        private_output,
        "PRIVATE_VALIDATION_ERROR_MUST_NOT_PERSIST",
        "PRIVATE_PIPELINE_ERROR_MUST_NOT_PERSIST",
    ):
        _assert(forbidden not in encoded, "full-round failure evidence retained private content")


def _assert_output_non_overwrite_contract(temp_dir: Path) -> None:
    report_path = temp_dir / "existing-report.json"
    export_path = temp_dir / "existing-export.docx"
    report_sentinel = b"existing report evidence"
    export_sentinel = b"existing export evidence"
    report_path.write_bytes(report_sentinel)
    export_path.write_bytes(export_sentinel)
    for expected_code, candidate_report, candidate_export in (
        ("report_path_already_exists", report_path, temp_dir / "fresh-export.docx"),
        ("export_path_already_exists", temp_dir / "fresh-report.json", export_path),
    ):
        try:
            runner._require_fresh_output_paths(candidate_report, candidate_export)
        except runner.E2EContractError as exc:
            _assert(exc.code == expected_code, "existing output returned the wrong stable code")
            _assert(exc.category == "input_error", "existing output returned the wrong category")
        else:
            raise AssertionError("existing E2E output path was accepted for replacement")
    _assert(report_path.read_bytes() == report_sentinel, "existing report evidence was modified")
    _assert(export_path.read_bytes() == export_sentinel, "existing export evidence was modified")


def _assert_format_anchor_resolution_contract() -> None:
    common_evidence = {
        "selectionPresent": True,
        "candidateSelectionV2": True,
        "selectedHashMatchesOutput": True,
        "readabilityPassed": True,
        "sourceRelativeStylePassed": True,
        "productionReleaseAssessmentPassed": True,
        "storesCandidateText": False,
    }
    summary = runner._summarize_format_anchor_executions(
        [
            {
                "chunkId": "format-preserved",
                "realCallCount": 2,
                "hardValidationPassed": True,
                "candidateReleaseEvidence": {
                    **common_evidence,
                    "eligible": False,
                    "publishedRewrite": False,
                    "baselinePreservedSafely": True,
                    "decision": "preserved_baseline",
                    "selectedOrigin": "baseline",
                },
            },
            {
                "chunkId": "format-published",
                "realCallCount": 1,
                "hardValidationPassed": True,
                "candidateReleaseEvidence": {
                    **common_evidence,
                    "eligible": True,
                    "publishedRewrite": True,
                    "baselinePreservedSafely": False,
                    "decision": "generated_selected",
                    "selectedOrigin": "model",
                },
            },
        ],
        {"format-preserved", "format-published"},
    )
    _assert(summary.get("formatAnchorExercisedTargetCount") == 2, "format-anchor real calls were not counted")
    _assert(summary.get("formatAnchorSafelyResolvedTargetCount") == 2, "safe format-anchor outcomes were not counted")
    _assert(summary.get("formatAnchorPublishedCandidateCount") == 1, "published format-anchor evidence was not counted")
    _assert(summary.get("formatAnchorBaselinePreservedCount") == 1, "safe format-anchor baseline preservation was not counted")

    unsafe = runner._summarize_format_anchor_executions(
        [
            {
                "chunkId": "format-unsafe",
                "realCallCount": 1,
                "hardValidationPassed": True,
                "candidateReleaseEvidence": {
                    **common_evidence,
                    "selectionPresent": False,
                    "baselinePreservedSafely": True,
                    "publishedRewrite": False,
                    "decision": "preserved_baseline",
                    "selectedOrigin": "baseline",
                },
            }
        ],
        {"format-unsafe"},
    )
    _assert(unsafe.get("formatAnchorExercisedTargetCount") == 1, "unsafe format target was not exercised")
    _assert(unsafe.get("formatAnchorSafelyResolvedTargetCount") == 0, "missing selection evidence failed open")


def _assert_pass_status_contract() -> None:
    _assert(
        runner._passed_e2e_status(manual_review=False, changed_rewrite_count=0)
        == "passed_baseline_preserved",
        "a safe no-gain run was incorrectly treated as a model-output failure",
    )
    _assert(
        runner._passed_e2e_status(manual_review=True, changed_rewrite_count=0)
        == "passed_baseline_preserved_with_manual_review",
        "a safe baseline-only run lost its manual-review disclosure",
    )
    _assert(
        runner._passed_e2e_status(manual_review=False, changed_rewrite_count=1) == "passed"
        and runner._passed_e2e_status(manual_review=True, changed_rewrite_count=1)
        == "passed_with_manual_review",
        "changed-candidate pass statuses drifted",
    )
    _assert(
        runner._passed_e2e_status(
            manual_review=True,
            changed_rewrite_count=0,
            source_fallback_count=2,
        ) == "passed_with_source_fallbacks_and_manual_review",
        "source-fallback pass status lost its manual-review disclosure",
    )
    _assert(
        runner._passed_e2e_status(
            manual_review=False,
            changed_rewrite_count=1,
            source_fallback_count=1,
        ) == "passed_with_source_fallbacks",
        "source-fallback pass status was treated as a normal rewrite pass",
    )


def _assert_full_round_fallback_outcome_contract() -> None:
    def selection(*, decision: str, origin: str, published: bool, attempts: int, reasons: list[str]) -> dict[str, object]:
        return {
            "decision": decision,
            "selectedOrigin": origin,
            "publishedRewrite": published,
            "runFailed": False,
            "modelAttemptCount": attempts,
            "reasonCodes": reasons,
        }

    snapshot = {
        "compare": {
            "chunks": [
                {
                    "chunkId": "published",
                    "inputText": "source-a",
                    "outputText": "changed-a",
                    "candidateSelection": selection(
                        decision="generated_selected",
                        origin="model",
                        published=True,
                        attempts=1,
                        reasons=[],
                    ),
                },
                {
                    "chunkId": "preserved",
                    "inputText": "source-b",
                    "outputText": "source-b",
                    "candidateSelection": selection(
                        decision="preserved_baseline",
                        origin="baseline",
                        published=False,
                        attempts=2,
                        reasons=["no_measurable_combined_style_gain"],
                    ),
                },
                {
                    "chunkId": "fallback",
                    "inputText": "source-c",
                    "outputText": "source-c",
                    "fallbackMode": "source",
                    "fallbackReason": round_service.HARD_VALIDATION_EXHAUSTED_SOURCE_PRESERVED,
                    "quality": {"needsReview": True, "flags": ["source_fallback"]},
                    "failedAttempts": [{"attempt": index} for index in range(round_service.MAX_TOTAL_MODEL_ATTEMPTS)],
                    "candidateSelection": selection(
                        decision="preserved_baseline",
                        origin="baseline",
                        published=False,
                        attempts=round_service.MAX_TOTAL_MODEL_ATTEMPTS,
                        reasons=[round_service.HARD_VALIDATION_EXHAUSTED_SOURCE_PRESERVED],
                    ),
                },
                {"chunkId": "frozen", "inputText": "source-d", "outputText": "source-d"},
            ]
        }
    }
    summary = runner._summarize_full_round_candidate_outcomes(snapshot)
    _assert(summary.get("outcomeDecompositionValid") is True, "full-round outcome decomposition rejected valid categories")
    _assert(summary.get("sourceFallbackEvidenceValid") is True, "valid source-fallback evidence was rejected")
    _assert(summary.get("publishedRewriteCount") == 1, "published full-round outcome count drifted")
    _assert(summary.get("preservedBaselineWithoutFallbackCount") == 1, "baseline-preserved count drifted")
    _assert(summary.get("sourceFallbackCount") == 1, "source-fallback count drifted")
    _assert(summary.get("frozenIdentityCount") == 1, "frozen identity count drifted")
    _assert(summary.get("modelAttemptCount") == 6, "full-round model attempt total drifted")
    _assert(summary.get("sourceFallbacksCountAsPublishedRewrites") is False, "source fallback was counted as a published rewrite")


def _run_lock_probe(lock_path: Path) -> dict[str, object]:
    probe = """
import json
from pathlib import Path
import sys
import time
import real_thesis_model_e2e as runner

started = time.monotonic()
try:
    with runner._execution_mode_run_lock("real_provider", lock_path=Path(sys.argv[1])):
        result = {"acquired": True, "code": "", "category": ""}
except runner.E2EContractError as exc:
    result = {"acquired": False, "code": exc.code, "category": exc.category}
result["lockAttemptMs"] = round((time.monotonic() - started) * 1000)
print(json.dumps(result, sort_keys=True))
"""
    completed = subprocess.run(
        [sys.executable, "-c", probe, str(lock_path)],
        cwd=ROOT_DIR,
        env={
            **os.environ,
            "PYTHONPATH": str(SCRIPTS_DIR),
        },
        check=True,
        capture_output=True,
        text=True,
        timeout=30,
    )
    value = json.loads(completed.stdout.strip())
    _assert(isinstance(value, dict), "lock probe returned a non-object")
    return value


def _assert_lock_file_has_no_runtime_metadata(lock_path: Path, message: str) -> None:
    # POSIX leaves an empty lock file.  Windows ``msvcrt`` needs one physical
    # byte to lock, so a single NUL sentinel is equally metadata-free.
    _assert(lock_path.read_bytes() in {b"", b"\0"}, message)


def _assert_cross_process_lock_contract(temp_dir: Path) -> None:
    lock_path = temp_dir / "real-e2e.lock"
    with runner._execution_mode_run_lock("real_provider", lock_path=lock_path):
        # The offline mode must remain usable even while a real-provider lock
        # is held; it never opens or competes for the advisory lock.
        with runner._execution_mode_run_lock("offline_fake", lock_path=lock_path):
            pass
        conflict = _run_lock_probe(lock_path)
        _assert(conflict.get("acquired") is False, "competing real-provider process acquired the lock")
        _assert(
            conflict.get("code") == runner.REAL_E2E_LOCK_CONFLICT_CODE,
            "lock contention did not return the stable code",
        )
        _assert(conflict.get("category") == "concurrency_conflict", "lock conflict category drifted")
        _assert(
            int(conflict.get("lockAttemptMs", 10000)) < 1000,
            "lock conflict did not fail immediately",
        )

    # Windows byte-range locks also deny a second handle from reading the
    # locked byte.  Inspect the metadata-free sentinel only after the owner
    # has released it; contention itself is proven by the subprocess above.
    _assert_lock_file_has_no_runtime_metadata(lock_path, "lock file persisted runtime metadata")

    class ExpectedProbeExit(RuntimeError):
        pass

    try:
        with runner._execution_mode_run_lock("real_provider", lock_path=lock_path):
            raise ExpectedProbeExit()
    except ExpectedProbeExit:
        pass
    released = _run_lock_probe(lock_path)
    _assert(released.get("acquired") is True, "exception exit did not release the advisory lock")
    _assert_lock_file_has_no_runtime_metadata(lock_path, "released lock file contains provider/run data")


def run_regression() -> dict[str, object]:
    _assert_snapshot_decision_coverage_contract()
    _assert_candidate_release_contract()
    _assert_completion_auditor_stream_contract()
    _assert_full_round_model_failure_audit_contract()
    _assert_format_anchor_resolution_contract()
    _assert_pass_status_contract()
    _assert_full_round_fallback_outcome_contract()
    finish_regression = ROOT_DIR / "finish" / "regression"
    finish_regression.mkdir(parents=True, exist_ok=True)
    previous_env = {key: os.environ.get(key) for key in ENV_VALUES}
    original_completion = app_service.llm_completion
    original_load_config = app_config.load_app_config
    fake = FakeCompletion()
    run_id = f"offline-{uuid.uuid4().hex[:12]}"
    work_source_paths: list[Path] = []
    provider_failure_cases: list[tuple[int, FakeProviderFailureCompletion, dict[str, object], Path]] = []
    cleanup_errors: list[str] = []
    with tempfile.TemporaryDirectory(prefix="real-thesis-e2e-offline-", dir=finish_regression) as temp_value:
        temp_dir = Path(temp_value)
        _assert_cross_process_lock_contract(temp_dir)
        _assert_output_non_overwrite_contract(temp_dir)
        real_template_scope = _assert_real_template_kind_aware_scope(temp_dir)
        sample_path = temp_dir / "university-thesis.docx"
        report_path = temp_dir / "report.json"
        export_path = temp_dir / "export.docx"
        _build_fixture(sample_path)
        try:
            os.environ.update(ENV_VALUES)
            app_service.llm_completion = fake

            def fail_if_config_is_read(*_args, **_kwargs):
                raise AssertionError("real thesis E2E must not read data/config/config.json")

            app_config.load_app_config = fail_if_config_is_read
            report = runner.run_e2e(
                sample_path=sample_path,
                source_url="https://example.invalid/public-university-thesis",
                source_commit="offline-fixture-v1",
                source_license="TEST-ONLY",
                report_path=report_path,
                export_path=export_path,
                run_id=run_id,
                max_real_targets=3,
                full_round=False,
                execution_mode="offline_fake",
            )
            work_source_paths.append(Path(str((report.get("source") or {}).get("workSourcePath", ""))))

            for status_code in (401, 403):
                provider_failure_fake = FakeProviderFailureCompletion(status_code)
                app_service.llm_completion = provider_failure_fake
                case_report_path = temp_dir / f"provider_failure_{status_code}.json"
                case_report = runner.run_e2e(
                    sample_path=sample_path,
                    source_url="https://example.invalid/public-university-thesis",
                    source_commit=f"offline-provider-failure-{status_code}",
                    source_license="TEST-ONLY",
                    report_path=case_report_path,
                    export_path=temp_dir / f"provider_failure_{status_code}.docx",
                    run_id=f"offline-provider-failure-{status_code}-{uuid.uuid4().hex[:8]}",
                    max_real_targets=3,
                    full_round=False,
                    execution_mode="offline_fake",
                )
                work_source_paths.append(
                    Path(str((case_report.get("source") or {}).get("workSourcePath", "")))
                )
                provider_failure_cases.append(
                    (status_code, provider_failure_fake, case_report, case_report_path)
                )

        finally:
            app_service.llm_completion = original_completion
            app_config.load_app_config = original_load_config
            for key, previous in previous_env.items():
                if previous is None:
                    os.environ.pop(key, None)
                else:
                    os.environ[key] = previous

        _assert(report.get("ok") is True, f"offline E2E failed: {report.get('failure')}")
        _assert(report.get("baselineMode") == "identity", "offline E2E did not use identity baseline")
        execution_isolation = report.get("executionIsolation") or {}
        _assert(
            execution_isolation.get("realProviderCrossProcessLockRequired") is False
            and execution_isolation.get("realProviderCrossProcessLockHeldForRun") is False,
            "offline fake report falsely claimed the real-provider lock",
        )
        _assert(
            int(execution_isolation.get("maxConcurrentRealProviderE2E", 0)) == 1
            and int(execution_isolation.get("configuredRewriteConcurrency", 0)) == 1,
            "single-concurrency execution policy evidence drifted",
        )
        _assert(
            execution_isolation.get("lockFileStoresRuntimeMetadata") is False
            and execution_isolation.get("lockPathStored") is False,
            "lock report exposes or claims to persist runtime metadata",
        )
        baseline = report.get("baselineExecution") or {}
        _assert(int(baseline.get("providerCallCount", -1)) == 0, "identity baseline claimed provider calls")
        real_model = report.get("realModel") or {}
        _assert(int(real_model.get("callCount", 0)) >= 3, "bounded fake provider was not called")
        _assert(int(real_model.get("sourceFallbackCount", -1)) == 0, "offline fake unexpectedly reported source fallbacks")
        _assert(real_model.get("sourceFallbacksCountAsSuccessfulCandidates") is False, "source fallback success semantics are undisclosed")
        _assert(real_model.get("callCountMatchesModelAttemptCount") is None, "bounded mode falsely claimed full-round attempt reconciliation")
        _assert(
            int(real_model.get("maxCallCount", 0)) == 3 * runner.E2E_MAX_COMPLETIONS_PER_TARGET,
            "bounded E2E completion budget drifted",
        )
        _assert(
            all(
                retries == runner.E2E_HTTP_RETRIES_PER_COMPLETION
                and timeout == runner.E2E_REQUEST_TIMEOUT_SECONDS
                for retries, timeout in fake.transport_policies
            ),
            "real-provider E2E inherited unbounded production retry/timeout policy",
        )
        _assert(fake.stream_requests, "bounded fake provider recorded no stream policy")
        _assert(
            all(requested is True for requested in fake.stream_requests),
            "bounded provider call did not dynamically receive stream=True",
        )
        _assert(
            all(fake.stream_callback_presence),
            "stream auditor did not install a metadata callback on every call",
        )
        _assert(
            real_model.get("allRealCallsRequestedStreaming") is True,
            "public report cannot prove every bounded call requested streaming",
        )
        _assert(
            int(real_model.get("nonStreamingCallCount", -1)) == 0,
            "public report counted a non-streaming bounded call",
        )
        _assert(
            real_model.get("allSuccessfulCallsStreamCompleted") is True
            and int(real_model.get("successfulStreamCompletedCallCount", -1))
            == int(real_model.get("successCount", -2)),
            "stream completion evidence does not cover every successful bounded call",
        )
        _assert(int(real_model.get("targetCount", 0)) == 3, "representative target count drifted")
        representative = report.get("representativeTargets") or {}
        _assert(int(representative.get("availableFormatAnchorCount", 0)) >= 1, "fixture lost its styled citation anchor")
        _assert(int(representative.get("formatAnchorTargetCount", 0)) >= 1, "styled anchor was not selected for real-model coverage")
        _assert(int(real_model.get("formatAnchorExercisedTargetCount", 0)) >= 1, "styled-anchor target made no real model call")
        _assert(int(real_model.get("formatAnchorSafelyResolvedTargetCount", 0)) >= 1, "styled-anchor target was not safely resolved")
        _assert(
            int(real_model.get("formatAnchorSafelyResolvedTargetCount", 0))
            == int(real_model.get("formatAnchorPublishedCandidateCount", 0))
            + int(real_model.get("formatAnchorBaselinePreservedCount", 0)),
            "format-anchor safe outcomes do not decompose into publish or preserve",
        )
        round_summary = report.get("round") or {}
        _assert(int(round_summary.get("immutableFormatAnchorCount", 0)) >= 1, "baseline worker did not bind format anchors")
        _assert(bool(real_model.get("baselineIdentityCallsExcluded")), "identity calls were mixed into real call totals")
        preflight = report.get("preflight") or {}
        preflight_contract = preflight.get("contract") or {}
        preflight_freeze = preflight.get("freeze") or {}
        for evidence in (preflight_contract, preflight_freeze):
            _assert(
                evidence.get("semanticRangeTopologyValid") is True,
                "semantic range topology proof is missing or invalid",
            )
            _assert(
                int(evidence.get("semanticRangeCount", -1)) >= 0,
                "semantic range count proof is missing",
            )
            _assert(
                int(evidence.get("semanticRangeCoveredUnitCount", -1)) >= 0,
                "semantic range covered-unit proof is missing",
            )
            _assert(
                int(evidence.get("editableSemanticRangeCoveredUnitCount", -1)) == 0,
                "a comment-range-covered unit entered editable model scope",
            )
            _assert(
                int(evidence.get("editableSemanticRangeAnchorUnitCount", -1)) == 0,
                "a semantic-range endpoint entered editable model scope",
            )
            _assert(
                int(evidence.get("editableBookmarkRangeInteriorUnitCount", -1)) >= 0,
                "bookmark-interior scope proof is missing",
            )
            _assert(
                evidence.get("kindAwareSemanticRangePolicy") is True,
                "kind-aware semantic-range policy disclosure is missing",
            )
        snapshot_summaries = report.get("snapshots") or {}
        for label in ("initial", "postModelPreReview", "reviewMaterialized"):
            snapshot_summary = snapshot_summaries.get(label) or {}
            candidate_evidence = snapshot_summary.get("candidateSelectionEvidence") or {}
            _assert(
                candidate_evidence.get("allSelectionsV2") is True,
                f"{label} snapshot lost candidate-selection v2",
            )
            _assert(
                candidate_evidence.get("allReleaseCriticalSourceRelativeEvidenceValid") is True
                and candidate_evidence.get("allSourceRelativeEvidenceValid") is True,
                f"{label} snapshot contains invalid source-relative evidence",
            )
            _assert(
                candidate_evidence.get("sourceRelativeDocumentDeltaPassed") is True,
                f"{label} snapshot lost the cumulative document delta",
            )
            _assert(
                int(candidate_evidence.get("sourcePatternProfileRegistryCount", 0)) > 0,
                f"{label} snapshot lost the content-addressed source profile",
            )
        executions = real_model.get("executions") or []
        _assert(len(executions) == 3, "target execution evidence is incomplete")
        _assert(any(item.get("failedAttempts") for item in executions), "validation retry evidence was not summarized")
        for item in executions:
            for attempt in item.get("failedAttempts") or []:
                _assert(attempt.get("schemaValid") is True, "failed attempt schema is invalid")
                _assert(attempt.get("safeEvidenceValid") is True, "failed attempt evidence is not privacy-safe")
                _assert(attempt.get("textStored") is False, "failed model output leaked into report")
                _assert(attempt.get("errorStored") is False, "raw validation error leaked into report")
                _assert(attempt.get("reasoningSuppressed") is True, "reasoning suppression proof is missing")
                _assert(attempt.get("providerContentStored") is False, "provider body leaked into report")
                _assert(attempt.get("forbiddenFieldObserved") is False, "private failed-attempt field was observed")
                _assert(bool(attempt.get("issueCodes")), "failed attempt has no canonical issue code")
                output_sha256 = str(attempt.get("outputTextSha256", "") or "")
                _assert(
                    not output_sha256 or runner.SHA256_RE.fullmatch(output_sha256) is not None,
                    "failed attempt output identity is invalid",
                )
        review = report.get("review") or {}
        _assert(int(review.get("sourceFallbackCount", -1)) == 0, "offline review falsely reported source fallback")
        _assert(int(review.get("decisionCount", 0)) == int(review.get("chunkCount", -1)), "review decisions are not full coverage")
        _assert(int(review.get("changedRewriteConfirmedCount", 0)) >= 1, "no changed fake candidate reached export")
        _assert(review.get("reviewActor") == "automated_e2e", "E2E CAS report omitted its automated actor")
        _assert(review.get("isHumanReview") is False, "E2E CAS report falsely claimed human review")
        _assert(
            review.get("allRewriteConfirmationsProductionSelected") is True
            and review.get("allRewriteConfirmationsHashBound") is True
            and review.get("allRewriteConfirmationsReadabilityPassed") is True,
            "automated rewrite confirmations lack release evidence",
        )
        export = report.get("export") or {}
        _assert(export.get("overallStatus") == "passed", "certified DOCX export did not pass")
        _assert(int(export.get("auditIssueCount", -1)) == 0, "protected text audit failed")
        _assert(int(export.get("ooxmlAuditIssueCount", -1)) == 0, "OOXML audit failed")
        _assert(int(export.get("formatLockIssueCount", -1)) == 0, "format lock failed")
        _assert(bool(export.get("modelInputMatchesEditableUnits")), "export lost exact editable scope")
        reviewed_snapshot = (report.get("snapshots") or {}).get("reviewMaterialized") or {}
        export_snapshot = export.get("roundArtifactSnapshot") or {}
        _assert(export.get("reviewMaterializationBound") is True, "DOCX export lacks review-materialization binding")
        for key in (
            "compareRevision",
            "reviewRevision",
            "contentRevision",
            "artifactSnapshotDigest",
            "effectiveTextSha256",
        ):
            _assert(
                export_snapshot.get(key) == reviewed_snapshot.get(key),
                f"DOCX export diverged from reviewed snapshot field {key}",
            )
        artifact_scan = report.get("artifactRuntimeParameterScan") or {}
        _assert(artifact_scan.get("ok") is True, f"runtime parameter artifact scan failed: {artifact_scan}")
        raw_report = report_path.read_text(encoding="utf-8")
        _assert_no_private_report_fields(report)
        _assert(
            (report.get("provider") or {}).get("credentialFingerprintStored") is False,
            "report claims to retain a credential fingerprint",
        )
        for value in (
            ENV_VALUES[runner.BASE_URL_ENV],
            ENV_VALUES[runner.API_KEY_ENV],
            ENV_VALUES[runner.MODEL_ENV],
        ):
            _assert(value not in raw_report, "runtime provider plaintext leaked into report")
        _assert("无效输出" not in raw_report, "failed model output leaked into report")
        _assert(
            FakeCompletion.PRIVATE_STREAM_EVENT_BODY not in raw_report,
            "stream event body leaked into report",
        )
        _assert(
            FakeCompletion.PRIVATE_REASONING_BODY not in raw_report,
            "thinking/reasoning event body leaked into report",
        )

        _assert(len(provider_failure_cases) == 2, "offline 401/403 cases did not both execute")
        provider_failure_statuses: list[int] = []
        for status_code, provider_failure_fake, case_report, case_report_path in provider_failure_cases:
            provider_failure_statuses.append(status_code)
            _assert(case_report.get("ok") is False, f"HTTP {status_code} case unexpectedly passed")
            _assert(
                case_report.get("status") == "provider_configuration_failure",
                f"HTTP {status_code} was not classified as provider configuration failure",
            )
            case_failure = case_report.get("failure") or {}
            _assert(
                case_failure.get("category") == "provider_configuration_failure",
                f"HTTP {status_code} top-level failure category drifted",
            )
            _assert(
                case_failure.get("code") == "bounded_target_provider_configuration_failure",
                f"HTTP {status_code} top-level failure code drifted",
            )
            top_provider_failure = case_failure.get("providerFailure") or {}
            _assert(
                top_provider_failure
                == {
                    "status": "provider_configuration_failure",
                    "category": "auth",
                    "statusCode": status_code,
                },
                f"HTTP {status_code} safe top-level provider descriptor is incomplete",
            )

            case_real_model = case_report.get("realModel") or {}
            _assert(
                provider_failure_fake.stream_requests
                and all(value is True for value in provider_failure_fake.stream_requests),
                f"HTTP {status_code} call did not request streaming",
            )
            _assert(
                case_real_model.get("allRealCallsRequestedStreaming") is True
                and int(case_real_model.get("nonStreamingCallCount", -1)) == 0,
                f"HTTP {status_code} stream request evidence drifted",
            )
            _assert(
                int(case_real_model.get("streamCompletedCallCount", -1)) == 0,
                f"HTTP {status_code} falsely claimed a completed provider stream",
            )
            _assert(
                int(case_real_model.get("providerConfigurationFailureTargetCount", 0)) == 3,
                f"HTTP {status_code} provider failure target count drifted",
            )
            _assert(
                int(case_real_model.get("externalFailureTargetCount", -1)) == 0
                and int(case_real_model.get("productFailureTargetCount", -1)) == 0,
                f"HTTP {status_code} was also misclassified as external/product failure",
            )
            case_executions = case_real_model.get("executions") or []
            _assert(len(case_executions) == 3, f"HTTP {status_code} execution evidence is incomplete")
            _assert(
                all(item.get("outcome") == "provider_configuration_failure" for item in case_executions),
                f"HTTP {status_code} execution was mislabeled as model-output failure",
            )
            _assert(
                all(
                    (item.get("providerFailure") or {}).get("category") == "auth"
                    and (item.get("providerFailure") or {}).get("statusCode") == status_code
                    and (item.get("providerFailure") or {}).get("status")
                    == "provider_configuration_failure"
                    for item in case_executions
                ),
                f"HTTP {status_code} execution omitted safe provider status/category/statusCode",
            )
            case_calls = case_real_model.get("calls") or []
            _assert(
                len(case_calls) == provider_failure_fake.call_count >= 3,
                f"HTTP {status_code} call audit count drifted",
            )
            _assert(
                all(
                    item.get("outcome") == "error"
                    and item.get("errorCategory") == "auth"
                    and item.get("statusCode") == status_code
                    for item in case_calls
                ),
                f"HTTP {status_code} call audit lost the canonical category/status code",
            )
            _assert(
                not any(item.get("outcome") == "model_output_hard_validation_failed" for item in case_executions),
                f"HTTP {status_code} still appears as a model-output validation failure",
            )
            case_raw_report = case_report_path.read_text(encoding="utf-8")
            _assert_no_private_report_fields(case_report)
            for forbidden in (
                FakeProviderFailureCompletion.PRIVATE_BODY,
                FakeProviderFailureCompletion.PRIVATE_ENDPOINT,
                ENV_VALUES[runner.BASE_URL_ENV],
                ENV_VALUES[runner.API_KEY_ENV],
                ENV_VALUES[runner.MODEL_ENV],
            ):
                _assert(forbidden not in case_raw_report, f"HTTP {status_code} report leaked provider details")

        checks = [
            "complete DOCX identity baseline is separated from bounded provider calls",
            "three representative targets use production rerun/strategy paths",
            "a superscript citation is hidden, restored, and accepted through a real-model target path",
            "network retries, request timeout, and completion calls have explicit E2E budgets",
            "every real-model call explicitly requests streaming and records text-free completion metadata",
            "existing stream callbacks remain chained while event bodies are discarded",
            "real-provider E2E uses a fail-closed cross-process advisory lock; offline mode is lock-free",
            "kind-aware semantic ranges freeze anchors/comments while allowing safe bookmark interiors",
            "failed-attempt v1 evidence keeps only stable enums, hash/count metadata, and suppression flags",
            "candidate-selection v2 source-relative and cumulative document evidence are consumed fail-closed",
            "candidate selections and frozen source identities cover every chunk exactly once",
            "only production-published, hash-bound, readability-passing candidates can be automated",
            "stale preserved-baseline evidence with changed output is rejected before export",
            "automated CAS is labeled automated_e2e and never presented as human review",
            "selected/output/review-materialized hashes bind to one certified DOCX snapshot generation",
            "offline HTTP 401/403 are provider-configuration failures, never model-output failures",
            "full review CAS and certified DOCX audits pass",
            "runtime provider plaintext is absent from changed artifacts and report",
        ]
        if real_template_scope.get("executed") is True:
            checks.insert(
                8,
                "the private real template remains authoritative at 396 total / 75 editable / 60 bookmark-interior units",
            )
        else:
            checks.insert(
                8,
                "the untracked private real-template inventory is explicitly skipped when unavailable",
            )

        result = {
            "ok": True,
            "checks": checks,
            "identityCallCount": int(baseline.get("identityTransformCallCount", 0) or 0),
            "boundedCallCount": int(real_model.get("callCount", 0) or 0),
            "providerFailureStatuses": provider_failure_statuses,
            "changedRewriteConfirmedCount": int(review.get("changedRewriteConfirmedCount", 0) or 0),
            "realTemplateScope": real_template_scope,
        }

    for work_source_path in work_source_paths:
        if not work_source_path.exists() or not work_source_path.is_file():
            continue
        try:
            work_source_path.resolve().relative_to(runner.RUN_SOURCE_ROOT.resolve())
        except ValueError:
            cleanup_errors.append("UnsafeCleanupPath")
            continue
        try:
            status = app_service.get_document_status(str(work_source_path))
            app_service.delete_document_history(
                str(status.get("docId", "") or ""),
                mode="records_artifacts_and_source",
            )
        except Exception as exc:  # cleanup must not hide the regression result
            cleanup_errors.append(type(exc).__name__)
        try:
            shutil.rmtree(work_source_path.parent, ignore_errors=False)
        except OSError as exc:
            cleanup_errors.append(type(exc).__name__)
    result["cleanupErrorTypes"] = cleanup_errors
    return result


def main() -> int:
    try:
        report = run_regression()
    except Exception as exc:
        print(json.dumps({"ok": False, "errorType": type(exc).__name__, "message": str(exc)}, ensure_ascii=False, indent=2))
        return 1
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
