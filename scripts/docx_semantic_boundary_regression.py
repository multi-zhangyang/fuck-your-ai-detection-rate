from __future__ import annotations

import json
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

import app_service  # noqa: E402
import fyadr_round_service as round_service  # noqa: E402
import round_helper as round_helper_module  # noqa: E402
from docx_audit import (  # noqa: E402
    _semantic_document_range_evidence,
    _semantic_marker_attachment_signature,
    audit_docx_export,
    audit_docx_format_lock,
    audit_docx_ooxml_integrity,
)
from docx_bodymap import load_docx_body_map, update_docx_body_map_texts  # noqa: E402
from docx_pipeline import (  # noqa: E402
    DOCX_SNAPSHOT_VERSION,
    _load_docx_snapshot,
    build_docx_scope_diagnostics,
    build_docx_snapshot,
    ensure_docx_processing_assets,
    get_docx_snapshot_path,
    rebuild_docx_from_body_map_units,
)
from document_edit_contract import build_document_edit_contract  # noqa: E402
from round_helper import run_document_round  # noqa: E402


REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_semantic_boundary_regression_report.json"
SEMANTIC_PREFIX = "语义范围前文保持稳定，"
COMMENT_TARGET = "批注关键结论"
BOOKMARK_TARGET = "书签术语"
SEMANTIC_SUFFIX = "，语义范围后文同样保持稳定。"
SEMANTIC_PARAGRAPH = f"{SEMANTIC_PREFIX}{COMMENT_TARGET}{BOOKMARK_TARGET}{SEMANTIC_SUFFIX}"
POINT_PREFIX = "独立批注落点前文保持稳定，"
POINT_SUFFIX = "独立批注落点后文保持稳定。"
POINT_PARAGRAPH = f"{POINT_PREFIX}{POINT_SUFFIX}"
ORDINARY_PARAGRAPH = "普通正文段落不含零宽语义标记，应继续进入模型并允许按既有规则改写。"
COMMENT_BODY = "仅用于验证批注部件不会进入模型或证据正文。"
CROSS_COMMENT_BODY = "仅用于验证跨段批注范围闭包，不得进入任何公开证据。"
CROSS_COMMENT_START = "跨段批注从该正文段开始，后续被选中的自然段必须全部冻结。"
CROSS_RANGE_PARAGRAPHS = (
    "跨段语义范围中间第一段保持稳定，不得发送给模型或由模型改写。",
    "跨段语义范围中间第二段保持完整，不得发送给模型或由模型改写。",
    "跨段语义范围中间第三段保持原样，不得发送给模型或由模型改写。",
)
CROSS_RANGE_ATTACK = CROSS_RANGE_PARAGRAPHS[1].replace("完整", "篡改")
CROSS_COMMENT_END = "跨段批注在该正文段结束，范围端点及批注落点必须保持稳定。"
BOOKMARK_ONLY_START = "跨段书签从该段建立导航起点，当前锚点段必须保持稳定。"
BOOKMARK_ONLY_PARAGRAPHS = (
    "跨段书签内部第一段包含完整正文证据，边界节点不在本段，因此允许进入模型处理。",
    "跨段书签内部第二段用于验证安全回填，正文可以改写而书签起止位置必须保持不变。",
)
BOOKMARK_ONLY_REWRITE = BOOKMARK_ONLY_PARAGRAPHS[1].replace("正文可以改写", "正文能够调整")
BOOKMARK_ONLY_END = "跨段书签在该段结束导航范围，当前锚点段同样必须保持稳定。"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _add_bookmark(run: Any, *, bookmark_id: int, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    run._r.addprevious(start)
    run._r.addnext(end)


def _add_bookmark_start(paragraph: Any, *, bookmark_id: int, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.append(start)


def _add_bookmark_end(paragraph: Any, *, bookmark_id: int) -> None:
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def _move_comment_end(
    source_paragraph: Any,
    destination_paragraph: Any,
    *,
    comment_id: int,
) -> None:
    matching_ends = [
        node
        for node in source_paragraph._p.xpath(".//w:commentRangeEnd")
        if str(node.get(qn("w:id")) or "") == str(comment_id)
    ]
    matching_references = [
        node
        for node in source_paragraph._p.xpath(".//w:commentReference")
        if str(node.get(qn("w:id")) or "") == str(comment_id)
    ]
    _assert(len(matching_ends) == 1, "cross-comment fixture has no unique range end")
    _assert(len(matching_references) == 1, "cross-comment fixture has no unique reference")
    range_end = matching_ends[0]
    reference_run = matching_references[0].getparent()
    _assert(reference_run is not None, "cross-comment reference has no run parent")
    range_end.getparent().remove(range_end)
    reference_run.getparent().remove(reference_run)
    destination_paragraph._p.append(range_end)
    destination_paragraph._p.append(reference_run)


def _create_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("零宽语义边界回归论文", style="Title")
    document.add_paragraph("摘 要")
    document.add_paragraph("摘要正文用于建立论文正文范围，并保持正常的保护边界识别。")
    document.add_paragraph("1 引言", style="Heading 1")

    paragraph = document.add_paragraph()
    paragraph.add_run(SEMANTIC_PREFIX)
    comment_run = paragraph.add_run(COMMENT_TARGET)
    bookmark_run = paragraph.add_run(BOOKMARK_TARGET)
    paragraph.add_run(SEMANTIC_SUFFIX)
    comment = document.add_comment(
        comment_run,
        text=COMMENT_BODY,
        author="FYADR regression",
        initials="FR",
    )
    _add_bookmark(bookmark_run, bookmark_id=42, name="fyadr_semantic_anchor")

    point_paragraph = document.add_paragraph()
    point_paragraph.add_run(POINT_PREFIX)
    point_run = point_paragraph.add_run()
    point_reference = OxmlElement("w:commentReference")
    point_reference.set(qn("w:id"), str(comment.comment_id))
    point_run._r.append(point_reference)
    point_paragraph.add_run(POINT_SUFFIX)

    cross_start_paragraph = document.add_paragraph()
    cross_start_run = cross_start_paragraph.add_run(CROSS_COMMENT_START)
    cross_comment = document.add_comment(
        cross_start_run,
        text=CROSS_COMMENT_BODY,
        author="FYADR cross-range regression",
        initials="FCR",
    )

    marker_only_start = document.add_paragraph()
    _add_bookmark_start(
        marker_only_start,
        bookmark_id=84,
        name="fyadr_cross_paragraph_anchor",
    )
    for range_paragraph in CROSS_RANGE_PARAGRAPHS:
        document.add_paragraph(range_paragraph)
    marker_only_end = document.add_paragraph()
    _add_bookmark_end(marker_only_end, bookmark_id=84)

    cross_end_paragraph = document.add_paragraph(CROSS_COMMENT_END)
    _move_comment_end(
        cross_start_paragraph,
        cross_end_paragraph,
        comment_id=cross_comment.comment_id,
    )

    bookmark_only_start = document.add_paragraph(BOOKMARK_ONLY_START)
    _add_bookmark_start(
        bookmark_only_start,
        bookmark_id=85,
        name="fyadr_bookmark_only_cross_paragraph",
    )
    for range_paragraph in BOOKMARK_ONLY_PARAGRAPHS:
        document.add_paragraph(range_paragraph)
    bookmark_only_end = document.add_paragraph(BOOKMARK_ONLY_END)
    _add_bookmark_end(bookmark_only_end, bookmark_id=85)

    document.add_paragraph(ORDINARY_PARAGRAPH)
    document.add_paragraph("致 谢")
    document.add_paragraph("感谢参与边界回归验证的人员。")
    document.add_paragraph("参考文献", style="Heading 1")
    document.add_paragraph("[1] Semantic boundary regression reference.")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _read_part(path: Path, name: str) -> bytes:
    with zipfile.ZipFile(str(path), "r") as archive:
        return archive.read(name)


def _rewrite_part(path: Path, destination: Path, name: str, transform: Any) -> None:
    with zipfile.ZipFile(str(path), "r") as source, zipfile.ZipFile(
        str(destination),
        "w",
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == name:
                data = transform(data)
            target.writestr(item, data)


def _move_comment_range_without_changing_paragraph_text(xml: bytes) -> bytes:
    prefix = SEMANTIC_PREFIX.encode("utf-8")
    target = COMMENT_TARGET.encode("utf-8")
    moved_character = COMMENT_TARGET[0].encode("utf-8")
    if xml.count(prefix) != 1 or xml.count(target) != 1:
        raise AssertionError("semantic-boundary fixture text is not unique in document.xml")
    mutated = xml.replace(prefix, prefix + moved_character, 1)
    mutated = mutated.replace(target, COMMENT_TARGET[1:].encode("utf-8"), 1)
    if mutated == xml:
        raise AssertionError("semantic-boundary attachment mutation did not change document.xml")
    return mutated


def _change_cross_range_middle_text(xml: bytes) -> bytes:
    original = CROSS_RANGE_PARAGRAPHS[1].encode("utf-8")
    replacement = CROSS_RANGE_ATTACK.encode("utf-8")
    _assert(len(CROSS_RANGE_PARAGRAPHS[1]) == len(CROSS_RANGE_ATTACK), "range attack must preserve character length")
    if xml.count(original) != 1:
        raise AssertionError("cross-range middle paragraph is not unique in document.xml")
    mutated = xml.replace(original, replacement, 1)
    if mutated == xml:
        raise AssertionError("cross-range content mutation did not change document.xml")
    return mutated


def _document_range_evidence(path: Path) -> dict[str, Any]:
    root = ET.fromstring(_read_part(path, "word/document.xml"))
    return _semantic_document_range_evidence(root)


def _create_invalid_range_fixture(path: Path, mode: str) -> None:
    document = Document()
    document.add_paragraph("坏范围拓扑回归论文", style="Title")
    document.add_paragraph("摘 要")
    document.add_paragraph("摘要正文用于建立正文范围。")
    document.add_paragraph("1 引言", style="Heading 1")
    paragraph = document.add_paragraph("正文不得在范围拓扑异常时进入模型。")

    def append_marker(name: str, marker_id: int) -> None:
        marker = OxmlElement(f"w:{name}")
        marker.set(qn("w:id"), str(marker_id))
        if name == "bookmarkStart":
            marker.set(qn("w:name"), f"invalid_{mode}")
        paragraph._p.append(marker)

    if mode == "unmatched":
        append_marker("bookmarkStart", 301)
    elif mode == "duplicate":
        append_marker("bookmarkStart", 302)
        append_marker("bookmarkStart", 302)
        append_marker("bookmarkEnd", 302)
    elif mode == "reversed":
        append_marker("bookmarkEnd", 303)
        append_marker("bookmarkStart", 303)
    else:
        raise AssertionError(f"unsupported invalid range mode: {mode}")
    document.add_paragraph(ORDINARY_PARAGRAPH)
    document.add_paragraph("致 谢")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _find_paragraph(document: Any, text: str) -> Any:
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def _read_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    _assert(isinstance(payload, dict), f"JSON evidence is not an object: {path}")
    return payload


def main() -> int:
    checks: list[str] = []
    work_root = ROOT_DIR / "finish" / "web_exports"
    work_root.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="docx-semantic-boundary-", dir=work_root) as temp_dir:
        work_dir = Path(temp_dir)
        source_path = work_dir / "source.docx"
        latest_export_path = work_dir / "export.docx"
        tampered_path = work_dir / "attachment-drift.docx"
        cross_tampered_path = work_dir / "cross-range-content-drift.docx"
        bookmark_rewrite_export_path = work_dir / "bookmark-interior-rewrite.docx"
        _create_fixture(source_path)

        model_inputs: list[str] = []

        def identity_transform(chunk_text: str, _prompt: str, _round: int, _chunk_id: str) -> str:
            model_inputs.append(str(chunk_text))
            return str(chunk_text).replace(
                BOOKMARK_ONLY_PARAGRAPHS[1],
                BOOKMARK_ONLY_REWRITE,
            )

        original_update_round = round_service.update_round
        original_helper_update_round = round_helper_module.update_round
        try:
            # This regression validates DOCX scope/export semantics, not the
            # shared history index. Keep it deterministic and independent from
            # external SQLite maintenance or concurrent history tests.
            round_service.update_round = lambda **kwargs: {"rounds": [kwargs]}
            round_helper_module.update_round = lambda **kwargs: {"rounds": [kwargs]}
            round_result = run_document_round(
                source_path,
                identity_transform,
                round_number=1,
                prompt_profile="cn",
            )
        finally:
            round_service.update_round = original_update_round
            round_helper_module.update_round = original_helper_update_round
        output_path = Path(str(round_result["output_path"]))
        snapshot_path = get_docx_snapshot_path(source_path)
        snapshot = _load_docx_snapshot(snapshot_path)
        _assert(snapshot is not None, "semantic-boundary snapshot was not created")
        _assert(snapshot.version == DOCX_SNAPSHOT_VERSION == 22, "semantic-boundary snapshot schema was not bumped")
        _assert(snapshot.semantic_range_topology_valid is True, "valid cross-paragraph ranges failed topology validation")
        _assert(snapshot.semantic_range_issue_count == 0, "valid cross-paragraph ranges emitted topology issues")
        _assert(snapshot.semantic_range_count >= 4, "snapshot omitted same-paragraph or cross-paragraph semantic ranges")

        semantic_unit = next((unit for unit in snapshot.units if unit.text == SEMANTIC_PARAGRAPH), None)
        point_unit = next((unit for unit in snapshot.units if unit.text == POINT_PARAGRAPH), None)
        ordinary_unit = next((unit for unit in snapshot.units if unit.text == ORDINARY_PARAGRAPH), None)
        cross_start_unit = next((unit for unit in snapshot.units if unit.text == CROSS_COMMENT_START), None)
        cross_end_unit = next((unit for unit in snapshot.units if unit.text == CROSS_COMMENT_END), None)
        cross_middle_units = [
            next((unit for unit in snapshot.units if unit.text == paragraph), None)
            for paragraph in CROSS_RANGE_PARAGRAPHS
        ]
        bookmark_only_start_unit = next(
            (unit for unit in snapshot.units if unit.text == BOOKMARK_ONLY_START),
            None,
        )
        bookmark_only_end_unit = next(
            (unit for unit in snapshot.units if unit.text == BOOKMARK_ONLY_END),
            None,
        )
        bookmark_only_middle_units = [
            next((unit for unit in snapshot.units if unit.text == paragraph), None)
            for paragraph in BOOKMARK_ONLY_PARAGRAPHS
        ]
        _assert(semantic_unit is not None, "semantic-boundary paragraph disappeared from scope diagnostics")
        _assert(point_unit is not None, "semantic point-reference paragraph disappeared from scope diagnostics")
        _assert(ordinary_unit is not None, "ordinary control paragraph disappeared from scope diagnostics")
        _assert(cross_start_unit is not None and cross_end_unit is not None, "cross-comment endpoints disappeared from scope diagnostics")
        _assert(all(unit is not None for unit in cross_middle_units), "cross-range middle paragraph disappeared from scope diagnostics")
        _assert(
            bookmark_only_start_unit is not None and bookmark_only_end_unit is not None,
            "bookmark-only endpoints disappeared from scope diagnostics",
        )
        _assert(
            all(unit is not None for unit in bookmark_only_middle_units),
            "bookmark-only interior disappeared from scope diagnostics",
        )
        _assert(semantic_unit.editable is False, "semantic-boundary paragraph entered model scope")
        _assert(
            semantic_unit.protect_reason == "semantic_range_anchor",
            f"semantic-boundary reason is not explicit: {semantic_unit.protect_reason}",
        )
        _assert(semantic_unit.has_semantic_range_anchor is True, "range-anchor evidence was not captured")
        _assert(semantic_unit.has_semantic_point_reference is True, "comment-reference evidence was not captured")
        _assert(semantic_unit.has_complex_inline is True, "semantic boundary did not join the complex-inline fail-safe")
        _assert(point_unit.editable is False, "semantic point-reference paragraph entered model scope")
        _assert(
            point_unit.protect_reason == "semantic_point_reference",
            f"point-reference reason is not explicit: {point_unit.protect_reason}",
        )
        _assert(point_unit.has_semantic_range_anchor is False, "point-only paragraph was misclassified as a range")
        _assert(point_unit.has_semantic_point_reference is True, "point-reference evidence was not captured")
        _assert(
            cross_start_unit.editable is False
            and cross_start_unit.protect_reason == "semantic_range_anchor"
            and cross_start_unit.inside_semantic_range is True,
            "cross-comment start endpoint entered model scope",
        )
        _assert(
            cross_end_unit.editable is False
            and cross_end_unit.protect_reason == "semantic_range_anchor"
            and cross_end_unit.inside_semantic_range is True,
            "cross-comment end endpoint entered model scope",
        )
        for cross_middle_unit in cross_middle_units:
            _assert(cross_middle_unit is not None, "cross-range middle unit is missing")
            _assert(
                cross_middle_unit.editable is False
                and cross_middle_unit.protect_reason == "semantic_range_span"
                and cross_middle_unit.inside_semantic_range is True
                and cross_middle_unit.inside_comment_range is True
                and cross_middle_unit.has_semantic_range_anchor is False,
                "unmarked cross-range middle paragraph entered model scope",
            )
        _assert(
            bookmark_only_start_unit.editable is False
            and bookmark_only_start_unit.has_bookmark_range_anchor is True
            and bookmark_only_start_unit.has_comment_range_anchor is False,
            "bookmark-only start anchor entered model scope",
        )
        _assert(
            bookmark_only_end_unit.editable is False
            and bookmark_only_end_unit.has_bookmark_range_anchor is True
            and bookmark_only_end_unit.has_comment_range_anchor is False,
            "bookmark-only end anchor entered model scope",
        )
        for bookmark_middle_unit in bookmark_only_middle_units:
            _assert(bookmark_middle_unit is not None, "bookmark-only interior unit is missing")
            _assert(
                bookmark_middle_unit.editable is True
                and bookmark_middle_unit.inside_bookmark_range is True
                and bookmark_middle_unit.inside_comment_range is False
                and bookmark_middle_unit.has_semantic_range_anchor is False,
                "marker-free bookmark interior did not remain eligible prose",
            )
        _assert(ordinary_unit.editable is True, "ordinary unmarked body text was globally disabled")

        joined_model_inputs = "\n\n".join(model_inputs)
        _assert(SEMANTIC_PARAGRAPH not in joined_model_inputs, "semantic-boundary text reached a model call")
        _assert(POINT_PARAGRAPH not in joined_model_inputs, "semantic point-reference text reached a model call")
        _assert(CROSS_COMMENT_START not in joined_model_inputs, "cross-comment start reached a model call")
        _assert(CROSS_COMMENT_END not in joined_model_inputs, "cross-comment end reached a model call")
        _assert(BOOKMARK_ONLY_START not in joined_model_inputs, "bookmark-only start anchor reached a model call")
        _assert(BOOKMARK_ONLY_END not in joined_model_inputs, "bookmark-only end anchor reached a model call")
        _assert(
            all(paragraph not in joined_model_inputs for paragraph in CROSS_RANGE_PARAGRAPHS),
            "a cross-range middle paragraph reached a model call",
        )
        _assert(
            all(paragraph in joined_model_inputs for paragraph in BOOKMARK_ONLY_PARAGRAPHS),
            "a marker-free bookmark interior paragraph did not reach the model",
        )
        _assert(ORDINARY_PARAGRAPH in joined_model_inputs, "ordinary control text did not reach the model")
        body_map = load_docx_body_map(Path(str(round_result.get("body_map_path", ""))))
        _assert(body_map is not None, "semantic-boundary round lost its body map")
        _assert(
            all(unit.original_text != SEMANTIC_PARAGRAPH for unit in body_map.units),
            "semantic-boundary paragraph entered the editable body map",
        )
        _assert(
            all(unit.original_text != POINT_PARAGRAPH for unit in body_map.units),
            "semantic point-reference paragraph entered the editable body map",
        )
        _assert(
            any(unit.original_text == ORDINARY_PARAGRAPH for unit in body_map.units),
            "ordinary control paragraph was removed from the editable body map",
        )
        _assert(
            all(
                unit.original_text not in {
                    CROSS_COMMENT_START,
                    CROSS_COMMENT_END,
                    *CROSS_RANGE_PARAGRAPHS,
                }
                for unit in body_map.units
            ),
            "cross-paragraph semantic range entered the editable body map",
        )
        _assert(
            all(
                any(unit.original_text == paragraph for unit in body_map.units)
                for paragraph in BOOKMARK_ONLY_PARAGRAPHS
            ),
            "marker-free bookmark interior was omitted from the editable body map",
        )
        bookmark_rewritten_texts = [
            text.replace(BOOKMARK_ONLY_PARAGRAPHS[1], BOOKMARK_ONLY_REWRITE)
            for text in body_map.current_texts()
        ]
        _assert(
            bookmark_rewritten_texts != body_map.current_texts(),
            "bookmark-only export fixture produced no rewrite",
        )
        bookmark_rewritten_body_map = update_docx_body_map_texts(
            body_map,
            bookmark_rewritten_texts,
            round_number=1,
        )
        rebuild_docx_from_body_map_units(
            bookmark_rewritten_body_map.units,
            source_path=source_path,
            export_path=bookmark_rewrite_export_path,
            preserve_format=True,
        )
        _assert(
            audit_docx_export(
                bookmark_rewrite_export_path,
                source_path=source_path,
                snapshot_path=snapshot_path,
            ).get("ok")
            is True,
            "protected-text audit rejected a safe bookmark interior rewrite",
        )
        _assert(
            audit_docx_format_lock(
                bookmark_rewrite_export_path,
                source_path=source_path,
                snapshot_path=snapshot_path,
            ).get("ok")
            is True,
            "format lock rejected a safe bookmark interior rewrite",
        )
        _assert(
            audit_docx_ooxml_integrity(
                bookmark_rewrite_export_path,
                source_path=source_path,
                snapshot_path=snapshot_path,
            ).get("ok")
            is True,
            "OOXML audit rejected a safe bookmark interior rewrite",
        )
        checks.append("bookmark/comment anchors and comment interiors freeze while marker-free bookmark interiors remain editable")

        legacy_snapshot_path = work_dir / "legacy-v18.snapshot.json"
        legacy_extracted_path = work_dir / "legacy-v18.extracted.txt"
        legacy_payload = json.loads(snapshot_path.read_text(encoding="utf-8"))
        legacy_payload["version"] = 18
        legacy_payload.pop("semantic_range_count", None)
        legacy_payload.pop("semantic_range_topology_valid", None)
        legacy_payload.pop("semantic_range_issue_count", None)
        legacy_payload.pop("semantic_range_issue_codes", None)
        for raw_unit in legacy_payload.get("units", []):
            if isinstance(raw_unit, dict) and raw_unit.get("text") in {
                SEMANTIC_PARAGRAPH,
                POINT_PARAGRAPH,
                CROSS_COMMENT_START,
                CROSS_COMMENT_END,
                *CROSS_RANGE_PARAGRAPHS,
            }:
                raw_unit["editable"] = True
                raw_unit["protect_reason"] = None
                raw_unit.pop("has_semantic_range_anchor", None)
                raw_unit.pop("inside_semantic_range", None)
                raw_unit.pop("has_semantic_point_reference", None)
        legacy_payload["editable_unit_count"] = sum(
            1
            for raw_unit in legacy_payload.get("units", [])
            if isinstance(raw_unit, dict) and bool(raw_unit.get("editable"))
        )
        legacy_snapshot_path.write_text(
            json.dumps(legacy_payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        _, _, refreshed_legacy_snapshot = ensure_docx_processing_assets(
            source_path,
            extracted_path=legacy_extracted_path,
            snapshot_path=legacy_snapshot_path,
            scope_diagnostics_path=work_dir / "legacy-v17.scope.json",
        )
        _assert(
            refreshed_legacy_snapshot.version == DOCX_SNAPSHOT_VERSION,
            "v18 snapshot cache was reused instead of source re-derivation",
        )
        refreshed_semantic = next(
            unit
            for unit in refreshed_legacy_snapshot.units
            if unit.text == SEMANTIC_PARAGRAPH
        )
        _assert(
            refreshed_semantic.editable is False
            and refreshed_semantic.protect_reason == "semantic_range_anchor",
            "v18 editable semantic range survived authoritative cache migration",
        )
        refreshed_point = next(
            unit
            for unit in refreshed_legacy_snapshot.units
            if unit.text == POINT_PARAGRAPH
        )
        _assert(
            refreshed_point.editable is False
            and refreshed_point.protect_reason == "semantic_point_reference",
            "v18 editable semantic point reference survived authoritative cache migration",
        )
        refreshed_cross_middle = next(
            unit
            for unit in refreshed_legacy_snapshot.units
            if unit.text == CROSS_RANGE_PARAGRAPHS[1]
        )
        _assert(
            refreshed_cross_middle.editable is False
            and refreshed_cross_middle.protect_reason == "semantic_range_span"
            and refreshed_cross_middle.inside_semantic_range is True,
            "v18 editable cross-range middle paragraph survived authoritative cache migration",
        )
        _assert(
            SEMANTIC_PARAGRAPH not in legacy_extracted_path.read_text(encoding="utf-8")
            and POINT_PARAGRAPH not in legacy_extracted_path.read_text(encoding="utf-8"),
            "v18 cache migration leaked semantic-boundary text into extracted model input",
        )
        _assert(
            all(
                paragraph not in legacy_extracted_path.read_text(encoding="utf-8")
                for paragraph in CROSS_RANGE_PARAGRAPHS
            ),
            "v18 cache migration leaked cross-range text into extracted model input",
        )
        checks.append("v18 caches re-derive to current v22 while retaining the v19 cross-range fail-closed scope")

        export_result = app_service.export_round_output(
            str(output_path),
            str(latest_export_path),
            "docx",
        )
        certified_path = Path(str(export_result.get("path", "")))
        manifest_path = Path(str(export_result.get("evidenceManifestPath", "")))
        contract_path = Path(str(export_result.get("contentContractPath", "")))
        _assert(certified_path.exists(), "semantic-boundary export was not published")
        _assert(manifest_path.exists(), "semantic-boundary evidence manifest is missing")
        _assert(contract_path.exists(), "semantic-boundary content contract is missing")

        contract = _read_json(contract_path)
        manifest = _read_json(manifest_path)
        _assert(contract.get("semanticRangeTopologyValid") is True, "contract marked valid semantic ranges invalid")
        _assert(contract.get("semanticRangeIssueCount") == 0, "contract emitted semantic range topology issues")
        _assert(contract.get("semanticRangeCount", 0) >= 4, "contract omitted document-level range count")
        _assert(contract.get("bookmarkRangeCount", 0) >= 2, "contract omitted bookmark range count")
        _assert(contract.get("commentRangeCount", 0) >= 2, "contract omitted comment range count")
        _assert(contract.get("semanticRangeAnchorUnitCount", 0) >= 1, "contract omitted range-anchor count")
        _assert(contract.get("semanticRangeCoveredUnitCount", 0) >= 3, "contract omitted protected comment-range interior units")
        _assert(contract.get("editableSemanticRangeCoveredUnitCount") == 0, "contract allowed editable text inside a comment range")
        _assert(contract.get("bookmarkRangeInteriorUnitCount", 0) >= 2, "contract omitted bookmark interior units")
        _assert(contract.get("editableBookmarkRangeInteriorUnitCount", 0) >= 2, "contract did not certify safe bookmark interior prose")
        _assert(contract.get("semanticPointReferenceUnitCount", 0) >= 1, "contract omitted point-reference count")
        _assert(contract.get("editableSemanticRangeAnchorUnitCount") == 0, "contract allowed an editable range anchor")
        _assert(contract.get("editableSemanticPointReferenceUnitCount") == 0, "contract allowed an editable point reference")
        boundary_summary = manifest.get("semanticBoundarySummary")
        _assert(isinstance(boundary_summary, dict), "evidence manifest omitted semantic-boundary summary")
        _assert(boundary_summary.get("rangeTopologyValid") is True, "manifest marked valid semantic ranges invalid")
        _assert(boundary_summary.get("rangeCount", 0) >= 4, "manifest omitted document-level range count")
        _assert(boundary_summary.get("rangeAnchorUnitCount", 0) >= 1, "manifest omitted range-anchor evidence")
        _assert(boundary_summary.get("rangeCoveredUnitCount", 0) >= 3, "manifest omitted protected comment-range evidence")
        _assert(boundary_summary.get("editableRangeCoveredUnitCount") == 0, "manifest exposed editable text inside a comment range")
        _assert(boundary_summary.get("pointReferenceUnitCount", 0) >= 1, "manifest omitted point-reference evidence")
        _assert(boundary_summary.get("editableRangeAnchorUnitCount") == 0, "manifest exposed an editable range anchor")
        _assert(boundary_summary.get("editablePointReferenceUnitCount") == 0, "manifest exposed an editable point reference")
        serialized_evidence = json.dumps(
            {"contract": contract, "manifest": manifest},
            ensure_ascii=False,
        )
        _assert(COMMENT_BODY not in serialized_evidence, "comment body leaked into contract/evidence")
        _assert(CROSS_COMMENT_BODY not in serialized_evidence, "cross-comment body leaked into contract/evidence")

        source_document = Document(str(source_path))
        export_document = Document(str(certified_path))
        bookmark_export_document = Document(str(bookmark_rewrite_export_path))
        source_semantic = _find_paragraph(source_document, SEMANTIC_PARAGRAPH)
        export_semantic = _find_paragraph(export_document, SEMANTIC_PARAGRAPH)
        source_point = _find_paragraph(source_document, POINT_PARAGRAPH)
        export_point = _find_paragraph(export_document, POINT_PARAGRAPH)
        export_bookmark_rewrite = _find_paragraph(
            bookmark_export_document,
            BOOKMARK_ONLY_REWRITE,
        )
        _assert(
            source_semantic._p.xml == export_semantic._p.xml,
            "successful export changed bookmark/comment range OOXML",
        )
        _assert(
            source_point._p.xml == export_point._p.xml,
            "successful export changed point-reference OOXML",
        )
        _assert(
            _read_part(source_path, "word/comments.xml")
            == _read_part(certified_path, "word/comments.xml"),
            "successful export changed the exact comments part",
        )
        _assert(
            _read_part(source_path, "word/comments.xml")
            == _read_part(bookmark_rewrite_export_path, "word/comments.xml"),
            "bookmark interior rewrite changed the exact comments part",
        )
        _assert(
            export_bookmark_rewrite.text == BOOKMARK_ONLY_REWRITE,
            "marker-free bookmark interior rewrite was not exported",
        )
        source_range_evidence = _document_range_evidence(source_path)
        export_range_evidence = _document_range_evidence(bookmark_rewrite_export_path)
        _assert(source_range_evidence.get("topologyValid") is True, "source global semantic range topology is invalid")
        _assert(export_range_evidence.get("topologyValid") is True, "export global semantic range topology is invalid")
        _assert(source_range_evidence.get("rangeCount", 0) >= 4, "global semantic range evidence omitted cross ranges")
        source_ranges = {
            (str(item.get("kind", "")), str(item.get("id", ""))): item
            for item in source_range_evidence.get("ranges", [])
            if isinstance(item, dict)
        }
        export_ranges = {
            (str(item.get("kind", "")), str(item.get("id", ""))): item
            for item in export_range_evidence.get("ranges", [])
            if isinstance(item, dict)
        }
        _assert(set(source_ranges) == set(export_ranges), "export changed bookmark/comment range identity")
        for key, source_range in source_ranges.items():
            export_range = export_ranges[key]
            _assert(
                source_range.get("start") == export_range.get("start")
                and source_range.get("end") == export_range.get("end"),
                f"export moved a semantic range endpoint: {key}",
            )
            if key[0] == "comment":
                _assert(
                    source_range.get("contentLength") == export_range.get("contentLength")
                    and source_range.get("contentSha256") == export_range.get("contentSha256"),
                    f"export changed protected comment-range content: {key}",
                )
        _assert(
            source_ranges[("bookmark", "85")].get("contentSha256")
            != export_ranges[("bookmark", "85")].get("contentSha256"),
            "bookmark-only interior rewrite did not change the expected bookmark content hash",
        )
        source_signature = _semantic_marker_attachment_signature(source_semantic._p)
        export_signature = _semantic_marker_attachment_signature(export_semantic._p)
        _assert(source_signature == export_signature and source_signature.get("eventCount", 0) >= 5, "semantic marker signature is incomplete")
        checks.append("bookmark interior rewrites preserve endpoint topology while comment content and exact comments part remain locked")

        _rewrite_part(
            certified_path,
            tampered_path,
            "word/document.xml",
            _move_comment_range_without_changing_paragraph_text,
        )
        tampered_document = Document(str(tampered_path))
        tampered_semantic = _find_paragraph(tampered_document, SEMANTIC_PARAGRAPH)
        _assert(
            tampered_semantic.text == source_semantic.text,
            "negative fixture changed visible paragraph text instead of only marker attachment",
        )
        _assert(
            _semantic_marker_attachment_signature(tampered_semantic._p) != source_signature,
            "negative fixture did not move the semantic marker offset",
        )

        format_report = audit_docx_format_lock(
            tampered_path,
            source_path=source_path,
            snapshot_path=snapshot_path,
        )
        format_issue_types = {
            str(item.get("type", ""))
            for item in format_report.get("issues", [])
            if isinstance(item, dict)
        }
        _assert(
            "format_lock_violation" in format_issue_types,
            f"format lock missed semantic attachment drift: {sorted(format_issue_types)}",
        )

        ooxml_report = audit_docx_ooxml_integrity(
            tampered_path,
            source_path=source_path,
            snapshot_path=snapshot_path,
        )
        ooxml_issue_types = {
            str(item.get("type", ""))
            for item in ooxml_report.get("issues", [])
            if isinstance(item, dict)
        }
        _assert(
            "semantic_marker_attachment_changed" in ooxml_issue_types,
            f"OOXML audit missed unchanged-text marker drift: {sorted(ooxml_issue_types)}",
        )
        ooxml_serialized = json.dumps(ooxml_report, ensure_ascii=False)
        _assert(COMMENT_BODY not in ooxml_serialized, "comment body leaked into OOXML audit evidence")
        _assert(CROSS_COMMENT_BODY not in ooxml_serialized, "cross-comment body leaked into OOXML audit evidence")
        checks.append("format lock and OOXML audit reject unchanged-text marker-offset drift")

        _rewrite_part(
            certified_path,
            cross_tampered_path,
            "word/document.xml",
            _change_cross_range_middle_text,
        )
        tampered_range_evidence = _document_range_evidence(cross_tampered_path)
        _assert(
            tampered_range_evidence.get("topologyValid") is True,
            "middle-text attack unexpectedly changed semantic range topology",
        )
        source_ranges = {
            (str(item.get("kind", "")), str(item.get("id", ""))): item
            for item in source_range_evidence.get("ranges", [])
            if isinstance(item, dict)
        }
        tampered_ranges = {
            (str(item.get("kind", "")), str(item.get("id", ""))): item
            for item in tampered_range_evidence.get("ranges", [])
            if isinstance(item, dict)
        }
        _assert(source_ranges.keys() == tampered_ranges.keys(), "middle-text attack changed range identities")
        _assert(
            any(
                source_ranges[key].get("start") == tampered_ranges[key].get("start")
                and source_ranges[key].get("end") == tampered_ranges[key].get("end")
                and source_ranges[key].get("contentLength") == tampered_ranges[key].get("contentLength")
                and source_ranges[key].get("contentSha256") != tampered_ranges[key].get("contentSha256")
                for key in source_ranges
            ),
            "middle-text attack did not isolate a same-topology range-content hash change",
        )

        protected_report = audit_docx_export(
            cross_tampered_path,
            source_path=source_path,
            snapshot_path=snapshot_path,
        )
        protected_issue_types = {
            str(item.get("type", ""))
            for item in protected_report.get("issues", [])
            if isinstance(item, dict)
        }
        _assert(
            "protected_text_changed" in protected_issue_types,
            f"protected-text audit missed cross-range middle drift: {sorted(protected_issue_types)}",
        )

        cross_ooxml_report = audit_docx_ooxml_integrity(
            cross_tampered_path,
            source_path=source_path,
            snapshot_path=snapshot_path,
        )
        cross_ooxml_issue_types = {
            str(item.get("type", ""))
            for item in cross_ooxml_report.get("issues", [])
            if isinstance(item, dict)
        }
        _assert(
            "semantic_range_content_changed" in cross_ooxml_issue_types,
            f"global OOXML range audit missed middle-text drift: {sorted(cross_ooxml_issue_types)}",
        )
        cross_ooxml_serialized = json.dumps(cross_ooxml_report, ensure_ascii=False)
        _assert(CROSS_RANGE_PARAGRAPHS[1] not in cross_ooxml_serialized, "source range prose leaked into OOXML audit evidence")
        _assert(CROSS_RANGE_ATTACK not in cross_ooxml_serialized, "tampered range prose leaked into OOXML audit evidence")
        _assert(CROSS_COMMENT_BODY not in cross_ooxml_serialized, "cross-comment body leaked into range audit evidence")
        checks.append("protected-text and global hash audits reject same-topology cross-range middle-text drift without prose leakage")

        invalid_cases = {
            "unmatched": "semantic_range_unmatched_start",
            "duplicate": "semantic_range_duplicate_start",
            "reversed": "semantic_range_reversed",
        }
        for mode, expected_code in invalid_cases.items():
            invalid_path = work_dir / f"invalid-{mode}.docx"
            _create_invalid_range_fixture(invalid_path, mode)
            invalid_snapshot = build_docx_snapshot(invalid_path)
            _assert(invalid_snapshot.semantic_range_topology_valid is False, f"{mode} topology was accepted")
            _assert(expected_code in invalid_snapshot.semantic_range_issue_codes, f"{mode} topology omitted {expected_code}")
            _assert(not any(unit.editable for unit in invalid_snapshot.units), f"{mode} topology did not freeze all units")
            invalid_diagnostics = build_docx_scope_diagnostics(invalid_snapshot)
            _assert(invalid_diagnostics.get("ok") is False, f"{mode} topology passed scope diagnostics")
            invalid_contract = build_document_edit_contract(
                invalid_path,
                stage=f"invalid_{mode}_probe",
            )
            invalid_contract_codes = {
                str(item.get("code", ""))
                for item in invalid_contract.get("issues", [])
                if isinstance(item, dict)
            }
            _assert(invalid_contract.get("ready") is False, f"{mode} topology passed document contract")
            _assert(
                "semantic_range_topology_invalid" in invalid_contract_codes,
                f"{mode} contract omitted semantic_range_topology_invalid",
            )
        checks.append("unmatched, duplicate, and reversed semantic ranges fail closed before model input")

    report = {
        "ok": True,
        "snapshotVersion": DOCX_SNAPSHOT_VERSION,
        "checks": checks,
    }
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
