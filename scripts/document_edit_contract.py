"""Fail-closed contract for body-only rewriting and DOCX format fidelity.

The rewrite pipeline has two different kinds of evidence:

* a frozen set of editable body paragraph targets; and
* post-export audits proving that every other text/OOXML/format surface stayed
  identical to the source document.

This module turns those checks into one stable, serialisable contract that can
be enforced by the round runner, the exporter and the Web report.  It does not
attempt to infer authorship or an AIGC score.
"""

from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path
from typing import Any, Sequence
import zipfile

from docx import Document  # type: ignore[import]

from docx_audit import _paragraph_format_signature, _xml_structure_hash
from docx_bodymap import DocxBodyMap, validate_docx_body_map
from docx_pipeline import (
    DOCX_EDITABLE_STRUCTURAL_ROLES,
    DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
    DocxSnapshot,
    _is_snapshot_current,
    _looks_like_heading,
    build_docx_scope_diagnostics,
    ensure_docx_processing_assets,
    get_docx_extracted_text_path,
    get_docx_snapshot_path,
    get_docx_structural_role_map_digest,
    get_docx_unit_edit_eligibility_evidence_digest,
    validate_docx_unit_edit_eligibility_evidence,
    verify_docx_snapshot_derivation,
)
from path_utils import build_document_artifact_stem


ROOT_DIR = Path(__file__).resolve().parents[1]
INTERMEDIATE_DIR = ROOT_DIR / "finish" / "intermediate"
DOCUMENT_EDIT_CONTRACT_VERSION = 3
DOCUMENT_EDIT_POLICY = "editable_body_text_only"
FORMAT_LOCK_POLICY = "source_ooxml_fidelity_lock"


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _sha256_json(value: Any) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return _sha256_bytes(payload.encode("utf-8"))


def get_document_edit_contract_path(source_path: Path) -> Path:
    artifact_stem = build_document_artifact_stem(root_dir=ROOT_DIR, source_path=source_path)
    return INTERMEDIATE_DIR / f"{artifact_stem}_content_contract.json"


def get_export_edit_contract_path(export_path: Path) -> Path:
    return export_path.with_suffix(".content_contract.json")


def write_document_edit_contract(report: dict[str, Any], path: Path) -> Path:
    normalized_path = path.resolve()
    normalized_path.parent.mkdir(parents=True, exist_ok=True)
    payload = dict(report)
    payload["reportPath"] = str(normalized_path)
    normalized_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    report["reportPath"] = str(normalized_path)
    return normalized_path


def _add_issue(
    issues: list[dict[str, Any]],
    code: str,
    message: str,
    *,
    severity: str = "error",
    **details: Any,
) -> None:
    issues.append(
        {
            "code": code,
            "severity": severity,
            "message": message,
            **details,
        }
    )


def _scope_digest(snapshot: DocxSnapshot) -> str:
    units = []
    for unit in snapshot.units:
        units.append(
            {
                "unitIndex": int(unit.unit_index),
                "target": dict(unit.target),
                "styleName": str(unit.style_name),
                "editable": bool(unit.editable),
                "protectReason": str(unit.protect_reason or ""),
                "structuralRole": str(unit.structural_role),
                "editEligibility": str(unit.edit_eligibility),
                "editEligibilityEvidenceDigest": get_docx_unit_edit_eligibility_evidence_digest(unit),
                "presentationSignalsDigest": _sha256_json(unit.presentation_signals),
                "hasSemanticRangeAnchor": bool(
                    getattr(unit, "has_semantic_range_anchor", False)
                ),
                "insideSemanticRange": bool(
                    getattr(unit, "inside_semantic_range", False)
                ),
                "hasBookmarkRangeAnchor": bool(
                    getattr(unit, "has_bookmark_range_anchor", False)
                ),
                "hasCommentRangeAnchor": bool(
                    getattr(unit, "has_comment_range_anchor", False)
                ),
                "insideBookmarkRange": bool(
                    getattr(unit, "inside_bookmark_range", False)
                ),
                "insideCommentRange": bool(
                    getattr(unit, "inside_comment_range", False)
                ),
                "hasSemanticPointReference": bool(
                    getattr(unit, "has_semantic_point_reference", False)
                ),
                "textSha256": _sha256_bytes(str(unit.text).encode("utf-8")),
                "sourceTextSha256": _sha256_bytes(unit.source_text().encode("utf-8")),
                "leadingWhitespaceSha256": _sha256_bytes(unit.leading_whitespace.encode("utf-8")),
                "trailingWhitespaceSha256": _sha256_bytes(unit.trailing_whitespace.encode("utf-8")),
            }
        )
    return _sha256_json(
        {
            "snapshotVersion": snapshot.version,
            "semanticRangeCount": int(getattr(snapshot, "semantic_range_count", 0)),
            "bookmarkRangeCount": int(getattr(snapshot, "bookmark_range_count", 0)),
            "commentRangeCount": int(getattr(snapshot, "comment_range_count", 0)),
            "semanticRangeTopologyValid": bool(
                getattr(snapshot, "semantic_range_topology_valid", True)
            ),
            "semanticRangeIssueCodes": list(
                getattr(snapshot, "semantic_range_issue_codes", [])
            ),
            "structuralRolePolicyVersion": snapshot.structural_role_policy_version,
            "structuralInventoryVersion": snapshot.structural_inventory_version,
            "structuralRoleMapDigest": get_docx_structural_role_map_digest(snapshot),
            "protectedStructuralUnits": [
                {
                    "unitIndex": unit.unit_index,
                    "target": unit.target,
                    "structuralRole": unit.structural_role,
                    "editEligibility": unit.edit_eligibility,
                    "evidenceDigest": get_docx_unit_edit_eligibility_evidence_digest(unit),
                    "sourceTextSha256": _sha256_bytes(unit.source_text().encode("utf-8")),
                }
                for unit in snapshot.protected_structural_units
            ],
            "units": units,
        }
    )


def _format_digest(source_path: Path) -> str:
    """Hash every paragraph/table/section format surface plus package assets.

    Direct editable ``w:t`` payloads are blanked by
    ``_paragraph_format_signature``.  Other package parts are included by raw
    hash so styles, numbering, headers/footers, media and relationships are
    represented in the baseline identity.
    """

    document = Document(str(source_path.resolve()))
    paragraph_signatures = [
        _paragraph_format_signature(paragraph._p)
        for paragraph in document.paragraphs
    ]
    table_signatures = [
        _xml_structure_hash(table._tbl.xml.encode("utf-8"))
        for table in document.tables
    ]
    section_signatures = [
        _xml_structure_hash(section._sectPr.xml.encode("utf-8"))
        for section in document.sections
    ]
    package_parts: list[tuple[str, str]] = []
    with zipfile.ZipFile(str(source_path.resolve()), "r") as archive:
        for item in archive.infolist():
            if item.is_dir() or item.filename == "word/document.xml":
                continue
            package_parts.append((item.filename, _sha256_bytes(archive.read(item.filename))))
    return _sha256_json(
        {
            "paragraphs": paragraph_signatures,
            "tables": table_signatures,
            "sections": section_signatures,
            "packageParts": sorted(package_parts),
        }
    )


def _body_map_contract_issues(
    body_map: DocxBodyMap,
    *,
    source_path: Path,
    snapshot_path: Path,
    authoritative_snapshot: DocxSnapshot,
    snapshot_derivation_report: dict[str, Any],
    issues: list[dict[str, Any]],
) -> dict[str, Any]:
    validation = validate_docx_body_map(
        body_map,
        source_path=source_path,
        snapshot_path=snapshot_path,
        authoritative_snapshot=authoritative_snapshot,
        snapshot_derivation_report=snapshot_derivation_report,
    )
    raw_blocking = validation.get("blockingIssues")
    for raw_issue in raw_blocking if isinstance(raw_blocking, list) else []:
        if not isinstance(raw_issue, dict):
            continue
        if str(raw_issue.get("code", "")) == "snapshot_authority_mismatch":
            # The document-level contract records this once with the same
            # source-derived evidence; avoid duplicating it under a body-map
            # prefix while still keeping bodyMapReady=False via validation.ok.
            continue
        _add_issue(
            issues,
            f"body_map_{str(raw_issue.get('code', 'invalid'))}",
            str(raw_issue.get("message", "DOCX 正文映射校验失败。")),
            details={key: value for key, value in raw_issue.items() if key not in {"code", "message"}},
        )
    raw_warnings = validation.get("warnings")
    for raw_warning in raw_warnings if isinstance(raw_warnings, list) else []:
        if not isinstance(raw_warning, dict):
            continue
        _add_issue(
            issues,
            f"body_map_{str(raw_warning.get('code', 'warning'))}",
            str(raw_warning.get("message", "DOCX 正文映射存在提示。")),
            severity="warning",
            details={key: value for key, value in raw_warning.items() if key not in {"code", "message"}},
        )
    return validation


def _build_docx_contract(
    source_path: Path,
    *,
    snapshot_path: Path | None,
    extracted_text_path: Path | None,
    body_map: DocxBodyMap | None,
    candidate_texts: Sequence[str] | None,
    stage: str,
    export_path: Path | None,
    export_evidence: dict[str, Any] | None,
    expected_source_sha256: str | None,
    provenance_source_path: Path | None,
) -> dict[str, Any]:
    source_sha256_before = _sha256_file(source_path)
    normalized_expected_source_sha256 = str(expected_source_sha256 or "").strip().lower()
    issues: list[dict[str, Any]] = []
    resolved_extracted, resolved_snapshot, snapshot = ensure_docx_processing_assets(
        source_path,
        extracted_path=extracted_text_path or get_docx_extracted_text_path(source_path),
        snapshot_path=snapshot_path or get_docx_snapshot_path(source_path),
    )
    cached_snapshot = snapshot
    snapshot_derivation_report, authoritative_snapshot = verify_docx_snapshot_derivation(
        cached_snapshot,
        source_path,
    )
    snapshot_current = _is_snapshot_current(cached_snapshot, source_path)
    if not snapshot_current:
        _add_issue(issues, "snapshot_stale", "原始 Word 已在正文范围冻结后发生变化，必须重新建立快照。")
    if not bool(snapshot_derivation_report.get("ok")):
        _add_issue(
            issues,
            "snapshot_authority_mismatch",
            "缓存的 Word 正文范围与从源 DOCX 重新派生的权威范围不一致，已阻止继续处理。",
            cachedDigest=str(snapshot_derivation_report.get("cachedDigest", "")),
            authoritativeDigest=str(snapshot_derivation_report.get("authoritativeDigest", "")),
            mismatchUnitIndexes=list(snapshot_derivation_report.get("mismatchUnitIndexes", []))[:40],
            authoritativeSourceCurrent=bool(snapshot_derivation_report.get("authoritativeSourceCurrent")),
        )

    # Every semantic decision below is based on a fresh parse of the source
    # OOXML. The persisted snapshot remains useful for diagnostics/caching, but
    # it can no longer define the editable allowlist by itself.
    snapshot = authoritative_snapshot

    editable_units = snapshot.editable_units()
    protected_units = [
        *[unit for unit in snapshot.units if not unit.editable],
        *snapshot.protected_structural_units,
    ]
    heading_units = [
        unit
        for unit in snapshot.units
        if unit.structural_role in {"document_title", "heading", "abstract_heading", "acknowledgement_heading"}
        or _looks_like_heading(
                unit.text,
                style_name=unit.style_name,
                has_numbering=unit.has_numbering,
                outline_level=unit.outline_level,
        )
    ]
    editable_heading_units = [unit for unit in heading_units if unit.editable]
    protected_heading_units = [unit for unit in heading_units if not unit.editable]
    semantic_range_anchor_units = [
        unit
        for unit in snapshot.units
        if bool(getattr(unit, "has_semantic_range_anchor", False))
    ]
    semantic_range_covered_units = [
        unit
        for unit in snapshot.units
        if bool(getattr(unit, "inside_comment_range", False))
    ]
    bookmark_range_interior_units = [
        unit
        for unit in snapshot.units
        if bool(getattr(unit, "inside_bookmark_range", False))
    ]
    semantic_point_units = [
        unit
        for unit in snapshot.units
        if bool(getattr(unit, "has_semantic_point_reference", False))
    ]
    editable_semantic_range_anchor_units = [
        unit for unit in semantic_range_anchor_units if unit.editable
    ]
    editable_semantic_range_covered_units = [
        unit for unit in semantic_range_covered_units if unit.editable
    ]
    editable_bookmark_range_interior_units = [
        unit for unit in bookmark_range_interior_units if unit.editable
    ]
    editable_semantic_point_units = [unit for unit in semantic_point_units if unit.editable]
    illegal_editable_role_units = [
        unit
        for unit in editable_units
        if unit.structural_role not in DOCX_EDITABLE_STRUCTURAL_ROLES
        or unit.edit_eligibility != "eligible"
    ]
    missing_edit_eligibility_evidence_units = [
        unit
        for unit in editable_units
        if not validate_docx_unit_edit_eligibility_evidence(unit)
    ]
    ambiguous_editable_units = [
        unit
        for unit in editable_units
        if unit.structural_role in {"unknown", "ambiguous_non_prose"}
    ]
    invalid_protected_structural_units = [
        unit
        for unit in snapshot.protected_structural_units
        if unit.editable
        or unit.edit_eligibility != "protected"
        or not validate_docx_unit_edit_eligibility_evidence(unit)
    ]
    structural_role_map_digest = get_docx_structural_role_map_digest(snapshot)
    if not editable_units:
        _add_issue(issues, "empty_editable_body", "没有识别到可编辑正文段落，已阻止把整篇 Word 送入模型。")
    if editable_heading_units:
        _add_issue(
            issues,
            "editable_heading_detected",
            "检测到标题被标记为可编辑；标题不得进入模型或被导出回填。",
            count=len(editable_heading_units),
            unitIndexes=[int(unit.unit_index) for unit in editable_heading_units[:20]],
        )
    if snapshot.structural_role_policy_version != DOCX_STRUCTURAL_ROLE_POLICY_VERSION:
        _add_issue(
            issues,
            "structural_role_policy_version_mismatch",
            "Word 结构角色策略版本缺失或过期，必须从源文件重新派生正文资格。",
            expected=DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
            actual=snapshot.structural_role_policy_version,
        )
    if illegal_editable_role_units:
        _add_issue(
            issues,
            "illegal_editable_structural_role",
            "检测到非正文结构角色被标记为可编辑，已阻止进入模型。",
            count=len(illegal_editable_role_units),
            unitIndexes=[unit.unit_index for unit in illegal_editable_role_units[:20]],
            roles=[unit.structural_role for unit in illegal_editable_role_units[:20]],
        )
    if missing_edit_eligibility_evidence_units:
        _add_issue(
            issues,
            "missing_edit_eligibility_evidence",
            "检测到可编辑单元缺少完整的源文件正文资格证据。",
            count=len(missing_edit_eligibility_evidence_units),
            unitIndexes=[unit.unit_index for unit in missing_edit_eligibility_evidence_units[:20]],
        )
    if ambiguous_editable_units:
        _add_issue(
            issues,
            "ambiguous_editable_unit",
            "无法确认为正文的歧义段落不得进入模型。",
            count=len(ambiguous_editable_units),
            unitIndexes=[unit.unit_index for unit in ambiguous_editable_units[:20]],
        )
    if invalid_protected_structural_units:
        _add_issue(
            issues,
            "protected_structural_evidence_invalid",
            "表格等显式保护结构缺少完整资格证据或被错误标记为可编辑。",
            count=len(invalid_protected_structural_units),
            unitIndexes=[unit.unit_index for unit in invalid_protected_structural_units[:20]],
        )
    if not bool(getattr(snapshot, "semantic_range_topology_valid", True)):
        _add_issue(
            issues,
            "semantic_range_topology_invalid",
            "检测到未配对、重复、反序或位于段落之外的书签/批注范围标记；正文范围已 fail closed。",
            issueCount=int(getattr(snapshot, "semantic_range_issue_count", 0)),
            issueCodes=list(getattr(snapshot, "semantic_range_issue_codes", [])),
        )
    if editable_semantic_range_anchor_units:
        _add_issue(
            issues,
            "editable_semantic_range_anchor_detected",
            "检测到带书签或批注范围锚点的段落被标记为可编辑；该类零宽语义范围不得进入模型。",
            count=len(editable_semantic_range_anchor_units),
            unitIndexes=[int(unit.unit_index) for unit in editable_semantic_range_anchor_units[:20]],
        )
    if editable_semantic_range_covered_units:
        _add_issue(
            issues,
            "editable_semantic_range_span_detected",
            "检测到跨段批注范围内的正文被标记为可编辑；批注所覆盖的完整范围不得进入模型。",
            count=len(editable_semantic_range_covered_units),
            unitIndexes=[
                int(unit.unit_index)
                for unit in editable_semantic_range_covered_units[:20]
            ],
        )
    if editable_semantic_point_units:
        _add_issue(
            issues,
            "editable_semantic_point_reference_detected",
            "检测到带批注、脚注或尾注落点的段落被标记为可编辑；零宽引用落点不得进入模型。",
            count=len(editable_semantic_point_units),
            unitIndexes=[int(unit.unit_index) for unit in editable_semantic_point_units[:20]],
        )

    scope_diagnostics = build_docx_scope_diagnostics(snapshot, snapshot_path=resolved_snapshot)
    raw_scope_issues = scope_diagnostics.get("issues")
    for raw_issue in raw_scope_issues if isinstance(raw_scope_issues, list) else []:
        if not isinstance(raw_issue, dict):
            continue
        severity = str(raw_issue.get("severity", "warning") or "warning")
        if severity != "error":
            continue
        _add_issue(
            issues,
            f"scope_{str(raw_issue.get('code', 'invalid'))}",
            str(raw_issue.get("message", "DOCX 正文边界诊断失败。")),
            unit=raw_issue.get("unit"),
        )

    expected_extracted_text = "\n\n".join(unit.text for unit in editable_units)
    extracted_text_exists = resolved_extracted.exists() and resolved_extracted.is_file()
    actual_extracted_text = ""
    if extracted_text_exists:
        try:
            actual_extracted_text = resolved_extracted.read_text(encoding="utf-8")
        except (OSError, UnicodeDecodeError):
            extracted_text_exists = False
    extracted_matches = extracted_text_exists and actual_extracted_text == expected_extracted_text
    if not extracted_text_exists:
        _add_issue(issues, "extracted_text_missing", "模型输入正文文件不存在或无法读取。")
    elif not extracted_matches:
        _add_issue(issues, "extracted_text_scope_mismatch", "模型输入与冻结的可编辑正文单元不完全一致。")

    body_map_validation: dict[str, Any] | None = None
    if body_map is not None:
        body_map_validation = _body_map_contract_issues(
            body_map,
            source_path=source_path,
            snapshot_path=resolved_snapshot,
            authoritative_snapshot=authoritative_snapshot,
            snapshot_derivation_report=snapshot_derivation_report,
            issues=issues,
        )

    candidate_list = [str(value) for value in candidate_texts] if candidate_texts is not None else None
    expected_candidate_count = len(body_map.units) if body_map is not None else len(editable_units)
    candidate_count_matches = candidate_list is None or len(candidate_list) == expected_candidate_count
    if not candidate_count_matches:
        _add_issue(
            issues,
            "candidate_paragraph_count_mismatch",
            "待处理正文段落数与冻结范围不一致。",
            expected=expected_candidate_count,
            actual=len(candidate_list or []),
        )
    candidate_body_map_matches = bool(
        candidate_list is None
        or body_map is None
        or candidate_list == body_map.current_texts()
    )
    if not candidate_body_map_matches:
        _add_issue(
            issues,
            "candidate_body_map_text_mismatch",
            "待处理正文与冻结 body map 当前文本不完全一致。",
        )
    inline_break_indexes = [
        index
        for index, text in enumerate(candidate_list or [])
        if "\n" in text or "\r" in text
    ]
    if inline_break_indexes:
        _add_issue(
            issues,
            "candidate_inline_break",
            "正文回填单元包含换行，可能改变 Word 自然段结构。",
            unitIndexes=inline_break_indexes[:20],
        )

    evidence = dict(export_evidence or {})
    for key, label in (
        ("textIntegrityOk", "正文文本回填一致性"),
        ("protectedTextAuditOk", "保护区文本审计"),
        ("ooxmlIntegrityOk", "Word OOXML 结构审计"),
        ("formatLockAuditOk", "格式保真审计"),
    ):
        if key in evidence and evidence.get(key) is not True:
            _add_issue(issues, f"export_{key}", f"{label}未通过。")

    source_sha256_after = _sha256_file(source_path)
    source_generation_stable = source_sha256_before == source_sha256_after
    if normalized_expected_source_sha256 and (
        source_sha256_before != normalized_expected_source_sha256
        or source_sha256_after != normalized_expected_source_sha256
    ):
        _add_issue(
            issues,
            "source_anchor_sha256_mismatch",
            "内容寻址的源 Word 与本次导出绑定的 SHA-256 不一致。",
            expectedSourceSha256=normalized_expected_source_sha256,
            sourceSha256Before=source_sha256_before,
            sourceSha256After=source_sha256_after,
        )
    if not source_generation_stable:
        _add_issue(
            issues,
            "source_changed_during_contract",
            "执行正文与格式契约期间，内容寻址的源 Word 发生变化。",
            sourceSha256Before=source_sha256_before,
            sourceSha256After=source_sha256_after,
        )

    error_count = sum(1 for issue in issues if issue.get("severity") == "error")
    warning_count = sum(1 for issue in issues if issue.get("severity") == "warning")
    model_input_matches = bool(
        snapshot_derivation_report.get("ok")
        and bool(getattr(snapshot, "semantic_range_topology_valid", True))
        and snapshot.structural_role_policy_version == DOCX_STRUCTURAL_ROLE_POLICY_VERSION
        and not illegal_editable_role_units
        and not missing_edit_eligibility_evidence_units
        and not ambiguous_editable_units
        and not invalid_protected_structural_units
        and not editable_semantic_range_covered_units
        and extracted_matches
        and candidate_count_matches
        and candidate_body_map_matches
        and not inline_break_indexes
        and (body_map_validation is None or bool(body_map_validation.get("ok")))
    )
    contract: dict[str, Any] = {
        "version": DOCUMENT_EDIT_CONTRACT_VERSION,
        "policy": DOCUMENT_EDIT_POLICY,
        "stage": str(stage or "snapshot"),
        "createdAt": _utc_now(),
        "sourceKind": ".docx",
        "sourcePath": str(source_path),
        "sourceSha256": authoritative_snapshot.source_sha256,
        "expectedSourceSha256": normalized_expected_source_sha256,
        "sourceGenerationStable": source_generation_stable,
        "provenanceSourcePath": (
            str(provenance_source_path.resolve())
            if provenance_source_path is not None
            else ""
        ),
        "snapshotPath": str(resolved_snapshot),
        "snapshotVersion": int(snapshot.version),
        "snapshotCurrent": snapshot_current,
        "snapshotAuthorityVerified": bool(snapshot_derivation_report.get("ok")),
        "cachedSnapshotDigest": str(snapshot_derivation_report.get("cachedDigest", "")),
        "authoritativeSnapshotDigest": str(snapshot_derivation_report.get("authoritativeDigest", "")),
        "snapshotAuthorityMismatchUnitIndexes": list(snapshot_derivation_report.get("mismatchUnitIndexes", []))[:40],
        "structuralRolePolicyVersion": snapshot.structural_role_policy_version,
        "structuralInventoryVersion": snapshot.structural_inventory_version,
        "structuralRoleMapDigest": structural_role_map_digest,
        "structuralRoleAuthorityVerified": bool(
            snapshot_derivation_report.get("ok")
            and snapshot.structural_role_policy_version == DOCX_STRUCTURAL_ROLE_POLICY_VERSION
            and not illegal_editable_role_units
            and not missing_edit_eligibility_evidence_units
            and not invalid_protected_structural_units
        ),
        "editableRoleAllowlist": sorted(DOCX_EDITABLE_STRUCTURAL_ROLES),
        "editableRoleCounts": {
            role: sum(1 for unit in editable_units if unit.structural_role == role)
            for role in sorted(DOCX_EDITABLE_STRUCTURAL_ROLES)
        },
        "illegalEditableRoleCount": len(illegal_editable_role_units),
        "ambiguousEditableUnitCount": len(ambiguous_editable_units),
        "missingEditEligibilityEvidenceCount": len(missing_edit_eligibility_evidence_units),
        "protectedStructuralUnitCount": len(snapshot.protected_structural_units),
        "protectedTableParagraphCount": sum(
            1
            for unit in snapshot.protected_structural_units
            if unit.structural_role == "table_content"
        ),
        "protectedStructuralEvidenceInvalidCount": len(invalid_protected_structural_units),
        "scopeDigest": _scope_digest(snapshot),
        "formatDigest": _format_digest(source_path),
        "formatLockPolicy": FORMAT_LOCK_POLICY,
        "formatLockApplicable": True,
        "formatLockReady": bool(
            snapshot_current
            and snapshot_derivation_report.get("ok")
            and bool(getattr(snapshot, "semantic_range_topology_valid", True))
            and not editable_heading_units
            and not illegal_editable_role_units
            and not missing_edit_eligibility_evidence_units
            and not ambiguous_editable_units
            and not invalid_protected_structural_units
            and not editable_semantic_range_anchor_units
            and not editable_semantic_range_covered_units
            and not editable_semantic_point_units
        ),
        "scopeReady": bool(error_count == 0),
        "editableUnitCount": len(editable_units),
        "protectedUnitCount": len(protected_units),
        "headingCount": len(heading_units),
        "protectedHeadingCount": len(protected_heading_units),
        "editableHeadingCount": len(editable_heading_units),
        "semanticRangeCount": int(getattr(snapshot, "semantic_range_count", 0)),
        "bookmarkRangeCount": int(getattr(snapshot, "bookmark_range_count", 0)),
        "commentRangeCount": int(getattr(snapshot, "comment_range_count", 0)),
        "semanticRangeTopologyValid": bool(
            getattr(snapshot, "semantic_range_topology_valid", True)
        ),
        "semanticRangeIssueCount": int(
            getattr(snapshot, "semantic_range_issue_count", 0)
        ),
        "semanticRangeIssueCodes": list(
            getattr(snapshot, "semantic_range_issue_codes", [])
        ),
        "semanticRangeAnchorUnitCount": len(semantic_range_anchor_units),
        "protectedSemanticRangeAnchorUnitCount": len(
            [unit for unit in semantic_range_anchor_units if not unit.editable]
        ),
        "editableSemanticRangeAnchorUnitCount": len(editable_semantic_range_anchor_units),
        "semanticRangeCoveredUnitCount": len(semantic_range_covered_units),
        "protectedSemanticRangeCoveredUnitCount": len(
            [unit for unit in semantic_range_covered_units if not unit.editable]
        ),
        "editableSemanticRangeCoveredUnitCount": len(
            editable_semantic_range_covered_units
        ),
        "bookmarkRangeInteriorUnitCount": len(bookmark_range_interior_units),
        "protectedBookmarkRangeInteriorUnitCount": len(
            [unit for unit in bookmark_range_interior_units if not unit.editable]
        ),
        "editableBookmarkRangeInteriorUnitCount": len(
            editable_bookmark_range_interior_units
        ),
        "semanticPointReferenceUnitCount": len(semantic_point_units),
        "protectedSemanticPointReferenceUnitCount": len(
            [unit for unit in semantic_point_units if not unit.editable]
        ),
        "editableSemanticPointReferenceUnitCount": len(editable_semantic_point_units),
        "modelInputUnitCount": len(candidate_list) if candidate_list is not None else len(editable_units),
        "modelInputMatchesEditableUnits": model_input_matches,
        "extractedTextPath": str(resolved_extracted),
        "extractedTextMatchesEditableUnits": extracted_matches,
        "bodyMapPresent": body_map is not None,
        "bodyMapReady": body_map_validation is None or bool(body_map_validation.get("ok")),
        "scopeDiagnosticsOk": bool(scope_diagnostics.get("ok")),
        "exportPath": str(export_path.resolve()) if export_path is not None else "",
        "exportSha256": _sha256_file(export_path.resolve()) if export_path is not None and export_path.exists() else "",
        "exportEvidence": evidence,
        "ready": error_count == 0 and model_input_matches,
        "issueCount": error_count,
        "warningCount": warning_count,
        "issues": issues[:80],
        "truncatedIssues": max(0, len(issues) - 80),
    }
    return contract


def _build_plain_text_contract(source_path: Path, *, stage: str) -> dict[str, Any]:
    text = source_path.read_text(encoding="utf-8")
    paragraphs = [item.strip() for item in text.replace("\r\n", "\n").replace("\r", "\n").split("\n\n") if item.strip()]
    return {
        "version": DOCUMENT_EDIT_CONTRACT_VERSION,
        "policy": "plain_text_content_guard",
        "stage": str(stage or "snapshot"),
        "createdAt": _utc_now(),
        "sourceKind": source_path.suffix.lower() or ".txt",
        "sourcePath": str(source_path),
        "sourceSha256": _sha256_file(source_path),
        "snapshotPath": "",
        "snapshotVersion": 0,
        "snapshotCurrent": True,
        "structuralRolePolicyVersion": 0,
        "structuralInventoryVersion": 0,
        "structuralRoleMapDigest": "",
        "structuralRoleAuthorityVerified": True,
        "editableRoleAllowlist": ["plain_text"],
        "editableRoleCounts": {"plain_text": len(paragraphs)},
        "illegalEditableRoleCount": 0,
        "ambiguousEditableUnitCount": 0,
        "missingEditEligibilityEvidenceCount": 0,
        "protectedStructuralUnitCount": 0,
        "protectedTableParagraphCount": 0,
        "protectedStructuralEvidenceInvalidCount": 0,
        "scopeDigest": _sha256_json(paragraphs),
        "formatDigest": "",
        "formatLockPolicy": "not_applicable_plain_text",
        "formatLockApplicable": False,
        "formatLockReady": True,
        "scopeReady": True,
        "editableUnitCount": len(paragraphs),
        "protectedUnitCount": 0,
        "headingCount": 0,
        "protectedHeadingCount": 0,
        "editableHeadingCount": 0,
        "semanticRangeCount": 0,
        "bookmarkRangeCount": 0,
        "commentRangeCount": 0,
        "semanticRangeTopologyValid": True,
        "semanticRangeIssueCount": 0,
        "semanticRangeIssueCodes": [],
        "semanticRangeAnchorUnitCount": 0,
        "protectedSemanticRangeAnchorUnitCount": 0,
        "editableSemanticRangeAnchorUnitCount": 0,
        "semanticRangeCoveredUnitCount": 0,
        "protectedSemanticRangeCoveredUnitCount": 0,
        "editableSemanticRangeCoveredUnitCount": 0,
        "bookmarkRangeInteriorUnitCount": 0,
        "protectedBookmarkRangeInteriorUnitCount": 0,
        "editableBookmarkRangeInteriorUnitCount": 0,
        "semanticPointReferenceUnitCount": 0,
        "protectedSemanticPointReferenceUnitCount": 0,
        "editableSemanticPointReferenceUnitCount": 0,
        "modelInputUnitCount": len(paragraphs),
        "modelInputMatchesEditableUnits": True,
        "extractedTextPath": str(source_path),
        "extractedTextMatchesEditableUnits": True,
        "bodyMapPresent": False,
        "bodyMapReady": True,
        "scopeDiagnosticsOk": True,
        "exportPath": "",
        "exportSha256": "",
        "exportEvidence": {},
        "ready": True,
        "issueCount": 0,
        "warningCount": 1,
        "issues": [
            {
                "code": "plain_text_has_no_word_format_layer",
                "severity": "warning",
                "message": "TXT 不包含 Word 版式层；标题冻结由运行时文本保护规则执行。",
            }
        ],
        "truncatedIssues": 0,
    }


def build_document_edit_contract(
    source_path: Path | str,
    *,
    snapshot_path: Path | None = None,
    extracted_text_path: Path | None = None,
    body_map: DocxBodyMap | None = None,
    candidate_texts: Sequence[str] | None = None,
    stage: str = "snapshot",
    export_path: Path | None = None,
    export_evidence: dict[str, Any] | None = None,
    expected_source_sha256: str | None = None,
    provenance_source_path: Path | None = None,
    report_path: Path | None = None,
) -> dict[str, Any]:
    normalized_source = Path(source_path).resolve()
    if not normalized_source.exists() or not normalized_source.is_file():
        raise ValueError(f"Source document does not exist: {normalized_source}")
    if normalized_source.suffix.lower() == ".docx":
        report = _build_docx_contract(
            normalized_source,
            snapshot_path=snapshot_path.resolve() if snapshot_path is not None else None,
            extracted_text_path=extracted_text_path.resolve() if extracted_text_path is not None else None,
            body_map=body_map,
            candidate_texts=candidate_texts,
            stage=stage,
            export_path=export_path,
            export_evidence=export_evidence,
            expected_source_sha256=expected_source_sha256,
            provenance_source_path=(
                provenance_source_path.resolve()
                if provenance_source_path is not None
                else None
            ),
        )
    else:
        report = _build_plain_text_contract(normalized_source, stage=stage)
    if report_path is not None:
        write_document_edit_contract(report, report_path)
    return report


def assert_document_edit_contract_ready(report: dict[str, Any], *, label: str) -> None:
    if bool(report.get("ready")):
        return
    raw_issues = report.get("issues")
    issues = raw_issues if isinstance(raw_issues, list) else []
    messages = [
        str(issue.get("message", "")).strip()
        for issue in issues
        if isinstance(issue, dict)
        and issue.get("severity") == "error"
        and str(issue.get("message", "")).strip()
    ]
    detail = "；".join(messages[:3]) or "正文范围或格式锁证据不完整。"
    report_path = str(report.get("reportPath", "") or "").strip()
    suffix = f" 报告：{report_path}" if report_path else ""
    raise ValueError(f"{label}已拦截：只改正文/格式固定契约未通过。{detail}{suffix}")
