#!/usr/bin/env python3
"""Regression for pre-model DOCX format-sensitive text anchors."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
import sys
from typing import Any

from docx import Document  # type: ignore[import]
from docx.enum.style import WD_STYLE_TYPE  # type: ignore[import]
from docx.oxml import OxmlElement  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import fyadr_round_service as service  # noqa: E402
import docx_bodymap as bodymap  # noqa: E402
import docx_pipeline as pipeline  # noqa: E402
from docx_audit import audit_docx_format_lock  # noqa: E402


REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_model_format_anchor_regression_report.json"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _append_rpr_element(run: Any, name: str, **attributes: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    element = OxmlElement(f"w:{name}")
    for attribute_name, value in attributes.items():
        element.set(qn(f"w:{attribute_name}"), str(value))
    r_pr.append(element)


def _rpr_sha256(run: Any) -> str:
    r_pr = getattr(run._r, "rPr", None)
    payload = b"" if r_pr is None else r_pr.xml.encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _swap_once(value: str, first: str, second: str) -> str:
    placeholder = "@@FYADR_FORMAT_SWAP_PROBE@@"
    _assert(placeholder not in value, "format swap placeholder collided with fixture text")
    return value.replace(first, placeholder, 1).replace(second, first, 1).replace(placeholder, second, 1)


def _run_previous_format_cache_regression(source_path: Path, work_dir: Path) -> list[str]:
    """Prove the immediately previous anchor caches are never authoritative."""

    anchor_text = "西文字体锚点甲"
    legacy_snapshot_path = work_dir / "format_anchor_v20.snapshot.json"
    legacy_extracted_path = work_dir / "format_anchor_v20.extracted.txt"
    legacy_scope_path = work_dir / "format_anchor_v20.scope.json"
    legacy_body_map_path = work_dir / "format_anchor_v7.body_map.json"

    authoritative_snapshot = pipeline.build_docx_snapshot(source_path)
    _assert(authoritative_snapshot.version == pipeline.DOCX_SNAPSHOT_VERSION == 22, "authoritative format snapshot version drifted")
    stale_snapshot_payload = authoritative_snapshot.to_dict()
    stale_snapshot_payload["version"] = 20
    cleared_snapshot_anchor = False
    for raw_unit in stale_snapshot_payload.get("units", []):
        if not isinstance(raw_unit, dict) or anchor_text not in str(raw_unit.get("text", "")):
            continue
        raw_anchors = raw_unit.get("format_anchors", [])
        if not isinstance(raw_anchors, list):
            continue
        retained = [
            anchor
            for anchor in raw_anchors
            if not (isinstance(anchor, dict) and str(anchor.get("text", "")) == anchor_text)
        ]
        cleared_snapshot_anchor = len(retained) != len(raw_anchors)
        raw_unit["format_anchors"] = retained
        raw_unit["format_anchor_ambiguous"] = any(
            bool(anchor.get("ambiguous"))
            for anchor in retained
            if isinstance(anchor, dict)
        )
        break
    _assert(cleared_snapshot_anchor, "v20 fixture did not remove the new font-only anchor")
    legacy_snapshot_path.write_text(
        json.dumps(stale_snapshot_payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    _, resolved_snapshot_path, refreshed_snapshot = pipeline.ensure_docx_processing_assets(
        source_path,
        extracted_path=legacy_extracted_path,
        snapshot_path=legacy_snapshot_path,
        scope_diagnostics_path=legacy_scope_path,
    )
    _assert(resolved_snapshot_path == legacy_snapshot_path, "v20 cache refresh changed the requested snapshot path")
    _assert(refreshed_snapshot.version == 22, "v20 format-anchor snapshot was reused instead of re-derived")
    refreshed_unit = next(unit for unit in refreshed_snapshot.units if anchor_text in unit.text)
    _assert(
        any(str(anchor.get("text", "")) == anchor_text for anchor in refreshed_unit.format_anchors),
        "v22 snapshot re-derivation did not restore the font-only anchor",
    )
    persisted_snapshot = pipeline._load_docx_snapshot(legacy_snapshot_path)
    _assert(persisted_snapshot is not None and persisted_snapshot.version == 22, "refreshed v22 snapshot was not persisted")

    fresh_body_map = bodymap.build_docx_body_map(
        source_path,
        snapshot_path=legacy_snapshot_path,
        prompt_profile="cn",
        round_number=1,
    )
    _assert(
        fresh_body_map.version == bodymap.DOCX_BODY_MAP_VERSION == 9
        and fresh_body_map.snapshot_version == 22,
        "fresh body-map did not bind the v22 snapshot under schema v9",
    )
    stale_body_map_payload = fresh_body_map.to_dict()
    stale_body_map_payload["version"] = 7
    stale_body_map_payload["snapshot_version"] = 20
    cleared_body_map_anchor = False
    for raw_unit in stale_body_map_payload.get("units", []):
        if not isinstance(raw_unit, dict) or anchor_text not in str(raw_unit.get("original_text", "")):
            continue
        raw_anchors = raw_unit.get("format_anchors", [])
        if not isinstance(raw_anchors, list):
            continue
        retained = [
            anchor
            for anchor in raw_anchors
            if not (isinstance(anchor, dict) and str(anchor.get("text", "")) == anchor_text)
        ]
        cleared_body_map_anchor = len(retained) != len(raw_anchors)
        raw_unit["format_anchors"] = retained
        break
    _assert(cleared_body_map_anchor, "v7 fixture did not remove the new font-only anchor")
    stale_body_map = bodymap.docx_body_map_from_payload(stale_body_map_payload)
    _assert(stale_body_map is not None, "v7 format body-map fixture was unreadable")
    stale_body_map.scope_signature = bodymap._build_scope_signature(stale_body_map.units)
    bodymap.save_docx_body_map(stale_body_map, legacy_body_map_path)

    loaded_stale_body_map = bodymap.load_docx_body_map(legacy_body_map_path)
    _assert(
        loaded_stale_body_map is not None and loaded_stale_body_map.version == 7,
        "body-map loader did not expose the v7 cache for authority validation",
    )
    stale_validation = bodymap.validate_docx_body_map(
        loaded_stale_body_map,
        source_path=source_path,
        snapshot_path=legacy_snapshot_path,
    )
    stale_codes = {
        str(issue.get("code", ""))
        for issue in stale_validation.get("blockingIssues", [])
        if isinstance(issue, dict)
    }
    _assert(stale_validation.get("ok") is False, "v7 body-map cache was treated as authoritative")
    _assert(
        {"body_map_version_stale", "snapshot_version_mismatch", "snapshot_scope_signature_drift"}.issubset(stale_codes),
        f"v7 body-map rejection evidence drifted: {sorted(stale_codes)}",
    )

    rebuilt_body_map = bodymap.build_docx_body_map(
        source_path,
        snapshot_path=legacy_snapshot_path,
        prompt_profile="cn",
        round_number=1,
    )
    bodymap.save_docx_body_map(rebuilt_body_map, legacy_body_map_path)
    loaded_rebuilt_body_map = bodymap.load_docx_body_map(legacy_body_map_path)
    _assert(
        loaded_rebuilt_body_map is not None
        and loaded_rebuilt_body_map.version == 9
        and loaded_rebuilt_body_map.snapshot_version == 22,
        "production body-map rebuild did not replace v7/v20 cache metadata",
    )
    rebuilt_unit = next(unit for unit in loaded_rebuilt_body_map.units if anchor_text in unit.original_text)
    _assert(
        any(str(anchor.get("text", "")) == anchor_text for anchor in rebuilt_unit.format_anchors),
        "v9 body-map rebuild did not restore the font-only anchor",
    )
    return ["v20 snapshot and v7 body-map caches re-derive to v22/v9 with the font-only anchor restored"]


def _run_local_run_format_regression(work_dir: Path) -> list[str]:
    checks: list[str] = []
    source_path = work_dir / "local_run_format_source.docx"
    export_path = work_dir / "local_run_format_export.docx"
    snapshot_path = work_dir / "local_run_format_snapshot.json"

    source_document = Document()
    heading = source_document.add_paragraph("1 引言")
    heading.style = source_document.styles["Heading 1"]

    def add_case(
        label: str,
        token_a: str,
        token_b: str,
        apply_format: Any,
        expected_kinds: set[str],
    ) -> dict[str, Any]:
        paragraph = source_document.add_paragraph()
        paragraph.add_run(f"围绕{label}开展前置论证，并将")
        first = paragraph.add_run(token_a)
        apply_format(first)
        paragraph.add_run("作为局部格式示例；随后以")
        second = paragraph.add_run(token_b)
        apply_format(second)
        paragraph.add_run("继续说明控制策略的稳定性、适用条件与论证边界。")
        return {
            "label": label,
            "paragraphIndex": len(source_document.paragraphs) - 1,
            "tokens": (token_a, token_b),
            "expectedKinds": expected_kinds,
        }

    def apply_western_font(run: Any) -> None:
        _append_rpr_element(run, "rFonts", ascii="Aptos Display", hAnsi="Aptos Display")

    def apply_font_size(run: Any) -> None:
        _append_rpr_element(run, "sz", val="30")
        _append_rpr_element(run, "szCs", val="30")

    def apply_spacing_position(run: Any) -> None:
        _append_rpr_element(run, "spacing", val="24")
        _append_rpr_element(run, "position", val="3")

    def apply_strike_small_caps(run: Any) -> None:
        _append_rpr_element(run, "strike")
        _append_rpr_element(run, "smallCaps")

    def apply_east_asia_font(run: Any) -> None:
        _append_rpr_element(run, "rFonts", eastAsia="黑体", hint="eastAsia")

    cases = [
        add_case("西文字体差异", "西文字体锚点甲", "西文字体锚点乙", apply_western_font, {"font_family"}),
        add_case("字号差异", "字号锚点甲", "字号锚点乙", apply_font_size, {"font_size"}),
        add_case(
            "字符间距与位置差异",
            "字距位置锚点甲",
            "字距位置锚点乙",
            apply_spacing_position,
            {"character_spacing", "character_position"},
        ),
        add_case(
            "删除线与小型大写差异",
            "删除线小型大写锚点甲",
            "删除线小型大写锚点乙",
            apply_strike_small_caps,
            {"strike", "small_caps"},
        ),
        add_case("东亚字体差异", "东亚字体锚点甲", "东亚字体锚点乙", apply_east_asia_font, {"font_family"}),
    ]

    uniform_paragraph = source_document.add_paragraph()
    for text in ("统一格式段落仍可", "完整改写且不会", "因多个 run 被过度冻结。"):
        run = uniform_paragraph.add_run(text)
        _append_rpr_element(
            run,
            "rFonts",
            ascii="Times New Roman",
            hAnsi="Times New Roman",
            eastAsia="宋体",
            hint="eastAsia",
        )
        _append_rpr_element(run, "sz", val="24")
        _append_rpr_element(run, "szCs", val="24")
    uniform_index = len(source_document.paragraphs) - 1

    # One compact descriptor probe covers the less common range properties.
    descriptor_probe = source_document.add_paragraph().add_run("综合字符属性探针")
    _append_rpr_element(
        descriptor_probe,
        "rFonts",
        ascii="Arial",
        hAnsi="Arial",
        eastAsia="黑体",
        cs="Arial",
        asciiTheme="majorAscii",
        hAnsiTheme="minorHAnsi",
        eastAsiaTheme="majorEastAsia",
        cstheme="minorBidi",
        hint="eastAsia",
    )
    for element_name, attributes in (
        ("sz", {"val": "28"}),
        ("szCs", {"val": "28"}),
        ("spacing", {"val": "20"}),
        ("position", {"val": "2"}),
        ("kern", {"val": "24"}),
        ("fitText", {"val": "720", "id": "1"}),
        ("b", {"val": "0"}),
        ("i", {"val": "false"}),
        ("u", {"val": "none"}),
        ("strike", {"val": "0"}),
        ("dstrike", {}),
        ("caps", {}),
        ("smallCaps", {}),
        ("outline", {}),
        ("shadow", {}),
        ("emboss", {}),
        ("imprint", {}),
        ("vanish", {}),
        ("webHidden", {}),
        ("bdr", {"val": "single", "sz": "4", "space": "0", "color": "222222"}),
        ("shd", {"val": "clear", "color": "auto", "fill": "E5E7EB"}),
    ):
        _append_rpr_element(descriptor_probe, element_name, **attributes)
    descriptor = pipeline._format_sensitive_run_descriptor(descriptor_probe)
    _assert(descriptor is not None, "comprehensive character property descriptor was empty")
    descriptor_values = dict(descriptor[0])
    descriptor_kinds = set(descriptor[1])
    required_descriptor_kinds = {
        "font_family",
        "font_size",
        "character_spacing",
        "character_position",
        "kerning",
        "fit_text",
        "bold",
        "italic",
        "underline",
        "strike",
        "double_strike",
        "caps",
        "small_caps",
        "outline",
        "shadow",
        "emboss",
        "imprint",
        "hidden",
        "web_hidden",
        "text_border",
        "text_shading",
    }
    _assert(
        required_descriptor_kinds.issubset(descriptor_kinds),
        f"character property descriptor missed kinds: {sorted(required_descriptor_kinds - descriptor_kinds)}",
    )
    _assert("false" in descriptor_values["bold"], "explicit false bold was discarded")
    _assert("false" in descriptor_values["italic"], "explicit false italic was discarded")
    _assert("val=none" in descriptor_values["underline"], "explicit false underline was discarded")
    _assert(descriptor_values["strike"] == "false", "explicit false strike was discarded")
    source_document.paragraphs[-1]._element.getparent().remove(source_document.paragraphs[-1]._element)
    checks.append("requested direct/effective rPr properties retain normalized values, including explicit false")

    cancellation_style = source_document.styles.add_style("FYADR inherited emphasis", WD_STYLE_TYPE.CHARACTER)
    cancellation_style.font.bold = True
    cancellation_style.font.italic = True
    cancellation_style.font.underline = True
    cancellation_style.font.strike = True
    cancellation_paragraph = source_document.add_paragraph()
    for text in ("继承格式前文", "局部取消格式", "继承格式后文"):
        run = cancellation_paragraph.add_run(text)
        run.style = cancellation_style
    cancellation_run = cancellation_paragraph.runs[1]
    cancellation_run.bold = False
    cancellation_run.italic = False
    cancellation_run.underline = False
    cancellation_run.font.strike = False
    cancellation_anchors = pipeline._extract_format_sensitive_anchors(cancellation_paragraph)
    _assert(
        [anchor.get("text") for anchor in cancellation_anchors] == ["局部取消格式"],
        f"explicit format cancellation did not create one local anchor: {cancellation_anchors}",
    )
    checks.append("explicit false emphasis locally overrides inherited character formatting")

    coverage_probe = Document().add_paragraph()
    long_plain_text = "较长的普通正文负责承载论文中的主要论证内容，因此字符覆盖范围应优先于被人为切碎的局部格式 run 数量。"
    coverage_probe.add_run(long_plain_text)
    for token in ("短甲", "短乙", "短丙"):
        run = coverage_probe.add_run(token)
        _append_rpr_element(run, "smallCaps")
    coverage_anchors = pipeline._extract_format_sensitive_anchors(coverage_probe)
    _assert(
        [anchor.get("text") for anchor in coverage_anchors] == ["短甲短乙短丙"],
        f"run-count bias misclassified the long ordinary range as local formatting: {coverage_anchors}",
    )
    checks.append("dominant format baseline prefers character coverage over fragmented run count")

    source_document.save(str(source_path))
    snapshot = pipeline.build_docx_snapshot(source_path)
    snapshot_path.write_text(json.dumps(snapshot.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    checks.extend(_run_previous_format_cache_regression(source_path, work_dir))

    for case in cases:
        paragraph = source_document.paragraphs[int(case["paragraphIndex"])]
        anchors = pipeline._extract_format_sensitive_anchors(paragraph)
        tokens = tuple(case["tokens"])
        _assert(
            [anchor.get("text") for anchor in anchors] == list(tokens),
            f"{case['label']} local run anchors drifted: {anchors}",
        )
        captured_kinds = {
            str(kind)
            for anchor in anchors
            for kind in anchor.get("kinds", [])
        }
        _assert(
            set(case["expectedKinds"]).issubset(captured_kinds),
            f"{case['label']} missed anchor kinds: {sorted(set(case['expectedKinds']) - captured_kinds)}",
        )
        source_text = paragraph.text
        invalid_cases = (
            (source_text.replace(tokens[0], "", 1), "format_anchor_missing"),
            (source_text.replace(tokens[0], f"{tokens[0]}与{tokens[0]}", 1), "format_anchor_duplicated"),
            (_swap_once(source_text, tokens[0], tokens[1]), "format_anchor_order_changed"),
        )
        for invalid_text, expected_code in invalid_cases:
            issues = pipeline.validate_format_sensitive_anchors(invalid_text, anchors)
            codes = {str(issue.get("code", "")) for issue in issues}
            _assert(expected_code in codes, f"{case['label']} accepted {expected_code}: {issues}")
            evidence_text = json.dumps(issues, ensure_ascii=False, sort_keys=True)
            _assert(all(token not in evidence_text for token in tokens), f"{case['label']} leaked anchor prose in evidence")
            _assert("anchorSha256" in evidence_text and "kinds" in evidence_text, f"{case['label']} lost hash/kind evidence")
    checks.append("font, size, spacing/position, strike/smallCaps, and East Asian font anchors reject delete/copy/swap")

    _assert(
        pipeline._extract_format_sensitive_anchors(source_document.paragraphs[uniform_index]) == [],
        "uniformly formatted multi-run paragraph was over-frozen",
    )

    export_document = Document(str(source_path))
    for case in cases:
        paragraph = export_document.paragraphs[int(case["paragraphIndex"])]
        token_a, token_b = tuple(case["tokens"])
        source_anchor_hashes = {
            run.text: _rpr_sha256(run)
            for run in paragraph.runs
            if run.text in {token_a, token_b}
        }
        rewritten = (
            f"重新组织{case['label']}的论证后，{token_a}仍对应原字符范围；"
            f"进一步分析表明，{token_b}继续用于标识局部格式语义。"
        )
        pipeline._replace_paragraph_text(paragraph, rewritten)
        exported_anchor_runs = {
            run.text: _rpr_sha256(run)
            for run in paragraph.runs
            if run.text in {token_a, token_b}
        }
        _assert(
            exported_anchor_runs == source_anchor_hashes,
            f"{case['label']} legal rewrite moved text away from its original rPr run",
        )

    uniform_rewrite = "统一格式段落已经完成整体重写，多个同签名 run 仍可共同承载新的学术正文。"
    pipeline._replace_paragraph_text(export_document.paragraphs[uniform_index], uniform_rewrite)
    _assert(export_document.paragraphs[uniform_index].text == uniform_rewrite, "uniform format paragraph was not editable")
    export_document.save(str(export_path))

    format_lock = audit_docx_format_lock(
        export_path,
        source_path=source_path,
        snapshot_path=snapshot_path,
    )
    _assert(format_lock.get("ok") is True, f"legal local-run rewrite failed OOXML format lock: {format_lock}")
    checks.append("legal surrounding rewrites preserve local run ownership and pass the OOXML format lock")
    checks.append("uniformly formatted multi-run paragraph remains fully editable")
    return checks


def run_regression() -> dict[str, Any]:
    checks: list[str] = []

    source = "前段论证保留格式术语，后段内容继续改写。综上所述，相关分析用于验证这一处理。"
    protected = service.protect_structure_tokens(source, exact_anchors=["格式术语"])
    _assert("格式术语" not in protected.text, "styled anchor leaked into model text")
    _assert(list(protected.token_types.values()) == ["FMT"], f"styled anchor type drifted: {protected.token_types}")
    placeholder = next(iter(protected.tokens))
    _assert(service.restore_structure_tokens(protected.text, protected.tokens) == source, "styled anchor did not restore exactly")
    for invalid in (protected.text.replace(placeholder, ""), protected.text.replace(placeholder, placeholder * 2)):
        try:
            service.validate_structure_placeholders(invalid, protected.tokens, "anchor_probe")
            raise AssertionError("missing/duplicated format placeholder was accepted")
        except ValueError:
            pass
    for invalid_output in (
        source.replace("格式术语", "", 1),
        source.replace("格式术语", "格式术语与格式术语", 1),
    ):
        try:
            service.validate_immutable_text_anchors(source, invalid_output, ["格式术语"], "anchor_probe")
            raise AssertionError("post-restore format-anchor count drift was accepted")
        except ValueError:
            pass
    checks.append("exact styled text becomes one position-bound FMT placeholder")

    citation = service.protect_structure_tokens("依据[2]开展验证。", exact_anchors=["[2]"])
    _assert(list(citation.token_types.values()) == ["REF"], "citation anchor should reuse the stronger REF token")
    containing = service.protect_structure_tokens(
        "采用YOLOv8 参数完成检测。",
        exact_anchors=["YOLOv8 参数"],
    )
    _assert(list(containing.token_types.values()) == ["FMT"], "format anchor should absorb nested technical tokens")
    _assert(service.restore_structure_tokens(containing.text, containing.tokens) == "采用YOLOv8 参数完成检测。", "nested token restore drifted")
    checks.append("overlapping citation/technical protection resolves to one non-nested placeholder")

    work_dir = ROOT_DIR / "finish" / "regression" / "docx_model_format_anchor"
    work_dir.mkdir(parents=True, exist_ok=True)
    checks.extend(_run_local_run_format_regression(work_dir))
    input_path = work_dir / "input.txt"
    output_path = work_dir / "output.txt"
    manifest_path = work_dir / "manifest.json"
    input_path.write_text(source, encoding="utf-8")

    calls: list[dict[str, Any]] = []

    def transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
        calls.append(
            {
                "chunkId": chunk_id,
                "rawAnchorVisible": "格式术语" in chunk_text,
                "fmtCount": chunk_text.count("@@FYADR_FMT_"),
                "guardVisible": "[WORD FORMAT ANCHOR LOCK]" in prompt_input,
            }
        )
        if len(calls) == 1:
            # First completion keeps the placeholder but also emits the raw
            # anchor, which would duplicate styled text after restoration.
            # Candidate validation must reject it before artifact acceptance.
            return chunk_text.replace("，", "格式术语，", 1)
        return (
            chunk_text.replace("前段论证保留", "重组论证后仍保留")
            .replace("后段内容继续改写", "其余正文已经完成改写")
            .replace("综上所述，", "")
        )

    original_update_round = service.update_round
    try:
        service.update_round = lambda **kwargs: {"rounds": [kwargs]}  # type: ignore[assignment]
        result = service.run_round(
            doc_id="finish/regression/docx_model_format_anchor/input.txt",
            round_number=1,
            input_path=input_path,
            output_path=output_path,
            manifest_path=manifest_path,
            transform=transform,
            prompt_profile="cn",
            immutable_format_anchors={0: ["格式术语"]},
        )
    finally:
        service.update_round = original_update_round

    _assert(len(calls) == 2, f"format placeholder violation should cause one bounded retry: {calls}")
    _assert(all(not item["rawAnchorVisible"] for item in calls), "model saw raw styled anchor text")
    _assert(all(item["fmtCount"] == 1 and item["guardVisible"] for item in calls), "model format guard was incomplete")
    output_text = output_path.read_text(encoding="utf-8")
    _assert(output_text.count("格式术语") == 1, "accepted output did not restore the anchor exactly once")
    _assert("重组论证后仍保留" in output_text and "其余正文已经完成改写" in output_text, "prose around anchor was not rewritten")
    validation_events = result.get("quality_summary", {}).get("validationEvents", [])
    if not validation_events:
        # The persisted compare is the authoritative retry ledger in older
        # quality payloads; keep the regression compatible with that schema.
        compare_payload = json.loads(Path(str(result["compare_path"])).read_text(encoding="utf-8"))
        validation_events = compare_payload.get("validationEvents", [])
    _assert(
        any(item.get("event") == "validation-retry" for item in validation_events if isinstance(item, dict)),
        "format placeholder retry was not recorded",
    )
    run_audit_text = json.dumps(result.get("run_audit", {}), ensure_ascii=False, sort_keys=True)
    _assert("格式术语" not in run_audit_text, "run audit persisted raw format-anchor text")
    _assert("immutableFormatAnchorPlanSha256" in run_audit_text, "run audit lost hashed anchor-plan evidence")
    checks.append("round worker hides anchors, retries invalid output, restores exact text, and rewrites surrounding prose")

    split_input_path = work_dir / "split_input.txt"
    split_input_path.write_text("前置前置前置锚点甲。锚点乙后置后置后置", encoding="utf-8")
    split_calls = 0

    def forbidden_transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
        nonlocal split_calls
        split_calls += 1
        return chunk_text

    try:
        service.run_round(
            doc_id="finish/regression/docx_model_format_anchor/split_input.txt",
            round_number=1,
            input_path=split_input_path,
            output_path=work_dir / "split_output.txt",
            manifest_path=work_dir / "split_manifest.json",
            transform=forbidden_transform,
            prompt_profile="cn",
            chunk_limit=10,
            immutable_format_anchors={0: ["锚点甲。锚点乙"]},
        )
        raise AssertionError("chunker accepted an anchor split across model chunks")
    except ValueError as exc:
        _assert("stopped before API execution" in str(exc), f"split-anchor error drifted: {exc}")
    _assert(split_calls == 0, "split format anchor reached the model transform")
    checks.append("anchor split across chunks hard-fails with zero model calls")

    report = {"ok": True, "checks": checks, "modelCallCount": len(calls), "splitModelCallCount": split_calls}
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> int:
    report = run_regression()
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
