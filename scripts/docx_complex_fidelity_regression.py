#!/usr/bin/env python3
"""Strict complex-DOCX round-trip fidelity regression.

Builds a local DOCX containing mixed run formatting, paragraph layout,
numbering, a hyperlink, fields, OMML, content controls, a merged table,
headers/footers, and mixed Chinese/English text. It executes the real body-map round/export path in
``preserve_original`` mode and independently compares OOXML structure.
"""

from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import json
import os
from pathlib import Path
import sys
import tempfile
from typing import Any, Callable
import zipfile

ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

# Keep this regression isolated from the user's real provider/config state.
os.environ["FYADR_APP_CONFIG_DIR"] = tempfile.mkdtemp(prefix="fyadr_docx_complex_")

from docx import Document  # noqa: E402
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402
from lxml import etree  # noqa: E402

import app_config  # noqa: E402
import app_service  # noqa: E402
from docx_audit import audit_docx_format_lock, audit_docx_ooxml_integrity  # noqa: E402
from docx_bodymap import (  # noqa: E402
    load_docx_body_map,
    update_docx_body_map_texts,
    validate_docx_body_map,
)
from document_edit_contract import build_document_edit_contract  # noqa: E402
from docx_pipeline import _load_docx_snapshot, _replace_paragraph_text, get_docx_snapshot_path  # noqa: E402
from round_helper import run_document_round  # noqa: E402


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_complex_fidelity_regression_report.json"


def _set_run_font(
    run: Any,
    *,
    size_pt: float,
    ascii_font: str,
    east_asia_font: str,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = ascii_font
    run.font.bold = bold
    run.font.italic = italic
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia_font)


def _append_field(paragraph: Any, instruction: str, result: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = f" {instruction.strip()} "
    instruction_run._r.append(instruction_node)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(result)
    _set_run_font(
        result_run,
        size_pt=9,
        ascii_font="Times New Roman",
        east_asia_font="宋体",
    )

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _append_omml_equation(paragraph: Any, expression: str) -> None:
    equation = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = expression
    math_run.append(math_text)
    equation.append(math_run)
    paragraph._p.append(equation)


def _add_hyperlink(paragraph: Any, text: str, address: str) -> None:
    relationship_id = paragraph.part.relate_to(address, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    run_properties.extend((fonts, color, underline))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_top_level_content_control(document: Any) -> None:
    """Add a body-level SDT that is protected wholesale, not rewritten."""

    content_control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "FidelityProtectedControl")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "fyadr-protected-sdt")
    properties.extend((alias, tag))
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = "受控内容：顶层内容控件必须原样保留。"
    run.append(text_node)
    paragraph.append(run)
    content.append(paragraph)
    content_control.extend((properties, content))
    body = document._element.body
    section_properties = body.sectPr
    if section_properties is None:
        body.append(content_control)
    else:
        body.insert(body.index(section_properties), content_control)


def _append_inline_content_control(paragraph: Any) -> None:
    paragraph.add_run("内联受控内容前文；")
    content_control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "fyadr-inline-protected-sdt")
    properties.append(tag)
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = "内联控件锁定文字"
    run.append(text_node)
    content.append(run)
    content_control.extend((properties, content))
    paragraph._p.append(content_control)
    paragraph.add_run("后文也必须保持原位。")


def _create_complex_source(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3.1)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.9)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.1)
    section.different_first_page_header_footer = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("复杂页眉 · Complex Header")
    _set_run_font(
        header_run,
        size_pt=9,
        ascii_font="Arial",
        east_asia_font="黑体",
        italic=True,
    )
    first_header = section.first_page_header.paragraphs[0]
    first_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_header_run = first_header.add_run("首页专用页眉")
    _set_run_font(
        first_header_run,
        size_pt=10,
        ascii_font="Arial",
        east_asia_font="楷体",
        bold=True,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    _append_field(footer, "PAGE", "1")
    footer.add_run(" 页")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("复杂格式保真测试论文")
    _set_run_font(
        title_run,
        size_pt=20,
        ascii_font="Times New Roman",
        east_asia_font="黑体",
        bold=True,
    )

    document.add_paragraph("1 引言", style="Heading 1")

    multi_run = document.add_paragraph()
    multi_run.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    multi_format = multi_run.paragraph_format
    multi_format.first_line_indent = Cm(0.83)
    multi_format.left_indent = Cm(0.2)
    multi_format.right_indent = Cm(0.1)
    multi_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    multi_format.line_spacing = Pt(19)
    multi_format.space_before = Pt(3)
    multi_format.space_after = Pt(5)
    _set_run_font(
        multi_run.add_run("本段包含"),
        size_pt=12,
        ascii_font="SimSun",
        east_asia_font="宋体",
    )
    _set_run_font(
        multi_run.add_run("粗体术语"),
        size_pt=14,
        ascii_font="SimHei",
        east_asia_font="黑体",
        bold=True,
    )
    _set_run_font(
        multi_run.add_run("、"),
        size_pt=12,
        ascii_font="SimSun",
        east_asia_font="宋体",
    )
    _set_run_font(
        multi_run.add_run("italic English"),
        size_pt=11,
        ascii_font="Times New Roman",
        east_asia_font="宋体",
        italic=True,
    )
    _set_run_font(
        multi_run.add_run(
            " 与中英混排正文，用于验证多 run 字体、字号和段落缩进。"
            "综上所述，相关说明用于验证格式映射。"
        ),
        size_pt=12,
        ascii_font="Times New Roman",
        east_asia_font="仿宋",
    )

    hyperlink_paragraph = document.add_paragraph()
    hyperlink_paragraph.paragraph_format.first_line_indent = Cm(0.83)
    hyperlink_paragraph.add_run("访问 ")
    _add_hyperlink(hyperlink_paragraph, "项目文档", "https://example.com/docs")
    hyperlink_paragraph.add_run(
        " 获取完整说明，这一正文段落应允许改写。"
        "综上所述，相关说明用于验证链接格式映射。"
    )

    numbered = document.add_paragraph(style="List Number")
    numbered.paragraph_format.left_indent = Cm(0.74)
    numbered.paragraph_format.first_line_indent = Cm(-0.37)
    numbered.add_run("自动编号正文包含足够信息，因此应参与改写并保留编号定义与缩进。")

    mixed = document.add_paragraph()
    mixed.paragraph_format.first_line_indent = Cm(0.83)
    _set_run_font(
        mixed.add_run("本研究使用 "),
        size_pt=12,
        ascii_font="Times New Roman",
        east_asia_font="宋体",
    )
    model_run = mixed.add_run("YOLOv8")
    _set_run_font(
        model_run,
        size_pt=11,
        ascii_font="Consolas",
        east_asia_font="宋体",
        bold=True,
    )
    model_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    mixed.add_run(" 处理 640px 图像，并在中文语境中保留 English terms，相关设置依据")
    citation_run = mixed.add_run("[2]")
    _set_run_font(
        citation_run,
        size_pt=8,
        ascii_font="Times New Roman",
        east_asia_font="宋体",
    )
    citation_run.font.superscript = True
    mixed.add_run("执行。")

    field_paragraph = document.add_paragraph()
    field_paragraph.add_run("生成日期：")
    _append_field(field_paragraph, "DATE \\@ \"yyyy-MM-dd\"", "2026-07-17")

    equation_paragraph = document.add_paragraph()
    equation_paragraph.add_run("公式对象：")
    _append_omml_equation(equation_paragraph, "E=mc²")

    inline_control_paragraph = document.add_paragraph()
    _append_inline_content_control(inline_control_paragraph)

    inline_tab_paragraph = document.add_paragraph()
    inline_tab_paragraph.add_run("行内制表符左侧")
    inline_tab_run = inline_tab_paragraph.add_run()
    inline_tab_run.add_tab()
    inline_tab_run.add_text("行内制表符右侧")

    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "合并表头"
    table.cell(0, 2).text = "指标"
    table.cell(1, 0).text = "模型"
    table.cell(1, 1).text = "数据集"
    table.cell(1, 2).text = "mAP"
    table.cell(2, 0).text = "YOLOv8"
    table.cell(2, 1).text = "Mixed-CN/EN"
    table.cell(2, 2).text = "87.2%"
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(3.2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("致 谢", style="Heading 1")
    document.add_paragraph("感谢导师与同学在测试过程中的帮助。")
    document.add_paragraph("参考文献", style="Heading 1")
    document.add_paragraph("[1] Example Author. Complex DOCX Fidelity. 2026.")
    _append_top_level_content_control(document)
    document.save(str(path))
    return path


def _rewrite_transform(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
    text = str(chunk_text)
    for source, replacement in (
        ("本段包含", "该段保留"),
        ("访问 ", "请访问 "),
        ("完整说明", "详细说明"),
        ("足够信息", "完整信息"),
        ("本研究使用", "本工作采用"),
        ("综上所述，", ""),
    ):
        text = text.replace(source, replacement)
    return text or chunk_text


def _canonical_xml(xml: bytes, *, strip_text_nodes: bool = False) -> bytes:
    root = etree.fromstring(xml)
    if strip_text_nodes:
        for text_node in root.xpath(".//w:t", namespaces={"w": W_NS}):
            parent = text_node.getparent()
            if parent is not None:
                parent.remove(text_node)
    return etree.tostring(root, method="c14n", with_comments=True)


def _hash_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _element_signature(element: Any, *, strip_text_nodes: bool = False) -> str:
    return _hash_bytes(_canonical_xml(element.xml.encode("utf-8"), strip_text_nodes=strip_text_nodes))


def _read_parts(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path, "r") as archive:
        return {item.filename: archive.read(item.filename) for item in archive.infolist() if not item.is_dir()}


def _xml_part_signature(parts: dict[str, bytes], name: str, *, strip_text_nodes: bool = False) -> str:
    value = parts.get(name)
    if value is None:
        return "missing"
    return _hash_bytes(_canonical_xml(value, strip_text_nodes=strip_text_nodes))


def _rewrite_zip_part(
    source_path: Path,
    target_path: Path,
    part_name: str,
    transform: Callable[[bytes], bytes],
) -> None:
    with zipfile.ZipFile(source_path, "r") as source, zipfile.ZipFile(target_path, "w") as target:
        found = False
        for item in source.infolist():
            value = source.read(item.filename)
            if item.filename == part_name:
                value = transform(value)
                found = True
            target.writestr(item, value)
    if not found:
        raise AssertionError(f"DOCX part not found for corruption probe: {part_name}")


def _remove_zip_part(source_path: Path, target_path: Path, part_name: str) -> None:
    with zipfile.ZipFile(source_path, "r") as source, zipfile.ZipFile(target_path, "w") as target:
        found = False
        for item in source.infolist():
            if item.filename == part_name:
                found = True
                continue
            target.writestr(item, source.read(item.filename))
    if not found:
        raise AssertionError(f"DOCX part not found for removal probe: {part_name}")


def _corrupt_hyperlink_target(xml: bytes) -> bytes:
    root = etree.fromstring(xml)
    relationships = root.findall(f"{{{REL_NS}}}Relationship")
    target = next((item for item in relationships if "example.com/docs" in str(item.get("Target", ""))), None)
    if target is None:
        raise AssertionError("hyperlink relationship missing from corruption probe")
    target.set("Target", "https://invalid.example/fidelity-drift")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _corrupt_table_style(xml: bytes) -> bytes:
    root = etree.fromstring(xml)
    table_properties = root.find(f".//{W}tblPr")
    if table_properties is None:
        raise AssertionError("table properties missing from corruption probe")
    table_style = table_properties.find(f"{W}tblStyle")
    if table_style is None:
        table_style = etree.Element(f"{W}tblStyle")
        table_properties.insert(0, table_style)
    table_style.set(f"{W}val", "LightShading-Accent1")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _corrupt_field_instruction(xml: bytes) -> bytes:
    root = etree.fromstring(xml)
    instruction = next(
        (
            item
            for item in root.findall(f".//{W}instrText")
            if "DATE" in str(item.text or "")
        ),
        None,
    )
    if instruction is None:
        raise AssertionError("DATE field instruction missing from corruption probe")
    instruction.text = ' AUTHOR \\@ "yyyy-MM-dd" '
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _corrupt_top_level_content_control(xml: bytes) -> bytes:
    root = etree.fromstring(xml)
    text_node = root.find(f".//{W}sdt/{W}sdtContent/{W}p/{W}r/{W}t")
    if text_node is None:
        raise AssertionError("top-level content control missing from corruption probe")
    text_node.text = "受控内容已被恶意篡改。"
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _find_paragraph(document: Any, needle: str) -> Any:
    return next(paragraph for paragraph in document.paragraphs if needle in paragraph.text)


def _run_regression() -> dict[str, Any]:
    failures: list[str] = []
    checks: list[str] = []
    work_dir = ROOT_DIR / "finish" / "regression" / "docx_complex_fidelity"
    work_dir.mkdir(parents=True, exist_ok=True)
    # Keep the fixture basename unique. DOCX processing assets currently live in
    # the shared intermediate directory and use the source stem in their names;
    # reusing ``complex_source.docx`` here would make this regression overwrite
    # the independent multi-round fidelity fixture's snapshot.
    source_path = _create_complex_source(work_dir / "complex_fidelity_source.docx")
    export_path = work_dir / "complex_export.docx"

    app_config.save_app_config({**app_config.load_app_config(), "formatMode": "preserve_original"})
    round_result = run_document_round(
        source_path,
        _rewrite_transform,
        round_number=1,
        prompt_profile="cn",
    )
    output_path = Path(str(round_result["output_path"]))
    snapshot_path = get_docx_snapshot_path(source_path)
    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        raise AssertionError("complex regression did not create a DOCX snapshot")

    captured_anchor_kinds = {
        str(kind)
        for unit in snapshot.units
        for anchor in getattr(unit, "format_anchors", [])
        for kind in anchor.get("kinds", [])
    }
    required_anchor_kinds = {"bold", "italic", "superscript", "color", "hyperlink"}
    if not required_anchor_kinds.issubset(captured_anchor_kinds):
        failures.append(
            "snapshot missed format-sensitive anchor kinds: "
            f"{sorted(required_anchor_kinds - captured_anchor_kinds)}"
        )
    else:
        checks.append("snapshot records bold, italic, superscript, color, and hyperlink text anchors")

    def assert_anchor_rewrite_blocked(
        needle: str,
        replacement_builder: Callable[[str], str],
        expected_code: str,
        label: str,
    ) -> None:
        probe_document = Document(str(source_path))
        probe_paragraph = _find_paragraph(probe_document, needle)
        original_text = probe_paragraph.text
        original_xml = probe_paragraph._p.xml
        blocked_error = ""
        try:
            _replace_paragraph_text(probe_paragraph, replacement_builder(original_text))
        except ValueError as exc:
            blocked_error = str(exc)
        if expected_code not in blocked_error:
            failures.append(f"{label} was not rejected with {expected_code}: {blocked_error!r}")
            return
        if probe_paragraph.text != original_text or probe_paragraph._p.xml != original_xml:
            failures.append(f"{label} mutated the paragraph before the hard failure")
            return
        checks.append(f"{label} hard-fails atomically")

    assert_anchor_rewrite_blocked(
        "粗体术语",
        lambda value: value.replace("粗体术语", "普通替代术语", 1),
        "format_anchor_missing",
        "bold anchor removal",
    )
    assert_anchor_rewrite_blocked(
        "粗体术语",
        lambda value: value.replace("粗体术语", "粗体术语与粗体术语", 1),
        "format_anchor_missing",
        "composite local-format anchor member duplication",
    )
    assert_anchor_rewrite_blocked(
        "粗体术语",
        lambda value: value.replace("粗体术语", "@@ANCHOR_A@@", 1)
        .replace("italic English", "粗体术语", 1)
        .replace("@@ANCHOR_A@@", "italic English", 1),
        "format_anchor_missing",
        "composite local-format anchor member reordering",
    )
    assert_anchor_rewrite_blocked(
        "项目文档",
        lambda value: value.replace("项目文档", "普通文字", 1),
        "format_anchor_missing",
        "hyperlink label rewrite",
    )
    assert_anchor_rewrite_blocked(
        "YOLOv8",
        lambda value: value.replace("[2]", "", 1),
        "format_anchor_missing",
        "superscript citation removal",
    )

    leading_probe = Document()
    leading_paragraph = leading_probe.add_paragraph()
    leading_anchor = leading_paragraph.add_run("段首锚点")
    leading_anchor.bold = True
    leading_paragraph.add_run("后的普通正文允许改写。")
    leading_before = leading_paragraph._p.xml
    try:
        _replace_paragraph_text(leading_paragraph, "新增前缀段首锚点后的正文。")
        failures.append("paragraph-leading styled anchor accepted an unmappable prefix")
    except ValueError as exc:
        if "format_anchor_boundary_shifted" not in str(exc) or leading_paragraph._p.xml != leading_before:
            failures.append(f"paragraph-leading anchor boundary failure drifted: {exc}")
        else:
            checks.append("paragraph-leading styled anchor blocks unmappable prefix insertion")

    trailing_probe = Document()
    trailing_paragraph = trailing_probe.add_paragraph()
    trailing_paragraph.add_run("普通正文位于")
    trailing_anchor = trailing_paragraph.add_run("段末锚点")
    trailing_anchor.italic = True
    trailing_before = trailing_paragraph._p.xml
    try:
        _replace_paragraph_text(trailing_paragraph, "改写后的普通正文位于段末锚点新增后缀")
        failures.append("paragraph-trailing styled anchor accepted an unmappable suffix")
    except ValueError as exc:
        if "format_anchor_boundary_shifted" not in str(exc) or trailing_paragraph._p.xml != trailing_before:
            failures.append(f"paragraph-trailing anchor boundary failure drifted: {exc}")
        else:
            checks.append("paragraph-trailing styled anchor blocks unmappable suffix insertion")

    adjacent_probe = Document()
    adjacent_paragraph = adjacent_probe.add_paragraph()
    adjacent_bold = adjacent_paragraph.add_run("粗体锚点")
    adjacent_bold.bold = True
    adjacent_super = adjacent_paragraph.add_run("上标锚点")
    adjacent_super.font.superscript = True
    adjacent_paragraph.add_run("之后仍有普通正文。")
    adjacent_before = adjacent_paragraph._p.xml
    try:
        _replace_paragraph_text(adjacent_paragraph, "粗体锚点插入内容上标锚点之后正文已改。")
        failures.append("adjacent styled runs accepted text without an ordinary run slot")
    except ValueError as exc:
        if "format_anchor_missing" not in str(exc) or adjacent_paragraph._p.xml != adjacent_before:
            failures.append(f"adjacent composite anchor failure drifted: {exc}")
        else:
            checks.append("adjacent differently styled runs form one immutable composite anchor")

    relocation_document = Document(str(source_path))
    relocation_paragraph = _find_paragraph(relocation_document, "粗体术语")
    _replace_paragraph_text(
        relocation_paragraph,
        "本段包含粗体术语、italic English，并在重新组织的中英混排论证中保留原字符范围，"
        "段末普通正文也已显著改写。",
    )
    relocation_run_texts = [run.text for run in relocation_paragraph.runs]
    if "粗体术语" not in relocation_run_texts or "italic English" not in relocation_run_texts:
        failures.append(f"large rewrite detached styled anchors from their runs: {relocation_run_texts}")
    else:
        checks.append("large surrounding rewrite keeps styled text attached to its original runs")

    body_map_path = Path(str(round_result.get("body_map_path", "")))
    body_map = load_docx_body_map(body_map_path)
    if body_map is None:
        failures.append("complex regression did not persist a versioned DOCX body map")
    else:
        anchored_unit_index = next(
            (index for index, unit in enumerate(body_map.units) if unit.format_anchors),
            None,
        )
        if anchored_unit_index is None:
            failures.append("DOCX body map lost all format-sensitive anchors")
        else:
            anchored_unit = body_map.units[anchored_unit_index]
            anchor_text = str(anchored_unit.format_anchors[0].get("text", ""))
            invalid_texts = body_map.current_texts()
            invalid_texts[anchored_unit_index] = invalid_texts[anchored_unit_index].replace(
                anchor_text,
                "格式锚点已被改写",
                1,
            )
            invalid_body_map = update_docx_body_map_texts(body_map, invalid_texts, round_number=1)
            invalid_validation = validate_docx_body_map(
                invalid_body_map,
                source_path=source_path,
                snapshot_path=snapshot_path,
            )
            invalid_codes = {
                str(issue.get("code", ""))
                for issue in invalid_validation.get("blockingIssues", [])
                if isinstance(issue, dict)
            }
            if "format_anchor_missing" not in invalid_codes or invalid_validation.get("ok") is not False:
                failures.append(f"body-map validation accepted a changed styled anchor: {sorted(invalid_codes)}")
            else:
                checks.append("body-map validation blocks styled-anchor text drift")

            invalid_contract = build_document_edit_contract(
                source_path,
                snapshot_path=snapshot_path,
                body_map=invalid_body_map,
                candidate_texts=invalid_body_map.current_texts(),
                stage="format_anchor_negative_probe",
            )
            contract_codes = {
                str(issue.get("code", ""))
                for issue in invalid_contract.get("issues", [])
                if isinstance(issue, dict)
            }
            if invalid_contract.get("ready") is not False or "body_map_format_anchor_missing" not in contract_codes:
                failures.append(f"document contract accepted styled-anchor drift: {sorted(contract_codes)}")
            else:
                checks.append("document edit contract blocks styled-anchor drift before export")

    export_result = app_service.export_round_output(str(output_path), str(export_path), "docx")

    for key in ("guardIssueCount", "auditIssueCount", "ooxmlAuditIssueCount", "formatLockIssueCount"):
        if int(export_result.get(key, 0) or 0) != 0:
            failures.append(f"{key} should be zero, got {export_result.get(key)!r}")
    if str(export_result.get("formatMode", "")) != "preserve_original":
        failures.append(f"formatMode should be preserve_original, got {export_result.get('formatMode')!r}")

    source_document = Document(str(source_path))
    export_document = Document(str(export_path))
    if len(source_document.paragraphs) != len(export_document.paragraphs):
        failures.append("top-level paragraph count changed")
    paragraph_drift: list[int] = []
    for index, (source_paragraph, export_paragraph) in enumerate(zip(source_document.paragraphs, export_document.paragraphs)):
        if _element_signature(source_paragraph._p, strip_text_nodes=True) != _element_signature(export_paragraph._p, strip_text_nodes=True):
            paragraph_drift.append(index)
    if paragraph_drift:
        failures.append(f"paragraph non-text OOXML drifted at indexes {paragraph_drift}")
    else:
        checks.append("all top-level paragraph non-text OOXML signatures match")

    if len(source_document.tables) != len(export_document.tables):
        failures.append("table count changed")
    table_drift = [
        index
        for index, (source_table, export_table) in enumerate(zip(source_document.tables, export_document.tables))
        if _element_signature(source_table._tbl) != _element_signature(export_table._tbl)
    ]
    if table_drift:
        failures.append(f"table OOXML drifted at indexes {table_drift}")
    else:
        checks.append("merged table structure, formatting, and text match")

    source_section = source_document._element.body.sectPr
    export_section = export_document._element.body.sectPr
    if source_section is None or export_section is None or _element_signature(source_section) != _element_signature(export_section):
        failures.append("section/page-layout OOXML drifted")
    else:
        checks.append("section margins and header/footer references match")

    source_parts = _read_parts(source_path)
    export_parts = _read_parts(export_path)
    critical_parts = ["word/styles.xml", "word/numbering.xml", "word/settings.xml", "word/_rels/document.xml.rels"]
    critical_parts.extend(sorted(name for name in source_parts if name.startswith(("word/header", "word/footer")) and name.endswith(".xml")))
    critical_drift = [
        name
        for name in critical_parts
        if _xml_part_signature(source_parts, name) != _xml_part_signature(export_parts, name)
    ]
    if critical_drift:
        failures.append(f"critical OOXML parts drifted: {critical_drift}")
    else:
        checks.append("styles, numbering, settings, relationships, headers, and footers match")

    source_multi = _find_paragraph(source_document, "粗体术语")
    export_multi = _find_paragraph(export_document, "粗体术语")
    source_run_texts = [run.text for run in source_multi.runs]
    export_run_texts = [run.text for run in export_multi.runs]
    if "粗体术语" not in export_run_texts or "italic English" not in export_run_texts:
        failures.append(f"styled anchors moved across runs: {export_run_texts}")
    bold_run = next((run for run in export_multi.runs if run.text == "粗体术语"), None)
    italic_run = next((run for run in export_multi.runs if run.text == "italic English"), None)
    if bold_run is None or bold_run.bold is not True or bold_run.font.size is None or abs(bold_run.font.size.pt - 14) > 0.01:
        failures.append("bold 14pt run formatting was not preserved")
    if italic_run is None or italic_run.italic is not True or italic_run.font.size is None or abs(italic_run.font.size.pt - 11) > 0.01:
        failures.append("italic 11pt English run formatting was not preserved")
    if (
        source_run_texts == export_run_texts
        or not export_multi.text.startswith("本段包含粗体术语、italic English")
    ):
        failures.append("multi-run body text was not actually rewritten")
    else:
        checks.append("rewritten text changed while bold/italic run anchors stayed intact")

    source_hyperlink_paragraph = _find_paragraph(source_document, "项目文档")
    export_hyperlink_paragraph = _find_paragraph(export_document, "项目文档")
    source_hyperlinks = [item for item in source_hyperlink_paragraph.iter_inner_content() if type(item).__name__ == "Hyperlink"]
    export_hyperlinks = [item for item in export_hyperlink_paragraph.iter_inner_content() if type(item).__name__ == "Hyperlink"]
    if len(source_hyperlinks) != 1 or len(export_hyperlinks) != 1:
        failures.append("hyperlink wrapper count changed")
    else:
        source_hyperlink = source_hyperlinks[0]
        export_hyperlink = export_hyperlinks[0]
        if source_hyperlink.address != export_hyperlink.address or export_hyperlink.address != "https://example.com/docs":
            failures.append("hyperlink target changed")
        if source_hyperlink.text != export_hyperlink.text or export_hyperlink.text != "项目文档":
            failures.append("hyperlink display text changed or duplicated")
        source_link_rpr = [_element_signature(run._r.get_or_add_rPr()) for run in source_hyperlink.runs]
        export_link_rpr = [_element_signature(run._r.get_or_add_rPr()) for run in export_hyperlink.runs]
        if source_link_rpr != export_link_rpr:
            failures.append("hyperlink run formatting changed")
        else:
            checks.append("hyperlink target, wrapper, text, color, and underline match")
    if not export_hyperlink_paragraph.text.startswith("请访问 ") or "获取详细说明" not in export_hyperlink_paragraph.text:
        failures.append("text surrounding hyperlink was not rewritten")

    field_units = [unit for unit in snapshot.units if unit.has_field_code]
    if not field_units or any(unit.editable for unit in field_units):
        failures.append("field-code paragraph was not protected")
    else:
        checks.append("main-document field paragraph remains protected")
    math_units = [unit for unit in snapshot.units if getattr(unit, "has_math", False)]
    if not math_units or any(unit.editable for unit in math_units):
        failures.append("OMML equation paragraph was not protected")
    else:
        checks.append("OMML equation paragraph remains protected")
    complex_inline_units = [
        unit for unit in snapshot.units if getattr(unit, "has_complex_inline", False)
    ]
    if not complex_inline_units or any(unit.editable for unit in complex_inline_units):
        failures.append("inline content-control paragraph was not protected")
    else:
        checks.append("inline content-control paragraph remains protected")
    inline_tab_units = [unit for unit in snapshot.units if "\t" in unit.source_text()]
    if (
        not inline_tab_units
        or any(unit.editable for unit in inline_tab_units)
        or any(not getattr(unit, "has_complex_inline", False) for unit in inline_tab_units)
    ):
        failures.append("inline tab paragraph entered the editable model scope")
    else:
        checks.append("inline tab and manual-break structure remains protected before model input")

    inline_probe = Document()
    inline_probe_paragraph = inline_probe.add_paragraph()
    inline_probe_paragraph.add_run("左侧")
    tab_run = inline_probe_paragraph.add_run()
    tab_run.add_tab()
    tab_run.add_text("右侧")
    unsupported_inline_blocked = False
    try:
        _replace_paragraph_text(inline_probe_paragraph, "改写后的左右正文")
    except ValueError:
        unsupported_inline_blocked = "\t" in inline_probe_paragraph.text
    if not unsupported_inline_blocked:
        failures.append("unsupported inline tab structure was flattened or silently accepted")
    else:
        checks.append("unsupported inline tab structure hard-fails without flattening")

    lock_report = audit_docx_format_lock(export_path, source_path=source_path, snapshot_path=snapshot_path)
    ooxml_report = audit_docx_ooxml_integrity(export_path, source_path=source_path, snapshot_path=snapshot_path)
    if not lock_report.get("ok"):
        failures.append(f"strict format lock failed: {lock_report.get('issues')}")
    if not ooxml_report.get("ok"):
        failures.append(f"OOXML integrity audit failed: {ooxml_report.get('issues')}")

    relationship_corrupt_path = work_dir / "corrupt_hyperlink_target.docx"
    _rewrite_zip_part(
        export_path,
        relationship_corrupt_path,
        "word/_rels/document.xml.rels",
        _corrupt_hyperlink_target,
    )
    relationship_report = audit_docx_ooxml_integrity(
        relationship_corrupt_path,
        source_path=source_path,
        snapshot_path=snapshot_path,
    )
    relationship_issue_types = {str(item.get("type", "")) for item in relationship_report.get("issues", [])}
    if "relationship_semantics_changed" not in relationship_issue_types:
        failures.append(f"OOXML audit missed hyperlink target drift: {sorted(relationship_issue_types)}")
    else:
        checks.append("OOXML audit detects hyperlink relationship target drift")

    missing_part_path = work_dir / "missing_referenced_package_part.docx"
    _remove_zip_part(export_path, missing_part_path, "docProps/app.xml")
    missing_part_report = audit_docx_ooxml_integrity(
        missing_part_path,
        source_path=source_path,
        snapshot_path=snapshot_path,
    )
    missing_part_issue_types = {str(item.get("type", "")) for item in missing_part_report.get("issues", [])}
    if "package_part_missing" not in missing_part_issue_types:
        failures.append(f"OOXML audit missed package part removal: {sorted(missing_part_issue_types)}")
    if "relationship_target_missing" not in missing_part_issue_types:
        failures.append(f"OOXML audit missed dangling relationship target: {sorted(missing_part_issue_types)}")
    if {"package_part_missing", "relationship_target_missing"}.issubset(missing_part_issue_types):
        checks.append("OOXML audit detects removed package parts and dangling internal relationships")

    table_corrupt_path = work_dir / "corrupt_table_style.docx"
    _rewrite_zip_part(export_path, table_corrupt_path, "word/document.xml", _corrupt_table_style)
    table_lock_report = audit_docx_format_lock(
        table_corrupt_path,
        source_path=source_path,
        snapshot_path=snapshot_path,
    )
    table_issue_types = {str(item.get("type", "")) for item in table_lock_report.get("issues", [])}
    if "format_lock_table_changed" not in table_issue_types:
        failures.append(f"format lock missed table style drift: {sorted(table_issue_types)}")
    else:
        checks.append("format lock detects protected table style drift")

    field_corrupt_path = work_dir / "corrupt_field_instruction.docx"
    _rewrite_zip_part(export_path, field_corrupt_path, "word/document.xml", _corrupt_field_instruction)
    field_lock_report = audit_docx_format_lock(
        field_corrupt_path,
        source_path=source_path,
        snapshot_path=snapshot_path,
    )
    field_issue_types = {str(item.get("type", "")) for item in field_lock_report.get("issues", [])}
    if "format_lock_violation" not in field_issue_types:
        failures.append(f"format lock missed protected field-instruction drift: {sorted(field_issue_types)}")
    else:
        checks.append("format lock detects protected field-instruction drift")

    content_control_corrupt_path = work_dir / "corrupt_content_control.docx"
    _rewrite_zip_part(
        export_path,
        content_control_corrupt_path,
        "word/document.xml",
        _corrupt_top_level_content_control,
    )
    content_control_report = audit_docx_ooxml_integrity(
        content_control_corrupt_path,
        source_path=source_path,
        snapshot_path=snapshot_path,
    )
    content_control_issue_types = {
        str(item.get("type", "")) for item in content_control_report.get("issues", [])
    }
    if "protected_body_container_changed" not in content_control_issue_types:
        failures.append(
            "OOXML audit missed protected top-level content-control drift: "
            f"{sorted(content_control_issue_types)}"
        )
    else:
        checks.append("OOXML audit detects protected top-level content-control drift")

    original_lock_audit = app_service.audit_docx_format_lock
    blocked_stage = ""
    try:
        def _forced_format_drift(*args: Any, **kwargs: Any) -> dict[str, Any]:
            return {
                "ok": False,
                "issueCount": 1,
                "issues": [{"type": "forced_format_drift"}],
                "editableChecked": len(snapshot.editable_units()),
                "reportPath": str(kwargs.get("report_path", "")),
            }

        app_service.audit_docx_format_lock = _forced_format_drift
        try:
            app_service.export_round_output(
                str(output_path),
                str(work_dir / "forced_format_drift_export.docx"),
                "docx",
            )
        except app_service.ExportRoundError as exc:
            blocked_stage = str(exc.export_failure.get("stage", ""))
    finally:
        app_service.audit_docx_format_lock = original_lock_audit
    if blocked_stage != "format-lock":
        failures.append(f"export did not hard-block failed format lock; stage={blocked_stage!r}")
    else:
        checks.append("export hard-blocks a failed preserve-format audit")

    return {
        "ok": not failures,
        "createdAt": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "sourcePath": str(source_path),
        "outputPath": str(output_path),
        "exportPath": str(export_path),
        "snapshotPath": str(snapshot_path),
        "editableUnitCount": len(snapshot.editable_units()),
        "totalTextUnitCount": snapshot.total_text_unit_count,
        "checks": checks,
        "failures": failures,
        "export": export_result,
        "formatLock": {
            "ok": bool(lock_report.get("ok")),
            "editableChecked": int(lock_report.get("editableChecked", 0) or 0),
            "protectedChecked": int(lock_report.get("protectedChecked", 0) or 0),
            "tableChecked": int(lock_report.get("tableChecked", 0) or 0),
            "issueCount": int(lock_report.get("issueCount", 0) or 0),
        },
        "ooxmlAuditIssueCount": int(ooxml_report.get("issueCount", 0) or 0),
        "negativeChecks": {
            "relationshipIssueTypes": sorted(relationship_issue_types),
            "missingPackagePartIssueTypes": sorted(missing_part_issue_types),
            "tableIssueTypes": sorted(table_issue_types),
            "fieldIssueTypes": sorted(field_issue_types),
            "contentControlIssueTypes": sorted(content_control_issue_types),
            "unsupportedInlineBlocked": unsupported_inline_blocked,
            "forcedExportBlockedStage": blocked_stage,
        },
    }


def main() -> int:
    report = _run_regression()
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
