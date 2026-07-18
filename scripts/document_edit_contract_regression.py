from __future__ import annotations

import json
import os
import sys
import tempfile
import zipfile
from pathlib import Path


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document  # noqa: E402
from docx.enum.style import WD_STYLE_TYPE  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402

import app_config  # noqa: E402
import app_service  # noqa: E402
from document_edit_contract import build_document_edit_contract  # noqa: E402
from docx_audit import audit_docx_export, audit_docx_format_lock, audit_docx_ooxml_integrity  # noqa: E402
from docx_bodymap import build_docx_body_map, update_docx_body_map_texts  # noqa: E402
from docx_pipeline import (  # noqa: E402
    _assign_unit_edit_decision,
    _is_snapshot_current,
    _load_docx_snapshot,
    ensure_docx_processing_assets,
    get_docx_extracted_text_path,
    get_docx_snapshot_path,
)
from round_helper import run_document_round  # noqa: E402


REFERENCE_VARIANT_CASES = (
    {
        "name": "reference_list",
        "heading": "Reference List",
        "style": None,
        "entries": (
            "[1] Smith J. Reliable control under temporal constraints[J]. Systems Journal, 2021, 12(3): 44-59.",
            "[2] Brown A. Boundary-aware evaluation for control systems[M]. Academic Press, 2022.",
        ),
    },
    {
        "name": "sources_consulted",
        "heading": "Sources Consulted",
        "style": None,
        "entries": (
            "Smith, J. (2021). Reliable Control under Temporal Constraints. Systems Journal, 12(3), 44-59.",
            "Brown, A. (2022). Boundary-aware Evaluation for Control Systems. Academic Press.",
        ),
    },
    {
        "name": "major_reference_materials_cn",
        "heading": "主要参考资料",
        "style": None,
        "entries": (
            "张三，李四. 时序约束下的稳定性研究[J]. 控制学报，2021，12(3)：44-52.",
            "王五. 异构控制系统的边界分析[M]. 北京：科学出版社，2022.",
        ),
    },
    {
        "name": "semantic_heading",
        "heading": "Selected Reference Materials for This Study",
        "style": "Heading 1",
        "ack_after_entries": True,
        "entries": (
            "Miller, T. (2020). Constraint-aware Scheduling. Control Engineering Journal, 8(2), 11-19.",
            "Wilson, R. (2023). Stable Evaluation Pipelines. Engineering Press.",
        ),
    },
    {
        "name": "entry_run_without_marker",
        "heading": None,
        "style": None,
        "entries": (
            "[1] Taylor K. Robust sequence estimation[J]. Automation Journal, 2020, 6(1): 1-12.",
            "[2] Clark M. Constraint propagation in practice[M]. Technical Press, 2021.",
        ),
    },
    {
        "name": "author_year_run_without_marker",
        "heading": None,
        "style": None,
        "entries": (
            "Taylor, K. (2020). Robust Sequence Estimation. Automation Journal, 6(1), 1-12.",
            "Clark, M. (2021). Constraint Propagation in Practice. Technical Press.",
        ),
    },
)


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _set_numbering(paragraph, *, level: int = 0, num_id: int = 1) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def _set_style_outline_level(style, *, level: int) -> None:
    p_pr = style._element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        style._element.append(p_pr)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    p_pr.append(outline)


def _add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在 Word 中更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def _add_math(paragraph) -> None:
    math_para = OxmlElement("m:oMathPara")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "E=mc²"
    math_run.append(math_text)
    math.append(math_run)
    math_para.append(math)
    paragraph._p.append(math_para)


def _build_fixture(path: Path) -> dict[str, str]:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.91)
    section.bottom_margin = Cm(2.13)
    section.left_margin = Cm(2.77)
    section.right_margin = Cm(2.39)
    section.header.paragraphs[0].text = "实验室内部论文页眉"
    section.footer.paragraphs[0].text = "固定页脚 · 不得改写"

    if "标题 1" not in [style.name for style in document.styles]:
        localized_heading = document.styles.add_style("标题 1", WD_STYLE_TYPE.PARAGRAPH)
        localized_heading.base_style = document.styles["Heading 1"]
    outline_style = document.styles.add_style("Research Landmark", WD_STYLE_TYPE.PARAGRAPH)
    outline_style.base_style = document.styles["Normal"]
    _set_style_outline_level(outline_style, level=1)

    title_text = "异构时序约束下的局部稳定性研究"
    title = document.add_paragraph(title_text, style="Title")
    title.runs[0].font.name = "Arial"
    title.runs[0].font.size = Pt(17)
    document.add_paragraph("作者：某某 · 学号：20260001")
    document.add_paragraph("目录", style="Heading 1")
    toc = document.add_paragraph()
    _add_toc_field(toc)

    document.add_paragraph("摘要", style="Heading 1")
    abstract_body = "本文讨论异构时序约束的局部稳定性，并说明实验边界与评价方法。"
    document.add_paragraph(abstract_body)

    first_heading = "第一章 任意命名的问题边界"
    document.add_paragraph(first_heading, style="标题 1")
    body_one = "控制器读取上一时刻的状态，再依据约束矩阵筛除不可行路径。该过程不改变原始观测值。"
    first_body_paragraph = document.add_paragraph(body_one)
    first_body_paragraph.paragraph_format.first_line_indent = Cm(0.81)
    first_body_paragraph.paragraph_format.line_spacing = Pt(18)
    first_body_paragraph.runs[0].font.name = "FangSong"
    first_body_paragraph.runs[0].font.size = Pt(13)

    numbered_heading_text = "误差传播与收敛边界"
    numbered_heading = document.add_paragraph(numbered_heading_text, style="Heading 2")
    _set_numbering(numbered_heading, level=1)
    body_two = "误差只在相邻时间窗之间传递。窗口结束后，缓存会记录本轮校验状态，失败项等待人工确认。"
    document.add_paragraph(body_two)

    outline_heading_text = "Cross-domain Evaluation Boundary Under Heterogeneous Temporal Constraints"
    document.add_paragraph(outline_heading_text, style="Research Landmark")
    english_caption_text = "Fig. 1 Experimental pipeline across three evaluation stages"
    document.add_paragraph(english_caption_text, style="Caption")
    plain_formula_text = "y = ax + b"
    document.add_paragraph(plain_formula_text)

    caption_text = "图1-1 约束传播流程"
    document.add_paragraph(caption_text, style="Caption")
    formula = document.add_paragraph("公式（1）")
    _add_math(formula)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "参数"
    table.cell(0, 1).text = "取值"
    table.cell(1, 0).text = "窗口"
    table.cell(1, 1).text = "12"

    document.add_paragraph("致谢", style="Heading 1")
    acknowledgement = "感谢指导教师对实验边界和术语使用提出的修改意见。"
    document.add_paragraph(acknowledgement)
    works_cited_heading = "Works Cited"
    document.add_paragraph(works_cited_heading)
    reference_text = "[1] Zhang A. Temporal constraints in control systems[J]. 2024."
    document.add_paragraph(reference_text)
    document.save(str(path))
    return {
        "title": title_text,
        "firstHeading": first_heading,
        "numberedHeading": numbered_heading_text,
        "outlineHeading": outline_heading_text,
        "abstractBody": abstract_body,
        "bodyOne": body_one,
        "bodyTwo": body_two,
        "caption": caption_text,
        "englishCaption": english_caption_text,
        "plainFormula": plain_formula_text,
        "worksCitedHeading": works_cited_heading,
        "reference": reference_text,
        "header": "实验室内部论文页眉",
        "footer": "固定页脚 · 不得改写",
    }


def _build_reference_variant_fixture(
    path: Path,
    *,
    heading: str | None,
    heading_style: str | None,
    entries: tuple[str, ...],
    acknowledgement_after_entries: bool = False,
) -> str:
    document = Document()
    document.add_paragraph("参考文献范围保护回归", style="Title")
    document.add_paragraph("摘 要", style="Heading 1")
    editable_body = "本文正文用于验证参考文献变体保护；只有这一段内容允许进入模型，参考资料标题和条目均不得出现。"
    document.add_paragraph(editable_body)
    if heading is not None:
        document.add_paragraph(heading, style=heading_style) if heading_style else document.add_paragraph(heading)
    for entry in entries:
        document.add_paragraph(entry)
    if acknowledgement_after_entries:
        document.add_paragraph("致 谢", style="Heading 1")
        document.add_paragraph("感谢指导教师对论文结构提出的修改意见。")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return editable_body


def _assert_reference_variants_never_reach_model(source_dir: Path) -> None:
    for case in REFERENCE_VARIANT_CASES:
        case_name = str(case["name"])
        heading = case.get("heading")
        heading_text = str(heading) if isinstance(heading, str) else None
        heading_style = case.get("style")
        acknowledgement_after_entries = bool(case.get("ack_after_entries", False))
        entries = tuple(str(value) for value in case["entries"])
        source_path = source_dir / f"reference-variant-{case_name}.docx"
        editable_body = _build_reference_variant_fixture(
            source_path,
            heading=heading_text,
            heading_style=str(heading_style) if isinstance(heading_style, str) else None,
            entries=entries,
            acknowledgement_after_entries=acknowledgement_after_entries,
        )
        model_inputs: list[str] = []

        def identity_transform(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
            model_inputs.append(str(chunk_text))
            return str(chunk_text)

        round_result = run_document_round(
            source_path,
            identity_transform,
            round_number=1,
            prompt_profile="cn",
        )
        extracted_text = get_docx_extracted_text_path(source_path).read_text(encoding="utf-8")
        serialized_model_input = Path(str(round_result["round_context"]["input_text_path"])).read_text(encoding="utf-8")
        actual_model_input = "\n\n".join(model_inputs)
        _assert(editable_body in extracted_text, f"{case_name}: editable body missing from extracted text")
        _assert(editable_body in serialized_model_input, f"{case_name}: editable body missing from serialized model input")
        _assert(editable_body in actual_model_input, f"{case_name}: editable body missing from actual model call")
        protected_values = tuple(value for value in (heading_text, *entries) if value is not None)
        for protected_value in protected_values:
            _assert(protected_value not in extracted_text, f"{case_name}: reference content leaked into extracted text")
            _assert(protected_value not in serialized_model_input, f"{case_name}: reference content leaked into model input file")
            _assert(protected_value not in actual_model_input, f"{case_name}: reference content reached the transform/model")
        _, _, case_snapshot = ensure_docx_processing_assets(source_path)
        protected_units_by_text = {unit.text: unit for unit in case_snapshot.units if unit.text in protected_values}
        for protected_value in protected_values:
            protected_unit = protected_units_by_text.get(protected_value)
            _assert(protected_unit is not None, f"{case_name}: protected reference unit missing from snapshot")
            _assert(not protected_unit.editable, f"{case_name}: protected reference unit remained editable")
        if case_name == "semantic_heading":
            for protected_value in protected_values:
                _assert(
                    protected_units_by_text[protected_value].protect_reason == "references",
                    f"semantic reference heading failed to switch references phase for {protected_value}",
                )
        _assert(
            bool(round_result.get("edit_contract", {}).get("ready")),
            f"{case_name}: protected reference fixture should keep the document contract ready",
        )


def _read_all_visible_text(path: Path) -> str:
    document = Document(str(path))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)


def main() -> int:
    checks: list[str] = []
    origin_root = ROOT_DIR / "origin"
    finish_root = ROOT_DIR / "finish" / "regression"
    origin_root.mkdir(parents=True, exist_ok=True)
    finish_root.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="content-contract-", dir=origin_root) as source_temp, tempfile.TemporaryDirectory(
        prefix="content-contract-", dir=finish_root
    ) as output_temp:
        source_path = Path(source_temp) / "scope-fixture.docx"
        fixture = _build_fixture(source_path)
        with zipfile.ZipFile(source_path, "a") as archive:
            archive.comment = b"fyadr-v10-a"
        extracted_path = Path(output_temp) / "extracted.txt"
        snapshot_path = Path(output_temp) / "snapshot.json"
        _, _, snapshot = ensure_docx_processing_assets(
            source_path,
            extracted_path=extracted_path,
            snapshot_path=snapshot_path,
            scope_diagnostics_path=Path(output_temp) / "scope.json",
        )
        contract = build_document_edit_contract(
            source_path,
            snapshot_path=snapshot_path,
            extracted_text_path=extracted_path,
            stage="regression",
        )
        _assert(contract["ready"] is True, f"fixture contract should be ready: {contract['issues']}")
        _assert(contract["editableHeadingCount"] == 0, "no title or heading may be editable")
        _assert(contract["protectedHeadingCount"] >= 6, "Title/Heading/localized/numbered/outline headings should all be protected")
        _assert(contract["extractedTextMatchesEditableUnits"] is True, "extracted model input must exactly equal editable units")
        checks.append("arbitrary, localized and numbered headings are frozen before model input")

        frozen_stat = source_path.stat()
        frozen_hash = snapshot.source_sha256
        with zipfile.ZipFile(source_path, "a") as archive:
            archive.comment = b"fyadr-v10-b"
        _assert(source_path.stat().st_size == frozen_stat.st_size, "fixture replacement must preserve source size")
        os.utime(source_path, ns=(frozen_stat.st_atime_ns, frozen_stat.st_mtime_ns))
        _assert(source_path.stat().st_mtime_ns == frozen_stat.st_mtime_ns, "fixture replacement must preserve source mtime")
        _assert(_is_snapshot_current(snapshot, source_path) is False, "snapshot identity must reject changed bytes with identical stat metadata")
        _, _, refreshed_snapshot = ensure_docx_processing_assets(
            source_path,
            extracted_path=extracted_path,
            snapshot_path=snapshot_path,
            scope_diagnostics_path=Path(output_temp) / "scope.json",
        )
        _assert(refreshed_snapshot.source_sha256 != frozen_hash, "same-stat replacement must rebuild the cached snapshot")
        _assert(_is_snapshot_current(refreshed_snapshot, source_path) is True, "rebuilt snapshot must match the replacement source")
        snapshot = refreshed_snapshot
        checks.append("snapshot cache invalidates changed bytes even when source size and mtime are preserved")

        foreign_path = Path(source_temp) / "foreign-scope-fixture.docx"
        _build_fixture(foreign_path)
        foreign_document = Document(str(foreign_path))
        foreign_body = next(
            paragraph
            for paragraph in foreign_document.paragraphs
            if paragraph.text == fixture["bodyOne"]
        )
        foreign_body.text = "这是另一份论文的正文，绝不能借用相同段落数量进入当前文档模型输入。"
        foreign_document.save(str(foreign_path))
        foreign_body_map = build_docx_body_map(foreign_path, round_number=1)
        foreign_contract = build_document_edit_contract(
            source_path,
            snapshot_path=snapshot_path,
            extracted_text_path=extracted_path,
            body_map=foreign_body_map,
            candidate_texts=foreign_body_map.current_texts(),
            stage="foreign_body_map",
        )
        foreign_issue_codes = {
            str(item.get("code", ""))
            for item in foreign_contract.get("issues", [])
            if isinstance(item, dict)
        }
        _assert(foreign_contract["ready"] is False, "a body map from another DOCX must fail closed")
        _assert(
            "body_map_source_path_mismatch" in foreign_issue_codes,
            f"foreign body map must expose source identity failure: {sorted(foreign_issue_codes)}",
        )
        _assert(
            foreign_contract["modelInputMatchesEditableUnits"] is False,
            "foreign body-map text must never be declared equal to the frozen model scope",
        )
        checks.append("foreign or stale body maps cannot substitute another paper's text into model input")

        extracted_text = extracted_path.read_text(encoding="utf-8")
        for protected_value in (
            fixture["title"],
            fixture["firstHeading"],
            fixture["numberedHeading"],
            fixture["outlineHeading"],
            fixture["caption"],
            fixture["englishCaption"],
            fixture["plainFormula"],
            fixture["worksCitedHeading"],
            fixture["reference"],
            fixture["header"],
            fixture["footer"],
            "参数",
            "E=mc²",
        ):
            _assert(protected_value not in extracted_text, f"protected content leaked into model input: {protected_value}")
        for editable_value in (fixture["abstractBody"], fixture["bodyOne"], fixture["bodyTwo"]):
            _assert(editable_value in extracted_text, f"editable body missing from model input: {editable_value}")
        checks.append("TOC, captions, formulas, tables, references and headers/footers never enter model input")

        _assert_reference_variants_never_reach_model(Path(source_temp))
        checks.append(
            "reference-list/source-consulted/主要参考资料 headings and consecutive entry runs stay outside model input"
        )

        # Prove the gate runs before any model call, not only at export time.
        _, default_snapshot_path, _ = ensure_docx_processing_assets(source_path)
        tampered_snapshot = _load_docx_snapshot(default_snapshot_path)
        _assert(tampered_snapshot is not None, "fixture snapshot could not be loaded for tamper test")
        tampered_units = tampered_snapshot.units
        tampered_heading = next(
            (
                unit
                for unit in tampered_units
                if not unit.editable
                and str(unit.style_name).lower().startswith("title")
            ),
            None,
        )
        tampered_abstract = next(
            (
                unit
                for unit in tampered_units
                if str(unit.text).replace(" ", "") == "摘要"
            ),
            None,
        )
        _assert(tampered_heading is not None, "fixture should expose a protected title")
        _assert(tampered_abstract is not None, "fixture should expose an abstract heading")

        # Forge a cache that keeps the real source path/stat/hash but disguises
        # the actual Title target as an auto-numbered body paragraph.  Hiding the
        # abstract marker makes this forged target the computed scope start.  A
        # cache-only gate would now report zero editable headings and permit a
        # later export to overwrite the real title.
        tampered_heading.text = "1 本文首先说明研究对象与方法。"
        tampered_heading.style_name = "Normal"
        tampered_heading.has_numbering = True
        tampered_heading.numbering_level = 0
        tampered_heading.outline_level = None
        _assign_unit_edit_decision(
            tampered_heading,
            structural_role="body_list",
            edit_eligibility="eligible",
            reason_codes=("forged_numbered_sentence_prose",),
            protect_reason=None,
        )
        tampered_abstract.text = "结构分隔符"
        tampered_abstract.style_name = "Normal"
        tampered_abstract.has_numbering = False
        tampered_abstract.numbering_level = None
        tampered_abstract.outline_level = None
        _assign_unit_edit_decision(
            tampered_abstract,
            structural_role="heading",
            edit_eligibility="protected",
            reason_codes=("forged_structural_heading",),
            protect_reason="heading",
        )
        tampered_snapshot.editable_unit_count = len(tampered_snapshot.editable_units())
        default_snapshot_path.write_text(
            json.dumps(tampered_snapshot.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        forged_contract = build_document_edit_contract(
            source_path,
            snapshot_path=default_snapshot_path,
            extracted_text_path=get_docx_extracted_text_path(source_path),
            stage="forged_snapshot_cache",
        )
        forged_issue_codes = {
            str(item.get("code", ""))
            for item in forged_contract.get("issues", [])
            if isinstance(item, dict) and item.get("severity") == "error"
        }
        _assert(forged_contract["ready"] is False, "a forged snapshot cache must fail closed")
        _assert(
            "snapshot_authority_mismatch" in forged_issue_codes,
            f"forged cache must expose source-derived authority mismatch: {sorted(forged_issue_codes)}",
        )
        _assert(
            forged_contract.get("snapshotAuthorityVerified") is False,
            "forged cache must never report source-derived authority verification",
        )
        model_call_count = 0

        def should_never_run(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
            nonlocal model_call_count
            model_call_count += 1
            return chunk_text

        try:
            run_document_round(source_path, should_never_run, round_number=1, prompt_profile="cn")
            raise AssertionError("forged snapshot cache should block before the round starts")
        except ValueError as exc:
            _assert("只改正文/格式固定契约未通过" in str(exc), f"unexpected pre-run block: {exc}")
        _assert(model_call_count == 0, "a failed scope contract must block before the first model request")
        checks.append("source-derived scope verification blocks a forged snapshot cache before any model request")
        default_snapshot_path.unlink(missing_ok=True)
        get_docx_extracted_text_path(source_path).unlink(missing_ok=True)

        body_map = build_docx_body_map(source_path, snapshot_path=snapshot_path, round_number=1)
        rewritten = [f"{text.rstrip('。！？.!?')}，已完成局部表达调整。" for text in body_map.current_texts()]
        updated_body_map = update_docx_body_map_texts(body_map, rewritten, round_number=1)
        round_output = Path(output_temp) / "round1.txt"
        round_output.write_text("\n\n".join(rewritten), encoding="utf-8")
        export_path = Path(output_temp) / "export.docx"

        originals = {
            "_find_origin_docx_for_output": app_service._find_origin_docx_for_output,
            "_load_body_map_for_output": app_service._load_body_map_for_output,
            "_load_compare_payload_for_output": app_service._load_compare_payload_for_output,
            "_find_validation_path_for_output": app_service._find_validation_path_for_output,
        }
        try:
            app_service._find_origin_docx_for_output = lambda _path: (source_path, snapshot_path)
            app_service._load_body_map_for_output = lambda _path: updated_body_map
            app_service._load_compare_payload_for_output = lambda _path: None
            app_service._find_validation_path_for_output = lambda _path: None
            export_result = app_service._export_docx_round(
                round_output,
                export_path,
                {},
                format_mode="school_rules",
            )
        finally:
            for name, value in originals.items():
                setattr(app_service, name, value)

        _assert(export_result["formatMode"] == "preserve_original", "legacy school_rules request must be migrated to preserve_original")
        _assert(export_result["contentContractReady"] is True, "post-export content contract must pass")
        _assert(export_result["editableHeadingCount"] == 0, "export evidence must report zero editable headings")
        _assert(export_result["modelInputMatchesEditableUnits"] is True, "export evidence must retain exact scope identity")
        _assert(app_config._normalize_format_mode("school_rules") == "preserve_original", "saved legacy mode must normalize to fidelity lock")
        _assert(app_service._resolve_docx_format_mode() == "preserve_original", "service mode resolver must be fail-closed")
        checks.append("legacy school_rules configuration and explicit overrides cannot mutate product exports")

        visible = _read_all_visible_text(export_path)
        for protected_value in (
            fixture["title"],
            fixture["firstHeading"],
            fixture["numberedHeading"],
            fixture["outlineHeading"],
            fixture["caption"],
            fixture["englishCaption"],
            fixture["plainFormula"],
            fixture["worksCitedHeading"],
            fixture["reference"],
            fixture["header"],
            fixture["footer"],
        ):
            _assert(protected_value in visible, f"protected text changed or disappeared after export: {protected_value}")
        for rewritten_text in rewritten:
            _assert(rewritten_text in visible, "rewritten body text was not placed back into its frozen target")

        protection_audit = audit_docx_export(export_path, source_path=source_path, snapshot_path=snapshot_path)
        ooxml_audit = audit_docx_ooxml_integrity(export_path, source_path=source_path, snapshot_path=snapshot_path)
        format_audit = audit_docx_format_lock(export_path, source_path=source_path, snapshot_path=snapshot_path)
        _assert(protection_audit["ok"] is True, f"protected text audit failed: {protection_audit['issues']}")
        _assert(ooxml_audit["ok"] is True, f"OOXML audit failed: {ooxml_audit['issues']}")
        _assert(format_audit["ok"] is True, f"format-lock audit failed: {format_audit['issues']}")
        checks.append("post-export protected text, OOXML structure and every paragraph format signature remain identical")

        malformed = json.loads(json.dumps(updated_body_map.to_dict(), ensure_ascii=False))
        malformed["units"][0]["target"] = {"kind": "paragraph", "paragraph_index": 0}
        malformed_path = Path(output_temp) / "malformed_body_map.json"
        malformed_path.write_text(json.dumps(malformed, ensure_ascii=False), encoding="utf-8")
        from docx_bodymap import load_docx_body_map

        malformed_body_map = load_docx_body_map(malformed_path)
        _assert(malformed_body_map is not None, "malformed test body map should remain parseable")
        blocked_contract = build_document_edit_contract(
            source_path,
            snapshot_path=snapshot_path,
            extracted_text_path=extracted_path,
            body_map=malformed_body_map,
            candidate_texts=malformed_body_map.current_texts(),
            stage="malicious_scope",
        )
        _assert(blocked_contract["ready"] is False, "a body map retargeted to the title must fail closed")
        _assert(blocked_contract["issueCount"] > 0, "blocked scope must expose actionable issues")
        checks.append("tampered body maps cannot retarget rewritten text into a title")

    report = {"ok": True, "checks": checks}
    report_path = finish_root / "document_edit_contract_regression_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
