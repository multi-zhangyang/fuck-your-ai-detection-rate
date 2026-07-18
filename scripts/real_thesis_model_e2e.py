#!/usr/bin/env python3
"""Auditable real-thesis DOCX end-to-end runner.

The default mode derives and processes the *complete* DOCX with an explicit
identity baseline, then spends a bounded real-provider budget on representative
thesis chunks.  This keeps a normal verification run small enough to audit
while still exercising the production targeted-rerun, RateAudit, review-CAS,
review materialisation, and certified DOCX-export paths.

Provider identity and credentials are read exclusively from these variables:

* ``FYADR_RUN_REAL_LLM``
* ``FYADR_BASE_URL``
* ``FYADR_API_KEY``
* ``FYADR_MODEL``
* ``FYADR_API_TYPE``

The report never contains prompts, model output, the provider URL/model in
clear text, or any credential/fingerprint derived from the API key.
"""

from __future__ import annotations

import argparse
from collections import Counter
from contextlib import contextmanager
from dataclasses import dataclass
from datetime import datetime, timezone
import errno
import hashlib
import json
import os
from pathlib import Path
import re
import shutil
import sys
import threading
import time
from typing import Any, Callable, Iterator
from urllib.parse import urlsplit
import uuid
import zipfile

try:  # POSIX advisory locking.
    import fcntl  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - exercised by the Windows CI runner.
    fcntl = None  # type: ignore[assignment]

try:  # Windows advisory locking.
    import msvcrt  # type: ignore[import-not-found]
except ImportError:  # pragma: no cover - exercised by POSIX runners.
    msvcrt = None  # type: ignore[assignment]

from docx import Document  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import app_service  # noqa: E402
from academic_readability import assess_academic_readability_delta  # noqa: E402
from document_edit_contract import (  # noqa: E402
    assert_document_edit_contract_ready,
    build_document_edit_contract,
)
from docx_pipeline import (  # noqa: E402
    build_docx_scope_diagnostics,
    ensure_docx_processing_assets,
    verify_docx_snapshot_derivation,
)
from llm_client import LLMRequestError  # noqa: E402
from path_utils import is_path_under  # noqa: E402
from fyadr_round_service import (  # noqa: E402
    CANDIDATE_SELECTION_SCHEMA,
    CANDIDATE_SELECTION_VERSION,
    FAILED_ATTEMPT_EVIDENCE_SCHEMA,
    FAILED_ATTEMPT_EVIDENCE_VERSION,
    FAILED_ATTEMPT_GUARD_CATEGORIES,
    FAILED_ATTEMPT_ISSUE_CODES,
    FAILED_ATTEMPT_PUBLIC_FORBIDDEN_KEYS,
)
from prompt_library import (  # noqa: E402
    ROUND_PERTURBATION_DIMENSIONS,
    get_rate_audit_dimension_definition,
)
from source_relative_style_delta import (  # noqa: E402
    source_relative_document_delta_passed,
    source_relative_style_delta_passed,
)


REAL_RUN_ENV = "FYADR_RUN_REAL_LLM"
BASE_URL_ENV = "FYADR_BASE_URL"
API_KEY_ENV = "FYADR_API_KEY"
MODEL_ENV = "FYADR_MODEL"
API_TYPE_ENV = "FYADR_API_TYPE"

ORIGIN_ROOT = (ROOT_DIR / "origin").resolve()
FINISH_ROOT = (ROOT_DIR / "finish").resolve()
RUN_SOURCE_ROOT = ORIGIN_ROOT / "real_thesis_e2e"
REAL_E2E_LOCK_PATH = FINISH_ROOT / ".real_thesis_model_e2e.lock"
REAL_E2E_LOCK_CONFLICT_CODE = "real_provider_e2e_already_running"
REPORT_VERSION = 1
DEFAULT_MAX_REAL_TARGETS = 3
MAX_REAL_TARGETS = 8
# A real-provider E2E must have a network-level budget, not merely a target
# count.  Production rewriting intentionally has a generous transient retry
# floor, but inheriting that policy here could turn three representative
# targets into dozens of HTTP requests.  The E2E transport therefore performs
# one HTTP attempt per completion and caps each request independently.  Model
# output validation may still request a fresh completion within the explicit
# completion-call budget below.
E2E_HTTP_RETRIES_PER_COMPLETION = 0
E2E_REQUEST_TIMEOUT_SECONDS = 180
E2E_MAX_COMPLETIONS_PER_TARGET = 4
RUN_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,79}$")
SHA256_RE = re.compile(r"^[0-9a-f]{64}$")
SECRET_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("openai_compatible_key", re.compile(r"\bsk-(?:proj-)?[A-Za-z0-9_-]{16,}\b")),
    ("bearer_token", re.compile(r"(?i)\bbearer\s+[A-Za-z0-9._~-]{20,}")),
    ("github_token", re.compile(r"\bgh[pousr]_[A-Za-z0-9_]{30,}\b")),
    ("jwt", re.compile(r"\beyJ[A-Za-z0-9_-]{20,}\.[A-Za-z0-9_-]{20,}\.[A-Za-z0-9_-]{10,}\b")),
)
PROVIDER_CONFIGURATION_ERROR_CATEGORIES = {"auth", "client_configuration"}
PROVIDER_CONFIGURATION_STATUS_CODES = {401, 403}
TECH_TERM_RE = re.compile(
    r"\b(?:AI|API|CNN|RNN|LSTM|Transformer|ResNet|YOLO|ROS|PID|IoU|mAP|GPU|CPU|"
    r"Python|Java|MATLAB|Simulink|OpenCV|PyTorch|TensorFlow|SQL|HTTP|TCP|UDP)\b",
    re.IGNORECASE,
)
CITATION_RE = re.compile(r"\[\s*\d+(?:\s*[-,，]\s*\d+)*\s*\]")
UNIT_RE = re.compile(r"(?i)(?:\d(?:\.\d+)?\s*(?:%|ms|s|kg|g|mm|cm|m|hz|mhz|ghz|v|a|w)\b)")


class E2EContractError(RuntimeError):
    """A report-safe, code-only E2E assertion failure."""

    def __init__(self, code: str, category: str = "product_pipeline_failure") -> None:
        super().__init__(code)
        self.code = str(code)
        self.category = str(category)


@dataclass(frozen=True)
class ProviderConfig:
    base_url: str
    api_key: str
    model: str
    api_type: str

    def neutral_app_model_config(self) -> dict[str, Any]:
        # Core round/rerun code may persist route metadata. Give it only an
        # inert identity; the in-memory transport adapter below swaps in the
        # environment provider at the final llm_completion boundary.
        return {
            "baseUrl": "https://identity.invalid/v1",
            "apiKey": "nonsecret-runtime-placeholder",
            "model": "identity-baseline",
            "apiType": "chat_completions",
            "temperature": 0.3,
            "requestTimeoutSeconds": 600,
            "maxRetries": 0,
            "streaming": True,
            "rewriteConcurrency": 1,
            "promptProfile": "cn_custom",
            # A one-round route makes "round 1" unambiguous: this is the real
            # sentence/rhythm prompt, not a hidden provider-specific workflow.
            "promptSequence": ["round1"],
        }

    def public_summary(self) -> dict[str, Any]:
        return {
            "configurationSource": "environment_only",
            "apiType": self.api_type,
            "baseUrlSha256": _sha256_text(self.base_url),
            "modelSha256": _sha256_text(self.model),
            "credentialPresent": bool(self.api_key),
            "credentialFingerprintStored": False,
        }


class EnvironmentProviderTransport:
    """Inject real environment parameters only at the HTTP call boundary."""

    def __init__(self, delegate: Callable[..., Any], provider: ProviderConfig) -> None:
        self.delegate = delegate
        self.provider = provider

    def __call__(self, *args: Any, **kwargs: Any) -> Any:
        forwarded = dict(kwargs)
        forwarded.update(
            {
                "model": self.provider.model,
                "api_key": self.provider.api_key,
                "base_url": self.provider.base_url,
                "api_type": self.provider.api_type,
                # Keep the real E2E request budget deterministic even though
                # the production rewrite path applies a larger retry floor.
                "max_retries": E2E_HTTP_RETRIES_PER_COMPLETION,
                "timeout": E2E_REQUEST_TIMEOUT_SECONDS,
            }
        )
        return self.delegate(*args, **forwarded)


class CompletionCallAuditor:
    """Wrap ``app_service.llm_completion`` without retaining prompt/output."""

    def __init__(self, delegate: Callable[..., Any], *, max_calls: int | None = None) -> None:
        self.delegate = delegate
        self.max_calls = int(max_calls) if max_calls is not None else None
        self._lock = threading.Lock()
        self._issued_calls = 0
        self.calls: list[dict[str, Any]] = []

    def __call__(self, *args: Any, **kwargs: Any) -> Any:
        with self._lock:
            if self.max_calls is not None and self._issued_calls >= self.max_calls:
                raise E2EContractError("real_provider_completion_budget_exhausted", "budget_exhausted")
            self._issued_calls += 1
        forwarded = dict(kwargs)
        stream_requested = forwarded.get("stream") is True
        upstream_stream_callback = forwarded.get("stream_callback")
        stream_state_lock = threading.Lock()
        stream_state = {"eventCount": 0, "done": False}

        if stream_requested:
            def audit_stream_callback(event: dict[str, object]) -> None:
                event_count = 0
                done = False
                if isinstance(event, dict):
                    try:
                        event_count = max(0, int(event.get("eventCount", 0) or 0))
                    except (TypeError, ValueError):
                        event_count = 0
                    done = event.get("done") is True
                with stream_state_lock:
                    stream_state["eventCount"] = max(stream_state["eventCount"], event_count)
                    stream_state["done"] = bool(stream_state["done"] or done)
                # Preserve the production callback chain exactly.  The auditor
                # keeps no reference to the event and records only the two
                # transport metadata fields above.
                if callable(upstream_stream_callback):
                    upstream_stream_callback(event)

            forwarded["stream_callback"] = audit_stream_callback
        prompt = args[0] if args else kwargs.get("prompt", "")
        prompt_chars = len(prompt) if isinstance(prompt, str) else len(str(prompt or ""))
        started = time.monotonic()
        call: dict[str, Any] = {
            "index": 0,
            "promptChars": prompt_chars,
            "outputChars": 0,
            "durationMs": 0,
            "outcome": "unknown",
            "errorCategory": "",
            "statusCode": None,
            "requestedStreaming": stream_requested,
            "streamEventCount": 0,
            "streamDone": False,
        }
        try:
            output = self.delegate(*args, **forwarded)
        except BaseException as exc:
            descriptor = _provider_error_descriptor(exc)
            call.update(
                {
                    "outcome": "error",
                    "errorCategory": descriptor["category"],
                    "statusCode": descriptor["statusCode"],
                }
            )
            raise
        else:
            call.update(
                {
                    "outcome": "success",
                    "outputChars": len(output) if isinstance(output, str) else len(str(output or "")),
                }
            )
            return output
        finally:
            with stream_state_lock:
                call["streamEventCount"] = int(stream_state["eventCount"])
                call["streamDone"] = bool(stream_state["done"])
            call["durationMs"] = round((time.monotonic() - started) * 1000)
            with self._lock:
                call["index"] = len(self.calls) + 1
                self.calls.append(call)

    def public_summary(self) -> dict[str, Any]:
        with self._lock:
            calls = [dict(item) for item in self.calls]
        non_streaming_call_count = sum(
            1 for item in calls if item.get("requestedStreaming") is not True
        )
        successful_calls = [item for item in calls if item.get("outcome") == "success"]
        successful_stream_completed_count = sum(
            1
            for item in successful_calls
            if item.get("requestedStreaming") is True and item.get("streamDone") is True
        )
        return {
            "callCount": len(calls),
            "maxCallCount": self.max_calls,
            "remainingCallCount": (
                max(0, self.max_calls - len(calls))
                if self.max_calls is not None
                else None
            ),
            "httpRetriesPerCompletion": E2E_HTTP_RETRIES_PER_COMPLETION,
            "requestTimeoutSeconds": E2E_REQUEST_TIMEOUT_SECONDS,
            "successCount": len(successful_calls),
            "errorCount": sum(1 for item in calls if item.get("outcome") == "error"),
            "totalDurationMs": sum(int(item.get("durationMs", 0) or 0) for item in calls),
            "totalPromptChars": sum(int(item.get("promptChars", 0) or 0) for item in calls),
            "totalOutputChars": sum(int(item.get("outputChars", 0) or 0) for item in calls),
            "allRealCallsRequestedStreaming": bool(calls) and non_streaming_call_count == 0,
            # Completion evidence is bound to the same call that returned a
            # successful result.  Counting a done event from an errored call
            # must never compensate for a successful call whose stream did
            # not reach its terminal event.
            "streamCompletedCallCount": successful_stream_completed_count,
            "successfulStreamCompletedCallCount": successful_stream_completed_count,
            "successfulStreamIncompleteCallCount": (
                len(successful_calls) - successful_stream_completed_count
            ),
            "allSuccessfulCallsStreamCompleted": (
                successful_stream_completed_count == len(successful_calls)
            ),
            "streamDoneObservedCallCount": sum(
                1 for item in calls if item.get("streamDone") is True
            ),
            "nonStreamingCallCount": non_streaming_call_count,
            "errorCategories": dict(
                sorted(
                    Counter(
                        str(item.get("errorCategory", "") or "unknown")
                        for item in calls
                        if item.get("outcome") == "error"
                    ).items()
                )
            ),
            "calls": calls,
            "storesPromptText": False,
            "storesOutputText": False,
        }


class ProgressAudit:
    """Count production progress events while dropping previews and messages."""

    def __init__(self) -> None:
        self._lock = threading.Lock()
        self.phase_counts: Counter[str] = Counter()
        self.retry_categories: Counter[str] = Counter()
        self.retry_status_classes: Counter[str] = Counter()

    def __call__(self, event: dict[str, Any]) -> None:
        phase = str(event.get("phase", "unknown") or "unknown")
        with self._lock:
            self.phase_counts[phase] += 1
            if phase == "provider-retry-wait":
                category = str(event.get("errorCategory", "") or "unknown")
                self.retry_categories[category] += 1
                status = event.get("statusCode")
                if isinstance(status, int):
                    self.retry_status_classes[f"{status // 100}xx"] += 1

    def public_summary(self) -> dict[str, Any]:
        with self._lock:
            return {
                "phaseCounts": dict(sorted(self.phase_counts.items())),
                "providerRetryCategories": dict(sorted(self.retry_categories.items())),
                "providerRetryStatusClasses": dict(sorted(self.retry_status_classes.items())),
            }


class IdentityCompletion:
    """Local baseline transform; deliberately never delegates to a provider."""

    def __init__(self) -> None:
        self.call_count = 0
        self.total_prompt_chars = 0
        self.total_output_chars = 0

    def __call__(self, *args: Any, **kwargs: Any) -> str:
        prompt = args[0] if args else kwargs.get("prompt", "")
        prompt_text = str(prompt or "")
        if "[INPUT TEXT]" not in prompt_text:
            raise E2EContractError("identity_prompt_missing_input_marker")
        output = prompt_text.rsplit("[INPUT TEXT]", 1)[1].lstrip("\r\n").strip()
        self.call_count += 1
        self.total_prompt_chars += len(prompt_text)
        self.total_output_chars += len(output)
        return output

    def public_summary(self) -> dict[str, Any]:
        return {
            "mode": "identity",
            "providerCallCount": 0,
            "identityTransformCallCount": self.call_count,
            "totalPromptChars": self.total_prompt_chars,
            "totalOutputChars": self.total_output_chars,
            "claim": "artifact_derivation_only_not_model_rewrite",
        }


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _canonical_json_sha256(value: Any) -> str:
    encoded = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _require(condition: bool, code: str, category: str = "product_pipeline_failure") -> None:
    if not condition:
        raise E2EContractError(code, category)


def _acquire_real_provider_file_lock(descriptor: int) -> None:
    if fcntl is not None:
        fcntl.flock(descriptor, fcntl.LOCK_EX | fcntl.LOCK_NB)
        return
    if msvcrt is not None:
        # ``msvcrt.locking`` cannot reliably lock a byte beyond EOF.  Keep a
        # single NUL sentinel on Windows so an empty, freshly created lock
        # file has a real byte-range to lock.  Writing before locking also
        # turns a write denial from the current owner into the same immediate
        # contention result as ``LK_NBLCK``.
        os.lseek(descriptor, 0, os.SEEK_SET)
        os.write(descriptor, b"\0")
        os.lseek(descriptor, 0, os.SEEK_SET)
        msvcrt.locking(descriptor, msvcrt.LK_NBLCK, 1)
        return
    raise OSError(errno.ENOSYS, "No supported advisory file-lock backend is available.")


def _release_real_provider_file_lock(descriptor: int) -> None:
    if fcntl is not None:
        fcntl.flock(descriptor, fcntl.LOCK_UN)
        return
    if msvcrt is not None:
        os.lseek(descriptor, 0, os.SEEK_SET)
        msvcrt.locking(descriptor, msvcrt.LK_UNLCK, 1)
        return
    raise OSError(errno.ENOSYS, "No supported advisory file-lock backend is available.")


@contextmanager
def _real_provider_e2e_lock(
    lock_path: Path = REAL_E2E_LOCK_PATH,
) -> Iterator[None]:
    """Hold the host-wide advisory lock for one real-provider E2E.

    The file contains no PID, endpoint, model, provider identity, credential,
    or run metadata.  POSIX keeps it empty; Windows may keep one NUL byte as
    the byte-range lock sentinel required by ``msvcrt``.  The OS-native
    advisory lock releases when the descriptor closes, including exception
    exits and process termination.
    """

    path = lock_path.expanduser().resolve()
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        descriptor = os.open(
            path,
            os.O_CREAT
            | os.O_RDWR
            | getattr(os, "O_CLOEXEC", 0)
            | getattr(os, "O_BINARY", 0),
            0o600,
        )
    except OSError as exc:
        raise E2EContractError(
            "real_provider_e2e_lock_unavailable",
            "runtime_environment_failure",
        ) from exc

    acquired = False
    try:
        try:
            _acquire_real_provider_file_lock(descriptor)
            acquired = True
        except OSError as exc:
            if exc.errno in {errno.EACCES, errno.EAGAIN, errno.EDEADLK}:
                raise E2EContractError(
                    REAL_E2E_LOCK_CONFLICT_CODE,
                    "concurrency_conflict",
                ) from exc
            raise E2EContractError(
                "real_provider_e2e_lock_unavailable",
                "runtime_environment_failure",
            ) from exc
        # Remove stale bytes from an older implementation only after this
        # process owns the advisory lock.  Windows must retain its one-byte
        # NUL sentinel; POSIX retains no file content at all.
        try:
            os.ftruncate(descriptor, 1 if msvcrt is not None else 0)
        except OSError as exc:
            raise E2EContractError(
                "real_provider_e2e_lock_unavailable",
                "runtime_environment_failure",
            ) from exc
        yield
    finally:
        if acquired:
            try:
                _release_real_provider_file_lock(descriptor)
            except OSError:
                pass
        os.close(descriptor)


@contextmanager
def _execution_mode_run_lock(
    execution_mode: str,
    *,
    lock_path: Path = REAL_E2E_LOCK_PATH,
) -> Iterator[None]:
    if execution_mode == "offline_fake":
        yield
        return
    if execution_mode != "real_provider":
        raise E2EContractError("execution_mode_invalid", "input_error")
    with _real_provider_e2e_lock(lock_path):
        yield


def _load_provider_from_environment(*, allow_offline_injection: bool) -> ProviderConfig:
    enabled = os.environ.get(REAL_RUN_ENV, "").strip() == "1"
    if not enabled and not allow_offline_injection:
        raise E2EContractError("real_provider_execution_not_enabled", "configuration_error")
    base_url = os.environ.get(BASE_URL_ENV, "").strip()
    api_key = os.environ.get(API_KEY_ENV, "").strip()
    model = os.environ.get(MODEL_ENV, "").strip()
    api_type = os.environ.get(API_TYPE_ENV, "").strip() or "chat_completions"
    if not base_url or not api_key or not model:
        raise E2EContractError("provider_environment_incomplete", "configuration_error")
    parsed = urlsplit(base_url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc or parsed.username or parsed.password:
        raise E2EContractError("provider_base_url_invalid", "configuration_error")
    if api_type not in {"chat_completions", "responses"}:
        raise E2EContractError("provider_api_type_unsupported", "configuration_error")
    return ProviderConfig(base_url=base_url, api_key=api_key, model=model, api_type=api_type)


def _validate_attribution(source_url: str, source_commit: str, source_license: str) -> dict[str, str]:
    url = str(source_url or "").strip()
    commit = str(source_commit or "").strip()
    license_name = str(source_license or "").strip()
    parsed = urlsplit(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc or parsed.username or parsed.password:
        raise E2EContractError("source_attribution_url_invalid", "input_error")
    if not commit or len(commit) > 160:
        raise E2EContractError("source_attribution_commit_invalid", "input_error")
    if not license_name or len(license_name) > 160:
        raise E2EContractError("source_attribution_license_invalid", "input_error")
    return {"url": url, "commit": commit, "license": license_name}


def _resolve_finish_path(value: Path, *, label: str, suffix: str) -> Path:
    path = value.expanduser().resolve()
    if not is_path_under(path, FINISH_ROOT):
        raise E2EContractError(f"{label}_outside_finish", "input_error")
    if path.suffix.lower() != suffix:
        raise E2EContractError(f"{label}_suffix_invalid", "input_error")
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _require_fresh_output_paths(report_path: Path, export_path: Path) -> None:
    """Fail closed instead of replacing evidence from an earlier E2E run."""

    _require(not report_path.exists(), "report_path_already_exists", "input_error")
    _require(not export_path.exists(), "export_path_already_exists", "input_error")


def _copy_unique_work_source(sample_path: Path, run_id: str) -> tuple[Path, str]:
    sample = sample_path.expanduser().resolve()
    if not sample.exists() or not sample.is_file() or sample.suffix.lower() != ".docx":
        raise E2EContractError("sample_docx_invalid", "input_error")
    if not RUN_ID_RE.fullmatch(run_id):
        raise E2EContractError("run_id_invalid", "input_error")
    unique_dir = (RUN_SOURCE_ROOT / run_id / uuid.uuid4().hex).resolve()
    _require(is_path_under(unique_dir, ORIGIN_ROOT), "work_source_path_invalid", "input_error")
    unique_dir.mkdir(parents=True, exist_ok=False)
    work_source = unique_dir / "source.docx"
    sample_sha256 = _sha256_file(sample)
    shutil.copy2(sample, work_source)
    _require(_sha256_file(work_source) == sample_sha256, "work_source_copy_sha256_mismatch", "input_error")
    return work_source, sample_sha256


def _provider_error_descriptor(error: BaseException) -> dict[str, Any]:
    current: BaseException | None = error
    seen: set[int] = set()
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        if isinstance(current, LLMRequestError):
            return {
                "category": str(current.category or "provider_error"),
                "statusCode": current.status_code if isinstance(current.status_code, int) else None,
                "retryable": bool(current.retryable),
            }
        if isinstance(current, TimeoutError):
            return {"category": "timeout", "statusCode": None, "retryable": True}
        if isinstance(current, (ConnectionError, OSError)):
            return {"category": "network", "statusCode": None, "retryable": True}
        current = current.__cause__ or current.__context__
    return {"category": "non_provider", "statusCode": None, "retryable": False}


def _is_external_provider_failure(error: BaseException) -> bool:
    descriptor = _provider_error_descriptor(error)
    category = str(descriptor.get("category", ""))
    status = descriptor.get("statusCode")
    return bool(
        category in {"network", "timeout", "rate_limit", "server"}
        or status in {408, 409, 425, 429}
        or isinstance(status, int) and status >= 500
    )


def _provider_configuration_failure_summary(
    descriptor: dict[str, Any],
) -> dict[str, Any] | None:
    """Return a report-safe provider/configuration failure descriptor.

    Authentication and client-configuration failures happen before a model
    response exists.  They must therefore never be reported as model-output
    validation failures.  Keep only the canonical category and HTTP status;
    endpoint, credential, provider body, and exception text are deliberately
    excluded.
    """

    category = str(
        descriptor.get("category", descriptor.get("errorCategory", "")) or "provider_configuration"
    ).strip().lower()
    status_code = descriptor.get("statusCode")
    normalized_status = status_code if isinstance(status_code, int) else None
    if (
        category not in PROVIDER_CONFIGURATION_ERROR_CATEGORIES
        and normalized_status not in PROVIDER_CONFIGURATION_STATUS_CODES
    ):
        return None
    return {
        "status": "provider_configuration_failure",
        "category": category,
        "statusCode": normalized_status,
    }


def _provider_configuration_failure_from_calls(
    calls: list[dict[str, Any]],
) -> dict[str, Any] | None:
    for call in calls:
        summary = _provider_configuration_failure_summary(call)
        if summary is not None:
            return summary
    return None


def _provider_configuration_failure_for_error(
    error: BaseException,
) -> dict[str, Any] | None:
    return _provider_configuration_failure_summary(_provider_error_descriptor(error))


@contextmanager
def _patched_completion(replacement: Callable[..., Any]) -> Iterator[None]:
    original = app_service.llm_completion
    app_service.llm_completion = replacement
    try:
        yield
    finally:
        app_service.llm_completion = original


def _read_json_object(path_value: str | Path) -> dict[str, Any]:
    path = Path(path_value)
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def _fresh_docx_preflight(work_source: Path) -> tuple[dict[str, Any], dict[str, Any], Path, Path, Any]:
    extracted_path, snapshot_path, cached_snapshot = ensure_docx_processing_assets(work_source)
    derivation, authoritative_snapshot = verify_docx_snapshot_derivation(cached_snapshot, work_source)
    _require(bool(derivation.get("ok")), "fresh_snapshot_derivation_failed", "document_scope_failure")
    contract = build_document_edit_contract(
        work_source,
        snapshot_path=snapshot_path,
        extracted_text_path=extracted_path,
        stage="real_thesis_e2e_preflight",
    )
    try:
        assert_document_edit_contract_ready(contract, label="real thesis E2E preflight")
    except ValueError as exc:
        raise E2EContractError("document_edit_contract_not_ready", "document_scope_failure") from exc
    _require(bool(contract.get("modelInputMatchesEditableUnits")), "model_input_scope_mismatch", "document_scope_failure")
    _require(
        contract.get("semanticRangeTopologyValid") is True,
        "contract_semantic_range_topology_invalid",
        "document_scope_failure",
    )
    _require(
        int(contract.get("editableSemanticRangeCoveredUnitCount", 0) or 0) == 0,
        "contract_comment_range_interior_entered_model_scope",
        "document_scope_failure",
    )
    _require(
        int(contract.get("editableSemanticRangeAnchorUnitCount", 0) or 0) == 0,
        "contract_semantic_range_anchor_entered_model_scope",
        "document_scope_failure",
    )
    _require(
        int(contract.get("editableSemanticPointReferenceUnitCount", 0) or 0) == 0,
        "contract_semantic_point_reference_entered_model_scope",
        "document_scope_failure",
    )
    expected_text = "\n\n".join(authoritative_snapshot.editable_texts())
    actual_text = extracted_path.read_text(encoding="utf-8")
    _require(actual_text == expected_text, "fresh_extracted_text_mismatch", "document_scope_failure")
    freeze = _build_freeze_summary(work_source, authoritative_snapshot, snapshot_path)
    _require(bool(freeze.get("allObservedSurfacesFrozen")), "protected_surface_entered_model_scope", "document_scope_failure")
    _require(
        int(freeze.get("editableAmbiguousFormatAnchorUnitCount", 0) or 0) == 0,
        "ambiguous_format_anchor_entered_model_scope",
        "document_scope_failure",
    )
    _require(
        freeze.get("semanticRangeTopologyValid") is True,
        "semantic_range_topology_invalid",
        "document_scope_failure",
    )
    _require(
        int(freeze.get("editableSemanticRangeCoveredUnitCount", 0) or 0) == 0,
        "comment_range_interior_entered_model_scope",
        "document_scope_failure",
    )
    _require(
        int(freeze.get("editableSemanticRangeAnchorUnitCount", 0) or 0) == 0,
        "semantic_range_anchor_entered_model_scope",
        "document_scope_failure",
    )
    _require(
        int(freeze.get("editableSemanticPointReferenceUnitCount", 0) or 0) == 0,
        "semantic_point_reference_entered_model_scope",
        "document_scope_failure",
    )
    public_contract = _summarize_contract(contract)
    public_contract.update(
        {
            "extractedTextSha256": _sha256_file(extracted_path),
            "snapshotSha256": _sha256_file(snapshot_path),
            "snapshotDerivationVerified": True,
            "cachedSnapshotDigest": str(derivation.get("cachedDigest", "") or ""),
            "authoritativeSnapshotDigest": str(derivation.get("authoritativeDigest", "") or ""),
        }
    )
    return public_contract, freeze, extracted_path, snapshot_path, authoritative_snapshot


def _build_freeze_summary(source_path: Path, snapshot: Any, snapshot_path: Path) -> dict[str, Any]:
    diagnostics = build_docx_scope_diagnostics(snapshot, snapshot_path=snapshot_path)
    unit_diagnostics = {
        int(item.get("unitIndex", -1)): item
        for item in diagnostics.get("units", [])
        if isinstance(item, dict)
    }
    categories: dict[str, list[Any]] = {
        "titlesAndHeadings": [],
        "toc": [],
        "formula": [],
        "references": [],
    }
    for unit in snapshot.units:
        diagnostic = unit_diagnostics.get(int(unit.unit_index), {})
        flags = diagnostic.get("flags") if isinstance(diagnostic.get("flags"), dict) else {}
        style_name = str(unit.style_name or "").casefold()
        if bool(flags.get("heading")) or "title" in style_name or "标题" in style_name:
            categories["titlesAndHeadings"].append(unit)
        if bool(flags.get("tocHeading")) or bool(flags.get("tocEntry")) or (
            bool(unit.has_field_code) and str(unit.protect_reason or "") in {"front_matter", "generated_field"}
        ):
            categories["toc"].append(unit)
        if bool(flags.get("formula")) or bool(unit.has_math):
            categories["formula"].append(unit)
        if (
            bool(flags.get("referencesHeading"))
            or bool(flags.get("referenceEntry"))
            or str(unit.protect_reason or "") == "references"
        ):
            categories["references"].append(unit)

    document = Document(str(source_path))
    table_paragraph_count = sum(
        1
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if paragraph.text.strip()
    )
    header_footer_paragraph_count = 0
    for section in document.sections:
        header_footer_paragraph_count += sum(1 for p in section.header.paragraphs if p.text.strip())
        header_footer_paragraph_count += sum(1 for p in section.footer.paragraphs if p.text.strip())

    public_categories: dict[str, Any] = {}
    all_frozen = True
    for name, units in categories.items():
        editable_count = sum(1 for unit in units if bool(unit.editable))
        frozen = editable_count == 0
        all_frozen = all_frozen and frozen
        public_categories[name] = {
            "observedCount": len(units),
            "editableCount": editable_count,
            "frozen": frozen,
            "status": "frozen" if units and frozen else "not_observed" if not units else "scope_violation",
        }

    editable_non_body_target_count = sum(
        1 for unit in snapshot.units if unit.editable and str(unit.target.get("kind", "")) != "paragraph"
    )
    table_frozen = editable_non_body_target_count == 0
    public_categories["tables"] = {
        "observedTableCount": len(document.tables),
        "observedTextParagraphCount": table_paragraph_count,
        "editableTargetCount": editable_non_body_target_count,
        "frozen": table_frozen,
        "status": "frozen_by_structural_exclusion" if document.tables else "not_observed",
    }
    public_categories["headersAndFooters"] = {
        "observedTextParagraphCount": header_footer_paragraph_count,
        "editableTargetCount": 0,
        "frozen": True,
        "status": "frozen_by_structural_exclusion" if header_footer_paragraph_count else "not_observed",
    }
    all_frozen = all_frozen and table_frozen
    format_anchor_count = sum(len(getattr(unit, "format_anchors", []) or []) for unit in snapshot.units)
    editable_format_anchor_count = sum(
        len(getattr(unit, "format_anchors", []) or [])
        for unit in snapshot.units
        if bool(unit.editable)
    )
    semantic_range_count = int(getattr(snapshot, "semantic_range_count", 0) or 0)
    bookmark_range_count = int(getattr(snapshot, "bookmark_range_count", 0) or 0)
    comment_range_count = int(getattr(snapshot, "comment_range_count", 0) or 0)
    semantic_range_topology_valid = bool(
        getattr(snapshot, "semantic_range_topology_valid", True)
    )
    semantic_range_anchor_unit_count = sum(
        1 for unit in snapshot.units if bool(getattr(unit, "has_semantic_range_anchor", False))
    )
    editable_semantic_range_anchor_unit_count = sum(
        1
        for unit in snapshot.units
        if bool(unit.editable) and bool(getattr(unit, "has_semantic_range_anchor", False))
    )
    # A bookmark is a positional Word feature: marker-free paragraphs between
    # its endpoints may be rewritten while the endpoint topology stays frozen.
    # Comment ranges are different because their covered content is part of the
    # annotation contract, so their interiors remain protected.
    semantic_range_covered_unit_count = sum(
        1 for unit in snapshot.units if bool(getattr(unit, "inside_comment_range", False))
    )
    editable_semantic_range_covered_unit_count = sum(
        1
        for unit in snapshot.units
        if bool(unit.editable) and bool(getattr(unit, "inside_comment_range", False))
    )
    bookmark_range_interior_unit_count = sum(
        1 for unit in snapshot.units if bool(getattr(unit, "inside_bookmark_range", False))
    )
    editable_bookmark_range_interior_unit_count = sum(
        1
        for unit in snapshot.units
        if bool(unit.editable) and bool(getattr(unit, "inside_bookmark_range", False))
    )
    semantic_point_reference_unit_count = sum(
        1 for unit in snapshot.units if bool(getattr(unit, "has_semantic_point_reference", False))
    )
    editable_semantic_point_reference_unit_count = sum(
        1
        for unit in snapshot.units
        if bool(unit.editable) and bool(getattr(unit, "has_semantic_point_reference", False))
    )
    all_frozen = bool(
        all_frozen
        and semantic_range_topology_valid
        and editable_semantic_range_anchor_unit_count == 0
        and editable_semantic_range_covered_unit_count == 0
        and editable_semantic_point_reference_unit_count == 0
    )
    return {
        "allObservedSurfacesFrozen": all_frozen,
        "totalTextUnitCount": int(snapshot.total_text_unit_count),
        "editableUnitCount": int(snapshot.editable_unit_count),
        "protectedUnitCount": int(snapshot.total_text_unit_count - snapshot.editable_unit_count),
        "scopeDiagnosticsOk": bool(diagnostics.get("ok")),
        "scopeErrorCount": int(diagnostics.get("errorCount", 0) or 0),
        "scopeWarningCount": int(diagnostics.get("warningCount", 0) or 0),
        "formatAnchorCount": format_anchor_count,
        "editableFormatAnchorCount": editable_format_anchor_count,
        "semanticRangeCount": semantic_range_count,
        "bookmarkRangeCount": bookmark_range_count,
        "commentRangeCount": comment_range_count,
        "semanticRangeTopologyValid": semantic_range_topology_valid,
        "semanticRangeAnchorUnitCount": semantic_range_anchor_unit_count,
        "editableSemanticRangeAnchorUnitCount": editable_semantic_range_anchor_unit_count,
        "semanticRangeCoveredUnitCount": semantic_range_covered_unit_count,
        "editableSemanticRangeCoveredUnitCount": editable_semantic_range_covered_unit_count,
        "bookmarkRangeInteriorUnitCount": bookmark_range_interior_unit_count,
        "editableBookmarkRangeInteriorUnitCount": editable_bookmark_range_interior_unit_count,
        "commentRangeInteriorUnitCount": semantic_range_covered_unit_count,
        "editableCommentRangeInteriorUnitCount": editable_semantic_range_covered_unit_count,
        "semanticPointReferenceUnitCount": semantic_point_reference_unit_count,
        "editableSemanticPointReferenceUnitCount": editable_semantic_point_reference_unit_count,
        "kindAwareSemanticRangePolicy": True,
        "ambiguousFormatAnchorUnitCount": sum(
            1 for unit in snapshot.units if bool(getattr(unit, "format_anchor_ambiguous", False))
        ),
        "editableAmbiguousFormatAnchorUnitCount": sum(
            1
            for unit in snapshot.units
            if bool(unit.editable) and bool(getattr(unit, "format_anchor_ambiguous", False))
        ),
        "categories": public_categories,
        "storesProtectedText": False,
    }


def _summarize_contract(contract: dict[str, Any]) -> dict[str, Any]:
    return {
        "ready": bool(contract.get("ready")),
        "policy": str(contract.get("policy", "") or ""),
        "sourceKind": str(contract.get("sourceKind", "") or ""),
        "sourceSha256": str(contract.get("sourceSha256", "") or ""),
        "scopeDigest": str(contract.get("scopeDigest", "") or ""),
        "formatDigest": str(contract.get("formatDigest", "") or ""),
        "snapshotVersion": int(contract.get("snapshotVersion", 0) or 0),
        "snapshotCurrent": bool(contract.get("snapshotCurrent")),
        "snapshotAuthorityVerified": bool(contract.get("snapshotAuthorityVerified")),
        "structuralRolePolicyVersion": int(contract.get("structuralRolePolicyVersion", 0) or 0),
        "structuralInventoryVersion": int(contract.get("structuralInventoryVersion", 0) or 0),
        "scopeReady": bool(contract.get("scopeReady")),
        "formatLockReady": bool(contract.get("formatLockReady")),
        "modelInputMatchesEditableUnits": bool(contract.get("modelInputMatchesEditableUnits")),
        "extractedTextMatchesEditableUnits": bool(contract.get("extractedTextMatchesEditableUnits")),
        "editableUnitCount": int(contract.get("editableUnitCount", 0) or 0),
        "protectedUnitCount": int(contract.get("protectedUnitCount", 0) or 0),
        "editableHeadingCount": int(contract.get("editableHeadingCount", 0) or 0),
        "semanticRangeCount": int(contract.get("semanticRangeCount", 0) or 0),
        "bookmarkRangeCount": int(contract.get("bookmarkRangeCount", 0) or 0),
        "commentRangeCount": int(contract.get("commentRangeCount", 0) or 0),
        "semanticRangeTopologyValid": bool(contract.get("semanticRangeTopologyValid")),
        "semanticRangeAnchorUnitCount": int(
            contract.get("semanticRangeAnchorUnitCount", 0) or 0
        ),
        "editableSemanticRangeAnchorUnitCount": int(
            contract.get("editableSemanticRangeAnchorUnitCount", 0) or 0
        ),
        "semanticRangeCoveredUnitCount": int(
            contract.get("semanticRangeCoveredUnitCount", 0) or 0
        ),
        "editableSemanticRangeCoveredUnitCount": int(
            contract.get("editableSemanticRangeCoveredUnitCount", 0) or 0
        ),
        "bookmarkRangeInteriorUnitCount": int(
            contract.get("bookmarkRangeInteriorUnitCount", 0) or 0
        ),
        "editableBookmarkRangeInteriorUnitCount": int(
            contract.get("editableBookmarkRangeInteriorUnitCount", 0) or 0
        ),
        "semanticPointReferenceUnitCount": int(
            contract.get("semanticPointReferenceUnitCount", 0) or 0
        ),
        "editableSemanticPointReferenceUnitCount": int(
            contract.get("editableSemanticPointReferenceUnitCount", 0) or 0
        ),
        "kindAwareSemanticRangePolicy": True,
        "issueCount": int(contract.get("issueCount", 0) or 0),
        "warningCount": int(contract.get("warningCount", 0) or 0),
    }


def _summarize_round_result(result: dict[str, Any], preflight_contract: dict[str, Any]) -> dict[str, Any]:
    contract = result.get("editContract") if isinstance(result.get("editContract"), dict) else {}
    _require(int(result.get("round", 0) or 0) == 1, "baseline_round_number_mismatch")
    _require(Path(str(result.get("outputPath", ""))).exists(), "baseline_output_missing")
    _require(bool(contract.get("ready")), "round_edit_contract_not_ready", "document_scope_failure")
    _require(bool(contract.get("modelInputMatchesEditableUnits")), "round_model_input_scope_mismatch", "document_scope_failure")
    _require(
        int(result.get("paragraphCount", 0) or 0) == int(preflight_contract.get("editableUnitCount", -1)),
        "round_paragraph_scope_count_mismatch",
        "document_scope_failure",
    )
    for key in ("sourceSha256", "scopeDigest", "formatDigest"):
        _require(
            str(contract.get(key, "") or "") == str(preflight_contract.get(key, "") or ""),
            f"round_contract_{key}_mismatch",
            "document_scope_failure",
        )
    quality = result.get("qualitySummary") if isinstance(result.get("qualitySummary"), dict) else {}
    run_audit = result.get("runAudit") if isinstance(result.get("runAudit"), dict) else {}
    return {
        "round": 1,
        "outputPath": str(Path(str(result.get("outputPath", ""))).resolve()),
        "comparePath": str(Path(str(result.get("comparePath", ""))).resolve()),
        "manifestPath": str(Path(str(result.get("manifestPath", ""))).resolve()),
        "bodyMapPath": str(Path(str(result.get("bodyMapPath", ""))).resolve()),
        "validationPath": str(Path(str(result.get("validationPath", ""))).resolve()),
        "editContractPath": str(Path(str(result.get("editContractPath", ""))).resolve()),
        "chunkCount": int(run_audit.get("chunkCount", result.get("outputSegmentCount", 0)) or 0),
        "paragraphCount": int(result.get("paragraphCount", 0) or 0),
        "inputSegmentCount": int(result.get("inputSegmentCount", 0) or 0),
        "outputSegmentCount": int(result.get("outputSegmentCount", 0) or 0),
        "validationRetryCount": int(quality.get("validationRetryCount", 0) or 0),
        "validationEventCount": int(quality.get("validationEventCount", 0) or 0),
        "sourceFallbackCount": int(quality.get("sourceFallbackCount", 0) or 0),
        "immutableFormatAnchorCount": int(run_audit.get("immutableFormatAnchorCount", 0) or 0),
        "immutableFormatAnchorPlanSha256": str(run_audit.get("immutableFormatAnchorPlanSha256", "") or ""),
        "formatPlaceholderCount": int((quality.get("protectedTokenTypes") or {}).get("FMT", 0) or 0),
        "contract": _summarize_contract(contract),
    }


def _summarize_snapshot(snapshot: dict[str, Any]) -> dict[str, Any]:
    compare = snapshot.get("compare") if isinstance(snapshot.get("compare"), dict) else {}
    review = snapshot.get("review") if isinstance(snapshot.get("review"), dict) else {}
    chunks = [item for item in compare.get("chunks", []) if isinstance(item, dict)]
    selections = [
        item.get("candidateSelection")
        for item in chunks
        if isinstance(item.get("candidateSelection"), dict)
    ]
    v2_selections = [
        item
        for item in selections
        if item.get("schema") == CANDIDATE_SELECTION_SCHEMA
        and item.get("schemaVersion") == CANDIDATE_SELECTION_VERSION
    ]
    source_relative_values: list[object] = []
    for selection in v2_selections:
        source_relative_values.append(selection.get("resultSourceRelativeStyleDelta"))
        raw_candidates = selection.get("candidates")
        if isinstance(raw_candidates, list):
            source_relative_values.extend(
                candidate.get("sourceRelativeStyleDelta")
                for candidate in raw_candidates
                if isinstance(candidate, dict)
            )
    valid_source_relative_count = sum(
        1 for value in source_relative_values if source_relative_style_delta_passed(value)
    )
    raw_profile_registry = compare.get("sourcePatternProfiles")
    profile_registry_count = len(raw_profile_registry) if isinstance(raw_profile_registry, dict) else 0
    document_delta = compare.get("sourceRelativeDocumentDelta")
    return {
        "version": int(snapshot.get("version", 0) or 0),
        "outputPath": str(snapshot.get("outputPath", "") or ""),
        "docId": str(snapshot.get("docId", "") or ""),
        "round": int(snapshot.get("round", 0) or 0),
        "chunkCount": len(chunks),
        "decisionCount": len(review.get("decisions", {}) if isinstance(review.get("decisions"), dict) else {}),
        "compareRevision": str(snapshot.get("compareRevision", "") or ""),
        "reviewRevision": str(snapshot.get("reviewRevision", "") or ""),
        "contentRevision": str(snapshot.get("contentRevision", "") or ""),
        "artifactSnapshotDigest": str(snapshot.get("artifactSnapshotDigest", "") or ""),
        "compareSha256": str(snapshot.get("compareSha256", "") or ""),
        "reviewSha256": snapshot.get("reviewSha256"),
        "effectiveTextSha256": str(snapshot.get("effectiveTextSha256", "") or ""),
        "outputSha256": str(snapshot.get("outputSha256", "") or ""),
        "bodyMapSha256": snapshot.get("bodyMapSha256"),
        "manifestSha256": snapshot.get("manifestSha256"),
        "effectiveTextChars": int((snapshot.get("effectivePreview") or {}).get("totalChars", 0) or 0),
        "rawOutputMatchesEffective": bool(snapshot.get("rawOutputMatchesEffective")),
        "bodyMapMatchesEffective": snapshot.get("bodyMapMatchesEffective"),
        "reviewLinkReady": bool(review.get("reviewLinkReady")),
        "materializationSource": str(snapshot.get("materializationSource", "") or ""),
        "candidateSelectionEvidence": {
            "schema": CANDIDATE_SELECTION_SCHEMA,
            "schemaVersion": CANDIDATE_SELECTION_VERSION,
            "selectionCount": len(selections),
            "v2SelectionCount": len(v2_selections),
            "allSelectionsV2": len(selections) == len(v2_selections),
            "isAiDetectorFalseCount": sum(
                1 for item in v2_selections if item.get("isAiDetector") is False
            ),
            "claimsDetectionRateFalseCount": sum(
                1 for item in v2_selections if item.get("claimsDetectionRate") is False
            ),
            "sourceRelativeEvidenceCount": len(source_relative_values),
            "validSourceRelativeEvidenceCount": valid_source_relative_count,
            "allSourceRelativeEvidenceValid": (
                bool(source_relative_values)
                and valid_source_relative_count == len(source_relative_values)
            ),
            "sourcePatternProfileRegistryCount": profile_registry_count,
            "sourceRelativeDocumentDeltaPresent": isinstance(document_delta, dict),
            "sourceRelativeDocumentDeltaPassed": source_relative_document_delta_passed(document_delta),
            "heuristicOnly": True,
            "isAiDetector": False,
            "claimsDetectionRate": False,
            "storesCandidateText": False,
        },
    }


def _require_v2_snapshot_candidate_evidence(summary: dict[str, Any]) -> None:
    evidence = (
        summary.get("candidateSelectionEvidence")
        if isinstance(summary.get("candidateSelectionEvidence"), dict)
        else {}
    )
    chunk_count = int(summary.get("chunkCount", 0) or 0)
    selection_count = int(evidence.get("selectionCount", 0) or 0)
    v2_count = int(evidence.get("v2SelectionCount", 0) or 0)
    _require(chunk_count > 0, "candidate_evidence_snapshot_empty")
    _require(selection_count == chunk_count, "candidate_selection_evidence_incomplete")
    _require(
        evidence.get("allSelectionsV2") is True and v2_count == selection_count,
        "candidate_selection_v2_evidence_invalid",
    )
    _require(
        int(evidence.get("isAiDetectorFalseCount", 0) or 0) == selection_count
        and int(evidence.get("claimsDetectionRateFalseCount", 0) or 0) == selection_count,
        "candidate_selection_heuristic_claims_invalid",
    )
    _require(
        evidence.get("allSourceRelativeEvidenceValid") is True,
        "candidate_source_relative_style_evidence_invalid",
    )
    _require(
        int(evidence.get("sourcePatternProfileRegistryCount", 0) or 0) > 0,
        "candidate_source_pattern_profile_registry_missing",
    )
    _require(
        evidence.get("sourceRelativeDocumentDeltaPassed") is True,
        "candidate_source_relative_document_delta_invalid",
    )


def _summarize_rate_audit(report: dict[str, Any]) -> dict[str, Any]:
    baseline = report.get("baseline") if isinstance(report.get("baseline"), dict) else {}
    current = report.get("current") if isinstance(report.get("current"), dict) else {}
    delta = report.get("delta") if isinstance(report.get("delta"), dict) else {}
    plan = report.get("strategyPlan") if isinstance(report.get("strategyPlan"), dict) else {}
    binding = report.get("strategyBinding") if isinstance(report.get("strategyBinding"), dict) else {}
    return {
        "label": str(report.get("label", "") or ""),
        "isAiDetector": bool(report.get("isAiDetector")),
        "stageCount": int(report.get("stageCount", 0) or 0),
        "hotspotCount": int(report.get("hotspotCount", 0) or 0),
        "baseline": _summarize_rate_stage(baseline),
        "current": _summarize_rate_stage(current),
        "delta": {
            key: delta.get(key)
            for key in (
                "beforeRiskPoints",
                "afterRiskPoints",
                "riskPointChange",
                "beforeRiskCount",
                "afterRiskCount",
                "relativeRiskChangePercent",
                "improvedDimensionCount",
                "regressedDimensionCount",
                "stableDimensionCount",
            )
        },
        "strategy": {
            "decision": str(plan.get("decision", "") or ""),
            "dimensionId": str(plan.get("dimensionId", "") or ""),
            "recommendedPromptId": str(plan.get("recommendedPromptId", "") or ""),
            "targetChunkCount": int(plan.get("targetChunkCount", 0) or 0),
            "canExecute": bool(plan.get("canExecute")),
            "targetScope": str(plan.get("targetScope", "") or ""),
            "maxAttempts": int(plan.get("maxAttempts", 0) or 0),
            "plateauPolicy": str(plan.get("plateauPolicy", "") or ""),
        },
        "binding": {
            "ready": bool(binding.get("ready")),
            "blockedReason": str(binding.get("blockedReason", "") or ""),
            "planDigest": str(binding.get("planDigest", "") or ""),
            "compareRevision": str(binding.get("compareRevision", "") or ""),
            "targetChunkCount": len(binding.get("targetChunkIds", []) if isinstance(binding.get("targetChunkIds"), list) else []),
        },
        "storesHotspotExcerpts": False,
    }


def _summarize_rate_stage(stage: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": str(stage.get("id", "") or ""),
        "round": stage.get("round"),
        "originalCharCount": int(stage.get("originalCharCount", 0) or 0),
        "analyzedCharCount": int(stage.get("analyzedCharCount", 0) or 0),
        "truncated": bool(stage.get("truncated")),
        "riskCount": int(stage.get("riskCount", 0) or 0),
        "highRiskCount": int(stage.get("highRiskCount", 0) or 0),
        "riskPoints": int(stage.get("riskPoints", 0) or 0),
        "metrics": {
            str(key): value
            for key, value in (stage.get("metrics") or {}).items()
            if isinstance(value, (str, int, float, bool)) or value is None
        },
    }


def _build_strategy_request(source_path: Path, output_path: Path, binding: dict[str, Any]) -> dict[str, Any]:
    return {
        "sourcePath": str(source_path),
        "outputPath": str(output_path),
        "dimensionId": str(binding.get("dimensionId", "") or ""),
        "recommendedPromptId": str(binding.get("recommendedPromptId", "") or ""),
        "compareRevision": str(binding.get("compareRevision", "") or ""),
        "scopeDigest": str(binding.get("scopeDigest", "") or ""),
        "formatDigest": str(binding.get("formatDigest", "") or ""),
        "sourceSha256": str(binding.get("sourceSha256", "") or ""),
        "targetChunkIds": list(binding.get("targetChunkIds", []) or []),
        "planDigest": str(binding.get("planDigest", "") or ""),
    }


def _select_representative_targets(
    snapshot: dict[str, Any],
    audit: dict[str, Any],
    *,
    max_targets: int,
) -> list[dict[str, Any]]:
    compare = snapshot.get("compare") if isinstance(snapshot.get("compare"), dict) else {}
    raw_chunks = [item for item in compare.get("chunks", []) if isinstance(item, dict)]
    _require(bool(raw_chunks), "representative_target_selection_empty")
    chunk_by_id = {
        str(item.get("chunkId", "") or ""): item
        for item in raw_chunks
        if str(item.get("chunkId", "") or "")
    }
    hotspots = [item for item in audit.get("hotspots", []) if isinstance(item, dict)]
    hotspot_by_id = {
        str(item.get("chunkId", "") or ""): item
        for item in hotspots
        if str(item.get("chunkId", "") or "") in chunk_by_id
    }
    internal = snapshot.get("_internal") if isinstance(snapshot.get("_internal"), dict) else {}
    body_map_payload = internal.get("bodyMapPayload") if isinstance(internal.get("bodyMapPayload"), dict) else {}
    body_map_units = body_map_payload.get("units") if isinstance(body_map_payload.get("units"), list) else []

    def format_anchor_summary(chunk: dict[str, Any]) -> tuple[int, list[str]]:
        try:
            paragraph_index = int(chunk.get("paragraphIndex", -1))
        except (TypeError, ValueError):
            return 0, []
        if paragraph_index < 0 or paragraph_index >= len(body_map_units):
            return 0, []
        unit = body_map_units[paragraph_index]
        raw_anchors = unit.get("format_anchors") if isinstance(unit, dict) else []
        anchors = [item for item in raw_anchors if isinstance(item, dict)] if isinstance(raw_anchors, list) else []
        kinds = sorted(
            {
                str(kind)
                for anchor in anchors
                for kind in (anchor.get("kinds") if isinstance(anchor.get("kinds"), list) else [])
                if str(kind)
            }
        )
        return len(anchors), kinds

    def paragraph_position(chunk: dict[str, Any]) -> tuple[int, int]:
        try:
            return (int(chunk.get("paragraphIndex", 0)), int(chunk.get("chunkIndex", 0)))
        except (TypeError, ValueError):
            return (10**9, 10**9)

    abstract_candidate = min(raw_chunks, key=paragraph_position)

    def technical_score(chunk: dict[str, Any]) -> tuple[int, int, int, int, int, int]:
        text = str(chunk.get("inputText", "") or "")
        format_anchor_count, _format_anchor_kinds = format_anchor_summary(chunk)
        return (
            1 if format_anchor_count > 0 else 0,
            len(CITATION_RE.findall(text)),
            len(UNIT_RE.findall(text)),
            len(TECH_TERM_RE.findall(text)),
            len(re.findall(r"\d", text)),
            min(len(text), 5000),
        )

    technical_candidate = max(raw_chunks, key=technical_score)
    risk_candidate = next(
        (chunk_by_id[str(item.get("chunkId", "") or "")] for item in hotspots if str(item.get("chunkId", "") or "") in chunk_by_id),
        max(raw_chunks, key=lambda item: len(str(item.get("inputText", "") or ""))),
    )

    role_candidates = (
        ("abstract", abstract_candidate),
        ("technical_numeric_citation", technical_candidate),
        ("rate_audit_highest_risk", risk_candidate),
    )
    selected_ids: list[str] = []
    role_map: dict[str, list[str]] = {}
    for role, chunk in role_candidates:
        chunk_id = str(chunk.get("chunkId", "") or "")
        if not chunk_id:
            continue
        role_map.setdefault(chunk_id, []).append(role)
        format_anchor_count, _format_anchor_kinds = format_anchor_summary(chunk)
        if format_anchor_count > 0 and "format_sensitive_anchor" not in role_map[chunk_id]:
            role_map[chunk_id].append("format_sensitive_anchor")
        if chunk_id not in selected_ids and len(selected_ids) < max_targets:
            selected_ids.append(chunk_id)

    fallback_ranked = sorted(
        raw_chunks,
        key=lambda item: (
            int((hotspot_by_id.get(str(item.get("chunkId", "") or "")) or {}).get("riskPoints", 0) or 0),
            technical_score(item),
            len(str(item.get("inputText", "") or "")),
        ),
        reverse=True,
    )
    for chunk in fallback_ranked:
        if len(selected_ids) >= max_targets:
            break
        chunk_id = str(chunk.get("chunkId", "") or "")
        if chunk_id and chunk_id not in selected_ids:
            selected_ids.append(chunk_id)
            role_map.setdefault(chunk_id, []).append("coverage_fallback")

    targets: list[dict[str, Any]] = []
    for order, chunk_id in enumerate(selected_ids, start=1):
        chunk = chunk_by_id[chunk_id]
        text = str(chunk.get("inputText", "") or "")
        hotspot = hotspot_by_id.get(chunk_id, {})
        format_anchor_count, format_anchor_kinds = format_anchor_summary(chunk)
        dimension_ids = [str(value) for value in (hotspot.get("dimensionIds") or []) if str(value)]
        targets.append(
            {
                "order": order,
                "chunkId": chunk_id,
                "roles": role_map.get(chunk_id, ["coverage_fallback"]),
                "inputChars": len(text),
                "citationCount": len(CITATION_RE.findall(text)),
                "unitCount": len(UNIT_RE.findall(text)),
                "technicalTermCount": len(TECH_TERM_RE.findall(text)),
                "formatAnchorCount": format_anchor_count,
                "formatAnchorKinds": format_anchor_kinds,
                "digitCount": len(re.findall(r"\d", text)),
                "riskPoints": int(hotspot.get("riskPoints", 0) or 0),
                "highRiskCount": int(hotspot.get("highRiskCount", 0) or 0),
                "dimensionIds": dimension_ids,
                "selectionTextSha256": _sha256_text(text),
            }
        )
    _require(bool(targets), "representative_target_selection_failed")
    return targets


def _find_effective_chunk(snapshot: dict[str, Any], chunk_id: str) -> dict[str, Any] | None:
    internal = snapshot.get("_internal") if isinstance(snapshot.get("_internal"), dict) else {}
    for item in internal.get("effectiveChunks", []) if isinstance(internal.get("effectiveChunks"), list) else []:
        if isinstance(item, dict) and str(item.get("chunkId", "") or "") == chunk_id:
            return item
    return None


def _find_compare_chunk(snapshot: dict[str, Any], chunk_id: str) -> dict[str, Any] | None:
    compare = snapshot.get("compare") if isinstance(snapshot.get("compare"), dict) else {}
    for item in compare.get("chunks", []) if isinstance(compare.get("chunks"), list) else []:
        if isinstance(item, dict) and str(item.get("chunkId", "") or "") == chunk_id:
            return item
    return None


def _candidate_release_evidence(
    chunk: dict[str, Any] | None,
    *,
    baseline_text: str | None = None,
    compare: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Validate v2 selection and source-relative evidence without retaining text.

    The canonical document release assessment is deliberately reused here so
    the E2E cannot approve a candidate with a missing targeted baseline,
    unresolvable document pattern profile, stale delta, or forged positive
    source-relative evidence.  The additional fields below are a text-free
    explanation layer for the E2E report, not a second release policy.
    """

    if not isinstance(chunk, dict):
        return {
            "eligible": False,
            "reasonCodes": ["candidate_chunk_missing"],
            "selectionPresent": False,
            "publishedRewrite": False,
            "selectedHashMatchesOutput": False,
            "readabilityPassed": False,
            "readabilityIssueCodes": ["candidate_chunk_missing"],
            "candidateSelectionV2": False,
            "sourceRelativeStylePassed": False,
            "sourcePatternProfileResolved": False,
            "productionReleaseAssessmentPassed": False,
            "storesCandidateText": False,
        }

    output_text = str(chunk.get("outputText", "") or "").strip()
    input_text = str(chunk.get("inputText", "") or "").strip()
    raw_candidate_baseline = chunk.get("candidateBaselineText")
    candidate_baseline_canonical = bool(
        isinstance(raw_candidate_baseline, str)
        and raw_candidate_baseline.strip()
        and raw_candidate_baseline == raw_candidate_baseline.strip()
    )
    comparison_baseline = (
        str(raw_candidate_baseline)
        if candidate_baseline_canonical
        else input_text
    )
    expected_baseline = str(baseline_text or "").strip() if baseline_text is not None else ""
    expected_baseline_matches = bool(
        baseline_text is None or expected_baseline == comparison_baseline
    )
    baseline_sha256 = _sha256_text(comparison_baseline)
    output_sha256 = _sha256_text(output_text)
    raw_selection = chunk.get("candidateSelection")
    selection_present = isinstance(raw_selection, dict)
    selection = raw_selection if isinstance(raw_selection, dict) else {}
    core_reason_codes: list[str] = []
    publication_reason_codes: list[str] = []
    selected_sha256 = ""
    selected_candidate: dict[str, Any] | None = None

    if not candidate_baseline_canonical:
        core_reason_codes.append("candidate_baseline_text_missing_or_invalid")
    if not expected_baseline_matches:
        core_reason_codes.append("candidate_baseline_context_mismatch")

    if not selection_present:
        core_reason_codes.append("candidate_selection_missing")
    else:
        if str(selection.get("event", "") or "") != "candidate-selection":
            core_reason_codes.append("candidate_selection_event_mismatch")
        if str(selection.get("schema", "") or "") != CANDIDATE_SELECTION_SCHEMA:
            core_reason_codes.append("candidate_selection_schema_mismatch")
        if selection.get("schemaVersion") != CANDIDATE_SELECTION_VERSION:
            core_reason_codes.append("candidate_selection_version_mismatch")
        if selection.get("isAiDetector") is not False:
            core_reason_codes.append("candidate_selection_ai_detector_claim_invalid")
        if selection.get("claimsDetectionRate") is not False:
            core_reason_codes.append("candidate_selection_detection_rate_claim_invalid")
        retention_assessment = selection.get("retentionAssessment")
        if (
            not isinstance(retention_assessment, dict)
            or retention_assessment.get("isAiDetector") is not False
            or retention_assessment.get("claimsDetectionRate") is not False
            or retention_assessment.get("claimsSemanticEquivalence") is not False
        ):
            core_reason_codes.append("candidate_selection_retention_claims_invalid")
        if selection.get("publishedRewrite") is not True:
            publication_reason_codes.append("candidate_selection_not_published")
        if selection.get("runFailed") is not False:
            core_reason_codes.append("candidate_selection_run_failed")
        if str(selection.get("decision", "") or "") != "generated_selected":
            publication_reason_codes.append("candidate_selection_decision_mismatch")
        if str(selection.get("selectedOrigin", "") or "") != "model":
            publication_reason_codes.append("candidate_selection_origin_mismatch")

        selected_sha256 = str(selection.get("selectedTextSha256", "") or "").strip().lower()
        if SHA256_RE.fullmatch(selected_sha256) is None:
            core_reason_codes.append("candidate_selection_selected_hash_invalid")
        elif selected_sha256 != output_sha256:
            core_reason_codes.append("candidate_selection_output_hash_mismatch")

        for hash_key in ("publishedTextSha256", "resultTextSha256"):
            raw_hash = selection.get(hash_key)
            if raw_hash in (None, ""):
                continue
            normalized_hash = str(raw_hash or "").strip().lower()
            if SHA256_RE.fullmatch(normalized_hash) is None or normalized_hash != output_sha256:
                core_reason_codes.append(f"candidate_selection_{hash_key}_mismatch")

        selected_candidate_id = str(selection.get("selectedCandidateId", "") or "")
        raw_candidates = selection.get("candidates")
        if isinstance(raw_candidates, list) and selected_candidate_id:
            selected_candidate = next(
                (
                    item
                    for item in raw_candidates
                    if isinstance(item, dict)
                    and str(item.get("candidateId", "") or "") == selected_candidate_id
                ),
                None,
            )
        if not isinstance(selected_candidate, dict):
            core_reason_codes.append("candidate_selection_selected_candidate_missing")
        else:
            candidate_sha256 = str(selected_candidate.get("textSha256", "") or "").strip().lower()
            if candidate_sha256 != selected_sha256:
                core_reason_codes.append("candidate_selection_candidate_hash_mismatch")
            if selected_candidate.get("hardValid") is not True:
                core_reason_codes.append("candidate_selection_candidate_hard_invalid")
            if selected_candidate.get("safetyEligible") is not True:
                core_reason_codes.append("candidate_selection_candidate_safety_ineligible")
            if selected_candidate.get("changedFromBaseline") is not True:
                publication_reason_codes.append("candidate_selection_candidate_unchanged")

    if not output_text or output_text == comparison_baseline:
        publication_reason_codes.append("candidate_output_unchanged")

    source_pattern_profiles: dict[str, dict[str, object]] | None = None
    registry_issue_codes: list[str] = []
    document_context_available = isinstance(compare, dict)
    if isinstance(compare, dict):
        try:
            source_pattern_profiles, normalized_registry_issues = (
                app_service._normalize_source_pattern_profile_registry(
                    compare.get("sourcePatternProfiles")
                )
            )
            registry_issue_codes.extend(str(code) for code in normalized_registry_issues)
        except Exception:
            source_pattern_profiles = {}
            registry_issue_codes.append("source_pattern_profile_registry_assessment_failed")

    result_source_relative = selection.get("resultSourceRelativeStyleDelta")
    source_relative_style_passed = source_relative_style_delta_passed(result_source_relative)
    result_binding = (
        result_source_relative.get("binding")
        if isinstance(result_source_relative, dict)
        and isinstance(result_source_relative.get("binding"), dict)
        else {}
    )
    source_relative_context_scope = str(
        result_source_relative.get("contextScope", "")
        if isinstance(result_source_relative, dict)
        else ""
    )
    source_profile_sha256 = str(result_binding.get("sourceProfileSha256", "") or "")
    source_pattern_profile_bound = bool(
        source_relative_context_scope == "document"
        and SHA256_RE.fullmatch(source_profile_sha256)
    )
    source_pattern_profile_resolved = bool(
        source_pattern_profile_bound
        and isinstance(source_pattern_profiles, dict)
        and source_profile_sha256 in source_pattern_profiles
    )
    source_relative_claims = (
        result_source_relative.get("claims")
        if isinstance(result_source_relative, dict)
        and isinstance(result_source_relative.get("claims"), dict)
        else {}
    )
    source_relative_claims_valid = bool(
        source_relative_claims.get("heuristicOnly") is True
        and source_relative_claims.get("isAiDetector") is False
        and source_relative_claims.get("claimsAuthorshipDetection") is False
        and source_relative_claims.get("claimsDetectionRate") is False
        and source_relative_claims.get("claimsSemanticEquivalence") is False
        and source_relative_claims.get("storesInputText") is False
        and source_relative_claims.get("storesOutputText") is False
        and source_relative_claims.get("storesMatchedText") is False
    )
    if not source_relative_style_passed:
        core_reason_codes.append("candidate_source_relative_style_evidence_invalid")
    if not source_relative_claims_valid:
        core_reason_codes.append("candidate_source_relative_style_claims_invalid")
    if document_context_available and not source_pattern_profile_resolved:
        core_reason_codes.append("candidate_source_pattern_profile_unresolved")
    core_reason_codes.extend(registry_issue_codes)

    try:
        production_assessment = app_service._assess_document_release_chunk(
            chunk,
            "rewrite_confirmed",
            source_pattern_profiles=source_pattern_profiles,
            source_pattern_profile_registry_issues=registry_issue_codes,
            document_context_available=document_context_available,
        )
    except Exception:
        production_assessment = {
            "ok": False,
            "issueCodes": ["candidate_production_release_assessment_failed"],
        }
    production_release_issue_codes = [
        str(code)
        for code in production_assessment.get("issueCodes", [])
        if str(code or "").strip()
    ]
    if production_assessment.get("ok") is not True:
        core_reason_codes.extend(production_release_issue_codes)

    try:
        readability = assess_academic_readability_delta(comparison_baseline, output_text)
    except Exception:
        readability = {
            "ok": False,
            "issueCodes": ["academic_readability_assessment_failed"],
        }
    raw_readability_codes = readability.get("issueCodes") if isinstance(readability, dict) else None
    readability_codes = list(
        dict.fromkeys(
            str(value).strip()
            for value in (raw_readability_codes if isinstance(raw_readability_codes, list) else [])
            if str(value).strip()
        )
    )
    readability_passed = bool(isinstance(readability, dict) and readability.get("ok") is True)
    if not readability_passed:
        core_reason_codes.append("academic_readability_regression")
        if not readability_codes:
            readability_codes.append("academic_readability_assessment_failed")

    baseline_preserved_safely = bool(
        selection_present
        and not list(dict.fromkeys(core_reason_codes))
        and selection.get("publishedRewrite") is False
        and selection.get("runFailed") is False
        and str(selection.get("decision", "") or "") == "preserved_baseline"
        and str(selection.get("selectedOrigin", "") or "") == "baseline"
        and selected_sha256 == output_sha256
        and output_text == comparison_baseline
        and readability_passed
        and isinstance(selected_candidate, dict)
        and selected_candidate.get("hardValid") is True
        and selected_candidate.get("safetyEligible") is True
        and production_assessment.get("ok") is True
    )
    normalized_reasons = list(
        dict.fromkeys([*core_reason_codes, *publication_reason_codes])
    )
    return {
        "eligible": bool(selection.get("publishedRewrite") is True and not normalized_reasons),
        "reasonCodes": normalized_reasons,
        "selectionPresent": selection_present,
        "selectionSchema": str(selection.get("schema", "") or ""),
        "selectionSchemaVersion": selection.get("schemaVersion"),
        "candidateSelectionV2": bool(
            selection.get("schema") == CANDIDATE_SELECTION_SCHEMA
            and selection.get("schemaVersion") == CANDIDATE_SELECTION_VERSION
        ),
        "decision": str(selection.get("decision", "") or ""),
        "publishedRewrite": selection.get("publishedRewrite") is True,
        "selectedOrigin": str(selection.get("selectedOrigin", "") or ""),
        "selectedTextSha256": selected_sha256,
        "outputTextSha256": output_sha256,
        "selectedHashMatchesOutput": bool(selected_sha256 and selected_sha256 == output_sha256),
        "candidateBaselinePresent": candidate_baseline_canonical,
        "candidateBaselineTextSha256": baseline_sha256,
        "candidateBaselineMatchesExpected": expected_baseline_matches,
        "baselinePreservedSafely": baseline_preserved_safely,
        "readabilityPassed": readability_passed,
        "readabilityIssueCodes": readability_codes,
        "sourceRelativeStylePassed": source_relative_style_passed,
        "sourceRelativeContextScope": source_relative_context_scope,
        "sourceRelativeClaimsValid": source_relative_claims_valid,
        "sourcePatternProfileBound": source_pattern_profile_bound,
        "sourcePatternProfileResolved": source_pattern_profile_resolved,
        "sourcePatternProfileRegistryIssueCodes": list(dict.fromkeys(registry_issue_codes)),
        "productionReleaseAssessmentPassed": production_assessment.get("ok") is True,
        "productionReleaseIssueCodes": list(dict.fromkeys(production_release_issue_codes)),
        "heuristicOnly": True,
        "isAiDetector": False,
        "claimsDetectionRate": False,
        "storesCandidateText": False,
    }


def _validate_automated_review_release(
    snapshot: dict[str, Any],
    decisions: dict[str, str],
) -> dict[str, Any]:
    """Prove every automated rewrite confirmation materialized an approved candidate."""

    compare = snapshot.get("compare") if isinstance(snapshot.get("compare"), dict) else {}
    chunks = [item for item in compare.get("chunks", []) if isinstance(item, dict)]
    chunk_by_id = {
        str(item.get("chunkId", "") or ""): item
        for item in chunks
        if str(item.get("chunkId", "") or "")
    }
    rewrite_ids = [chunk_id for chunk_id, decision in decisions.items() if decision == "rewrite_confirmed"]
    _require(all(chunk_id in chunk_by_id for chunk_id in decisions), "automated_review_chunk_mismatch")

    for chunk_id in rewrite_ids:
        chunk = chunk_by_id[chunk_id]
        evidence = _candidate_release_evidence(chunk, compare=compare)
        _require(bool(evidence.get("eligible")), "automated_review_candidate_evidence_invalid", "model_output_failure")
        try:
            app_service.validate_chunk_output(
                str(chunk.get("inputText", "") or ""),
                str(chunk.get("outputText", "") or ""),
                chunk_id,
            )
        except ValueError as exc:
            raise E2EContractError("automated_review_candidate_hard_invalid", "model_output_failure") from exc
        effective = _find_effective_chunk(snapshot, chunk_id)
        _require(isinstance(effective, dict), "automated_review_effective_chunk_missing")
        effective_sha256 = _sha256_text(str((effective or {}).get("text", "") or "").strip())
        _require(
            effective_sha256 == str(evidence.get("outputTextSha256", "") or ""),
            "automated_review_materialized_hash_mismatch",
        )

    return {
        "rewriteConfirmedCandidateCount": len(rewrite_ids),
        "allRewriteConfirmationsProductionSelected": True,
        "allRewriteConfirmationsHashBound": True,
        "allRewriteConfirmationsReadabilityPassed": True,
    }


def _safe_direction_summary(direction: dict[str, Any]) -> dict[str, Any]:
    allowed_scalars = {
        "dimensionId",
        "direction",
        "primaryMetric",
        "before",
        "after",
        "variationBefore",
        "variationAfter",
        "openingConcentrationBefore",
        "openingConcentrationAfter",
        "ok",
        "satisfied",
        "effective",
    }
    result = {
        str(key): value
        for key, value in direction.items()
        if key in allowed_scalars and (isinstance(value, (str, int, float, bool)) or value is None)
    }
    structure = direction.get("structureDirection")
    if isinstance(structure, dict):
        result["structureDirection"] = {
            str(key): value
            for key, value in structure.items()
            if key in allowed_scalars and (isinstance(value, (str, int, float, bool)) or value is None)
        }
    return result


def _summarize_failed_attempts(chunk: dict[str, Any] | None) -> list[dict[str, Any]]:
    """Project failed-attempt v1 evidence without ever reading private prose."""

    if not isinstance(chunk, dict):
        return []
    raw_attempts = chunk.get("failedAttempts")
    if not isinstance(raw_attempts, list):
        return []
    summaries: list[dict[str, Any]] = []
    for raw in raw_attempts[-4:]:
        if not isinstance(raw, dict):
            continue
        raw_keys = {str(key).strip().lower() for key in raw}
        forbidden_field_observed = bool(
            raw_keys.intersection(FAILED_ATTEMPT_PUBLIC_FORBIDDEN_KEYS)
        )
        guard_category = str(raw.get("guardCategory", "") or "").strip().lower()
        issue_codes = [
            str(value).strip().lower()
            for value in (raw.get("issueCodes") if isinstance(raw.get("issueCodes"), list) else [])
            if str(value).strip().lower() in FAILED_ATTEMPT_ISSUE_CODES
        ]
        issue_codes = list(dict.fromkeys(issue_codes))
        try:
            output_chars = int(raw.get("outputCharCount", 0) or 0)
        except (TypeError, ValueError):
            output_chars = 0
        output_sha256 = str(raw.get("outputTextSha256", "") or "").strip().lower()
        output_sha256_valid = bool(
            not output_sha256 or SHA256_RE.fullmatch(output_sha256)
        )
        schema_valid = bool(
            raw.get("schema") == FAILED_ATTEMPT_EVIDENCE_SCHEMA
            and raw.get("schemaVersion") == FAILED_ATTEMPT_EVIDENCE_VERSION
        )
        safe_evidence_valid = bool(
            schema_valid
            and not forbidden_field_observed
            and guard_category in FAILED_ATTEMPT_GUARD_CATEGORIES
            and bool(issue_codes)
            and output_chars >= 0
            and output_sha256_valid
            and raw.get("textStored") is False
            and raw.get("errorStored") is False
            and raw.get("reasoningSuppressed") is True
            and raw.get("providerContentStored") is False
        )
        summaries.append(
            {
                "schema": str(raw.get("schema", "") or ""),
                "schemaVersion": raw.get("schemaVersion"),
                "schemaValid": schema_valid,
                "attempt": raw.get("attempt"),
                "guardCategory": guard_category,
                "issueCodes": issue_codes,
                "outputCharCount": output_chars,
                "outputTextSha256": output_sha256 if output_sha256_valid else "",
                "truncated": bool(raw.get("truncated")),
                "textStored": False,
                "errorStored": False,
                "reasoningSuppressed": True,
                "providerContentStored": False,
                "forbiddenFieldObserved": forbidden_field_observed,
                "safeEvidenceValid": safe_evidence_valid,
            }
        )
    return summaries


def _same_dimension_evaluation(input_text: str, output_text: str, dimension_ids: list[str]) -> dict[str, Any]:
    for dimension_id in dimension_ids:
        definition = get_rate_audit_dimension_definition(dimension_id)
        prompt_id = str(definition.get("repairPromptId", "") or "")
        evaluator = dict(ROUND_PERTURBATION_DIMENSIONS.get(prompt_id, {}))
        if not bool(definition.get("canExecute")) or not evaluator:
            continue
        converged, direction = app_service._dimension_converged(input_text, output_text, evaluator)
        return {
            "available": True,
            "dimensionId": dimension_id,
            "evaluatorDimensionId": str(definition.get("evaluatorDimensionId", "") or ""),
            "primaryMetric": str(definition.get("primaryMetric", "") or ""),
            "converged": bool(converged),
            "direction": _safe_direction_summary(direction if isinstance(direction, dict) else {}),
        }
    return {
        "available": False,
        "dimensionId": str(dimension_ids[0]) if dimension_ids else "",
        "evaluatorDimensionId": "",
        "primaryMetric": "",
        "converged": None,
        "direction": {},
    }


def _summarize_format_anchor_executions(
    executions: list[dict[str, Any]],
    format_anchor_target_ids: set[str],
) -> dict[str, int]:
    """Count real format-anchor coverage without requiring a needless rewrite.

    A protected-format target is successful when the real provider was
    exercised and production evidence proves either a publishable selected
    candidate *or* a hash-bound safe baseline preservation.  Requiring every
    format-anchor target to publish would contradict the candidate selector:
    an unchanged or no-gain model response must keep the accepted baseline.
    """

    exercised_ids: set[str] = set()
    safely_resolved_ids: set[str] = set()
    published_ids: set[str] = set()
    preserved_ids: set[str] = set()
    for execution in executions:
        chunk_id = str(execution.get("chunkId", "") or "")
        if chunk_id not in format_anchor_target_ids:
            continue
        try:
            real_call_count = int(execution.get("realCallCount", 0) or 0)
        except (TypeError, ValueError):
            real_call_count = 0
        if real_call_count <= 0:
            continue
        exercised_ids.add(chunk_id)
        evidence = execution.get("candidateReleaseEvidence")
        if not isinstance(evidence, dict) or execution.get("hardValidationPassed") is not True:
            continue
        common_ready = bool(
            evidence.get("selectionPresent") is True
            and evidence.get("candidateSelectionV2") is True
            and evidence.get("selectedHashMatchesOutput") is True
            and evidence.get("readabilityPassed") is True
            and evidence.get("sourceRelativeStylePassed") is True
            and evidence.get("productionReleaseAssessmentPassed") is True
            and evidence.get("storesCandidateText") is False
        )
        published_ready = bool(
            common_ready
            and evidence.get("eligible") is True
            and evidence.get("publishedRewrite") is True
            and str(evidence.get("selectedOrigin", "") or "") == "model"
        )
        preserved_ready = bool(
            common_ready
            and evidence.get("baselinePreservedSafely") is True
            and evidence.get("publishedRewrite") is False
            and str(evidence.get("decision", "") or "") == "preserved_baseline"
            and str(evidence.get("selectedOrigin", "") or "") == "baseline"
        )
        if published_ready:
            published_ids.add(chunk_id)
            safely_resolved_ids.add(chunk_id)
        elif preserved_ready:
            preserved_ids.add(chunk_id)
            safely_resolved_ids.add(chunk_id)
    return {
        "formatAnchorExercisedTargetCount": len(exercised_ids),
        "formatAnchorSafelyResolvedTargetCount": len(safely_resolved_ids),
        "formatAnchorPublishedCandidateCount": len(published_ids),
        "formatAnchorBaselinePreservedCount": len(preserved_ids),
    }


def _passed_e2e_status(*, manual_review: bool, changed_rewrite_count: int) -> str:
    """Describe a successful release without using rewrite count as a quota."""

    if changed_rewrite_count <= 0:
        return (
            "passed_baseline_preserved_with_manual_review"
            if manual_review
            else "passed_baseline_preserved"
        )
    return "passed_with_manual_review" if manual_review else "passed"


def _execute_bounded_real_targets(
    *,
    source_path: Path,
    output_path: Path,
    initial_snapshot: dict[str, Any],
    initial_audit: dict[str, Any],
    targets: list[dict[str, Any]],
    model_config: dict[str, Any],
    call_auditor: CompletionCallAuditor,
) -> tuple[list[dict[str, Any]], set[str], dict[str, Any], list[str], list[str], list[str]]:
    selected_ids = {str(item.get("chunkId", "") or "") for item in targets}
    target_by_id = {str(item.get("chunkId", "") or ""): item for item in targets}
    real_candidate_ids: set[str] = set()
    external_failure_ids: list[str] = []
    provider_configuration_failure_ids: list[str] = []
    product_failure_ids: list[str] = []
    executions: dict[str, dict[str, Any]] = {
        chunk_id: {
            "chunkId": chunk_id,
            "roles": list(target_by_id[chunk_id].get("roles", [])),
            "paths": [],
            "outcome": "pending",
            "realCallCount": 0,
            "hardValidationPassed": False,
            "sameDimension": {"available": False, "converged": None},
            "candidateChanged": False,
        }
        for chunk_id in selected_ids
    }

    binding = initial_audit.get("strategyBinding") if isinstance(initial_audit.get("strategyBinding"), dict) else {}
    bound_ids = [str(value) for value in (binding.get("targetChunkIds") or []) if str(value)]
    strategy_summary: dict[str, Any] = {
        "ready": bool(binding.get("ready")),
        "eligibleWithinRepresentativeBudget": bool(bound_ids and set(bound_ids).issubset(selected_ids)),
        "targetCount": len(bound_ids),
        "executed": False,
        "outcome": "not_ready" if not bool(binding.get("ready")) else "representative_budget_mismatch",
        "successCount": 0,
        "failureCount": 0,
        "canceled": False,
        "realCallCount": 0,
        "convergedTargetCount": 0,
        "nonConvergedTargetCount": 0,
    }
    completed_by_strategy: set[str] = set()
    if bool(binding.get("ready")) and bound_ids and set(bound_ids).issubset(selected_ids):
        before_calls = len(call_auditor.calls)
        strategy_summary["executed"] = True
        strategy_summary["outcome"] = "running"
        try:
            strategy_result = app_service.execute_rate_audit_strategy(
                _build_strategy_request(source_path, output_path, binding),
                model_config,
            )
        except Exception as exc:
            after_calls = len(call_auditor.calls)
            call_slice = call_auditor.calls[before_calls:after_calls]
            provider_configuration_failure = _provider_configuration_failure_from_calls(call_slice)
            if provider_configuration_failure is None:
                provider_configuration_failure = _provider_configuration_failure_summary(
                    _provider_error_descriptor(exc)
                )
            strategy_summary["realCallCount"] = after_calls - before_calls
            if provider_configuration_failure is not None:
                strategy_summary["outcome"] = "provider_configuration_failure"
                strategy_summary["providerFailure"] = provider_configuration_failure
                provider_configuration_failure_ids.extend(bound_ids)
            elif _is_external_provider_failure(exc):
                strategy_summary["outcome"] = "external_failure"
                external_failure_ids.extend(bound_ids)
            else:
                strategy_summary["outcome"] = "pipeline_failure"
                product_failure_ids.extend(bound_ids)
            for chunk_id in bound_ids:
                executions[chunk_id]["paths"].append("bound_strategy")
                executions[chunk_id]["realCallCount"] += max(0, after_calls - before_calls)
                if provider_configuration_failure is not None:
                    executions[chunk_id]["outcome"] = "provider_configuration_failure"
                    executions[chunk_id]["providerFailure"] = provider_configuration_failure
        else:
            after_calls = len(call_auditor.calls)
            success_ids = {str(value) for value in (strategy_result.get("successChunkIds") or []) if str(value)}
            strategy_summary.update(
                {
                    "outcome": "converged" if success_ids and not strategy_result.get("failures") else "non_converged",
                    "successCount": int(strategy_result.get("successCount", 0) or 0),
                    "failureCount": int(strategy_result.get("failureCount", 0) or 0),
                    "canceled": bool(strategy_result.get("canceled")),
                    "realCallCount": after_calls - before_calls,
                    "convergedTargetCount": len(success_ids),
                    "nonConvergedTargetCount": max(0, len(bound_ids) - len(success_ids)),
                }
            )
            latest_snapshot = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
            for chunk_id in bound_ids:
                execution = executions[chunk_id]
                execution["paths"].append("bound_strategy")
                execution["realCallCount"] += max(0, after_calls - before_calls)
                chunk = _find_compare_chunk(latest_snapshot, chunk_id)
                execution["failedAttempts"] = _summarize_failed_attempts(chunk)
                if chunk_id in success_ids and isinstance(chunk, dict):
                    # A successful strategy transaction is complete even when
                    # its candidate is later rejected by the independent E2E
                    # release checks.  Do not fall through to the legacy rerun:
                    # the pending strategy-review flag intentionally blocks it.
                    completed_by_strategy.add(chunk_id)
                    hard_ok = False
                    try:
                        app_service.validate_chunk_output(
                            str(chunk.get("inputText", "") or ""),
                            str(chunk.get("outputText", "") or ""),
                            chunk_id,
                        )
                        hard_ok = True
                    except ValueError:
                        hard_ok = False
                    converged = chunk.get("rerunDimensionConverged") is True
                    baseline_chunk = _find_effective_chunk(initial_snapshot, chunk_id)
                    baseline_text = (
                        str((baseline_chunk or {}).get("text", "") or "")
                        if isinstance(baseline_chunk, dict)
                        else str(chunk.get("inputText", "") or "")
                    )
                    release_evidence = _candidate_release_evidence(
                        chunk,
                        baseline_text=baseline_text,
                        compare=(
                            latest_snapshot.get("compare")
                            if isinstance(latest_snapshot.get("compare"), dict)
                            else {}
                        ),
                    )
                    release_ready = bool(hard_ok and converged and release_evidence.get("eligible"))
                    if release_ready:
                        outcome = "converged"
                    elif not hard_ok:
                        outcome = "strategy_candidate_hard_validation_failed"
                    elif not converged:
                        outcome = "strategy_non_converged"
                    elif release_evidence.get("readabilityPassed") is not True:
                        outcome = "candidate_readability_failed"
                    elif release_evidence.get("publishedRewrite") is not True:
                        outcome = "candidate_not_published"
                    else:
                        outcome = "candidate_selection_evidence_invalid"
                    execution.update(
                        {
                            "outcome": outcome,
                            "hardValidationPassed": hard_ok,
                            "sameDimension": {
                                "available": True,
                                "dimensionId": str(binding.get("dimensionId", "") or ""),
                                "evaluatorDimensionId": str(binding.get("evaluatorDimensionId", "") or ""),
                                "primaryMetric": str(binding.get("primaryMetric", "") or ""),
                                "converged": converged,
                                "direction": {},
                            },
                            "candidateChanged": str(chunk.get("outputText", "") or "") != str(chunk.get("inputText", "") or ""),
                            "candidateSha256": _sha256_text(str(chunk.get("outputText", "") or "").strip()),
                            "candidateReleaseEvidence": release_evidence,
                        }
                    )
                    if release_ready:
                        real_candidate_ids.add(chunk_id)
                else:
                    execution["outcome"] = "strategy_non_converged"

    for target in targets:
        chunk_id = str(target.get("chunkId", "") or "")
        if chunk_id in completed_by_strategy:
            continue
        before_snapshot = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
        effective = _find_effective_chunk(before_snapshot, chunk_id)
        _require(isinstance(effective, dict), "target_effective_chunk_missing")
        effective_input = str((effective or {}).get("text", "") or "")
        before_calls = len(call_auditor.calls)
        execution = executions[chunk_id]
        execution["paths"].append("legacy_targeted_rerun")
        feedback = (
            "REAL_THESIS_E2E_BOUNDED_REVIEW: perform one conservative local repair only. "
            "Preserve facts, claims, citations, numbers, units, technical terms, paragraph count, and scope. "
            "Return body text only; if the source is already natural, keep changes minimal."
        )
        try:
            rerun_result = app_service.rerun_compare_chunk(
                str(output_path),
                chunk_id,
                model_config,
                feedback,
            )
        except Exception as exc:
            after_calls = len(call_auditor.calls)
            call_count = after_calls - before_calls
            execution["realCallCount"] += call_count
            call_slice = call_auditor.calls[before_calls:after_calls]
            provider_configuration_failure = _provider_configuration_failure_from_calls(call_slice)
            if provider_configuration_failure is None:
                provider_configuration_failure = _provider_configuration_failure_summary(
                    _provider_error_descriptor(exc)
                )
            has_external_call_error = any(
                str(item.get("errorCategory", "") or "") in {"network", "timeout", "rate_limit", "server"}
                or isinstance(item.get("statusCode"), int) and int(item["statusCode"]) >= 500
                for item in call_slice
            )
            if provider_configuration_failure is not None:
                execution["outcome"] = "provider_configuration_failure"
                execution["providerFailure"] = provider_configuration_failure
                provider_configuration_failure_ids.append(chunk_id)
            elif has_external_call_error or _is_external_provider_failure(exc):
                execution["outcome"] = "external_failure"
                external_failure_ids.append(chunk_id)
            elif call_count > 0:
                execution["outcome"] = "model_output_hard_validation_failed"
            else:
                execution["outcome"] = "pipeline_failure_before_model_call"
                product_failure_ids.append(chunk_id)
            try:
                failed_snapshot = app_service.read_round_artifact_snapshot(output_path, include_internal=False)
            except Exception:
                failed_chunk = None
            else:
                failed_chunk = _find_compare_chunk(failed_snapshot, chunk_id)
            execution["failedAttempts"] = _summarize_failed_attempts(failed_chunk)
            continue

        after_calls = len(call_auditor.calls)
        execution["realCallCount"] += after_calls - before_calls
        chunk = rerun_result.get("chunk") if isinstance(rerun_result.get("chunk"), dict) else {}
        latest_snapshot = app_service.read_round_artifact_snapshot(
            output_path,
            include_internal=True,
        )
        persisted_chunk = _find_compare_chunk(latest_snapshot, chunk_id)
        if isinstance(persisted_chunk, dict):
            chunk = persisted_chunk
        candidate_output = str(chunk.get("outputText", "") or "")
        hard_ok = False
        try:
            app_service.validate_chunk_output(effective_input, candidate_output, chunk_id)
            hard_ok = True
        except ValueError:
            hard_ok = False
        same_dimension = _same_dimension_evaluation(
            effective_input,
            candidate_output,
            [str(value) for value in (target.get("dimensionIds") or []) if str(value)],
        )
        release_evidence = _candidate_release_evidence(
            chunk,
            baseline_text=effective_input,
            compare=(
                latest_snapshot.get("compare")
                if isinstance(latest_snapshot.get("compare"), dict)
                else {}
            ),
        )
        dimension_ready = not (
            same_dimension.get("available") is True
            and same_dimension.get("converged") is False
        )
        release_ready = bool(hard_ok and dimension_ready and release_evidence.get("eligible"))
        if release_ready:
            outcome = "candidate_ready"
        elif not hard_ok:
            outcome = "candidate_hard_validation_failed"
        elif not dimension_ready:
            outcome = "candidate_dimension_not_converged"
        elif release_evidence.get("baselinePreservedSafely") is True:
            outcome = "baseline_preserved"
        elif release_evidence.get("readabilityPassed") is not True:
            outcome = "candidate_readability_failed"
        elif release_evidence.get("publishedRewrite") is not True:
            outcome = "candidate_stale_preserved_evidence"
        else:
            outcome = "candidate_selection_evidence_invalid"
        execution.update(
            {
                "outcome": outcome,
                "hardValidationPassed": hard_ok,
                "sameDimension": same_dimension,
                "candidateChanged": candidate_output != effective_input,
                "candidateSha256": _sha256_text(candidate_output) if candidate_output else "",
                "candidateReleaseEvidence": release_evidence,
                "failedAttempts": _summarize_failed_attempts(chunk),
            }
        )
        if release_ready:
            real_candidate_ids.add(chunk_id)
        elif not hard_ok:
            product_failure_ids.append(chunk_id)

    return (
        [executions[str(item.get("chunkId", "") or "")] for item in targets],
        real_candidate_ids,
        strategy_summary,
        sorted(set(provider_configuration_failure_ids)),
        sorted(set(external_failure_ids)),
        sorted(set(product_failure_ids)),
    )


def _build_review_decisions(
    snapshot: dict[str, Any],
    *,
    baseline_mode: str,
    real_candidate_ids: set[str],
    executions: list[dict[str, Any]],
) -> tuple[dict[str, str], dict[str, Any]]:
    compare = snapshot.get("compare") if isinstance(snapshot.get("compare"), dict) else {}
    chunks = [item for item in compare.get("chunks", []) if isinstance(item, dict)]
    execution_by_id = {str(item.get("chunkId", "") or ""): item for item in executions}
    decisions: dict[str, str] = {}
    reasons: Counter[str] = Counter()
    hard_validation_failures = 0
    changed_rewrite_count = 0
    candidate_evidence_failures = 0
    readability_failures = 0
    preserved_baseline_count = 0
    for chunk in chunks:
        chunk_id = str(chunk.get("chunkId", "") or "")
        _require(bool(chunk_id), "review_chunk_id_missing")
        input_text = str(chunk.get("inputText", "") or "")
        output_text = str(chunk.get("outputText", "") or "")
        hard_ok = False
        try:
            app_service.validate_chunk_output(input_text, output_text, chunk_id)
            hard_ok = True
        except ValueError:
            hard_validation_failures += 1

        execution = execution_by_id.get(chunk_id, {})
        same_dimension = execution.get("sameDimension") if isinstance(execution.get("sameDimension"), dict) else {}
        strategy_pending = chunk.get("rateAuditStrategyReviewRequired") is True
        strategy_converged = chunk.get("rerunDimensionConverged") is True
        default = app_service._default_export_decision_for_chunk(chunk)
        release_evidence = _candidate_release_evidence(chunk, compare=compare)
        candidate_release_ready = bool(release_evidence.get("eligible"))
        if release_evidence.get("readabilityPassed") is not True and output_text.strip() != input_text.strip():
            readability_failures += 1
        if (
            isinstance(chunk.get("candidateSelection"), dict)
            and str((chunk.get("candidateSelection") or {}).get("decision", "") or "") == "preserved_baseline"
        ):
            preserved_baseline_count += 1

        if chunk_id not in real_candidate_ids and baseline_mode == "identity":
            decision = "source_confirmed"
            reason = "identity_baseline_unselected"
        elif strategy_pending:
            if chunk_id in real_candidate_ids and hard_ok and strategy_converged and candidate_release_ready:
                decision = "rewrite_confirmed"
                reason = "strategy_production_selected_hash_bound_and_readable"
            else:
                decision = "source_confirmed"
                if release_evidence.get("readabilityPassed") is not True:
                    reason = "strategy_academic_readability_failed"
                elif not candidate_release_ready:
                    reason = "strategy_candidate_selection_not_release_eligible"
                else:
                    reason = "strategy_not_safe_to_confirm"
                candidate_evidence_failures += 1
        elif chunk_id in real_candidate_ids:
            # A same-dimension failure is an explicit stop/manual-review signal.
            # Do not export that candidate merely because its generic hard
            # validation passed.
            if not candidate_release_ready:
                decision = "source_confirmed"
                if release_evidence.get("readabilityPassed") is not True:
                    reason = "academic_readability_failed"
                else:
                    reason = "production_candidate_selection_not_release_eligible"
                candidate_evidence_failures += 1
            elif same_dimension.get("available") is True and same_dimension.get("converged") is False:
                decision = "source_confirmed"
                reason = "same_dimension_not_converged"
            elif default == "rewrite" and hard_ok:
                decision = "rewrite_confirmed"
                reason = "production_selected_hash_bound_readable_and_server_rewrite"
            else:
                decision = "source_confirmed"
                reason = "server_default_source_or_hard_invalid"
        elif baseline_mode == "full_real_round":
            if candidate_release_ready and default == "rewrite" and hard_ok:
                decision = "rewrite_confirmed"
                reason = "full_real_round_production_selected_hash_bound_and_readable"
            else:
                decision = "source_confirmed"
                if release_evidence.get("readabilityPassed") is not True:
                    reason = "full_real_round_academic_readability_failed"
                elif not candidate_release_ready:
                    reason = "full_real_round_candidate_selection_not_release_eligible"
                else:
                    reason = "full_real_round_server_default_source_or_hard_invalid"
                candidate_evidence_failures += 1
        else:
            decision = "source_confirmed"
            reason = "server_default_source"

        decisions[chunk_id] = decision
        reasons[reason] += 1
        if decision == "rewrite_confirmed" and output_text != input_text:
            changed_rewrite_count += 1

    _require(len(decisions) == len(chunks), "review_decisions_not_full_coverage")
    return decisions, {
        "chunkCount": len(chunks),
        "decisionCount": len(decisions),
        "sourceConfirmedCount": sum(1 for value in decisions.values() if value == "source_confirmed"),
        "rewriteConfirmedCount": sum(1 for value in decisions.values() if value == "rewrite_confirmed"),
        "changedRewriteConfirmedCount": changed_rewrite_count,
        "hardValidationFailureCount": hard_validation_failures,
        "candidateEvidenceFailureCount": candidate_evidence_failures,
        "academicReadabilityFailureCount": readability_failures,
        "preservedBaselineSelectionCount": preserved_baseline_count,
        "reasonCounts": dict(sorted(reasons.items())),
        "allDecisionsExplicit": all(value in {"source_confirmed", "rewrite_confirmed"} for value in decisions.values()),
        "reviewActor": "automated_e2e",
        "isHumanReview": False,
        "decisionAuthority": "production_candidate_selection_plus_release_gates",
    }


def _export_and_verify(
    output_path: Path,
    requested_export_path: Path,
    preflight_contract: dict[str, Any],
    reviewed_snapshot: dict[str, Any],
) -> dict[str, Any]:
    result = app_service.export_round_output(
        str(output_path),
        str(requested_export_path),
        "docx",
        "preserve_original",
        expected_doc_id=str(reviewed_snapshot.get("docId", "") or ""),
        expected_round=int(reviewed_snapshot.get("round", 0) or 0),
        expected_compare_revision=str(reviewed_snapshot.get("compareRevision", "") or ""),
        expected_content_revision=str(reviewed_snapshot.get("contentRevision", "") or ""),
        expected_artifact_snapshot_digest=str(reviewed_snapshot.get("artifactSnapshotDigest", "") or ""),
    )
    actual_path = Path(str(result.get("path", ""))).resolve()
    _require(actual_path.exists() and actual_path.is_file(), "certified_export_missing")
    required_checks = {
        "document_generation",
        "format_preflight",
        "pre_export_guard",
        "content_contract",
        "text_integrity",
        "protected_text_audit",
        "ooxml_integrity",
        "format_lock",
        "post_export_contract",
    }
    performed = {str(value) for value in (result.get("checksPerformed") or []) if str(value)}
    _require(str(result.get("overallStatus", "")) == "passed", "export_overall_status_failed")
    _require(str(result.get("sourceKind", "")) == "original_docx", "export_lost_docx_provenance")
    _require(str(result.get("contentContractStatus", "")) == "passed", "export_content_contract_status_failed")
    _require(str(result.get("formatLockStatus", "")) == "passed", "export_format_lock_status_failed")
    _require(required_checks.issubset(performed), "export_required_checks_missing")
    for key in (
        "auditIssueCount",
        "ooxmlAuditIssueCount",
        "formatLockIssueCount",
        "contentContractIssueCount",
        "preflightIssueCount",
        "guardIssueCount",
    ):
        _require(int(result.get(key, 0) or 0) == 0, f"export_{key}_nonzero")
    _require(bool(result.get("contentContractReady")), "export_content_contract_not_ready")
    _require(bool(result.get("modelInputMatchesEditableUnits")), "export_model_input_scope_mismatch")
    snapshot_fields = (
        "outputPath",
        "docId",
        "round",
        "compareRevision",
        "reviewRevision",
        "contentRevision",
        "artifactSnapshotDigest",
        "effectiveTextSha256",
    )
    for key in snapshot_fields:
        _require(
            result.get(key) == reviewed_snapshot.get(key),
            f"export_reviewed_snapshot_{key}_mismatch",
        )
    artifact_sha256 = _sha256_file(actual_path)
    _require(artifact_sha256 == str(result.get("artifactSha256", "") or ""), "certified_export_sha256_mismatch")

    post_contract = _read_json_object(str(result.get("contentContractPath", "") or ""))
    _require(bool(post_contract.get("ready")), "post_export_contract_report_not_ready")
    for key in ("sourceSha256", "scopeDigest", "formatDigest"):
        _require(
            str(post_contract.get(key, "") or "") == str(preflight_contract.get(key, "") or ""),
            f"post_export_{key}_mismatch",
        )

    evidence_manifest_path = Path(str(result.get("evidenceManifestPath", ""))).resolve()
    evidence_manifest = _read_json_object(evidence_manifest_path)
    _require(evidence_manifest_path.exists(), "export_evidence_manifest_missing")
    _require(str(evidence_manifest.get("status", "")) == "passed", "export_evidence_manifest_failed")
    _require(str(evidence_manifest.get("artifactSha256", "")) == artifact_sha256, "export_evidence_sha256_mismatch")
    manifest_snapshot = (
        evidence_manifest.get("roundArtifactSnapshot")
        if isinstance(evidence_manifest.get("roundArtifactSnapshot"), dict)
        else {}
    )
    for key in snapshot_fields:
        manifest_key = "version" if key == "roundSnapshotVersion" else key
        _require(
            manifest_snapshot.get(manifest_key) == reviewed_snapshot.get(key),
            f"export_manifest_reviewed_snapshot_{key}_mismatch",
        )

    reports: dict[str, Any] = {}
    for key in ("guardPath", "auditPath", "ooxmlAuditPath", "formatLockPath", "contentContractPath", "preflightPath"):
        report_path = Path(str(result.get(key, ""))).resolve()
        _require(report_path.exists() and report_path.is_file(), f"export_report_missing_{key}")
        reports[key] = {"path": str(report_path), "sha256": _sha256_file(report_path)}

    return {
        "requestedAliasPath": str(requested_export_path),
        "certifiedArtifactPath": str(actual_path),
        "artifactSha256": artifact_sha256,
        "artifactBytes": actual_path.stat().st_size,
        "evidenceManifestPath": str(evidence_manifest_path),
        "evidenceManifestSha256": _sha256_file(evidence_manifest_path),
        "overallStatus": str(result.get("overallStatus", "") or ""),
        "sourceKind": str(result.get("sourceKind", "") or ""),
        "contentContractStatus": str(result.get("contentContractStatus", "") or ""),
        "formatLockStatus": str(result.get("formatLockStatus", "") or ""),
        "checksPerformed": sorted(performed),
        "auditIssueCount": int(result.get("auditIssueCount", 0) or 0),
        "ooxmlAuditIssueCount": int(result.get("ooxmlAuditIssueCount", 0) or 0),
        "formatLockIssueCount": int(result.get("formatLockIssueCount", 0) or 0),
        "contentContractIssueCount": int(result.get("contentContractIssueCount", 0) or 0),
        "preflightIssueCount": int(result.get("preflightIssueCount", 0) or 0),
        "guardIssueCount": int(result.get("guardIssueCount", 0) or 0),
        "contentContractReady": bool(result.get("contentContractReady")),
        "modelInputMatchesEditableUnits": bool(result.get("modelInputMatchesEditableUnits")),
        "roundArtifactSnapshot": {
            "outputPath": str(result.get("outputPath", "") or ""),
            "docId": str(result.get("docId", "") or ""),
            "round": int(result.get("round", 0) or 0),
            "compareRevision": str(result.get("compareRevision", "") or ""),
            "reviewRevision": str(result.get("reviewRevision", "") or ""),
            "contentRevision": str(result.get("contentRevision", "") or ""),
            "artifactSnapshotDigest": str(result.get("artifactSnapshotDigest", "") or ""),
            "effectiveTextSha256": str(result.get("effectiveTextSha256", "") or ""),
        },
        "reviewMaterializationBound": True,
        "reports": reports,
        "contractDigestMatch": {
            "sourceSha256": True,
            "scopeDigest": True,
            "formatDigest": True,
        },
    }


def _secret_scan_report(payload: dict[str, Any], *, api_key: str) -> dict[str, Any]:
    serialized = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    findings: list[str] = []
    if api_key and api_key in serialized:
        findings.append("exact_runtime_credential")
    for code, pattern in SECRET_PATTERNS:
        if pattern.search(serialized):
            findings.append(code)
    forbidden_key_names = {"apikey", "api_key", "authorization", "password", "secret", "token"}

    def walk(value: Any) -> None:
        if isinstance(value, dict):
            for key, item in value.items():
                if str(key).casefold() in forbidden_key_names:
                    findings.append("forbidden_secret_field")
                walk(item)
        elif isinstance(value, list):
            for item in value:
                walk(item)

    walk(payload)
    return {
        "ok": not findings,
        "findingCount": len(set(findings)),
        "findingCodes": sorted(set(findings)),
        "exactRuntimeCredentialAbsent": "exact_runtime_credential" not in findings,
        "promptTextStored": False,
        "modelOutputTextStored": False,
    }


def _capture_finish_file_state() -> dict[str, tuple[int, int]]:
    state: dict[str, tuple[int, int]] = {}
    if not FINISH_ROOT.exists():
        return state
    for path in FINISH_ROOT.rglob("*"):
        try:
            if not path.is_file():
                continue
            stat = path.stat()
        except OSError:
            continue
        state[str(path.resolve())] = (int(stat.st_size), int(stat.st_mtime_ns))
    return state


def _runtime_parameter_counts_for_path(path: Path, *, provider: ProviderConfig) -> tuple[Counter[str], int, int]:
    needles = {
        "credential": provider.api_key.encode("utf-8"),
        "baseUrl": provider.base_url.encode("utf-8"),
        "model": provider.model.encode("utf-8"),
    }
    counts: Counter[str] = Counter()
    scanned_bytes = 0
    archive_part_count = 0

    def scan_bytes(payload: bytes) -> None:
        nonlocal scanned_bytes
        scanned_bytes += len(payload)
        for label, needle in needles.items():
            if needle:
                counts[label] += payload.count(needle)
        decoded = payload.decode("utf-8", errors="ignore")
        counts["genericSecret"] += sum(len(pattern.findall(decoded)) for _code, pattern in SECRET_PATTERNS)

    try:
        payload = path.read_bytes()
    except OSError:
        return counts, scanned_bytes, archive_part_count
    scan_bytes(str(path).encode("utf-8"))
    scan_bytes(payload)
    if path.suffix.lower() in {".docx", ".zip"}:
        try:
            with zipfile.ZipFile(path, "r") as archive:
                for member in archive.infolist():
                    if member.is_dir() or member.file_size > 32 * 1024 * 1024:
                        continue
                    archive_part_count += 1
                    scan_bytes(archive.read(member.filename))
        except (OSError, zipfile.BadZipFile, RuntimeError):
            pass
    return +counts, scanned_bytes, archive_part_count


def _capture_finish_runtime_hit_baseline(
    provider: ProviderConfig,
    *,
    paths: list[Path],
) -> dict[str, dict[str, int]]:
    baseline: dict[str, dict[str, int]] = {}
    for path in paths:
        try:
            normalized = path.resolve()
            if not normalized.is_file():
                continue
        except OSError:
            continue
        counts, _scanned_bytes, _archive_parts = _runtime_parameter_counts_for_path(normalized, provider=provider)
        if counts:
            baseline[str(normalized)] = dict(counts)
    return baseline


def _scan_new_or_changed_artifacts(
    before: dict[str, tuple[int, int]],
    *,
    provider: ProviderConfig,
    baseline_hits: dict[str, dict[str, int]],
) -> dict[str, Any]:
    """Scan this run's changed artifacts for plaintext runtime parameters."""

    after = _capture_finish_file_state()
    changed_paths = [
        Path(path)
        for path, identity in after.items()
        if before.get(path) != identity
    ]
    hit_counts: Counter[str] = Counter()
    preexisting_hit_counts: Counter[str] = Counter()
    scanned_bytes = 0
    archive_part_count = 0

    for path in changed_paths:
        current_counts, current_bytes, current_archive_parts = _runtime_parameter_counts_for_path(path, provider=provider)
        scanned_bytes += current_bytes
        archive_part_count += current_archive_parts
        previous_counts = baseline_hits.get(str(path.resolve()), {})
        for label in {"credential", "baseUrl", "model", "genericSecret"}:
            previous = int(previous_counts.get(label, 0) or 0)
            current = int(current_counts.get(label, 0) or 0)
            preexisting_hit_counts[label] += min(previous, current)
            hit_counts[label] += max(0, current - previous)

    new_hit_total = sum(value for value in hit_counts.values() if value > 0)
    result = {
        "ok": new_hit_total == 0,
        "changedArtifactFileCount": len(changed_paths),
        "scannedByteCount": scanned_bytes,
        "scannedArchivePartCount": archive_part_count,
        "exactCredentialHitCount": int(hit_counts.get("credential", 0)),
        "exactBaseUrlHitCount": int(hit_counts.get("baseUrl", 0)),
        "exactModelHitCount": int(hit_counts.get("model", 0)),
        "genericSecretPatternNewHitCount": int(hit_counts.get("genericSecret", 0)),
        "preexistingHitCountIgnored": sum(preexisting_hit_counts.values()),
        "runtimeParameterPlaintextPersisted": new_hit_total > 0,
        "matchedPathsStored": False,
    }
    return result


def _write_report_atomically(report_path: Path, report: dict[str, Any], *, api_key: str) -> None:
    scan = _secret_scan_report(report, api_key=api_key)
    report["security"] = {"preWriteSecretScan": scan}
    final_scan = _secret_scan_report(report, api_key=api_key)
    report["security"]["preWriteSecretScan"] = final_scan
    if not final_scan.get("ok"):
        raise E2EContractError("report_secret_scan_failed", "security_failure")
    payload = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
    temporary = report_path.with_name(f".{report_path.name}.{uuid.uuid4().hex}.tmp")
    try:
        temporary.write_text(payload, encoding="utf-8")
        temporary.replace(report_path)
    finally:
        temporary.unlink(missing_ok=True)


def _run_e2e_unlocked(
    *,
    sample_path: Path,
    source_url: str,
    source_commit: str,
    source_license: str,
    report_path: Path,
    export_path: Path,
    run_id: str,
    max_real_targets: int = DEFAULT_MAX_REAL_TARGETS,
    full_round: bool = False,
    execution_mode: str = "real_provider",
) -> dict[str, Any]:
    """Run the E2E and always return a prompt/output-free report object.

    ``execution_mode='offline_fake'`` exists solely for the local regression;
    the CLI never exposes it and always requires ``FYADR_RUN_REAL_LLM=1``.
    """

    if not 1 <= int(max_real_targets) <= MAX_REAL_TARGETS:
        raise E2EContractError("max_real_targets_out_of_range", "input_error")
    attribution = _validate_attribution(source_url, source_commit, source_license)
    normalized_report_path = _resolve_finish_path(report_path, label="report_path", suffix=".json")
    normalized_export_path = _resolve_finish_path(export_path, label="export_path", suffix=".docx")
    _require(normalized_report_path != normalized_export_path, "report_export_path_collision", "input_error")
    _require_fresh_output_paths(normalized_report_path, normalized_export_path)
    provider = _load_provider_from_environment(allow_offline_injection=execution_mode == "offline_fake")
    finish_state_before = _capture_finish_file_state()
    history_db_path = FINISH_ROOT / "fyadr_history.sqlite3"
    finish_runtime_hits_before = _capture_finish_runtime_hit_baseline(
        provider,
        paths=[
            normalized_report_path,
            normalized_export_path,
            history_db_path,
            Path(f"{history_db_path}-wal"),
            Path(f"{history_db_path}-shm"),
        ],
    )

    started = time.monotonic()
    report: dict[str, Any] = {
        "version": REPORT_VERSION,
        "ok": False,
        "skipped": False,
        "status": "running",
        "runId": run_id,
        "createdAt": _utc_now(),
        "completedAt": "",
        "durationMs": 0,
        "executionMode": execution_mode,
        "executionIsolation": {
            "realProviderCrossProcessLockRequired": execution_mode == "real_provider",
            "realProviderCrossProcessLockHeldForRun": execution_mode == "real_provider",
            "maxConcurrentRealProviderE2E": 1,
            "configuredRewriteConcurrency": 1,
            "lockFileStoresRuntimeMetadata": False,
            "lockPathStored": False,
        },
        "fullRoundRealProvider": bool(full_round),
        "maxRealTargets": int(max_real_targets),
        "sourceAttribution": attribution,
        "provider": provider.public_summary(),
        "failure": None,
    }
    work_source: Path | None = None
    real_call_auditor: CompletionCallAuditor | None = None
    try:
        phase = "copy_source"
        work_source, sample_sha256 = _copy_unique_work_source(sample_path, run_id)
        report["source"] = {
            "samplePath": str(sample_path.expanduser().resolve()),
            "sampleSha256": sample_sha256,
            "workSourcePath": str(work_source),
            "workSourceSha256": _sha256_file(work_source),
            "workSourceBytes": work_source.stat().st_size,
            "uniqueOriginCopy": True,
        }

        phase = "fresh_docx_preflight"
        preflight_contract, freeze, extracted_path, snapshot_path, authoritative_snapshot = _fresh_docx_preflight(work_source)
        report["preflight"] = {
            "contract": preflight_contract,
            "freeze": freeze,
            "extractedTextPath": str(extracted_path),
            "snapshotPath": str(snapshot_path),
        }
        _require(int(authoritative_snapshot.editable_unit_count) > 0, "docx_has_no_editable_body", "document_scope_failure")

        model_config = provider.neutral_app_model_config()
        _require(
            model_config.get("streaming") is True,
            "real_e2e_streaming_configuration_disabled",
            "configuration_error",
        )
        _require(
            int(model_config.get("rewriteConcurrency", 0) or 0) == 1,
            "real_e2e_rewrite_concurrency_not_single",
            "configuration_error",
        )
        baseline_progress = ProgressAudit()
        phase = "full_document_baseline"
        if full_round:
            baseline_mode = "full_real_round"
            baseline_auditor = CompletionCallAuditor(
                EnvironmentProviderTransport(app_service.llm_completion, provider),
                max_calls=max(
                    1,
                    int(authoritative_snapshot.editable_unit_count)
                    * E2E_MAX_COMPLETIONS_PER_TARGET,
                ),
            )
            with _patched_completion(baseline_auditor):
                round_result = app_service.run_round_for_app(
                    str(work_source),
                    model_config,
                    round_number=1,
                    progress_callback=baseline_progress,
                )
            baseline_provider_calls = baseline_auditor.public_summary()
            if execution_mode == "real_provider":
                _require(
                    baseline_provider_calls.get("allRealCallsRequestedStreaming") is True,
                    "full_round_real_provider_non_streaming_call_detected",
                )
                _require(
                    baseline_provider_calls.get("allSuccessfulCallsStreamCompleted") is True,
                    "full_round_stream_completion_evidence_incomplete",
                )
            baseline_execution = {
                "mode": "full_real_round",
                "providerCalls": baseline_provider_calls,
                "claim": "complete_round_used_real_provider",
            }
        else:
            baseline_mode = "identity"
            identity = IdentityCompletion()
            with _patched_completion(identity):
                round_result = app_service.run_round_for_app(
                    str(work_source),
                    model_config,
                    round_number=1,
                    progress_callback=baseline_progress,
                )
            baseline_execution = identity.public_summary()
        report["baselineMode"] = baseline_mode
        report["baselineExecution"] = {**baseline_execution, "progress": baseline_progress.public_summary()}
        report["round"] = _summarize_round_result(round_result, preflight_contract)
        output_path = Path(str(round_result["outputPath"])).resolve()

        phase = "initial_unified_snapshot"
        initial_snapshot = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
        _require(initial_snapshot.get("bodyMapSha256") is not None, "docx_snapshot_body_map_missing")
        _require(initial_snapshot.get("manifestSha256") is not None, "docx_snapshot_manifest_missing")
        _require(bool(initial_snapshot.get("rawOutputMatchesEffective")), "baseline_raw_output_cache_mismatch")
        _require(initial_snapshot.get("bodyMapMatchesEffective") is True, "baseline_body_map_cache_mismatch")
        initial_snapshot_summary = _summarize_snapshot(initial_snapshot)
        _require_v2_snapshot_candidate_evidence(initial_snapshot_summary)
        report["snapshots"] = {"initial": initial_snapshot_summary}

        phase = "initial_rate_audit"
        initial_audit = app_service.get_document_rate_audit(str(work_source), str(output_path))
        report["rateAudit"] = {"beforeRealTargets": _summarize_rate_audit(initial_audit)}

        phase = "representative_target_selection"
        targets = _select_representative_targets(initial_snapshot, initial_audit, max_targets=int(max_real_targets))
        internal_body_map = (initial_snapshot.get("_internal") or {}).get("bodyMapPayload")
        internal_units = internal_body_map.get("units") if isinstance(internal_body_map, dict) else []
        available_format_anchor_count = sum(
            len(unit.get("format_anchors", []))
            for unit in internal_units
            if isinstance(unit, dict) and isinstance(unit.get("format_anchors"), list)
        ) if isinstance(internal_units, list) else 0
        format_anchor_target_ids = {
            str(target.get("chunkId", "") or "")
            for target in targets
            if int(target.get("formatAnchorCount", 0) or 0) > 0
        }
        if available_format_anchor_count > 0:
            _require(bool(format_anchor_target_ids), "format_anchor_representative_target_missing")
        report["representativeTargets"] = {
            "count": len(targets),
            "availableFormatAnchorCount": available_format_anchor_count,
            "formatAnchorTargetCount": len(format_anchor_target_ids),
            "selectionDigest": _canonical_json_sha256(targets),
            "targets": targets,
            "storesInputText": False,
        }

        phase = "bounded_real_model_execution"
        real_call_auditor = CompletionCallAuditor(
            EnvironmentProviderTransport(app_service.llm_completion, provider),
            max_calls=int(max_real_targets) * E2E_MAX_COMPLETIONS_PER_TARGET,
        )
        with _patched_completion(real_call_auditor):
            (
                executions,
                real_candidate_ids,
                strategy_summary,
                provider_configuration_failure_ids,
                external_failure_ids,
                product_failure_ids,
            ) = _execute_bounded_real_targets(
                source_path=work_source,
                output_path=output_path,
                initial_snapshot=initial_snapshot,
                initial_audit=initial_audit,
                targets=targets,
                model_config=model_config,
                call_auditor=real_call_auditor,
            )
        real_call_summary = real_call_auditor.public_summary()
        format_anchor_execution_summary = _summarize_format_anchor_executions(
            executions,
            format_anchor_target_ids,
        )
        failed_attempt_evidence = [
            attempt
            for execution in executions
            for attempt in (
                execution.get("failedAttempts")
                if isinstance(execution.get("failedAttempts"), list)
                else []
            )
            if isinstance(attempt, dict)
        ]
        invalid_failed_attempt_evidence_count = sum(
            1
            for attempt in failed_attempt_evidence
            if attempt.get("safeEvidenceValid") is not True
        )
        report["realModel"] = {
            **real_call_summary,
            "targetCount": len(targets),
            "successfulCandidateCount": len(real_candidate_ids),
            **format_anchor_execution_summary,
            "providerConfigurationFailureTargetCount": len(provider_configuration_failure_ids),
            "externalFailureTargetCount": len(external_failure_ids),
            "productFailureTargetCount": len(product_failure_ids),
            "executions": executions,
            "boundStrategy": strategy_summary,
            "baselineIdentityCallsExcluded": not full_round,
            "failedAttemptEvidenceSchema": FAILED_ATTEMPT_EVIDENCE_SCHEMA,
            "failedAttemptEvidenceSchemaVersion": FAILED_ATTEMPT_EVIDENCE_VERSION,
            "failedAttemptEvidenceCount": len(failed_attempt_evidence),
            "invalidFailedAttemptEvidenceCount": invalid_failed_attempt_evidence_count,
            "failedAttemptTextStored": False,
            "failedAttemptErrorStored": False,
            "failedAttemptReasoningSuppressed": True,
        }
        _require(
            invalid_failed_attempt_evidence_count == 0,
            "failed_attempt_evidence_invalid",
            "security_failure",
        )
        if execution_mode == "real_provider":
            _require(int(report["realModel"]["callCount"]) > 0, "bounded_real_provider_made_zero_calls")
            _require(
                report["realModel"].get("allRealCallsRequestedStreaming") is True,
                "bounded_real_provider_non_streaming_call_detected",
            )
            _require(
                report["realModel"].get("allSuccessfulCallsStreamCompleted") is True,
                "bounded_stream_completion_evidence_incomplete",
            )

        phase = "post_model_rate_audit"
        post_model_snapshot = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
        post_model_audit = app_service.get_document_rate_audit(str(work_source), str(output_path))
        post_model_snapshot_summary = _summarize_snapshot(post_model_snapshot)
        _require_v2_snapshot_candidate_evidence(post_model_snapshot_summary)
        report["snapshots"]["postModelPreReview"] = post_model_snapshot_summary
        report["rateAudit"]["postModelPreReview"] = _summarize_rate_audit(post_model_audit)

        phase = "automated_review_cas"
        decisions, decision_summary = _build_review_decisions(
            post_model_snapshot,
            baseline_mode=baseline_mode,
            real_candidate_ids=real_candidate_ids,
            executions=executions,
        )
        save_result = app_service.save_review_decisions(
            str(output_path),
            decisions,
            expected_compare_revision=str(post_model_snapshot.get("compareRevision", "") or ""),
            require_compare_revision=True,
        )
        _require(bool(save_result.get("reviewLinkReady")), "review_cas_link_not_ready")
        _require(
            len(save_result.get("decisions", {}) if isinstance(save_result.get("decisions"), dict) else {}) == len(decisions),
            "review_cas_decision_count_mismatch",
        )
        reviewed_snapshot = app_service.read_round_artifact_snapshot(output_path, include_internal=True)
        _require(bool((reviewed_snapshot.get("review") or {}).get("reviewLinkReady")), "reviewed_snapshot_link_not_ready")
        _require(
            len((reviewed_snapshot.get("review") or {}).get("decisions", {})) == decision_summary["decisionCount"],
            "reviewed_snapshot_decision_count_mismatch",
        )
        automated_release = _validate_automated_review_release(reviewed_snapshot, decisions)
        report["review"] = {
            **decision_summary,
            **automated_release,
            "compareRevisionBefore": str(post_model_snapshot.get("compareRevision", "") or ""),
            "compareRevisionAfter": str(reviewed_snapshot.get("compareRevision", "") or ""),
            "contentRevisionAfter": str(reviewed_snapshot.get("contentRevision", "") or ""),
            "reviewRevisionAfter": str(reviewed_snapshot.get("reviewRevision", "") or ""),
            "reviewLinkReady": True,
            "casEnforced": True,
            # This is a deterministic E2E release decision, not a human review.
            # Keep the actor on the CAS evidence itself so report consumers
            # cannot mistake explicit-state syntax for human authorship.
            "reviewActor": "automated_e2e",
            "isHumanReview": False,
            "storesDecisionMap": False,
        }
        reviewed_snapshot_summary = _summarize_snapshot(reviewed_snapshot)
        _require_v2_snapshot_candidate_evidence(reviewed_snapshot_summary)
        report["snapshots"]["reviewMaterialized"] = reviewed_snapshot_summary

        phase = "post_review_rate_audit"
        final_audit = app_service.get_document_rate_audit(str(work_source), str(output_path))
        report["rateAudit"]["postReview"] = _summarize_rate_audit(final_audit)
        final_decision = str((final_audit.get("strategyPlan") or {}).get("decision", "") or "")
        report["convergence"] = {
            "strategyDecision": final_decision,
            "automaticStop": final_decision == "stop",
            "manualReviewRequired": final_decision in {"manual_review", "targeted_rerun", "blocked"}
            or any(
                isinstance(item.get("sameDimension"), dict)
                and item["sameDimension"].get("available") is True
                and item["sameDimension"].get("converged") is False
                for item in executions
            ),
            "sameDimensionEvaluatedCount": sum(
                1
                for item in executions
                if isinstance(item.get("sameDimension"), dict) and item["sameDimension"].get("available") is True
            ),
            "sameDimensionConvergedCount": sum(
                1
                for item in executions
                if isinstance(item.get("sameDimension"), dict) and item["sameDimension"].get("converged") is True
            ),
            "sameDimensionNonConvergedCount": sum(
                1
                for item in executions
                if isinstance(item.get("sameDimension"), dict) and item["sameDimension"].get("converged") is False
            ),
            "plateauPolicy": "stop_preserve_source_then_manual_review",
        }

        phase = "pre_export_candidate_release_gate"
        if provider_configuration_failure_ids:
            raise E2EContractError(
                "bounded_target_provider_configuration_failure",
                "provider_configuration_failure",
            )
        if product_failure_ids:
            raise E2EContractError("bounded_target_pipeline_failure")
        if external_failure_ids:
            raise E2EContractError("bounded_target_external_provider_failure", "external_unavailable")
        hard_model_failures = [
            item
            for item in executions
            if str(item.get("outcome", ""))
            in {
                "model_output_hard_validation_failed",
                "candidate_hard_validation_failed",
                "strategy_candidate_hard_validation_failed",
            }
        ]
        if hard_model_failures:
            raise E2EContractError("bounded_target_model_output_failure", "model_output_failure")
        release_rejections = [
            item
            for item in executions
            if str(item.get("outcome", ""))
            in {
                "candidate_stale_preserved_evidence",
                "candidate_selection_evidence_invalid",
                "candidate_readability_failed",
            }
        ]
        if release_rejections:
            raise E2EContractError("bounded_target_candidate_release_rejected", "model_output_failure")
        if (
            available_format_anchor_count > 0
            and int(report["realModel"].get("formatAnchorSafelyResolvedTargetCount", 0) or 0) <= 0
        ):
            raise E2EContractError("format_anchor_real_target_not_safely_resolved", "model_output_failure")
        phase = "certified_full_docx_export"
        report["export"] = _export_and_verify(
            output_path,
            normalized_export_path,
            preflight_contract,
            reviewed_snapshot,
        )
        _require(_sha256_file(work_source) == sample_sha256, "work_source_changed_during_e2e")

        phase = "runtime_parameter_artifact_scan"
        artifact_scan = _scan_new_or_changed_artifacts(
            finish_state_before,
            provider=provider,
            baseline_hits=finish_runtime_hits_before,
        )
        report["artifactRuntimeParameterScan"] = artifact_scan
        _require(bool(artifact_scan.get("ok")), "runtime_parameter_plaintext_persisted", "security_failure")

        manual_review = bool(report["convergence"]["manualReviewRequired"])
        changed_rewrite_count = int(
            (report.get("review") or {}).get("changedRewriteConfirmedCount", 0) or 0
        )
        report["review"]["baselineOnlyPass"] = changed_rewrite_count <= 0
        report["review"]["rewriteQuotaRequiredForPass"] = False
        report["ok"] = True
        report["status"] = _passed_e2e_status(
            manual_review=manual_review,
            changed_rewrite_count=changed_rewrite_count,
        )
    except Exception as exc:
        provider_configuration_failure = _provider_configuration_failure_for_error(exc)
        if isinstance(exc, E2EContractError):
            code = exc.code
            category = exc.category
        elif provider_configuration_failure is not None:
            code = "provider_configuration_failure"
            category = "provider_configuration_failure"
        elif _is_external_provider_failure(exc):
            code = "external_provider_unavailable"
            category = "external_unavailable"
        else:
            code = "unhandled_pipeline_exception"
            category = "product_pipeline_failure"
        report["ok"] = False
        report["skipped"] = False
        report["status"] = category
        report["failure"] = {
            "phase": locals().get("phase", "initialization"),
            "category": category,
            "code": code,
            "errorType": type(exc).__name__,
            "messageStored": False,
        }
        if category == "provider_configuration_failure":
            if provider_configuration_failure is None:
                executions = (report.get("realModel") or {}).get("executions")
                if isinstance(executions, list):
                    for execution in executions:
                        if not isinstance(execution, dict):
                            continue
                        candidate = execution.get("providerFailure")
                        if isinstance(candidate, dict):
                            provider_configuration_failure = _provider_configuration_failure_summary(candidate)
                        if provider_configuration_failure is not None:
                            break
            if provider_configuration_failure is not None:
                report["failure"]["providerFailure"] = provider_configuration_failure
    finally:
        report["completedAt"] = _utc_now()
        report["durationMs"] = round((time.monotonic() - started) * 1000)
        if work_source is not None and work_source.exists():
            report.setdefault("source", {})["workSourceFinalSha256"] = _sha256_file(work_source)
        if real_call_auditor is not None and "realModel" not in report:
            report["realModel"] = real_call_auditor.public_summary()
        _write_report_atomically(normalized_report_path, report, api_key=provider.api_key)
        report["reportPath"] = str(normalized_report_path)
        report["reportSha256"] = _sha256_file(normalized_report_path)
    return report


def run_e2e(
    *,
    sample_path: Path,
    source_url: str,
    source_commit: str,
    source_license: str,
    report_path: Path,
    export_path: Path,
    run_id: str,
    max_real_targets: int = DEFAULT_MAX_REAL_TARGETS,
    full_round: bool = False,
    execution_mode: str = "real_provider",
) -> dict[str, Any]:
    """Run one E2E, serializing real-provider executions across processes."""

    with _execution_mode_run_lock(execution_mode):
        return _run_e2e_unlocked(
            sample_path=sample_path,
            source_url=source_url,
            source_commit=source_commit,
            source_license=source_license,
            report_path=report_path,
            export_path=export_path,
            run_id=run_id,
            max_real_targets=max_real_targets,
            full_round=full_round,
            execution_mode=execution_mode,
        )


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Run an auditable real-university-thesis DOCX E2E with bounded real model calls."
    )
    parser.add_argument("--sample", type=Path, required=True, help="University thesis DOCX sample.")
    parser.add_argument("--source-url", required=True, help="Public source/repository URL.")
    parser.add_argument("--source-commit", required=True, help="Pinned source commit or release identifier.")
    parser.add_argument("--source-license", required=True, help="Source sample license identifier.")
    parser.add_argument("--report", type=Path, required=True, help="JSON report path under finish/.")
    parser.add_argument("--export", type=Path, required=True, help="DOCX export alias path under finish/.")
    parser.add_argument("--run-id", required=True, help="Stable audit run label (ASCII letters/digits/._-).")
    parser.add_argument(
        "--max-real-targets",
        type=int,
        default=DEFAULT_MAX_REAL_TARGETS,
        help=f"Maximum representative chunks sent to the real provider (default {DEFAULT_MAX_REAL_TARGETS}, max {MAX_REAL_TARGETS}).",
    )
    parser.add_argument(
        "--full-round",
        action="store_true",
        help="Explicitly send the complete round to the real provider; expensive and off by default.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        report = run_e2e(
            sample_path=args.sample,
            source_url=args.source_url,
            source_commit=args.source_commit,
            source_license=args.source_license,
            report_path=args.report,
            export_path=args.export,
            run_id=args.run_id,
            max_real_targets=args.max_real_targets,
            full_round=args.full_round,
            execution_mode="real_provider",
        )
    except E2EContractError as exc:
        print(
            json.dumps(
                {
                    "ok": False,
                    "skipped": False,
                    "status": exc.category,
                    "failure": {"code": exc.code, "messageStored": False},
                },
                ensure_ascii=False,
            )
        )
        return 2
    print(
        json.dumps(
            {
                "ok": bool(report.get("ok")),
                "skipped": False,
                "status": str(report.get("status", "") or ""),
                "reportPath": str(report.get("reportPath", "") or ""),
                "reportSha256": str(report.get("reportSha256", "") or ""),
                "realCallCount": int((report.get("realModel") or {}).get("callCount", 0) or 0),
                "certifiedArtifactPath": str((report.get("export") or {}).get("certifiedArtifactPath", "") or ""),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if report.get("ok") else 1


if __name__ == "__main__":
    raise SystemExit(main())
