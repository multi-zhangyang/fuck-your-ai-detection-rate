from __future__ import annotations

import argparse
from collections.abc import Iterator, Mapping
import hashlib
import json
import posixpath
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urlsplit
import xml.etree.ElementTree as ET
import zipfile

from docx import Document  # type: ignore[import]

from docx_pipeline import _load_docx_snapshot, _resolve_target_paragraph
from docx_security import (
    MAX_DOCX_ENTRY_UNCOMPRESSED_BYTES,
    MAX_DOCX_XML_PART_BYTES,
    validate_docx_package,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
REL = f"{{{REL_NS}}}"
EXACT_PART_NAMES = {
    "word/fontTable.xml",
    "word/webSettings.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/commentsIds.xml",
    "word/people.xml",
}
EXACT_PART_PREFIXES = (
    "customXml/item",
    "word/media/",
    "word/embeddings/",
    "word/activeX/",
    "word/charts/",
    "word/comments",
    "word/diagrams/",
    "word/glossary/",
    "word/tasks/",
    "word/theme/",
)
SEMANTIC_RANGE_MARKER_NAMES = {
    "bookmarkStart",
    "bookmarkEnd",
    "commentRangeStart",
    "commentRangeEnd",
}
SEMANTIC_POINT_REFERENCE_NAMES = {
    "commentReference",
    "footnoteReference",
    "endnoteReference",
}
SEMANTIC_MARKER_NAMES = SEMANTIC_RANGE_MARKER_NAMES | SEMANTIC_POINT_REFERENCE_NAMES
SEMANTIC_RANGE_MARKER_ACTIONS = {
    "bookmarkStart": ("bookmark", "start"),
    "bookmarkEnd": ("bookmark", "end"),
    "commentRangeStart": ("comment", "start"),
    "commentRangeEnd": ("comment", "end"),
}
SEMANTIC_BLOCK_LEVEL_BOOKMARK_PARENTS = frozenset({"body", "tc"})


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _append_source_generation_issues(
    issues: list[dict[str, Any]],
    *,
    source_sha256_before: str,
    source_sha256_after: str,
    expected_source_sha256: str,
) -> None:
    if expected_source_sha256 and (
        source_sha256_before != expected_source_sha256
        or source_sha256_after != expected_source_sha256
    ):
        issues.append(
            {
                "type": "source_anchor_sha256_mismatch",
                "expectedSourceSha256": expected_source_sha256,
                "sourceSha256Before": source_sha256_before,
                "sourceSha256After": source_sha256_after,
            }
        )
    if source_sha256_before != source_sha256_after:
        issues.append(
            {
                "type": "source_changed_during_audit",
                "sourceSha256Before": source_sha256_before,
                "sourceSha256After": source_sha256_after,
            }
        )


def _clone_without_text_nodes(node: ET.Element) -> ET.Element:
    """Clone OOXML while blanking only WordprocessingML ``w:t`` values.

    Keep the text node itself (and every non-``w:t`` value, such as field
    instructions or DrawingML ``a:t`` text).  Removing every element named
    ``t`` and failing to copy other element text made the old structural hash
    blind to text-node cardinality, field instructions, and non-Word drawing
    text.
    """

    clone = ET.Element(node.tag, dict(node.attrib))
    raw_text = node.text
    if node.tag == f"{W}t":
        clone.text = None
    # python-docx's lxml subclasses expose synthetic aggregate ``.text`` on
    # composite CT_P/CT_R nodes.  Only leaf-node character data is literal XML
    # content and belongs in this clone.
    elif len(node) == 0 and raw_text and raw_text.strip():
        clone.text = raw_text
    for child in node:
        clone.append(_clone_without_text_nodes(child))
    return clone


def _clone_paragraph_for_format_lock(
    node: ET.Element,
    ancestors: tuple[str, ...] = (),
    *,
    freeze_direct_text: bool = False,
) -> ET.Element:
    """Clone a paragraph, blanking only text nodes the rewrite path can edit.

    ``docx_pipeline`` updates direct paragraph runs and direct hyperlink runs.
    Text nested in content controls, custom XML, text boxes, tracked changes,
    field instructions, or DrawingML is never a rewrite target and therefore
    remains part of the strict fidelity signature.
    """

    clone = ET.Element(node.tag, dict(node.attrib))
    is_plain_run_text = (
        node.tag == f"{W}t"
        and len(ancestors) >= 2
        and ancestors[-1] == f"{W}r"
        and ancestors[-2] == f"{W}p"
    )
    is_hyperlink_run_text = (
        node.tag == f"{W}t"
        and len(ancestors) >= 3
        and ancestors[-1] == f"{W}r"
        and ancestors[-2] == f"{W}hyperlink"
        and ancestors[-3] == f"{W}p"
    )
    if (is_plain_run_text or is_hyperlink_run_text) and not freeze_direct_text:
        clone.text = None
        # xml:space follows the rewritten text's leading/trailing whitespace;
        # it is text semantics, not a formatting mutation.
        clone.attrib.pop("{http://www.w3.org/XML/1998/namespace}space", None)
    elif len(node) == 0 and node.text and node.text.strip():
        clone.text = node.text
    for child in node:
        clone.append(
            _clone_paragraph_for_format_lock(
                child,
                (*ancestors, node.tag),
                freeze_direct_text=freeze_direct_text,
            )
        )
    return clone


def _element_structure_hash(element: ET.Element, *, exclude_text_nodes: bool = False) -> str:
    normalized = _clone_without_text_nodes(element) if exclude_text_nodes else element
    return hashlib.sha256(ET.tostring(normalized, encoding="utf-8")).hexdigest()


def _xml_structure_hash(xml: bytes, *, exclude_text_nodes: bool = False) -> str:
    return _element_structure_hash(ET.fromstring(xml), exclude_text_nodes=exclude_text_nodes)


def _paragraph_format_signature(paragraph: ET.Element) -> str:
    """Sha256 signature of a paragraph's non-text OOXML, text-agnostic.

    Captures the complete paragraph tree, including pPr/rPr, hyperlink wrappers,
    bookmarks, tabs, breaks, drawings, fields, and nested containers. Only the
    values (plus text-dependent ``xml:space``) of direct run/hyperlink ``w:t``
    nodes are normally blanked; their count and every non-rewriteable text
    value remain locked. Paragraphs carrying zero-width bookmark/comment/note
    markers keep direct text in the signature as well, because those paragraphs
    are scope-frozen and marker-to-text attachment is semantic evidence.
    """

    freeze_direct_text = any(
        _local_name(node.tag) in SEMANTIC_MARKER_NAMES
        for node in paragraph.iter()
    )
    return hashlib.sha256(
        ET.tostring(
            _clone_paragraph_for_format_lock(
                paragraph,
                freeze_direct_text=freeze_direct_text,
            ),
            encoding="utf-8",
        )
    ).hexdigest()


def audit_docx_format_lock(
    export_path: Path,
    *,
    source_path: Path,
    snapshot_path: Path,
    expected_source_sha256: str | None = None,
    provenance_source_path: Path | None = None,
    report_path: Path | None = None,
) -> dict[str, Any]:
    """Assert fidelity-lock mode preserved every editable paragraph's format.

    For each editable unit in the snapshot, the exported paragraph's pPr/rPr
    signature must equal the source paragraph's signature. A mismatch means
    the export mutated format-bearing OOXML rather than only replacing text.
    """

    normalized_expected_source_sha256 = str(expected_source_sha256 or "").strip().lower()
    source_sha256_before = _sha256_file(source_path)
    issues: list[dict[str, Any]] = []
    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        raise ValueError(f"DOCX snapshot not found: {snapshot_path}")

    source_document = Document(str(source_path.resolve()))
    export_document = Document(str(export_path.resolve()))
    editable_indexes = {
        int(unit.target.get("paragraph_index"))
        for unit in snapshot.units
        if unit.editable
        and str(unit.target.get("kind", "")) == "paragraph"
    }
    editable_checked = 0
    protected_checked = 0

    if len(source_document.paragraphs) != len(export_document.paragraphs):
        issues.append(
            {
                "type": "format_lock_paragraph_count_changed",
                "expected": len(source_document.paragraphs),
                "actual": len(export_document.paragraphs),
            }
        )
    for paragraph_index, (source_paragraph, export_paragraph) in enumerate(
        zip(source_document.paragraphs, export_document.paragraphs)
    ):
        if paragraph_index in editable_indexes:
            editable_checked += 1
        else:
            protected_checked += 1
        source_sig = _paragraph_format_signature(source_paragraph._p)
        export_sig = _paragraph_format_signature(export_paragraph._p)
        if source_sig != export_sig:
            issues.append(
                {
                    "type": "format_lock_violation",
                    "paragraphIndex": paragraph_index,
                    "editable": paragraph_index in editable_indexes,
                    "sourceSignature": source_sig,
                    "exportSignature": export_sig,
                }
            )

    if len(source_document.tables) != len(export_document.tables):
        issues.append(
            {
                "type": "format_lock_table_count_changed",
                "expected": len(source_document.tables),
                "actual": len(export_document.tables),
            }
        )
    table_checked = 0
    for table_index, (source_table, export_table) in enumerate(zip(source_document.tables, export_document.tables)):
        table_checked += 1
        source_sig = _xml_structure_hash(source_table._tbl.xml.encode("utf-8"))
        export_sig = _xml_structure_hash(export_table._tbl.xml.encode("utf-8"))
        if source_sig != export_sig:
            issues.append(
                {
                    "type": "format_lock_table_changed",
                    "tableIndex": table_index,
                    "sourceSignature": source_sig,
                    "exportSignature": export_sig,
                }
            )

    source_sect_pr = source_document._element.body.sectPr
    export_sect_pr = export_document._element.body.sectPr
    source_section_sig = _xml_structure_hash(source_sect_pr.xml.encode("utf-8")) if source_sect_pr is not None else ""
    export_section_sig = _xml_structure_hash(export_sect_pr.xml.encode("utf-8")) if export_sect_pr is not None else ""
    if source_section_sig != export_section_sig:
        issues.append(
            {
                "type": "format_lock_section_changed",
                "sourceSignature": source_section_sig,
                "exportSignature": export_section_sig,
            }
        )

    source_sha256_after = _sha256_file(source_path)
    _append_source_generation_issues(
        issues,
        source_sha256_before=source_sha256_before,
        source_sha256_after=source_sha256_after,
        expected_source_sha256=normalized_expected_source_sha256,
    )

    report = {
        "ok": not issues,
        "sourcePath": str(source_path.resolve()),
        "sourceSha256": source_sha256_before,
        "expectedSourceSha256": normalized_expected_source_sha256,
        "sourceGenerationStable": source_sha256_before == source_sha256_after,
        "provenanceSourcePath": (
            str(provenance_source_path.resolve())
            if provenance_source_path is not None
            else ""
        ),
        "exportPath": str(export_path.resolve()),
        "snapshotPath": str(snapshot_path.resolve()),
        "editableChecked": editable_checked,
        "protectedChecked": protected_checked,
        "tableChecked": table_checked,
        "issueCount": len(issues),
        "issues": issues[:50],
        "truncatedIssues": max(0, len(issues) - 50),
    }
    if report_path is not None:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        report["reportPath"] = str(report_path.resolve())
    return report


def audit_docx_export(
    export_path: Path,
    *,
    source_path: Path,
    snapshot_path: Path,
    expected_source_sha256: str | None = None,
    provenance_source_path: Path | None = None,
    report_path: Path | None = None,
) -> dict[str, Any]:
    normalized_expected_source_sha256 = str(expected_source_sha256 or "").strip().lower()
    source_sha256_before = _sha256_file(source_path)
    issues: list[dict[str, Any]] = []
    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        raise ValueError(f"DOCX snapshot not found: {snapshot_path}")

    source_document = Document(str(source_path.resolve()))
    export_document = Document(str(export_path.resolve()))
    protected_checked = 0
    editable_checked = 0
    table_checked = 0
    audited_units = [*snapshot.units, *getattr(snapshot, "protected_structural_units", [])]
    semantic_range_unit_indexes = {
        int(unit.unit_index)
        for unit in audited_units
        if bool(getattr(unit, "inside_semantic_range", False))
    }

    for unit in audited_units:
        expected_text = _read_target_text(source_document, unit.target)
        actual_text = _read_target_text(export_document, unit.target)
        target_kind = str(unit.target.get("kind", ""))
        if target_kind == "table_cell_paragraph":
            table_checked += 1
        if unit.editable:
            editable_checked += 1
            continue
        protected_checked += 1
        if actual_text != expected_text:
            issue = {
                "type": "protected_text_changed",
                "unitIndex": unit.unit_index,
                "target": unit.target,
                "protectReason": unit.protect_reason,
                "structuralRole": getattr(unit, "structural_role", "unknown"),
            }
            if int(unit.unit_index) in semantic_range_unit_indexes:
                issue.update(_text_hash_evidence(expected_text, actual_text))
            else:
                issue.update({"expected": expected_text, "actual": actual_text})
            issues.append(issue)

    table_structure_issues = _audit_table_structure(source_document, export_document)
    issues.extend(table_structure_issues)

    source_sha256_after = _sha256_file(source_path)
    _append_source_generation_issues(
        issues,
        source_sha256_before=source_sha256_before,
        source_sha256_after=source_sha256_after,
        expected_source_sha256=normalized_expected_source_sha256,
    )

    report = {
        "ok": not issues,
        "sourcePath": str(source_path.resolve()),
        "sourceSha256": source_sha256_before,
        "expectedSourceSha256": normalized_expected_source_sha256,
        "sourceGenerationStable": source_sha256_before == source_sha256_after,
        "provenanceSourcePath": (
            str(provenance_source_path.resolve())
            if provenance_source_path is not None
            else ""
        ),
        "exportPath": str(export_path.resolve()),
        "snapshotPath": str(snapshot_path.resolve()),
        "protectedChecked": protected_checked,
        "editableChecked": editable_checked,
        "tableParagraphChecked": table_checked,
        "issueCount": len(issues),
        "issues": issues[:50],
        "truncatedIssues": max(0, len(issues) - 50),
    }
    if report_path is not None:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        report["reportPath"] = str(report_path.resolve())
    return report


def get_docx_audit_report_path(export_path: Path) -> Path:
    return export_path.with_suffix(".audit.json")


def get_docx_ooxml_audit_report_path(export_path: Path) -> Path:
    return export_path.with_suffix(".ooxml_audit.json")


def get_docx_format_lock_report_path(export_path: Path) -> Path:
    return export_path.with_suffix(".format_lock.json")


def audit_docx_ooxml_integrity(
    export_path: Path,
    *,
    source_path: Path,
    snapshot_path: Path,
    expected_source_sha256: str | None = None,
    provenance_source_path: Path | None = None,
    report_path: Path | None = None,
) -> dict[str, Any]:
    normalized_expected_source_sha256 = str(expected_source_sha256 or "").strip().lower()
    source_sha256_before = _sha256_file(source_path)
    issues: list[dict[str, Any]] = []
    snapshot = _load_docx_snapshot(snapshot_path)
    if snapshot is None:
        raise ValueError(f"DOCX snapshot not found: {snapshot_path}")

    with _read_docx_parts(source_path) as source_parts, _read_docx_parts(export_path) as export_parts:
        if "word/document.xml" not in source_parts or "word/document.xml" not in export_parts:
            issues.append({"type": "document_xml_missing"})
        else:
            issues.extend(_audit_document_xml(source_parts["word/document.xml"], export_parts["word/document.xml"], snapshot))

        issues.extend(_audit_part_inventory(source_parts, export_parts))
        issues.extend(_audit_exact_parts(source_parts, export_parts))
        issues.extend(_audit_header_footer_parts(source_parts, export_parts))
        issues.extend(_audit_styles_part(source_parts, export_parts))
        issues.extend(_audit_numbering_part(source_parts, export_parts))
        issues.extend(_audit_settings_part(source_parts, export_parts))
        issues.extend(_audit_relationship_parts(source_parts, export_parts))
        issues.extend(_audit_relationship_target_integrity(source_parts, export_parts))
        checked_parts = len(source_parts)

    source_sha256_after = _sha256_file(source_path)
    _append_source_generation_issues(
        issues,
        source_sha256_before=source_sha256_before,
        source_sha256_after=source_sha256_after,
        expected_source_sha256=normalized_expected_source_sha256,
    )

    report = {
        "ok": not issues,
        "sourcePath": str(source_path.resolve()),
        "sourceSha256": source_sha256_before,
        "expectedSourceSha256": normalized_expected_source_sha256,
        "sourceGenerationStable": source_sha256_before == source_sha256_after,
        "provenanceSourcePath": (
            str(provenance_source_path.resolve())
            if provenance_source_path is not None
            else ""
        ),
        "exportPath": str(export_path.resolve()),
        "snapshotPath": str(snapshot_path.resolve()),
        "issueCount": len(issues),
        "issues": issues[:50],
        "truncatedIssues": max(0, len(issues) - 50),
        "checkedParts": checked_parts,
    }
    if report_path is not None:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        report["reportPath"] = str(report_path.resolve())
    return report


def _read_target_text(document: Any, target: dict[str, Any]) -> str:
    return _resolve_target_paragraph(document, target).text


def _audit_table_structure(source_document: Any, export_document: Any) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    if len(source_document.tables) != len(export_document.tables):
        issues.append(
            {
                "type": "table_count_changed",
                "expected": len(source_document.tables),
                "actual": len(export_document.tables),
            }
        )
        return issues

    for table_index, source_table in enumerate(source_document.tables):
        export_table = export_document.tables[table_index]
        if len(source_table.rows) != len(export_table.rows):
            issues.append(
                {
                    "type": "table_row_count_changed",
                    "tableIndex": table_index,
                    "expected": len(source_table.rows),
                    "actual": len(export_table.rows),
                }
            )
            continue
        for row_index, source_row in enumerate(source_table.rows):
            export_row = export_table.rows[row_index]
            if len(source_row.cells) != len(export_row.cells):
                issues.append(
                    {
                        "type": "table_cell_count_changed",
                        "tableIndex": table_index,
                        "rowIndex": row_index,
                        "expected": len(source_row.cells),
                        "actual": len(export_row.cells),
                    }
                )
                continue
            for cell_index, source_cell in enumerate(source_row.cells):
                export_cell = export_row.cells[cell_index]
                source_text = "\n".join(paragraph.text for paragraph in source_cell.paragraphs)
                export_text = "\n".join(paragraph.text for paragraph in export_cell.paragraphs)
                if source_text != export_text:
                    issues.append(
                        {
                            "type": "table_text_changed",
                            "tableIndex": table_index,
                            "rowIndex": row_index,
                            "cellIndex": cell_index,
                            "expected": source_text,
                            "actual": export_text,
                        }
                    )
    return issues


class _BoundedDocxParts(Mapping[str, bytes]):
    """Lazy package mapping: XML is read on demand; binary hashes stream."""

    def __init__(self, path: Path) -> None:
        self.path = path.resolve()
        validate_docx_package(self.path)
        self.archive = zipfile.ZipFile(self.path, "r")
        self.infos = {
            item.filename: item
            for item in self.archive.infolist()
            if not item.is_dir()
        }

    def __len__(self) -> int:
        return len(self.infos)

    def __iter__(self) -> Iterator[str]:
        return iter(self.infos)

    def __getitem__(self, name: str) -> bytes:
        info = self.infos[name]
        limit = (
            MAX_DOCX_XML_PART_BYTES
            if name == "[Content_Types].xml" or name.endswith((".xml", ".rels"))
            else MAX_DOCX_ENTRY_UNCOMPRESSED_BYTES
        )
        if info.file_size > limit:
            raise ValueError(f"DOCX audit part exceeds its read bound: {name}")
        with self.archive.open(info, "r") as handle:
            payload = handle.read(limit + 1)
        if len(payload) > limit or len(payload) != info.file_size:
            raise ValueError(f"DOCX audit part size is inconsistent: {name}")
        return payload

    def sha256(self, name: str) -> str:
        info = self.infos[name]
        digest = hashlib.sha256()
        size = 0
        with self.archive.open(info, "r") as handle:
            for block in iter(lambda: handle.read(1024 * 1024), b""):
                size += len(block)
                if size > MAX_DOCX_ENTRY_UNCOMPRESSED_BYTES:
                    raise ValueError(f"DOCX audit part exceeds its hash bound: {name}")
                digest.update(block)
        if size != info.file_size:
            raise ValueError(f"DOCX audit part size is inconsistent: {name}")
        return digest.hexdigest()

    def close(self) -> None:
        self.archive.close()

    def __enter__(self) -> "_BoundedDocxParts":
        return self

    def __exit__(self, _exc_type: object, _exc: object, _traceback: object) -> None:
        self.close()


def _read_docx_parts(path: Path) -> _BoundedDocxParts:
    return _BoundedDocxParts(path)


def _docx_part_sha256(parts: Mapping[str, bytes], name: str) -> str:
    if isinstance(parts, _BoundedDocxParts):
        return parts.sha256(name)
    return hashlib.sha256(parts[name]).hexdigest()


def _audit_part_inventory(source_parts: Mapping[str, bytes], export_parts: Mapping[str, bytes]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    source_names = set(source_parts)
    export_names = set(export_parts)
    for name in sorted(source_names - export_names):
        issues.append({"type": "package_part_missing", "part": name})
    for name in sorted(export_names - source_names):
        issues.append({"type": "package_part_added", "part": name})
    return issues


def _audit_exact_parts(source_parts: Mapping[str, bytes], export_parts: Mapping[str, bytes]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    source_critical = {name for name in source_parts if _is_exact_part(name)}
    export_critical = {name for name in export_parts if _is_exact_part(name)}
    for name in sorted(source_critical - export_critical):
        issues.append({"type": "ooxml_part_missing", "part": name})
    for name in sorted(export_critical - source_critical):
        issues.append({"type": "ooxml_part_added", "part": name})
    for name in sorted(source_critical & export_critical):
        source_hash = _docx_part_sha256(source_parts, name)
        export_hash = _docx_part_sha256(export_parts, name)
        if source_hash != export_hash:
            issues.append(
                {
                    "type": "ooxml_part_changed",
                    "part": name,
                    "sourceHash": source_hash,
                    "exportHash": export_hash,
                }
            )
    return issues


def _is_exact_part(name: str) -> bool:
    return name in EXACT_PART_NAMES or any(name.startswith(prefix) for prefix in EXACT_PART_PREFIXES)


def _audit_header_footer_parts(source_parts: dict[str, bytes], export_parts: dict[str, bytes]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    source_names = _part_names_with_prefixes(source_parts, ("word/header", "word/footer"))
    export_names = _part_names_with_prefixes(export_parts, ("word/header", "word/footer"))
    for name in sorted(source_names - export_names):
        issues.append({"type": "header_footer_part_missing", "part": name})
    for name in sorted(export_names - source_names):
        issues.append({"type": "header_footer_part_added", "part": name})
    for name in sorted(source_names & export_names):
        source_signature = _header_footer_signature(source_parts[name])
        export_signature = _header_footer_signature(export_parts[name])
        if source_signature != export_signature:
            issues.append(
                {
                    "type": "header_footer_semantics_changed",
                    "part": name,
                    "expected": source_signature,
                    "actual": export_signature,
                }
            )
    return issues


def _audit_styles_part(source_parts: dict[str, bytes], export_parts: dict[str, bytes]) -> list[dict[str, Any]]:
    return _audit_optional_xml_signature(
        source_parts,
        export_parts,
        "word/styles.xml",
        "styles_identity_changed",
        _styles_signature,
    )


def _audit_numbering_part(source_parts: dict[str, bytes], export_parts: dict[str, bytes]) -> list[dict[str, Any]]:
    return _audit_optional_xml_signature(
        source_parts,
        export_parts,
        "word/numbering.xml",
        "numbering_identity_changed",
        _numbering_signature,
    )


def _audit_settings_part(source_parts: dict[str, bytes], export_parts: dict[str, bytes]) -> list[dict[str, Any]]:
    return _audit_optional_xml_signature(
        source_parts,
        export_parts,
        "word/settings.xml",
        "settings_identity_changed",
        _settings_signature,
    )


def _audit_optional_xml_signature(
    source_parts: dict[str, bytes],
    export_parts: dict[str, bytes],
    name: str,
    issue_type: str,
    signature_fn: Any,
) -> list[dict[str, Any]]:
    if name not in source_parts and name not in export_parts:
        return []
    if name not in source_parts:
        return [{"type": "ooxml_part_added", "part": name}]
    if name not in export_parts:
        return [{"type": "ooxml_part_missing", "part": name}]
    source_signature = signature_fn(source_parts[name])
    export_signature = signature_fn(export_parts[name])
    if source_signature == export_signature:
        return []
    return [
        {
            "type": issue_type,
            "part": name,
            "expected": source_signature,
            "actual": export_signature,
        }
    ]


def _part_names_with_prefixes(parts: dict[str, bytes], prefixes: tuple[str, ...]) -> set[str]:
    return {name for name in parts if any(name.startswith(prefix) and name.endswith(".xml") for prefix in prefixes)}


def _header_footer_signature(xml: bytes) -> dict[str, Any]:
    root = ET.fromstring(xml)
    semantic_markers = [
        signature
        for paragraph in root.findall(f".//{W}p")
        if (signature := _semantic_marker_attachment_signature(paragraph))
    ]
    return {
        "text": _element_text(root),
        "structureHash": _element_structure_hash(root, exclude_text_nodes=True),
        "fieldCount": len(root.findall(f".//{W}fldChar")),
        "instrText": sorted(" ".join(str(node.text or "").split()) for node in root.findall(f".//{W}instrText")),
        "drawingCount": len(root.findall(f".//{W}drawing")),
        "pictCount": len(root.findall(f".//{W}pict")),
        "tableSignatures": [_table_signature(table) for table in root.findall(f".//{W}tbl")],
        "semanticMarkerAttachments": semantic_markers,
    }


def _styles_signature(xml: bytes) -> dict[str, Any]:
    root = ET.fromstring(xml)
    styles: list[tuple[str, str, str, str]] = []
    for style in root.findall(f"{W}style"):
        styles.append(
            (
                _attr(style, "styleId"),
                _attr(style, "type"),
                _attr(style, "default"),
                _child_attr(style, "name", "val"),
            )
        )
    return {"styles": sorted(styles), "structureHash": _element_structure_hash(root)}


def _numbering_signature(xml: bytes) -> dict[str, Any]:
    root = ET.fromstring(xml)
    abstract_numbers: list[dict[str, Any]] = []
    for abstract in root.findall(f"{W}abstractNum"):
        levels: list[tuple[str, str, str, str]] = []
        for level in abstract.findall(f"{W}lvl"):
            levels.append(
                (
                    _attr(level, "ilvl"),
                    _child_attr(level, "numFmt", "val"),
                    _child_attr(level, "lvlText", "val"),
                    _child_attr(level, "pStyle", "val"),
                )
            )
        abstract_numbers.append({"id": _attr(abstract, "abstractNumId"), "levels": sorted(levels)})
    nums: list[tuple[str, str]] = []
    for num in root.findall(f"{W}num"):
        nums.append((_attr(num, "numId"), _child_attr(num, "abstractNumId", "val")))
    return {
        "abstractNums": sorted(abstract_numbers, key=lambda item: str(item.get("id", ""))),
        "nums": sorted(nums),
        "structureHash": _element_structure_hash(root),
    }


def _settings_signature(xml: bytes) -> dict[str, Any]:
    root = ET.fromstring(xml)
    return {
        "footnotePr": _xml_child_names(root.find(f"{W}footnotePr")),
        "endnotePr": _xml_child_names(root.find(f"{W}endnotePr")),
        "docGrid": _element_attrs(root.find(f"{W}docGrid")),
        "defaultTabStop": _element_attrs(root.find(f"{W}defaultTabStop")),
        "evenAndOddHeaders": root.find(f"{W}evenAndOddHeaders") is not None,
        "mirrorMargins": root.find(f"{W}mirrorMargins") is not None,
        "structureHash": _element_structure_hash(root),
    }


def _audit_relationship_parts(source_parts: dict[str, bytes], export_parts: dict[str, bytes]) -> list[dict[str, Any]]:
    # Relationships outside ``word/_rels`` can still carry document semantics
    # (package roots, custom XML bindings, embedded packages, and metadata).
    # Audit the complete OPC relationship inventory, not only document.xml's
    # immediate relationships.
    source_names = {name for name in source_parts if name.endswith(".rels")}
    export_names = {name for name in export_parts if name.endswith(".rels")}
    issues: list[dict[str, Any]] = []
    for name in sorted(source_names - export_names):
        issues.append({"type": "relationship_part_missing", "part": name})
    for name in sorted(export_names - source_names):
        issues.append({"type": "relationship_part_added", "part": name})
    for name in sorted(source_names & export_names):
        source_signature = _relationships_signature(source_parts[name])
        export_signature = _relationships_signature(export_parts[name])
        if source_signature != export_signature:
            issues.append(
                {
                    "type": "relationship_semantics_changed",
                    "part": name,
                    "expected": source_signature,
                    "actual": export_signature,
                }
            )
    return issues


def _relationships_signature(xml: bytes) -> list[tuple[str, str, str, str]]:
    root = ET.fromstring(xml)
    return sorted(
        (
            str(relationship.attrib.get("Id", "")),
            str(relationship.attrib.get("Type", "")),
            str(relationship.attrib.get("Target", "")),
            str(relationship.attrib.get("TargetMode", "")),
        )
        for relationship in root.findall(f"{REL}Relationship")
    )


def _audit_relationship_target_integrity(
    source_parts: dict[str, bytes],
    export_parts: dict[str, bytes],
) -> list[dict[str, Any]]:
    source_missing = {
        (item["part"], item["relationshipId"], item["target"], item["resolvedTarget"])
        for item in _missing_internal_relationship_targets(source_parts)
    }
    issues: list[dict[str, Any]] = []
    for item in _missing_internal_relationship_targets(export_parts):
        key = (item["part"], item["relationshipId"], item["target"], item["resolvedTarget"])
        if key in source_missing:
            continue
        issues.append({"type": "relationship_target_missing", **item})
    return issues


def _missing_internal_relationship_targets(parts: dict[str, bytes]) -> list[dict[str, str]]:
    missing: list[dict[str, str]] = []
    for rels_name in sorted(name for name in parts if name.endswith(".rels")):
        try:
            root = ET.fromstring(parts[rels_name])
        except ET.ParseError:
            continue
        base_dir = _relationship_base_dir(rels_name)
        for relationship in root.findall(f"{REL}Relationship"):
            if str(relationship.attrib.get("TargetMode", "")).casefold() == "external":
                continue
            raw_target = str(relationship.attrib.get("Target", "") or "")
            resolved_target = _resolve_relationship_target(base_dir, raw_target)
            if not resolved_target or resolved_target in parts:
                continue
            missing.append(
                {
                    "part": rels_name,
                    "relationshipId": str(relationship.attrib.get("Id", "") or ""),
                    "target": raw_target,
                    "resolvedTarget": resolved_target,
                }
            )
    return missing


def _relationship_base_dir(rels_name: str) -> str:
    if rels_name == "_rels/.rels":
        return ""
    rels_dir = posixpath.dirname(rels_name)
    if posixpath.basename(rels_dir) != "_rels":
        return posixpath.dirname(rels_name)
    return posixpath.dirname(rels_dir)


def _resolve_relationship_target(base_dir: str, raw_target: str) -> str:
    target_path = unquote(urlsplit(raw_target).path or "")
    if not target_path:
        return ""
    if target_path.startswith("/"):
        return posixpath.normpath(target_path.lstrip("/"))
    return posixpath.normpath(posixpath.join(base_dir, target_path))


def _xml_child_names(element: ET.Element | None) -> list[str]:
    if element is None:
        return []
    return [_local_name(child.tag) for child in list(element)]


def _element_attrs(element: ET.Element | None) -> dict[str, str]:
    if element is None:
        return {}
    return {_local_name(key): value for key, value in sorted(element.attrib.items())}


def _child_attr(element: ET.Element, child_name: str, attr_name: str) -> str:
    child = element.find(f"{W}{child_name}")
    return _attr(child, attr_name) if child is not None else ""


def _attr(element: ET.Element | None, name: str) -> str:
    if element is None:
        return ""
    return str(element.attrib.get(f"{W}{name}", ""))


def _semantic_document_range_evidence(root: ET.Element) -> dict[str, Any]:
    """Pair body bookmark/comment ranges and hash their enclosed text.

    Evidence intentionally excludes document and comment prose.  Endpoints are
    described by structural paragraph ordinals plus local visible-text offsets;
    the selected content is represented only by length and SHA-256.  Absolute
    document offsets are used internally for slicing but are not reported,
    because legitimate edits before a frozen range may shift them.
    """

    body = root.find(f"{W}body")
    if body is None:
        return {
            "topologyValid": False,
            "issueCount": 1,
            "issueCodes": ["semantic_range_document_body_missing"],
            "rangeCount": 0,
            "ranges": [],
        }

    canonical_parts: list[str] = []
    canonical_length = 0
    paragraph_ordinal = -1
    active: dict[tuple[str, str], dict[str, Any]] = {}
    seen_starts: set[tuple[str, str]] = set()
    seen_ends: set[tuple[str, str]] = set()
    completed: list[dict[str, Any]] = []
    issues: list[str] = []

    def append_visible(value: str, paragraph_state: dict[str, int] | None) -> None:
        nonlocal canonical_length
        if not value:
            return
        canonical_parts.append(value)
        canonical_length += len(value)
        if paragraph_state is not None:
            paragraph_state["offset"] += len(value)

    def marker_position(paragraph_state: dict[str, int] | None) -> dict[str, int]:
        if paragraph_state is None:
            # A block-level bookmark endpoint sits between paragraphs. Bind it
            # to the last visited paragraph ordinal and use -1 for the
            # non-paragraph offset so endpoint movement remains observable
            # without serializing document prose.
            return {"paragraphOrdinal": int(paragraph_ordinal), "offset": -1}
        return {
            "paragraphOrdinal": int(paragraph_state["ordinal"]),
            "offset": int(paragraph_state["offset"]),
        }

    def visit(
        node: ET.Element,
        paragraph_state: dict[str, int] | None,
        parent_local_name: str = "",
    ) -> None:
        nonlocal paragraph_ordinal

        local_name = _local_name(node.tag)
        current_state = paragraph_state
        if node.tag == f"{W}p":
            paragraph_ordinal += 1
            current_state = {"ordinal": paragraph_ordinal, "offset": 0}

        marker = SEMANTIC_RANGE_MARKER_ACTIONS.get(local_name)
        if marker is not None:
            kind, action = marker
            marker_id = str(node.attrib.get(f"{W}id", "") or "")
            if current_state is None and not (
                kind == "bookmark"
                and parent_local_name in SEMANTIC_BLOCK_LEVEL_BOOKMARK_PARENTS
            ):
                issues.append("semantic_range_marker_outside_paragraph")
            if not marker_id:
                issues.append("semantic_range_missing_id")
            else:
                key = (kind, marker_id)
                if action == "start":
                    if key in seen_starts or key in active:
                        issues.append("semantic_range_duplicate_start")
                    elif key in seen_ends:
                        issues.append("semantic_range_reversed")
                    else:
                        seen_starts.add(key)
                        active[key] = {
                            "kind": kind,
                            "id": marker_id,
                            "start": marker_position(current_state),
                            "contentStart": canonical_length,
                        }
                elif key in seen_ends:
                    issues.append("semantic_range_duplicate_end")
                else:
                    seen_ends.add(key)
                    start = active.pop(key, None)
                    if start is None:
                        issues.append("semantic_range_reversed")
                    else:
                        completed.append(
                            {
                                **start,
                                "end": marker_position(current_state),
                                "contentEnd": canonical_length,
                            }
                        )

        if node.tag == f"{W}t":
            append_visible(str(node.text or ""), current_state)
        elif local_name == "tab":
            append_visible("\t", current_state)
        elif local_name in {"br", "cr"}:
            append_visible("\n", current_state)
        else:
            for child in list(node):
                visit(child, current_state, local_name)

        if node.tag == f"{W}p":
            # A paragraph separator makes cross-paragraph range hashes sensitive
            # to content moving between paragraphs without exposing that content.
            append_visible("\u2029", None)

    visit(body, None, "")
    if active:
        issues.extend("semantic_range_unmatched_start" for _key in active)

    canonical_text = "".join(canonical_parts)
    ranges: list[dict[str, Any]] = []
    for raw_range in completed:
        content_start = max(0, min(int(raw_range["contentStart"]), len(canonical_text)))
        content_end = max(content_start, min(int(raw_range["contentEnd"]), len(canonical_text)))
        selected = canonical_text[content_start:content_end]
        ranges.append(
            {
                "kind": str(raw_range["kind"]),
                "id": str(raw_range["id"]),
                "start": dict(raw_range["start"]),
                "end": dict(raw_range["end"]),
                "contentLength": len(selected),
                "contentSha256": hashlib.sha256(selected.encode("utf-8")).hexdigest(),
            }
        )
    ranges.sort(
        key=lambda item: (
            str(item["kind"]),
            str(item["id"]),
            int(item["start"]["paragraphOrdinal"]),
            int(item["start"]["offset"]),
        )
    )
    return {
        "topologyValid": not issues,
        "issueCount": len(issues),
        "issueCodes": sorted(set(issues)),
        "rangeCount": len(ranges),
        "ranges": ranges,
    }


def _audit_semantic_document_ranges(
    source_root: ET.Element,
    export_root: ET.Element,
) -> list[dict[str, Any]]:
    source = _semantic_document_range_evidence(source_root)
    exported = _semantic_document_range_evidence(export_root)
    issues: list[dict[str, Any]] = []
    if not bool(source.get("topologyValid")):
        issues.append(
            {
                "type": "source_semantic_range_topology_invalid",
                "issueCount": int(source.get("issueCount", 0) or 0),
                "issueCodes": list(source.get("issueCodes", [])),
            }
        )
    if not bool(exported.get("topologyValid")):
        issues.append(
            {
                "type": "semantic_range_topology_invalid",
                "issueCount": int(exported.get("issueCount", 0) or 0),
                "issueCodes": list(exported.get("issueCodes", [])),
            }
        )

    source_ranges = {
        (str(item.get("kind", "")), str(item.get("id", ""))): item
        for item in source.get("ranges", [])
        if isinstance(item, dict)
    }
    export_ranges = {
        (str(item.get("kind", "")), str(item.get("id", ""))): item
        for item in exported.get("ranges", [])
        if isinstance(item, dict)
    }
    for key in sorted(set(source_ranges) | set(export_ranges)):
        expected = source_ranges.get(key)
        actual = export_ranges.get(key)
        kind, marker_id = key
        if expected is None or actual is None:
            issues.append(
                {
                    "type": "semantic_range_topology_changed",
                    "kind": kind,
                    "id": marker_id,
                    "expectedPresent": expected is not None,
                    "actualPresent": actual is not None,
                }
            )
            continue
        if expected.get("start") != actual.get("start") or expected.get("end") != actual.get("end"):
            issues.append(
                {
                    "type": "semantic_range_topology_changed",
                    "kind": kind,
                    "id": marker_id,
                    "expectedStart": expected.get("start"),
                    "actualStart": actual.get("start"),
                    "expectedEnd": expected.get("end"),
                    "actualEnd": actual.get("end"),
                }
            )
            continue
        # Bookmark boundaries are navigation topology. Marker paragraphs are
        # protected separately, but marker-free cross-paragraph interior prose
        # may change without moving either endpoint. Comment ranges denote the
        # reviewed selection itself and therefore retain the strict content
        # hash lock.
        if kind == "comment" and (
            int(expected.get("contentLength", -1)) != int(actual.get("contentLength", -1))
            or str(expected.get("contentSha256", "")) != str(actual.get("contentSha256", ""))
        ):
            issues.append(
                {
                    "type": "semantic_range_content_changed",
                    "kind": kind,
                    "id": marker_id,
                    "expectedLength": int(expected.get("contentLength", 0) or 0),
                    "actualLength": int(actual.get("contentLength", 0) or 0),
                    "expectedSha256": str(expected.get("contentSha256", "")),
                    "actualSha256": str(actual.get("contentSha256", "")),
                }
            )
    return issues


def _audit_document_xml(source_xml: bytes, export_xml: bytes, snapshot: Any) -> list[dict[str, Any]]:
    source_root = ET.fromstring(source_xml)
    export_root = ET.fromstring(export_xml)
    source_body = source_root.find(f"{W}body")
    export_body = export_root.find(f"{W}body")
    issues: list[dict[str, Any]] = []
    if source_body is None or export_body is None:
        return [{"type": "document_body_missing"}]

    issues.extend(_audit_semantic_document_ranges(source_root, export_root))

    source_blocks = [_local_name(child.tag) for child in list(source_body)]
    export_blocks = [_local_name(child.tag) for child in list(export_body)]
    if source_blocks != export_blocks:
        issues.append(
            {
                "type": "body_block_sequence_changed",
                "expected": source_blocks[:80],
                "actual": export_blocks[:80],
            }
        )

    # Direct body containers not represented by python-docx's
    # ``Document.paragraphs``/``tables`` collections (for example top-level
    # content controls, customXml, tracked-change containers, or altChunk) are
    # protected wholesale. ``sectPr`` is excluded from this structural pass
    # because the stricter format-lock audit verifies it byte-for-byte.
    for block_index, (source_block, export_block) in enumerate(zip(list(source_body), list(export_body))):
        if source_block.tag in {f"{W}p", f"{W}tbl", f"{W}sectPr"}:
            continue
        source_signature = _element_structure_hash(source_block)
        export_signature = _element_structure_hash(export_block)
        if source_signature != export_signature:
            issues.append(
                {
                    "type": "protected_body_container_changed",
                    "blockIndex": block_index,
                    "blockType": _local_name(source_block.tag),
                    "expected": source_signature,
                    "actual": export_signature,
                }
            )

    editable_paragraph_indexes = {
        int(unit.target.get("paragraph_index"))
        for unit in snapshot.units
        if bool(getattr(unit, "editable", False))
        and isinstance(getattr(unit, "target", None), dict)
        and str(unit.target.get("kind", "")) == "paragraph"
    }
    semantic_range_paragraph_indexes = {
        int(unit.target.get("paragraph_index"))
        for unit in snapshot.units
        if bool(getattr(unit, "inside_semantic_range", False))
        and isinstance(getattr(unit, "target", None), dict)
        and str(unit.target.get("kind", "")) == "paragraph"
    }
    source_paragraphs = [child for child in list(source_body) if child.tag == f"{W}p"]
    export_paragraphs = [child for child in list(export_body) if child.tag == f"{W}p"]
    if len(source_paragraphs) != len(export_paragraphs):
        issues.append({"type": "paragraph_count_changed", "expected": len(source_paragraphs), "actual": len(export_paragraphs)})
    for index, (source_paragraph, export_paragraph) in enumerate(zip(source_paragraphs, export_paragraphs)):
        source_text = _element_text(source_paragraph)
        export_text = _element_text(export_paragraph)
        source_marker_signature = _semantic_marker_attachment_signature(source_paragraph)
        export_marker_signature = _semantic_marker_attachment_signature(export_paragraph)
        if source_marker_signature != export_marker_signature:
            issues.append(
                {
                    "type": "semantic_marker_attachment_changed",
                    "paragraphIndex": index,
                    "expected": source_marker_signature,
                    "actual": export_marker_signature,
                }
            )
        if index not in editable_paragraph_indexes and source_text != export_text:
            issue = {
                "type": "protected_paragraph_text_changed",
                "paragraphIndex": index,
            }
            if index in semantic_range_paragraph_indexes:
                issue.update(_text_hash_evidence(source_text, export_text))
            else:
                issue.update(
                    {
                        "expected": _sample(source_text),
                        "actual": _sample(export_text),
                    }
                )
            issues.append(issue)
        source_signature = _paragraph_nontext_signature(source_paragraph)
        export_signature = _paragraph_nontext_signature(export_paragraph)
        if source_signature != export_signature and index not in editable_paragraph_indexes:
            issues.append(
                {
                    "type": "protected_paragraph_structure_changed",
                    "paragraphIndex": index,
                    "expected": source_signature,
                    "actual": export_signature,
                }
            )

    source_tables = [child for child in list(source_body) if child.tag == f"{W}tbl"]
    export_tables = [child for child in list(export_body) if child.tag == f"{W}tbl"]
    if len(source_tables) != len(export_tables):
        issues.append({"type": "table_count_changed", "expected": len(source_tables), "actual": len(export_tables)})
    for index, (source_table, export_table) in enumerate(zip(source_tables, export_tables)):
        source_signature = _table_signature(source_table)
        export_signature = _table_signature(export_table)
        if source_signature != export_signature:
            issues.append(
                {
                    "type": "table_structure_or_text_changed",
                    "tableIndex": index,
                    "expected": source_signature,
                    "actual": export_signature,
                }
            )

    for code, tag in (("field_code_count_changed", "fldChar"), ("field_instr_count_changed", "instrText"), ("drawing_count_changed", "drawing")):
        source_count = len(source_root.findall(f".//{W}{tag}"))
        export_count = len(export_root.findall(f".//{W}{tag}"))
        if source_count != export_count:
            issues.append({"type": code, "expected": source_count, "actual": export_count})
    return issues


def _paragraph_nontext_signature(paragraph: ET.Element) -> dict[str, Any]:
    return {
        "fieldCount": len(paragraph.findall(f".//{W}fldChar")),
        "instrCount": len(paragraph.findall(f".//{W}instrText")),
        "drawingCount": len(paragraph.findall(f".//{W}drawing")),
        "pictCount": len(paragraph.findall(f".//{W}pict")),
        "bookmarkStartCount": len(paragraph.findall(f".//{W}bookmarkStart")),
        "bookmarkEndCount": len(paragraph.findall(f".//{W}bookmarkEnd")),
        "commentRangeStartCount": len(paragraph.findall(f".//{W}commentRangeStart")),
        "commentRangeEndCount": len(paragraph.findall(f".//{W}commentRangeEnd")),
        "commentReferenceCount": len(paragraph.findall(f".//{W}commentReference")),
        "footnoteReferenceCount": len(paragraph.findall(f".//{W}footnoteReference")),
        "endnoteReferenceCount": len(paragraph.findall(f".//{W}endnoteReference")),
    }


def _semantic_marker_attachment_signature(paragraph: ET.Element) -> dict[str, Any]:
    """Bind zero-width semantic marker IDs to reproducible text offsets.

    The signature deliberately stores no paragraph/comment text.  Offsets are
    measured against deterministic Word text nodes (plus tab/break characters)
    in XML document order, which detects unchanged-node-count attacks that move
    the prose enclosed by a bookmark/comment range or preceding a point note
    reference.
    """

    text_offset = 0
    events: list[dict[str, Any]] = []
    for node in paragraph.iter():
        local_name = _local_name(node.tag)
        if local_name == "t":
            text_offset += len(str(node.text or ""))
            continue
        if local_name in {"tab", "br", "cr"}:
            text_offset += 1
            continue
        if local_name not in SEMANTIC_MARKER_NAMES:
            continue
        event = {
            "type": local_name,
            "id": str(node.attrib.get(f"{W}id", "")),
            "offset": text_offset,
        }
        events.append(event)
    if not events:
        return {}
    return {
        "textLength": text_offset,
        "eventCount": len(events),
        "events": events,
    }


def _table_signature(table: ET.Element) -> dict[str, Any]:
    rows: list[list[list[str]]] = []
    for row in table.findall(f"{W}tr"):
        row_cells: list[list[str]] = []
        for cell in row.findall(f"{W}tc"):
            row_cells.append([_element_text(paragraph) for paragraph in cell.findall(f"{W}p")])
        rows.append(row_cells)
    return {
        "rowCount": len(rows),
        "cellCounts": [len(row) for row in rows],
        "text": rows,
    }


def _element_text(element: ET.Element) -> str:
    return "".join(node.text or "" for node in element.iter(f"{W}t"))


def _text_hash_evidence(expected: str, actual: str) -> dict[str, Any]:
    """Return non-content-leaking evidence for a protected text mismatch."""

    expected_value = str(expected)
    actual_value = str(actual)
    return {
        "expectedLength": len(expected_value),
        "actualLength": len(actual_value),
        "expectedSha256": hashlib.sha256(expected_value.encode("utf-8")).hexdigest(),
        "actualSha256": hashlib.sha256(actual_value.encode("utf-8")).hexdigest(),
    }


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _sample(text: str, limit: int = 160) -> str:
    compact = " ".join(str(text or "").split())
    return compact[:limit] + ("..." if len(compact) > limit else "")


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="Audit exported DOCX against original protected regions")
    parser.add_argument("source", type=Path)
    parser.add_argument("snapshot", type=Path)
    parser.add_argument("export", type=Path)
    parser.add_argument("--report", type=Path)
    parser.add_argument("--ooxml", action="store_true")
    args = parser.parse_args(argv)
    audit_fn = audit_docx_ooxml_integrity if args.ooxml else audit_docx_export
    print(
        json.dumps(
            audit_fn(
                args.export,
                source_path=args.source,
                snapshot_path=args.snapshot,
                report_path=args.report,
            ),
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
