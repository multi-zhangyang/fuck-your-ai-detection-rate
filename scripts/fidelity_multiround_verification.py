#!/usr/bin/env python3
"""Real multi-round end-to-end fidelity verification.

Uses the complex docx, runs a REAL two-round rewrite chain (round 2 feeds off
round 1's output via the project's own round-linking), then exports round 2 in
preserve_original mode and independently (lxml, not trusting audit functions)
asserts:
  - editable paragraph count is stable across rounds (no drift),
  - every editable paragraph's pPr+rPr signature still byte-matches the SOURCE
    docx (fidelity lock holds after 2 rounds, not just 1),
  - round 2 text differs from round 1 text (the chain actually progressed),
  - protected content (title/refs/acknowledgement/table) text is unchanged.
This catches regressions that a single-round toy test would miss: e.g. a
second-round body-map re-resolution that breaks paragraph targeting.
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
from pathlib import Path
from typing import Any

ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from fidelity_real_verification import (  # noqa: E402
    _create_complex_sample,
    _collect_editable_signatures,
)
import app_config as ac  # noqa: E402
import app_service  # noqa: E402
from docx import Document  # noqa: E402
from docx_pipeline import ensure_docx_processing_assets, get_docx_snapshot_path  # noqa: E402
from fyadr_round_service import get_round_compare_path  # noqa: E402
from round_helper import run_document_round  # noqa: E402

os.environ.setdefault("FYADR_APP_CONFIG_DIR", tempfile.mkdtemp(prefix="fyadr_multiround_"))


ROUND1_SOURCE_TEXT = (
    "首先，实验组记录检测误差。其次，实验组记录训练时间。"
    "再次，实验组记录内存占用。此外，实验组记录推理延迟。"
    "最后，实验组记录准确率变化。综上，实验组记录各项结果。"
)
ROUND1_REWRITTEN_TEXT = (
    "首先，实验组记录检测误差。其次，训练时间由日志同步记录。"
    "再次，内存占用纳入采样项，以核对不同批次的资源波动。"
    "此外，测试结束后实验组补记推理延迟。最后，准确率变化仍按相同指标计算。"
    "综上，各项结果汇入对照表用于逐项复核。"
)
# Round 2 clears the mechanical connector burst while deliberately distributing
# its six sentences across plain, article, passive, in-context, and domain
# openings; otherwise the fidelity fixture itself would trigger the production
# document-level repeated-opening gate instead of testing multi-round format.
ROUND2_REWRITTEN_TEXT = (
    "检测误差由实验组记录。本研究中，训练时间由日志同步记录。"
    "内存占用被纳入采样项，以核对不同批次的资源波动。"
    "在测试结束阶段，实验组补记推理延迟。准确率变化得以按相同指标计算。"
    "数据汇入对照表，供后续逐项复核。"
)
ROUND1_CHAIN_MARKER = "训练时间由日志同步记录"
ROUND2_CHAIN_MARKER = "供后续逐项复核"


def _create_multiround_sample(path: Path) -> Path:
    """Put a measurable two-dimension risk in an unformatted body run.

    The source has six near-equal, repeated-frame sentences plus six mechanical
    connectors.  Round 1 can therefore improve sentence rhythm without
    removing the connector risk that round 2 is responsible for.  The target
    deliberately uses the inherited single-run paragraph: splitting this text
    over the mixed-format five-run paragraph would create immutable format
    anchors inside the exact transform target.  The untouched five-run body
    remains in the sample so final export still verifies its pPr/rPr fidelity.
    """

    _create_complex_sample(path)
    document = Document(str(path))
    mixed_run_body = next((p for p in document.paragraphs if "本节研究" in p.text), None)
    if mixed_run_body is None or len(mixed_run_body.runs) != 5:
        raise AssertionError("complex fixture lost its five-run format-fidelity body paragraph")
    target = next((p for p in document.paragraphs if "完全依赖 Normal 样式继承" in p.text), None)
    if target is None or len(target.runs) != 1 or target.runs[0]._element.rPr is not None:
        raise AssertionError("multi-round target must remain a single inherited run without direct rPr")
    target.runs[0].text = ROUND1_SOURCE_TEXT
    if target.text != ROUND1_SOURCE_TEXT:
        raise AssertionError("multi-round source fixture text did not assemble deterministically")
    document.save(str(path))
    return path


def _transform_round1(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
    text = str(chunk_text)
    return text.replace(ROUND1_SOURCE_TEXT, ROUND1_REWRITTEN_TEXT) or chunk_text


def _transform_round2(chunk_text: str, prompt_input: str, round_number: int, chunk_id: str) -> str:
    text = str(chunk_text)
    return text.replace(ROUND1_REWRITTEN_TEXT, ROUND2_REWRITTEN_TEXT) or chunk_text


def _load_published_marker_chunk(output_path: Path, marker: str) -> tuple[dict[str, Any] | None, list[dict[str, Any]]]:
    compare_path = get_round_compare_path(output_path)
    payload = json.loads(compare_path.read_text(encoding="utf-8"))
    chunks = [item for item in payload.get("chunks", []) if isinstance(item, dict)]
    matches = [
        item
        for item in chunks
        if marker in str(item.get("outputText", ""))
        and isinstance(item.get("candidateSelection"), dict)
        and item["candidateSelection"].get("publishedRewrite") is True
    ]
    return (matches[0] if len(matches) == 1 else None), matches


def _selected_direction(chunk: dict[str, Any]) -> dict[str, Any]:
    selection = chunk.get("candidateSelection")
    if not isinstance(selection, dict):
        return {}
    selected_id = str(selection.get("selectedCandidateId", ""))
    candidates = selection.get("candidates")
    if not isinstance(candidates, list):
        return {}
    for candidate in candidates:
        if isinstance(candidate, dict) and str(candidate.get("candidateId", "")) == selected_id:
            direction = candidate.get("sameDimensionDirection")
            return direction if isinstance(direction, dict) else {}
    return {}


def _target_format_anchors(round_result: dict[str, Any], marker: str) -> list[dict[str, Any]] | None:
    body_map_path = Path(str(round_result.get("body_map_path", "")))
    payload = json.loads(body_map_path.read_text(encoding="utf-8"))
    units = payload.get("units")
    if not isinstance(units, list):
        return None
    matches = [
        unit
        for unit in units
        if isinstance(unit, dict) and marker in str(unit.get("current_text", ""))
    ]
    if len(matches) != 1:
        return None
    anchors = matches[0].get("format_anchors")
    return anchors if isinstance(anchors, list) else None


def run() -> dict[str, Any]:
    failures: list[str] = []
    work = ROOT_DIR / "finish" / "regression" / "fidelity_multiround"
    work.mkdir(parents=True, exist_ok=True)
    sample = work / "complex_source.docx"
    _create_multiround_sample(sample)
    ac.save_app_config({**ac.load_app_config(), "formatMode": "preserve_original"})

    snapshot_path = get_docx_snapshot_path(sample)
    # Materialize/refresh the snapshot before inspecting it.  Reading the
    # stem-based intermediate path directly made this regression depend on a
    # previous test having created the right ``complex_source`` fixture.
    ensure_docx_processing_assets(sample, snapshot_path=snapshot_path)
    source_sigs = _collect_editable_signatures(sample, snapshot_path)
    source_doc = Document(str(sample))
    protected_baseline = {
        i: p.text
        for i, p in enumerate(source_doc.paragraphs)
        if any(k in p.text for k in ["复杂格式真实样本论文", "参考文献", "致 谢", "图 1 检测结果对比"])
    }

    # Round 1
    r1 = run_document_round(sample, _transform_round1, round_number=1, prompt_profile="cn")
    out1 = Path(str(r1["output_path"]))
    target_anchors = _target_format_anchors(r1, "实验组记录检测误差")
    if target_anchors is None:
        failures.append("round 1 body map did not contain exactly one measurable target unit")
    elif target_anchors:
        failures.append(f"round 1 target unexpectedly acquired {len(target_anchors)} immutable format anchor(s)")
    else:
        print("  [fixture boundary] progressing target has no immutable format anchors.")
    round1_chunk, round1_matches = _load_published_marker_chunk(out1, ROUND1_CHAIN_MARKER)
    if round1_chunk is None:
        failures.append(
            "round 1 did not publish exactly one measurable sentence-rhythm rewrite "
            f"(published marker matches={len(round1_matches)})"
        )
    else:
        direction = _selected_direction(round1_chunk)
        if direction.get("dimensionId") != "sentence_structure" or direction.get("ok") is not True:
            failures.append(f"round 1 selected candidate lacks converged sentence-structure evidence: {direction}")
        elif float(direction.get("variationAfter", 0) or 0) <= float(direction.get("variationBefore", 0) or 0):
            failures.append(f"round 1 sentence-length variation did not measurably improve: {direction}")
        else:
            print(
                "  [round 1 selection] publishedRewrite=true; sentence-length variation "
                f"{direction.get('variationBefore')} -> {direction.get('variationAfter')}."
            )
    # Round 2 (chains off round 1 output via the project's round linking)
    if round1_chunk is not None:
        app_service.save_review_decisions(
            str(out1),
            {str(round1_chunk.get("chunkId", "")): "rewrite_confirmed"},
        )
    parent_snapshot = app_service.read_round_artifact_snapshot(out1, include_internal=True)
    effective_parent_text = str(parent_snapshot.get("_internal", {}).get("effectiveText", ""))
    if ROUND1_CHAIN_MARKER not in effective_parent_text:
        failures.append("round 1 rewrite was not materialized into the immutable parent snapshot")
    r2 = run_document_round(
        sample,
        _transform_round2,
        round_number=2,
        prompt_profile="cn",
        parent_artifact_snapshot=parent_snapshot,
    )
    out2 = Path(str(r2["output_path"]))
    round2_chunk, round2_matches = _load_published_marker_chunk(out2, ROUND2_CHAIN_MARKER)
    if round2_chunk is None:
        failures.append(
            "round 2 did not publish exactly one measurable connector rewrite "
            f"(published marker matches={len(round2_matches)})"
        )
    else:
        direction = _selected_direction(round2_chunk)
        round2_input = str(round2_chunk.get("inputText", ""))
        round2_output = str(round2_chunk.get("outputText", ""))
        if ROUND1_CHAIN_MARKER not in round2_input or ROUND1_CHAIN_MARKER not in round2_output:
            failures.append("round 2 did not consume and retain the round 1 chain marker")
        if direction.get("dimensionId") != "connector_detail" or direction.get("ok") is not True:
            failures.append(f"round 2 selected candidate lacks converged connector evidence: {direction}")
        elif not direction.get("riskCodesBefore") or direction.get("riskCodesAfter"):
            failures.append(f"round 2 connector risks were not measurably cleared: {direction}")
        else:
            print(
                "  [round 2 selection] publishedRewrite=true; connector density "
                f"{direction.get('before')} -> {direction.get('after')}."
            )

    t1 = out1.read_text(encoding="utf-8")
    t2 = out2.read_text(encoding="utf-8")
    if t1 == t2:
        failures.append("round 2 output identical to round 1 — chain did not progress")
    else:
        print("  [round chain] round 2 differs from round 1 — progressed.")

    # Export round 2 in preserve mode
    export = work / "complex_export_r2_preserve.docx"
    res = app_service.export_round_output(str(out2), str(export), "docx")
    if str(res.get("formatMode", "")) != "preserve_original":
        failures.append(f"export formatMode={res.get('formatMode')!r}, expected preserve_original")

    # Paragraph count stability: editable units must equal source.
    exp_sigs = _collect_editable_signatures(export, snapshot_path)
    if len(exp_sigs) != len(source_sigs):
        failures.append(f"editable count drifted: source={len(source_sigs)} export={len(exp_sigs)}")
    else:
        print(f"  [paragraph count] stable across 2 rounds: {len(source_sigs)} editable paragraphs.")

    # Fidelity lock after 2 rounds: every editable signature byte-matches SOURCE.
    mism = 0
    for src, exp in zip(source_sigs, exp_sigs):
        if src["ppr_sig"] != exp["ppr_sig"] or src["rpr_sigs"] != exp["rpr_sigs"]:
            mism += 1
            failures.append(f"fidelity lock broke after 2 rounds at unit {src['unit_index']}: ppr {src['ppr_sig'][:8]} vs {exp['ppr_sig'][:8]}")
    if mism == 0:
        print(f"  [fidelity lock] all {len(source_sigs)} editable signatures byte-match SOURCE after 2 rounds.")

    # Protected content unchanged after 2 rounds + export.
    exp_doc = Document(str(export))
    exp_texts = [p.text for p in exp_doc.paragraphs]
    for i, baseline_text in protected_baseline.items():
        if i < len(exp_texts) and exp_texts[i] != baseline_text:
            failures.append(f"protected paragraph {i} changed: {baseline_text[:20]!r} -> {exp_texts[i][:20]!r}")
    if not any("protected" in f for f in failures):
        print(f"  [protected content] {len(protected_baseline)} protected paragraphs unchanged.")

    # The positive markers prove that round 2 consumed round 1 rather than
    # independently rewriting the source, and that export retained both edits.
    final_body = "\n".join(p.text for p in exp_doc.paragraphs)
    missing_markers = [
        marker
        for marker in (ROUND1_CHAIN_MARKER, ROUND2_CHAIN_MARKER)
        if marker not in final_body
    ]
    if missing_markers:
        failures.append(f"export body text lost rewritten markers across 2 rounds: {missing_markers}")
    elif ROUND1_SOURCE_TEXT in final_body or ROUND1_REWRITTEN_TEXT in final_body:
        failures.append("export body retained an obsolete pre-round-2 target text")
    else:
        print("  [round chain markers] round 1 and round 2 markers both survived DOCX export.")

    return {"ok": not failures, "failures": failures, "editableCount": len(source_sigs)}


if __name__ == "__main__":
    r = run()
    if r["ok"]:
        print(f"\nMULTI-ROUND FIDELITY: PASS (editable={r['editableCount']})")
        sys.exit(0)
    print("\nMULTI-ROUND FIDELITY: FAIL")
    for f in r["failures"]:
        print(f"  - {f}")
    sys.exit(1)
