#!/usr/bin/env python3
"""Protect Normal-style numbered subheadings while keeping real list prose editable."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document  # type: ignore[import]
from docx.enum.style import WD_STYLE_TYPE  # type: ignore[import]
from docx.oxml import OxmlElement  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]

ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx_audit import audit_docx_format_lock  # noqa: E402
from docx_bodymap import build_docx_body_map, update_docx_body_map_texts  # noqa: E402
from docx_pipeline import (  # noqa: E402
    DOCX_SNAPSHOT_VERSION,
    ensure_docx_processing_assets,
    rebuild_docx_from_body_map_units,
)
from fyadr_round_service import (  # noqa: E402
    _extract_leading_structure_marker,
    protect_structure_tokens,
    restore_structure_tokens,
    validate_chunk_output,
)


STRUCTURE_LABELS = (
    "（1）基于STM32F407微控制器的控制系统设计",
    "（2）数字PID增量型控制算法",
    "（3）测试实验",
    "（4）单电机双闭环测试",
)
BODY_LIST_ITEMS = (
    "（1）系统驱动总功率提高；",
    "（2）单个电机成本降低。",
)
AUTO_NUMBERED_LABEL = "误差分析与系统设计"
AUTO_NUMBERED_BODY = "控制器读取当前状态并核对采样周期，随后将计算结果写入实验记录。"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _set_word_numbering(paragraph, *, level: int = 0, num_id: int = 1) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering_properties = OxmlElement("w:numPr")
    numbering_level = OxmlElement("w:ilvl")
    numbering_level.set(qn("w:val"), str(level))
    numbering_id = OxmlElement("w:numId")
    numbering_id.set(qn("w:val"), str(num_id))
    numbering_properties.append(numbering_level)
    numbering_properties.append(numbering_id)
    paragraph_properties.append(numbering_properties)


def _create_source(path: Path) -> None:
    document = Document()
    custom_label_style = document.styles.add_style("Thesis Minor Label", WD_STYLE_TYPE.PARAGRAPH)
    custom_label_style.base_style = document.styles["Normal"]
    document.add_heading("编号结构冻结测试论文", 0)
    document.add_paragraph("摘要")
    document.add_paragraph("本文围绕控制系统设计展开分析，并记录实验条件与结果。")
    document.add_paragraph("1 引言")
    document.add_paragraph(STRUCTURE_LABELS[0])
    document.add_paragraph("控制系统采用既定硬件结构，并按照设计顺序完成模块连接。")
    document.add_paragraph(BODY_LIST_ITEMS[0])
    document.add_paragraph(BODY_LIST_ITEMS[1])
    document.add_paragraph(STRUCTURE_LABELS[1], style=custom_label_style)
    document.add_paragraph("增量型算法根据当前偏差计算控制量，并保留原有参数定义。")
    document.add_paragraph(STRUCTURE_LABELS[2])
    document.add_paragraph("测试过程记录输入信号、响应时间和稳定状态。")
    document.add_paragraph(STRUCTURE_LABELS[3])
    document.add_paragraph("闭环测试按照相同采样周期执行，并核对两组结果。")
    auto_numbered_label = document.add_paragraph(AUTO_NUMBERED_LABEL, style=custom_label_style)
    _set_word_numbering(auto_numbered_label)
    auto_numbered_body = document.add_paragraph(AUTO_NUMBERED_BODY)
    _set_word_numbering(auto_numbered_body)
    document.add_paragraph("致谢")
    document.add_paragraph("感谢参与测试的教师与同学。")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 测试文献，2026.")
    document.save(path)


def main() -> int:
    work_dir = ROOT_DIR / "finish" / "regression" / "docx_numbered_structure_boundary"
    shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)
    source_path = work_dir / "numbered_structure_source.docx"
    export_path = work_dir / "numbered_structure_export.docx"
    _create_source(source_path)

    marker_text = "（1）测试实验"
    _assert(
        _extract_leading_structure_marker(marker_text) == "（1）",
        "canonical leading-marker grammar did not recognize full-width Arabic numbering",
    )
    protected_marker = protect_structure_tokens(marker_text)
    _assert(
        list(protected_marker.tokens.values()) == ["（1）"],
        "full-width parenthesized numbering was split instead of protected atomically",
    )
    _assert(
        restore_structure_tokens(protected_marker.text, protected_marker.tokens) == marker_text,
        "atomic full-width numbering token did not restore exactly",
    )
    try:
        validate_chunk_output(marker_text, "测试实验（1）", "numbered-marker-boundary")
    except ValueError as exc:
        _assert("leading numbering marker" in str(exc), f"wrong marker validation failure: {exc}")
    else:
        raise AssertionError("moving a full-width leading numbering marker was not rejected")

    extracted_path, snapshot_path, snapshot = ensure_docx_processing_assets(source_path)
    _assert(
        snapshot.version == DOCX_SNAPSHOT_VERSION and DOCX_SNAPSHOT_VERSION >= 15,
        "numbered-heading fix did not invalidate stale snapshots",
    )
    by_text = {unit.text: unit for unit in snapshot.units}
    for label in STRUCTURE_LABELS:
        unit = by_text.get(label)
        _assert(unit is not None, f"missing numbered structure label: {label}")
        _assert(unit.editable is False, f"Normal-style numbered subheading entered model scope: {label}")
        _assert(unit.protect_reason == "heading", f"numbered subheading has the wrong protection reason: {label}")
    for item in BODY_LIST_ITEMS:
        unit = by_text.get(item)
        _assert(unit is not None and unit.editable is True, f"punctuated body list item was over-frozen: {item}")
    auto_label_unit = by_text.get(AUTO_NUMBERED_LABEL)
    _assert(
        auto_label_unit is not None
        and auto_label_unit.has_numbering
        and auto_label_unit.style_name == "Thesis Minor Label"
        and auto_label_unit.editable is False
        and auto_label_unit.protect_reason == "heading",
        "custom-style w:numPr short label entered model scope",
    )
    auto_body_unit = by_text.get(AUTO_NUMBERED_BODY)
    _assert(
        auto_body_unit is not None and auto_body_unit.has_numbering and auto_body_unit.editable is True,
        "punctuated long w:numPr body prose was over-frozen",
    )

    model_input = extracted_path.read_text(encoding="utf-8")
    for label in STRUCTURE_LABELS:
        _assert(label not in model_input, f"numbered structure label leaked into model input: {label}")
    for item in BODY_LIST_ITEMS:
        _assert(item in model_input, f"real body list prose disappeared from model input: {item}")
    _assert(AUTO_NUMBERED_LABEL not in model_input, "custom-style w:numPr short label leaked into model input")
    _assert(AUTO_NUMBERED_BODY in model_input, "long w:numPr body prose disappeared from model input")

    body_map = build_docx_body_map(source_path, snapshot_path=snapshot_path, round_number=1)
    rewritten = [text.replace("按照", "依照").replace("记录", "保存") for text in body_map.current_texts()]
    updated = update_docx_body_map_texts(body_map, rewritten, round_number=1)
    rebuild_docx_from_body_map_units(
        updated.units,
        source_path=source_path,
        export_path=export_path,
        preserve_format=True,
    )
    exported_texts = [paragraph.text for paragraph in Document(str(export_path)).paragraphs]
    for label in STRUCTURE_LABELS:
        _assert(label in exported_texts, f"protected numbered subheading changed during export: {label}")
    _assert(AUTO_NUMBERED_LABEL in exported_texts, "protected custom-style w:numPr label changed during export")
    format_report = audit_docx_format_lock(
        export_path,
        source_path=source_path,
        snapshot_path=snapshot_path,
    )
    _assert(format_report.get("ok") is True, f"format lock failed after numbered-boundary export: {format_report}")

    print("DOCX numbered structure boundary regression passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
