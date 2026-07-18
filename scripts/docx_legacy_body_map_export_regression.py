from __future__ import annotations

import json
import sys
from pathlib import Path

try:
    from docx import Document  # type: ignore[import]
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency python-docx. Install it with: pip install python-docx") from exc

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR / "scripts") not in sys.path:
    sys.path.insert(0, str(ROOT_DIR / "scripts"))

import app_service  # noqa: E402
from docx_bodymap import build_docx_body_map, save_docx_body_map, update_docx_body_map_texts  # noqa: E402
from docx_pipeline import get_docx_snapshot_path  # noqa: E402
from fyadr_round_service import (  # noqa: E402
    _build_candidate_selection_event,
    _evaluate_rewrite_candidate,
    _select_rewrite_candidate,
)


WORK_DIR = ROOT_DIR / "finish" / "regression" / "legacy_body_map_export"
SOURCE_PATH = WORK_DIR / "legacy_scope_source.docx"
BODY_MAP_PATH = WORK_DIR / "round1_body_map.json"
OUTPUT_PATH = WORK_DIR / "round1.txt"
COMPARE_PATH = WORK_DIR / "round1_compare.json"
EXPORT_PATH = WORK_DIR / "round1_export.docx"
REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_legacy_body_map_export_regression_report.json"

SOURCE_REWRITES = (
    (
        "首先，系统读取用户提交的论文段落并建立任务记录。其次，服务按照段落边界生成改写单元并保留原有编号。"
        "此外，校验模块核对术语、数值和引用标记。因此，只有通过事实与格式检查的文本才会进入结果页，"
        "审阅者仍需确认每处修改。",
        "系统读取用户提交的论文段落并建立任务记录，服务再按照段落边界生成改写单元，同时保留原有编号。"
        "校验模块负责核对术语、数值和引用标记；只有文本通过事实与格式检查后才进入结果页，审阅者仍需确认每处修改。",
    ),
    (
        "首先，任务队列读取待处理记录并检查文档标识。其次，调度模块按照提交顺序分配执行单元。"
        "此外，异常处理模块保存失败原因和重试状态。因此，管理员可以依据完整日志核对每项任务的处理过程。",
        "任务队列读取待处理记录并检查文档标识，调度模块随后按照提交顺序分配执行单元。"
        "异常处理模块保存失败原因和重试状态，管理员可依据完整日志核对每项任务的处理过程。",
    ),
    (
        "首先，本研究得到导师在方案设计方面的指导。其次，同学协助完成实验环境核验。"
        "此外，学院提供了必要设备。因此，本文谨向上述人员和机构表示感谢。",
        "导师为本研究的方案设计提供指导，同学协助完成实验环境核验。"
        "学院提供了必要设备，本文谨向上述人员和机构表示感谢。",
    ),
)

EDITABLE_REWRITES = SOURCE_REWRITES[:2]


def zh(*codes: int) -> str:
    return "".join(chr(code) for code in codes)


def create_source_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(zh(0x6d4b, 0x8bd5, 0x8bba, 0x6587))
    document.add_paragraph(zh(0x76ee, 0x20, 0x20, 0x5f55))
    document.add_paragraph(zh(0x6458, 0x20, 0x20, 0x8981))
    document.add_paragraph(SOURCE_REWRITES[0][0])
    document.add_paragraph("1 " + zh(0x7eea, 0x8bba))
    document.add_paragraph(SOURCE_REWRITES[1][0])
    document.add_paragraph(zh(0x81f4, 0x20, 0x20, 0x8c22))
    document.add_paragraph(SOURCE_REWRITES[2][0])
    document.add_paragraph(zh(0x9644, 0x5f55, 0x20, 0x41))
    document.add_paragraph("appendix must stay unchanged")
    document.save(str(path))


def build_candidate_selection(
    input_text: str,
    output_text: str,
    chunk_id: str,
    *,
    global_style_profile: dict[str, object],
) -> dict[str, object]:
    neutral_dimension: dict[str, object] = {"id": "neutral", "primaryMetric": ""}
    app_service.validate_chunk_output(input_text, input_text, chunk_id)
    app_service.validate_chunk_output(input_text, output_text, chunk_id)
    baseline = _evaluate_rewrite_candidate(
        input_text=input_text,
        output_text=input_text,
        candidate_id="baseline",
        origin="baseline",
        attempt=0,
        hard_valid=True,
        round_dimension=neutral_dimension,
        global_style_profile=global_style_profile,
    )
    generated = _evaluate_rewrite_candidate(
        input_text=input_text,
        output_text=output_text,
        candidate_id="model-attempt-1",
        origin="model",
        attempt=1,
        hard_valid=True,
        round_dimension=neutral_dimension,
        global_style_profile=global_style_profile,
    )
    selected, reason_codes = _select_rewrite_candidate(
        [baseline, generated],
        round_dimension=neutral_dimension,
    )
    expected_published = output_text.strip() != input_text.strip()
    if expected_published and selected is not generated:
        raise AssertionError(f"fixture {chunk_id} rewrite did not win the production candidate selector")
    if not expected_published and selected is not baseline:
        raise AssertionError(f"fixture {chunk_id} identity baseline was not preserved")
    selection = _build_candidate_selection_event(
        chunk_id=chunk_id,
        round_number=1,
        candidates=[baseline, generated],
        selected=selected,
        reason_codes=reason_codes,
        conditional_retry_count=0,
    )
    if selection.get("publishedRewrite") is not expected_published:
        raise AssertionError(f"fixture {chunk_id} publication decision drifted")
    if selection.get("schemaVersion") != 2:
        raise AssertionError(f"fixture {chunk_id} did not use candidate-selection v2")
    result_delta = selection.get("resultSourceRelativeStyleDelta")
    if not isinstance(result_delta, dict) or result_delta.get("contextScope") != "document":
        raise AssertionError(f"fixture {chunk_id} did not bind document-relative style evidence")
    return selection


def expand_snapshot_scope(snapshot_path: Path) -> str:
    payload = json.loads(snapshot_path.read_text(encoding="utf-8"))
    units = payload.get("units")
    if not isinstance(units, list):
        raise AssertionError("snapshot units missing")
    selected_text = ""
    for unit in units:
        if isinstance(unit, dict) and not bool(unit.get("editable")) and str(unit.get("text", "")).strip():
            unit["editable"] = True
            unit["protect_reason"] = None
            selected_text = str(unit.get("text", ""))
            break
    if not selected_text:
        raise AssertionError("sample must contain at least one protected paragraph to simulate scope expansion")
    payload["editable_unit_count"] = sum(1 for unit in units if isinstance(unit, dict) and bool(unit.get("editable")))
    snapshot_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return selected_text


def main() -> int:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    EXPORT_PATH.unlink(missing_ok=True)
    create_source_docx(SOURCE_PATH)
    body_map = build_docx_body_map(SOURCE_PATH, prompt_profile="cn_custom", round_number=1)
    expected_sources = [source for source, _rewrite in EDITABLE_REWRITES]
    actual_sources = [unit.original_text for unit in body_map.units]
    if actual_sources != expected_sources:
        raise AssertionError(f"unexpected editable body-map scope: expected {len(expected_sources)}, got {len(actual_sources)}")
    if any(unit.original_text == SOURCE_REWRITES[2][0] for unit in body_map.units):
        raise AssertionError("acknowledgement body unexpectedly entered the editable body-map scope")
    # Keep one authoritative identity chunk so this two-paragraph legacy
    # fixture does not itself introduce a document-wide cluster of plain
    # sentence openings.  The changed first chunk still exercises a published
    # v2 model candidate; the second exercises a v2 preserved baseline.
    rewritten = [EDITABLE_REWRITES[0][1], EDITABLE_REWRITES[1][0]]
    global_style_profile = app_service.build_global_style_profile_from_texts(expected_sources)
    document_pattern_profile = global_style_profile.get("documentPatternBaseline")
    if not isinstance(document_pattern_profile, dict):
        raise AssertionError("legacy fixture did not produce a document pattern baseline")
    profile_sha256 = str(document_pattern_profile.get("profileSha256", "") or "")
    if not profile_sha256:
        raise AssertionError("legacy fixture document pattern baseline has no digest")
    body_map = update_docx_body_map_texts(body_map, rewritten, round_number=1)
    save_docx_body_map(body_map, BODY_MAP_PATH)
    OUTPUT_PATH.write_text("\n\n".join(rewritten), encoding="utf-8")
    COMPARE_PATH.write_text(
        json.dumps(
            {
                "chunkCount": len(rewritten),
                "paragraphCount": len(rewritten),
                "chunks": [
                    {
                        "chunkId": f"p{index}_c0",
                        "chunkIndex": 0,
                        "paragraphIndex": index,
                        "inputText": body_map.units[index].original_text,
                        "outputText": text,
                        "candidateBaselineText": body_map.units[index].original_text,
                        "candidateSelection": build_candidate_selection(
                            body_map.units[index].original_text,
                            text,
                            f"p{index}_c0",
                            global_style_profile=global_style_profile,
                        ),
                    }
                    for index, text in enumerate(rewritten)
                ],
                "sourcePatternProfiles": {
                    profile_sha256: document_pattern_profile,
                },
                "sourceRelativeDocumentDelta": app_service.assess_source_relative_document_delta(
                    expected_sources,
                    rewritten,
                ),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    snapshot_path = get_docx_snapshot_path(SOURCE_PATH)
    expanded_text = expand_snapshot_scope(snapshot_path)

    record_entry = {
        "origin_path": str(SOURCE_PATH.relative_to(ROOT_DIR)).replace("\\", "/"),
        "rounds": [
            {
                "round": 1,
                "prompt": "prompts/classical-rewrite.md",
                "prompt_profile": "cn_custom",
                "prompt_sequence": ["classical"],
                "input_path": str(SOURCE_PATH.relative_to(ROOT_DIR)).replace("\\", "/"),
                "output_path": str(OUTPUT_PATH.relative_to(ROOT_DIR)).replace("\\", "/"),
                "compare_path": str(COMPARE_PATH.relative_to(ROOT_DIR)).replace("\\", "/"),
                "body_map_path": str(BODY_MAP_PATH.relative_to(ROOT_DIR)).replace("\\", "/"),
                "input_segment_count": len(rewritten),
                "output_segment_count": len(rewritten),
            }
        ],
    }
    original_list_records = app_service.list_records
    app_service.list_records = lambda: {record_entry["origin_path"]: record_entry}
    export: dict[str, object] = {}
    export_failure: dict[str, object] = {}
    blocked = False
    try:
        try:
            export = app_service.export_round_output(
                str(OUTPUT_PATH.relative_to(ROOT_DIR)).replace("\\", "/"),
                str(EXPORT_PATH),
                "docx",
            )
        except app_service.ExportRoundError as exc:
            blocked = True
            export_failure = dict(exc.export_failure)
    finally:
        app_service.list_records = original_list_records

    failures: list[str] = []
    if not blocked:
        failures.append("tampered legacy snapshot unexpectedly exported instead of failing closed")
    if str(export_failure.get("stage", "")) != "guard":
        failures.append(f"unexpected hard-block stage: {export_failure.get('stage')}")
    report_path = Path(str(export_failure.get("reportPath", ""))) if export_failure.get("reportPath") else None
    if report_path is None or not report_path.exists():
        failures.append("content-contract block did not persist an audit report")
    else:
        contract = json.loads(report_path.read_text(encoding="utf-8"))
        issue_codes = {
            str(issue.get("code", ""))
            for issue in contract.get("blockingIssues", [])
            if isinstance(issue, dict)
        }
        if not {
            "body_map_snapshot_authority_mismatch",
            "body_map_snapshot_scope_signature_drift",
        }.intersection(issue_codes):
            failures.append(f"tampered legacy scope did not report the frozen scope drift: {sorted(issue_codes)}")
    if EXPORT_PATH.exists():
        failures.append("blocked tampered legacy export left a downloadable DOCX behind")

    report = {
        "ok": not failures,
        "failures": failures,
        "blocked": blocked,
        "export": export,
        "exportFailure": export_failure,
        "bodyMapUnitCount": len(body_map.units),
        "expandedScopeText": expanded_text,
    }
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if failures:
        print(json.dumps(report, ensure_ascii=False, indent=2), file=sys.stderr)
        return 1
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
