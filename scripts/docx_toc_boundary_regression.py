#!/usr/bin/env python3
"""Regression for generated-TOC markers that resemble real tail headings.

Chinese university templates commonly render entries such as ``致 谢 22``
and ``参考文献 23`` in the table of contents.  Those visible labels must never
close the editable scope before the actual thesis body begins.
"""

from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

from docx import Document

ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx_pipeline import (  # noqa: E402
    _looks_like_toc_entry,
    build_docx_scope_diagnostics,
    build_docx_snapshot,
)


WORK_DIR = ROOT_DIR / "finish" / "regression" / "docx_toc_boundary"
REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_toc_boundary_report.json"


def _build_probe(path: Path) -> None:
    document = Document()
    document.add_paragraph("机器人轨迹规划方法研究", style="Title")
    document.add_paragraph("摘 要")
    document.add_paragraph("本文围绕机器人轨迹规划开展研究，并验证所提出方法的可行性。")
    document.add_paragraph("目 录")
    document.add_paragraph("第1章 绪论 1")
    document.add_paragraph("致 谢 22")
    document.add_paragraph("参考文献 23")
    document.add_paragraph("第1章 绪论", style="Heading 1")
    document.add_paragraph("机器人轨迹规划需要同时满足运动连续性、约束一致性与执行稳定性。")
    document.add_paragraph("控制电路如图3.3所示。…………。")
    document.add_paragraph("实验结果显示，规划轨迹在关节空间内保持连续，且未超出速度约束。")
    document.add_paragraph("致 谢")
    document.add_paragraph("感谢指导教师在研究方案与实验设计方面给予帮助。")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 张三. 机器人轨迹规划方法研究[J]. 自动化学报, 2024.")
    document.save(path)


def main() -> int:
    shutil.rmtree(WORK_DIR, ignore_errors=True)
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    probe_path = WORK_DIR / "toc-tail-markers.docx"
    _build_probe(probe_path)

    snapshot = build_docx_snapshot(probe_path.resolve())
    diagnostics = build_docx_scope_diagnostics(snapshot)
    units = list(snapshot.units)
    by_text = {unit.text: unit for unit in units}
    failures: list[str] = []

    toc_ack = by_text.get("致 谢 22")
    toc_refs = by_text.get("参考文献 23")
    real_ack = by_text.get("致 谢")
    real_refs = by_text.get("参考文献")
    body_units = [
        by_text.get("机器人轨迹规划需要同时满足运动连续性、约束一致性与执行稳定性。"),
        by_text.get("控制电路如图3.3所示。…………。"),
        by_text.get("实验结果显示，规划轨迹在关节空间内保持连续，且未超出速度约束。"),
    ]

    if toc_ack is None or toc_refs is None or real_ack is None or real_refs is None:
        failures.append("probe markers were not extracted")
    else:
        scope = diagnostics.get("scope", {})
        if scope.get("acknowledgementIndex") != real_ack.unit_index:
            failures.append("TOC acknowledgement entry was mistaken for the real acknowledgement heading")
        if scope.get("postAcknowledgementBoundaryIndex") != real_refs.unit_index:
            failures.append("TOC references entry was mistaken for the real references boundary")
        if int(scope.get("endIndex", -1)) < real_ack.unit_index:
            failures.append("editable scope ended before the real acknowledgement section")
        if toc_ack.editable or toc_refs.editable:
            failures.append("TOC tail entries must remain protected")

    if any(unit is None or not unit.editable for unit in body_units):
        failures.append("thesis body paragraphs after the TOC must remain editable")
    if _looks_like_toc_entry("控制电路如图3.3所示。…………。"):
        failures.append("ordinary body ellipsis without a page token was misclassified as a TOC entry")
    if diagnostics.get("errorCount"):
        failures.append(f"scope diagnostics reported errors: {diagnostics.get('errorCount')}")

    report = {
        "ok": not failures,
        "failures": failures,
        "totalTextUnitCount": snapshot.total_text_unit_count,
        "editableUnitCount": snapshot.editable_unit_count,
        "scope": diagnostics.get("scope", {}),
    }
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
