#!/usr/bin/env python3
"""Production-chain regression for source-derived DOCX structural roles."""

from __future__ import annotations

from copy import deepcopy
from dataclasses import replace
import json
import sys
import tempfile
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore[import]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
from docx.oxml import OxmlElement  # type: ignore[import]
from docx.oxml.ns import qn  # type: ignore[import]
from docx.shared import Pt  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import round_helper  # noqa: E402
from document_edit_contract import (  # noqa: E402
    DOCUMENT_EDIT_CONTRACT_VERSION,
    build_document_edit_contract,
)
from docx_audit import (  # noqa: E402
    _paragraph_format_signature,
    audit_docx_export,
    audit_docx_format_lock,
    audit_docx_ooxml_integrity,
)
from docx_bodymap import (  # noqa: E402
    DOCX_BODY_MAP_SCOPE_SIGNATURE_VERSION,
    DOCX_BODY_MAP_VERSION,
    _build_scope_signature,
    build_docx_body_map,
    docx_body_map_from_payload,
    update_docx_body_map_texts,
    validate_docx_body_map,
)
from docx_pipeline import (  # noqa: E402
    DOCX_SCOPE_DIAGNOSTICS_VERSION,
    DOCX_SNAPSHOT_VERSION,
    DOCX_STRUCTURAL_INVENTORY_VERSION,
    DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
    _assign_unit_edit_decision,
    _docx_snapshot_derivation_digest,
    _load_docx_snapshot,
    build_docx_snapshot,
    build_docx_scope_diagnostics,
    ensure_docx_processing_assets,
    get_docx_unit_edit_eligibility_evidence_digest,
    rebuild_docx_from_body_map_units,
    rebuild_docx_from_snapshot,
)
from docx_protection_map import build_docx_protection_map  # noqa: E402


REPORT_PATH = ROOT_DIR / "finish" / "regression" / "docx_structural_role_contract_report.json"

ABSTRACT_BODY = "本文构建结构角色回归样本，并说明正文资格与格式锁之间的约束关系。"
CN_BODY = "实验采用统一采样条件，并记录多组评价结果，以验证结构角色判定的稳定性。"
EN_BODY = "The experiment uses a fixed sampling interval and records every evaluation result for comparison."
BODY_AFTER_TABLE = "表格之后的正文仍采用相同评价口径，并保留完整的实验条件说明。"
AUTO_NUMBERED_BODY = "控制器读取当前状态并核对采样周期，随后将计算结果写入实验记录。"
MANUAL_FORMAT_HEADING = "复杂动态环境下面向多源异构信息融合的鲁棒感知方法研究"
AMBIGUOUS_LONG_LABEL = "面向复杂动态环境的多源异构感知与跨尺度协同推断框架边界约束机制综合分析方法"
COLON_TITLE = "多源异构信息融合方法：复杂动态环境下的跨尺度协同感知与鲁棒推断框架研究"
ACK_BODY = "感谢指导教师在实验设计和资料校验方面提供帮助。"
TABLE_MARKERS = ("表内摘要标记", "表内参考文献标记", "表格实验数据", "嵌套表格保护数据")


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _set_body_format(paragraph: Any) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(10.5)


def _set_numbering(paragraph: Any) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.extend((ilvl, num_id))
    p_pr.append(num_pr)


def _create_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("结构角色资格回归论文", style="Title")
    document.add_paragraph("摘要")
    abstract_body = document.add_paragraph(ABSTRACT_BODY)
    _set_body_format(abstract_body)
    document.add_paragraph("1 引言", style="Heading 1")
    cn_body = document.add_paragraph(CN_BODY)
    _set_body_format(cn_body)

    manual_heading = document.add_paragraph()
    manual_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    manual_run = manual_heading.add_run(MANUAL_FORMAT_HEADING)
    manual_run.bold = True
    manual_run.font.size = Pt(16)

    ambiguous = document.add_paragraph(AMBIGUOUS_LONG_LABEL)
    _set_body_format(ambiguous)

    colon_title = document.add_paragraph(COLON_TITLE)
    _set_body_format(colon_title)

    english_body = document.add_paragraph(EN_BODY)
    _set_body_format(english_body)

    numbered = document.add_paragraph(AUTO_NUMBERED_BODY)
    _set_body_format(numbered)
    _set_numbering(numbered)

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = TABLE_MARKERS[0]
    table.cell(1, 0).text = TABLE_MARKERS[1]
    table.cell(1, 1).text = TABLE_MARKERS[2]
    nested_table = table.cell(1, 1).add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = TABLE_MARKERS[3]

    trailing_body = document.add_paragraph(BODY_AFTER_TABLE)
    _set_body_format(trailing_body)
    document.add_paragraph("致谢")
    document.add_paragraph(ACK_BODY)
    document.add_paragraph("参考文献", style="Heading 1")
    document.add_paragraph("[1] Structural role contract regression, 2026.")
    document.save(str(path))


def _paragraph_by_text(document: Any, text: str) -> Any:
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def _visible_table_texts(document: Any) -> list[str]:
    values: list[str] = []
    def collect_table(table: Any) -> None:
        seen_cells: set[Any] = set()
        for row in table.rows:
            for cell in row.cells:
                identity = cell._tc
                if identity in seen_cells:
                    continue
                seen_cells.add(identity)
                values.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
                for nested_table in cell.tables:
                    collect_table(nested_table)

    for table in document.tables:
        collect_table(table)
    return values


def main() -> int:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    checks: list[str] = []
    with tempfile.TemporaryDirectory(
        prefix="docx-structural-role-v6-",
        dir=REPORT_PATH.parent,
    ) as temp_dir:
        work_dir = Path(temp_dir)
        source_path = work_dir / "source.docx"
        extracted_path = work_dir / "extracted.txt"
        snapshot_path = work_dir / "snapshot.json"
        diagnostics_path = work_dir / "scope.json"
        export_path = work_dir / "export.docx"
        _create_fixture(source_path)

        _, _, snapshot = ensure_docx_processing_assets(
            source_path,
            extracted_path=extracted_path,
            snapshot_path=snapshot_path,
            scope_diagnostics_path=diagnostics_path,
        )
        first_fresh_snapshot = build_docx_snapshot(source_path)
        second_fresh_snapshot = build_docx_snapshot(source_path)
        _assert(
            _docx_snapshot_derivation_digest(first_fresh_snapshot)
            == _docx_snapshot_derivation_digest(second_fresh_snapshot),
            "consecutive fresh parses of the same DOCX produced different authority digests",
        )
        _assert(snapshot.version == DOCX_SNAPSHOT_VERSION == 22, "snapshot schema did not advance to v22")
        _assert(
            snapshot.structural_role_policy_version == DOCX_STRUCTURAL_ROLE_POLICY_VERSION == 6,
            "snapshot lost structural-role policy v6",
        )
        _assert(
            snapshot.structural_inventory_version == DOCX_STRUCTURAL_INVENTORY_VERSION == 3,
            "snapshot lost kind-aware recursive structural inventory v3",
        )
        diagnostics = build_docx_scope_diagnostics(snapshot, snapshot_path=snapshot_path)
        _assert(diagnostics.get("version") == DOCX_SCOPE_DIAGNOSTICS_VERSION == 5, "diagnostics schema did not advance to v5")

        by_text = {unit.text: unit for unit in snapshot.units}
        manual_unit = by_text[MANUAL_FORMAT_HEADING]
        _assert(manual_unit.structural_role == "heading" and not manual_unit.editable, "Normal direct-format heading entered model scope")
        manual_reasons = set(manual_unit.edit_eligibility_evidence.get("reasonCodes", []))
        _assert(
            {
                "presentation_structural_heading",
                "paragraph_centered_or_right_structural",
                "font_size_above_body_baseline",
                "all_visible_runs_bold_structural",
            }.issubset(manual_reasons),
            f"manual heading evidence is incomplete: {sorted(manual_reasons)}",
        )
        _assert(manual_unit.presentation_signals.get("alignment") == 1, "manual heading alignment evidence drifted")
        _assert(manual_unit.presentation_signals.get("maximumFontSizePt") == 16.0, "manual heading size evidence drifted")

        ambiguous_unit = by_text[AMBIGUOUS_LONG_LABEL]
        _assert(
            ambiguous_unit.structural_role == "ambiguous_non_prose" and not ambiguous_unit.editable,
            "long unformatted non-prose label did not fail closed",
        )
        colon_title_unit = by_text[COLON_TITLE]
        _assert(
            colon_title_unit.structural_role == "ambiguous_non_prose" and not colon_title_unit.editable,
            "single-colon Normal-style thesis heading was treated as affirmative prose",
        )
        for text, expected_role in (
            (ABSTRACT_BODY, "abstract_body"),
            (CN_BODY, "body_prose"),
            (EN_BODY, "body_prose"),
            (BODY_AFTER_TABLE, "body_prose"),
            (AUTO_NUMBERED_BODY, "body_list"),
        ):
            unit = by_text[text]
            _assert(unit.editable and unit.edit_eligibility == "eligible", f"affirmative prose was frozen: {text}")
            _assert(unit.structural_role == expected_role, f"wrong prose role for {text}: {unit.structural_role}")
            _assert(bool(unit.edit_eligibility_evidence.get("evidenceDigest")), f"prose role evidence missing: {text}")
        _assert(by_text[ACK_BODY].structural_role == "acknowledgement_body" and not by_text[ACK_BODY].editable, "acknowledgement prose remained editable")
        checks.append("presentation-only headings and unproven labels fail closed while Chinese, English, abstract and numbered prose remain eligible")
        checks.append("consecutive fresh parses produce one deterministic style/evidence authority digest")

        table_units = list(snapshot.protected_structural_units)
        _assert(len(table_units) == 4, f"merged/nested table evidence was not complete and deduplicated: {len(table_units)}")
        _assert(len({json.dumps(unit.target, sort_keys=True) for unit in table_units}) == 4, "table targets are not unique")
        _assert(
            all(
                unit.structural_role == "table_content"
                and unit.edit_eligibility == "protected"
                and not unit.editable
                and unit.target.get("body_block_index", -1) >= 0
                for unit in table_units
            ),
            "table structural evidence is incomplete or editable",
        )
        _assert({unit.text for unit in table_units} == set(TABLE_MARKERS), "table marker evidence drifted")
        nested_unit = next(unit for unit in table_units if unit.text == TABLE_MARKERS[3])
        _assert(nested_unit.target.get("table_depth") == 1, "nested table depth evidence was not recorded")
        _assert(
            "nested_table_content" in nested_unit.edit_eligibility_evidence.get("reasonCodes", []),
            "nested table protection reason was not recorded",
        )
        _assert(all(marker not in by_text for marker in TABLE_MARKERS), "table markers polluted top-level scope units")
        _assert(diagnostics.get("protectedTableParagraphCount") == 4, "scope diagnostics omitted protected table paragraphs")
        protection_map = build_docx_protection_map(source_path)
        _assert(protection_map["summary"]["tableUnits"] == 4, "protection map omitted explicit table units")
        _assert(protection_map["summary"]["structuralRolePolicyVersion"] == 6, "protection map omitted role policy")
        _assert(protection_map["summary"]["structuralInventoryVersion"] == 3, "protection map omitted kind-aware recursive inventory version")
        checks.append("merged cells and nested tables are recursively inventoried as protected evidence and cannot affect body scope")

        body_map = build_docx_body_map(source_path, snapshot_path=snapshot_path, round_number=1)
        _assert(body_map.version == DOCX_BODY_MAP_VERSION == 9, "body-map schema did not advance to v9")
        _assert(body_map.scope_signature.get("version") == DOCX_BODY_MAP_SCOPE_SIGNATURE_VERSION == 5, "scope signature did not advance to v5")
        _assert(
            all(
                unit.structural_role in {"abstract_body", "body_prose", "body_list"}
                and unit.edit_eligibility == "eligible"
                and bool(unit.edit_eligibility_evidence_digest)
                for unit in body_map.units
            ),
            "body-map lost role/evidence binding",
        )
        validation = validate_docx_body_map(body_map, source_path=source_path, snapshot_path=snapshot_path)
        _assert(validation.get("ok") is True, f"fresh role-bound body-map failed validation: {validation.get('blockingIssues')}")
        contract = build_document_edit_contract(
            source_path,
            snapshot_path=snapshot_path,
            extracted_text_path=extracted_path,
            body_map=body_map,
            candidate_texts=body_map.current_texts(),
            stage="role-regression",
        )
        _assert(contract.get("version") == DOCUMENT_EDIT_CONTRACT_VERSION == 3, "document contract did not advance to v3")
        _assert(contract.get("ready") is True, f"fresh structural-role contract is not ready: {contract.get('issues')}")
        _assert(contract.get("illegalEditableRoleCount") == 0, "contract exposed an illegal editable role")
        _assert(contract.get("ambiguousEditableUnitCount") == 0, "contract exposed an editable ambiguous unit")
        _assert(contract.get("protectedTableParagraphCount") == 4, "contract omitted protected table evidence")
        checks.append("body-map v9, scope signature v5 and document contract v3 bind source role/evidence")

        forged_heading_authority = deepcopy(manual_unit)
        _assign_unit_edit_decision(
            forged_heading_authority,
            structural_role="body_prose",
            edit_eligibility="eligible",
            reason_codes=("forged_positive_body_evidence",),
            protect_reason=None,
        )
        forged_first_unit = replace(
            body_map.units[0],
            unit_id=f"u{forged_heading_authority.unit_index}",
            unit_index=forged_heading_authority.unit_index,
            target=dict(forged_heading_authority.target),
            style_name=forged_heading_authority.style_name,
            original_text=forged_heading_authority.text,
            current_text="伪造的标题改写结果不得写回文档。",
            leading_whitespace=forged_heading_authority.leading_whitespace,
            trailing_whitespace=forged_heading_authority.trailing_whitespace,
            format_anchors=[dict(anchor) for anchor in forged_heading_authority.format_anchors],
            structural_role=forged_heading_authority.structural_role,
            edit_eligibility=forged_heading_authority.edit_eligibility,
            edit_eligibility_evidence_digest=get_docx_unit_edit_eligibility_evidence_digest(
                forged_heading_authority
            ),
        )
        forged_units = [forged_first_unit, *body_map.units[1:]]
        forged_body_map = replace(
            body_map,
            units=forged_units,
            editable_unit_count=len(forged_units),
            scope_signature=_build_scope_signature(forged_units),
        )
        forged_validation = validate_docx_body_map(
            forged_body_map,
            source_path=source_path,
            snapshot_path=snapshot_path,
        )
        forged_validation_codes = {
            str(issue.get("code", ""))
            for issue in forged_validation.get("blockingIssues", [])
            if isinstance(issue, dict)
        }
        _assert(forged_validation.get("ok") is False, "self-consistent forged body-map passed validation")
        _assert(
            "snapshot_scope_signature_drift" in forged_validation_codes,
            f"forged body-map did not expose source scope drift: {sorted(forged_validation_codes)}",
        )
        forged_contract = build_document_edit_contract(
            source_path,
            snapshot_path=snapshot_path,
            extracted_text_path=extracted_path,
            body_map=forged_body_map,
            candidate_texts=forged_body_map.current_texts(),
            stage="forged-body-map-role-regression",
        )
        _assert(forged_contract.get("ready") is False, "contract accepted a forged editable heading body-map")

        direct_forged_export_path = work_dir / "forged_body_map_export.docx"
        try:
            rebuild_docx_from_body_map_units(
                forged_body_map.units,
                source_path=source_path,
                export_path=direct_forged_export_path,
                preserve_format=True,
            )
        except ValueError as exc:
            _assert(
                "fresh source-derived prose authority" in str(exc),
                f"direct forged body-map export failed for the wrong reason: {exc}",
            )
        else:
            raise AssertionError("direct body-map rebuild accepted a forged editable heading")
        _assert(not direct_forged_export_path.exists(), "direct forged body-map rejection still wrote a DOCX")

        pre_run_calls = {"write": 0, "run": 0, "model": 0}
        original_prepare_body_map = round_helper._prepare_docx_body_map
        original_write_body_map_input = round_helper.write_docx_body_map_input
        original_run_round = round_helper.run_round

        def forged_prepare(_context: Any, *, prepared_parent: Any = None) -> Any:
            del prepared_parent
            return forged_body_map

        def counted_write(*args: Any, **kwargs: Any) -> Any:
            pre_run_calls["write"] += 1
            return original_write_body_map_input(*args, **kwargs)

        def counted_run(*args: Any, **kwargs: Any) -> Any:
            pre_run_calls["run"] += 1
            return original_run_round(*args, **kwargs)

        def blocked_model(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
            pre_run_calls["model"] += 1
            return chunk_text

        round_helper._prepare_docx_body_map = forged_prepare
        round_helper.write_docx_body_map_input = counted_write
        round_helper.run_round = counted_run
        try:
            try:
                round_helper.run_document_round(
                    source_path,
                    blocked_model,
                    round_number=1,
                    prompt_profile="cn",
                )
            except ValueError:
                pass
            else:
                raise AssertionError("production pre-run contract accepted a forged body-map")
        finally:
            round_helper._prepare_docx_body_map = original_prepare_body_map
            round_helper.write_docx_body_map_input = original_write_body_map_input
            round_helper.run_round = original_run_round
        _assert(
            pre_run_calls == {"write": 0, "run": 0, "model": 0},
            f"forged body-map crossed the pre-run side-effect boundary: {pre_run_calls}",
        )
        checks.append("forged body-map roles fail validation, contract, direct rebuild and production pre-run before any write/model call")

        model_inputs: list[str] = []

        def model_spy(chunk_text: str, _prompt_input: str, _round: int, _chunk_id: str) -> str:
            model_inputs.append(str(chunk_text))
            return str(chunk_text)

        original_update_round = round_helper.update_round
        round_helper.update_round = lambda **_kwargs: {"rounds": []}
        try:
            round_helper.run_document_round(
                source_path,
                model_spy,
                round_number=1,
                prompt_profile="cn",
            )
        finally:
            round_helper.update_round = original_update_round
        serialized_model_input = "\n\n".join(model_inputs)
        for protected_text in (MANUAL_FORMAT_HEADING, AMBIGUOUS_LONG_LABEL, COLON_TITLE, ACK_BODY, *TABLE_MARKERS):
            _assert(protected_text not in serialized_model_input, f"protected structural text reached model input: {protected_text}")
        for prose_text in (ABSTRACT_BODY, CN_BODY, EN_BODY, AUTO_NUMBERED_BODY, BODY_AFTER_TABLE):
            _assert(prose_text in serialized_model_input, f"eligible prose did not reach model input: {prose_text}")
        checks.append("production round model input contains only source-certified prose")

        rewritten = [
            text.replace("记录多组", "保存多组").replace("records every", "retains every")
            for text in body_map.current_texts()
        ]
        _assert(rewritten != body_map.current_texts(), "export fixture produced no body rewrite")
        updated_body_map = update_docx_body_map_texts(body_map, rewritten, round_number=1)
        rebuild_docx_from_body_map_units(
            updated_body_map.units,
            source_path=source_path,
            export_path=export_path,
            preserve_format=True,
        )
        source_document = Document(str(source_path))
        export_document = Document(str(export_path))
        for protected_text in (MANUAL_FORMAT_HEADING, AMBIGUOUS_LONG_LABEL, COLON_TITLE, ACK_BODY):
            source_paragraph = _paragraph_by_text(source_document, protected_text)
            export_paragraph = _paragraph_by_text(export_document, protected_text)
            _assert(source_paragraph.text == export_paragraph.text, f"protected paragraph text changed: {protected_text}")
            _assert(
                _paragraph_format_signature(source_paragraph._p) == _paragraph_format_signature(export_paragraph._p),
                f"protected paragraph format changed: {protected_text}",
            )
        _assert(_visible_table_texts(source_document) == _visible_table_texts(export_document), "table text changed during body-only export")
        _assert(any(text in "\n".join(paragraph.text for paragraph in export_document.paragraphs) for text in rewritten), "rewritten prose was not placed into the export")
        protection_audit = audit_docx_export(export_path, source_path=source_path, snapshot_path=snapshot_path)
        format_audit = audit_docx_format_lock(export_path, source_path=source_path, snapshot_path=snapshot_path)
        ooxml_audit = audit_docx_ooxml_integrity(export_path, source_path=source_path, snapshot_path=snapshot_path)
        _assert(protection_audit.get("ok") is True, f"protected-text audit failed: {protection_audit.get('issues')}")
        _assert(protection_audit.get("tableParagraphChecked") == 4, "protected-text audit omitted table evidence")
        _assert(format_audit.get("ok") is True, f"format-lock audit failed: {format_audit.get('issues')}")
        _assert(ooxml_audit.get("ok") is True, f"OOXML audit failed: {ooxml_audit.get('issues')}")
        checks.append("body-only export changes eligible prose while title, ambiguous label, acknowledgement, tables and all formatting remain locked")

        tampered_snapshot_path = work_dir / "tampered_snapshot.json"
        tampered_export_path = work_dir / "tampered_export.docx"
        tampered_snapshot = _load_docx_snapshot(snapshot_path)
        _assert(tampered_snapshot is not None, "fresh snapshot could not be loaded for tamper regression")
        tampered_heading = next(
            unit
            for unit in tampered_snapshot.units
            if unit.text == MANUAL_FORMAT_HEADING
        )
        _assign_unit_edit_decision(
            tampered_heading,
            structural_role="body_prose",
            edit_eligibility="eligible",
            reason_codes=("forged_positive_body_evidence",),
            protect_reason=None,
        )
        tampered_snapshot.editable_unit_count = len(tampered_snapshot.editable_units())
        tampered_snapshot_path.write_text(
            json.dumps(tampered_snapshot.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        try:
            rebuild_docx_from_snapshot(
                tampered_snapshot.editable_texts(),
                source_path=source_path,
                snapshot_path=tampered_snapshot_path,
                export_path=tampered_export_path,
                preserve_format=True,
            )
        except ValueError as exc:
            _assert(
                "fresh source-derived structural-role map" in str(exc),
                f"tampered snapshot failed for the wrong reason: {exc}",
            )
        else:
            raise AssertionError("self-consistent forged snapshot bypassed direct export authority")
        _assert(not tampered_export_path.exists(), "rejected forged snapshot still produced an export")
        checks.append("direct snapshot export re-derives source authority and rejects a self-consistent forged editable heading")

        stale_snapshot_payload = json.loads(snapshot_path.read_text(encoding="utf-8"))
        stale_snapshot_payload["version"] = 21
        stale_snapshot_payload.pop("structural_role_policy_version", None)
        stale_snapshot_payload.pop("structural_inventory_version", None)
        stale_snapshot_payload.pop("protected_structural_units", None)
        for raw_unit in stale_snapshot_payload.get("units", []):
            if isinstance(raw_unit, dict) and raw_unit.get("text") == MANUAL_FORMAT_HEADING:
                raw_unit["editable"] = True
                raw_unit["protect_reason"] = None
                raw_unit.pop("structural_role", None)
                raw_unit.pop("edit_eligibility", None)
                raw_unit.pop("edit_eligibility_evidence", None)
                raw_unit.pop("presentation_signals", None)
        snapshot_path.write_text(json.dumps(stale_snapshot_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        _, _, refreshed_snapshot = ensure_docx_processing_assets(
            source_path,
            extracted_path=extracted_path,
            snapshot_path=snapshot_path,
            scope_diagnostics_path=diagnostics_path,
        )
        _assert(refreshed_snapshot.version == 22, "v21 snapshot was reused instead of re-derived")
        refreshed_manual = next(unit for unit in refreshed_snapshot.units if unit.text == MANUAL_FORMAT_HEADING)
        _assert(refreshed_manual.structural_role == "heading" and not refreshed_manual.editable, "v21 editable heading survived v22 migration")
        _assert(len(refreshed_snapshot.protected_structural_units) == 4, "v21 migration did not restore recursive table evidence")

        stale_body_map_payload = body_map.to_dict()
        stale_body_map_payload["version"] = 8
        stale_body_map_payload["snapshot_version"] = 21
        stale_body_map_payload.pop("structural_role_policy_version", None)
        stale_body_map_payload.pop("structural_role_map_digest", None)
        for raw_unit in stale_body_map_payload.get("units", []):
            if isinstance(raw_unit, dict):
                raw_unit.pop("structural_role", None)
                raw_unit.pop("edit_eligibility", None)
                raw_unit.pop("edit_eligibility_evidence_digest", None)
        stale_body_map = docx_body_map_from_payload(stale_body_map_payload)
        _assert(stale_body_map is not None, "v8 body-map fixture did not parse for explicit rejection")
        stale_validation = validate_docx_body_map(
            stale_body_map,
            source_path=source_path,
            snapshot_path=snapshot_path,
        )
        stale_codes = {
            str(issue.get("code", ""))
            for issue in stale_validation.get("blockingIssues", [])
            if isinstance(issue, dict)
        }
        _assert(
            {
                "body_map_version_stale",
                "snapshot_version_mismatch",
                "structural_role_policy_version_mismatch",
                "structural_role_map_digest_mismatch",
            }.issubset(stale_codes),
            f"v8 body-map did not fail closed with role evidence: {sorted(stale_codes)}",
        )
        checks.append("v21 snapshot re-derives to v22 and v8 body-map fails closed under v9 role binding")

    report = {
        "ok": True,
        "snapshotVersion": DOCX_SNAPSHOT_VERSION,
        "structuralRolePolicyVersion": DOCX_STRUCTURAL_ROLE_POLICY_VERSION,
        "structuralInventoryVersion": DOCX_STRUCTURAL_INVENTORY_VERSION,
        "bodyMapVersion": DOCX_BODY_MAP_VERSION,
        "scopeSignatureVersion": DOCX_BODY_MAP_SCOPE_SIGNATURE_VERSION,
        "documentContractVersion": DOCUMENT_EDIT_CONTRACT_VERSION,
        "scopeDiagnosticsVersion": DOCX_SCOPE_DIAGNOSTICS_VERSION,
        "checks": checks,
    }
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
