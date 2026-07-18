from __future__ import annotations

import base64
from concurrent.futures import ThreadPoolExecutor
import json
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore[import]


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT_DIR / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import app_service  # noqa: E402
import round_helper  # noqa: E402
import web_app  # noqa: E402
from docx_pipeline import (  # noqa: E402
    INTERMEDIATE_DIR,
    build_docx_snapshot,
    ensure_docx_processing_assets,
    get_docx_extracted_text_path,
    get_docx_scope_diagnostics_path,
    get_docx_snapshot_path,
)
from fyadr_round_service import get_round_compare_path, relative_to_root  # noqa: E402


REPORT_PATH = ROOT_DIR / "finish" / "regression" / "document_artifact_isolation_regression_report.json"
WORK_DIR = ROOT_DIR / "finish" / "regression" / "document_artifact_isolation"
SOURCE_NAME = "p0-same-document-identity.docx"


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def _create_docx(path: Path, marker: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(f"{marker} 标题")
    document.add_paragraph("1 绪论")
    document.add_paragraph(f"{marker} 正文内容必须与另一份同名文档隔离。")
    document.add_paragraph("致谢")
    document.save(str(path))


def _remove_path(path: Path) -> None:
    try:
        path.unlink(missing_ok=True)
    except OSError:
        pass


def _cleanup_upload(path: Path) -> None:
    _remove_path(path)
    try:
        path.parent.rmdir()
    except OSError:
        pass


def _write_round_artifacts(output_path: Path, *, doc_id: str | None) -> dict[str, Any]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("legacy round output", encoding="utf-8")
    compare_path = get_round_compare_path(output_path)
    payload: dict[str, Any] = {
        "chunkCount": 1,
        "paragraphCount": 1,
        "chunks": [
            {
                "chunkId": "p0_c0",
                "chunkIndex": 0,
                "paragraphIndex": 0,
                "inputText": "input",
                "outputText": "legacy round output",
            }
        ],
    }
    if doc_id is not None:
        payload["docId"] = doc_id
    compare_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "round": 1,
        "prompt_profile": "cn_custom",
        "prompt_sequence": ["classical"],
        "output_path": relative_to_root(output_path),
        "compare_path": relative_to_root(compare_path),
        "input_segment_count": 1,
        "output_segment_count": 1,
    }


def _exercise_history_provenance(
    source_a: Path,
    source_b: Path,
    cleanup_paths: set[Path],
    checks: list[str],
) -> None:
    doc_a = relative_to_root(source_a)
    doc_b = relative_to_root(source_b)
    legacy_output = WORK_DIR / "legacy" / "p0-same-document-identity_round1.txt"
    cleanup_paths.update({legacy_output, get_round_compare_path(legacy_output)})
    unique_round = _write_round_artifacts(legacy_output, doc_id=None)

    original_load_records = round_helper.load_records_normalized
    original_list_records = app_service.list_records
    try:
        unique_records = {
            doc_a: {"origin_path": doc_a, "rounds": [dict(unique_round)]},
        }
        round_helper.load_records_normalized = lambda: unique_records
        restored = round_helper._previous_round_output_path(  # noqa: SLF001 - regression targets compatibility boundary
            doc_a,
            1,
            prompt_profile="cn_custom",
            prompt_sequence=["classical"],
        )
        _assert(restored.resolve() == legacy_output.resolve(), "a uniquely owned legacy artifact must remain readable")
        checks.append("uniquely owned pre-identity history artifacts remain readable")

        ambiguous_records = {
            doc_a: {"origin_path": doc_a, "rounds": [dict(unique_round)]},
            doc_b: {"origin_path": doc_b, "rounds": [dict(unique_round)]},
        }
        round_helper.load_records_normalized = lambda: ambiguous_records
        for doc_id in (doc_a, doc_b):
            try:
                round_helper._previous_round_output_path(  # noqa: SLF001
                    doc_id,
                    1,
                    prompt_profile="cn_custom",
                    prompt_sequence=["classical"],
                )
            except ValueError:
                pass
            else:
                raise AssertionError("an untagged legacy artifact shared by two documents must be rejected")
        app_service.list_records = lambda: ambiguous_records
        _assert(
            app_service._find_record_context_for_output(legacy_output) is None,  # noqa: SLF001
            "ambiguous legacy output lookup must not select the first document",
        )
        checks.append("ambiguous legacy history paths are rejected instead of first-match restored")

        tagged_round = _write_round_artifacts(legacy_output, doc_id=doc_b)
        tagged_records = {
            doc_a: {"origin_path": doc_a, "rounds": [dict(tagged_round)]},
            doc_b: {"origin_path": doc_b, "rounds": [dict(tagged_round)]},
        }
        round_helper.load_records_normalized = lambda: tagged_records
        try:
            round_helper._previous_round_output_path(  # noqa: SLF001
                doc_a,
                1,
                prompt_profile="cn_custom",
                prompt_sequence=["classical"],
            )
        except ValueError:
            pass
        else:
            raise AssertionError("compare docId mismatch must block another document from restoring the artifact")
        restored_b = round_helper._previous_round_output_path(  # noqa: SLF001
            doc_b,
            1,
            prompt_profile="cn_custom",
            prompt_sequence=["classical"],
        )
        _assert(restored_b.resolve() == legacy_output.resolve(), "matching compare docId must select the correct document")
        app_service.list_records = lambda: tagged_records
        context = app_service._find_record_context_for_output(legacy_output)  # noqa: SLF001
        _assert(context is not None and context[0].get("origin_path") == doc_b, "output lookup must honor compare docId")
        checks.append("compare docId proves legacy artifact ownership and blocks cross-document reuse")
    finally:
        round_helper.load_records_normalized = original_load_records
        app_service.list_records = original_list_records


def _exercise_upload_isolation(cleanup_uploads: set[Path], checks: list[str]) -> None:
    text_name = "p0-upload-same-name.txt"
    first_text = "first upload must remain immutable"
    second_text = "second upload has different content"
    with ThreadPoolExecutor(max_workers=2) as executor:
        first_future = executor.submit(web_app.write_uploaded_file, text_name, first_text)
        second_future = executor.submit(web_app.write_uploaded_file, text_name, second_text)
        first_path = first_future.result()
        second_path = second_future.result()
    cleanup_uploads.update({first_path, second_path})
    _assert(first_path != second_path, "same-name uploads with different content must have different source paths")
    _assert(first_path.name == second_path.name == text_name, "content directories must preserve the display filename")
    _assert(first_path.read_text(encoding="utf-8") == first_text, "the first text upload was overwritten")
    _assert(second_path.read_text(encoding="utf-8") == second_text, "the second text upload was not stored intact")
    _assert(web_app.write_uploaded_file(text_name, first_text) == first_path, "identical text upload must reuse its stable path")
    checks.append("same-name text uploads are content-addressed, immutable, and deduplicated")

    binary_name = "p0-upload-same-name.docx"
    first_binary = b"PK\x03\x04first-binary-document"
    second_binary = b"PK\x03\x04second-binary-document"
    with ThreadPoolExecutor(max_workers=2) as executor:
        first_future = executor.submit(
            web_app.write_uploaded_binary_file,
            binary_name,
            base64.b64encode(first_binary).decode("ascii"),
        )
        second_future = executor.submit(
            web_app.write_uploaded_binary_file,
            binary_name,
            base64.b64encode(second_binary).decode("ascii"),
        )
        first_binary_path = first_future.result()
        second_binary_path = second_future.result()
    cleanup_uploads.update({first_binary_path, second_binary_path})
    _assert(first_binary_path != second_binary_path, "same-name binary uploads must not collide")
    _assert(first_binary_path.read_bytes() == first_binary, "the first binary upload was overwritten")
    _assert(second_binary_path.read_bytes() == second_binary, "the second binary upload was not stored intact")
    _assert(
        web_app.write_uploaded_binary_file(binary_name, base64.b64encode(first_binary).decode("ascii")) == first_binary_path,
        "identical binary upload must reuse its stable path",
    )
    checks.append("same-name binary uploads are atomically isolated without changing display names")


def main() -> int:
    shutil.rmtree(WORK_DIR, ignore_errors=True)
    source_a = WORK_DIR / "directory-a" / SOURCE_NAME
    source_b = WORK_DIR / "directory-b" / SOURCE_NAME
    _create_docx(source_a, "ALPHA")
    _create_docx(source_b, "BRAVO")

    cleanup_paths: set[Path] = set()
    cleanup_uploads: set[Path] = set()
    checks: list[str] = []
    failures: list[str] = []
    details: dict[str, Any] = {}

    snapshot_a = get_docx_snapshot_path(source_a)
    snapshot_b = get_docx_snapshot_path(source_b)
    extracted_a = get_docx_extracted_text_path(source_a)
    extracted_b = get_docx_extracted_text_path(source_b)
    scope_a = get_docx_scope_diagnostics_path(source_a)
    scope_b = get_docx_scope_diagnostics_path(source_b)
    cleanup_paths.update({snapshot_a, snapshot_b, extracted_a, extracted_b, scope_a, scope_b})

    legacy_snapshot_path = INTERMEDIATE_DIR / f"{source_a.stem}_docx_snapshot.json"
    legacy_extracted_path = INTERMEDIATE_DIR / f"{source_a.stem}_extracted.txt"
    legacy_scope_path = INTERMEDIATE_DIR / f"{source_a.stem}_scope_diagnostics.json"
    cleanup_paths.update({legacy_snapshot_path, legacy_extracted_path, legacy_scope_path})

    try:
        _assert(snapshot_a != snapshot_b, "same-named DOCX snapshots must use different paths")
        _assert(extracted_a != extracted_b, "same-named DOCX extracted text must use different paths")
        _assert(scope_a != scope_b, "same-named DOCX diagnostics must use different paths")
        _assert(get_docx_snapshot_path(source_a) == snapshot_a, "the same source path must deterministically hit one snapshot")
        _assert("directory-a" not in snapshot_a.name and "directory-b" not in snapshot_b.name, "artifact names leaked source directories")
        checks.append("DOCX processing paths use stable non-leaking document identities")

        for path in cleanup_paths:
            _remove_path(path)
        legacy_payload = build_docx_snapshot(source_a).to_dict()
        units = legacy_payload.get("units")
        if isinstance(units, list) and units and isinstance(units[0], dict):
            units[0]["style_name"] = "LEGACY_MIGRATION_SENTINEL"
        legacy_snapshot_path.parent.mkdir(parents=True, exist_ok=True)
        legacy_snapshot_path.write_text(json.dumps(legacy_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        legacy_extracted_path.write_text("UNTRUSTED LEGACY EXTRACTED SENTINEL", encoding="utf-8")
        legacy_scope_path.write_text('{"sentinel":"legacy scope must remain read-only"}', encoding="utf-8")
        legacy_bytes_before = {
            legacy_snapshot_path: legacy_snapshot_path.read_bytes(),
            legacy_extracted_path: legacy_extracted_path.read_bytes(),
            legacy_scope_path: legacy_scope_path.read_bytes(),
        }

        with ThreadPoolExecutor(max_workers=2) as executor:
            result_a_future = executor.submit(ensure_docx_processing_assets, source_a)
            result_b_future = executor.submit(ensure_docx_processing_assets, source_b)
            result_a = result_a_future.result()
            result_b = result_b_future.result()
        _assert(result_a[0] == extracted_a and result_a[1] == snapshot_a, "source A resolved unexpected DOCX assets")
        _assert(result_b[0] == extracted_b and result_b[1] == snapshot_b, "source B resolved unexpected DOCX assets")
        payload_a = json.loads(snapshot_a.read_text(encoding="utf-8"))
        payload_b = json.loads(snapshot_b.read_text(encoding="utf-8"))
        _assert(Path(payload_a["source_path"]).resolve() == source_a.resolve(), "source A snapshot provenance drifted")
        _assert(Path(payload_b["source_path"]).resolve() == source_b.resolve(), "source B reused source A's legacy snapshot")
        _assert("ALPHA" in extracted_a.read_text(encoding="utf-8"), "source A extracted text was crossed")
        _assert("BRAVO" in extracted_b.read_text(encoding="utf-8"), "source B extracted text was crossed")
        _assert("UNTRUSTED LEGACY" not in extracted_a.read_text(encoding="utf-8"), "legacy extracted text was copied without provenance")
        _assert("UNTRUSTED LEGACY" not in extracted_b.read_text(encoding="utf-8"), "foreign legacy extracted text was reused")
        migrated_units = payload_a.get("units")
        _assert(
            isinstance(migrated_units, list)
            and migrated_units
            and isinstance(migrated_units[0], dict)
            and migrated_units[0].get("style_name") == "LEGACY_MIGRATION_SENTINEL",
            "a current, provenance-matching legacy snapshot was not migrated",
        )
        for path, before in legacy_bytes_before.items():
            _assert(path.read_bytes() == before, f"legacy compatibility path was modified: {path.name}")
        checks.append("legacy snapshot migration requires exact provenance and leaves old paths read-only")
        checks.append("same-name DOCX assets remain isolated under concurrent interleaving")

        context_a = round_helper.build_round_context(
            source_a,
            round_number=1,
            prompt_profile="cn_custom",
            prompt_sequence=["classical"],
        )
        context_b = round_helper.build_round_context(
            source_b,
            round_number=1,
            prompt_profile="cn_custom",
            prompt_sequence=["classical"],
        )
        for field in ("input_text_path", "output_text_path", "manifest_path", "body_map_path", "validation_path"):
            _assert(getattr(context_a, field) != getattr(context_b, field), f"round artifact collision in {field}")
        _assert(
            round_helper.build_round_artifact_stem(context_a.doc_id, "cn_custom", ["classical"])
            == round_helper.build_round_artifact_stem(context_a.doc_id, "cn_custom", ["classical"]),
            "round artifact identity is not stable",
        )
        checks.append("round input/output/manifest/body-map/validation stems isolate same-named documents")

        context_a.output_text_path.parent.mkdir(parents=True, exist_ok=True)
        context_a.output_text_path.write_text("A reset target", encoding="utf-8")
        context_b.output_text_path.write_text("B must survive A reset", encoding="utf-8")
        cleanup_paths.update({context_a.output_text_path, context_b.output_text_path})
        original_delete_rounds = app_service.delete_rounds
        app_service.delete_rounds = lambda *_args, **_kwargs: {
            "deletedRounds": [],
            "deletedFiles": [],
            "remainingRounds": [],
        }
        try:
            app_service.reset_round_progress(
                str(source_a),
                "cn_custom",
                1,
                prompt_sequence=["classical"],
            )
        finally:
            app_service.delete_rounds = original_delete_rounds
        _assert(not context_a.output_text_path.exists(), "reset did not delete the selected document artifact")
        _assert(context_b.output_text_path.read_text(encoding="utf-8") == "B must survive A reset", "reset cross-deleted another document")
        checks.append("round reset deletes only the selected document's identity stem")

        _exercise_history_provenance(source_a, source_b, cleanup_paths, checks)
        _exercise_upload_isolation(cleanup_uploads, checks)

        details = {
            "sourceA": relative_to_root(source_a),
            "sourceB": relative_to_root(source_b),
            "snapshotA": relative_to_root(snapshot_a),
            "snapshotB": relative_to_root(snapshot_b),
            "roundStemA": round_helper.build_round_artifact_stem(context_a.doc_id, "cn_custom", ["classical"]),
            "roundStemB": round_helper.build_round_artifact_stem(context_b.doc_id, "cn_custom", ["classical"]),
        }
    except Exception as exc:
        failures.append(f"{type(exc).__name__}: {exc}")
    finally:
        for path in cleanup_paths:
            _remove_path(path)
        for path in cleanup_uploads:
            _cleanup_upload(path)

    report = {
        "ok": not failures,
        "failures": failures,
        "checks": checks,
        "details": details,
    }
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    stream = sys.stdout if not failures else sys.stderr
    print(json.dumps(report, ensure_ascii=False, indent=2), file=stream)
    return 0 if not failures else 1


if __name__ == "__main__":
    raise SystemExit(main())
